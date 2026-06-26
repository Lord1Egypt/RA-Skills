"""
文档可视化主入口

导出函数:
    run(input_path: str, theme: str = "THEME_BLUE") -> dict
        完整流程：解析 → 分析 → 生成HTML → 导出PDF/长图
        返回: {"html": str, "pdf": str, "png": str, "output_dir": str}
"""

import sys
import os
import json
import re
import argparse
from datetime import datetime

# 添加skill路径
SKILL_DIR = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, SKILL_DIR)

def parse_args():
    """解析命令行参数"""
    if len(sys.argv) < 2:
        print("用法: python doc_visualizer.py <文件路径或URL或飞书文档链接>")
        sys.exit(1)
    return sys.argv[1]

def detect_input_type(input_str: str) -> dict:
    """检测输入类型"""
    input_str = input_str.strip()
    
    if input_str.startswith("https://") or input_str.startswith("http://"):
        if "feishu.cn" in input_str or "larksuite" in input_str:
            return {"type": "feishu_doc", "value": input_str}
        elif input_str.endswith((".pdf", ".docx", ".xlsx", ".txt", ".csv")):
            return {"type": "url_file", "value": input_str}
        else:
            return {"type": "url", "value": input_str}
    elif os.path.isfile(input_str):
        ext = os.path.splitext(input_str)[1].lower()
        if ext == ".pdf":
            return {"type": "pdf", "value": input_str}
        elif ext == ".docx":
            return {"type": "docx", "value": input_str}
        elif ext in [".xlsx", ".xls"]:
            return {"type": "excel", "value": input_str}
        elif ext in [".txt", ".csv"]:
            return {"type": "text", "value": input_str}
        elif ext == ".html":
            return {"type": "html", "value": input_str}
        else:
            return {"type": "unknown", "value": input_str}
    else:
        return {"type": "text_raw", "value": input_str}

def parse_freshdesk_doc(url: str) -> dict:
    """解析飞书文档"""
    from feishu_fetch_doc import fetch_doc
    # 实际上用subprocess调用更稳定，这里简化
    doc_id = extract_doc_id(url)
    return fetch_doc(doc_id)

def extract_doc_id(url: str) -> str:
    """从URL提取飞书文档ID"""
    # https://my.feishu.cn/docx/RaAUdNosQoUmB7xiPcwcPaTXnxh
    match = re.search(r'/docx/([a-zA-Z0-9]+)', url)
    if match:
        return match.group(1)
    return url

def parse_pdf(file_path: str) -> dict:
    """解析PDF文件"""
    try:
        import pdfplumber
    except ImportError:
        print("需要安装: pip install pdfplumber")
        sys.exit(1)
    
    tables = []
    text_blocks = []
    
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                text_blocks.append(text)
            page_tables = page.extract_tables()
            tables.extend(page_tables)
    
    return {"text_blocks": text_blocks, "tables": tables, "metadata": {"source": "pdf"}}

def parse_docx(file_path: str) -> dict:
    """解析Word文件"""
    try:
        from docx import Document
    except ImportError:
        print("需要安装: pip install python-docx")
        sys.exit(1)
    
    doc = Document(file_path)
    text_blocks = [p.text for p in doc.paragraphs if p.text.strip()]
    tables = []
    for table in doc.tables:
        tbl = []
        for row in table.rows:
            tbl.append([cell.text.strip() for cell in row.cells])
        tables.append(tbl)
    
    return {"text_blocks": text_blocks, "tables": tables, "metadata": {"source": "docx"}}

def parse_excel(file_path: str) -> dict:
    """解析Excel文件"""
    try:
        import openpyxl
    except ImportError:
        print("需要安装: pip install openpyxl")
        sys.exit(1)
    
    wb = openpyxl.load_workbook(file_path)
    tables = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        tbl = []
        for row in ws.iter_rows(values_only=True):
            tbl.append([str(c) if c is not None else "" for c in row])
        if tbl:
            tables.append({"sheet": sheet_name, "data": tbl})
    
    return {"tables": tables, "text_blocks": [], "metadata": {"source": "excel", "sheets": wb.sheetnames}}

def parse_text(file_path: str) -> dict:
    """解析文本/CSV文件"""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        content = f.read()
    
    is_csv = os.path.splitext(file_path)[1].lower() == ".csv"
    if is_csv:
        lines = content.split("\n")
        tables = []
        for line in lines:
            if line.strip():
                tables.append([c.strip() for c in line.split(",")])
        return {"tables": [tables], "text_blocks": [], "metadata": {"source": "csv"}}
    else:
        return {"text_blocks": content.split("\n"), "tables": [], "metadata": {"source": "text"}}

def analyze_data(parsed: dict) -> dict:
    """
    分析解析后的数据，自动识别类型并生成可视化配置
    这是核心智能分析逻辑
    """
    text = " ".join(parsed.get("text_blocks", []))
    tables = parsed.get("tables", [])
    meta = parsed.get("metadata", {})
    source = meta.get("source", "unknown")
    
    # 通用数据提取
    title = extract_title(text, tables, source)
    sections = []
    
    # 策略1: 检测是否为竞品/商业分析文档
    if contains_keywords(text, ["竞品", "竞争对手", "市场", "营收", "财务", "SWOT", "战略"]):
        sections = build_business_analysis(text, tables)
    # 策略2: 检测是否为时间序列/事件流
    elif contains_keywords(text, ["时间", "事件", "动态", "进展", "2024", "2025", "2026"]):
        sections = build_timeline_section(text, tables)
    # 策略3: 检测是否为对比数据
    elif len(tables) > 0 and has_comparable_data(tables):
        sections = build_compare_section(tables)
    # 策略4: 通用表格数据
    elif len(tables) > 0:
        sections = build_table_section(tables)
    # 策略5: 纯文本摘要
    else:
        sections = build_text_summary(text)
    
    # 始终尝试提取财务卡片（如果有任何数字数据）
    fin_cards = extract_financial_cards(text, tables)
    if fin_cards:
        sections.insert(0, {"type": "fin_grid", "data": {"cards": fin_cards}})
    
    return {
        "title": title,
        "subtitle": f"自动生成 | {source} | {datetime.now().strftime('%Y-%m-%d')}",
        "meta": {
            "数据来源": source.upper(),
            "生成时间": datetime.now().strftime("%Y-%m-%d %H:%M"),
        },
        "sections": sections
    }

def contains_keywords(text: str, keywords: list) -> bool:
    text_lower = text.lower()
    return any(kw in text_lower for kw in keywords)

def has_comparable_data(tables: list) -> bool:
    # tables格式: list of {"sheet": str, "data": [[]]} or just [[]]
    if not tables:
        return False
    for tbl in tables:
        if isinstance(tbl, dict):
            data = tbl.get("data", [])
        else:
            data = tbl
        if data and len(data) > 1 and len(data[0]) >= 2:
            return True
    return False

def extract_title(text: str, tables: list, source: str) -> str:
    # 尝试从文本提取标题
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    if lines:
        # 取第一行作为标题
        title = lines[0]
        if len(title) > 5 and len(title) < 50:
            return title
    # 回退
    return f"{source.upper()} 数据可视化看板"

def extract_financial_cards(text: str, tables: list) -> list:
    """从文本和表格中提取财务指标"""
    cards = []
    patterns = [
        (r'营收[^\d]*([\d.]+)[亿万元]', '营收', 'val'),
        (r'收入[^\d]*([\d.]+)[亿万元]', '收入', 'val'),
        (r'利润[^\d]*([\d.]+)[亿万元]', '净利润', 'val'),
        (r'同比[+]?([\d.]+)%', '同比', 'sub'),
        (r'市占率[^\d]*([\d.]+)%', '市占率', 'sub'),
    ]
    
    for pat, label, val_type in patterns:
        m = re.search(pat, text)
        if m:
            val = m.group(1)
            if val_type == 'val':
                cards.append({"label": label, "val": f"¥{val}亿", "sub": "数据来源提取", "sub_color": "green"})
            else:
                cards.append({"label": label, "val": f"+{val}%", "sub": "同比增长", "sub_color": "green"})
    
    return cards[:8]  # 最多8个卡片

def build_business_analysis(text: str, tables: list) -> list:
    """构建商业分析可视化"""
    sections = []
    
    # 时间轴
    timeline_events = extract_timeline_events(text)
    if timeline_events:
        sections.append({"type": "timeline", "data": {"events": timeline_events}})
    
    # SWOT
    swot_data = extract_swot(text)
    if any(swot_data.values()):
        sections.append({"type": "swot", "data": swot_data})
    
    # 四维分析
    four_dim = extract_four_dim(text)
    if any(four_dim.values()):
        sections.append({"type": "four_dim", "data": four_dim})
    
    # 表格
    if tables:
        sections.extend(build_table_section(tables))
    
    return sections

def extract_timeline_events(text: str) -> list:
    """从文本提取时间轴事件"""
    events = []
    # 匹配日期+事件
    date_pattern = r'(202[0-9][年/-]?\d{1,2}[月/-]?\d{0,2})[^。\n]{2,30}'
    matches = re.findall(date_pattern, text)
    for i, m in enumerate(matches[:10]):
        tag_colors = []
        tags = []
        if any(kw in m for kw in ["重大", "崩", "重整"]):
            tags.append("重大风险"); tag_colors.append("red")
        elif any(kw in m for kw in ["新品", "发布", "产品"]):
            tags.append("产品动向"); tag_colors.append("blue")
        elif any(kw in m for kw in ["战略", "合作"]):
            tags.append("战略合作"); tag_colors.append("green")
        events.append({
            "time": m[0][:10] if m[0] else f"2026-{i+1}",
            "tags": tags or ["事件"],
            "tag_colors": tag_colors or ["blue"],
            "event": m[0][10:] if len(m[0]) > 10 else m[0]
        })
    return events

def extract_swot(text: str) -> dict:
    """提取SWOT分析"""
    return {
        "S": extract_bullet_list(text, ["优势", "Strength", "S "]),
        "W": extract_bullet_list(text, ["劣势", "Weakness", "W "]),
        "O": extract_bullet_list(text, ["机会", "Opportunity", "O "]),
        "T": extract_bullet_list(text, ["威胁", "Threat", "T "]),
    }

def extract_four_dim(text: str) -> dict:
    """提取四维战略分析"""
    return {
        "attack": extract_bullet_list(text, ["进攻", "扩张", "增长"]),
        "defend": extract_bullet_list(text, ["防守", "风险", "压力"]),
        "opportunity": extract_bullet_list(text, ["机会", "机遇"]),
        "threat": extract_bullet_list(text, ["威胁", "风险"]),
    }

def extract_bullet_list(text: str, keywords: list) -> list:
    """提取列表项"""
    lines = text.split("\n")
    result = []
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if any(kw in line for kw in keywords):
            # 清理
            cleaned = re.sub(r'^[-*•🔵🔴🟢🟡🔶✅⚠️🌟🚨]\s*', '', line)
            if cleaned and len(cleaned) > 3:
                result.append(cleaned[:100])
    return result[:5]

def build_timeline_section(text: str, tables: list) -> list:
    events = extract_timeline_events(text)
    result = [
        {"type": "timeline", "data": {"events": events}},
    ]
    if tables:
        result.extend(build_table_section(tables))
    else:
        result.append({"type": "section_header", "data": {"text": "📋 内容摘要", "color": ""}})
    return result

def build_compare_section(tables: list) -> list:
    if not tables:
        return []
    tbl = tables[0] if isinstance(tables[0], list) else tables[0].get("data", [])
    if not tbl:
        return []
    headers = tbl[0]
    headers = ["维度"] + headers if headers else ["项目", "数据"]
    rows = []
    for row in tbl[1:]:
        if row and any(c.strip() for c in row):
            label = row[0] if row else ""
            vals = row[1:] if len(row) > 1 else [""]
            rows.append({"label": label, "values": vals, "class": "val-neutral"})
    return [{"type": "compare", "data": {"headers": headers, "rows": rows}}]

def build_table_section(tables: list) -> list:
    sections = []
    for tbl in tables:
        if isinstance(tbl, dict):
            sheet_name = tbl.get("sheet", "")
            data = tbl.get("data", [])
        else:
            sheet_name = ""
            data = tbl
        if not data:
            continue
        headers = data[0] if data else []
        rows = data[1:] if len(data) > 1 else []
        sec = {"type": "table", "data": {"headers": headers, "rows": rows}}
        if sheet_name:
            sec["data"]["sheet"] = sheet_name
        sections.append(sec)
    return sections

def build_text_summary(text: str) -> list:
    lines = [l.strip() for l in text.split("\n") if l.strip() and len(l.strip()) > 10]
    summary_lines = lines[:20]
    content = "\n".join(f"<p>{l}</p>" for l in summary_lines)
    return [{"type": "section_header", "data": {"text": "📋 内容摘要", "color": ""}}]

def interactive_theme_select(available_themes: list) -> str:
    """交互式主题选择"""
    print("\n" + "="*50)
    print("🎨 请选择配色主题:")
    print("="*50)
    for i, t in enumerate(available_themes, 1):
        name, desc, color = t
        print(f"  {i}. [{color}] {name} - {desc}")
    print("="*50)
    
    while True:
        try:
            choice = input("请输入数字 (1-5): ").strip()
            idx = int(choice) - 1
            if 0 <= idx < len(available_themes):
                return available_themes[idx][0]
        except ValueError:
            pass
        print("无效输入，请重新输入")

AVAILABLE_THEMES = [
    ("THEME_BLUE", "科技蓝（默认）", "🔵"),
    ("THEME_RED", "商务红", "🔴"),
    ("THEME_GREEN", "自然绿", "🟢"),
    ("THEME_PURPLE", "皇家紫", "🟣"),
    ("THEME_ORANGE", "活力橙", "🟠"),
]

def run(input_path: str, theme: str = "THEME_BLUE"):
    """主运行函数"""
    print(f"\n📂 输入: {input_path}\n")
    
    # Step 1: 检测输入类型
    info = detect_input_type(input_path)
    input_type = info["type"]
    value = info["value"]
    print(f"📌 检测类型: {input_type}")
    
    # Step 2: 解析文档
    print("⏳ 解析文档...")
    if input_type == "feishu_doc":
        # 通过飞书API获取
        try:
            from feishu_fetch_doc import fetch_doc
            doc_id = extract_doc_id(value)
            result = fetch_doc(doc_id)
            parsed = {"text_blocks": [result.get("markdown", "")], "tables": [], "metadata": {"source": "feishu"}}
        except Exception as e:
            print(f"飞书文档解析失败: {e}")
            parsed = {"text_blocks": [value], "tables": [], "metadata": {"source": "feishu"}}
    elif input_type == "pdf":
        parsed = parse_pdf(value)
    elif input_type == "docx":
        parsed = parse_docx(value)
    elif input_type == "excel":
        parsed = parse_excel(value)
    elif input_type == "text" or input_type == "csv":
        parsed = parse_text(value)
    elif input_type == "html":
        with open(value, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        parsed = {"text_blocks": [content], "tables": [], "metadata": {"source": "html"}}
    else:
        parsed = {"text_blocks": [value], "tables": [], "metadata": {"source": "text"}}
    
    print(f"  ✅ 解析完成: {len(parsed.get('tables',[]))}个表格, {len(parsed.get('text_blocks',[]))}个文本块")
    
    # Step 3: 数据分析
    print("⏳ 分析数据结构...")
    viz_config = analyze_data(parsed)
    print(f"  ✅ 生成 {len(viz_config['sections'])} 个可视化组件")
    
    # Step 4: 主题选择
    print("\n" + "="*50)
    print("🎨 可视化主题选择")
    print("="*50)
    for i, (name, desc, emoji) in enumerate(AVAILABLE_THEMES, 1):
        print(f"  {i}. {emoji} {name} - {desc}")
    print("="*50)
    
    # theme由调用方传入，此处不再重复解析
    pass
    
    print(f"  ✅ 选择主题: {theme}")
    
    # Step 5: 生成HTML
    print("⏳ 生成HTML...")
    sys.path.insert(0, os.path.join(SKILL_DIR, "generator"))
    from html_generator import generate_html
    
    html_content = generate_html(viz_config, theme=theme, title=viz_config.get("title", os.path.basename(input_path)))
    
    # 保存HTML
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_name = re.sub(r'[^\w\u4e00-\u9fff]', '_', os.path.basename(input_path).rsplit(".", 1)[0] or "output")[:40]
    output_dir = os.path.join(SKILL_DIR, "..", "..", "visual_exports", timestamp)
    os.makedirs(output_dir, exist_ok=True)
    
    html_path = os.path.join(output_dir, f"{safe_name}.html")
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html_content)
    print(f"  ✅ HTML已保存: {html_path}")
    
    # Step 6: 导出PDF和长图
    print("⏳ 导出PDF和长图...")
    sys.path.insert(0, os.path.join(SKILL_DIR, "exporter"))
    from exporter import export_pdf_and_screenshot
    
    try:
        exports = export_pdf_and_screenshot(html_path, output_dir)
        print(f"  ✅ PDF已保存: {exports['pdf']}")
        print(f"  ✅ 长图已保存: {exports['png']}")
    except Exception as e:
        print(f"  ⚠️ 导出失败: {e}")
        exports = {"pdf": None, "png": None}
    
    # Step 7: 输出结果
    print("\n" + "="*50)
    print("📦 生成结果")
    print("="*50)
    print(f"  📄 HTML: {html_path}")
    if exports.get("pdf"):
        print(f"  📕 PDF: {exports['pdf']}")
    if exports.get("png"):
        print(f"  🖼️ 长图: {exports['png']}")
    print("="*50)
    
    return {
        "html": html_path,
        "pdf": exports.get("pdf"),
        "png": exports.get("png"),
        "output_dir": output_dir
    }

if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("input_path", nargs="?", default=None)
    parser.add_argument("--theme", "-t", default="THEME_BLUE")
    args = parser.parse_args()
    if not args.input_path:
        print("Usage: python doc_visualizer.py <file> [--theme THEME_XXX]")
        print("Themes: THEME_BLUE / THEME_RED / THEME_GREEN / THEME_PURPLE / THEME_ORANGE")
        sys.exit(1)
    tm = {"1":"THEME_BLUE","2":"THEME_RED","3":"THEME_GREEN","4":"THEME_PURPLE","5":"THEME_ORANGE"}
    resolved = tm.get(args.theme, args.theme if args.theme in [t[0] for t in AVAILABLE_THEMES] else "THEME_BLUE")
    result = run(args.input_path, theme=resolved)
    print(f"Done! Output: {result['output_dir']}")
