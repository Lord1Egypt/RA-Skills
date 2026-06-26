#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
CMMI V3.0 文档生成器 - 配置驱动版本
支持全部31个PA（实践域）的文档自动生成
使用方法: python generate_cmmi_docx.py --pa <PA代码> --output <输出文件>
"""

import argparse
import json
import os
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ========== 工具函数 ==========
def set_font(run, size=10, bold=False, color=None, name='宋体'):
    """设置字体样式"""
    run.font.size = Pt(size)
    run.font.name = name
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run

def add_heading(doc, text, level=1, indent=False):
    """添加标题"""
    if level == 1:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(text), size=18, bold=True, color=(26, 26, 110))
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        p = doc.add_paragraph()
        if indent:
            p.paragraph_format.left_indent = Inches(0.5)
        set_font(p.add_run(text), size=14, bold=True, color=(26, 26, 110))
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
    return p

def add_text(doc, text, indent=False):
    """添加正文"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.5)
    set_font(p.add_run(text), size=10)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_bullets(doc, items, indent=False):
    """添加项目符号列表"""
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        if indent:
            p.paragraph_format.left_indent = Inches(0.5)
        set_font(p.add_run(item), size=10)
        p.paragraph_format.space_after = Pt(2)

def add_table(doc, headers, rows=None):
    """添加表格"""
    if not rows:
        rows = [['' for _ in headers]]
    
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_font(run, size=9, bold=True, color=(255, 255, 255))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 数据行
    for r_idx, row_data in enumerate(rows, 1):
        for c_idx, cell_data in enumerate(row_data):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = str(cell_data)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_font(run, size=9)
    
    doc.add_paragraph()
    return table

def add_toc_placeholder(doc):
    """添加目录占位符"""
    p = doc.add_paragraph()
    p.add_run('目录').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('（自动生成目录）')
    doc.add_page_break()

# ========== 通用文档生成函数 ==========
def generate_pa_document(pa_code, output_file, config_file='pa_config.json'):
    """
    根据PA配置生成CMMI V3.0合规文档
    
    Args:
        pa_code: PA代码（如 'RDM', 'PLAN' 等）
        output_file: 输出文件路径
        config_file: 配置文件路径
    """
    # 读取配置
    try:
        with open(config_file, 'r', encoding='utf-8') as f:
            config = json.load(f)
    except Exception as e:
        print(f"❌ 无法读取配置文件: {e}")
        return False
    
    if pa_code not in config:
        print(f"❌ 未找到PA代码: {pa_code}")
        print(f"   支持的PA: {', '.join(sorted(config.keys()))}")
        return False
    
    pa_config = config[pa_code]
    pa_name = pa_config['name']
    doc_title = pa_config['doc_title']
    sections = pa_config['sections']
    
    print(f"📄 开始生成文档: {pa_code} - {pa_name}")
    print(f"   输出文件: {output_file}")
    print(f"   章节数量: {len(sections)}")
    
    # 创建文档
    doc = Document()
    
    # 封面
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run('CMMI V3.0 合规文档'), size=24, bold=True, color=(26, 26, 110))
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(doc_title), size=20, bold=True)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(f'实践域: {pa_code} - {pa_name}'), size=14)
    
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run('文档版本: V1.0'), size=12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run('生成日期: ________年__月__日'), size=12)
    
    doc.add_page_break()
    
    # 目录（占位符）
    add_toc_placeholder(doc)
    
    # 正文章节
    for section in sections:
        title = section['title']
        section_type = section['type']
        
        # 章节标题
        add_heading(doc, title, level=2)
        
        # 根据类型添加内容
        if section_type == 'text':
            add_text(doc, '【待补充】根据CMMI V3.0要求填写此部分内容。')
        
        elif section_type == 'bullets':
            add_bullets(doc, [
                '【待补充项目1】',
                '【待补充项目2】',
                '【待补充项目3】'
            ])
        
        elif section_type == 'table':
            headers = section.get('headers', ['字段1', '字段2', '字段3'])
            add_table(doc, headers)
        
        else:
            add_text(doc, f'【未知类型: {section_type}】')
    
    # 保存文档
    try:
        doc.save(output_file)
        print(f"✅ 文档生成成功: {output_file}")
        return True
    except Exception as e:
        print(f"❌ 保存文档失败: {e}")
        return False

# ========== 主函数 ==========
def main():
    parser = argparse.ArgumentParser(description='CMMI V3.0 文档生成器（配置驱动版本）')
    parser.add_argument('--pa', required=True, help='PA代码（如: RDM, PLAN, CM 等）')
    parser.add_argument('--output', required=True, help='输出文件路径（.docx）')
    parser.add_argument('--config', default='pa_config.json', help='配置文件路径（默认: pa_config.json）')
    
    args = parser.parse_args()
    
    # 检查配置文件
    if not os.path.exists(args.config):
        print(f"❌ 配置文件不存在: {args.config}")
        sys.exit(1)
    
    # 生成文档
    success = generate_pa_document(args.pa.upper(), args.output, args.config)
    
    if success:
        print(f"\n🎉 完成！文档已生成: {args.output}")
        sys.exit(0)
    else:
        print(f"\n❌ 文档生成失败")
        sys.exit(1)

if __name__ == '__main__':
    main()
