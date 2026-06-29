#!/usr/bin/env python3
"""
test_full_workflow.py - Step 12 全链路集成测试（修订 v2）

端到端验证 Step 1-11 全部组件协同工作。
"""

import sys
import os
import json

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from orchestrator_v2 import (
    orchestrate_phase1_1,
    confirm_phase1,
    orchestrate_phase1_3,
    update_node_content_hint,
    confirm_phase1_3,
    synthesize_chapter_summary,
    load_orchestrate_state,
    _get_orchestrate_state_path,
)
from context_builder import build_prompt_package
from state_manager_v2 import (
    outline_save,
    outline_load,
    outline_update_status,
    outline_get_context,
    _get_state_path,
)


TEST_PAPER = "test_full_workflow_paper"


def cleanup():
    for p in [_get_state_path(TEST_PAPER), _get_orchestrate_state_path(TEST_PAPER)]:
        if os.path.exists(p):
            try:
                os.remove(p)
            except FileNotFoundError:
                pass


def create_proposal():
    """创建真实开题报告 docx"""
    try:
        import docx
    except ImportError:
        return None

    doc = docx.Document()
    doc.add_heading("论文大纲", level=1)
    chapters = [
        ("第1章 绪论", [("1.1 研究背景", "AI 时代背景"),
                     ("1.2 研究内容", "本文聚焦差异化战略")]),
        ("第2章 理论基础", [("2.1 竞争战略理论", "Porter 三种基本战略"),
                       ("2.2 文献综述", "近年 AI 时代战略文献")]),
        ("第3章 外部环境", [("3.1 宏观环境", "PEST 模型分析"),
                       ("3.2 行业结构", "波特五力模型")]),
        ("第4章 内部环境", [("4.1 资源与能力", "VRIN 资源分析")]),
        ("第5章 战略选择", [("5.1 QSPM 矩阵", "QSPM 量化评估")]),
        ("第6章 战略实施", [("6.1 实施路径", "三阶段实施")]),
        ("第7章 结论", [("7.1 研究结论", "总结研究结论")]),
    ]
    for ch_title, sections in chapters:
        doc.add_heading(ch_title, level=2)
        for sec_title, sec_content in sections:
            doc.add_heading(sec_title, level=3)
            doc.add_paragraph(sec_content)
    doc.add_heading("参考文献", level=1)
    path = f"/tmp/{TEST_PAPER}_proposal.docx"
    doc.save(path)
    return path


def chapter_aware_summary_llm(prompt: str) -> str:
    """章节感知 mock：根据 prompt 中的章节标题返回不同摘要"""
    if "绪论" in prompt:
        return "本章系统提出差异化战略研究问题，奠定全文框架。"
    if "理论基础" in prompt:
        return "本章系统梳理竞争战略理论，为后续分析提供方法论。"
    if "外部环境" in prompt:
        return "本章从宏观与行业结构两个维度分析外部环境。"
    return "本章进行相关分析。"


# ============================================================
# 测试 A：完整 docx 流程（happy path）
# ============================================================

def test_full_workflow_happy_path():
    """测试 A：完整 docx 流程"""
    print("\n=== 测试 A：完整 docx 流程（happy path） ===")
    cleanup()

    fake_path = create_proposal()
    if not fake_path:
        print("   ⚠️ 跳过（python-docx 未安装）")
        return

    # === Phase 1.1: docx 解析 ===
    print("   [1] Phase 1.1 docx 解析...")
    r1 = orchestrate_phase1_1(TEST_PAPER, "docx", fake_path)
    assert r1["ok"] is True
    outline_state = outline_load(TEST_PAPER)
    virtual_count = sum(1 for n in outline_state["outline"]["outline_tree"]["nodes"]
                        if n.get("is_virtual"))
    assert virtual_count == 7
    print(f"      ✅ 解析成功（{virtual_count} 个虚拟摘要节点）")

    # === Phase 1.2: 确认 ===
    print("   [2] Phase 1.2 目录确认...")
    r2 = confirm_phase1(TEST_PAPER)
    assert r2["phase"] == "phase1"

    # === Phase 1.3: 归因 ===
    print("   [3] Phase 1.3 开题报告归因...")
    r3 = orchestrate_phase1_3(TEST_PAPER)
    if r3["ok"]:
        print(f"      ✅ 归因成功")

    # === 用户调整 hint ===
    print("   [4] 用户调整 content_hint...")
    r4 = update_node_content_hint(TEST_PAPER, "1.1", "用户对研究背景的深度自定义提示")
    assert r4["ok"]
    node_11 = next((n for n in outline_load(TEST_PAPER)["outline"]["outline_tree"]["nodes"]
                    if n["id"] == "1.1"), None)
    assert node_11["content_hint"] == "用户对研究背景的深度自定义提示"
    print(f"      ✅ 用户调整生效")

    # === Phase 1.3 确认 → Phase 2 ===
    print("   [5] Phase 1.3 确认 → Phase 2...")
    r5 = confirm_phase1_3(TEST_PAPER)
    assert r5["phase"] == "phase2"

    # === Phase 2 写作（增强项1 触发章节摘要） ===
    print("   [6] Phase 2 写作（第一章所有 L2）...")
    # 模拟 1.1 + 1.2 完成 → 触发 ch1 摘要
    outline_update_status(TEST_PAPER, "1.1", "completed", key_conclusion="AI 时代背景")
    outline_update_status(TEST_PAPER, "1.2", "completed", key_conclusion="研究内容")

    r_ch1_summary = synthesize_chapter_summary(
        TEST_PAPER, "ch1", llm_func=chapter_aware_summary_llm
    )
    assert r_ch1_summary["ok"], f"ch1 摘要失败: {r_ch1_summary.get('error', '')}"
    print(f"      ✅ ch1 摘要合成: {r_ch1_summary['summary'][:40]}...")

    # === 验证 ch1 摘要已写入 ===
    outline_state = outline_load(TEST_PAPER)
    ch1_summary_node = next((n for n in outline_state["outline"]["outline_tree"]["nodes"]
                              if n["id"] == "__ch1_summary__"), None)
    assert ch1_summary_node["writing_status"] == "completed"
    print(f"      ✅ __ch1_summary__ 节点已 completed")

    print("   ✅ 全链路集成通过")


# ============================================================
# 测试 B：失败回退流程
# ============================================================

def test_full_workflow_failure_recovery():
    """测试 B：解析失败 → 重新输入 → 成功"""
    print("\n=== 测试 B：解析失败 → 重新输入 → 成功 ===")
    cleanup()

    print("   [1] 第一次输入：完全无效 text...")
    r1 = orchestrate_phase1_1(TEST_PAPER, "text", "完全是垃圾内容")
    assert r1["ok"] is False
    assert r1["action"] == "input_required"

    print("   [2] 第二次输入：还是无效...")
    r2 = orchestrate_phase1_1(TEST_PAPER, "text", "还是没有章节结构")
    assert r2["ok"] is False

    print("   [3] 第三次：切换到 docx...")
    fake_path = create_proposal()
    if not fake_path:
        print("      ⚠️ 跳过")
        return

    r3 = orchestrate_phase1_1(TEST_PAPER, "docx", fake_path)
    assert r3["ok"] is True

    confirm_phase1(TEST_PAPER)
    print(f"      ✅ 失败回退 → 成功")


# ============================================================
# 测试 C：章节摘要 + bridge 跨章节（增强项1 + Step 11 协同）
# ============================================================

def test_chapter_summary_bridge_integration():
    """测试 C：章节摘要合成 + bridge 跨章节引用"""
    print("\n=== 测试 C：章节摘要 + bridge 跨章节 ===")
    cleanup()

    fake_path = create_proposal()
    if not fake_path:
        print("   ⚠️ 跳过")
        return

    orchestrate_phase1_1(TEST_PAPER, "docx", fake_path)
    confirm_phase1(TEST_PAPER)

    # 模拟 ch1 + ch2 全部完成
    outline_update_status(TEST_PAPER, "1.1", "completed", key_conclusion="AI 时代背景")
    outline_update_status(TEST_PAPER, "1.2", "completed", key_conclusion="研究内容")
    r1 = synthesize_chapter_summary(TEST_PAPER, "ch1", chapter_aware_summary_llm)
    assert r1["ok"], f"ch1 summary fail: {r1.get('error', '')}"

    outline_update_status(TEST_PAPER, "2.1", "completed", key_conclusion="战略理论")
    outline_update_status(TEST_PAPER, "2.2", "completed", key_conclusion="文献综述")
    r2 = synthesize_chapter_summary(TEST_PAPER, "ch2", chapter_aware_summary_llm)
    assert r2["ok"], f"ch2 summary fail: {r2.get('error', '')}"

    # === 验证 ch2 摘要 ===
    outline_state = outline_load(TEST_PAPER)
    ch2_summary_node = next((n for n in outline_state["outline"]["outline_tree"]["nodes"]
                              if n["id"] == "__ch2_summary__"), None)
    assert ch2_summary_node["key_conclusion"] is not None
    print(f"   [1] __ch2_summary__: {ch2_summary_node['key_conclusion'][:50]}")

    # === 验证增强项1 P3 fallback：3.1 应引用 ch2 摘要 ===
    print("   [2] 验证 3.1 (ch3 首节点) prev_chapter_summary...")
    ctx = outline_get_context(TEST_PAPER, "3.1")
    prev_summary = ctx.get("prev_chapter_summary")
    assert prev_summary is not None, f"3.1 应有 prev_chapter_summary"
    assert prev_summary["chapter_id"] == "ch2"
    assert prev_summary["chapter_title"] == "理论基础"
    # 关键是内容应该是 ch2 摘要，不是 ch1
    assert "理论" in prev_summary["key_conclusion"] or "方法论" in prev_summary["key_conclusion"]
    print(f"      ✅ prev_chapter_summary = {prev_summary['chapter_title']}")

    # === 验证 bridge P3 fallback：3.1 prompt 应包含 ch2 摘要 ===
    print("   [3] 验证 3.1 prompt bridge...")
    pkg = build_prompt_package(TEST_PAPER, "3.1")
    bridge = pkg.get("bridge_paragraph")
    assert bridge is not None, "3.1 应有 bridge"
    assert "理论基础" in bridge, f"bridge 应包含 ch2 章节标题"
    print(f"      ✅ bridge: {bridge[:80]}...")

    # === 验证 content_hint 增强项4 + Step 11 协同 ===
    print("   [4] 验证 content_hint 集成...")
    content_hint = pkg.get("content_hint")
    print(f"      ✅ content_hint: {content_hint[:60] if content_hint else '<empty>'}...")

    print("   ✅ 增强项1 + 增强项4 + Step 11 协同工作")


def main():
    """主测试入口"""
    print("=" * 60)
    print("Step 12 全链路集成测试（修订 v2）")
    print("=" * 60)

    tests = [
        test_full_workflow_happy_path,
        test_full_workflow_failure_recovery,
        test_chapter_summary_bridge_integration,
    ]

    passed = 0
    failed = 0
    for t in tests:
        try:
            t()
            passed += 1
        except AssertionError as e:
            print(f"\n❌ 测试失败: {e}")
            failed += 1
        except Exception as e:
            print(f"\n❌ 异常: {e}")
            import traceback
            traceback.print_exc()
            failed += 1
        finally:
            cleanup()

    print("\n" + "=" * 60)
    print(f"结果：{passed} passed, {failed} failed")
    print("=" * 60)
    return 0 if failed == 0 else 1


if __name__ == "__main__":
    sys.exit(main())
