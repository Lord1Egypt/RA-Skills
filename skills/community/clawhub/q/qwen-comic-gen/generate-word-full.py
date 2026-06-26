# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 创建文档
doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(10)

# 标题
heading = doc.add_heading('感知与行动中心 - 专家解决方案 v2.1', level=0)
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
heading_run = heading.runs[0]
heading_run.font.name = 'Microsoft YaHei'
heading_run.font.size = Pt(16)
heading_run.font.bold = True

# 文档信息
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.LEFT
info.add_run('文档维护：阿福（AI 助理） | 更新频率：每周（周五） | 下次更新：2026-03-13\n')
info.add_run('版本号：v2.1（2026-03-06 修订）\n')
info.add_run('本文档由阿福自动生成，基于项目知识库和行业最佳实践')

doc.add_paragraph('_' * 80)

# 项目基本信息
doc.add_heading('📋 项目基本信息摘要', level=1)

doc.add_heading('项目概述', level=2)
doc.add_paragraph('分期实施：')
doc.add_paragraph('• 一期（2026-03 至 2026-06）：感知侧 - 数据采集、监控看板、数据治理', style='List Bullet')
doc.add_paragraph('• 二期（2026-07 至 2026-09）：行动侧 - 风险预警、工单闭环、组织保障', style='List Bullet')

doc.add_paragraph('一期目标：')
doc.add_paragraph('• 115 家供应商产能数据采集完成', style='List Bullet')
doc.add_paragraph('• 6 大核心指标监控看板上线', style='List Bullet')
doc.add_paragraph('• 数据字段标准化（23 个字段）', style='List Bullet')
doc.add_paragraph('• 风险识别规则配置', style='List Bullet')

doc.add_paragraph('关键点：')
doc.add_paragraph('• 数据治理先行（字段标准化）', style='List Bullet')
doc.add_paragraph('• 风险预警机制（红/黄/绿三级）', style='List Bullet')
doc.add_paragraph('• 工单闭环流程（商务/技术/资源/质量）', style='List Bullet')

# 第一部分
doc.add_heading('第一部分 业务背景和需求', level=1)

doc.add_heading('1.1 业务背景', level=2)
doc.add_paragraph('业务现状：')
doc.add_paragraph('小米汽车产能管理涉及 115 家供应商，涵盖电池、电机、电控等核心零部件。当前采用人工数据采集方式，通过飞书文档/表格记录供应商产能信息。')

doc.add_paragraph('面临的变化：')
doc.add_paragraph('• 产能爬坡期，供应商产能需快速提升', style='List Bullet')
doc.add_paragraph('• 多车型并行，产能复杂度增加', style='List Bullet')
doc.add_paragraph('• 供应链风险频发（疫情、地缘政治、原材料涨价）', style='List Bullet')

doc.add_heading('1.2 当前现状和痛点', level=2)
doc.add_paragraph('业务目标：')
doc.add_paragraph('• 产能达标率≥85%', style='List Bullet')
doc.add_paragraph('• SDC 签约及时率≥90%（v2.1 更新：SBC→SDC）', style='List Bullet')
doc.add_paragraph('• GP9 完成率≥90%', style='List Bullet')
doc.add_paragraph('• 风险预警准确率≥95%', style='List Bullet')

doc.add_paragraph('业务现状：')
doc.add_paragraph('• 产能达标率 68%（距离 85% 目标差距 17%）', style='List Bullet')
doc.add_paragraph('• SDC 签约及时率 82%（距离 90% 目标差距 8%）', style='List Bullet')
doc.add_paragraph('• GP9 完成率 71.3%（距离 90% 目标差距 18.7%）', style='List Bullet')
doc.add_paragraph('• 28 家供应商未签 SDC（包含弗迪等头部供应商）', style='List Bullet')

doc.add_paragraph('业务痛点：')
doc.add_paragraph('• 数据孤岛 - 数据散落在飞书文档、Excel、邮件中', style='List Bullet')
doc.add_paragraph('• 风险滞后 - 无法实时识别产能风险', style='List Bullet')
doc.add_paragraph('• 闭环缺失 - 问题发现后无有效跟踪机制', style='List Bullet')
doc.add_paragraph('• 效率低下 - 大量人工操作，自动化程度低', style='List Bullet')

doc.add_heading('1.5 业务价值', level=2)
doc.add_paragraph('价值 1：提升产能达标率')
doc.add_paragraph('论点：风险预警提前发现问题，工单闭环确保问题解决')
doc.add_paragraph('论据：试点供应商产能达标率从 68% 提升至 85%+')

doc.add_paragraph('价值 2：降低供应链风险')
doc.add_paragraph('论点：实时预警机制提前识别风险')
doc.add_paragraph('论据：风险识别时间从平均 7 天缩短至 1 天')

# 第二部分
doc.add_heading('第二部分 数据侧/解决方案', level=1)

doc.add_heading('2.1 方案目标', level=2)
doc.add_paragraph('目标 1：数据采集自动化 - 飞书文档自动读取，减少人工操作 90%')
doc.add_paragraph('目标 2：风险预警实时化 - 风险识别时间<1 小时，预警准确率≥95%')
doc.add_paragraph('目标 3：工单闭环系统化 - 工单自动派发，SLA 管理，闭环率 100%')
doc.add_paragraph('目标 4：监控看板可视化 - 6 大 KPI 实时展示，支持钻取分析')

# 第三部分
doc.add_heading('第三部分 项目规划', level=1)

doc.add_heading('3.1 组织架构', level=2)
doc.add_paragraph('项目指导委员会：')
doc.add_paragraph('主任：供应链管理部负责人')
doc.add_paragraph('职责：重大决策、资源协调、升级事项处理')

doc.add_heading('3.2 实施计划', level=2)
doc.add_paragraph('时间安排：')
doc.add_paragraph('开始时间：2026-03-06')
doc.add_paragraph('结束时间：2026-12-31')
doc.add_paragraph('总周期：10 个月')

# v2.1 更新内容
doc.add_heading('📋 v2.1 修订记录', level=1)
doc.add_paragraph('根据第 1 轮批注修订：')
doc.add_paragraph('• SBC→SDC（全局替换）', style='List Bullet')
doc.add_paragraph('• 补充工单流程细节', style='List Bullet')
doc.add_paragraph('• 补充预警规则阈值', style='List Bullet')
doc.add_paragraph('• 补充供应商分级标准', style='List Bullet')

doc.add_heading('工单闭环流程（v2.1 新增）', level=2)
doc.add_paragraph('5 阶段流转：')
doc.add_paragraph('1. 创建 - 系统自动创建（风险预警触发）', style='List Number')
doc.add_paragraph('2. 派发 - 项目经理派发（商务/技术/资源/质量）', style='List Number')
doc.add_paragraph('3. 执行 - 责任人执行解决（每日更新进展）', style='List Number')
doc.add_paragraph('4. 验收 - 管理员审核验收', style='List Number')
doc.add_paragraph('5. 关闭 - 系统自动关闭', style='List Number')

doc.add_heading('风险预警规则（v2.1 新增）', level=2)
doc.add_paragraph('红/黄/绿三级阈值：')
doc.add_paragraph('🟢 绿色：产能达标率≥80% → 正常监控', style='List Bullet')
doc.add_paragraph('🟡 黄色：产能达标率 60%-80% → 48 小时响应', style='List Bullet')
doc.add_paragraph('🔴 红色：产能达标率<60% → 24 小时响应', style='List Bullet')

doc.add_heading('供应商分级标准（v2.1 新增）', level=2)
doc.add_paragraph('A 类（战略供应商）：产能占比≥10% 或独家供应，约 15 家', style='List Bullet')
doc.add_paragraph('B 类（核心供应商）：产能占比 5%-10%，约 30 家', style='List Bullet')
doc.add_paragraph('C 类（一般供应商）：产能占比<5%，约 70 家', style='List Bullet')

# 附录
doc.add_heading('📚 附录', level=1)
doc.add_paragraph('术语表（v2.1 更新）：')
doc.add_paragraph('SDC（Supplier Daily Capacity）供应商日产能（v2.1 更新：原 SBC）', style='List Bullet')
doc.add_paragraph('GP9（产能实测验证）', style='List Bullet')
doc.add_paragraph('LCR（场地要求确认）', style='List Bullet')
doc.add_paragraph('SLA（服务等级协议）', style='List Bullet')

doc.add_paragraph('参考文档：')
doc.add_paragraph('元数据表：https://feishu.cn/docx/NOhWdjxFAoNpXqxK3Woc68MxnRd', style='List Bullet')

# 保存文档
save_path = r'C:\Users\Xiabi\.openclaw\workspace\感知与行动中心 - 专家解决方案 v2.1.docx'
doc.save(save_path)
print(f'Word document saved to: {save_path}')
print(f'File size: {len(open(save_path, "rb").read())} bytes')
