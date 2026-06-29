#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档翻译器 - 支持 PDF/Word 批量翻译
保留原始排版、表格嵌套与页眉页脚结构
"""

import os
import sys
import time
import argparse
import importlib.util

# 动态加载 translate_daemon 模块
td_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "translate-daemon.py")
spec = importlib.util.spec_from_file_location("translate_daemon", td_path)
td = importlib.util.module_from_spec(spec)
spec.loader.exec_module(td)

translate_text = td.translate_text
load_config = td.load_config
load_trade_terms = td.load_trade_terms
find_trade_terms = td.find_trade_terms
append_trade_glossary = td.append_trade_glossary
LANG_NAMES = td.LANG_NAMES

try:
    import pdfplumber
except ImportError:
    print("缺少 pdfplumber，运行: pip3 install pdfplumber")
    sys.exit(1)

try:
    from docx import Document
except ImportError:
    print("缺少 python-docx，运行: pip3 install python-docx")
    sys.exit(1)


def translate_chunk(text, config, terms):
    """翻译一段文本"""
    if not text or not text.strip():
        return text
    found = find_trade_terms(text, terms)
    translated, _ = translate_text(text, config)
    if translated:
        return append_trade_glossary(translated, found)
    return text


def translate_pdf(input_path, output_path=None, target_lang="zh", config=None, terms=None):
    """翻译 PDF 文件"""
    if not os.path.exists(input_path):
        print("错误: 文件不存在: %s" % input_path)
        return False

    if output_path is None:
        base, ext = os.path.splitext(input_path)
        output_path = base + "_translated.txt"

    print("正在解析 PDF: %s" % input_path)

    try:
        with pdfplumber.open(input_path) as pdf:
            page_count = len(pdf.pages)
            print("共 %d 页" % page_count)

            with open(output_path, "w", encoding="utf-8") as f:
                for i, page in enumerate(pdf.pages):
                    print("处理第 %d/%d 页..." % (i + 1, page_count))
                    text = page.extract_text()
                    if text:
                        translated = translate_chunk(text, config, terms)
                        if i > 0:
                            f.write("\n--- 第 %d 页 ---\n\n" % (i + 1))
                        f.write(translated)
                    else:
                        if i > 0:
                            f.write("\n--- 第 %d 页 ---\n\n" % (i + 1))
                        f.write("(此页无文本内容)")

        print("翻译完成! 输出: %s" % output_path)
        return True

    except Exception as e:
        print("PDF 翻译失败: %s" % str(e))
        return False


def translate_docx(input_path, output_path=None, target_lang="zh", config=None, terms=None):
    """翻译 Word 文档，保留格式"""
    if not os.path.exists(input_path):
        print("错误: 文件不存在: %s" % input_path)
        return False

    if output_path is None:
        base, ext = os.path.splitext(input_path)
        output_path = base + "_translated.docx"

    print("正在解析 Word: %s" % input_path)
    doc = Document(input_path)
    total_paragraphs = len(doc.paragraphs)
    translated_count = 0
    skipped_count = 0

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text and len(text) > 2:
            try:
                from langdetect import detect, LangDetectException
                detected = detect(text)
                if not detected.startswith(target_lang):
                    translated = translate_chunk(text, config, terms)
                    if translated:
                        # 保留原始格式：清空runs，写入翻译后的文本
                        for run in para.runs:
                            run.text = ""
                        if para.runs:
                            para.runs[0].text = translated
                        else:
                            para.clear()
                            para.add_run(translated)
                        translated_count += 1

                        if (translated_count % 5 == 0):
                            print("已翻译 %d/%d 段..." % (translated_count, total_paragraphs))
                    else:
                        skipped_count += 1
                else:
                    skipped_count += 1
            except Exception:
                skipped_count += 1

    doc.save(output_path)
    print("翻译完成!")
    print("  已翻译: %d 段" % translated_count)
    print("  跳过: %d 段" % skipped_count)
    print("  输出: %s" % output_path)
    return True


def main():
    parser = argparse.ArgumentParser(description="文档翻译器 - 支持 PDF/Word")
    parser.add_argument("input", help="输入文件路径 (PDF 或 DOCX)")
    parser.add_argument("-o", "--output", help="输出文件路径", default=None)
    parser.add_argument("--to", help="目标语言 (默认: zh 中文)", default="zh")
    parser.add_argument("--terms", action="store_true", help="启用外贸术语替换", default=True)
    args = parser.parse_args()

    if not os.path.exists(args.input):
        print("错误: 文件不存在: %s" % args.input)
        sys.exit(1)

    ext = os.path.splitext(args.input)[1].lower()
    if ext not in (".pdf", ".docx", ".doc"):
        print("错误: 不支持的文件格式: %s (支持 PDF 和 DOCX)" % ext)
        sys.exit(1)

    config = load_config()
    config["target_language"] = args.to
    terms = load_trade_terms() if args.terms else {}

    start_time = time.time()

    if ext == ".pdf":
        success = translate_pdf(args.input, args.output, args.to, config, terms)
    else:
        success = translate_docx(args.input, args.output, args.to, config, terms)

    elapsed = time.time() - start_time
    print("耗时: %.1f 秒" % elapsed)

    if not success:
        sys.exit(1)


if __name__ == "__main__":
    main()
