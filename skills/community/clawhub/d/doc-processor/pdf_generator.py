#!/usr/bin/env python3
"""PDF 生成模块 - v2.6.0"""

import logging
from pathlib import Path
from typing import Dict, Any

logger = logging.getLogger(__name__)

class PDFGenerator:
    def __init__(self, backend: str = 'auto'):
        self.backend = backend
        self._available_backends = self._check_backends()
        if backend == 'auto':
            if 'docx2pdf' in self._available_backends:
                self.backend = 'docx2pdf'
            elif 'weasyprint' in self._available_backends:
                self.backend = 'weasyprint'
            else:
                raise RuntimeError("没有可用的 PDF 生成后端")
        logger.info(f"PDF 生成器已初始化，后端：{self.backend}")
    
    def _check_backends(self) -> list:
        available = []
        try:
            import docx2pdf
            available.append('docx2pdf')
        except ImportError:
            pass
        try:
            import weasyprint
            available.append('weasyprint')
        except ImportError:
            pass
        return available
    
    def generate(self, output_path: str, content: Dict[str, Any]) -> str:
        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        logger.info(f"开始生成 PDF: {output.name}")
        if self.backend == 'docx2pdf':
            return self._generate_via_docx2pdf(output, content)
        elif self.backend == 'weasyprint':
            return self._generate_via_weasyprint(output, content)
        else:
            raise ValueError(f"不支持的后端：{self.backend}")
    
    def _generate_via_docx2pdf(self, output: Path, content: Dict) -> str:
        try:
            from docx2pdf import convert
        except ImportError:
            raise RuntimeError("docx2pdf 未安装：pip install docx2pdf")
        
        temp_docx = output.with_suffix('.docx')
        try:
            from docx import Document
            from docx.shared import Pt, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            section = doc.sections[0]
            section.page_width = Cm(21)
            section.page_height = Cm(29.7)
            
            if 'title' in content:
                heading = doc.add_heading(content['title'], 0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if 'paragraphs' in content:
                for para_text in content['paragraphs']:
                    doc.add_paragraph(para_text)
            
            doc.save(temp_docx)
            convert(str(temp_docx), str(output))
            logger.info(f"PDF 生成完成：{output.name}")
            return str(output.absolute())
        finally:
            if temp_docx.exists():
                temp_docx.unlink()
    
    def _generate_via_weasyprint(self, output: Path, content: Dict) -> str:
        try:
            from weasyprint import HTML
        except ImportError:
            raise RuntimeError("WeasyPrint 未安装：pip install weasyprint")
        
        html_content = self._generate_html(content)
        HTML(string=html_content).write_pdf(str(output))
        logger.info(f"PDF 生成完成：{output.name}")
        return str(output.absolute())
    
    def _generate_html(self, content: Dict) -> str:
        title = content.get('title', '')
        title_html = f"<h1>{title}</h1>" if title else ''
        paragraphs_html = ''.join(f"<p>{para}</p>" for para in content.get('paragraphs', []))
        return f"""<!DOCTYPE html><html lang="zh-CN"><head><meta charset="UTF-8"><title>{title}</title><style>@page {{ size: A4; margin: 2.54cm; }} body {{ font-family: "SimSun", serif; font-size: 12pt; line-height: 1.8; }} h1 {{ font-size: 22pt; text-align: center; }}</style></head><body>{title_html}{paragraphs_html}</body></html>"""

def generate_pdf(output_path: str, content: Dict, backend: str = 'auto') -> str:
    generator = PDFGenerator(backend=backend)
    return generator.generate(output_path, content)
