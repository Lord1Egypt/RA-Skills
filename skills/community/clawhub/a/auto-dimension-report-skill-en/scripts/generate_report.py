#!/usr/bin/env python3
"""
Review Report Generation Script v4
- Use shared rules to analyze XLSX / DOCX / OCR results
- Added checks: tight-boundary, precision inconsistency, bias, missing fields
- Supports md / json / summary output formats
"""

from __future__ import annotations

import argparse
import json
import os
import re
import sys
from collections import Counter, defaultdict
from pathlib import Path

import openpyxl
from docx import Document

from inspection_utils import (
    PASS_VALUES,
    build_merged_map,
    classify_judge_text,
    clean_text,
    contains_numeric_spec,
    count_decimals,
    evaluate_value_against_tolerance,
    extract_rows_from_sheet,
    get_cell_value,
    maybe_parse_float,
    parse_tolerance,
    summarize_precision,
)


OCR_KEYWORDS = ("seal", "stamp", "signature")
SOURCE_EXTS = {".pdf", ".xlsx", ".xlsm", ".docx"}
DEFAULT_METADATA_FIELDS = [
    ("Part Number", "I55"),
    ("Responsible Person", "J53"),
    ("Supplier Name", "J52"),
    ("Vehicle Model", "B55"),
]


def fmt_path(abspath, base_dir):
    try:
        rel = os.path.relpath(abspath, base_dir)
        return rel.replace("\\", "/")
    except ValueError:
        return abspath


def expected_output_name(source_name):
    source_path = Path(source_name)
    if source_path.suffix.lower() == ".pdf":
        return "{0}.docx".format(source_name)
    return source_name


def load_image_index(index_path):
    rows = []
    missing = []
    workbook = openpyxl.load_workbook(index_path, data_only=True)
    try:
        ws = workbook.active
        for row_idx in range(2, ws.max_row + 1):
            src = ws.cell(row_idx, 1).value or ""
            relp = ws.cell(row_idx, 2).value or ""
            absp = ws.cell(row_idx, 3).value or ""
            rows.append({"src": src, "rel": relp, "abs": absp})
            if absp and not os.path.exists(absp):
                missing.append(absp)
    finally:
        workbook.close()
    return rows, missing


def scan_ocr_keywords(task_dir):
    results = []
    imagetomd_dir = task_dir / "imagetomd"
    if not imagetomd_dir.exists():
        return results
    for md_path in sorted(imagetomd_dir.rglob("*.md")):
        try:
            content = md_path.read_text(encoding="utf-8")
        except OSError:
            continue
        hits = [keyword for keyword in OCR_KEYWORDS if keyword.lower() in content.lower()]
        if not hits:
            continue
        preview = ""
        for keyword in hits:
            match = re.search(r".{0,20}" + re.escape(keyword) + r".{0,20}", content, re.DOTALL | re.IGNORECASE)
            if match:
                preview = match.group(0).replace("\n", " ")
                break
        results.append(
            {
                "path": fmt_path(str(md_path), str(task_dir)),
                "keywords": hits,
                "preview": preview,
            }
        )
    return results


def parse_metadata_fields(items):
    if not items:
        return list(DEFAULT_METADATA_FIELDS)
    parsed = []
    for item in items:
        if ":" not in item:
            continue
        label, cell = item.split(":", 1)
        label = label.strip()
        cell = cell.strip().upper()
        if label and cell:
            parsed.append((label, cell))
    return parsed or list(DEFAULT_METADATA_FIELDS)


def detect_missing_metadata_xlsx(ws, merged_map, metadata_fields):
    issues = []
    for label, cell_ref in metadata_fields:
        value = clean_text(ws[cell_ref].value)
        if value:
            continue
        row_idx = ws[cell_ref].row
        col_idx = ws[cell_ref].column
        left_value = clean_text(get_cell_value(ws, row_idx, max(col_idx - 1, 1), merged_map))
        above_value = clean_text(get_cell_value(ws, max(row_idx - 1, 1), col_idx, merged_map))
        context = left_value or above_value
        issues.append(
            {
                "field": label,
                "cell": cell_ref,
                "context": context,
                "message": "Field is missing or empty",
            }
        )
    return issues


def analyze_measurement_rows(rows, checks):
    issues = []
    method_precision = summarize_precision(rows)
    for row in rows:
        tolerance = parse_tolerance(row.get("std"))
        numeric_values = []
        pass_like_values = []
        decimals = []
        for value in row.get("values", []):
            text = clean_text(value)
            if not text:
                continue
            number = maybe_parse_float(value)
            if number is not None:
                numeric_values.append(number)
                decimal_count = count_decimals(value)
                if decimal_count is not None:
                    decimals.append(decimal_count)
            elif text in PASS_VALUES or text.upper() in PASS_VALUES:
                pass_like_values.append(text)

        if checks["numeric_std_but_ok"] and tolerance and pass_like_values and not numeric_values:
            issues.append(
                {
                    "type": "numeric_std_but_ok",
                    "severity": "medium",
                    "seq": row["seq"],
                    "sheet": row["sheet"],
                    "item": row["item"],
                    "message": "Standard is numeric tolerance, but data column only contains OK/checkmark pass symbols",
                }
            )

        if not tolerance or not numeric_values:
            continue

        margins = []
        failed_values = []
        for number in numeric_values:
            in_range, margin = evaluate_value_against_tolerance(number, tolerance)
            if not in_range:
                failed_values.append(number)
            if margin is not None:
                margins.append(margin)

        if checks["tight_limit"] and margins:
            min_margin = min(margins)
            if min_margin <= checks["tight_limit_threshold"]:
                issues.append(
                    {
                        "type": "tight_limit",
                        "severity": "medium",
                        "seq": row["seq"],
                        "sheet": row["sheet"],
                        "item": row["item"],
                        "message": "Measured value close to tolerance boundary, min margin {0:.3f}mm".format(min_margin),
                    }
                )

        if checks["precision_inconsistency"] and decimals and row.get("method"):
            method_info = method_precision.get(row["method"])
            if method_info:
                dominant = method_info["dominant_decimals"]
                if any(abs(value - dominant) > 1 for value in decimals):
                    issues.append(
                        {
                            "type": "precision_inconsistency",
                            "severity": "low",
                            "seq": row["seq"],
                            "sheet": row["sheet"],
                            "item": row["item"],
                            "message": "Inconsistent decimals within same method, current={0}, dominant={1}".format(decimals, dominant),
                        }
                    )

        nominal = tolerance.get("nominal")
        lower = tolerance.get("lower")
        upper = tolerance.get("upper")
        if checks["all_bias"] and nominal is not None and len(numeric_values) >= 3 and lower is not None and upper is not None:
            if all(value > nominal for value in numeric_values) or all(value < nominal for value in numeric_values):
                bias_direction = "All positive bias" if all(value > nominal for value in numeric_values) else "All negative bias"
                tolerance_range = upper - lower
                avg_deviation = abs(sum(numeric_values) / len(numeric_values) - nominal)
                deviation_pct = (avg_deviation / tolerance_range * 100.0) if tolerance_range else 0.0
                if deviation_pct >= checks["all_bias_threshold_pct"]:
                    issues.append(
                        {
                            "type": "all_bias",
                            "severity": "medium",
                            "seq": row["seq"],
                            "sheet": row["sheet"],
                            "item": row["item"],
                            "message": "{0}, avg deviation occupies {1:.1f}% of tolerance band".format(bias_direction, deviation_pct),
                        }
                    )

        judge_kind, normalized_judge = classify_judge_text(row.get("judge"))
        if failed_values and judge_kind == "pass":
            issues.append(
                {
                    "type": "judge_mismatch",
                    "severity": "high",
                    "seq": row["seq"],
                    "sheet": row["sheet"],
                    "item": row["item"],
                    "message": "Measured value out of tolerance {0}, but judgement is {1}".format(failed_values, normalized_judge),
                }
            )
    return issues


def analyze_xlsx(path, base_dir, metadata_fields, checks):
    result = {
        "path": fmt_path(str(path), str(base_dir)),
        "filename": path.name,
        "sheets": [],
        "issues": [],
        "rows": [],
        "stats": {},
    }
    workbook = openpyxl.load_workbook(path, data_only=True)
    try:
        for sheet_name in workbook.sheetnames:
            ws = workbook[sheet_name]
            merged_map = build_merged_map(ws)
            extracted = extract_rows_from_sheet(ws, sheet_name)
            rows = extracted["rows"]
            sheet_info = {
                "name": sheet_name,
                "rows": ws.max_row,
                "cols": ws.max_column,
                "merged": len(ws.merged_cells.ranges),
                "layout": extracted["layout"],
                "data_rows": len(rows),
                "metadata_issues": detect_missing_metadata_xlsx(ws, merged_map, metadata_fields) if checks["missing_metadata"] else [],
            }
            result["sheets"].append(sheet_info)
            result["rows"].extend(rows)
            for issue in sheet_info["metadata_issues"]:
                result["issues"].append(
                    {
                        "type": "missing_metadata",
                        "severity": "high",
                        "sheet": sheet_name,
                        "item": issue["field"],
                        "message": "{0} {1} is empty".format(issue["field"], issue["cell"]),
                    }
                )
    finally:
        workbook.close()

    result["issues"].extend(analyze_measurement_rows(result["rows"], checks))
    result["stats"] = {
        "row_count": len(result["rows"]),
        "issue_count": len(result["issues"]),
        "methods": dict(Counter((row.get("method") or "Not specified") for row in result["rows"])),
        "tolerance_types": dict(
            Counter(
                (parse_tolerance(row.get("std")) or {}).get("kind", "unknown")
                for row in result["rows"]
                if row.get("std")
            )
        ),
    }
    return result


def collect_doc_text(doc):
    text_chunks = []
    for paragraph in doc.paragraphs:
        if clean_text(paragraph.text):
            text_chunks.append(clean_text(paragraph.text))
    for table in doc.tables:
        for row in table.rows:
            cells = [clean_text(cell.text) for cell in row.cells]
            merged = " | ".join([cell for cell in cells if cell])
            if merged:
                text_chunks.append(merged)
    return "\n".join(text_chunks)


def analyze_docx(path, base_dir, checks):
    result = {
        "path": fmt_path(str(path), str(base_dir)),
        "filename": path.name,
        "tables": 0,
        "images": 0,
        "paragraphs": 0,
        "total_chars": 0,
        "issues": [],
        "samples": [],
    }
    doc = Document(path)
    result["tables"] = len(doc.tables)
    result["paragraphs"] = len(doc.paragraphs)
    result["total_chars"] = sum(len(paragraph.text) for paragraph in doc.paragraphs)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            images = run._element.findall(
                ".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline"
            )
            result["images"] += len(images)

    full_text = collect_doc_text(doc)
    for table_index, table in enumerate(doc.tables[:2]):
        sample_rows = []
        for row in table.rows[:4]:
            sample_rows.append([clean_text(cell.text)[:25] for cell in row.cells[:8]])
        result["samples"].append(
            {
                "table_idx": table_index,
                "rows": len(table.rows),
                "cols": len(table.columns),
                "sample": sample_rows,
            }
        )

    if checks["missing_metadata"]:
        for keyword in ["Part Number", "Responsible Person", "Supplier", "Vehicle Model"]:
            if keyword.lower() in full_text.lower():
                pattern = re.compile(r"{0}\s*[:：]?\s*$".format(re.escape(keyword)), re.MULTILINE | re.IGNORECASE)
                if pattern.search(full_text):
                    result["issues"].append(
                        {
                            "type": "missing_metadata",
                            "severity": "medium",
                            "item": keyword,
                            "message": "Document contains `{0}` label, but value appears unfilled".format(keyword),
                        }
                    )
            else:
                result["issues"].append(
                    {
                        "type": "missing_metadata",
                        "severity": "low",
                        "item": keyword,
                        "message": "Document does not contain `{0}` field".format(keyword),
                    }
                )
    return result


def build_review_data(task_dir, metadata_fields, checks):
    output_dir = task_dir / "output"
    files = sorted(
        file_path.name for file_path in output_dir.iterdir() if file_path.is_file() and not file_path.name.startswith("~")
    )
    xlsx_files = [name for name in files if name.lower().endswith(".xlsx")]
    docx_files = [name for name in files if name.lower().endswith(".docx")]
    src_files = sorted(
        file_path.name
        for file_path in task_dir.iterdir()
        if file_path.is_file() and file_path.suffix.lower() in SOURCE_EXTS and not file_path.name.startswith("~")
    )

    data = {
        "task_dir": str(task_dir),
        "output_dir": str(output_dir),
        "files": [],
        "xlsx": [],
        "docx": [],
        "image_index": {"rows": [], "missing": []},
        "ocr_hits": scan_ocr_keywords(task_dir),
        "source_mapping": [],
        "summary": {},
    }

    for file_name in files:
        file_path = output_dir / file_name
        size_bytes = file_path.stat().st_size
        data["files"].append(
            {
                "filename": file_name,
                "type": file_path.suffix.lower(),
                "size_bytes": size_bytes,
            }
        )

    for file_name in xlsx_files:
        data["xlsx"].append(analyze_xlsx(output_dir / file_name, task_dir, metadata_fields, checks))
    for file_name in docx_files:
        data["docx"].append(analyze_docx(output_dir / file_name, task_dir, checks))

    idx_fp = output_dir / "_ImageIndex.xlsx"
    if idx_fp.exists():
        rows, missing = load_image_index(idx_fp)
        data["image_index"] = {"rows": rows, "missing": missing}

    unmatched_sources = []
    matched_count = 0
    for src_name in src_files:
        expected = expected_output_name(src_name)
        matched = expected in files
        if matched:
            matched_count += 1
        else:
            unmatched_sources.append(src_name)
        data["source_mapping"].append(
            {
                "source": src_name,
                "expected_output": expected,
                "matched": matched,
            }
        )

    xlsx_issue_count = sum(len(item["issues"]) for item in data["xlsx"])
    docx_issue_count = sum(len(item["issues"]) for item in data["docx"])
    data["summary"] = {
        "file_count": len(files),
        "xlsx_count": len(data["xlsx"]),
        "docx_count": len(data["docx"]),
        "ocr_hit_count": len(data["ocr_hits"]),
        "image_index_missing": len(data["image_index"]["missing"]),
        "unmatched_sources": len(unmatched_sources),
        "matched_sources": matched_count,
        "xlsx_issue_count": xlsx_issue_count,
        "docx_issue_count": docx_issue_count,
        "total_issue_count": xlsx_issue_count + docx_issue_count + len(data["image_index"]["missing"]) + len(unmatched_sources),
    }
    return data


def render_markdown(data):
    lines = []

    def w(text=""):
        lines.append(text)

    w("# Dimension Inspection Report — Review Report")
    w()
    w("**Data directory:** `{0}/`  ".format(fmt_path(data["task_dir"], data["task_dir"])))
    w("**Output directory:** `{0}/`  ".format(fmt_path(data["output_dir"], data["task_dir"])))
    w("**Files involved:** {0}".format(len(data["files"])))
    w()
    w("---")
    w()

    w("## 1. File Overview")
    w()
    w("| File | Type | Size |")
    w("|------|------|------|")
    for item in data["files"]:
        size_bytes = item["size_bytes"]
        size_str = "{0:.1f}KB".format(size_bytes / 1024.0) if size_bytes < 1024 * 1024 else "{0:.1f}MB".format(size_bytes / 1024.0 / 1024.0)
        w("| {0} | {1} | {2} |".format(item["filename"], item["type"].upper(), size_str))
    w()

    w("## 2. Spreadsheet Analysis (XLSX)")
    w()
    if not data["xlsx"]:
        w("No XLSX output")
        w()
    for item in data["xlsx"]:
        w("### {0}".format(item["filename"]))
        w()
        w("- Valid data rows: {0}".format(item["stats"]["row_count"]))
        w("- Anomalies: {0}".format(item["stats"]["issue_count"]))
        if item["stats"]["methods"]:
            w("- Method distribution: " + "; ".join("{0}={1}".format(k, v) for k, v in sorted(item["stats"]["methods"].items())))
        if item["stats"]["tolerance_types"]:
            w("- Tolerance type distribution: " + "; ".join("{0}={1}".format(k, v) for k, v in sorted(item["stats"]["tolerance_types"].items())))
        for sheet in item["sheets"]:
            layout = sheet["layout"]
            w(
                "- Sheet `{0}`: rows={1} cols={2} merged={3} header={4} data_row={5} data_col={6}-{7} judge={8}".format(
                    sheet["name"],
                    sheet["rows"],
                    sheet["cols"],
                    sheet["merged"],
                    layout["header_row"],
                    layout["data_start_row"],
                    layout["data_start"],
                    layout["data_start"] + layout["data_cols"] - 1,
                    layout["judge_col"],
                )
            )
        if item["issues"]:
            w()
            w("**Key anomalies:**")
            for issue in item["issues"][:30]:
                w(
                    "- [{0}] `{1}` sheet=`{2}` item=`{3}` -> {4}".format(
                        issue["severity"].upper(),
                        issue["type"],
                        issue.get("sheet", ""),
                        issue.get("item", "")[:24],
                        issue["message"],
                    )
                )
        w()

    w("## 3. Word Document Analysis (DOCX)")
    w()
    if not data["docx"]:
        w("No DOCX output")
        w()
    for item in data["docx"]:
        w("### {0}".format(item["filename"]))
        w()
        w("- Paragraphs: {0} | Total chars: {1} | Tables: {2} | Images: {3}".format(item["paragraphs"], item["total_chars"], item["tables"], item["images"]))
        for sample in item["samples"]:
            w("- Table #{0}: {1} rows x {2} cols".format(sample["table_idx"] + 1, sample["rows"], sample["cols"]))
        if item["issues"]:
            w("- Metadata risks: {0}".format(len(item["issues"])))
            for issue in item["issues"][:10]:
                w("  - [{0}] `{1}` -> {2}".format(issue["severity"].upper(), issue.get("item", ""), issue["message"]))
        w()

    w("## 4. Image Index Verification")
    w()
    idx_rows = data["image_index"]["rows"]
    missing = data["image_index"]["missing"]
    if idx_rows:
        w("- Index entries: {0}".format(len(idx_rows)))
        by_src = defaultdict(int)
        for row in idx_rows:
            by_src[row["src"]] += 1
        for src, count in sorted(by_src.items()):
            w("- `{0}` -> {1} images".format(src[:50], count))
        if missing:
            w("- Missing image files: {0}".format(len(missing)))
            for missing_path in missing[:5]:
                w("  - `{0}`".format(missing_path))
        else:
            w("- All indexed image paths exist")
    else:
        w("`_ImageIndex.xlsx` not found. Cannot verify image index integrity.")
    w()

    w("## 5. Seal/Signature OCR Detection")
    w()
    if data["ocr_hits"]:
        w("**{0}** potential seal/signature keyword hits detected:".format(len(data["ocr_hits"])))
        for hit in data["ocr_hits"][:20]:
            preview = " | snippet: `{0}`".format(hit["preview"]) if hit["preview"] else ""
            w("- `{0}` | hits: `{1}`{2}".format(hit["path"], ", ".join(hit["keywords"]), preview))
        w("- Manual review of hit images is recommended; do not treat OCR keywords as conclusive seal evidence.")
    else:
        imagetomd_dir = Path(data["task_dir"]) / "imagetomd"
        w("No seal/stamp/signature keywords detected in OCR text" if imagetomd_dir.exists() else "`imagetomd/` does not exist — OCR detection not yet performed")
    w()

    w("## 6. Source-to-Output Mapping")
    w()
    for item in data["source_mapping"]:
        if item["matched"]:
            w("- `{0}` -> `output/{1}`".format(item["source"], item["expected_output"]))
        else:
            w("- `{0}` -> `output/{1}` not found".format(item["source"], item["expected_output"]))
    w()

    w("---")
    w()
    w("## 7. Review Summary")
    w()
    issues = []
    if data["summary"]["unmatched_sources"]:
        issues.append("Source files without corresponding output: {0}".format(data["summary"]["unmatched_sources"]))
    if data["summary"]["image_index_missing"]:
        issues.append("Missing image index files: {0}".format(data["summary"]["image_index_missing"]))
    if data["summary"]["xlsx_issue_count"]:
        issues.append("XLSX risk items: {0}".format(data["summary"]["xlsx_issue_count"]))
    if data["summary"]["docx_issue_count"]:
        issues.append("DOCX metadata risk items: {0}".format(data["summary"]["docx_issue_count"]))
    if data["summary"]["ocr_hit_count"]:
        issues.append("OCR keyword hits: {0}".format(data["summary"]["ocr_hit_count"]))

    if issues:
        w("Review found the following issues — manual re-check recommended:")
        for idx, issue in enumerate(issues, start=1):
            w("{0}. {1}".format(idx, issue))
    else:
        w("No blocking issues found — basic structure check and file mapping passed.")
    w()
    w("| Check | Status |")
    w("|--------|------|")
    w("| File integrity | {0} |".format("Missing outputs" if data["summary"]["unmatched_sources"] else "All aligned"))
    w("| Image index | {0} |".format("Missing images" if data["summary"]["image_index_missing"] else ("Accessible" if idx_rows else "Index not generated")))
    w("| XLSX deep check | {0} |".format("Risk items found" if data["summary"]["xlsx_issue_count"] else ("No anomalies" if data["xlsx"] else "No XLSX")))
    w("| DOCX metadata check | {0} |".format("Risk items found" if data["summary"]["docx_issue_count"] else ("Parsed successfully" if data["docx"] else "No DOCX")))
    w("| Seal/Signature OCR | {0} |".format("Keywords detected" if data["summary"]["ocr_hit_count"] else ("No keywords detected" if (Path(data["task_dir"]) / "imagetomd").exists() else "OCR not executed")))
    w()
    w("---")
    w("*Report auto-generated by script -- `scripts/generate_report.py`*")
    return "\n".join(lines)


def render_summary(data):
    lines = []
    lines.append("Dimension Review Summary")
    lines.append("Files: {0}, XLSX: {1}, DOCX: {2}".format(data["summary"]["file_count"], data["summary"]["xlsx_count"], data["summary"]["docx_count"]))
    lines.append("Total risk items: {0}".format(data["summary"]["total_issue_count"]))
    lines.append("XLSX risks: {0}, DOCX metadata risks: {1}".format(data["summary"]["xlsx_issue_count"], data["summary"]["docx_issue_count"]))
    lines.append("Image index missing: {0}, Source files not aligned: {1}".format(data["summary"]["image_index_missing"], data["summary"]["unmatched_sources"]))
    lines.append("OCR keyword hits: {0}".format(data["summary"]["ocr_hit_count"]))
    top_xlsx_issues = []
    for item in data["xlsx"]:
        top_xlsx_issues.extend(item["issues"][:3])
    if top_xlsx_issues:
        lines.append("Key items:")
        for issue in top_xlsx_issues[:5]:
            lines.append("- [{0}] seq={1} item={2} -> {3}".format(issue["type"], issue.get("seq", "-"), issue.get("item", ""), issue["message"]))
    return "\n".join(lines) + "\n"


def main():
    parser = argparse.ArgumentParser(description="Generate Review Report")
    parser.add_argument("--dir", "-d", required=True, help="Task folder path (containing output/)")
    parser.add_argument("--output", "-o", default="", help="Report output path (default: task-dir/ReviewReport.md)")
    parser.add_argument("--format", choices=["md", "json", "summary"], default="md", help="Output format")
    parser.add_argument("--metadata-field", action="append", default=[], help="Required metadata fields, format: Field:Cell e.g. PartNumber:I55")
    parser.add_argument("--tight-limit-threshold", type=float, default=0.05, help="Tight boundary threshold, default 0.05")
    parser.add_argument("--all-bias-threshold-pct", type=float, default=50.0, help="Bias alert threshold, default 50")
    parser.add_argument("--disable-tight-limit", action="store_true", help="Disable tight boundary detection")
    parser.add_argument("--disable-precision-inconsistency", action="store_true", help="Disable precision inconsistency detection")
    parser.add_argument("--disable-all-bias", action="store_true", help="Disable all-positive/all-negative bias detection")
    parser.add_argument("--disable-numeric-std-but-ok", action="store_true", help="Disable numeric tolerance but OK-only check")
    parser.add_argument("--disable-missing-metadata", action="store_true", help="Disable required field check")
    args = parser.parse_args()

    task_dir = Path(args.dir).resolve()
    output_dir = task_dir / "output"
    if not output_dir.is_dir():
        print("[Error] output/ not found: {0}".format(output_dir), file=sys.stderr)
        sys.exit(1)

    checks = {
        "tight_limit": not args.disable_tight_limit,
        "tight_limit_threshold": args.tight_limit_threshold,
        "precision_inconsistency": not args.disable_precision_inconsistency,
        "all_bias": not args.disable_all_bias,
        "all_bias_threshold_pct": args.all_bias_threshold_pct,
        "numeric_std_but_ok": not args.disable_numeric_std_but_ok,
        "missing_metadata": not args.disable_missing_metadata,
    }
    metadata_fields = parse_metadata_fields(args.metadata_field)
    review_data = build_review_data(task_dir, metadata_fields, checks)

    if args.output:
        report_path = Path(args.output).resolve()
    else:
        suffix = {"md": ".md", "json": ".json", "summary": ".txt"}[args.format]
        report_path = task_dir / ("ReviewReport" + suffix)

    if args.format == "json":
        report_path.write_text(json.dumps(review_data, ensure_ascii=False, indent=2), encoding="utf-8")
    elif args.format == "summary":
        report_path.write_text(render_summary(review_data), encoding="utf-8")
    else:
        report_path.write_text(render_markdown(review_data), encoding="utf-8")

    print("[Done] Review report generated: {0}".format(report_path))
    print()
    print("[File Locations] (Copy and paste into File Explorer)")
    print("  [Report] Review report:    {0}".format(report_path))
    print("  [Source] Source directory:  {0}".format(task_dir))
    print("  [Output] output:            {0}".format(output_dir))
    print("  [Images] image:             {0}".format(task_dir / "image"))
    print("  [OCR]    imagetomd:         {0}".format(task_dir / "imagetomd"))
    print()
    print("[Tip] Select the path above -> Ctrl+C -> paste into File Explorer address bar")
    print()


if __name__ == "__main__":
    main()
