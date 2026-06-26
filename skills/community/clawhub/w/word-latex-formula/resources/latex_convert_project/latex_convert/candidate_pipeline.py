from __future__ import annotations

from copy import deepcopy
from dataclasses import dataclass, asdict, field
import hashlib
import json
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
import re
import ssl
import time
from typing import Any, Callable
from urllib import error, request
from zipfile import ZIP_DEFLATED, ZipFile

from lxml import etree

from .detect import split_formula_spans
from .math_parser import parse_formula
from .omml import node_to_omath, W_NS, w_tag
from .ooxml_transform import (
    PROCESSABLE_PARTS,
    XML_SPACE,
    _join_split_nary_scripts,
    _make_text_run,
    _run_is_simple_text,
    is_bibliography_entry,
    is_bibliography_heading,
    merge_adjacent_text_runs,
)


MODES = {"aggressive", "balanced", "conservative", "display-only"}


@dataclass
class Candidate:
    id: str
    part: str
    paragraph_index: int
    run_index: int
    segment_index: int
    start: int
    end: int
    text: str
    paragraph_text: str
    context: str
    confidence: float
    default_action: str
    reason: str


@dataclass
class ScanResult:
    input: str
    prepared_docx: str
    engine: str
    mode: str
    source_sha256: str
    candidates: list[Candidate] = field(default_factory=list)
    paragraphs_skipped: int = 0
    runs_merged: int = 0

    def to_dict(self) -> dict[str, Any]:
        return {
            "input": self.input,
            "prepared_docx": self.prepared_docx,
            "engine": self.engine,
            "mode": self.mode,
            "source_sha256": self.source_sha256,
            "summary": {
                "candidates": len(self.candidates),
                "default_convert": sum(1 for c in self.candidates if c.default_action == "convert"),
                "default_review": sum(1 for c in self.candidates if c.default_action == "review"),
                "paragraphs_skipped": self.paragraphs_skipped,
                "runs_merged": self.runs_merged,
            },
            "candidates": [asdict(candidate) for candidate in self.candidates],
        }


@dataclass
class ApplyStats:
    formulas_converted: int = 0
    formulas_kept: int = 0
    placeholders_inserted: int = 0
    candidates_seen: int = 0
    paragraphs_skipped: int = 0
    runs_merged: int = 0
    parts_changed: list[str] = field(default_factory=list)
    failed: list[dict[str, str]] = field(default_factory=list)
    samples: list[str] = field(default_factory=list)


def scan_docx_candidates(
    src: Path,
    *,
    input_path: Path | None = None,
    engine: str = "already-docx",
    mode: str = "balanced",
    skip_bibliography: bool = True,
) -> ScanResult:
    _validate_mode(mode)
    result = ScanResult(
        input=str((input_path or src).resolve()),
        prepared_docx=str(src.resolve()),
        engine=engine,
        mode=mode,
        source_sha256=_sha256(src),
    )
    counter = 0
    with ZipFile(src, "r") as zin:
        for name in _dynamic_parts(zin):
            try:
                root = etree.fromstring(zin.read(name))
            except etree.XMLSyntaxError:
                continue
            result.runs_merged += merge_adjacent_text_runs(root)
            _join_split_nary_scripts(root)
            in_bibliography = False
            for p_idx, paragraph in enumerate(root.xpath(".//w:p", namespaces={"w": W_NS}), 1):
                paragraph_text = _paragraph_text(paragraph)
                skip = skip_bibliography and (
                    in_bibliography
                    or is_bibliography_heading(paragraph_text.strip())
                    or is_bibliography_entry(paragraph_text.strip())
                )
                if skip:
                    in_bibliography = in_bibliography or is_bibliography_heading(paragraph_text.strip())
                    result.paragraphs_skipped += 1
                    continue
                for run_idx, run in enumerate(paragraph.xpath("./w:r", namespaces={"w": W_NS}), 1):
                    t = _single_text_child_for_scan(run)
                    if t is None:
                        continue
                    segment_idx = 0
                    for value, is_formula, start, end in split_formula_spans(t.text or ""):
                        if not is_formula:
                            continue
                        segment_idx += 1
                        counter += 1
                        confidence, reason = score_candidate(value, paragraph_text)
                        default_action = default_action_for(value, paragraph_text, confidence, mode)
                        result.candidates.append(
                            Candidate(
                                id=f"F{counter:05d}",
                                part=name,
                                paragraph_index=p_idx,
                                run_index=run_idx,
                                segment_index=segment_idx,
                                start=start,
                                end=end,
                                text=value,
                                paragraph_text=paragraph_text[:1000],
                                context=_context(paragraph_text, value),
                                confidence=confidence,
                                default_action=default_action,
                                reason=reason,
                            )
                        )
    return result


def write_scan_json(result: ScanResult, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(result.to_dict(), ensure_ascii=False, indent=2), encoding="utf-8")


def load_decisions(path: Path) -> dict[str, str]:
    payload = json.loads(path.read_text(encoding="utf-8"))
    if isinstance(payload, list):
        items = payload
    else:
        items = payload.get("candidates", [])
    decisions: dict[str, str] = {}
    for item in items:
        if not isinstance(item, dict) or "id" not in item:
            continue
        action = item.get("action") or item.get("decision") or item.get("default_action") or "review"
        decisions[str(item["id"])] = normalize_action(str(action))
    return decisions


def apply_docx_decisions(
    src: Path,
    dst: Path,
    decisions: dict[str, str] | None = None,
    *,
    mode: str = "balanced",
    skip_bibliography: bool = True,
    placeholder: bool = False,
    formula_text_overrides: dict[str, str] | None = None,
) -> ApplyStats:
    _validate_mode(mode)
    decisions = decisions or {}
    formula_text_overrides = formula_text_overrides or {}
    stats = ApplyStats()
    dst.parent.mkdir(parents=True, exist_ok=True)
    counter = 0
    with ZipFile(src, "r") as zin, ZipFile(dst, "w", ZIP_DEFLATED) as zout:
        changed_payloads: dict[str, bytes] = {}
        for name in _dynamic_parts(zin):
            try:
                root = etree.fromstring(zin.read(name))
            except etree.XMLSyntaxError:
                continue
            before = (stats.formulas_converted, stats.placeholders_inserted, stats.runs_merged)
            stats.runs_merged += merge_adjacent_text_runs(root)
            _join_split_nary_scripts(root)
            in_bibliography = False
            for paragraph in root.xpath(".//w:p", namespaces={"w": W_NS}):
                paragraph_text = _paragraph_text(paragraph)
                skip = skip_bibliography and (
                    in_bibliography
                    or is_bibliography_heading(paragraph_text.strip())
                    or is_bibliography_entry(paragraph_text.strip())
                )
                if skip:
                    in_bibliography = in_bibliography or is_bibliography_heading(paragraph_text.strip())
                    stats.paragraphs_skipped += 1
                    continue
                counter = _apply_paragraph(
                    paragraph,
                    counter,
                    decisions,
                    stats,
                    mode=mode,
                    placeholder=placeholder,
                    formula_text_overrides=formula_text_overrides,
                )
            after = (stats.formulas_converted, stats.placeholders_inserted, stats.runs_merged)
            if after != before:
                stats.parts_changed.append(name)
                changed_payloads[name] = etree.tostring(
                    root,
                    xml_declaration=True,
                    encoding="UTF-8",
                    standalone=False,
                )
        for item in zin.infolist():
            payload = changed_payloads.get(item.filename)
            zout.writestr(item, payload if payload is not None else zin.read(item.filename))
    return stats


def write_review_docx(scan: ScanResult, path: Path) -> None:
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover - environment fallback
        raise RuntimeError("python-docx is required to write review DOCX files") from exc

    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("Formula Candidate Review", level=1)
    summary = scan.to_dict()["summary"]
    doc.add_paragraph(
        f"Candidates: {summary['candidates']} | Default convert: {summary['default_convert']} | "
        f"Default review: {summary['default_review']} | Skipped paragraphs: {summary['paragraphs_skipped']}"
    )
    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    headers = ["ID", "Action", "Confidence", "Formula", "Context", "Part", "Reason"]
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
    for candidate in scan.candidates:
        row = table.add_row().cells
        row[0].text = candidate.id
        row[1].text = candidate.default_action
        row[2].text = f"{candidate.confidence:.2f}"
        row[3].text = candidate.text
        row[4].text = candidate.context
        row[5].text = f"{candidate.part} p{candidate.paragraph_index} r{candidate.run_index}"
        row[6].text = candidate.reason
    doc.save(path)


def decide_with_ai(
    scan: ScanResult,
    *,
    api_key: str,
    base_url: str,
    model: str,
    timeout_seconds: int = 60,
    batch_size: int = 10,
    max_workers: int = 5,
    retries: int = 1,
    failure_fallback: str = "rule",
    progress_callback: Callable[[int, int, Candidate, str], None] | None = None,
    failure_callback: Callable[[int, list[Candidate], Exception, str], None] | None = None,
) -> dict[str, str]:
    candidates = scan.candidates
    if not candidates:
        return {}
    failure_fallback = _normalize_ai_failure_fallback(failure_fallback)
    chunks = [candidates[start : start + batch_size] for start in range(0, len(candidates), batch_size)]
    workers = max(1, min(max_workers, len(chunks)))
    decisions: dict[str, str] = {}
    completed = 0
    total = len(candidates)
    with ThreadPoolExecutor(max_workers=workers) as pool:
        futures = {
            pool.submit(
                _decide_candidates_with_ai,
                chunk,
                api_key=api_key,
                base_url=base_url,
                model=model,
                timeout_seconds=timeout_seconds,
                retries=retries,
            ): index
            for index, chunk in enumerate(chunks)
        }
        for future in as_completed(futures):
            index = futures[future]
            chunk = chunks[index]
            progress_suffix = ""
            try:
                chunk_decisions = future.result()
            except Exception as exc:
                chunk_decisions = _fallback_decisions_for_chunk(chunk, failure_fallback)
                progress_suffix = f" (fallback-{failure_fallback})"
                if failure_callback:
                    failure_callback(index + 1, chunk, exc, failure_fallback)
            decisions.update(chunk_decisions)
            if progress_callback:
                for candidate in chunk:
                    completed += 1
                    progress_callback(
                        completed,
                        total,
                        candidate,
                        chunk_decisions.get(candidate.id, "review") + progress_suffix,
                    )
    return decisions


def _decide_candidates_with_ai(
    candidates: list[Candidate],
    *,
    api_key: str,
    base_url: str,
    model: str,
    timeout_seconds: int,
    retries: int,
) -> dict[str, str]:
    endpoint = base_url.rstrip("/") + "/chat/completions"
    compact = [
        {
            "id": c.id,
            "text": c.text,
            "context": c.context,
            "confidence": c.confidence,
            "default_action": c.default_action,
        }
        for c in candidates
    ]
    prompt = (
        "You classify candidate strings from academic Word manuscripts. "
        "Return strict JSON only: an array of objects with id and action. "
        "action must be one of convert, keep, review. Convert real math/science formulas; "
        "keep DOI, URL, references, page ranges, ordinary prose, dates, and citations.\n\n"
        + json.dumps(compact, ensure_ascii=False)
    )
    body = json.dumps(
        {
            "model": model,
            "messages": [
                {"role": "system", "content": "You are a precise formula-candidate reviewer."},
                {"role": "user", "content": prompt},
            ],
            "temperature": 0,
        },
        ensure_ascii=False,
    ).encode("utf-8")
    req = request.Request(
        endpoint,
        data=body,
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        method="POST",
    )
    payload = _post_ai_request(req, timeout_seconds=timeout_seconds, retries=retries)
    content = payload["choices"][0]["message"]["content"]
    decisions_raw = json.loads(_extract_json(content))
    return {str(item["id"]): normalize_action(str(item["action"])) for item in decisions_raw}


def _post_ai_request(request_obj: request.Request, *, timeout_seconds: int, retries: int) -> dict[str, Any]:
    last_exc: Exception | None = None
    for attempt in range(max(0, retries) + 1):
        try:
            with request.urlopen(request_obj, timeout=timeout_seconds, context=_ssl_context()) as resp:
                return json.loads(resp.read().decode("utf-8"))
        except (error.HTTPError, error.URLError, TimeoutError) as exc:
            last_exc = exc
            if attempt >= retries:
                break
            time.sleep(min(2**attempt, 8))
    assert last_exc is not None
    raise last_exc


def _normalize_ai_failure_fallback(value: str) -> str:
    value = value.strip().lower()
    if value not in {"rule", "keep", "review"}:
        raise ValueError("failure_fallback must be one of rule, keep, review")
    return value


def _fallback_decisions_for_chunk(candidates: list[Candidate], strategy: str) -> dict[str, str]:
    if strategy == "rule":
        return {candidate.id: normalize_action(candidate.default_action) for candidate in candidates}
    return {candidate.id: strategy for candidate in candidates}


def _ssl_context() -> ssl.SSLContext:
    try:
        import certifi

        return ssl.create_default_context(cafile=certifi.where())
    except ImportError:
        return ssl.create_default_context()


def write_decisions_json(scan: ScanResult, decisions: dict[str, str], path: Path) -> None:
    payload = scan.to_dict()
    for item in payload["candidates"]:
        item["action"] = decisions.get(item["id"], item.get("default_action", "review"))
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def normalize_action(action: str) -> str:
    action = action.strip().lower()
    aliases = {
        "yes": "convert",
        "true": "convert",
        "formula": "convert",
        "ai_convert": "convert",
        "no": "keep",
        "false": "keep",
        "ignore": "keep",
        "skip": "keep",
    }
    action = aliases.get(action, action)
    if action not in {"convert", "keep", "review", "placeholder"}:
        return "review"
    return action


def score_candidate(text: str, paragraph_text: str) -> tuple[float, str]:
    stripped = text.strip()
    score = 0.35
    reasons: list[str] = ["formula-like characters"]
    if _looks_like_feature_code_text(stripped):
        score -= 0.35
        reasons.append("feature-code-like")
    if _looks_like_prose_candidate(stripped):
        score -= 0.40
        reasons.append("prose-like")
    if re.search(r"[_^]\{?[^{}\s]+\}?", stripped):
        score += 0.22
        reasons.append("script")
    if re.search(r"[=<>≤≥≈]", stripped):
        score += 0.18
        reasons.append("relation")
    if "/" in stripped and re.search(r"[A-Za-zΑ-Ωα-ωϑℓ0-9]\s*/\s*[A-Za-zΑ-Ωα-ωϑℓ0-9]", stripped):
        score += 0.12
        reasons.append("slash fraction")
    if re.search(r"[∑∏∂∞α-ωΑ-Ωϑϖ℘ℓℛ𝓡]", stripped):
        score += 0.16
        reasons.append("math symbol")
    if stripped == paragraph_text.strip() or len(stripped) / max(len(paragraph_text.strip()), 1) > 0.65:
        score += 0.12
        reasons.append("display-like")
    if re.search(r"https?://|doi\.org|arxiv:|www\.|@", stripped, re.I):
        score -= 0.55
        reasons.append("link-like")
    if is_bibliography_entry(paragraph_text.strip()):
        score -= 0.55
        reasons.append("reference-like")
    if re.fullmatch(r"[\d\s.,;:()\\/\-–—]+", stripped):
        score -= 0.35
        reasons.append("numeric-only")
    return max(0.0, min(1.0, round(score, 3))), ", ".join(reasons)


def default_action_for(text: str, paragraph_text: str, confidence: float, mode: str) -> str:
    if _looks_like_feature_code_text(text) or _looks_like_prose_candidate(text):
        return "review"
    if mode == "aggressive":
        return "convert" if confidence >= 0.35 else "review"
    if mode == "balanced":
        return "convert" if confidence >= 0.55 else "review"
    if mode == "conservative":
        return "convert" if confidence >= 0.72 else "review"
    if mode == "display-only":
        stripped = paragraph_text.strip()
        display_like = text.strip() == stripped or len(text.strip()) / max(len(stripped), 1) > 0.75
        return "convert" if display_like and confidence >= 0.55 else "review"
    raise ValueError(f"Unknown mode: {mode}")


def _looks_like_feature_code_text(text: str) -> bool:
    stripped = text.strip()
    if not stripped:
        return False
    codes = re.findall(r"\b[A-Z]{2,}[A-Z0-9]*_\d+\b", stripped)
    if not codes:
        return False
    if re.search(r"[=<>≤≥≈∑∏∂∞α-ωΑ-Ωϑϖ℘ℓℛ𝓡]", stripped):
        return False
    letters = re.findall(r"[A-Za-z_0-9]+", stripped)
    if not letters:
        return False
    code_chars = sum(len(code) for code in codes)
    alpha_chars = sum(len(token) for token in letters)
    return code_chars / max(alpha_chars, 1) > 0.45


def _looks_like_prose_candidate(text: str) -> bool:
    math_marks = len(re.findall(r"[_^=<>≤≥≈∑∏∂∞{}α-ωΑ-Ωϑϖ℘ℓℛ𝓡]", text))
    has_relation_or_symbol = bool(re.search(r"[=<>≤≥≈∑∏∂∞]", text))
    if math_marks >= 4 and has_relation_or_symbol:
        return False
    words = re.findall(r"[A-Za-z]{2,}", text)
    lower_words = [word for word in words if any(ch.islower() for ch in word)]
    if len(lower_words) < 2:
        return False
    if not re.search(r"\s", text) and math_marks:
        return False
    if re.search(r"\b(with|without|after|before|include|includes|including|feature|features|metric|metrics|split|train|testing|validation|overlaps|dropped|where|common|accuracy|precision|recall|curve)\b", text, re.I):
        return True
    if len(words) >= 5 and not re.search(r"[=<>≤≥≈∑∏∂∞α-ωΑ-Ωϑϖ℘ℓℛ𝓡]", text):
        return True
    return False


def _apply_paragraph(
    paragraph: etree._Element,
    counter: int,
    decisions: dict[str, str],
    stats: ApplyStats,
    *,
    mode: str,
    placeholder: bool,
    formula_text_overrides: dict[str, str],
) -> int:
    paragraph_text = _paragraph_text(paragraph)
    for t in list(paragraph.xpath(".//w:t", namespaces={"w": W_NS})):
        run = t.getparent()
        if run is None or run.tag != w_tag("r") or not _run_is_simple_text(run, t):
            continue
        text = t.text or ""
        parts = split_formula_spans(text)
        if not any(is_formula for _, is_formula, _, _ in parts):
            continue
        parent = run.getparent()
        if parent is None:
            continue
        rpr = run.find(w_tag("rPr"))
        new_nodes: list[etree._Element] = []
        changed = False
        for value, is_formula, _, _ in parts:
            if not is_formula:
                new_nodes.append(_make_text_run(value, rpr))
                continue
            counter += 1
            candidate_id = f"F{counter:05d}"
            stats.candidates_seen += 1
            confidence, _ = score_candidate(value, paragraph_text)
            action = normalize_action(decisions.get(candidate_id, default_action_for(value, paragraph_text, confidence, mode)))
            if placeholder:
                new_nodes.append(_make_text_run(f"⟦FORMULA_{candidate_id}⟧", rpr))
                stats.placeholders_inserted += 1
                changed = True
                continue
            if action == "convert":
                formula_text = formula_text_overrides.get(candidate_id, value)
                try:
                    new_nodes.append(node_to_omath(parse_formula(formula_text)))
                    stats.formulas_converted += 1
                    if len(stats.samples) < 20:
                        stats.samples.append(formula_text)
                    changed = True
                except Exception as exc:
                    if formula_text != value:
                        try:
                            new_nodes.append(node_to_omath(parse_formula(value)))
                            stats.formulas_converted += 1
                            if len(stats.samples) < 20:
                                stats.samples.append(value)
                            changed = True
                            stats.failed.append({"id": candidate_id, "text": formula_text, "fallback": value, "error": str(exc)})
                            continue
                        except Exception:
                            pass
                    new_nodes.append(_make_text_run(value, rpr))
                    stats.failed.append({"id": candidate_id, "text": formula_text, "error": str(exc)})
                    stats.formulas_kept += 1
            else:
                new_nodes.append(_make_text_run(value, rpr))
                stats.formulas_kept += 1
        if changed:
            insert_at = parent.index(run)
            parent.remove(run)
            for offset, node in enumerate(new_nodes):
                parent.insert(insert_at + offset, node)
    return counter


def _dynamic_parts(zin: ZipFile) -> list[str]:
    return [
        name
        for name in zin.namelist()
        if name in PROCESSABLE_PARTS
        or name.startswith("word/header")
        or name.startswith("word/footer")
    ]


def _paragraph_text(paragraph: etree._Element) -> str:
    return "".join(paragraph.xpath(".//w:t/text()", namespaces={"w": W_NS}))


def _single_text_child_for_scan(run: etree._Element) -> etree._Element | None:
    text_children = [child for child in run if child.tag == w_tag("t")]
    non_text_children = [child for child in run if child.tag not in {w_tag("rPr"), w_tag("t")}]
    if len(text_children) != 1 or non_text_children:
        return None
    return text_children[0]


def _context(paragraph_text: str, value: str, width: int = 90) -> str:
    index = paragraph_text.find(value)
    if index < 0:
        return paragraph_text[: width * 2]
    start = max(0, index - width)
    end = min(len(paragraph_text), index + len(value) + width)
    return paragraph_text[start:end]


def _sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def _extract_json(content: str) -> str:
    content = content.strip()
    if content.startswith("```"):
        content = re.sub(r"^```(?:json)?\s*", "", content)
        content = re.sub(r"\s*```$", "", content)
    return content


def _validate_mode(mode: str) -> None:
    if mode not in MODES:
        raise ValueError(f"mode must be one of {', '.join(sorted(MODES))}")
