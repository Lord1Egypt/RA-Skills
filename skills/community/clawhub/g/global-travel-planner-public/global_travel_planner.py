#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Global Travel & Work Planner v1.9.0
出国游全球通版 v1.9.0

One-stop bilingual travel guide for 48 countries.
一键查询48个国家的中英文双语旅行方案。

Output: Word (.docx) format — fully bilingual (CN/EN).
输出：Word文档（.docx）格式，中英双语对照。

⚠️ ALL DATA FOR REFERENCE ONLY. Always verify via official sources before travel.
所有信息仅供参考，出行前请通过官方渠道核实。
面向世界各国公民，如需领事协助请联系本国使领馆。
"""
import os
import sys
import json
import datetime

# ── Language mapping: ISO code → native name ──
# Determines which language label appears alongside English in the output.
# "source_lang" specifies the user's language (place of use).
# Output format: "EN: xxx / {native_name}: xxx"
LANG_MAP = {
    "zh": "中文",
    "en": "English",
    "ko": "한국어",
    "ja": "日本語",
    "fr": "Français",
    "de": "Deutsch",
    "es": "Español",
    "pt": "Português",
    "ru": "Русский",
    "ar": "العربية",
    "th": "ไทย",
    "vi": "Tiếng Việt",
    "id": "Bahasa Indonesia",
    "ms": "Bahasa Melayu",
    "hi": "हिन्दी",
    "nl": "Nederlands",
    "it": "Italiano",
    "tr": "Türkçe",
    "pl": "Polski",
    "sv": "Svenska",
    "tl": "Filipino",
    "mn": "Монгол",
    "km": "ភាសាខ្មែរ",
    "lo": "ລາວ",
    "my": "မြန်မာ",
    "ar": "العربية",
    "ru": "Русский",
    "th": "ไทย",
    "hi": "हिन्दी",
    "el": "Ελληνικά",
    "he": "עברית",
    "ur": "اردو",
    "bn": "বাংলা",
    "ne": "नेपाली",
    "am": "አማርኛ",      # Amharic (Ethiopia)
    "ti": "ትግርኛ",      # Tigrinya (Eritrea, Ethiopia)
    "om": "Afaan Oromo",  # Oromo (Ethiopia)
    "ha": "Hausa",       # Hausa (Nigeria, Niger)
    "sw": "Kiswahili",    # Swahili (East Africa)
    "yo": "Yorùbá",      # Yoruba (Nigeria)
    "ig": "Igbo",        # Igbo (Nigeria)
    "zu": "isiZulu",     # Zulu (South Africa)
    "af": "Afrikaans",   # Afrikaans (South Africa)
    "st": "Sesotho",     # Sotho (Lesotho, South Africa)
    "sn": "chiShona",    # Shona (Zimbabwe)
    "ny": "Chichewa",    # Chewa (Malawi, Zambia)
    "rw": "Kinyarwanda", # Kinyarwanda (Rwanda)
    "mg": "Malagasy",    # Malagasy (Madagascar)
    "ff": "Fulfulde",   # Fulani (West Africa)
    "bm": "Bamanankan",  # Bambara (Mali)
    "ka": "ქართული",      # Georgian
    "hy": "Հայերեն",      # Armenian
    "si": "සිංහල",        # Sinhala (Sri Lanka)
    "bo": "བོད་སྐད།",    # Tibetan
    "my": "မြန်မာ",       # Burmese (Myanmar)
    "km": "ភាសាខ្មែរ",   # Khmer (Cambodia)
    "lo": "ລາວ",          # Lao
    "jv": "Basa Jawa",    # Javanese
    "dv": "ދިވެހި",       # Dhivehi (Maldives)
    "chr": "ᏣᎳᎩ",        # Cherokee
    "iu": "ᐃᓄᒃᑎᑐᑦ",    # Inuktitut (Canadian Aboriginal)
    "mni": "Meitei",      # Meitei/Manipuri (India)
    "ii": "ꆈꌠꉙ",        # Yi (Sichuan, China)
}

_CURRENT_SOURCE_LANG = "zh"
_INPUT_LANG = "zh"
_OUTPUT_DEST_LANG = "zh"


def _native_name(code):
    """Return the native name for a language code (fallback to English name)."""
    return LANG_MAP.get(code, code.upper())



# ── Auto-detect: user language + destination language ──
# 大叔亲授：审题→判断输入语言，质检→匹配目的地语言
# 输出：输入语言 + 目的地语言（如中国人在日本=中文+日本語）
# 使用规则：
#   中国人查日本 → 中日双语（中文+日本語）
#   日本人查中国 → 日中双语（日本語+中文）
#   韩国人去日本 → 한일双语（한국어+日本語）
#   法国人去韩国 → français+한국어
AUTO_COUNTRY_LANG = {
    # East Asia
    "CN": "zh", "TW": "zh", "HK": "zh", "MO": "zh",
    "JP": "ja", "KR": "ko", "MN": "mn",
    # Southeast Asia
    "TH": "th", "VN": "vi", "ID": "id", "MY": "ms",
    "PH": "tl", "SG": "en", "KH": "km", "LA": "lo", "MM": "my",
    # South Asia
    "IN": "hi", "PK": "ur", "BD": "bn", "NP": "ne", "LK": "si",
    # Europe
    "FR": "fr", "DE": "de", "IT": "it", "ES": "es", "PT": "pt",
    "GB": "en", "IE": "en", "NL": "nl", "BE": "nl", "CH": "de",
    "AT": "de", "SE": "sv", "NO": "no", "DK": "da", "FI": "fi",
    "PL": "pl", "CZ": "cs", "SK": "sk", "HU": "hu", "RO": "ro",
    "GR": "el", "TR": "tr", "RU": "ru", "UA": "uk",
    # Americas
    "US": "en", "CA": "en", "MX": "es", "BR": "pt", "AR": "es",
    # Middle East / Africa
    "AE": "ar", "SA": "ar", "EG": "ar", "IL": "he", "ZA": "en",
    # Oceania
    "AU": "en", "NZ": "en",
}


def _detect_input_lang(input_text=""):
    """Detect input language from text content (大叔亲授·审题机制).
    
    只看输入的文字识别语言，不猜用户的国籍。
    Chinese input → zh, Japanese kana → ja, Hangul → ko,
    Arabic → ar, Cyrillic → ru, Thai → th, Devanagari → hi,
    Greek → el, Hebrew → he, etc.
    
    Args:
        input_text: the user's raw input text
    
    Returns:
        ISO language code
    """
    if not input_text:
        return "zh"
    
    # Unicode block ranges for world scripts
    # --- CJK (Chinese characters) ---
    cjk = sum(1 for c in input_text if 0x4E00 <= ord(c) <= 0x9FFF or 0x3400 <= ord(c) <= 0x4DBF)
    # --- Japanese kana (Hiragana + Katakana) ---
    kana = sum(1 for c in input_text if 0x3040 <= ord(c) <= 0x309F or 0x30A0 <= ord(c) <= 0x30FF or 0xFF66 <= ord(c) <= 0xFF9F)
    # --- Korean Hangul ---
    hangul = sum(1 for c in input_text if 0xAC00 <= ord(c) <= 0xD7AF or 0x1100 <= ord(c) <= 0x11FF)
    # --- Arabic ---
    arabic = sum(1 for c in input_text if 0x0600 <= ord(c) <= 0x06FF or 0x0750 <= ord(c) <= 0x077F or 0x08A0 <= ord(c) <= 0x08FF or 0xFB50 <= ord(c) <= 0xFDFF or 0xFE70 <= ord(c) <= 0xFEFF)
    # --- Cyrillic (Russian, Ukrainian, Bulgarian, etc.) ---
    cyrillic = sum(1 for c in input_text if 0x0400 <= ord(c) <= 0x04FF or 0x0500 <= ord(c) <= 0x052F)
    # --- Thai ---
    thai = sum(1 for c in input_text if 0x0E00 <= ord(c) <= 0x0E7F)
    # --- Devanagari (Hindi, Sanskrit, Marathi, Nepali) ---
    devanagari = sum(1 for c in input_text if 0x0900 <= ord(c) <= 0x097F)
    # --- Greek ---
    greek = sum(1 for c in input_text if 0x0370 <= ord(c) <= 0x03FF)
    # --- Hebrew ---
    hebrew = sum(1 for c in input_text if 0x0590 <= ord(c) <= 0x05FF)
    # --- Ethiopic / Ge'ez (Amharic, Tigrinya, Oromo) ---
    ethiopic = sum(1 for c in input_text if 0x1200 <= ord(c) <= 0x137F or 0x1380 <= ord(c) <= 0x139F or 0x2D80 <= ord(c) <= 0x2DDF or 0xAB00 <= ord(c) <= 0xAB2F)
    # --- N'Ko (Manding languages: Bambara, Malinké, etc. - Guinea, Mali) ---
    nko = sum(1 for c in input_text if 0x07C0 <= ord(c) <= 0x07FF)
    # --- Vai (Liberia, Sierra Leone) ---
    vai = sum(1 for c in input_text if 0xA500 <= ord(c) <= 0xA63F)
    # --- Tifinagh (Berber / Tuareg - North Africa, Sahara) ---
    tifinagh = sum(1 for c in input_text if 0x2D30 <= ord(c) <= 0x2D7F)
    # --- Bamum (Cameroon) ---
    bamum = sum(1 for c in input_text if 0xA6A0 <= ord(c) <= 0xA6FF)
    # --- Adlam (Fulani / Fulfulde - West Africa: Guinea, Nigeria, Cameroon) ---
    adlam = sum(1 for c in input_text if 0x1E900 <= ord(c) <= 0x1E95F)
    # --- Osmanya (Somali) ---
    osmanya = sum(1 for c in input_text if 0x10480 <= ord(c) <= 0x104AF)
    # --- Georgian (Mkhedruli / Asomtavruli) ---
    georgian = sum(1 for c in input_text if 0x10A0 <= ord(c) <= 0x10FF or 0x2D00 <= ord(c) <= 0x2D25)
    # --- Armenian ---
    armenian = sum(1 for c in input_text if 0x0530 <= ord(c) <= 0x058F or 0xFB00 <= ord(c) <= 0xFB17)
    # --- Sinhala (Sri Lanka) ---
    sinhala = sum(1 for c in input_text if 0x0D80 <= ord(c) <= 0x0DFF or 0x111E0 <= ord(c) <= 0x111FF)
    # --- Tibetan ---
    tibetan = sum(1 for c in input_text if 0x0F00 <= ord(c) <= 0x0FFF)
    # --- Myanmar / Burmese ---
    myanmar = sum(1 for c in input_text if 0x1000 <= ord(c) <= 0x109F or 0xAA60 <= ord(c) <= 0xAA7F)
    # --- Khmer (Cambodian) ---
    khmer = sum(1 for c in input_text if 0x1780 <= ord(c) <= 0x17FF or 0x19E0 <= ord(c) <= 0x19FF)
    # --- Lao ---
    lao = sum(1 for c in input_text if 0x0E80 <= ord(c) <= 0x0EFF)
    # --- Mongolian (Mongolian script, not Cyrillic) ---
    mongolian = sum(1 for c in input_text if 0x1800 <= ord(c) <= 0x18AF)
    # --- Javanese (Indonesia) ---
    javanese = sum(1 for c in input_text if 0xA980 <= ord(c) <= 0xA9DF)
    # --- Thaana (Dhivehi / Maldives) ---
    thaana = sum(1 for c in input_text if 0x0780 <= ord(c) <= 0x07BF)
    # --- Cherokee (USA) ---
    cherokee = sum(1 for c in input_text if 0x13A0 <= ord(c) <= 0x13FF or 0xAB70 <= ord(c) <= 0xABBF)
    # --- Unified Canadian Aboriginal Syllabics (Inuktitut, Cree, Ojibwe, etc.) ---
    canadian_aboriginal = sum(1 for c in input_text if 0x1400 <= ord(c) <= 0x167F or 0x18B0 <= ord(c) <= 0x18FF)
    # --- Yi (Sichuan, China) ---
    yi = sum(1 for c in input_text if 0xA000 <= ord(c) <= 0xA4CF)
    # --- Meitei Mayek (Manipuri, India) ---
    meitei = sum(1 for c in input_text if 0xAAE0 <= ord(c) <= 0xAAFF or 0xABC0 <= ord(c) <= 0xABFF)
    # --- Latin/ASCII (English, French, Spanish, German, Swahili, Hausa, etc.) ---
    ascii_letters = sum(1 for c in input_text if c.isascii() and c.isalpha())
    
    # Build list of detectable scripts and their counts
    scripts = [
        (kana, "ja"),        # Japanese kana
        (hangul, "ko"),      # Korean Hangul
        (arabic, "ar"),      # Arabic
        (cyrillic, "ru"),    # Cyrillic → Russian (fallback)
        (thai, "th"),        # Thai
        (devanagari, "hi"),  # Devanagari → Hindi (fallback)
        (greek, "el"),       # Greek
        (hebrew, "he"),      # Hebrew
        (cjk, "zh"),         # Chinese characters
        # --- African scripts ---
        (ethiopic, "am"),    # Ethiopic → Amharic (Ethiopia, Eritrea)
        (nko, "bm"),         # N'Ko → Bambara / Manding (Guinea, Mali)
        (adlam, "ff"),       # Adlam → Fulani / Fulfulde (West Africa)
        (vai, "vai"),        # Vai (Liberia)
        (tifinagh, "ber"),   # Tifinagh → Berber / Tuareg (North Africa)
        (bamum, "bam"),      # Bamum (Cameroon)
        (osmanya, "so"),     # Osmanya → Somali
        # --- Eurasian scripts ---
        (georgian, "ka"),    # Georgian Mkhedruli (Georgia)
        (armenian, "hy"),    # Armenian
        (sinhala, "si"),     # Sinhala (Sri Lanka)
        (tibetan, "bo"),     # Tibetan
        (myanmar, "my"),     # Burmese (Myanmar)
        (khmer, "km"),       # Khmer (Cambodia)
        (lao, "lo"),         # Lao
        (mongolian, "mn"),   # Mongolian script
        (javanese, "jv"),    # Javanese (Java, Indonesia)
        (thaana, "dv"),      # Thaana → Dhivehi (Maldives)
        (cherokee, "chr"),   # Cherokee (USA)
        (canadian_aboriginal, "iu"),  # Inuktitut / Cree (Canada)
        (yi, "ii"),          # Yi (Sichuan, China)
        (meitei, "mni"),     # Meitei Mayek (Manipur, India)
    ]
    
    # Total identifiable characters
    total = sum(s[0] for s in scripts) + ascii_letters
    
    if total == 0:
        return "zh"
    
    # Find dominant script (if any script > 10% of total)
    for count, lang in scripts:
        if count > total * 0.1:
            return lang
    
    # Only ASCII letters → English (or Latin-alphabet language)
    if ascii_letters > total * 0.4:
        return "en"
    
    return "zh"  # Default


def _match_dest_lang(country_code):
    """Look up destination country's native language (大叔亲授·质检机制).
    
    Args:
        country_code: ISO code of destination country
    
    Returns:
        ISO language code
    """
    return AUTO_COUNTRY_LANG.get(country_code.upper(), "en")


BASE_DIR = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, BASE_DIR)

DATA_PATH = os.path.join(BASE_DIR, "global_data", "countries_base.json")
OUTPUT_DIR = os.path.join(BASE_DIR, "output")
_country_cache = None


def _load_data():
    """Load country database from JSON file."""
    global _country_cache
    if _country_cache is not None:
        return _country_cache
    try:
        with open(DATA_PATH, "r", encoding="utf-8") as f:
            _country_cache = json.load(f)
        return _country_cache
    except (FileNotFoundError, json.JSONDecodeError) as e:
        print(f"Error: Country database load failed: {e}")
        return {}


def search_country(query):
    """Search country by code, CN name, or EN name."""
    data = _load_data()
    q = query.strip().upper()
    if q in data:
        return data[q]
    for code, info in data.items():
        if q.lower() in info["name_cn"].lower():
            return info
        if q.lower() in info["name_en"].lower():
            return info
        if q == info.get("alpha3", "").upper():
            return info
    return None


# ---------------------------------------------------------------------------
# Helper: fully bilingual table rows
# ---------------------------------------------------------------------------

def _visa_rows(c):
    """Return list of (item_en_cn, value_en, value_cn) for visa."""
    v = c.get("visa_cn", {})
    return [
        ("Type / 类型", v.get("type_en", "N/A"), v.get("type_cn", "N/A")),
        ("Processing / 办理周期", f"{v.get('processing_days', 'N/A')} business days", f"{v.get('processing_days', 'N/A')}个工作日"),
        ("Fee / 费用", v.get("fee_en", "N/A"), v.get("fee_cn", "N/A")),
        ("Validity / 有效期", v.get("validity_en", "N/A"), v.get("validity_cn", "N/A")),
    ]


def _laws_rows(c):
    """Return list of (item_en_cn, value_en, value_cn) for laws."""
    l = c.get("laws_cn", {})
    return [
        ("Drinking Age / 饮酒年龄", f"{l.get('drinking_age', 'N/A')} years old", f"{l.get('drinking_age', 'N/A')}岁"),
        ("Smoking / 吸烟", l.get("smoking_regulation", "N/A"), l.get("smoking_regulation", "N/A")),
        ("Drug Penalty / 毒品处罚", l.get("drug_penalty", "N/A"), l.get("drug_penalty", "N/A")),
        ("Photo Restrictions / 拍照限制", l.get("photo_restriction", "None"), l.get("photo_restriction", "None")),
        ("Tipping Culture / 小费文化", l.get("tipping_culture_en", "Not required"), l.get("tipping_culture_cn", "Not required")),
        ("Public Conduct / 公共行为", l.get("public_behavior", "Follow rules"), l.get("public_behavior", "Follow rules")),
        ("Driving License / 驾照", l.get("driving_license_en", "CN license valid"), l.get("driving_license_cn", "CN license valid")),
    ]


def _culture_rows(c):
    """Culture table rows."""
    cv = c.get("culture_cn", {})
    return [
        ("Greeting / 见面礼仪", cv.get("greeting_en", "Handshake"), cv.get("greeting", "Handshake")),
        ("Dining / 用餐礼仪", cv.get("dining_etiquette_en", "Knife & fork"), cv.get("dining_etiquette", "Knife & fork")),
        ("Taboos / 禁忌", cv.get("taboo_en", "None"), cv.get("taboo_cn", "None")),
        ("Dress Code / 着装", cv.get("dress_code_en", "Casual"), cv.get("dress_code", "Casual")),
        ("Gifts / 送礼", cv.get("gift_en", "Small tokens"), cv.get("gift_cn", "Small tokens")),
    ]


def _safety_rows(c):
    """Safety table rows."""
    s = c.get("safety_cn", {})
    return [
        ("Safety Level / 安全等级", s.get("level_en", s.get("level", "N/A")), s.get("level", "N/A")),
        ("Common Scams / 常见骗局", s.get("common_scams_en", s.get("common_scams", "None")), s.get("common_scams", "None")),
        ("Night Safety / 夜间安全", s.get("night_safety_en", s.get("night_safety", "Caution")), s.get("night_safety", "Caution")),
        ("Female Safety / 女性安全", s.get("female_safety_en", s.get("female_safety", "Moderate")), s.get("female_safety", "Moderate")),
        ("Emergency Phrases / 应急用语", s.get("emergency_words_en", s.get("emergency_words", "Help!/Police!")), s.get("emergency_words", "Help!/Police!")),
        ("Must-have Apps / 必备APP", s.get("must_have_apps_en", s.get("must_have_apps", "")), s.get("must_have_apps", "")),
    ]


def _travel_rows(c):
    """Travel plan table rows."""
    plugs = ", ".join(c.get("plug_type", []))
    voltage = c.get("voltage", "")
    return [
        ("Capital / 首都", c.get("capital_en", ""), c.get("capital_cn", "")),
        ("Budget Level / 预算等级", "Medium", "中等"),
        ("Accommodation / 住宿", "3-star hotel", "3星级酒店"),
        ("Avg Meal Cost / 平均餐费", c.get("avg_meal_en", "N/A"), c.get("avg_meal_cn", "N/A")),
        ("Avg Hotel / 住宿费", c.get("avg_hotel_en", "N/A"), c.get("avg_hotel_cn", "N/A")),
        ("Timezone / 时区", f"UTC{c.get('timezone_offset', '')}", f"UTC{c.get('timezone_offset', '')}"),
        ("Currency / 货币", f"{c.get('currency_en', '')} ({c.get('currency_symbol', '')})", f"{c.get('currency_cn', '')} ({c.get('currency_symbol', '')})"),
        ("Language / 语言", c.get("language_en", "N/A"), c.get("language_cn", "N/A")),
        ("Plug / Voltage / 插头电压", f"{plugs} ({voltage})", f"{plugs} ({voltage})"),
    ]


def _accommodation_rows(accommodation_data, tier_label_en, tier_label_cn):
    """Accommodation table rows for one tier."""
    rows = []
    for a in accommodation_data:
        name = f"{a[0]} / {a[1]}"
        price = f"{a[2]} / {a[3]}"
        features = a[4] if len(a) > 4 else ""
        features_cn = a[5] if len(a) > 5 else ""
        recommend = ""
        if len(a) > 7 and a[6]:
            if a[6] == "★":
                recommend = "⭐ Recommended / 推荐"
            elif a[6] == "★★":
                recommend = "⭐⭐ Highly Recommended / 强烈推荐"
            elif a[6] == "★★★":
                recommend = "⭐⭐⭐ Top Pick / 首选"
        if features or features_cn:
            detail = f"💰 {price} | 📍 {features} / {features_cn}"
        else:
            detail = f"💰 {price}"
        if recommend:
            detail = f"{recommend} | {detail}"
        rows.append((name, detail, detail))
    return rows


def _work_rows(c):
    """Work guide table rows."""
    return [
        ("Work Visa / 工作签证", "Employer-sponsored work visa required", "需雇主担保的工作签证"),
        ("Language / 语言", c.get("language_en", "N/A"), c.get("language_cn", "N/A")),
        ("Currency / 货币", f"{c.get('currency_en', '')} ({c.get('currency_symbol', '')})", f"{c.get('currency_cn', '')} ({c.get('currency_symbol', '')})"),
        ("Timezone / 时区", f"UTC{c.get('timezone_offset', '')}", f"UTC{c.get('timezone_offset', '')}"),
    ]


def _study_rows(c):
    """Study guide table rows."""
    return [
        ("Visa Type / 签证类型", "Student visa (requires admission letter)", "学生签证（需录取通知书）"),
        ("Language / 语言", c.get("language_en", "N/A"), c.get("language_cn", "N/A")),
        ("Currency / 货币", f"{c.get('currency_en', '')} ({c.get('currency_symbol', '')})", f"{c.get('currency_cn', '')} ({c.get('currency_symbol', '')})"),
        ("Living Cost / 生活费", f"~{c.get('cost_level_en', 'Medium')}/month", f"~{c.get('cost_level_cn', '中等')}/月"),
    ]


# ---------------------------------------------------------------------------
# Docx generation
# ---------------------------------------------------------------------------

def _add_bilingual_table(doc, heading, rows):
    """Add a dual-language table: user's native language + destination's language.
    
    Uses _INPUT_LANG and _OUTPUT_DEST_LANG (module-level variables set by generate_docx).
    
    大叔亲授逻辑：
      - 中国人查日本 → 中文+日本語（输入语言+目的地语言）
      - 日本人查中国 → 日本語+中文
      - 韩国人去日本 → 한국어+日本語
      - 同语言（如中国人查中国）→ 单语
    
    Args:
        doc: Document object
        heading: section heading string
        rows: list of (item_label, value_en, value_cn)
    """
    input_lang = _INPUT_LANG
    dest_lang = _OUTPUT_DEST_LANG
    input_native = _native_name(input_lang)
    dest_native = _native_name(dest_lang)
    
    if heading:
        doc.add_heading(heading, level=1)

    table = doc.add_table(rows=len(rows) + 1, cols=2, style="Light Shading Accent 1")

    # Header
    hdr = table.rows[0].cells
    p = hdr[0].paragraphs[0]
    run = p.add_run("Item / " + input_native if input_lang != "zh" else "Item / 项目")
    run.bold = True
    p = hdr[1].paragraphs[0]
    if input_lang == dest_lang:
        run = p.add_run(input_native)
    else:
        run = p.add_run(f"{input_native} / {dest_native}")
    run.bold = True

    for i, (item, v_en, v_cn) in enumerate(rows):
        cells = table.rows[i + 1].cells
        cells[0].text = item
        
        # Input language content
        if input_lang == "zh":
            input_content = v_cn
        else:
            input_content = v_en
        
        # Destination's language content
        if dest_lang == "zh":
            dest_content = v_cn
        else:
            dest_content = v_en
        
        if input_lang == dest_lang:
            cells[1].text = f"{input_native}: {input_content}"
        else:
            cells[1].text = f"{input_native}: {input_content} / {dest_native}: {dest_content}"


def _add_warning(doc, text_en, text_cn):
    """Add a colored warning paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(f"⚠️ {text_en}")
    run.font.color.rgb = __import__("docx.shared", fromlist=["RGBColor"]).RGBColor(0xCC, 0x66, 0x00)
    p2 = doc.add_paragraph()
    run2 = p2.add_run(f"⚠️ {text_cn}")
    run2.font.color.rgb = __import__("docx.shared", fromlist=["RGBColor"]).RGBColor(0xCC, 0x66, 0x00)


def _add_text_section(doc, heading, lines):
    """Add a section with bullet-point text lines."""
    doc.add_heading(heading, level=1)
    for line in lines:
        if line.startswith("#"):
            p = doc.add_paragraph()
            run = p.add_run(line.lstrip("# "))
            run.bold = True
        else:
            doc.add_paragraph(line, style="List Bullet")


def generate_docx(country_code, modules=None, nationality=None, days=7, supplement=None, user_input="", source_lang="zh"):
    """Generate a bilingual Word document for a country.

    Args:
        country_code: ISO code or name
        modules: list of module names, or None for all
        nationality: ISO code (default from env or 'CN')
        days: number of travel days
        source_lang: (deprecated, kept for compat) use user_input instead
        user_input: raw user query text (used when source_lang="auto" for language detection)
        supplement: dict with optional keys:
            - attractions: list of (en_name, cn_name, ticket_en, ticket_cn, desc_en, desc_cn)
            - foods: list of (en_name, cn_name, desc_en, desc_cn)
            - transport: list of (method_en, method_cn, desc_en, desc_cn)
            - itinerary: list of (day_label, activity_en, activity_cn)
            - accommodation: dict with optional keys:
                - budget: list of (name_en, name_cn, price_en, price_cn, features_en, features_cn, recommend_stars)
                - midrange: same format
                - luxury: same format
              recommend_stars: "★", "★★", or "★★★" (blank for none)
            - visa_details: dict with optional keys:
                - how_to_apply: list of (step_en, step_cn)
                - where_to_apply: list of (place_en, place_cn, address_en, address_cn, link)
                - process: list of (step_en, step_cn, detail_en, detail_cn)
                - documents: list of (doc_en, doc_cn, notes_en, notes_cn)
                - tips: list of (tip_en, tip_cn)
            - entry_exit: dict with optional keys:
                - arrival: list of (step_en, step_cn) — arrival process
                - customs: list of (item_en, item_cn, rule_en, rule_cn) — customs regulations
                - departure: list of (step_en, step_cn) — departure process
                - tips: list of (tip_en, tip_cn)
            - phrases: list of (situation_en, situation_cn, phrase_en, pronunciation, phrase_cn)
            - weather: dict with optional keys:
                - overview: (overview_en, overview_cn)
                - best_time: (best_time_en, best_time_cn)
                - seasons: list of (season_en, season_cn, temp_en, temp_cn, clothing_en, clothing_cn)
            - budget_detail: list of (category_en, category_cn, amount_en, amount_cn, pct, notes_en_cn)
            - insurance: dict with optional keys:
                - recommendation: (rec_en, rec_cn)
                - types: list of (type_en, type_cn, desc_en, desc_cn, suggested_en, suggested_cn)
                - tips: list of (tip_en, tip_cn)
            - shopping: dict with optional keys:
                - specialties: list of (item_en, item_cn, price_en_cn, where_en, where_cn)
                - areas: list of (area_en, area_cn, desc_en, desc_cn)
                - tax_free: (tax_free_en, tax_free_cn)
                - tips: list of (tip_en, tip_cn)
            - communication: dict with optional keys:
                - sim_card: (info_en, info_cn)
                - wifi: (info_en, info_cn)
                - roaming: (info_en, info_cn)
                - useful_apps: list of (app_name, desc_en, desc_cn)
            - predeparture: dict with optional keys:
                - documents: list of (item_en, item_cn, notes_en, notes_cn)
                - luggage: list of (item_en, item_cn, notes_en, notes_cn)
                - money: list of (tip_en, tip_cn)
                - health: list of (tip_en, tip_cn)
                - checklist: list of (item_en, item_cn)

    Returns:
        Path to generated .docx file, or None on error.
    """
    global _CURRENT_SOURCE_LANG, _INPUT_LANG, _OUTPUT_DEST_LANG
    
    # ── 大叔审题机制：从输入判断输入语言 ──
    _INPUT_LANG = _detect_input_lang(user_input)
    
    # ── 大叔质检机制：匹配目的地国家的语言 ──
    _OUTPUT_DEST_LANG = _match_dest_lang(country_code)
    
    # For backward compat: set _CURRENT_SOURCE_LANG to user's language
    _CURRENT_SOURCE_LANG = _INPUT_LANG
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor

    if supplement is None:
        supplement = {}

    c = search_country(country_code)
    if not c:
        print(f"Country not found: {country_code}")
        return None

    if modules is None:
        modules = ["visa", "laws", "culture", "safety", "consular", "travel", "work", "study", "attractions", "food", "transport", "itinerary", "accommodation", "predeparture", "entry_exit", "phrases", "weather", "budget_detail", "insurance", "shopping", "communication"]
    if nationality is None:
        nationality = os.environ.get("CONSULAR_NATIONALITY", "CN")

    doc = Document()

    # --- Title ---
    title = doc.add_heading("", 0)
    run = title.add_run(f"🌍 {c['name_cn']} / {c['name_en']}")
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)

    sub = doc.add_paragraph()
    run = sub.add_run(
        "出国游全球通版 v1.9.0  |  Global Travel & Work Planner"
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    sub2 = doc.add_paragraph(
        f"{datetime.datetime.now().strftime('%Y-%m-%d')}  |  "
        "🌍 面向世界各国公民 / For users of ALL nationalities"
    )

    doc.add_page_break()

    # --- Basic Info ---
    plugs = ", ".join(c.get("plug_type", []))
    voltage = c.get("voltage", "")
    driving = "靠右 / Right" if c.get("driving_side") != "left" else "靠左 / Left"
    _add_bilingual_table(doc, "📋 Basic Info / 基本信息", [
        ("Name / 名称", c["name_en"], c["name_cn"]),
        ("Capital / 首都", c.get("capital_en", ""), c.get("capital_cn", "")),
        ("Continent / 大洲", c.get("continent_en", c.get("continent", "")), c.get("continent_cn", c.get("continent", ""))),
        ("Currency / 货币", f"{c.get('currency_en', '')} ({c.get('currency_symbol', '')})", f"{c.get('currency_cn', '')} ({c.get('currency_symbol', '')})"),
        ("Timezone / 时区", f"UTC{c.get('timezone_offset', '')}", f"UTC{c.get('timezone_offset', '')}"),
        ("Language / 语言", c.get("language_en", ""), c.get("language_cn", "")),
        ("Plug & Voltage / 插头电压", f"{plugs} ({voltage})", f"{plugs} ({voltage})"),
        ("Driving Side / 行驶方向", driving, driving),
    ])

    # --- Pre-Departure Checklist (from supplement) ---
    if "predeparture" in modules and supplement.get("predeparture"):
        pd = supplement["predeparture"]
        doc.add_heading("📝 Pre-Departure Checklist / 行前准备清单", level=1)

        if pd.get("documents"):
            doc.add_heading("📄 Documents / 证件文件", level=2)
            rows = []
            for d in pd["documents"]:
                name = f"{d[0]} / {d[1]}"
                notes = d[2] if len(d) > 2 else ""
                if notes:
                    rows.append((name, notes, notes))
                else:
                    rows.append((name, "", ""))
            if rows:
                _add_bilingual_table(doc, "", rows)

        if pd.get("luggage"):
            doc.add_heading("🧳 Luggage / 行李准备", level=2)
            rows = []
            for l in pd["luggage"]:
                name = f"{l[0]} / {l[1]}"
                notes = l[2] if len(l) > 2 else ""
                if notes:
                    rows.append((name, notes, notes))
                else:
                    rows.append((name, "", ""))
            if rows:
                _add_bilingual_table(doc, "", rows)

        if pd.get("money"):
            doc.add_heading("💰 Money / 资金准备", level=2)
            for t in pd["money"]:
                if len(t) >= 2:
                    doc.add_paragraph(f"• {t[0]} / {t[1]}", style="List Bullet")
                else:
                    doc.add_paragraph(f"• {t[0]}", style="List Bullet")

        if pd.get("health"):
            doc.add_heading("💊 Health / 健康准备", level=2)
            for t in pd["health"]:
                if len(t) >= 2:
                    doc.add_paragraph(f"• {t[0]} / {t[1]}", style="List Bullet")
                else:
                    doc.add_paragraph(f"• {t[0]}", style="List Bullet")

        if pd.get("checklist"):
            doc.add_heading("✅ Final Checklist / 出发前检查清单", level=2)
            rows = []
            for ci in pd["checklist"]:
                name = f"{c[0]} / {c[1]}"
                rows.append((name, "☐", "☐"))
            _add_bilingual_table(doc, "", rows)

        p = doc.add_paragraph()
        run = p.add_run("💡 Check all items before departure. / 出发前逐项检查确认。")
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # --- Travel Risk & Warning (from supplement) ---
    # supplement["risk"] = dict with:
    #   "overall_risk": (level_en, level_cn, rating)  — e.g. ("Medium", "中等", "3/5")
    #   "warnings": list of [(category_en, category_cn, warning_en, warning_cn), ...]
    #   "tips": list of [(title_en, title_cn, tip_en, tip_cn), ...]
    if "risk" in modules and supplement.get("risk"):
        r = supplement["risk"]
        doc.add_heading("⚠️ Travel Risk Warning / 旅游风险预警", level=1)

        # Overall risk rating
        if r.get("overall_risk"):
            level_en, level_cn, rating = r["overall_risk"]
            emoji = {"low": "🟢", "medium": "🟡", "high": "🟠", "very high": "🔴", "extreme": "⛔"}
            key = level_en.lower().strip()
            e = emoji.get(key, "⚠️")
            p = doc.add_paragraph()
            run = p.add_run(f"{e} Overall Risk Level / 综合风险等级: {level_en} ({rating}) / {level_cn} ({rating})")
            run.bold = True
            run.font.size = Pt(14)

        # Risk categories table
        if r.get("warnings"):
            rows = []
            for w in r["warnings"]:
                if len(w) >= 4:
                    cat = f"{w[0]} / {w[1]}"
                    val = f"{w[2]} / {w[3]}"
                else:
                    cat = w[0]
                    val = w[1]
                rows.append((cat, val, val))
            table = doc.add_table(rows=len(rows) + 1, cols=2, style="Light Shading Accent 1")
            hdr = table.rows[0].cells
            p = hdr[0].paragraphs[0]
            run = p.add_run("Category / 类别")
            run.bold = True
            p = hdr[1].paragraphs[0]
            run = p.add_run("Risk Details / 风险详情")
            run.bold = True
            for i, (cat, val, _) in enumerate(rows):
                cells = table.rows[i + 1].cells
                cells[0].text = cat
                cells[1].text = val

        # Safety tips
        if r.get("tips"):
            doc.add_heading("🛡️ Safety Tips / 安全建议", level=2)
            rows = []
            for tip in r["tips"]:
                if len(tip) >= 4:
                    title = f"{tip[0]} / {tip[1]}"
                    val = f"{tip[2]} / {tip[3]}"
                else:
                    title = tip[0]
                    val = tip[1]
                rows.append((title, val, val))
            table = doc.add_table(rows=len(rows) + 1, cols=2, style="Light Shading Accent 1")
            hdr = table.rows[0].cells
            p = hdr[0].paragraphs[0]
            run = p.add_run("Aspect / 方面")
            run.bold = True
            p = hdr[1].paragraphs[0]
            run = p.add_run("Advice / 建议")
            run.bold = True
            for i, (title, val, _) in enumerate(rows):
                cells = table.rows[i + 1].cells
                cells[0].text = title
                cells[1].text = val

    # --- History & Culture (from supplement) ---
    # supplement["history"] = dict with:
    #   "overview": list of [(en_para, cn_para), ...]  — overview paragraphs
    #   "timeline": list of [(year_en_cn, event_en, event_cn), ...]  — key historical events
    #   "culture_tips": list of [(title_en, title_cn, tip_en, tip_cn), ...]  — cultural insights
    if "history" in modules and supplement.get("history"):
        h = supplement["history"]
        doc.add_heading("📜 History & Culture / 历史人文", level=1)

        # Overview paragraphs
        if h.get("overview"):
            for en_text, cn_text in h["overview"]:
                p = doc.add_paragraph()
                run = p.add_run(en_text)
                run.italic = True
                doc.add_paragraph(cn_text)

        # Historical timeline
        # items: (period_label, event_bilingual) or (period_label, event_en, event_cn)
        if h.get("timeline"):
            doc.add_heading("📅 Key Historical Events / 关键历史事件", level=2)
            rows = []
            for item in h["timeline"]:
                label = item[0]
                if len(item) >= 3:
                    val = f"{item[1]} / {item[2]}"
                else:
                    val = item[1]
                rows.append((label, val, val))
            table = doc.add_table(rows=len(rows) + 1, cols=2, style="Light Shading Accent 1")
            hdr = table.rows[0].cells
            p = hdr[0].paragraphs[0]
            run = p.add_run("Period / 时期")
            run.bold = True
            p = hdr[1].paragraphs[0]
            run = p.add_run("Events / 事件")
            run.bold = True
            for i, (label, val, _) in enumerate(rows):
                cells = table.rows[i + 1].cells
                cells[0].text = label
                cells[1].text = val

        # Cultural insights
        if h.get("culture_tips"):
            doc.add_heading("🎨 Cultural Highlights / 文化特色", level=2)
            rows = []
            for item in h["culture_tips"]:
                # items: (title_bilingual, desc_bilingual) or (title_en, title_cn, desc_en, desc_cn)
                if len(item) >= 4:
                    title = f"{item[0]} / {item[1]}"
                    val = f"{item[2]} / {item[3]}"
                else:
                    title = item[0]
                    val = item[1]
                rows.append((title, val, val))
            table = doc.add_table(rows=len(rows) + 1, cols=2, style="Light Shading Accent 1")
            hdr = table.rows[0].cells
            p = hdr[0].paragraphs[0]
            run = p.add_run("Aspect / 方面")
            run.bold = True
            p = hdr[1].paragraphs[0]
            run = p.add_run("Details / 详情")
            run.bold = True
            for i, (title, val, _) in enumerate(rows):
                cells = table.rows[i + 1].cells
                cells[0].text = title
                cells[1].text = val

    # --- Visa ---
    if "visa" in modules:
        doc.add_heading("🛂 Visa / 签证信息", level=1)

        # Basic visa info (from country data)
        _add_bilingual_table(doc, "📋 Basic Info / 基本概况", _visa_rows(c))

        # Detailed visa guide (from supplement)
        # supplement["visa_details"] = {
        #   "how_to_apply": [(step_en, step_cn), ...],
        #   "where_to_apply": [(place_en, place_cn, address_en, address_cn, link_en_cn), ...],
        #   "process": [(step_en, step_cn, detail_en, detail_cn), ...],
        #   "documents": [(doc_en, doc_cn, notes_en, notes_cn), ...],
        #   "tips": [(tip_en, tip_cn), ...]
        # }
        vd = supplement.get("visa_details", {})

        if vd:
            # How to apply
            if vd.get("how_to_apply"):
                doc.add_heading("📝 How to Apply / 如何申请", level=2)
                for step in vd["how_to_apply"]:
                    if len(step) >= 2:
                        doc.add_paragraph(f"• {step[0]} / {step[1]}", style="List Bullet")
                    else:
                        doc.add_paragraph(f"• {step[0]}", style="List Bullet")

            # Where to apply
            if vd.get("where_to_apply"):
                doc.add_heading("📍 Where to Apply / 去哪里申请", level=2)
                rows = []
                for w in vd["where_to_apply"]:
                    if len(w) >= 5:
                        name = f"{w[0]} / {w[1]}"
                        addr = f"{w[2]} / {w[3]}"
                        link = f"🔗 {w[4]}"
                        rows.append((name, f"{addr}\n{link}", f"{addr}\n{link}"))
                    else:
                        name = w[0]
                        rows.append((name, w[1], w[1]))
                if rows:
                    _add_bilingual_table(doc, "", rows)

            # Visa process flow
            if vd.get("process"):
                doc.add_heading("🔄 Visa Process / 签证流程", level=2)
                rows = []
                for i, p in enumerate(vd["process"]):
                    step_label = f"Step {i+1} / 第{i+1}步"
                    if len(p) >= 4:
                        val = f"{p[0]} / {p[1]}\n💡 {p[2]} / {p[3]}"
                    elif len(p) >= 2:
                        val = f"{p[0]} / {p[1]}"
                    else:
                        val = p[0]
                    rows.append((step_label, val, val))
                table = doc.add_table(rows=len(rows) + 1, cols=2, style="Light Shading Accent 1")
                hdr = table.rows[0].cells
                p = hdr[0].paragraphs[0]
                run = p.add_run("Step / 步骤")
                run.bold = True
                p = hdr[1].paragraphs[0]
                run = p.add_run("Action / 操作")
                run.bold = True
                for i, (step_label, val, _) in enumerate(rows):
                    cells = table.rows[i + 1].cells
                    cells[0].text = step_label
                    cells[1].text = val

            # Required documents
            if vd.get("documents"):
                doc.add_heading("📄 Required Documents / 所需材料", level=2)
                rows = []
                for d in vd["documents"]:
                    if len(d) >= 4:
                        name = f"{d[0]} / {d[1]}"
                        notes = f"{d[2]} / {d[3]}"
                        rows.append((name, notes, notes))
                    elif len(d) >= 2:
                        rows.append((f"{d[0]} / {d[1]}", "", ""))
                if rows:
                    _add_bilingual_table(doc, "", rows)

            # Tips
            if vd.get("tips"):
                doc.add_heading("💡 Tips / 小贴士", level=2)
                for t in vd["tips"]:
                    if len(t) >= 2:
                        doc.add_paragraph(f"• {t[0]} / {t[1]}", style="List Bullet")
                    else:
                        doc.add_paragraph(f"• {t[0]}", style="List Bullet")

        _add_warning(
            doc,
            "Visa policies change frequently. Verify via official sources before travel.",
            "签证政策随时变化，出行前请通过官方渠道核实最新要求。",
        )

    # --- Laws ---
    if "laws" in modules:
        _add_bilingual_table(doc, "⚖️ Laws & Regulations / 法律法规", _laws_rows(c))
        _add_warning(
            doc,
            "Observe all local laws. Ignorance is not a defence.",
            "各国法律不同，请遵守当地法规，不知情不是免责理由。",
        )

    # --- Culture ---
    if "culture" in modules:
        _add_bilingual_table(doc, "🎭 Culture & Etiquette / 文化与礼仪", _culture_rows(c))

    # --- Phrases / Local Language (from supplement) ---
    if "phrases" in modules and supplement.get("phrases"):
        doc.add_heading("🗣️ Useful Phrases / 当地常用语", level=1)
        doc.add_heading("Common Phrases / 常用表达", level=2)
        rows = []
        for ph in supplement["phrases"]:
            if len(ph) >= 5:
                situation = f"{ph[0]} / {ph[1]}"
                pron = ph[3] if ph[3] else ""
                if pron:
                    val = f"{ph[2]} [{pron}] / {ph[4]}"
                else:
                    val = f"{ph[2]} / {ph[4]}"
                rows.append((situation, val, val))
            elif len(ph) >= 4:
                situation = f"{ph[0]} / {ph[1]}"
                val = f"{ph[2]} / {ph[3]}"
                rows.append((situation, val, val))
            else:
                rows.append((ph[0], ph[1], ph[1]))
        if rows:
            _add_bilingual_table(doc, "", rows)
            p = doc.add_paragraph()
            run = p.add_run("💡 Learn a few local phrases. Locals appreciate the effort. / 学几句当地话，当地人会更友善。")
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # --- Weather & Best Time (from supplement) ---
    if "weather" in modules and supplement.get("weather"):
        w = supplement["weather"]
        doc.add_heading("🌤️ Weather & Best Time to Visit / 天气与最佳旅行季节", level=1)

        if w.get("overview"):
            overview_en, overview_cn = w["overview"]
            p = doc.add_paragraph()
            run = p.add_run(f"{overview_en} / {overview_cn}")
            run.font.size = Pt(11)

        if w.get("best_time"):
            best_en, best_cn = w["best_time"]
            p = doc.add_paragraph()
            run = p.add_run(f"✅ Best Time / 最佳旅行时间: {best_en} / {best_cn}")
            run.bold = True

        if w.get("seasons"):
            rows = [("Season / 季节", "Temp / 温度", "Clothing / 着装建议")]
            for s in w["seasons"]:
                season = f"{s[0]} / {s[1]}"
                temp = f"{s[2]} / {s[3]}" if len(s) > 3 else s[2]
                clothing = f"{s[4]} / {s[5]}" if len(s) > 5 else ""
                rows.append((season, temp, clothing))
            table = doc.add_table(rows=len(rows), cols=3, style="Light Shading Accent 1")
            for i, (col0, col1, col2) in enumerate(rows):
                cells = table.rows[i].cells
                cells[0].text = col0
                cells[1].text = col1
                cells[2].text = col2
                if i == 0:
                    for cell in cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

        p = doc.add_paragraph()
        run = p.add_run("💡 Check weather forecast before departure. / 出发前请查询实时天气预报。")
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # --- Safety ---
    if "safety" in modules:
        _add_bilingual_table(doc, "🚨 Safety Guide / 安全指南", _safety_rows(c))

    # --- Consular ---
    if "consular" in modules:
        doc.add_heading("🏛️ Consular Assistance / 领事协助", level=1)

        # Emergency procedures
        doc.add_heading("Emergency Procedures / 通用紧急流程", level=2)
        procedures = [
            (
                "Lost Passport / 护照丢失",
                [
                    "File police report immediately / 立即报警",
                    "Contact your embassy or consulate / 联系本国使领馆",
                    "Apply for emergency travel document / 申请紧急旅行证件",
                ],
            ),
            (
                "Detention / 被拘留",
                [
                    "Request consular notification / 要求通知本国使领馆",
                    "Embassy arranges consular visit / 使馆安排领事探视",
                    "Legal assistance / 寻求法律援助",
                ],
            ),
            (
                "Medical Emergency / 紧急医疗",
                [
                    "Call local emergency number / 拨打当地急救电话",
                    "Contact travel insurance / 联系旅行保险",
                    "Contact embassy if needed / 必要时联系使领馆",
                ],
            ),
        ]
        for title_en_cn, steps in procedures:
            p = doc.add_paragraph()
            run = p.add_run(f"📌 {title_en_cn}")
            run.bold = True
            for step in steps:
                doc.add_paragraph(step, style="List Bullet")

        # China embassy reference
        doc.add_heading(
            "China Embassy Reference / 中国驻外使领馆参考", level=2
        )
        p = doc.add_paragraph()
        run = p.add_run(
            "📌 For other nationalities, please contact your own embassy. / 其他国籍请联系本国使领馆。"
        )
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        hotline = os.environ.get("CONSULAR_HOTLINE", "+86-10-12308")
        emb = c.get("china_embassy_cn", {})
        if emb:
            emb_table_data = [
                ("Embassy City / 所在城市", emb.get("city", ""), emb.get("city", "")),
                ("Address / 地址", emb.get("address", ""), emb.get("address", "")),
                ("Phone / 电话", emb.get("phone", ""), emb.get("phone", "")),
            ]
            if emb.get("after_hours"):
                emb_table_data.append(
                    ("Emergency / 紧急热线", emb.get("after_hours", ""), emb.get("after_hours", ""))
                )
            _add_bilingual_table(doc, "", emb_table_data)

        consulates = c.get("china_consulates_cn", [])
        if consulates:
            cons_rows = []
            for cs in consulates:
                label = f"Consulate ({cs.get('region','')}) / 总领馆（{cs.get('region','')}）"
                val = f"{cs.get('city','')} / {cs.get('phone','')}"
                cons_rows.append((label, val, val))
            _add_bilingual_table(doc, "", cons_rows)

        p = doc.add_paragraph()
        run = p.add_run(f"MFA Hotline / 外交部领事保护热线: {hotline} (24/7)")
        run.bold = True

        # Guide to find own embassy
        doc.add_heading(
            "How to Find Your Own Embassy / 如何查找本国使领馆", level=2
        )
        tips = [
            "Search online: [your nationality] embassy in [country name] / 在线搜索：[您的国籍]驻[目的地国家]大使馆",
            "Register with your embassy before departure / 出发前在本国使领馆登记",
            "Save your embassy phone number / 保存本国使馆电话",
            "Take a photo of your passport / 护照拍照留备份",
        ]
        for tip in tips:
            doc.add_paragraph(tip, style="List Bullet")

    # --- Entry & Exit / Customs (from supplement) ---
    if "entry_exit" in modules and supplement.get("entry_exit"):
        ee = supplement["entry_exit"]
        doc.add_heading("🛂 Entry & Exit / 出入境与海关", level=1)

        if ee.get("arrival"):
            doc.add_heading("📥 Arrival Process / 入境流程", level=2)
            for a in ee["arrival"]:
                if len(a) >= 2:
                    doc.add_paragraph(f"• {a[0]} / {a[1]}", style="List Bullet")
                else:
                    doc.add_paragraph(f"• {a[0]}", style="List Bullet")

        if ee.get("customs"):
            doc.add_heading("📦 Customs Regulations / 海关规定", level=2)
            rows = []
            for cu in ee["customs"]:
                item = f"{cu[0]} / {cu[1]}"
                rule = f"{cu[2]} / {cu[3]}" if len(cu) > 3 else cu[2]
                rows.append((item, rule, rule))
            if rows:
                _add_bilingual_table(doc, "", rows)

        if ee.get("departure"):
            doc.add_heading("📤 Departure Process / 出境流程", level=2)
            for d in ee["departure"]:
                if len(d) >= 2:
                    doc.add_paragraph(f"• {d[0]} / {d[1]}", style="List Bullet")
                else:
                    doc.add_paragraph(f"• {d[0]}", style="List Bullet")

        if ee.get("tips"):
            doc.add_heading("💡 Tips / 小贴士", level=2)
            for t in ee["tips"]:
                if len(t) >= 2:
                    doc.add_paragraph(f"• {t[0]} / {t[1]}", style="List Bullet")
                else:
                    doc.add_paragraph(f"• {t[0]}", style="List Bullet")

    # --- Travel Plan ---
    if "travel" in modules:
        _add_bilingual_table(doc, f"🗺️ Travel Plan / 旅行计划 ({days}-Day / {days}天)", _travel_rows(c))
        p = doc.add_paragraph()
        run = p.add_run(
            "Cost estimates for reference. Actual costs vary by season and personal spending. / "
            "费用为参考值，实际因季节和个人消费而异。"
        )
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # --- Budget Detail (from supplement) ---
    if "budget_detail" in modules and supplement.get("budget_detail"):
        doc.add_heading("💰 Detailed Budget / 详细预算", level=1)
        rows = []
        total_pct = 0
        for b in supplement["budget_detail"]:
            category = f"{b[0]} / {b[1]}"
            amount = f"{b[2]} / {b[3]}"
            pct = b[4] if len(b) > 4 else ""
            notes = b[5] if len(b) > 5 else ""
            detail = amount
            if notes:
                detail += f" | {notes}"
            rows.append((category, detail, detail))
            if pct:
                try:
                    total_pct += float(pct.rstrip("%"))
                except:
                    pass
        _add_bilingual_table(doc, "", rows)
        if total_pct > 0:
            p = doc.add_paragraph()
            run = p.add_run(f"📊 Estimated total: ~{total_pct:.0f}% of trip budget shown above. / 以上预算约占旅行总费用的{total_pct:.0f}%。")
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        p = doc.add_paragraph()
        run = p.add_run("💡 Budget estimates vary by season and personal spending. / 预算因季节和个人消费而异。")
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # --- Work Guide ---
    if "work" in modules:
        _add_bilingual_table(doc, "💼 Work Guide / 工作指南", _work_rows(c))

    # --- Study Guide ---
    if "study" in modules:
        _add_bilingual_table(doc, "🎓 Study Guide / 留学指南", _study_rows(c))

    # --- Emergency Guide (from supplement) ---
    # supplement["emergency"] = {
    #   "phones": [(en_name, cn_name, number), ...],
    #   "injury": [step_en_cn, ...],
    #   "hospitals": [(en_name, cn_name, address_en_cn, phone), ...]
    # }
    if "emergency" in modules and supplement.get("emergency"):
        em = supplement["emergency"]
        doc.add_heading("🆘 Emergency Guide / 应急指南", level=1)

        # Emergency phone numbers
        if em.get("phones"):
            doc.add_heading("📞 Emergency Numbers / 紧急报警电话", level=2)
            rows = []
            for ph in em["phones"]:
                name = f"{ph[0]} / {ph[1]}"
                rows.append((name, ph[2], ph[2]))
            _add_bilingual_table(doc, "", rows)
            p = doc.add_paragraph()
            run = p.add_run("💡 Save these numbers before travel. / 出行前请保存这些号码。")
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # Injury treatment
        if em.get("injury"):
            doc.add_heading("🤕 Injury & Illness / 受伤与生病处理", level=2)
            for step in em["injury"]:
                doc.add_paragraph(step, style="List Bullet")

        # Nearby hospitals
        if em.get("hospitals"):
            doc.add_heading("🏥 Nearby Hospitals / 附近医院信息", level=2)
            rows = []
            for h in em["hospitals"]:
                name = f"{h[0]} / {h[1]}"
                addr = h[2]
                phone = h[3]
                rows.append((name, f"{addr} | {phone}", f"{addr} | {phone}"))
            _add_bilingual_table(doc, "", rows)
            _add_warning(doc,
                "For medical emergencies, call local ambulance first. / 紧急医疗请先拨打当地急救电话。",
                "医疗紧急情况请先拨打当地急救电话。")

    # --- Attractions (from supplement) ---
    # supplement["attractions"] = list of (name_en, name_cn, ticket_en_cn, desc_en_cn, how_to_get_en, how_to_get_cn)
    if "attractions" in modules and supplement.get("attractions"):
        doc.add_heading("🏛️ Attractions & Tickets / 景点与门票", level=1)
        rows = []
        for a in supplement["attractions"]:
            name = f"{a[0]} / {a[1]}"
            ticket = a[2] if len(a) > 2 else ""
            desc = a[3] if len(a) > 3 else ""
            # How to get there
            if len(a) > 5 and a[4]:
                how = f"🚇 {a[4]} / {a[5]}"
            else:
                how = ""
            if how:
                detail = f"🎫 {ticket} | 📝 {desc} | {how}"
            else:
                detail = f"🎫 {ticket} | 📝 {desc}"
            rows.append((name, detail, detail))
        _add_bilingual_table(doc, "", rows)
        p = doc.add_paragraph()
        run = p.add_run("💡 Ticket prices for reference. Transport info may change. Check before travel.")
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run2 = p.add_run(" / 门票为参考价，交通信息可能有变，出行前请核实。")
        run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # --- Food (from supplement) ---
    # supplement["foods"] = list of (name_en, name_cn, desc_en_cn)
    if "food" in modules and supplement.get("foods"):
        doc.add_heading("🍽️ Local Food & Cuisine / 当地美食", level=1)
        rows = []
        for f in supplement["foods"]:
            name = f"{f[0]} / {f[1]}"
            desc = f[2]
            rows.append((name, desc, desc))
        _add_bilingual_table(doc, "", rows)

    # --- Transport (from supplement) ---
    # supplement["transport"] = list of (method_en, method_cn, desc_en_cn) or (method_bilingual, desc)
    if "transport" in modules and supplement.get("transport"):
        doc.add_heading("🚇 Transport Guide / 交通指南", level=1)
        rows = []
        for t in supplement["transport"]:
            if len(t) >= 3:
                method = f"{t[0]} / {t[1]}"
                desc = t[2]
            else:
                method = t[0]
                desc = t[1]
            rows.append((method, desc, desc))
        _add_bilingual_table(doc, "", rows)

    # --- Itinerary (from supplement) ---
    # supplement["itinerary"] = list of (day_label, activity_bilingual)
    #   or (day_label, activity_en, activity_cn)
    if "itinerary" in modules and supplement.get("itinerary"):
        doc.add_heading("📅 Suggested Itinerary / 推荐行程安排", level=1)
        rows = []
        for it in supplement["itinerary"]:
            day_label = it[0]
            if len(it) >= 3:
                # pre-merge bilingual into one string for the value
                val = f"{it[1]} / {it[2]}"
            else:
                val = it[1]
            rows.append((day_label, val, val))
        table = doc.add_table(rows=len(rows) + 1, cols=2, style="Light Shading Accent 1")
        hdr = table.rows[0].cells
        p = hdr[0].paragraphs[0]
        run = p.add_run("Day / 日期")
        run.bold = True
        p = hdr[1].paragraphs[0]
        run = p.add_run("Activities / 活动安排")
        run.bold = True
        for i, (day_label, act, _) in enumerate(rows):
            cells = table.rows[i + 1].cells
            cells[0].text = day_label
            cells[1].text = act

    # --- Accommodation (from supplement) ---
    # supplement["accommodation"] = {
    #   "budget": [(name_en, name_cn, price_en, price_cn, features_en, features_cn, stars), ...],
    #   "midrange": same format,
    #   "luxury": same format
    # }
    if "accommodation" in modules and supplement.get("accommodation"):
        acc = supplement["accommodation"]
        doc.add_heading("🏨 Accommodation / 住宿推荐", level=1)

        tier_config = [
            ("budget", "💰 Budget / 经济型", "经济型"),
            ("midrange", "💵 Mid-Range / 中档", "舒适型"),
            ("luxury", "👑 Luxury / 高档", "豪华型"),
        ]
        for tier_key, tier_heading_en, tier_heading_cn in tier_config:
            if acc.get(tier_key):
                doc.add_heading(f"{tier_heading_en}\n{tier_heading_cn}", level=2)
                rows = []
                for a in acc[tier_key]:
                    name = f"{a[0]} / {a[1]}"
                    price = f"💰 {a[2]} / {a[3]}"
                    features = a[4] if len(a) > 4 else ""
                    features_cn = a[5] if len(a) > 5 else ""
                    stars = a[6] if len(a) > 6 else ""
                    recommend = ""
                    if stars == "★":
                        recommend = "⭐ "
                    elif stars == "★★":
                        recommend = "⭐⭐ "
                    elif stars == "★★★":
                        recommend = "⭐⭐⭐ "
                    detail = price
                    if features or features_cn:
                        detail += f" | 📍 {features} / {features_cn}"
                    detail = f"{recommend}{detail}"
                    rows.append((name, detail, detail))
                if rows:
                    _add_bilingual_table(doc, "", rows)

        p = doc.add_paragraph()
        run = p.add_run("💡 Prices for reference. Book early for best rates. / 价格为参考价，建议提前预订。")
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # --- Insurance (from supplement) ---
    if "insurance" in modules and supplement.get("insurance"):
        ins = supplement["insurance"]
        doc.add_heading("🛡️ Travel Insurance / 旅行保险", level=1)

        if ins.get("recommendation"):
            rec_en, rec_cn = ins["recommendation"]
            p = doc.add_paragraph()
            run = p.add_run(f"✅ {rec_en} / {rec_cn}")
            run.bold = True

        if ins.get("types"):
            rows = []
            for t in ins["types"]:
                type_name = f"{t[0]} / {t[1]}"
                desc = f"{t[2]} / {t[3]}" if len(t) > 3 else t[2]
                suggested = f"{t[4]} / {t[5]}" if len(t) > 5 else ""
                detail = desc
                if suggested:
                    detail += f"\n💡 {suggested}"
                rows.append((type_name, detail, detail))
            if rows:
                _add_bilingual_table(doc, "", rows)

        if ins.get("tips"):
            doc.add_heading("💡 Tips / 小贴士", level=2)
            for t in ins["tips"]:
                if len(t) >= 2:
                    doc.add_paragraph(f"• {t[0]} / {t[1]}", style="List Bullet")
                else:
                    doc.add_paragraph(f"• {t[0]}", style="List Bullet")

        p = doc.add_paragraph()
        run = p.add_run("💡 Travel insurance is strongly recommended for all trips. / 强烈建议为所有行程购买旅行保险。")
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # --- Shopping (from supplement) ---
    if "shopping" in modules and supplement.get("shopping"):
        shop = supplement["shopping"]
        doc.add_heading("🛍️ Shopping Guide / 购物指南", level=1)

        if shop.get("areas"):
            doc.add_heading("📍 Shopping Areas / 购物地点", level=2)
            rows = []
            for a in shop["areas"]:
                name = f"{a[0]} / {a[1]}"
                desc = f"{a[2]} / {a[3]}" if len(a) > 3 else a[2]
                rows.append((name, desc, desc))
            if rows:
                _add_bilingual_table(doc, "", rows)

        if shop.get("specialties"):
            doc.add_heading("🎁 Local Specialties / 当地特色商品", level=2)
            rows = []
            for s in shop["specialties"]:
                item = f"{s[0]} / {s[1]}"
                price = s[2] if len(s) > 2 else ""
                where = f"📍 {s[3]} / {s[4]}" if len(s) > 4 else ""
                detail = price
                if where:
                    detail += f" | {where}"
                rows.append((item, detail, detail))
            if rows:
                _add_bilingual_table(doc, "", rows)

        if shop.get("tax_free"):
            tax_en, tax_cn = shop["tax_free"]
            p = doc.add_paragraph()
            run = p.add_run(f"📄 Tax-Free / 退税: {tax_en} / {tax_cn}")
            run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)

        if shop.get("tips"):
            doc.add_heading("💡 Tips / 购物小贴士", level=2)
            for t in shop["tips"]:
                if len(t) >= 2:
                    doc.add_paragraph(f"• {t[0]} / {t[1]}", style="List Bullet")
                else:
                    doc.add_paragraph(f"• {t[0]}", style="List Bullet")

    # --- Communication & Network (from supplement) ---
    if "communication" in modules and supplement.get("communication"):
        comm = supplement["communication"]
        doc.add_heading("📱 Communication & Network / 通讯与网络", level=1)

        if comm.get("sim_card"):
            sim_en, sim_cn = comm["sim_card"]
            p = doc.add_paragraph()
            run = p.add_run(f"📶 Local SIM / 当地SIM卡: {sim_en} / {sim_cn}")
            run.font.size = Pt(11)

        if comm.get("wifi"):
            wifi_en, wifi_cn = comm["wifi"]
            p = doc.add_paragraph()
            run = p.add_run(f"📡 WiFi: {wifi_en} / {wifi_cn}")
            run.font.size = Pt(11)

        if comm.get("roaming"):
            roam_en, roam_cn = comm["roaming"]
            p = doc.add_paragraph()
            run = p.add_run(f"📞 Roaming / 漫游: {roam_en} / {roam_cn}")
            run.font.size = Pt(11)

        # Huawei SkyTone (天际通) recommendation for Huawei phone users
        p = doc.add_paragraph()
        run = p.add_run("📱 Huawei SkyTone (天际通) / 华为天际通")
        run.bold = True
        p = doc.add_paragraph(
            "🔹 Huawei phone users: Use the built-in SkyTone (天际通) app for global data roaming — "
            "no SIM swap needed, one-click activation, daily/package plans available. "
            "Supports 100+ countries, pre-download before departure for best experience. "
            "/ 华为手机用户：使用系统自带的天际通APP即可开通全球境外上网——不需换卡，一键激活，支持日套餐/流量包，覆盖100+国家地区。出发前下载好效果最佳。"
        )

        if comm.get("useful_apps"):
            doc.add_heading("📲 Useful Apps / 实用App", level=2)
            rows = []
            for app in comm["useful_apps"]:
                app_name = app[0]
                desc = f"{app[1]} / {app[2]}" if len(app) > 2 else app[1]
                rows.append((app_name, desc, desc))
            if rows:
                _add_bilingual_table(doc, "", rows)

        p = doc.add_paragraph()
        run = p.add_run("💡 Buy local SIM or activate roaming before departure. / 出发前购买当地SIM卡或开通国际漫游。")
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # --- Footer ---
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("⚠️ DISCLAIMER / 免责声明")
    run.bold = True
    doc.add_paragraph(
        "All data for reference only. Verify via official sources before travel. "
        "Users of all nationalities: contact your own embassy for consular assistance."
    )
    doc.add_paragraph(
        "所有信息仅供参考，出行前请通过官方渠道核实。"
        "所有国籍用户如需领事协助请联系本国使领馆。"
    )

    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("© 出国游全球通版 v1.9.0 | 媳妇智投Pro出品，必属精品 🦊")
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.size = Pt(9)

    # Save
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    safe_name = c.get("code", country_code).upper()
    fname = f"旅行方案_{safe_name}_{datetime.datetime.now().strftime('%Y-%m-%d')}.docx"
    fpath = os.path.join(OUTPUT_DIR, fname)
    doc.save(fpath)
    print(f"✅ 方案已生成: {fpath}")
    print(f"✅ Report saved: {fpath}")
    return fpath


# ---------------------------------------------------------------------------
# CLI entry
# ---------------------------------------------------------------------------

def main():
    if len(sys.argv) < 2:
        print("Usage / 用法:")
        print("  Python3 global_travel_planner.py <country> [module] [...modules]")
        print()
        print("Examples / 示例:")
        print("  Python3 global_travel_planner.py JP")
        print("      → Full docx for Japan / 日本完整方案")
        print("  Python3 global_travel_planner.py TH visa safety")
        print("      → Thailand visa + safety only / 泰国签证+安全")
        print("  Python3 global_travel_planner.py list")
        print("      → List all countries / 列出所有国家")
        print()
        print("Modules / 模块: visa, laws, culture, safety, consular, travel, work, study")
        print("Set nationality / 设置国籍: CONSULAR_NATIONALITY=US python3 ...")
        print()
        data = _load_data()
        print(f"Countries available / 已收录: {len(data)}")
        return

    query = sys.argv[1]
    if query.upper() == "LIST":
        data = _load_data()
        print(f"Countries available / 已收录: {len(data)}\n")
        for code in sorted(data.keys()):
            c = data[code]
            print(f"  {code} - {c['name_cn']} / {c['name_en']}")
        return

    modules = sys.argv[2:] if len(sys.argv) > 2 else None
    generate_docx(query, modules)


if __name__ == "__main__":
    main()
