#!/usr/bin/env python3
"""Convert Markdown file to a formatted Word (.docx) document.

Usage:
    python3 md_to_docx.py <input.md> [output.docx]
    If output is omitted, writes to the same path with .docx extension.

Dependencies: python-docx (pip install python-docx)
"""

import sys
import re
import os

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx is required. Install it with: pip install python-docx", file=sys.stderr)
    sys.exit(1)


FONT_CN = 'PingFang SC'
FONT_EN = 'Helvetica Neue'
FONT_MONO = 'SF Mono'
PRIMARY_COLOR = RGBColor(0x24, 0x71, 0xA3)
TEXT_COLOR = RGBColor(0x2C, 0x3E, 0x50)
LIGHT_COLOR = RGBColor(0x5D, 0x6D, 0x7E)


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elm)


def set_run_font(run, size=None, bold=False, color=None, font_name=None):
    run.font.name = font_name or FONT_EN
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CN)
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def add_formatted_para(doc, text, style=None, font_size=None, bold=False, color=None):
    """Add a paragraph with formatted text."""
    p = doc.add_paragraph(style=style)
    if text is None:
        return p
    run = p.add_run(text)
    set_run_font(run, size=font_size, bold=bold, color=color)
    return p


def parse_inline_runs(paragraph, text):
    """Parse inline Markdown and add runs to paragraph."""
    # Split text into segments: code, bold, italic, links, plain
    segments = []
    i = 0
    while i < len(text):
        # Inline code
        if text[i] == '`':
            end = text.find('`', i + 1)
            if end > i:
                segments.append(('code', text[i+1:end]))
                i = end + 1
                continue
        # Bold
        if text[i:i+2] == '**':
            end = text.find('**', i + 2)
            if end > i:
                segments.append(('bold', text[i+2:end]))
                i = end + 2
                continue
        # Italic
        if text[i] == '*' and (i == 0 or text[i-1] != '*'):
            end = text.find('*', i + 1)
            if end > i and (end == len(text)-1 or text[end+1] != '*'):
                segments.append(('italic', text[i+1:end]))
                i = end + 1
                continue
        # Links
        link_match = re.match(r'\[([^\]]+)\]\(([^)]+)\)', text[i:])
        if link_match:
            segments.append(('link', link_match.group(1)))
            i += len(link_match.group(0))
            continue
        # Plain text - collect until next special char
        j = i
        while j < len(text) and text[j] not in '*`[':
            j += 1
        if j > i:
            segments.append(('plain', text[i:j]))
            i = j
        else:
            segments.append(('plain', text[i]))
            i += 1

    for seg_type, seg_text in segments:
        run = paragraph.add_run(seg_text)
        set_run_font(run, size=11)
        if seg_type == 'code':
            run.font.name = FONT_MONO
            run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
            run.font.size = Pt(9.5)
        elif seg_type == 'bold':
            run.font.bold = True
        elif seg_type == 'italic':
            run.font.italic = True
        elif seg_type == 'link':
            run.font.color.rgb = PRIMARY_COLOR
            run.font.underline = True


def convert_md_to_docx(md_text, output_path, title):
    doc = Document()

    # ── Configure default style ──
    style = doc.styles['Normal']
    style.font.name = FONT_EN
    style.font.size = Pt(11)
    style.font.color.rgb = LIGHT_COLOR
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CN)

    # ── Title ──
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(title)
    set_run_font(title_run, size=28, bold=True, color=PRIMARY_COLOR)
    doc.add_paragraph()  # spacer

    lines = md_text.split('\n')
    i = 0
    in_code_block = False
    code_buffer = []

    while i < len(lines):
        line = lines[i]

        # Code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                code_text = '\n'.join(code_buffer)
                code_p = doc.add_paragraph()
                code_run = code_p.add_run(code_text)
                code_run.font.name = FONT_MONO
                code_run.font.size = Pt(9)
                code_run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
                # Add shading
                pPr = code_p._element.get_or_add_pPr()
                shd = pPr.makeelement(qn('w:shd'), {
                    qn('w:fill'): 'F7F9FB',
                    qn('w:val'): 'clear',
                })
                pPr.append(shd)
                code_buffer = []
                in_code_block = False
                i += 1
                continue
            else:
                in_code_block = True
                i += 1
                continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # Skip first h1 (already used as title)
        h1_match = re.match(r'^#\s+(.+)$', line)
        h2_match = re.match(r'^##\s+(.+)$', line)
        h3_match = re.match(r'^###\s+(.+)$', line)

        if h1_match and i < 3:
            i += 1
            continue

        if h2_match:
            add_formatted_para(doc, h2_match.group(1), font_size=18, bold=True, color=PRIMARY_COLOR)
            i += 1
            continue

        if h3_match:
            add_formatted_para(doc, h3_match.group(1), font_size=14, bold=True, color=TEXT_COLOR)
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^[-*_]{3,}\s*$', line.strip()):
            doc.add_paragraph('─' * 40)
            i += 1
            continue

        # Tables
        if '|' in line and line.strip().startswith('|'):
            table_rows = []
            headers = []
            while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
                row = lines[i]
                if re.match(r'^\|[\s\-:|]+\|$', row.strip()):
                    i += 1
                    continue
                cells = [c.strip() for c in row.strip().strip('|').split('|')]
                if not headers:
                    headers = cells
                else:
                    table_rows.append(cells)
                i += 1

            if headers or table_rows:
                all_rows = [headers] + table_rows if headers else table_rows
                table = doc.add_table(rows=len(all_rows), cols=len(all_rows[0]))
                table.style = 'Table Grid'

                for ri, row in enumerate(all_rows):
                    for ci, cell_text in enumerate(row):
                        cell = table.cell(ri, ci)
                        cell.text = ''
                        p = cell.paragraphs[0]
                        run = p.add_run(cell_text)
                        if ri == 0 and headers:
                            set_run_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
                            set_cell_shading(cell, '2471A3')
                        else:
                            set_run_font(run, size=10, color=LIGHT_COLOR)

                doc.add_paragraph()  # spacer after table
            continue

        # Ordered list
        ol_match = re.match(r'^(\d+)\.\s+(.+)$', line)
        if ol_match:
            p = doc.add_paragraph(style='List Number')
            # Clear default text, add ours
            p.clear()
            parse_inline_runs(p, ol_match.group(2))
            i += 1
            continue

        # Unordered list
        ul_match = re.match(r'^[-*+]\s+(.+)$', line)
        if ul_match:
            p = doc.add_paragraph(style='List Bullet')
            p.clear()
            parse_inline_runs(p, ul_match.group(1))
            i += 1
            continue

        # Paragraph
        if line.strip():
            p = doc.add_paragraph()
            parse_inline_runs(p, line.strip())
            i += 1
            continue

        # Empty line
        i += 1

    doc.save(output_path)
    print(f'DOCX generated: {output_path}')


def main():
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else os.path.splitext(input_path)[0] + '.docx'

    with open(input_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    title_match = re.search(r'^#\s+(.+)$', md_text, re.MULTILINE)
    title = title_match.group(1) if title_match else os.path.splitext(os.path.basename(input_path))[0]

    convert_md_to_docx(md_text, output_path, title)


if __name__ == '__main__':
    main()
