#!/usr/bin/env python3
"""
docx Linter — 格式质量门

从 YAML 规范读取期望值，逐项检查 docx 实际格式。
不依赖 python-docx API 对 None/0 的差异，直接读 XML 对比。

用法:
  python docx_linter.py 论文.docx --spec 苏州科技大学.yaml [--template-docx 母版模板.docx]
"""

import re
import sys
import yaml
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

# ── XPath helpers ──

def _spacing(p_elem):
    """获取段落的 w:spacing 元素"""
    pPr = p_elem.find(qn('w:pPr'))
    if pPr is None:
        return None
    return pPr.find(qn('w:spacing'))

def _ind(p_elem):
    pPr = p_elem.find(qn('w:pPr'))
    if pPr is None:
        return None
    return pPr.find(qn('w:ind'))

def _eastasia(run):
    rPr = run._element.find(qn('w:rPr'))
    if rPr is None:
        return None
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        return None
    return rFonts.get(qn('w:eastAsia'))

def _borders(p_elem):
    """获取段落边框 w:pBdr"""
    pPr = p_elem.find(qn('w:pPr'))
    if pPr is None:
        return None
    return pPr.find(qn('w:pBdr'))

class Issue:
    def __init__(self, level, check, detail, fix_hint=""):
        self.level = level  # "error" or "warning"
        self.check = check
        self.detail = detail
        self.fix_hint = fix_hint

    def __repr__(self):
        icon = "🔴" if self.level == "error" else "🟡"
        return f"  {icon} [{self.check}] {self.detail}"


class DocxLinter:
    """从 YAML 规范驱动的 docx 格式检查器"""

    def __init__(self, doc_path, spec_path=None, template_docx=None):
        self.doc = Document(doc_path)
        self.doc_path = Path(doc_path)
        self.issues = []
        self.score = 100

        # 加载规范
        self.spec = {}
        if spec_path:
            with open(spec_path, 'r', encoding='utf-8') as f:
                self.spec = yaml.safe_load(f)

        # 加载模板（用于对比样式期望）
        self.tpl = None
        if template_docx:
            self.tpl = Document(template_docx)

    # ═══════════════════════════════════════════
    # 入口
    # ═══════════════════════════════════════════

    def check_all(self):
        self._check_margins()
        self._check_header()
        self._check_structure()
        self._check_appendix()
        self._check_special_paragraphs()
        self._check_tables()
        self._check_citations()
        self._calc_score()
        return self.issues

    # ═══════════════════════════════════════════
    # 1. 页边距
    # ═══════════════════════════════════════════

    def _check_margins(self):
        page = self.spec.get('page', {}).get('margins', {})
        if not page or not self.doc.sections:
            return
        s = self.doc.sections[0]
        checks = {
            'top': (s.top_margin, page.get('top')),
            'bottom': (s.bottom_margin, page.get('bottom')),
            'left': (s.left_margin, page.get('left')),
            'right': (s.right_margin, page.get('right')),
        }
        for side, (actual, expected_str) in checks.items():
            if not expected_str:
                continue
            expected_emu = _parse_length(expected_str)
            if expected_emu is None:
                continue
            diff = abs(actual - expected_emu)
            if diff > 50000:  # ~1mm tolerance
                # Word 可能自动四舍五入
                self.issues.append(Issue(
                    "error", "页边距", f"{side}: {actual} vs 期望 {expected_emu} ({expected_str})",
                    f"检查 section.{side}_margin"
                ))

    # ═══════════════════════════════════════════
    # 2. 页眉
    # ═══════════════════════════════════════════

    def _check_header(self):
        header_cfg = self.spec.get('header', {})
        if not header_cfg or not self.doc.sections:
            return
        expected_text = header_cfg.get('text', '')
        s = self.doc.sections[0]
        h = s.header
        if not h.paragraphs:
            self.issues.append(Issue("error", "页眉", "页眉为空"))
            return

        p = h.paragraphs[0]
        if expected_text and expected_text not in p.text:
            self.issues.append(Issue("warning", "页眉", f"文本不匹配: '{p.text[:30]}' vs '{expected_text}'"))

        # 检查页眉底线
        pBdr = _borders(p._element)
        if pBdr is None:
            self.issues.append(Issue("warning", "页眉", "无页眉底线（w:pBdr）"))

        # 检查字体
        for r in p.runs:
            ea = _eastasia(r)
            expected_ea = header_cfg.get('font_eastasia', '宋体')
            if ea and expected_ea and ea != expected_ea:
                self.issues.append(Issue("warning", "页眉", f"中文字体: {ea} vs 期望 {expected_ea}"))

    # ═══════════════════════════════════════════
    # 3. 结构检查
    # ═══════════════════════════════════════════

    def _check_structure(self):
        """检查附录区段落结构：附录标题→空行→标题→空行→正文"""
        structure = self.spec.get('structure', [])
        if not structure:
            return

        paras = list(self.doc.paragraphs)
        appendix_a_text = self.spec.get('label_map', {}).get('appendix_a', '附录A')
        a_idx = None
        for i, p in enumerate(paras):
            if appendix_a_text in p.text and '译文' in p.text:
                a_idx = i
                break

        if a_idx is None:
            return

        cn_title = self.spec.get('label_map', {}).get('cn_title', '')

        # 跳过连续空行，找到下一个非空段落
        def skip_blanks(start, max_skip=3):
            idx = start + 1
            skipped = 0
            while idx < len(paras) and skipped < max_skip and not paras[idx].text.strip():
                idx += 1
                skipped += 1
            return idx if idx < len(paras) and skipped > 0 else start + 1

        # 附录标签后第一个非空行应该是标题
        title_pos = skip_blanks(a_idx)
        if title_pos < len(paras):
            p = paras[title_pos]
            if cn_title and cn_title not in p.text:
                self.issues.append(Issue(
                    "error", "结构", f"附录A +{title_pos-a_idx} 不是标题: '{p.text[:40]}'",
                    "确保 附录标题 → 空行 → 标题 顺序"
                ))

        # 标题后第一个非空行应该是正文
        body_pos = skip_blanks(title_pos)
        if body_pos < len(paras) and body_pos == title_pos + 1:
            # 标题和正文之间没有空行
            self.issues.append(Issue(
                "warning", "结构", f"标题(p{title_pos})后缺少空行，正文在 p{body_pos}"
            ))

    # ═══════════════════════════════════════════
    # 4. 附录标题 / 文章标题 格式检查
    # ═══════════════════════════════════════════

    def _check_appendix(self):
        """检查附录标题和文章标题的格式（样式名、字体、字号、spacing）"""
        spec_styles = self.spec.get('styles', {})

        for i, p in enumerate(self.doc.paragraphs):
            text = p.text.strip()
            if not text:
                continue

            is_appendix = (
                ('附录A' in text or '附录B' in text)
                and ('译文' in text or '原文' in text)
            )
            cn_title = self.spec.get('label_map', {}).get('cn_title', '')
            en_title = self.spec.get('label_map', {}).get('en_title', '')
            is_title = (cn_title and cn_title in text) or (en_title and en_title in text)

            if not (is_appendix or is_title):
                continue

            label = "附录标题" if is_appendix else "文章标题"
            expected_style = spec_styles.get(label, {})

            # 样式名
            if p.style.name != label and p.style.name != 'Normal':
                self.issues.append(Issue(
                    "error", label, f"p{i} 样式名 '{p.style.name}' != '{label}'",
                    f"p.style = styles['{label}']"
                ))

            # 字体/字号（看第一个 run）
            if p.runs:
                r = p.runs[0]
                if r.font.size:
                    expected_size = _parse_length(expected_style.get('size'))
                    if expected_size and abs(r.font.size - expected_size) > 10000:
                        self.issues.append(Issue(
                            "warning", label, f"p{i} 字号 {r.font.size//12700}pt vs 期望 {expected_style.get('size')}"
                        ))
                if r.font.bold != expected_style.get('bold'):
                    self.issues.append(Issue(
                        "warning", label, f"p{i} bold={r.font.bold} vs 期望 {expected_style.get('bold')}"
                    ))

            # spacing（读 XML，不依赖 API）
            sp = _spacing(p._element)
            if sp is not None:
                bl = sp.get(qn('w:beforeLines'))
                al = sp.get(qn('w:afterLines'))
                line = sp.get(qn('w:line'))
                if is_appendix:
                    if bl != '50' or al != '50':
                        self.issues.append(Issue(
                            "error", label, f"p{i} bL={bl} aL={al} 期望 bL=50 aL=50 (0.5行)",
                            "模板样式设 beforeLines=50, afterLines=50"
                        ))
                    if line != '240':
                        self.issues.append(Issue(
                            "error", label, f"p{i} 行距 line={line} 期望 240 (单倍行距)",
                        ))

    # ═══════════════════════════════════════════
    # 5. 特殊段落
    # ═══════════════════════════════════════════

    def _check_special_paragraphs(self):
        specials = self.spec.get('special_paragraphs', [])
        if not specials:
            return

        for rule in specials:
            keywords = rule.get('keywords', [])
            fmt = rule.get('format', {})
            name = rule.get('name', '?')

            for i, p in enumerate(self.doc.paragraphs):
                text = p.text.strip()
                if not text:
                    continue

                matched = False
                for kw in keywords:
                    if kw in text:
                        matched = True
                        break
                if not matched:
                    continue

                # 检查格式
                # 首行缩进
                if 'first_line_indent' in fmt:
                    fi_val = fmt['first_line_indent']
                    if fi_val is None or fi_val == 'null' or fi_val == '0':
                        # 期望无缩进
                        ind = _ind(p._element)
                        if ind is not None:
                            fl = ind.get(qn('w:firstLine'))
                            if fl and int(fl) > 0:
                                self.issues.append(Issue(
                                    "error", f"特殊段落/{name}", f"p{i} 期望无缩进但 firstLine={fl}",
                                    "在 YAML 设 first_line_indent: null"
                                ))

                # 行距
                expected_ls = fmt.get('line_spacing')
                if expected_ls is not None:
                    expected_line = int(float(expected_ls) * 240)
                    sp = _spacing(p._element)
                    if sp is not None:
                        actual_line = sp.get(qn('w:line'))
                        if actual_line and abs(int(actual_line) - expected_line) > 10:
                            self.issues.append(Issue(
                                "warning", f"特殊段落/{name}", f"p{i} 行距 line={actual_line} 期望 {expected_line}",
                            ))

                # 段间距（读 XML 是 twips，需转 EMU）
                for sa_key, attr_name in [('space_before', 'before'), ('space_after', 'after')]:
                    expected = fmt.get(sa_key)
                    if expected is not None:
                        expected_emu = _parse_length(str(expected))
                        sp = _spacing(p._element)
                        if sp is not None:
                            actual = sp.get(qn(f'w:{attr_name}'))
                            if actual and expected_emu:
                                # w:before/after 是 twips → EMU
                                actual_emu = int(actual) * 635
                                diff = abs(actual_emu - expected_emu)
                                if diff > 20000:  # ~1.5pt tolerance for twips rounding
                                    self.issues.append(Issue(
                                        "warning", f"特殊段落/{name}",
                                        f"p{i} {sa_key}={actual}twips≈{actual_emu}emu 期望 {expected_emu}"
                                    ))

    # ═══════════════════════════════════════════
    # 6. 表格（三线表）
    # ═══════════════════════════════════════════

    def _check_tables(self):
        if not self.doc.tables:
            return
        for t_idx, table in enumerate(self.doc.tables):
            rows = table.rows
            if len(rows) < 2:
                continue
            # 顶线
            top_sz = _cell_border_sz(table, 0, 0, 'top')
            if top_sz is None:
                top_sz = _tbl_border_sz(table, 'top')
            if top_sz is None or int(top_sz) < 8:
                self.issues.append(Issue("error", "三线表", f"表{t_idx+1} 顶线缺失/太细 (sz={top_sz})",
                    "顶线应为 1.5pt (sz=12)"))
            # 底线
            last_row = len(rows) - 1
            bottom_sz = _cell_border_sz(table, last_row, 0, 'bottom')
            if bottom_sz is None:
                bottom_sz = _tbl_border_sz(table, 'bottom')
            if bottom_sz is None or int(bottom_sz) < 8:
                self.issues.append(Issue("error", "三线表", f"表{t_idx+1} 底线缺失/太细",
                    "底线应为 1.5pt (sz=12)"))

    # ═══════════════════════════════════════════
    # 7. 引用上标
    # ═══════════════════════════════════════════

    def _check_citations(self):
        citation_spec = self.spec.get('references', {}).get('citation', {})
        if not isinstance(citation_spec.get('superscript'), bool):
            return
        expected_superscript = citation_spec['superscript']

        total = 0
        not_super = 0
        in_ref = False

        for p in self.doc.paragraphs:
            text = p.text.strip()
            if '参考文献' == text or text.startswith('参考文献'):
                in_ref = True
                continue
            if in_ref:
                if '附录' in text or '致谢' in text:
                    in_ref = False
                continue
            for run in p.runs:
                if run.text and re.search(r'\[\d+', run.text):
                    total += 1
                    if not run.font.superscript:
                        not_super += 1

        mismatched = not_super if expected_superscript else total - not_super
        if total > 0 and mismatched > 0:
            matched_rate = (total - mismatched) / total
            level = "error" if matched_rate < 0.9 else "warning"
            expected = "上标" if expected_superscript else "不上标"
            self.issues.append(Issue(level, "引用上下标",
                f"{mismatched}/{total} 与规范不符，期望{expected}，符合率 {matched_rate:.0%}",
                "按规范确认引用是否上标"))

    # ═══════════════════════════════════════════
    # 统计
    # ═══════════════════════════════════════════

    def _calc_score(self):
        errors = sum(1 for i in self.issues if i.level == "error")
        warnings = sum(1 for i in self.issues if i.level == "warning")
        self.score = max(0, 100 - errors * 8 - warnings * 3)

    def report(self):
        lines = []
        lines.append("=" * 65)
        lines.append(f"  📋 docx Linter — {self.doc_path.name}")
        lines.append(f"  🎯 总分: {self.score}/100  |  🔴 {sum(1 for i in self.issues if i.level=='error')} 错误  🟡 {sum(1 for i in self.issues if i.level=='warning')} 警告")
        lines.append("=" * 65)
        for issue in self.issues:
            lines.append(str(issue))
            if issue.fix_hint:
                lines.append(f"      → {issue.fix_hint}")
        if not self.issues:
            lines.append("  ✅ 零问题，文档格式规范")
        return "\n".join(lines)


# ── XML 辅助 ──

def _parse_length(val):
    """'12pt'|'2.54cm' → EMU"""
    if val is None:
        return None
    if isinstance(val, (int, float)):
        return int(val) if val > 1000 else int(val * 12700)
    val = str(val).strip()
    m = re.match(r'^([\d.]+)\s*(pt|cm|mm)$', val, re.I)
    if m:
        num = float(m.group(1))
        unit = m.group(2).lower()
        if unit == 'pt': return int(num * 12700)
        if unit == 'cm': return int(num * 360000)
        if unit == 'mm': return int(num * 36000)
    return None


def _cell_border_sz(table, row, col, edge):
    try:
        cell = table.rows[row].cells[col]
    except IndexError:
        return None
    tcPr = cell._tc.find(qn('w:tcPr'))
    if tcPr is None:
        return None
    borders = tcPr.find(qn('w:tcBorders'))
    if borders is None:
        return None
    b = borders.find(qn(f'w:{edge}'))
    if b is None:
        return None
    return b.get(qn('w:sz'))


def _tbl_border_sz(table, edge):
    tblPr = table._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        return None
    borders = tblPr.find(qn('w:tblBorders'))
    if borders is None:
        return None
    b = borders.find(qn(f'w:{edge}'))
    if b is None:
        return None
    return b.get(qn('w:sz'))


def main():
    import argparse
    p = argparse.ArgumentParser(description="docx 格式质量门")
    p.add_argument("doc", help="docx 文件路径")
    p.add_argument("--spec", help="YAML 规范文件")
    p.add_argument("--template-docx", help="母版模板 .docx")
    args = p.parse_args()

    linter = DocxLinter(args.doc, args.spec, args.template_docx)
    linter.check_all()
    print(linter.report())
    sys.exit(0 if linter.score >= 70 else 1)


if __name__ == "__main__":
    main()
