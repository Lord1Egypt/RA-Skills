#!/usr/bin/env python3
"""
文档格式校正脚本
功能：对 DOCX 文档进行纯格式层面的校正，不修改任何文字内容。
依赖：pip install python-docx
"""

import sys
import copy
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("需要安装 python-docx: pip install python-docx", file=sys.stderr)
    sys.exit(1)


# ── 常用格式预设 ──────────────────────────────────────────────

PRESETS = {
    "gb-t9704": {  # 国家标准 GB/T 9704（公文格式）
        "page_width_cm": 21.0,
        "page_height_cm": 29.7,
        "margin_top_cm": 3.7,
        "margin_bottom_cm": 3.5,
        "margin_left_cm": 2.8,
        "margin_right_cm": 2.6,
        "title_font": "方正小标宋简体",
        "title_size_pt": 22,
        "body_font": "仿宋_GB2312",
        "body_size_pt": 16,
        "line_spacing_pt": 28.95,  # 3号字行距
        "first_line_indent_chars": 2,
    },
    "thesis-cn": {  # 中文论文通用
        "page_width_cm": 21.0,
        "page_height_cm": 29.7,
        "margin_top_cm": 2.54,
        "margin_bottom_cm": 2.54,
        "margin_left_cm": 3.17,
        "margin_right_cm": 3.17,
        "title_font": "黑体",
        "title_size_pt": 22,
        "body_font": "宋体",
        "body_size_pt": 12,
        "line_spacing_pt": 20,
        "first_line_indent_chars": 2,
    },
    "general": {  # 通用文档
        "page_width_cm": 21.0,
        "page_height_cm": 29.7,
        "margin_top_cm": 2.54,
        "margin_bottom_cm": 2.54,
        "margin_left_cm": 3.18,
        "margin_right_cm": 3.18,
        "title_font": "微软雅黑",
        "title_size_pt": 16,
        "body_font": "宋体",
        "body_size_pt": 12,
        "line_spacing_pt": 20,
        "first_line_indent_chars": 2,
    },
}


def apply_page_layout(doc, preset: dict):
    """设置页面尺寸与页边距"""
    for section in doc.sections:
        section.page_width = Cm(preset["page_width_cm"])
        section.page_height = Cm(preset["page_height_cm"])
        section.top_margin = Cm(preset["margin_top_cm"])
        section.bottom_margin = Cm(preset["margin_bottom_cm"])
        section.left_margin = Cm(preset["margin_left_cm"])
        section.right_margin = Cm(preset["margin_right_cm"])


def clean_paragraph_formatting(para, preset: dict, is_title: bool = False):
    """清理段落格式：去除多余空格、修正行距、首行缩进等"""
    pf = para.paragraph_format

    # 修正行距
    pf.line_spacing = Pt(preset["line_spacing_pt"])
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    # 首行缩进（仅正文段落）
    if not is_title and para.text.strip():
        pf.first_line_indent = Pt(preset["body_size_pt"] * preset["first_line_indent_chars"])
    elif is_title:
        pf.first_line_indent = Pt(0)
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER


def clean_run_formatting(run, preset: dict, is_title: bool = False):
    """清理 run 级别格式：字体、字号，不改动文字"""
    if not run.text or not run.text.strip():
        return

    target_font = preset["title_font"] if is_title else preset["body_font"]
    target_size = Pt(preset["title_size_pt"] if is_title else preset["body_size_pt"])

    run.font.name = target_font
    run.font.size = target_size
    # 中文字体设置
    run._element.rPr.rFonts.set(qn('w:eastAsia'), target_font)


def is_title_paragraph(para) -> bool:
    """启发式判断标题段落"""
    text = para.text.strip()
    if not text:
        return False
    style_name = (para.style.name or "").lower()
    if any(kw in style_name for kw in ["heading", "title", "标题"]):
        return True
    # 检测中文序号标题（一、 / 1. / 第一章 等）
    import re
    if re.match(r'^[一二三四五六七八九十百千]+[、．.]', text):
        return True
    if re.match(r'^第[一二三四五六七八九十百千]+[章节篇部]', text):
        return True
    if re.match(r'^\d+[\.、．]\s', text):
        return True
    return False


def remove_trailing_spaces(para):
    """去除段落末尾多余空格，不改动可见文字"""
    for run in para.runs:
        if run.text:
            run.text = run.text.rstrip(' \t') + ('' if not run.text.rstrip(' \t') else '')
            # 保留换行符
            if run.text.endswith('\n'):
                cleaned = run.text.rstrip(' \t\n')
                run.text = cleaned + '\n' if cleaned else '\n'
            else:
                run.text = run.text.rstrip(' \t')


def remove_empty_paragraphs(doc):
    """去除文档中完全空白的段落（无文字、无图片）"""
    to_remove = []
    for para in doc.paragraphs:
        if not para.text.strip() and not para._element.findall(qn('w:r')):
            # 保留至少一个空行用于段落分隔
            to_remove.append(para)
    # 只删除连续空白段落中多余的（保留单个空行）
    prev_empty = False
    for para in to_remove:
        if prev_empty:
            para._element.getparent().remove(para._element)
        prev_empty = True


def fix_format(input_path: str, output_path: str, preset_name: str = "general"):
    """主校正函数"""
    preset = PRESETS.get(preset_name, PRESETS["general"])
    doc = Document(input_path)

    changes = []

    # 1. 页面布局
    apply_page_layout(doc, preset)
    changes.append("页面尺寸与边距已按预设调整")

    # 2. 逐段落处理
    for para in doc.paragraphs:
        title_flag = is_title_paragraph(para)
        clean_paragraph_formatting(para, preset, is_title=title_flag)
        remove_trailing_spaces(para)
        for run in para.runs:
            clean_run_formatting(run, preset, is_title=title_flag)

    changes.append("段落行距、缩进、对齐已校正")
    changes.append("字体与字号已统一")
    changes.append("末尾多余空格已清除")

    # 3. 去除冗余空行
    remove_empty_paragraphs(doc)
    changes.append("连续空白段落已合并")

    # 4. 表格内文字格式统一
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        clean_run_formatting(run, preset, is_title=False)
    changes.append("表格内字体字号已统一")

    doc.save(output_path)
    return changes


def main():
    if len(sys.argv) < 3:
        print("用法: python fix_format.py <输入文件> <输出文件> [预设名称]")
        print(f"可用预设: {', '.join(PRESETS.keys())}")
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2]
    preset_name = sys.argv[3] if len(sys.argv) > 3 else "general"

    if not Path(input_path).exists():
        print(f"文件不存在: {input_path}", file=sys.stderr)
        sys.exit(1)

    changes = fix_format(input_path, output_path, preset_name)
    print(f"✅ 格式校正完成，输出文件: {output_path}")
    print("调整项目:")
    for c in changes:
        print(f"  • {c}")


if __name__ == "__main__":
    main()
