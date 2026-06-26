import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

class ResumeBeautifier:
    def __init__(self):
        # 核心：自动初始化环境，规避平台二进制限制
        self._ensure_templates()

    def _ensure_templates(self):
        """在云端动态创建模板文件夹和基础文档"""
        template_dir = 'templates'
        if not os.path.exists(template_dir):
            os.makedirs(template_dir)
        
        styles = {
            '简约': (0, 0, 0),
            '商务': (44, 62, 80),
            '极客': (0, 51, 102)
        }
        
        for name, color in styles.items():
            path = os.path.join(template_dir, f'{name}.docx')
            if not os.path.exists(path):
                doc = Document()
                title = doc.add_heading(f'{name}风格简历模板', 0)
                run = title.runs[0]
                run.font.color.rgb = RGBColor(*color)
                doc.save(path)

    def _create_docx(self, text, output_path, mode="Standard_Mode"):
        """执行视觉渲染：加粗、变色、布局适配"""
        doc = Document()
        section = doc.sections[0]
        
        # 布局适配
        if mode == "Compact_Mode":
            section.top_margin = Pt(20)
            line_spacing = 1.0
        else:
            section.top_margin = Pt(40)
            line_spacing = 1.25

        # 技术栈高亮正则
        tech_keywords = r'(Java|Spring|Python|Redis|MySQL|Kafka|Docker|K8s|微服务|分布式|架构|云原生)'
        
        for line in text.split('\n'):
            if not line.strip(): continue
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = line_spacing
            
            parts = re.split(tech_keywords, line, flags=re.IGNORECASE)
            for part in parts:
                run = p.add_run(part)
                if re.match(tech_keywords, part, re.IGNORECASE):
                    run.bold = True
                    run.font.color.rgb = RGBColor(0, 80, 158) # 专业深蓝
                    run.font.size = Pt(11)
                else:
                    run.font.size = Pt(10.5)
                    
        doc.save(output_path)
        return output_path