#!/usr/bin/env python3
"""
专利 Markdown 转 Word 文档
位置: skills/patent-agents/agents/patent-converter/convert_patents.py

流程：
1. 搜索目录下的 专利*.md 文件
2. Pandoc 转 md 为 HTML（提取结构化内容）
3. 正则提取标题、各章节内容
4. 替换模板中的占位符 {{ title }}, {{ chapter1 }} ~ {{ chapter7 }}, {{ year }}, {{ month }}, {{ day }}
5. 转换 Mermaid 图表并插入原位置
6. 输出到源文件同目录

用法：
    python convert_patents.py [目录路径]
    
    如果不指定目录，默认处理 /root/workspace/patent/new 目录
"""

import os
import re
import subprocess
import tempfile
import shutil
import sys
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 脚本所在目录
SCRIPT_DIR = Path(__file__).parent.resolve()

# 模板路径（固定）
TEMPLATE_PATH = SCRIPT_DIR / "templates" / "template.docx"

# 默认搜索目录
DEFAULT_SEARCH_DIR = "/root/workspace/patent/new"


def find_patent_files(directory: str) -> list:
    """
    在目录中搜索 专利*.md 文件
    返回文件路径列表
    """
    directory = Path(directory)
    
    if not directory.exists():
        print(f"错误: 目录不存在 - {directory}")
        return []
    
    # 搜索 专利*.md 文件
    patent_files = list(directory.glob("专利*.md"))
    # 也搜索子目录
    patent_files.extend(directory.glob("*/专利*.md"))
    
    # 去重并排序
    patent_files = sorted(set(str(f) for f in patent_files))
    
    return patent_files


def md_to_html(md_content: str) -> str:
    """使用 Pandoc 将 Markdown 转换为 HTML"""
    try:
        result = subprocess.run(
            ['pandoc', '-f', 'markdown', '-t', 'html', '--wrap=none'],
            input=md_content,
            capture_output=True,
            text=True,
            timeout=30
        )
        return result.stdout if result.returncode == 0 else md_content
    except Exception as e:
        print(f"Pandoc 错误: {e}")
        return md_content


def extract_mermaid_blocks(content: str) -> tuple:
    """提取 Mermaid 代码块，返回 (清理后的内容, mermaid块列表)"""
    pattern = r'```mermaid\s*\n(.*?)```'
    matches = list(re.finditer(pattern, content, re.DOTALL))
    
    mermaid_blocks = [m.group(1).strip() for m in matches]
    
    # 替换 Mermaid 块为占位符
    cleaned = re.sub(pattern, '[[MERMAID_IMAGE]]', content, flags=re.DOTALL)
    
    return cleaned, mermaid_blocks


def mermaid_to_png(code: str, output_path: str) -> bool:
    """将 Mermaid 代码转换为 PNG
    
    安全说明：
    - 推荐在非 root 用户下运行
    - 如果必须以 root 运行，mmdc 需要 --no-sandbox 参数
    - 这会降低浏览器沙箱安全性，请在隔离环境中使用
    """
    try:
        with tempfile.NamedTemporaryFile(mode='w', suffix='.mmd', delete=False, encoding='utf-8') as f:
            f.write(code)
            mmd_path = f.name
        
        # 基础命令
        cmd = [
            'mmdc',
            '-i', mmd_path,
            '-o', output_path,
            '-b', 'white',
            '-w', '800',
            '--quiet'
        ]
        
        # 检测是否以 root 用户运行
        is_root = os.geteuid() == 0 if hasattr(os, 'geteuid') else False
        
        if is_root:
            # root 用户需要 --no-sandbox（安全警告：降低沙箱保护）
            # 建议：在 Docker 容器或非 root 用户下运行
            puppeteer_config = {
                "args": ["--no-sandbox", "--disable-setuid-sandbox"]
            }
            with tempfile.NamedTemporaryFile(mode='w', suffix='.json', delete=False, encoding='utf-8') as f:
                import json
                json.dump(puppeteer_config, f)
                config_path = f.name
            cmd.extend(['-p', config_path])
        
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
        
        # 清理临时文件
        os.unlink(mmd_path)
        if is_root:
            os.unlink(config_path)
        
        return result.returncode == 0 and os.path.exists(output_path)
    except Exception as e:
        print(f"Mermaid 转换错误: {e}")
        return False


def parse_patent_sections(md_path: str) -> dict:
    """解析专利 Markdown 文件，返回 {title, chapter1~chapter7, mermaid_blocks}"""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 提取标题
    title_match = re.search(r'^#\s+(.+?)$', content, re.MULTILINE)
    title = title_match.group(1).strip() if title_match else "未知标题"
    
    # 章节标题
    chapter_titles = [
        "一、相关的现有技术及现有技术的缺陷或不足",
        "二、为克服上述缺陷本提案的技术改进点",
        "三、技术改进点的其他替代方案",
        "四、详细的技术方案具体实施例",
        "五、本提案相对现有技术的优点",
        "六、相关附图",
        "七、权利要求书",
    ]
    
    chapters = {}
    all_mermaid_blocks = []
    
    for i, section_title in enumerate(chapter_titles):
        chapter_key = f'chapter{i+1}'
        
        # 查找章节位置
        idx = content.find(section_title)
        if idx == -1:
            chapters[chapter_key] = ""
            continue
        
        # 找到章节结束位置
        start = content.find('\n', idx) + 1
        end = len(content)
        
        for next_title in chapter_titles[i+1:]:
            next_idx = content.find(next_title)
            if next_idx > start:
                end = content.rfind('\n', 0, next_idx)
                if end == -1:
                    end = next_idx
                break
        
        chapter_content = content[start:end].strip()
        
        # 提取 Mermaid 块
        chapter_content, mermaid_blocks = extract_mermaid_blocks(chapter_content)
        all_mermaid_blocks.extend(mermaid_blocks)
        
        chapters[chapter_key] = chapter_content
    
    return {
        'title': title,
        'chapters': chapters,
        'mermaid_blocks': all_mermaid_blocks
    }


def add_formatted_text(cell, text: str):
    """将文本添加到单元格，处理 Markdown 格式"""
    cell._element.clear_content()
    
    paragraphs = text.split('\n\n')
    
    for para_text in paragraphs:
        if not para_text.strip():
            continue
        
        # 处理标题（从高到低匹配）
        if para_text.startswith('#### '):
            p = cell.add_paragraph()
            run = p.add_run(para_text[5:])
            run.bold = True
        elif para_text.startswith('### '):
            p = cell.add_paragraph()
            run = p.add_run(para_text[4:])
            run.bold = True
        elif para_text.startswith('## '):
            p = cell.add_paragraph()
            run = p.add_run(para_text[3:])
            run.bold = True
        elif para_text.startswith('# '):
            p = cell.add_paragraph()
            run = p.add_run(para_text[2:])
            run.bold = True
        # 处理表格
        elif para_text.startswith('|') and '\n|' in para_text:
            add_table(cell, para_text)
        # 处理无序列表
        elif para_text.startswith('- ') or '\n- ' in para_text:
            add_list_items(cell, para_text, ordered=False)
        # 处理有序列表
        elif re.match(r'^\d+\.\s', para_text) or re.search(r'\n\d+\.\s', para_text):
            add_list_items(cell, para_text, ordered=True)
        else:
            p = cell.add_paragraph()
            add_run_with_format(p, para_text)


def add_run_with_format(paragraph, text: str):
    """添加带格式的文本到段落，处理 **加粗** 和 `代码`"""
    # 处理 **加粗** 格式
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # 加粗文本
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # 普通文本
            paragraph.add_run(part)


def add_list_items(cell, text: str, ordered: bool = False):
    """添加列表项"""
    lines = text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # 去除列表标记
        if line.startswith('- '):
            content = line[2:]
        elif ordered:
            match = re.match(r'^\d+\.\s+(.+)$', line)
            content = match.group(1) if match else line
        else:
            content = line
        
        # 添加带项目符号前缀的段落
        p = cell.add_paragraph()
        prefix = '• ' if not ordered else f'{lines.index(line)+1}. '
        add_run_with_format(p, prefix + content)


def add_table(cell, table_text: str):
    """将 Markdown 表格添加到单元格"""
    lines = [l.strip() for l in table_text.split('\n') if l.strip() and not re.match(r'^\|[-\s|]+\|$', l.strip())]
    
    if len(lines) < 2:
        return
    
    rows_data = []
    for line in lines:
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if cells:
            # 去除 Markdown 格式标记
            cells = [re.sub(r'\*\*([^*]+)\*\*', r'\1', c) for c in cells]
            rows_data.append(cells)
    
    if not rows_data:
        return
    
    num_cols = len(rows_data[0])
    table = cell.add_table(rows=len(rows_data), cols=num_cols)
    
    for i, row_data in enumerate(rows_data):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                table.rows[i].cells[j].text = cell_text


def convert_patent(md_path: str) -> str:
    """
    转换单个专利文件
    返回输出的 docx 文件路径
    """
    md_path = Path(md_path)
    print(f"\n处理: {md_path.name}")
    
    # 解析专利文件
    data = parse_patent_sections(str(md_path))
    print(f"  标题: {data['title'][:50]}...")
    print(f"  Mermaid 图表: {len(data['mermaid_blocks'])} 个")
    
    # 加载模板
    doc = Document(str(TEMPLATE_PATH))
    table = doc.tables[0]
    
    # 替换日期占位符
    today = datetime.now()
    for para in doc.paragraphs:
        for run in para.runs:
            run.text = run.text.replace('{{ year }}', str(today.year))
            run.text = run.text.replace('{{ month }}', str(today.month).zfill(2))
            run.text = run.text.replace('{{ day }}', str(today.day).zfill(2))
    
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.text = run.text.replace('{{ year }}', str(today.year))
                    run.text = run.text.replace('{{ month }}', str(today.month).zfill(2))
                    run.text = run.text.replace('{{ day }}', str(today.day).zfill(2))
    
    # 替换标题 (行0, 列1)
    table.rows[0].cells[1].text = data['title']
    
    # 替换各章节 (行7-13, 列1)
    chapter_rows = [7, 8, 9, 10, 11, 12, 13]
    
    for i, row_idx in enumerate(chapter_rows):
        chapter_key = f'chapter{i+1}'
        chapter_content = data['chapters'].get(chapter_key, '')
        
        if not chapter_content:
            table.rows[row_idx].cells[1].text = ""
            continue
        
        # 转换 Mermaid 图表
        if '[[MERMAID_IMAGE]]' in chapter_content:
            print(f"  章节 {i+1} 有图表，开始转换...")
            
            temp_dir = tempfile.mkdtemp()
            parts = chapter_content.split('[[MERMAID_IMAGE]]')
            img_idx = 0
            
            cell = table.rows[row_idx].cells[1]
            cell._element.clear_content()
            
            for part_idx, part in enumerate(parts):
                if part.strip():
                    for para_text in part.strip().split('\n\n'):
                        if para_text.strip():
                            if para_text.startswith('#### '):
                                p = cell.add_paragraph()
                                run = p.add_run(para_text[5:])
                                run.bold = True
                            elif para_text.startswith('### '):
                                p = cell.add_paragraph()
                                add_run_with_format(p, para_text[4:])
                            elif para_text.startswith('## '):
                                p = cell.add_paragraph()
                                add_run_with_format(p, para_text[3:])
                            elif para_text.startswith('|') and '\n|' in para_text:
                                add_table(cell, para_text)
                            elif para_text.startswith('- ') or '\n- ' in para_text:
                                add_list_items(cell, para_text, ordered=False)
                            elif re.match(r'^\d+\.\s', para_text) or re.search(r'\n\d+\.\s', para_text):
                                add_list_items(cell, para_text, ordered=True)
                            else:
                                p = cell.add_paragraph()
                                add_run_with_format(p, para_text)
                
                if part_idx < len(parts) - 1 and img_idx < len(data['mermaid_blocks']):
                    mermaid_code = data['mermaid_blocks'][img_idx]
                    png_path = os.path.join(temp_dir, f"diagram_{img_idx}.png")
                    
                    if mermaid_to_png(mermaid_code, png_path):
                        try:
                            p = cell.add_paragraph()
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = p.add_run()
                            run.add_picture(png_path, width=Inches(5.5))
                            print(f"    图 {img_idx+1} ✓")
                        except Exception as e:
                            print(f"    图 {img_idx+1} ✗ ({e})")
                    else:
                        print(f"    图 {img_idx+1} ✗ (转换失败)")
                    
                    img_idx += 1
            
            shutil.rmtree(temp_dir, ignore_errors=True)
        else:
            add_formatted_text(table.rows[row_idx].cells[1], chapter_content)
            print(f"  章节 {i+1} ✓")
    
    # 输出到源文件同目录
    output_path = md_path.with_suffix('.docx')
    doc.save(str(output_path))
    print(f"  输出: {output_path.name}")
    
    return str(output_path)


def main():
    print("=" * 60)
    print("专利 Markdown 转 Word 文档")
    print("=" * 60)
    
    # 检查模板
    if not TEMPLATE_PATH.exists():
        print(f"\n错误: 模板文件不存在 - {TEMPLATE_PATH}")
        sys.exit(1)
    
    print(f"\n模板: {TEMPLATE_PATH}")
    
    # 获取搜索目录
    if len(sys.argv) > 1:
        search_dir = sys.argv[1]
    else:
        search_dir = DEFAULT_SEARCH_DIR
    
    print(f"搜索目录: {search_dir}")
    
    # 检查工具
    pandoc_ok = subprocess.run(['which', 'pandoc'], capture_output=True).returncode == 0
    mmdc_ok = subprocess.run(['which', 'mmdc'], capture_output=True).returncode == 0
    
    if not pandoc_ok or not mmdc_ok:
        print("\n错误: 缺少必要工具")
        if not pandoc_ok:
            print("  - Pandoc 未安装")
        if not mmdc_ok:
            print("  - Mermaid CLI 未安装")
        sys.exit(1)
    
    # 查找专利文件
    patent_files = find_patent_files(search_dir)
    
    if not patent_files:
        print(f"\n未找到专利文件 (专利*.md)")
        sys.exit(0)
    
    print(f"\n找到 {len(patent_files)} 个专利文件:\n")
    for f in patent_files:
        print(f"  - {Path(f).name}")
    
    # 转换
    print(f"\n开始转换...\n")
    
    success_count = 0
    for md_path in patent_files:
        try:
            convert_patent(md_path)
            success_count += 1
        except Exception as e:
            print(f"  错误: {e}")
    
    print("\n" + "=" * 60)
    print(f"转换完成！成功: {success_count}/{len(patent_files)}")
    print("=" * 60)


if __name__ == "__main__":
    main()