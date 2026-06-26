"""
付款截图 Word 排版工具
功能：
  - 批量选取图片文件夹
  - OCR 自动识别最终实付金额（支持手动修正）
  - 生成 Word 文档，每页 6 张图（2列×3行），图片下方标注金额
  - 每 3 页（18张）插入小计行
  - 文档末尾插入总计
"""

from __future__ import annotations

import warnings
# 过滤 PyTorch DataLoader 的 pin_memory 无用警告（无 GPU 时正常触发）
warnings.filterwarnings('ignore', message=".*pin_memory.*", category=UserWarning)

import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import os
import re
import threading
from pathlib import Path
from PIL import Image, ImageTk
import io

# ─────────────────────────────────────────────
# OCR 模块
# ─────────────────────────────────────────────
_ocr_reader = None


def get_ocr_reader():
    """延迟初始化 easyocr（首次调用会下载模型，约300MB）"""
    global _ocr_reader
    if _ocr_reader is None:
        import easyocr
        _ocr_reader = easyocr.Reader(['ch_sim', 'en'], gpu=False, verbose=False)
    return _ocr_reader


def _extract_amount_from_text(text: str) -> float:
    """
    从 OCR 文本中提取最终实付金额。
    策略优先级：
      1. 显式"实付/应付/合计/总计"关键词（贪心找最近邻货币符号+金额）
      2. 订单金额 - 立减优惠（支付宝账单详情页）
      3. 负金额（账单格式，含中文 一 作为负号）
      3.5. 先用后付场景（确认收货后再付款）
      4. 正金额最大值兜底（含货币符号格式）
    """
    # ── 优先级1：显式"实付/应付/合计/总计" ──────────────────────────────
    for label in ["实付", "应付", "合计", "总计", "实际支付"]:
        label_pos = text.find(label)
        if label_pos < 0:
            continue
        # 在关键词后200字符内贪心匹配货币符号+金额（跳过 *0 等OCR干扰）
        search_scope = text[label_pos:label_pos + 200]
        # 贪心：匹配到最后一个[*￥¥半●]（实付金额通常在关键词附近最后出现）
        greedy_match = re.search(r'[*￥¥半●][\s]*(\d+\.?\d*)', search_scope)
        if greedy_match:
            try:
                v = float(greedy_match.group(1))
                if 0.01 <= v <= 99999:
                    return v
            except ValueError:
                pass
        # 备选：关键词后直接跟数字（无货币符号格式）
        direct_match = re.search(r'\s+(\d{1,5}\.\d{1,2})', search_scope)
        if direct_match:
            try:
                v = float(direct_match.group(1))
                if 0.01 <= v <= 99999:
                    return v
            except ValueError:
                pass

    # ── 优先级2：订单金额 - 立减优惠（支付宝账单详情页）────────────────
    order_match = re.search(r"订单金额\s*[：:￥¥半]?\s*(\d+\.\d{1,2})", text)
    discount_match = re.search(r"(?:到店)?支付立减\s*[：:￥¥半-]?\s*(\d+\.\d{1,2})", text)
    if order_match:
        order_amount = float(order_match.group(1))
        if discount_match:
            paid = order_amount - float(discount_match.group(1))
            if 0.01 <= paid <= 99999:
                return paid
        if 0.01 <= order_amount <= 9999:
            return order_amount

    # ── 优先级3：负金额（账单格式，含中文 一 作为负号）─────────────────
    # 取文本中第一个负金额（通常在开头 = 实际交易；后面的是立减），返回绝对值
    neg_positions = [(m.start(), float(m.group(1)))
                     for m in re.finditer(r'[-一](\d+\.\d{1,2})', text)
                     if 0.01 <= float(m.group(1)) <= 9999]
    if neg_positions:
        return abs(neg_positions[0][1])

    # ── 优先级3.5：先用后付场景（金额待确认）───────────────────────────
    # 检测"先用后付/确认收货后再付款"等关键词
    if re.search(r'先用后付|确认收货后再付款|付款金额待确认', text):
        # 尝试找订单金额或商品总价作为参考
        order_match = re.search(r'订单金额[：:]\s*[￥¥半]?\s*(\d+\.\d{1,2})', text)
        if order_match:
            v = float(order_match.group(1))
            if 0.01 <= v <= 9999:
                return v  # 返回订单金额供参考
        # 如果找不到订单金额，返回0但后续用户需手动确认
        return 0.0

    # ── 优先级4：所有金额，取最大正值（兜底）────────────────────────────
    # 1. 标准金额（含货币符号的，如半24.21、¥12.34）
    currency_amounts = re.findall(r'[￥¥半●][\s]*(\d+\.\d{1,2})', text)
    # 2. 普通金额（数字前后有边界）
    plain_amounts = re.findall(r'(?:^|[^\d￥¥半●])(\d{1,5}\.\d{1,2})(?=\s|$|[^\d￥¥半●])', text)
    all_amounts = currency_amounts + plain_amounts
    candidates = [float(a) for a in all_amounts if 0.01 <= float(a) <= 9999]
    return max(candidates) if candidates else 0.0

def extract_amount_from_image(image_path: str) -> tuple[str, float]:
    """
    从图片中 OCR 识别最终实付金额。
    返回 (原始识别文本摘要, 金额float)
    """
    import numpy as np
    # Windows 下 OpenCV imread 不支持含中文/特殊字符的路径
    # 改用 Pillow 读取图片后转为 numpy 数组传给 EasyOCR，彻底绕过此限制
    with Image.open(image_path) as pil_img:
        # 转为 RGB（去掉 alpha 通道，EasyOCR 只接受 RGB/灰度）
        img_array = np.array(pil_img.convert('RGB'))

    reader = get_ocr_reader()
    results = reader.readtext(img_array, detail=0)
    text = ' '.join(results)

    amount = _extract_amount_from_text(text)
    return text[:80], amount


def add_page_number_footer(section):
    """为文档节添加页脚居中页码：第 X 页"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    footer = section.footer
    footer.is_linked_to_previous = False

    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.clear()

    # "第 "
    r1 = para.add_run('第 ')

    # { PAGE } 域
    run = para.add_run()
    # begin
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld_begin)
    # instrText
    instr = OxmlElement('w:instrText')
    instr.text = 'PAGE'
    run._r.append(instr)
    # end
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_end)

    # " 页"
    para.add_run(' 页')


# ─────────────────────────────────────────────
# Word 生成模块
# ─────────────────────────────────────────────
def generate_word(image_amounts: list[tuple[str, float]], output_path: str):
    """
    image_amounts: [(image_path, amount), ...]
    排版规则：
      - A4 纵向，页边距 1.5cm
      - 每页 6 张图（2列×3行），图片下方居中显示金额
      - 每 18 张图（3页）后插入小计行（蓝色加粗）
      - 文档末尾插入总计（红色加粗）
    """
    from docx import Document
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import copy

    doc = Document()

    # 设置页面边距
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # 添加页脚页码
    add_page_number_footer(section)

    # 列宽计算：(21 - 3) / 2 ≈ 9cm 每列
    col_w = Cm(9.0)
    img_h = Cm(7.0)  # 图片高度

    total = 0.0
    subtotal = 0.0
    subtotal_start = 1
    n = len(image_amounts)

    # 每页6张 = 3行×2列，用一个大表存所有图
    # 策略：每6张图创建一个3行×2列的表格
    page_idx = 0  # 已处理完整页数

    i = 0
    while i < n:
        # 一批最多6张
        batch = image_amounts[i:i+6]
        batch_size = len(batch)

        # 创建 3行×2列 表格
        tbl = doc.add_table(rows=3, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.style = 'Table Grid'

        # 设置列宽
        for col in tbl.columns:
            for cell in col.cells:
                cell.width = col_w

        # 填充图片和金额
        for j, (img_path, amount) in enumerate(batch):
            row_idx = j // 2
            col_idx = j % 2
            cell = tbl.cell(row_idx, col_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # 清空默认段落
            cell.paragraphs[0].clear()

            # 插入图片
            img_para = cell.paragraphs[0]
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                # 计算图片尺寸（保持比例）
                with Image.open(img_path) as pil_img:
                    w, h = pil_img.size
                    ratio = w / h if h > 0 else 1
                    # 限制在单元格内
                    max_w = Cm(8.5)
                    max_h = img_h
                    if ratio >= 1:
                        fit_w = min(max_w, max_h * ratio)
                        fit_h = fit_w / ratio
                    else:
                        fit_h = min(max_h, max_w / ratio)
                        fit_w = fit_h * ratio
                run = img_para.add_run()
                run.add_picture(img_path, width=fit_w, height=fit_h)
            except Exception as e:
                img_para.add_run(f'[图片加载失败: {os.path.basename(img_path)}]')

            # 插入金额行
            amt_para = cell.add_paragraph()
            amt_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            amt_run = amt_para.add_run(f'¥{amount:.2f}')
            amt_run.bold = True
            amt_run.font.size = Pt(12)
            amt_run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

            subtotal += amount
            total += amount

        # 空余格填充（不足6张时）
        filled = batch_size
        for j in range(filled, 6):
            row_idx = j // 2
            col_idx = j % 2
            cell = tbl.cell(row_idx, col_idx)
            cell.paragraphs[0].clear()

        page_idx += 1
        i += 6

        # 每3页（18张）后插入小计
        if page_idx % 3 == 0:
            end_idx = i  # 已处理总数
            subtotal_para = doc.add_paragraph()
            subtotal_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r = subtotal_para.add_run(
                f'◆ 第 {subtotal_start} ~ {end_idx} 张小计：¥{subtotal:.2f}'
            )
            r.bold = True
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(0x00, 0x45, 0x99)
            subtotal = 0.0
            subtotal_start = end_idx + 1
            # 分页符（让下一组从新页开始）
            if i < n:
                doc.add_page_break()

    # 若有剩余未计入小计的部分（不足18张）
    if subtotal > 0:
        subtotal_para = doc.add_paragraph()
        subtotal_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = subtotal_para.add_run(
            f'◆ 第 {subtotal_start} ~ {n} 张小计：¥{subtotal:.2f}'
        )
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0x00, 0x45, 0x99)

    # 总计
    doc.add_paragraph()
    total_para = doc.add_paragraph()
    total_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = total_para.add_run(f'★ 总计金额：¥{total:.2f}')
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    doc.save(output_path)
    return total


# ─────────────────────────────────────────────
# GUI 主界面
# ─────────────────────────────────────────────
SUPPORTED_EXTS = {'.jpg', '.jpeg', '.png', '.bmp', '.gif', '.webp', '.tiff', '.tif'}


class ReceiptWordApp:
    def __init__(self, root: tk.Tk):
        self.root = root
        self.root.title('付款截图 Word 排版工具')
        self.root.geometry('920x680')
        self.root.configure(bg='#f5f5f5')

        self.image_paths: list[str] = []
        self.amounts: list[tk.StringVar] = []
        self.thumb_labels: list[tk.Label] = []
        self.ocr_status: list[tk.StringVar] = []

        self._build_ui()

    # ── UI 构建 ──────────────────────────────
    def _build_ui(self):
        # 顶部工具栏
        toolbar = tk.Frame(self.root, bg='#2c3e50', height=50)
        toolbar.pack(fill='x')
        toolbar.pack_propagate(False)

        tk.Label(toolbar, text='📋 付款截图 Word 排版工具',
                 bg='#2c3e50', fg='white',
                 font=('微软雅黑', 13, 'bold')).pack(side='left', padx=16, pady=10)

        btn_style = {'bg': '#3498db', 'fg': 'white', 'relief': 'flat',
                     'font': ('微软雅黑', 10), 'padx': 12, 'pady': 4, 'cursor': 'hand2'}

        tk.Button(toolbar, text='📁 选择图片文件夹', command=self._select_folder,
                  **btn_style).pack(side='left', padx=8, pady=10)
        tk.Button(toolbar, text='➕ 追加图片', command=self._append_files,
                  **btn_style).pack(side='left', padx=4, pady=10)
        tk.Button(toolbar, text='🔍 识别所有金额', command=self._ocr_all,
                  bg='#27ae60', fg='white', relief='flat',
                  font=('微软雅黑', 10), padx=12, pady=4, cursor='hand2').pack(side='left', padx=8, pady=10)
        tk.Button(toolbar, text='📄 生成 Word', command=self._generate,
                  bg='#e67e22', fg='white', relief='flat',
                  font=('微软雅黑', 10), padx=12, pady=4, cursor='hand2').pack(side='left', padx=4, pady=10)
        tk.Button(toolbar, text='🗑 清空', command=self._clear,
                  bg='#e74c3c', fg='white', relief='flat',
                  font=('微软雅黑', 10), padx=12, pady=4, cursor='hand2').pack(side='right', padx=12, pady=10)

        # 主内容区
        main = tk.Frame(self.root, bg='#f5f5f5')
        main.pack(fill='both', expand=True, padx=10, pady=6)

        # 状态栏
        self.status_var = tk.StringVar(value='请先选择图片文件夹')
        status_bar = tk.Label(self.root, textvariable=self.status_var,
                              bg='#ecf0f1', fg='#555', anchor='w',
                              font=('微软雅黑', 9), padx=10)
        status_bar.pack(fill='x', side='bottom')

        # 进度条
        self.progress = ttk.Progressbar(self.root, mode='determinate')
        self.progress.pack(fill='x', side='bottom', padx=0)

        # 滚动区域
        canvas_frame = tk.Frame(main, bg='#f5f5f5')
        canvas_frame.pack(fill='both', expand=True)

        self.canvas = tk.Canvas(canvas_frame, bg='#f5f5f5', highlightthickness=0)
        scrollbar = ttk.Scrollbar(canvas_frame, orient='vertical', command=self.canvas.yview)
        self.canvas.configure(yscrollcommand=scrollbar.set)

        scrollbar.pack(side='right', fill='y')
        self.canvas.pack(side='left', fill='both', expand=True)

        self.list_frame = tk.Frame(self.canvas, bg='#f5f5f5')
        self.canvas_window = self.canvas.create_window((0, 0), window=self.list_frame, anchor='nw')

        self.list_frame.bind('<Configure>', self._on_frame_configure)
        self.canvas.bind('<Configure>', self._on_canvas_configure)
        self.canvas.bind('<MouseWheel>', self._on_mousewheel)

        # 表头
        self._build_header()

    def _build_header(self):
        header = tk.Frame(self.list_frame, bg='#bdc3c7')
        header.pack(fill='x', pady=(0, 2))
        tk.Label(header, text='序号', width=4, bg='#bdc3c7', font=('微软雅黑', 9, 'bold')).pack(side='left', padx=4)
        tk.Label(header, text='缩略图', width=10, bg='#bdc3c7', font=('微软雅黑', 9, 'bold')).pack(side='left', padx=4)
        tk.Label(header, text='文件名', width=28, bg='#bdc3c7', font=('微软雅黑', 9, 'bold'), anchor='w').pack(side='left', padx=4)
        tk.Label(header, text='OCR识别结果(摘要)', width=24, bg='#bdc3c7', font=('微软雅黑', 9, 'bold'), anchor='w').pack(side='left', padx=4)
        tk.Label(header, text='金额(可修改)', width=14, bg='#bdc3c7', font=('微软雅黑', 9, 'bold')).pack(side='left', padx=4)
        tk.Label(header, text='操作', width=8, bg='#bdc3c7', font=('微软雅黑', 9, 'bold')).pack(side='left', padx=4)

    # ── 事件处理 ──────────────────────────────
    def _on_frame_configure(self, event):
        self.canvas.configure(scrollregion=self.canvas.bbox('all'))

    def _on_canvas_configure(self, event):
        self.canvas.itemconfig(self.canvas_window, width=event.width)

    def _on_mousewheel(self, event):
        self.canvas.yview_scroll(int(-1 * (event.delta / 120)), 'units')

    def _select_folder(self):
        folder = filedialog.askdirectory(title='选择包含图片的文件夹')
        if not folder:
            return
        files = sorted([
            os.path.join(folder, f) for f in os.listdir(folder)
            if Path(f).suffix.lower() in SUPPORTED_EXTS
        ])
        if not files:
            messagebox.showwarning('提示', '所选文件夹中没有支持的图片文件。')
            return
        self._clear()
        self._add_images(files)

    def _append_files(self):
        files = filedialog.askopenfilenames(
            title='选择图片文件',
            filetypes=[('图片文件', '*.jpg *.jpeg *.png *.bmp *.gif *.webp *.tiff *.tif'), ('所有文件', '*.*')]
        )
        if files:
            self._add_images(list(files))

    def _add_images(self, files: list[str]):
        start_idx = len(self.image_paths)
        for fp in files:
            self.image_paths.append(fp)
            amt_var = tk.StringVar(value='0.00')
            self.amounts.append(amt_var)
            ocr_var = tk.StringVar(value='未识别')
            self.ocr_status.append(ocr_var)
            self._add_row(len(self.image_paths) - 1, fp, amt_var, ocr_var)
        self._update_status()

    def _add_row(self, idx: int, img_path: str, amt_var: tk.StringVar, ocr_var: tk.StringVar):
        row_bg = '#ffffff' if idx % 2 == 0 else '#f9f9f9'
        row = tk.Frame(self.list_frame, bg=row_bg, pady=2)
        row.pack(fill='x', padx=2)

        # 序号
        tk.Label(row, text=str(idx + 1), width=4, bg=row_bg, font=('微软雅黑', 9)).pack(side='left', padx=4)

        # 缩略图（点击可查看大图）
        thumb_label = tk.Label(row, bg=row_bg, width=80, height=60, cursor='hand2')
        thumb_label.pack(side='left', padx=4)
        self._load_thumbnail(img_path, thumb_label)
        thumb_label.bind('<Button-1>', lambda e, path=img_path: self._show_image_preview(path))

        # 文件名
        fname = os.path.basename(img_path)
        fname_show = fname[:30] + '...' if len(fname) > 30 else fname
        tk.Label(row, text=fname_show, width=28, bg=row_bg,
                 font=('微软雅黑', 9), anchor='w').pack(side='left', padx=4)

        # OCR结果
        ocr_label = tk.Label(row, textvariable=ocr_var, width=24, bg=row_bg,
                             font=('微软雅黑', 8), anchor='w', fg='#666', wraplength=190)
        ocr_label.pack(side='left', padx=4)

        # 金额输入框
        amt_frame = tk.Frame(row, bg=row_bg)
        amt_frame.pack(side='left', padx=4)
        tk.Label(amt_frame, text='¥', bg=row_bg, font=('微软雅黑', 10, 'bold'), fg='#c0392b').pack(side='left')
        entry = tk.Entry(amt_frame, textvariable=amt_var, width=10,
                         font=('微软雅黑', 10), justify='right')
        entry.pack(side='left')

        # 单张识别按钮
        tk.Button(row, text='识别', bg='#95a5a6', fg='white', relief='flat',
                  font=('微软雅黑', 8), padx=6, cursor='hand2',
                  command=lambda i=idx: self._ocr_single(i)).pack(side='left', padx=4)

        self.thumb_labels.append(thumb_label)

    def _load_thumbnail(self, img_path: str, label: tk.Label):
        try:
            with Image.open(img_path) as img:
                img.thumbnail((80, 60), Image.LANCZOS)
                photo = ImageTk.PhotoImage(img)
                label.configure(image=photo)
                label.image = photo  # 防止GC
        except Exception:
            label.configure(text='❌', fg='red')

    def _show_image_preview(self, img_path: str):
        """点击缩略图，弹出原图预览窗口"""
        preview_win = tk.Toplevel(self.root)
        fname = os.path.basename(img_path)
        preview_win.title(f'📷 {fname}')
        preview_win.configure(bg='#2c3e50')

        # 取屏幕尺寸，让窗口不超出屏幕
        screen_w = self.root.winfo_screenwidth()
        screen_h = self.root.winfo_screenheight()
        max_w = int(screen_w * 0.85)
        max_h = int(screen_h * 0.85)

        # 加载原图并按比例缩放
        try:
            with Image.open(img_path) as pil_img:
                orig_w, orig_h = pil_img.size
                ratio = min(max_w / orig_w, max_h / orig_h, 1.0)
                new_w = int(orig_w * ratio)
                new_h = int(orig_h * ratio)
                pil_img_preview = pil_img.resize((new_w, new_h), Image.LANCZOS)
                photo = ImageTk.PhotoImage(pil_img_preview)

            img_label = tk.Label(preview_win, image=photo, bg='#2c3e50', cursor='hand2')
            img_label.image = photo  # 防止GC
            img_label.pack(padx=10, pady=(10, 4))

            # 右下角文件名标签
            info = tk.Label(preview_win, text=fname,
                           bg='#2c3e50', fg='#bdc3c7',
                           font=('微软雅黑', 9))
            info.pack(pady=(0, 8))

            # 关闭按钮
            tk.Button(preview_win, text='✕ 关闭',
                     bg='#e74c3c', fg='white', relief='flat',
                     font=('微软雅黑', 10), cursor='hand2',
                     command=preview_win.destroy).pack(pady=(0, 10))

            # 窗口尺寸跟随图片
            win_w = min(new_w + 20, max_w)
            win_h = new_h + 80
            preview_win.geometry(f'{win_w}x{win_h}')
        except Exception as e:
            tk.Label(preview_win, text=f'图片加载失败：{e}',
                    bg='#2c3e50', fg='white',
                    font=('微软雅黑', 12)).pack(padx=20, pady=40)

        # 居中显示
        preview_win.update_idletasks()
        x = (screen_w - preview_win.winfo_width()) // 2
        y = (screen_h - preview_win.winfo_height()) // 2
        preview_win.geometry(f'+{x}+{y}')

        preview_win.transient(self.root)
        preview_win.attributes('-topmost', True)
        preview_win.bind('<Escape>', lambda e: preview_win.destroy())
        preview_win.focus_force()

    def _ocr_single(self, idx: int):
        """单张OCR识别"""
        def task():
            self.root.after(0, lambda: self.ocr_status[idx].set('识别中…'))
            try:
                text, amount = extract_amount_from_image(self.image_paths[idx])
                self.root.after(0, lambda: self.amounts[idx].set(f'{amount:.2f}'))
                short = text[:40].replace('\n', ' ') if text else '无文字'
                self.root.after(0, lambda: self.ocr_status[idx].set(short))
            except Exception as e:
                self.root.after(0, lambda: self.ocr_status[idx].set(f'错误: {e}'))
        threading.Thread(target=task, daemon=True).start()

    def _ocr_all(self):
        """批量OCR识别（后台线程）"""
        if not self.image_paths:
            messagebox.showwarning('提示', '请先添加图片。')
            return

        total = len(self.image_paths)

        def task():
            self.root.after(0, lambda: self.status_var.set('正在初始化 OCR 模型，首次使用需下载约300MB模型，请耐心等待…'))
            self.root.after(0, lambda: self.progress.config(maximum=total))
            for i, img_path in enumerate(self.image_paths):
                self.root.after(0, lambda i=i: self.ocr_status[i].set('识别中…'))
                self.root.after(0, lambda i=i: self.status_var.set(f'正在识别第 {i+1}/{total} 张…'))
                try:
                    text, amount = extract_amount_from_image(img_path)
                    self.root.after(0, lambda i=i, amount=amount: (self.amounts[i].set(f'{amount:.2f}'), None))
                    short = text[:40].replace('\n', ' ') if text else '无文字'
                    self.root.after(0, lambda i=i, short=short: self.ocr_status[i].set(short))
                except Exception as e:
                    self.root.after(0, lambda i=i, e=e: self.ocr_status[i].set(f'错误:{str(e)[:30]}'))
                self.root.after(0, lambda i=i: self.progress.config(value=i + 1))
                self.root.after(0, lambda: self.root.update_idletasks())
            self.root.after(0, lambda: self.status_var.set(f'识别完成，共 {total} 张。请检查金额后点击【生成 Word】。'))
            self.root.after(0, lambda: self.progress.config(value=0))

        threading.Thread(target=task, daemon=True).start()

    def _generate(self):
        """生成 Word 文档"""
        if not self.image_paths:
            messagebox.showwarning('提示', '请先添加图片。')
            return

        # 收集金额
        image_amounts = []
        for i, fp in enumerate(self.image_paths):
            try:
                amt = float(self.amounts[i].get())
            except ValueError:
                amt = 0.0
            image_amounts.append((fp, amt))

        # 选择保存路径
        output_path = filedialog.asksaveasfilename(
            defaultextension='.docx',
            filetypes=[('Word文档', '*.docx')],
            initialfile='付款记录.docx',
            title='保存 Word 文件'
        )
        if not output_path:
            return

        def task():
            self.root.after(0, lambda: self.status_var.set('正在生成 Word 文档，请稍候…'))
            try:
                total = generate_word(image_amounts, output_path)
                self.root.after(0, lambda: self.status_var.set(f'✅ 生成成功！总计金额：¥{total:.2f}  →  {output_path}'))
                self.root.after(0, lambda: messagebox.showinfo('完成', f'Word 文档已生成！\n\n总计金额：¥{total:.2f}\n\n保存路径：\n{output_path}'))
                self.root.after(0, lambda p=output_path: os.startfile(p))
            except Exception as e:
                self.root.after(0, lambda: self.status_var.set(f'❌ 生成失败：{e}'))
                self.root.after(0, lambda: messagebox.showerror('错误', f'生成失败：{e}'))

        threading.Thread(target=task, daemon=True).start()

    def _clear(self):
        """清空所有数据"""
        self.image_paths.clear()
        self.amounts.clear()
        self.ocr_status.clear()
        self.thumb_labels.clear()
        # 清除列表区域（保留表头）
        for widget in self.list_frame.winfo_children():
            widget.destroy()
        self._build_header()
        self.status_var.set('已清空，请重新选择图片。')

    def _update_status(self):
        n = len(self.image_paths)
        pages = (n + 5) // 6
        self.status_var.set(
            f'已加载 {n} 张图片，共 {pages} 页（每页6张）。'
            f'  点击【识别所有金额】自动提取金额，或手动输入后直接【生成 Word】。'
        )


# ─────────────────────────────────────────────
# 入口
# ─────────────────────────────────────────────
if __name__ == '__main__':
    root = tk.Tk()
    app = ReceiptWordApp(root)
    root.mainloop()
