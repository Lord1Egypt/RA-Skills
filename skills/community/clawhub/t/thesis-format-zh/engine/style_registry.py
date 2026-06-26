# =============================================================================
# Style Registry — 样式注册 + 模板驱动 + 格式验证
# =============================================================================
# 两种模式：
#   1. YAML 驱动：从 spec_config.yaml 注册样式（代码定义格式）
#   2. 模板驱动：从 .docx 模板复制样式（Word 里定义格式，更直观）
# =============================================================================

import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

import yaml
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm, Emu, Inches
from docx.styles.style import BaseStyle


# ── 单位解析 ──────────────────────────────────────────────────

def parse_length(value: Union[str, int, float, None]) -> Optional[int]:
    """将 '12pt' / '2.54cm' / '24pt' 等字符串转为 EMU 整数。
    也支持纯数字（当作 pt 处理），空值返回 None。
    """
    if value is None:
        return None
    if isinstance(value, (int, float)):
        return int(Pt(float(value)))
    value = str(value).strip()
    if not value:
        return None

    m = re.match(r'^([\d.]+)\s*(pt|cm|inch|in|mm)$', value, re.I)
    if m:
        num = float(m.group(1))
        unit = m.group(2).lower()
        if unit == 'pt':
            return Pt(num)
        elif unit in ('cm',):
            return Cm(num)
        elif unit in ('inch', 'in'):
            return Inches(num)
        elif unit == 'mm':
            return int(num * 36000)
    try:
        return Pt(float(value))
    except ValueError:
        return None


ALIGN_MAP: Dict[str, int] = {
    'left': WD_ALIGN_PARAGRAPH.LEFT,
    'center': WD_ALIGN_PARAGRAPH.CENTER,
    'right': WD_ALIGN_PARAGRAPH.RIGHT,
    'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
}


# ═══════════════════════════════════════════════════════════════
# StyleRegistry — 样式注册表
# ═══════════════════════════════════════════════════════════════

class StyleRegistry:
    """管理样式注册与应用。支持 YAML 驱动和模板驱动两种模式。"""

    def __init__(self, doc: Document, config: Dict[str, Any]):
        self.doc = doc
        self.config = config
        self._style_cache: Dict[str, BaseStyle] = {}

    @classmethod
    def from_yaml(cls, doc: Document, yaml_path: Union[str, Path]) -> 'StyleRegistry':
        """模式 1：从 YAML 配置创建"""
        with open(yaml_path, 'r', encoding='utf-8') as f:
            config = yaml.safe_load(f)
        return cls(doc, config)

    @classmethod
    def from_template(cls, doc: Document, template_path: Union[str, Path]) -> 'StyleRegistry':
        """模式 2：从 .docx 模板复制样式（推荐）
        
        从模板 .docx 复制所有自定义样式到目标文档。
        模板里的样式可以在 Word 里直接编辑，不需要改代码。
        """
        tpl = Document(template_path)

        # 从模板复制样式定义
        copied = []
        for tpl_style in tpl.styles:
            if tpl_style.type == WD_STYLE_TYPE.PARAGRAPH and not tpl_style.builtin:
                try:
                    # 检查目标是否已有同名样式
                    target_style = doc.styles[tpl_style.name]
                    # 已存在，更新属性
                except KeyError:
                    # 不存在，创建
                    target_style = doc.styles.add_style(
                        tpl_style.name, WD_STYLE_TYPE.PARAGRAPH
                    )

                # 复制字体属性
                if tpl_style.font.name:
                    target_style.font.name = tpl_style.font.name
                if tpl_style.font.size:
                    target_style.font.size = tpl_style.font.size
                if tpl_style.font.bold is not None:
                    target_style.font.bold = tpl_style.font.bold

                # 复制 eastAsia 字体
                tpl_rPr = tpl_style.element.find(qn('w:rPr'))
                if tpl_rPr is not None:
                    tpl_rFonts = tpl_rPr.find(qn('w:rFonts'))
                    if tpl_rFonts is not None:
                        ea = tpl_rFonts.get(qn('w:eastAsia'))
                        if ea:
                            set_eastasia(target_style, ea)

                # 复制段落属性
                tpl_pf = tpl_style.paragraph_format
                tgt_pf = target_style.paragraph_format
                if tpl_pf.alignment is not None:
                    tgt_pf.alignment = tpl_pf.alignment
                if tpl_pf.line_spacing is not None:
                    tgt_pf.line_spacing = tpl_pf.line_spacing
                if tpl_pf.space_before is not None:
                    tgt_pf.space_before = tpl_pf.space_before
                if tpl_pf.space_after is not None:
                    tgt_pf.space_after = tpl_pf.space_after
                if tpl_pf.first_line_indent is not None:
                    tgt_pf.first_line_indent = tpl_pf.first_line_indent

                # 复制行级间距属性（beforeLines/afterLines，Word 里显示为"0.5行"）
                tpl_pPr = tpl_style.element.find(qn('w:pPr'))
                tgt_pPr = target_style.element.find(qn('w:pPr'))
                if tpl_pPr is not None and tgt_pPr is not None:
                    tpl_spacing = tpl_pPr.find(qn('w:spacing'))
                    if tpl_spacing is not None:
                        tgt_spacing = tgt_pPr.find(qn('w:spacing'))
                        if tgt_spacing is None:
                            tgt_spacing = OxmlElement('w:spacing')
                            tgt_pPr.insert(0, tgt_spacing)
                        for attr in ('beforeLines', 'afterLines'):
                            val = tpl_spacing.get(qn(f'w:{attr}'))
                            if val is not None:
                                tgt_spacing.set(qn(f'w:{attr}'), val)

                copied.append(tpl_style.name)

        # 从模板复制页面设置
        for tpl_section in tpl.sections:
            for section in doc.sections:
                section.page_width = tpl_section.page_width
                section.page_height = tpl_section.page_height
                section.top_margin = tpl_section.top_margin
                section.bottom_margin = tpl_section.bottom_margin
                section.left_margin = tpl_section.left_margin
                section.right_margin = tpl_section.right_margin

        # 从模板复制页眉
        for tpl_section in tpl.sections:
            tpl_header = tpl_section.header
            for section in doc.sections:
                header = section.header
                header.is_linked_to_previous = False
                for p in header.paragraphs:
                    p.clear()
                # 复制页眉内容
                for tpl_p in tpl_header.paragraphs:
                    hp = header.paragraphs[0] if header.paragraphs[0].text == '' else header.add_paragraph()
                    hp.alignment = tpl_p.alignment
                    # 复制段落边框
                    tpl_pPr = tpl_p._element.find(qn('w:pPr'))
                    if tpl_pPr is not None:
                        tpl_pBdr = tpl_pPr.find(qn('w:pBdr'))
                        if tpl_pBdr is not None:
                            pPr = hp._element.get_or_add_pPr()
                            pPr.append(tpl_pBdr)
                    for r in tpl_p.runs:
                        hr = hp.add_run(r.text)
                        hr.font.name = r.font.name
                        if r.font.size:
                            hr.font.size = r.font.size
                        # 复制 eastAsia
                        rPr = r._element.find(qn('w:rPr'))
                        if rPr is not None:
                            rFonts = rPr.find(qn('w:rFonts'))
                            if rFonts is not None:
                                ea = rFonts.get(qn('w:eastAsia'))
                                if ea:
                                    set_eastasia(hr, ea)

        config = {'_template': str(template_path), '_copied_styles': copied}
        return cls(doc, config)

    # ── 页面设置 ──────────────────────────────────────────────

    def apply_page_setup(self):
        """应用页面尺寸和页边距（仅 YAML 模式需要，模板模式已自动复制）"""
        if '_template' in self.config:
            return  # 模板模式：已从模板复制

        page = self.config.get('page', {})
        margins = page.get('margins', {})
        for section in self.doc.sections:
            if page.get('size', 'A4') == 'A4':
                section.page_width = Cm(21.0)
                section.page_height = Cm(29.7)
            for side in ('top', 'bottom', 'left', 'right'):
                val = margins.get(side)
                if val:
                    emu = parse_length(val)
                    if emu:
                        setattr(section, f'{side}_margin', emu)

    # ── 页眉 ──────────────────────────────────────────────────

    def apply_header(self):
        """应用页眉（仅 YAML 模式需要，模板模式已自动复制）"""
        if '_template' in self.config:
            return  # 模板模式：已从模板复制

        header_cfg = self.config.get('header', {})
        if not header_cfg:
            return
        text = header_cfg.get('text', '')
        if not text:
            return

        for section in self.doc.sections:
            header = section.header
            header.is_linked_to_previous = False
            for p in header.paragraphs:
                p.clear()
            p = header.paragraphs[0]
            p.alignment = ALIGN_MAP.get(header_cfg.get('alignment', 'center'),
                                        WD_ALIGN_PARAGRAPH.CENTER)
            border = header_cfg.get('border_bottom')
            if border:
                pPr = p._element.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), border.get('val', 'single'))
                bottom.set(qn('w:sz'), str(border.get('sz', 4)))
                bottom.set(qn('w:space'), str(border.get('space', 1)))
                bottom.set(qn('w:color'), border.get('color', '000000'))
                pBdr.append(bottom)
                pPr.append(pBdr)
            run = p.add_run(text)
            run.font.name = header_cfg.get('font_ascii', '宋体')
            size = parse_length(header_cfg.get('size', '10.5pt'))
            if size:
                run.font.size = size
            cn_font = header_cfg.get('font_eastasia', '宋体')
            set_eastasia(run, cn_font)

    # ── 样式注册（YAML 模式）───────────────────────────────────

    def register_all(self):
        """注册 YAML 中定义的所有样式（仅 YAML 模式）"""
        if '_template' in self.config:
            return self  # 模板模式：样式已从模板复制
        styles_cfg = self.config.get('styles', {})
        for name, spec in styles_cfg.items():
            self.register_style(name, spec)
        return self

    def register_style(self, name: str, spec: Dict[str, Any]) -> BaseStyle:
        """注册单个样式"""
        style_type = spec.get('type', 'paragraph')
        wd_type = WD_STYLE_TYPE.PARAGRAPH if style_type == 'paragraph' else WD_STYLE_TYPE.CHARACTER

        try:
            style = self.doc.styles[name]
        except KeyError:
            style = self.doc.styles.add_style(name, wd_type)

        base_on = spec.get('base_on')
        if base_on:
            try:
                style.base_style = self.doc.styles[base_on]
            except KeyError:
                pass

        font_ascii = spec.get('font_ascii')
        if font_ascii:
            style.font.name = font_ascii

        size = parse_length(spec.get('size'))
        if size:
            style.font.size = size

        bold = spec.get('bold')
        if bold is not None:
            style.font.bold = bold

        alignment = spec.get('alignment')
        if alignment:
            style.paragraph_format.alignment = ALIGN_MAP.get(alignment)

        ls = spec.get('line_spacing')
        if ls is not None:
            if isinstance(ls, str):
                if ls.endswith('_at_least'):
                    val = parse_length(ls.replace('_at_least', ''))
                    if val:
                        pPr = style.element.get_or_add_pPr()
                        spacing = pPr.find(qn('w:spacing'))
                        if spacing is None:
                            spacing = OxmlElement('w:spacing')
                            pPr.append(spacing)
                        spacing.set(qn('w:line'), str(val))
                        spacing.set(qn('w:lineRule'), 'atLeast')
                else:
                    val = parse_length(ls)
                    if val:
                        style.paragraph_format.line_spacing = Pt(val / 12700)
            else:
                style.paragraph_format.line_spacing = float(ls)

        fi = spec.get('first_line_indent')
        if fi is not None:
            val = parse_length(fi)
            if val:
                style.paragraph_format.first_line_indent = val

        for attr in ('space_before', 'space_after'):
            val = spec.get(attr)
            if val is not None:
                emu = parse_length(str(val))
                if emu:
                    setattr(style.paragraph_format, attr, emu)

        cn_font = spec.get('font_eastasia')
        if cn_font:
            set_eastasia(style, cn_font)

        self._style_cache[name] = style
        return style

    def get(self, name: str) -> BaseStyle:
        """获取已注册的样式"""
        if name in self._style_cache:
            return self._style_cache[name]
        return self.doc.styles[name]


# ═══════════════════════════════════════════════════════════════
# StyleApplier — 段落检测 + 样式应用
# ═══════════════════════════════════════════════════════════════

class StyleApplier:
    """扫描全文，检测段落类型并应用对应样式"""

    def __init__(self, doc: Document, registry: StyleRegistry, config: Dict[str, Any]):
        self.doc = doc
        self.registry = registry
        self.config = config
        self._label_map = config.get('label_map', {})
        self._special_rules = config.get('special_paragraphs', [])
    def apply(self):
        """扫描全文，应用样式"""
        in_orig = False
        header_text = self.config.get('header', {}).get('text', '')

        for p in self.doc.paragraphs:
            text = p.text.strip()

            # 跳过旧页眉文本
            if header_text and text == header_text:
                for r in p.runs:
                    r.text = ''
                continue

            if not text:
                continue

            # 边界检测
            appendix_a = self._label_map.get('appendix_a', '附录A')
            appendix_b = self._label_map.get('appendix_b', '附录B')

            if appendix_a in text and '译文' in text:
                p.style = self.registry.get('附录标题')
                self._apply_para_format(p, '附录标题')
                continue

            if appendix_b in text and ('原文' in text or '外文原文' in text):
                pPr = p._element.get_or_add_pPr()
                br = OxmlElement('w:pageBreakBefore')
                pPr.insert(0, br)
                p.style = self.registry.get('附录标题')
                self._apply_para_format(p, '附录标题')
                in_orig = True
                continue

            cn_title = self._label_map.get('cn_title', '')
            en_title = self._label_map.get('en_title', '')
            if cn_title and cn_title in text:
                p.style = self.registry.get('文章标题')
                self._apply_para_format(p, '文章标题')
                continue
            if en_title and en_title in text:
                p.style = self.registry.get('文章标题')
                self._apply_para_format(p, '文章标题')
                continue

            if in_orig:
                special = self._match_special(text)
                if special:
                    p.style = self.registry.get(special['style'])
                    self._apply_special_format(p, special['format'])
                else:
                    p.style = self.registry.get('英文正文')
                    self._apply_para_format(p, '英文正文')
            else:
                # 先检查是否匹配特殊段落规则
                special = self._match_special(text)
                if special:
                    p.style = self.registry.get(special['style'])
                    self._apply_special_format(p, special['format'])
                else:
                    p.style = self.registry.get('中文正文')
                    self._apply_para_format(p, '中文正文')

        return self

    def _apply_para_format(self, p, style_name: str):
        """将样式的首行缩进、行距、段间距直接应用到段落 pPr"""
        style = self.doc.styles[style_name]
        spf = style.paragraph_format
        ppf = p.paragraph_format

        if spf.first_line_indent is not None:
            ppf.first_line_indent = spf.first_line_indent
        if spf.line_spacing is not None:
            ppf.line_spacing = spf.line_spacing
        if spf.space_before is not None:
            ppf.space_before = spf.space_before
        if spf.space_after is not None:
            ppf.space_after = spf.space_after

        # 也复制 beforeLines/afterLines（行级单位，Word 里显示为"0.5行"）
        spPr = style.element.find(qn('w:pPr'))
        if spPr is None:
            return
        s_spacing = spPr.find(qn('w:spacing'))
        if s_spacing is None:
            return
        
        ppPr = p._element.get_or_add_pPr()
        p_spacing = ppPr.find(qn('w:spacing'))
        if p_spacing is None:
            p_spacing = OxmlElement('w:spacing')
            ppPr.insert(0, p_spacing)

        for attr in ('beforeLines', 'afterLines'):
            val = s_spacing.get(qn(f'w:{attr}'))
            if val is not None:
                p_spacing.set(qn(f'w:{attr}'), val)

    def _match_special(self, text: str) -> Optional[Dict]:
        """匹配特殊段落规则，返回第一个命中的规则"""
        for rule in self._special_rules:
            for kw in rule.get('keywords', []):
                if kw in text:
                    return rule
        return None

    def _apply_special_format(self, p, fmt: Dict):
        """应用特殊格式覆盖（first_line_indent/line_spacing/space_before/space_after）"""
        ppf = p.paragraph_format

        if 'first_line_indent' in fmt:
            fi = fmt['first_line_indent']
            if fi is None or fi == '0' or fi == 'null':
                # 显式清除首行缩进：直接删段落 XML 里的 w:firstLine
                ppf.first_line_indent = None
                pPr = p._element.find(qn('w:pPr'))
                if pPr is not None:
                    ind = pPr.find(qn('w:ind'))
                    if ind is not None:
                        for attr in (qn('w:firstLine'), qn('w:firstLineChars')):
                            if attr in ind.attrib:
                                del ind.attrib[attr]
                        if len(ind.attrib) == 0:
                            pPr.remove(ind)
            else:
                val = parse_length(str(fi))
                if val:
                    ppf.first_line_indent = val

        if 'line_spacing' in fmt:
            ls = fmt['line_spacing']
            if ls is None:
                pass  # 不覆盖
            else:
                ppf.line_spacing = float(ls)

        if 'space_before' in fmt:
            sb = fmt['space_before']
            val = parse_length(str(sb))
            if val:
                ppf.space_before = val

        if 'space_after' in fmt:
            sa = fmt['space_after']
            val = parse_length(str(sa))
            if val:
                ppf.space_after = val

        # 也复制 beforeLines/afterLines（行级单位）— 与 _apply_para_format 保持一致
        pPr = p._element.get_or_add_pPr()
        p_spacing = pPr.find(qn('w:spacing'))
        if p_spacing is None:
            p_spacing = OxmlElement('w:spacing')
            pPr.insert(0, p_spacing)
        for attr in ('beforeLines', 'afterLines'):
            val = fmt.get(attr)
            if val is not None:
                p_spacing.set(qn(f'w:{attr}'), str(val))

    def insert_title_paragraphs(self):
        """在附录标题后插入文章标题段落（如果不存在）
        
        关键：倒序处理（先附录B后附录A），避免插入后索引偏移。
        同时确保标题和正文之间有空行。
        """
        cn_title = self._label_map.get('cn_title', '')
        en_title = self._label_map.get('en_title', '')
        appendix_a = self._label_map.get('appendix_a', '附录A')
        appendix_b = self._label_map.get('appendix_b', '附录B')

        # 先收集所有需要插入的位置（倒序扫描）
        insertions = []  # [(idx, title_text, style_name)]
        for i, p in enumerate(self.doc.paragraphs):
            text = p.text.strip()
            if not text:
                continue
            if appendix_b in text and ('原文' in text or '外文原文' in text) and en_title:
                insertions.append((i, en_title, '文章标题'))
            if appendix_a in text and '译文' in text and cn_title:
                insertions.append((i, cn_title, '文章标题'))

        # 倒序插入（从后往前，避免索引偏移）
        insertions.sort(key=lambda x: x[0], reverse=True)
        for idx, title_text, style_name in insertions:
            # 1. 确保标题存在
            title_idx = self._ensure_next_paragraph_is(idx, title_text, style_name)
            # 2. 在附录标题和标题之间插入空行（仅当没有现成空行时）
            next_p = self.doc.paragraphs[idx + 1] if idx + 1 < len(self.doc.paragraphs) else None
            if next_p and next_p.text.strip():
                self._insert_blank_para(self.doc.paragraphs[idx]._element, idx)
                title_idx += 1  # 标题后移了一位
            # 3. 在标题和正文之间插入空行（仅当没有现成空行时）
            self._ensure_blank_after(title_idx)

    def _ensure_next_paragraph_is(self, idx: int, title_text: str, style_name: str) -> int:
        """确保 idx 的下一个段落是标题文本；不存在则插入。返回标题所在索引。"""
        for j in range(idx + 1, min(idx + 4, len(self.doc.paragraphs))):
            p = self.doc.paragraphs[j]
            t = p.text.strip()
            if t == title_text:
                p.style = self.registry.get(style_name)
                return j
            if t and t != title_text:
                break

        target_elem = self.doc.paragraphs[idx]._element
        parent = target_elem.getparent()
        pos = list(parent).index(target_elem)

        new_p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), style_name)
        pPr.append(pStyle)
        new_p.append(pPr)

        r = OxmlElement('w:r')
        t_node = OxmlElement('w:t')
        t_node.text = title_text
        r.append(t_node)
        new_p.append(r)

        parent.insert(pos + 1, new_p)
        return idx + 1  # 标题在附录标签后一个位置

    def _ensure_blank_after(self, title_idx: int):
        """确保标题后有一个空段落"""
        # 检查下一个段落是否为空
        next_idx = title_idx + 1
        if next_idx >= len(self.doc.paragraphs):
            # 文档末尾，插入空段落
            self._insert_blank_para(None, title_idx)
            return
        
        p = self.doc.paragraphs[next_idx]
        if not p.text.strip():
            # 已经是空段落
            return
        
        # 下一个段落有文字，在前面插入空段落
        self._insert_blank_para(p._element, title_idx)

    def _insert_blank_para(self, after_elem, after_idx: int):
        """插入一个空段落"""
        if after_elem is None:
            # 在文档末尾插入
            after_elem = self.doc.paragraphs[-1]._element
        
        parent = after_elem.getparent()
        pos = list(parent).index(after_elem)
        
        new_p = OxmlElement('w:p')
        parent.insert(pos + 1, new_p)


# ═══════════════════════════════════════════════════════════════
# FormatAuditor — 格式验证器（纯 Python diff）
# ═══════════════════════════════════════════════════════════════

class FormatAuditor:
    """对比格式化前后的文档，输出差异报告"""

    def __init__(self, original_path: str, formatted_path: str):
        self.original = Document(original_path)
        self.formatted = Document(formatted_path)

    def audit(self) -> Dict[str, Any]:
        """对比两份文档，返回差异报告"""
        changes = []

        orig_paras = [(p.text.strip(), self._get_para_info(p))
                      for p in self.original.paragraphs]
        fmt_paras = [(p.text.strip(), self._get_para_info(p))
                     for p in self.formatted.paragraphs]

        # 按文本内容匹配段落
        orig_by_text = {}
        for i, (text, info) in enumerate(orig_paras):
            if text:
                orig_by_text[text] = info

        for i, (text, info) in enumerate(fmt_paras):
            if not text:
                continue
            orig_info = orig_by_text.get(text)
            if orig_info and orig_info != info:
                diff = self._diff_info(orig_info, info, text[:30])
                if diff:
                    changes.append(diff)

        return {
            'total_changes': len(changes),
            'changes': changes,
        }

    def _get_para_info(self, p) -> Dict:
        """提取段落格式信息"""
        r = p.runs[0] if p.runs else None
        cn_font = '?'
        if r:
            rPr = r._element.find(qn('w:rPr'))
            if rPr is not None:
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is not None:
                    cn_font = rFonts.get(qn('w:eastAsia')) or '?'

        return {
            'style': p.style.name,
            'font': r.font.name if r else None,
            'cn_font': cn_font,
            'size': r.font.size if r and r.font.size else None,
            'bold': r.font.bold if r else None,
            'alignment': p.paragraph_format.alignment,
            'line_spacing': p.paragraph_format.line_spacing,
            'first_indent': p.paragraph_format.first_line_indent,
        }

    def _diff_info(self, orig: Dict, fmt: Dict, label: str) -> Optional[Dict]:
        """对比两个段落信息，返回差异"""
        diffs = {}
        for key in ('style', 'font', 'cn_font', 'size', 'bold',
                     'alignment', 'line_spacing', 'first_indent'):
            if orig.get(key) != fmt.get(key):
                diffs[key] = {'from': orig.get(key), 'to': fmt.get(key)}
        return {'paragraph': label, 'diffs': diffs} if diffs else None

    def print_report(self):
        """打印差异报告"""
        result = self.audit()
        print(f"\n{'='*60}")
        print(f"📋 格式化审计报告: {result['total_changes']} 处变化")
        print(f"{'='*60}")
        for change in result['changes']:
            print(f"\n  「{change['paragraph']}」")
            for key, vals in change['diffs'].items():
                print(f"    {key}: {vals['from']} → {vals['to']}")
        if not result['changes']:
            print("  ✅ 无差异")
        return result


# ── 辅助 ──────────────────────────────────────────────────────

def set_eastasia(element, font_name: str):
    """为样式或 run 设置中文字体"""
    if hasattr(element, 'element'):
        elem = element.element
    else:
        elem = element._element if hasattr(element, '_element') else element
    rPr = elem.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        elem.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
