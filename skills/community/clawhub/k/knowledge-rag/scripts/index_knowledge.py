#!/usr/bin/env python3
"""
index_knowledge.py — 知识仓库索引构建

扫描 ~/workspace/knowledge/ 下所有子目录中的文档，
按文件类型解析、切片、向量化，统一存入索引。

用法：
  .venv/bin/python3 scripts/index_knowledge.py                    # 增量更新
  .venv/bin/python3 scripts/index_knowledge.py --force            # 全量重建
  .venv/bin/python3 scripts/index_knowledge.py --source bilibili  # 只更新某个来源
"""

import os, sys, json, re, hashlib, sqlite3, concurrent.futures
from pathlib import Path

import numpy as np
import requests

try:
    from PyPDF2 import PdfReader
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# ===== 配置 =====
KNOWLEDGE_DIR = os.path.expanduser("~/workspace/knowledge")
RAG_DIR = os.path.join(KNOWLEDGE_DIR, ".rag_data")
OLLAMA_URL = "http://localhost:11434/api/embeddings"
EMBED_MODEL = "qwen3-embedding:0.6b"
CHUNK_SIZE = 512
MODEL_META_FILE = os.path.join(RAG_DIR, "model_meta.json")
CHUNK_OVERLAP = 64


def _get_config():
    config_file = os.path.join(KNOWLEDGE_DIR, ".knowledge-config.json")
    try:
        with open(config_file, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}

def get_embed_model():
    """从配置文件读取 embed_model，否则用默认值"""
    cfg = _get_config()
    return cfg.get("embed_model", EMBED_MODEL)

def load_model_meta():
    if not os.path.exists(MODEL_META_FILE):
        return None, None
    with open(MODEL_META_FILE, "r", encoding="utf-8") as f:
        meta = json.load(f)
    return meta.get("model"), meta.get("dim")

def save_model_meta(model_name, dim):
    with open(MODEL_META_FILE, "w", encoding="utf-8") as f:
        json.dump({"model": model_name, "dim": dim}, f)

os.makedirs(RAG_DIR, exist_ok=True)

# 来源类型 = 子目录名
SOURCE_MAP = {
    "bilibili": "bilibili",
    "notes": "note",
    "wechat-articles": "wechat",
    "other": "other",
}


def _get_db():
    db_path = os.path.join(RAG_DIR, "chunks.db")
    conn = sqlite3.connect(db_path)
    conn.execute("PRAGMA journal_mode=WAL")
    conn.execute("""CREATE TABLE IF NOT EXISTS chunks (
        id TEXT PRIMARY KEY,
        text TEXT NOT NULL,
        title TEXT NOT NULL DEFAULT '',
        author TEXT NOT NULL DEFAULT '',
        date TEXT NOT NULL DEFAULT '',
        bvid TEXT NOT NULL DEFAULT '',
        source_type TEXT NOT NULL DEFAULT '',
        source_label TEXT NOT NULL DEFAULT '',
        filename TEXT NOT NULL DEFAULT '',
        chunk_index INTEGER NOT NULL DEFAULT 0,
        hash TEXT NOT NULL DEFAULT '',
        vector BLOB
    )""")
    conn.execute("CREATE INDEX IF NOT EXISTS idx_chunks_fn ON chunks(filename)")
    conn.execute("CREATE INDEX IF NOT EXISTS idx_chunks_st ON chunks(source_type)")
    return conn


def load_sources():
    config_file = os.path.join(KNOWLEDGE_DIR, ".knowledge-config.json")
    try:
        with open(config_file, "r", encoding="utf-8") as f:
            cfg = json.load(f)
            return cfg.get("sources", [])
    except Exception:
        return ["bilibili", "notes", "wechat-articles", "other"]


def hash_file(path):
    with open(path, "rb") as f:
        return hashlib.md5(f.read()).hexdigest()


def extract_metadata_bilibili(text, filename):
    """从B站转录文件头提取元数据"""
    meta = {"title": "", "author": "", "date": "", "bvid": "", "source_type": "bilibili"}
    for line in text.split("\n")[:20]:
        m = re.match(r"📹 视频标题[：:]\s*(.+)", line)
        if m: meta["title"] = m.group(1).strip()
        m = re.match(r"👤 作者[：:]\s*(.+)", line)
        if m: meta["author"] = m.group(1).strip()
        m = re.match(r"📅 发布时间[：:]\s*(.+)", line)
        if m: meta["date"] = m.group(1).strip()
        m = re.search(r"BV[a-zA-Z0-9]+", filename)
        if m: meta["bvid"] = m.group(0)
    # 提取正文
    body_start = text.find("第二部分：完整原文")
    if body_start == -1:
        body_start = text.find("完整原文")
    body = text[body_start:] if body_start > 0 else text
    body = re.sub(r"^.*第.部分.*\n?", "", body)
    body = re.sub(r"文档结束.*", "", body)
    return meta, body.strip()


def extract_metadata_note(text, filename):
    """通用笔记/文章提取"""
    title = filename
    for line in text.strip().split("\n"):
        stripped = line.strip()
        if stripped.startswith("# ") and not stripped.startswith("##"):
            title = stripped[2:].strip()
            break
    meta = {"title": title, "author": "", "date": "", "source_type": "notes"}
    # 尝试从文件名提取日期（YYYY-MM-DD 开头）
    m = re.match(r"(\d{4}-\d{2}-\d{2})", filename)
    if m: meta["date"] = m.group(1)
    return meta, text.strip()


def extract_metadata_generic(text, filename, source_type):
    title = filename
    for line in text.strip().split("\n"):
        stripped = line.strip()
        if stripped.startswith("# ") and not stripped.startswith("##"):
            title = stripped[2:].strip()
            break
    meta = {"title": title, "author": "", "date": "", "source_type": source_type}
    return meta, text.strip()


def extract_pdf_text(filepath):
    if not HAS_PYPDF2 and not HAS_PDFPLUMBER:
        raise RuntimeError("PDF 解析库未安装，请先创建虚拟环境: cd <skill_dir> && python3 -m venv .venv && .venv/bin/pip install PyPDF2 pdfplumber")

    pages = []

    if HAS_PYPDF2:
        reader = PdfReader(filepath)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text and page_text.strip():
                pages.append(page_text.strip())

    if not pages and HAS_PDFPLUMBER:
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    pages.append(page_text.strip())

    if not pages:
        if not HAS_PDFPLUMBER:
            raise RuntimeError(
                "PDF 未提取到文字内容（可能是纯图像PDF，"
                "安装 pdfplumber 可能解决: .venv/bin/pip install pdfplumber"
            )
        raise RuntimeError("PDF 未提取到文字内容（可能是纯图像PDF）")
    return "\n\n".join(pages)


def extract_docx_text(filepath):
    if not HAS_DOCX:
        raise RuntimeError("python-docx 未安装，请先创建虚拟环境: .venv/bin/pip install python-docx")
    doc = Document(filepath)
    parts = []

    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    if not parts:
        raise RuntimeError("DOCX 未提取到文字内容")
    return "\n\n".join(parts)


def extract_metadata_pdf(text, filename):
    title = filename
    for line in text.strip().split("\n"):
        stripped = line.strip()
        if stripped and len(stripped) < 100 and not stripped.endswith("."):
            title = stripped[:80]
            break
    meta = {"title": title, "author": "", "date": "", "source_type": "pdf"}
    m = re.match(r"(\d{4}-\d{2}-\d{2})", filename)
    if m:
        meta["date"] = m.group(1)
    return meta, text.strip()


def extract_metadata_docx(text, filename):
    title = filename
    for line in text.strip().split("\n"):
        stripped = line.strip()
        if stripped and len(stripped) < 100 and not stripped.endswith("."):
            title = stripped[:80]
            break
    meta = {"title": title, "author": "", "date": "", "source_type": "docx"}
    m = re.match(r"(\d{4}-\d{2}-\d{2})", filename)
    if m:
        meta["date"] = m.group(1)
    return meta, text.strip()


def chunk_text(text, max_chars=CHUNK_SIZE, overlap=CHUNK_OVERLAP):
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    chunks = []
    current = ""
    for line in lines:
        if line.startswith("==="):
            continue
        if len(line) > max_chars:
            line = line[:max_chars]
        if len(current) + len(line) + 1 < max_chars:
            current += "\n" + line if current else line
        else:
            if current:
                chunks.append(current)
            if overlap > 0 and current:
                overlap_lines = current.split("\n")
                overlap_text = ""
                for ol in reversed(overlap_lines):
                    if len(overlap_text) + len(ol) + 1 < overlap:
                        overlap_text = ol + ("\n" + overlap_text if overlap_text else "")
                    else:
                        break
                current = (overlap_text + "\n" + line) if overlap_text else line
            else:
                current = line
    if current:
        chunks.append(current)
    return chunks


def embed(text):
    """调 Ollama 向量化"""
    resp = requests.post(OLLAMA_URL, json={
        "model": get_embed_model(),
        "prompt": text,
    }, timeout=30)
    resp.raise_for_status()
    return resp.json()["embedding"]


def scan_dir(source_type, existing_hashes, force=False):
    src_dir = os.path.join(KNOWLEDGE_DIR, source_type)
    if not os.path.isdir(src_dir):
        return [], False, set()

    has_update = False
    all_chunks = []
    scanned_fnames = set()

    text_files = list(Path(src_dir).glob("*.txt")) + list(Path(src_dir).glob("*.md"))
    pdf_files = list(Path(src_dir).glob("*.pdf"))
    docx_files = list(Path(src_dir).glob("*.docx"))
    all_files = sorted(text_files + pdf_files + docx_files, key=lambda p: p.name)

    for fp in all_files:
        try:
            fhash = hash_file(fp)
        except Exception as e:
            print(f"  [{source_type}] {fp.name} → ⚠️ 无法读取: {e}")
            continue
        fname = fp.name
        scanned_fnames.add(fname)

        if not force and fname in existing_hashes and existing_hashes[fname] == fhash:
            print(f"  [{source_type}] {fname} → (未变化，跳过)")
            continue

        suffix = fp.suffix.lower()
        try:
            if suffix == ".pdf":
                raw = extract_pdf_text(str(fp))
                meta, body = extract_metadata_pdf(raw, fname)
            elif suffix == ".docx":
                raw = extract_docx_text(str(fp))
                meta, body = extract_metadata_docx(raw, fname)
            else:
                with open(fp, "r", encoding="utf-8") as f:
                    raw = f.read()
                if source_type == "bilibili":
                    meta, body = extract_metadata_bilibili(raw, fname)
                elif source_type == "notes":
                    meta, body = extract_metadata_note(raw, fname)
                else:
                    meta, body = extract_metadata_generic(raw, fname, source_type)
        except Exception as e:
            print(f"  [{source_type}] {fname} → ⚠️ 解析失败: {e}")
            continue

        chunks = chunk_text(body)
        for i, chunk in enumerate(chunks):
            all_chunks.append({
                "id": f"{fname}#chunk{i}",
                "text": chunk,
                "title": meta.get("title", fname),
                "author": meta.get("author", ""),
                "date": meta.get("date", ""),
                "bvid": meta.get("bvid", ""),
                "source_type": source_type,
                "source_label": SOURCE_MAP.get(source_type, source_type),
                "filename": fname,
                "chunk_index": i,
                "hash": fhash,
            })
        print(f"  [{source_type}] {fname} → {len(chunks)} 段")
        has_update = True

    return all_chunks, has_update, scanned_fnames


def _get_existing_hashes(conn):
    existing = {}
    cursor = conn.execute("SELECT filename, hash FROM chunks")
    for fname, fhash in cursor:
        existing[fname] = fhash
    return existing


def build_index(force=False, source_only=None):
    current_model = get_embed_model()
    stored_model, stored_dim = load_model_meta()

    if stored_model and stored_model != current_model and not force:
        print(f"\n⚠️  检测到 Embedding 模型变化！")
        print(f"   旧模型: {stored_model} ({stored_dim}维)")
        print(f"   新模型: {current_model}")
        print(f"   → 自动触发全量重建索引...\n")
        force = True
        stored_dim = None

    conn = _get_db()
    existing_hashes = _get_existing_hashes(conn)

    # mtime 快速判断
    if not force and existing_hashes:
        db_path = os.path.join(RAG_DIR, "chunks.db")
        db_mtime = os.path.getmtime(db_path)
        latest_src_mtime = 0
        for src_type in load_sources():
            src_dir = os.path.join(KNOWLEDGE_DIR, src_type)
            if not os.path.isdir(src_dir):
                continue
            for fp in Path(src_dir).glob("*"):
                if fp.is_file() and fp.suffix.lower() in (".txt", ".md", ".pdf", ".docx"):
                    src_mtime = fp.stat().st_mtime
                    if src_mtime > latest_src_mtime:
                        latest_src_mtime = src_mtime
        if db_mtime > latest_src_mtime:
            cursor = conn.execute("SELECT COUNT(*) FROM chunks")
            count = cursor.fetchone()[0]
            conn.close()
            print("\n✅ 没有新的或变化的文件，索引已是最新")
            print(f"   总段落数: {count}")
            return

    if source_only:
        dirs_to_scan = [source_only]
    else:
        dirs_to_scan = [d for d in load_sources()
                        if os.path.isdir(os.path.join(KNOWLEDGE_DIR, d))]

    print(f"📂 扫描知识仓库: {KNOWLEDGE_DIR}")
    if source_only:
        print(f"   (仅扫描: {source_only})")

    new_chunks = []
    all_scanned_fnames = set()
    any_update = False

    for src_type in dirs_to_scan:
        chunks, updated, scanned = scan_dir(src_type, existing_hashes, force)
        new_chunks.extend(chunks)
        any_update = any_update or updated
        all_scanned_fnames.update(scanned)

    if not any_update and not force:
        cursor = conn.execute("SELECT COUNT(*) FROM chunks")
        count = cursor.fetchone()[0]
        if count > 0 or not new_chunks:
            conn.close()
            if count > 0:
                print("\n✅ 没有新的或变化的文件，索引已是最新")
                print(f"   总段落数: {count}")
            else:
                print("   (没有找到文件)")
            return

    # 删除已变更文件的旧 chunk
    changed_fnames = {c["filename"] for c in new_chunks}
    for fname in changed_fnames:
        conn.execute("DELETE FROM chunks WHERE filename = ?", (fname,))

    # 删除已不存在的文件
    if not force:
        db_fnames = set(row[0] for row in conn.execute("SELECT DISTINCT filename FROM chunks"))
        deleted_fnames = db_fnames - all_scanned_fnames
        for fname in deleted_fnames:
            conn.execute("DELETE FROM chunks WHERE filename = ?", (fname,))

    if force:
        conn.execute("DELETE FROM chunks")

    # 向量化新的 chunk
    embed_errors = 0
    if new_chunks:
        print(f"\n🧠 正在向量化 {len(new_chunks)} 个段落...")
        embed_dim = stored_dim
        with concurrent.futures.ThreadPoolExecutor(max_workers=4) as executor:
            future_map = {}
            for i, chunk in enumerate(new_chunks):
                future = executor.submit(embed, chunk["text"])
                future_map[future] = (i, chunk)
            for i, future in enumerate(concurrent.futures.as_completed(future_map)):
                idx, chunk = future_map[future]
                try:
                    vec = future.result()
                    if embed_dim is None:
                        embed_dim = len(vec)
                    vec_bytes = np.array(vec, dtype=np.float32).tobytes()
                except Exception as e:
                    embed_errors += 1
                    print(f"  ⚠️ 向量化失败 [{chunk.get('filename','?')} chunk {chunk.get('chunk_index','?')}]: {e}")
                    if embed_dim is None:
                        embed_dim = 768
                    vec_bytes = np.zeros(embed_dim, dtype=np.float32).tobytes()
                conn.execute(
                    "INSERT OR REPLACE INTO chunks "
                    "(id, text, title, author, date, bvid, source_type, source_label, filename, chunk_index, hash, vector) "
                    "VALUES (?,?,?,?,?,?,?,?,?,?,?,?)",
                    (chunk["id"], chunk["text"], chunk["title"], chunk["author"],
                     chunk["date"], chunk["bvid"], chunk["source_type"],
                     chunk["source_label"], chunk["filename"],
                     chunk["chunk_index"], chunk["hash"], vec_bytes)
                )
                if (i + 1) % 10 == 0 or i == len(new_chunks) - 1:
                    print(f"  {i+1}/{len(new_chunks)}")

        if embed_dim is not None:
            save_model_meta(current_model, embed_dim)

    # 统计（在关闭前查）
    cursor = conn.execute("SELECT source_label, COUNT(*) FROM chunks GROUP BY source_label")
    sources = dict(cursor.fetchall())
    total = sum(sources.values())

    conn.commit()
    conn.close()

    print(f"\n✅ 索引完成！")
    print(f"   总段落数: {total}")
    if embed_errors:
        print(f"   ⚠️ 向量化失败: {embed_errors} 段（已填零向量）")
    for s, n in sorted(sources.items()):
        print(f"   {s}: {n} 段")
    print(f"   数据库: {os.path.join(RAG_DIR, 'chunks.db')}")


if __name__ == "__main__":
    force = "--force" in sys.argv
    json_output = "--json" in sys.argv
    source_only = None
    for i, a in enumerate(sys.argv):
        if a == "--source" and i + 1 < len(sys.argv):
            source_only = sys.argv[i + 1]
    build_index(force=force, source_only=source_only)
    if json_output:
        conn = _get_db()
        cursor = conn.execute("SELECT source_label, COUNT(*) FROM chunks GROUP BY source_label")
        sources = dict(cursor.fetchall())
        total = sum(sources.values())
        conn.close()
        print(json.dumps({
            "total_chunks": total,
            "source_stats": sources,
        }, ensure_ascii=False))
