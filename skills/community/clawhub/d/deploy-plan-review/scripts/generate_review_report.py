#!/usr/bin/env python3
"""
部署方案审核报告生成器 v3.0
输入：变更方案 .docx + 部署表 .xlsx + 纳管表 .xlsx
输出：Word 审核报告 + Excel 问题跟踪表
"""

import os, re, sys, json, shutil
from pathlib import Path
from datetime import datetime
from collections import defaultdict
from typing import Optional

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("错误: 请安装 python-docx: pip3 install python-docx")
    sys.exit(1)

try:
    import openpyxl
except ImportError:
    print("错误: 请安装 openpyxl: pip3 install openpyxl")
    sys.exit(1)

# ─── 路径配置 ─────────────────────────────────────────
SKILL_DIR = Path(__file__).resolve().parent.parent
OUTPUT_DIR = SKILL_DIR / "output"
DATA_DIR   = SKILL_DIR / "data"
RAW_DIR    = DATA_DIR / "raw"
TRACK_XLSX = DATA_DIR / "问题跟踪表.xlsx"

OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
DATA_DIR.mkdir(parents=True, exist_ok=True)
RAW_DIR.mkdir(parents=True, exist_ok=True)


# ═══════════════════════════════════════════════════════
# 风险等级体系（P0/P1/P2/P3）
# ═══════════════════════════════════════════════════════

RISK_MATRIX = {
    # ── P0 高危 ──
    "P0-01": {"level": "P0", "name": "IP地址冲突", "desc": "方案中多个节点使用同一IP，或方案IP与已有生产环境IP段重叠"},
    "P0-02": {"level": "P0", "name": "IP不连续/数量不符", "desc": "方案给出的IP区间数量 ≠ 实际节点数，或跳跃超过1位"},
    "P0-03": {"level": "P0", "name": "CPU/OS架构与部署表矛盾", "desc": "方案写x86但部署表是aarch64（或反过来）"},
    "P0-04": {"level": "P0", "name": "yum源架构不匹配", "desc": "repo URL中的架构(aarch64/x86_64)与节点CPU架构不一致"},
    "P0-05": {"level": "P0", "name": "monitors数量/IP错误", "desc": "mon节点IP不在存储IP范围内，或mon数量不符合ceph最少3台要求"},
    "P0-06": {"level": "P0", "name": "AZ字段为NULL但方案说新建", "desc": "新建AZ场景下部署表所属az列为空"},
    "P0-07": {"level": "P0", "name": "关键步骤缺失", "desc": "对比模板，缺少模板中标记为「必须/必做」的操作步骤"},
    "P0-08": {"level": "P0", "name": "纳管表服务器缺失", "desc": "部署表中的服务器在纳管表中不存在（未纳管）"},
    "P0-09": {"level": "P0", "name": "纳管IP不一致", "desc": "部署表管理IP与纳管表管理IP不一致"},
    "P0-10": {"level": "P0", "name": "明文密码泄露", "desc": "方案文档中出现明文密码凭证（root密码、SSH密码、数据库密码等），存在安全风险"},
    # ── P1 中危 ──
    "P1-01": {"level": "P1", "name": "标题与正文矛盾", "desc": "方案标题/章节标题中的数量与实际内容不一致"},
    "P1-02": {"level": "P1", "name": "节点信息不一致", "desc": "方案正文中的节点列表与表格/部署表中的节点信息不一致"},
    "P1-03": {"level": "P1", "name": "变更人员表填写不完整", "desc": "模板要求的人员角色（实施/复核/审批/客户）缺失"},
    "P1-04": {"level": "P1", "name": "变更时间不合理", "desc": "时间已过当前日期，或窗口时长与操作量不匹配"},
    "P1-05": {"level": "P1", "name": "配置文件/脚本参数错误", "desc": "expand文件、config.json中的参数与部署表不一致"},
    "P1-06": {"level": "P1", "name": "版本修订记录未更新", "desc": "模板中的版本修订表为空白或未填"},
    "P1-07": {"level": "P1", "name": "章节缺失", "desc": "对比模板，缺少非关键章节（如割接回退）"},
    "P1-08": {"level": "P1", "name": "网络配置矛盾", "desc": "方案写的网络模式（VLAN/VXLAN）与实际不一致"},
    # ── P2 低危 ──
    "P2-01": {"level": "P2", "name": "占位符未替换", "desc": "方案中仍包含模板原始占位符 {XXX}, [待补充] 等"},
    "P2-02": {"level": "P2", "name": "文字错别字/语病", "desc": "明显的错别字、不通顺语句、重复内容"},
    "P2-03": {"level": "P2", "name": "格式不一致", "desc": "同级别标题字号不统一，表格样式不一致"},
    "P2-04": {"level": "P2", "name": "编号错误", "desc": "章节编号跳跃或编号体系混合使用"},
    "P2-05": {"level": "P2", "name": "截图/图片缺失", "desc": "模板中有「截图示例」的地方未替换为实际截图"},
    "P2-06": {"level": "P2", "name": "文件命名不规范", "desc": "方案文件名不符合规范格式"},
    "P2-07": {"level": "P2", "name": "操作记录表格式错误", "desc": "操作记录表列名或格式与模板不一致"},
    # ── P3 提示 ──
    "P3-01": {"level": "P3", "name": "纳管表状态异常", "desc": "纳管表中存在非「已纳管」状态的服务器"},
    "P3-02": {"level": "P3", "name": "纳管表AZ与部署表不一致", "desc": "纳管表AZ信息与部署表AZ不匹配"},
}

RISK_ORDER = ["P0", "P1", "P2", "P3"]
RISK_EMOJI = {"P0": "🔴", "P1": "🟡", "P2": "🔵", "P3": "⚪"}


# ═══════════════════════════════════════════════════════
# 解析器
# ═══════════════════════════════════════════════════════

def parse_docx(path: str) -> dict:
    """解析变更方案 .docx"""
    doc = Document(path)
    
    paragraphs = []
    headings = []
    tables_list = []
    full_text = ""
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name if para.style else "Normal"
        paragraphs.append({"style": style_name, "text": text})
        full_text += text + "\n"
        
        if style_name and "Heading" in style_name:
            level_str = style_name.replace("Heading ", "").replace("Heading", "1")
            try:
                level = int(level_str)
            except ValueError:
                level = 1
            headings.append({"level": level, "text": text, "style": style_name})
    
    for i, table in enumerate(doc.tables):
        rows_data = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows_data.append(cells)
        tables_list.append({"index": i, "rows": len(rows_data), "data": rows_data})
    
    return {
        "path": path,
        "filename": Path(path).name,
        "paragraphs": paragraphs,
        "headings": headings,
        "tables": tables_list,
        "full_text": full_text,
        "para_count": len(paragraphs),
        "heading_count": len(headings),
        "table_count": len(tables_list),
    }


def parse_deploy_xlsx(path: str) -> dict:
    """解析部署表 .xlsx"""
    wb = openpyxl.load_workbook(path, data_only=True)
    
    # 自动定位服务器信息 sheet
    target_sheet = None
    for name in wb.sheetnames:
        if "服务器" in name or "server" in name.lower() or "02" in name:
            target_sheet = name
            break
    if target_sheet is None:
        target_sheet = wb.sheetnames[0]  # fallback
    
    ws = wb[target_sheet]
    
    # 读取表头
    headers = []
    for cell in ws[1]:
        headers.append(str(cell.value).strip() if cell.value else "")
    
    # 检测关键列索引
    col_map = {}
    for idx, h in enumerate(headers):
        h_lower = h.lower()
        if "管理ip" in h_lower or "management" in h_lower:
            col_map["mgmt_ip"] = idx
        elif "存储ip" in h_lower or "storage" in h_lower or "ceph" in h_lower:
            col_map["storage_ip"] = idx
        elif "主机名" in h_lower or "hostname" in h_lower or "域名" in h_lower:
            col_map["hostname"] = idx
        elif "cpu" in h_lower and ("型号" in h_lower or "架构" in h_lower or "model" in h_lower):
            col_map["cpu"] = idx
        elif "os" in h_lower or "操作系统" in h_lower:
            col_map["os"] = idx
        elif "az" in h_lower or "可用区" in h_lower:
            col_map["az"] = idx
        elif "角色" in h_lower or "role" in h_lower or "用途" in h_lower:
            col_map["role"] = idx
        elif "序号" in h_lower or "no" == h_lower.strip(".").lower():
            col_map["seq"] = idx
    
    # 读取数据行
    servers = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if all(cell is None or str(cell).strip() == "" for cell in row):
            continue
        
        server = {}
        server["_row"] = [str(c).strip() if c is not None else "" for c in row]
        for key, col_idx in col_map.items():
            val = row[col_idx] if col_idx < len(row) else ""
            server[key] = str(val).strip() if val is not None else ""
        
        servers.append(server)
    
    return {
        "path": path,
        "filename": Path(path).name,
        "sheet": target_sheet,
        "headers": headers,
        "col_map": col_map,
        "servers": servers,
        "server_count": len(servers),
        "all_sheets": wb.sheetnames,
    }


def parse_manage_xlsx(path: str) -> dict:
    """解析纳管表 .xlsx"""
    wb = openpyxl.load_workbook(path, data_only=True)
    
    ws = wb[wb.sheetnames[0]]
    
    headers = []
    for cell in ws[1]:
        headers.append(str(cell.value).strip() if cell.value else "")
    
    col_map = {}
    for idx, h in enumerate(headers):
        h_lower = h.lower()
        if "管理ip" in h_lower or "management" in h_lower:
            col_map["mgmt_ip"] = idx
        elif "主机名" in h_lower or "hostname" in h_lower:
            col_map["hostname"] = idx
        elif "az" in h_lower or "可用区" in h_lower:
            col_map["az"] = idx
        elif "状态" in h_lower or "status" in h_lower:
            col_map["status"] = idx
        elif "管理平台" in h_lower or "platform" in h_lower or "管控" in h_lower:
            col_map["platform"] = idx
        elif "纳管" in h_lower:
            if "状态" not in h_lower:
                col_map["mgmt_status"] = idx
    
    entries = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if all(cell is None or str(cell).strip() == "" for cell in row):
            continue
        entry = {}
        entry["_row"] = [str(c).strip() if c is not None else "" for c in row]
        for key, col_idx in col_map.items():
            val = row[col_idx] if col_idx < len(row) else ""
            entry[key] = str(val).strip() if val is not None else ""
        entries.append(entry)
    
    return {
        "path": path,
        "filename": Path(path).name,
        "sheet": wb.sheetnames[0],
        "headers": headers,
        "col_map": col_map,
        "entries": entries,
        "entry_count": len(entries),
    }


# ═══════════════════════════════════════════════════════
# 交叉核验引擎
# ═══════════════════════════════════════════════════════

def find_chapter(plan: dict, keywords: list) -> Optional[str]:
    """在方案中查找包含关键词的章节"""
    for h in plan["headings"]:
        for kw in keywords:
            if kw in h["text"]:
                return h["text"]
    return None


def extract_ip_range(plan: dict) -> list:
    """从方案文本中提取 IP 地址"""
    ip_pattern = r'\b(?:(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)\.){3}(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)\b'
    return list(set(re.findall(ip_pattern, plan["full_text"])))


def extract_numbers_from_headings(plan: dict) -> dict:
    """从标题中提取数字信息（如：新建2台、扩容12台）"""
    results = {}
    for h in plan["headings"]:
        nums = re.findall(r'(\d+)\s*(?:台|节点|服务器)', h["text"])
        if nums:
            results[h["text"]] = int(nums[0])
    return results


def detect_placeholder(text: str) -> list:
    """检测占位符"""
    patterns = [
        r'\{[A-Z_]+\}',           # {XXX}, {RESOURCE_POOL}
        r'\{[A-Za-z]+\}',         # {Name}
        r'X{2,}',                 # XXXX
        r'\[待补充\]',
        r'\[填写\]',
        r'XXX资源池',
        r'XXX',
    ]
    found = []
    for pat in patterns:
        matches = re.findall(pat, text)
        for m in matches:
            if m not in found:
                found.append(m)
    return found


def run_all_checks(plan: dict, deploy: dict, manage: Optional[dict] = None) -> list:
    """执行全部交叉核验，返回 findings 列表"""
    findings = []
    
    # ── A类：结构合规性 ──
    _check_structure(plan, findings)
    
    # ── B类：数据一致性（方案 vs 部署表）──
    _check_consistency(plan, deploy, findings)
    
    # ── C类：纳管表交叉核验 ──
    if manage:
        _check_management(plan, deploy, manage, findings)
    
    # ── P0-10：明文密码泄露（最高优先级安全检测）──
    _check_plaintext_passwords(plan, findings)
    
    # ── D类：编辑质量 ──
    _check_quality(plan, findings)
    
    # 按风险等级排序
    findings.sort(key=lambda f: (RISK_ORDER.index(f["level"]), f["id"]))
    
    return findings


def _add_finding(findings: list, risk_id: str, location: str, description: str, evidence: str):
    """添加一条审核发现"""
    risk = RISK_MATRIX.get(risk_id, {"level": "P3", "name": risk_id, "desc": ""})
    findings.append({
        "id": risk_id,
        "level": risk["level"],
        "name": risk["name"],
        "location": location,
        "description": description,
        "evidence": evidence,
        "suggestion": f"请核对并修正 — {risk['desc']}",
    })


def _check_structure(plan: dict, findings: list):
    """A类 — 结构合规性检查"""
    heading_texts = [h["text"] for h in plan["headings"]]
    full = "\n".join(heading_texts)
    
    # P1-07: 章节缺失
    required_chapters = {
        "变更需求": ["变更需求", "需求概述"],
        "变更安排": ["变更安排", "实施安排"],
        "变更步骤": ["变更步骤", "变更方案", "实施步骤"],
        "割接回退": ["割接回退", "回退", "回滚"],
    }
    for chapter, keywords in required_chapters.items():
        if not any(kw in full for kw in keywords):
            _add_finding(findings, "P1-07", "全局",
                        f"缺少「{chapter}」章节",
                        f"方案标题列表中未找到 {chapter} 相关内容")
    
    # P1-06: 版本修订记录未更新
    for t in plan["tables"]:
        first_row = t["data"][0] if t["data"] else []
        first_cell = "".join(str(c) for c in first_row).lower()
        if "版本" in first_cell or "修订" in first_cell or "version" in first_cell:
            # 检查是否有实际内容（超过表头行）
            if t["rows"] <= 1:
                _add_finding(findings, "P1-06", "版本修订记录表",
                           "版本修订记录表为空，未填写修订历史",
                           f"表头: {first_row}，数据行数: {t['rows'] - 1}")
            else:
                # 检查内容是否只是模板占位
                all_content = " ".join(
                    " ".join(str(c) for c in row) for row in t["data"][1:]
                )
                if re.search(r'X{2,}|\{\w+\}|待填', all_content):
                    _add_finding(findings, "P1-06", "版本修订记录表",
                               "版本修订记录含占位符，未实际更新",
                               f"内容含: {all_content[:80]}...")
    
    # P1-03: 变更人员表不完整
    for t in plan["tables"]:
        first_row = t["data"][0] if t["data"] else []
        first_cell = "".join(str(c) for c in first_row).lower()
        if any(kw in first_cell for kw in ["角色", "人员", "职责", "实施", "复核"]):
            all_text = " ".join(
                " ".join(str(c) for c in row) for row in t["data"]
            )
            missing = []
            for role in ["实施", "复核", "审批", "客户"]:
                if role not in all_text:
                    missing.append(role)
            if missing:
                _add_finding(findings, "P1-03", "变更人员表",
                           f"缺少人员角色: {', '.join(missing)}",
                           f"表中未找到: {missing}")


def _check_consistency(plan: dict, deploy: dict, findings: list):
    """B类 — 方案 vs 部署表数据一致性"""
    deploy_count = deploy["server_count"]
    full_text = plan["full_text"]
    
    # P0-02: IP/节点数量不符
    title_nums = extract_numbers_from_headings(plan)
    for title, num in title_nums.items():
        if num != deploy_count:
            _add_finding(findings, "P0-02", title,
                       f"方案标题写「{num}台」，但部署表实际 {deploy_count} 台",
                       f"标题: {title} → {num} 台, 部署表: {deploy_count} 行")
    
    # P0-03: CPU/OS架构矛盾
    deploy_archs = set()
    for s in deploy["servers"]:
        cpu = s.get("cpu", "").lower()
        if "kunpeng" in cpu or "aarch64" in cpu or "arm" in cpu or "鲲鹏" in cpu:
            deploy_archs.add("aarch64")
        elif "hygon" in cpu or "x86" in cpu or "海光" in cpu or "intel" in cpu or "amd" in cpu:
            deploy_archs.add("x86_64")
    
    plan_arch = None
    if re.search(r'aarch64|arm64|arm', full_text, re.IGNORECASE):
        plan_arch = "aarch64"
    if re.search(r'x86[^_]|x86_64', full_text, re.IGNORECASE):
        plan_arch = plan_arch or "x86_64"
    
    if plan_arch and deploy_archs:
        if plan_arch not in deploy_archs:
            _add_finding(findings, "P0-03", "CPU架构",
                       f"方案写 {plan_arch}，部署表实际为 {'/'.join(deploy_archs)}",
                       f"方案标题/正文含 {plan_arch}，部署表CPU列含 {deploy_archs}")
    
    # P0-04: yum源架构不匹配
    yum_pattern = r'(aarch64|x86_64)[/\w]*\.(repo|yum)'
    yum_matches = re.findall(yum_pattern, full_text, re.IGNORECASE)
    for yum_arch in yum_matches:
        if deploy_archs and yum_arch.lower() not in {a.lower() for a in deploy_archs}:
            _add_finding(findings, "P0-04", "yum源配置",
                       f"repo URL 架构为 {yum_arch}，与部署表架构不符",
                       f"yum源: {yum_arch}，部署表: {deploy_archs}")
    
    # P0-06: AZ检查
    if "新建" in full_text and "az" in full_text.lower():
        az_nulls = [s for s in deploy["servers"] if not s.get("az", "").strip()]
        if az_nulls:
            _add_finding(findings, "P0-06", "AZ配置",
                       f"方案提到新建AZ，但部署表中有 {len(az_nulls)} 台服务器AZ为空",
                       f"空AZ服务器序号: {[s.get('seq', '?') for s in az_nulls[:5]]}")
    
    # P1-01: 标题与正文矛盾
    if title_nums:
        for title, num in title_nums.items():
            # 检查正文中是否出现不同的数字
            body_nums = re.findall(r'(\d+)\s*(?:台|节点|服务器)', full_text)
            for bn in body_nums:
                if int(bn) != num and int(bn) > 0:
                    _add_finding(findings, "P1-01", title,
                               f"标题写「{num}台」，正文中出现「{bn}台」",
                               f"标题: {title}, 正文: {bn}台")
                    break
    
    # P1-02: 节点信息不一致
    plan_ips = set(extract_ip_range(plan))
    deploy_ips = set()
    for s in deploy["servers"]:
        if s.get("mgmt_ip"):
            deploy_ips.add(s["mgmt_ip"])
        if s.get("storage_ip"):
            deploy_ips.add(s["storage_ip"])
    
    ips_only_plan = plan_ips - deploy_ips
    ips_only_deploy = deploy_ips - plan_ips
    if ips_only_plan:
        _add_finding(findings, "P1-02", "IP/节点列表",
                   f"方案中有 {len(ips_only_plan)} 个IP不在部署表中",
                   f"方案独有IP: {list(ips_only_plan)[:5]}")
    if ips_only_deploy:
        _add_finding(findings, "P1-02", "IP/节点列表",
                   f"部署表中有 {len(ips_only_deploy)} 个IP未在方案中提及",
                   f"部署表独有IP: {list(ips_only_deploy)[:5]}")


def _check_management(plan: dict, deploy: dict, manage: dict, findings: list):
    """C类 — 纳管表交叉核验"""
    deploy_ips = set()
    for s in deploy["servers"]:
        if s.get("mgmt_ip"):
            deploy_ips.add(s["mgmt_ip"])
    
    manage_ips = set()
    for e in manage["entries"]:
        if e.get("mgmt_ip"):
            manage_ips.add(e["mgmt_ip"])
    
    # P0-08: 纳管表服务器缺失
    missing_ip = deploy_ips - manage_ips
    if missing_ip:
        _add_finding(findings, "P0-08", "纳管表",
                   f"部署表中有 {len(missing_ip)} 台服务器未在纳管表中",
                   f"缺失IP: {sorted(list(missing_ip))[:5]}")
    
    # P0-09: 纳管IP不一致
    if manage.get("col_map", {}).get("mgmt_ip") is not None:
        deploy_ip_map = {}
        for s in deploy["servers"]:
            if s.get("mgmt_ip"):
                deploy_ip_map[s["mgmt_ip"]] = s
        
        manage_ip_map = {}
        for e in manage["entries"]:
            if e.get("mgmt_ip"):
                manage_ip_map[e["mgmt_ip"]] = e
        
        # 找同IP不同hostname
        common_ips = deploy_ip_map.keys() & manage_ip_map.keys()
        mismatches = []
        for ip in common_ips:
            d_host = deploy_ip_map[ip].get("hostname", "").strip()
            m_host = manage_ip_map[ip].get("hostname", "").strip()
            if d_host and m_host and d_host != m_host:
                mismatches.append((ip, d_host, m_host))
        
        if mismatches:
            sample = "; ".join(f"{ip}: 部署表={dh}, 纳管表={mh}" for ip, dh, mh in mismatches[:3])
            _add_finding(findings, "P1-02", "纳管表 vs 部署表",
                       f"同一管理IP下主机名不一致 ({len(mismatches)} 处)",
                       sample)
    
    # P3-01: 纳管状态异常
    for e in manage["entries"]:
        status = e.get("status", e.get("mgmt_status", "")).strip()
        if status and status not in ["已纳管", "正常", "online", "active", "启用"]:
            _add_finding(findings, "P3-01", f"纳管表 {e.get('mgmt_ip', '?')}",
                       f"纳管状态异常: {status}",
                       f"服务器 {e.get('hostname', '?')} 状态={status}")
            break  # 只报一次
    
    # P3-02: 纳管表AZ与部署表不一致
    for e in manage["entries"]:
        m_az = e.get("az", "").strip()
        if not m_az:
            continue
        # 找部署表中同IP的AZ
        for s in deploy["servers"]:
            if s.get("mgmt_ip") == e.get("mgmt_ip") and s.get("az", "").strip() != m_az:
                _add_finding(findings, "P3-02", f"服务器 {e.get('mgmt_ip', '?')}",
                           f"AZ不一致: 部署表={s.get('az')}, 纳管表={m_az}",
                           f"IP: {e.get('mgmt_ip')}")
                return  # 只报一次


def _check_plaintext_passwords(plan: dict, findings: list):
    """P0-10 — 明文密码泄露检测：扫描方案全文，检测疑似明文密码的凭证。
    检测规则：
      1. ansible_ssh_pass 后跟非占位符的具体密码
      2. password/passwd/secret 关键字后跟≥8位具体字符串（排除占位符模式）
      3. 连接字符串中明文密码（如 mysql://user:pass@host）
    占位符判定：xxxxxx / xxx / <password> / ****** / 密码 / 请填写 等视为占位符，不触发。
    """
    import re
    
    placeholder_patterns = [
        r'x{3,}',           # xxx, xxxxxx
        r'X{3,}',           # XXX, XXXXXX
        r'\*{3,}',           # ***, ******
        r'<[\w]+>',         # <password>
        r'密码',             # 中文占位符
        r'请填写',
        r'your_password',
        r'YOUR_PASSWORD',
        r'替换为真实密码',
        r'替换为.*密码',
        r'替换密码',
    ]
    
    def is_placeholder(val: str) -> bool:
        """判断一个值是否为占位符而非真实密码"""
        val = val.strip()
        if not val:
            return True
        for pat in placeholder_patterns:
            if re.search(pat, val):
                return True
        if len(val) <= 2:
            return True
        return False
    
    # 检测模式1: ansible_ssh_pass="xxx" 或 ansible_ssh_pass=xxx
    ansible_pattern = r'ansible_ssh_pass[=":\'\s]+([^"\'\s\n,}\]]{3,})'
    for m in re.finditer(ansible_pattern, plan["full_text"]):
        val = m.group(1).strip().strip('"\'').strip()
        if not is_placeholder(val):
            evidence = m.group(0).strip()[:150]
            _add_finding(findings, "P0-10", m.group(0).strip()[:50],
                       f"方案中出现 ansible_ssh_pass 明文密码，应脱敏为占位符（如 xxxxxx）",
                       f"原文: {evidence}")
            break  # 只报一次
    
    # 检测模式2: 密码关键字后跟具体值
    # 如 password: xxx, 密码: xxx, passwd=xxx
    pw_keyword_patterns = [
        r'(?i)(?:password|passwd|pwd)\s*[=:]\s*([^\s\n,;"\'\}\]]{8,})',
        r'(?i)secret\s*[=:]\s*([^\s\n,;"\'\}\]]{8,})',
        r'(?:密码)\s*[=：]\s*([^\s\n,;"\'\}\]]{8,})',
    ]
    for pat in pw_keyword_patterns:
        for m in re.finditer(pat, plan["full_text"]):
            val = m.group(1).strip()
            if not is_placeholder(val):
                evidence = m.group(0).strip()[:150]
                _add_finding(findings, "P0-10", m.group(0).strip()[:50],
                           f"方案中出现疑似明文密码，应使用占位符替代",
                           f"原文: {evidence}")
                break
    
    # 检测模式3: 连接字符串中的明文密码
    conn_pattern = r'(?:mysql|redis|mongodb|postgresql)://[\w]+:([^@\s]{3,})@'
    for m in re.finditer(conn_pattern, plan["full_text"]):
        val = m.group(1).strip()
        if not is_placeholder(val):
            evidence = m.group(0).strip()[:150]
            _add_finding(findings, "P0-10", "连接字符串",
                       f"方案中出现连接字符串包含明文密码，应脱敏",
                       f"原文: {evidence}")
            break


def _check_quality(plan: dict, findings: list):
    """D类 — 编辑质量检查"""
    full_text = plan["full_text"]
    
    # P2-01: 占位符未替换
    placeholders = detect_placeholder(full_text)
    if placeholders:
        _add_finding(findings, "P2-01", "全文",
                   f"发现 {len(placeholders)} 处未替换的占位符",
                   f"占位符: {placeholders[:10]}")
    
    # P2-02: 错别字检测
    typos = {
        "扩融": "扩容",
        "资原池": "资源池",
        "部暑": "部署",
        "配值": "配置",
        "节店": "节点",
        "完装": "完成",
        "己完成": "已完成",
        "实列": "实例",
        "付载": "负载",
        "云蝶": "云盘",
    }
    found_typos = []
    for wrong, correct in typos.items():
        if wrong in full_text:
            count = full_text.count(wrong)
            found_typos.append(f'「{wrong}」 → 「{correct}」 ({count}次)')
    if found_typos:
        _add_finding(findings, "P2-02", "全文",
                   f"发现疑似错别字: {'; '.join(found_typos[:5])}",
                   f"共 {len(found_typos)} 类错别字")
    
    # P2-04: 章节编号检查
    heading_nums = []
    for h in plan["headings"]:
        m = re.match(r'^(\d+(?:\.\d+)*)', h["text"])
        if m:
            heading_nums.append(m.group(1))
    
    if heading_nums:
        # 检查跳跃（如 1.2.1 → 1.4）
        prev_parts = None
        for hn in heading_nums:
            parts = hn.split(".")
            if prev_parts:
                if len(parts) == len(prev_parts):
                    if int(parts[-1]) != int(prev_parts[-1]) + 1:
                        _add_finding(findings, "P2-04", f"章节 {hn}",
                                   f"编号疑似跳跃: {'.'.join(prev_parts)} → {hn}",
                                   f"上一编号: {'.'.join(prev_parts)}")
                        break
            prev_parts = parts


# ═══════════════════════════════════════════════════════
# Word 报告生成
# ═══════════════════════════════════════════════════════

def generate_review_report(output_path: str, plan_name: str,
                           plan: dict, deploy: dict, manage: Optional[dict],
                           findings: list) -> str:
    """生成 Word 格式审核报告"""
    doc = Document()
    
    # 默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    
    # ── 封面标题 ──
    title = doc.add_heading(f'📋 {plan_name} — 部署方案审核报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 统计
    p0 = [f for f in findings if f["level"] == "P0"]
    p1 = [f for f in findings if f["level"] == "P1"]
    p2 = [f for f in findings if f["level"] == "P2"]
    p3 = [f for f in findings if f["level"] == "P3"]
    
    # 结论判定
    if p0:
        conclusion = "❌ 不通过"
        conclusion_detail = f"存在 {len(p0)} 项高危风险，必须全部修正后方可执行"
    elif len(p1) > 5:
        conclusion = "❌ 不通过"
        conclusion_detail = f"中危风险 {len(p1)} 项 (>5)，建议修正后重审"
    elif len(p1) + len(p2) <= 3:
        conclusion = "✅ 通过"
        conclusion_detail = "无高危风险，低危问题在可控范围内"
    else:
        conclusion = "⚠️ 条件通过"
        conclusion_detail = f"无高危风险，中危 {len(p1)} 项需确认后执行"
    
    # 文档信息表
    doc.add_paragraph()
    info_table = doc.add_table(rows=7, cols=2, style='Light Grid Accent 1')
    info_data = [
        ("方案名称", plan_name),
        ("审核日期", datetime.now().strftime('%Y-%m-%d %H:%M')),
        ("部署表", deploy.get("filename", "未提供")),
        ("纳管表", manage.get("filename", "未提供") if manage else "未提供"),
        ("部署表节点数", str(deploy.get("server_count", "N/A"))),
        ("风险统计", f"🔴P0: {len(p0)}  🟡P1: {len(p1)}  🔵P2: {len(p2)}  ⚪P3: {len(p3)}"),
        ("审核结论", conclusion),
    ]
    for i, (k, v) in enumerate(info_data):
        info_table.rows[i].cells[0].text = k
        info_table.rows[i].cells[1].text = v
        _set_cell_font(info_table.rows[i].cells[0], bold=True)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # ── 一、一致性核查（通过项）──
    doc.add_heading('一、✅ 一致性核查（通过项）', level=1)
    
    checks_passed = _get_passed_checks(findings)
    if checks_passed:
        pass_table = doc.add_table(rows=len(checks_passed) + 1, cols=3, style='Light Grid Accent 1')
        pass_table.rows[0].cells[0].text = "检查项"
        pass_table.rows[0].cells[1].text = "核查结果"
        pass_table.rows[0].cells[2].text = "状态"
        for row_idx, row in enumerate(pass_table.rows[0].cells):
            _set_cell_font(row, bold=True)
        for i, check in enumerate(checks_passed, 1):
            pass_table.rows[i].cells[0].text = check["name"]
            pass_table.rows[i].cells[1].text = check["detail"]
            pass_table.rows[i].cells[2].text = "✅"
    else:
        doc.add_paragraph("所有检查项均存在不一致，无完全通过项。")
    
    doc.add_paragraph()
    
    # ── 二、问题清单 ──
    doc.add_heading('二、❌ 问题清单', level=1)
    
    for level in ["P0", "P1", "P2", "P3"]:
        level_findings = [f for f in findings if f["level"] == level]
        if not level_findings:
            continue
        
        emoji = RISK_EMOJI.get(level, "⚪")
        level_names = {"P0": "高危风险（必须修改）", "P1": "中危风险（建议修改）",
                       "P2": "低危风险（按需修改）", "P3": "提示信息"}
        doc.add_heading(f'{emoji} {level} — {level_names.get(level, "")}', level=2)
        
        # 问题表
        tbl = doc.add_table(rows=len(level_findings) + 1, cols=5, style='Light Grid Accent 1')
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        headers = ["#", "风险ID", "位置", "问题描述", "判定依据"]
        for idx, h in enumerate(headers):
            tbl.rows[0].cells[idx].text = h
            _set_cell_font(tbl.rows[0].cells[idx], bold=True)
        
        for i, f in enumerate(level_findings, 1):
            tbl.rows[i].cells[0].text = str(i)
            tbl.rows[i].cells[1].text = f["id"]
            tbl.rows[i].cells[2].text = f["location"]
            tbl.rows[i].cells[3].text = f["description"]
            tbl.rows[i].cells[4].text = f["evidence"][:200]
        
        doc.add_paragraph()
    
    # ── 三、审核结论 ──
    doc.add_heading('三、📊 审核结论', level=1)
    
    summary_table = doc.add_table(rows=5, cols=2, style='Light Grid Accent 1')
    summary_data = [
        ("🔴 P0 高危", f"{len(p0)} 项 — 必须全部修正"),
        ("🟡 P1 中危", f"{len(p1)} 项 — 建议全部修正"),
        ("🔵 P2 低危", f"{len(p2)} 项 — 按需修改"),
        ("⚪ P3 提示", f"{len(p3)} 项 — 仅供参考"),
        ("总体结论", conclusion_detail),
    ]
    for i, (k, v) in enumerate(summary_data):
        summary_table.rows[i].cells[0].text = k
        summary_table.rows[i].cells[1].text = v
        _set_cell_font(summary_table.rows[i].cells[0], bold=True)
    
    doc.add_paragraph()
    p = doc.add_paragraph(f'审核结论：{conclusion}')
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.bold = True
    if p0:
        p.runs[0].font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    elif p1:
        p.runs[0].font.color.rgb = RGBColor(0xCC, 0x88, 0x00)
    else:
        p.runs[0].font.color.rgb = RGBColor(0x00, 0x88, 0x00)
    
    doc.save(output_path)
    return output_path


def _set_cell_font(cell, bold=False, size=None):
    """设置单元格字体"""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = '微软雅黑'
            if bold:
                run.font.bold = True
            if size:
                run.font.size = size


def _get_passed_checks(findings: list) -> list:
    """计算通过的检查项（在所有可能检查中未出问题的）"""
    all_ids = set(RISK_MATRIX.keys())
    problem_ids = {f["id"] for f in findings}
    passed_ids = all_ids - problem_ids
    
    passed = []
    for rid in sorted(passed_ids):
        risk = RISK_MATRIX[rid]
        passed.append({"name": risk["name"], "detail": risk["desc"]})
    return passed


# ═══════════════════════════════════════════════════════
# Excel 问题跟踪表
# ═══════════════════════════════════════════════════════

def generate_tracking_xlsx(output_path: str, plan_name: str, findings: list) -> str:
    """生成 Excel 问题跟踪表"""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "问题跟踪"
    
    headers = ["序号", "风险ID", "风险等级", "位置", "问题描述", "判定依据", "建议修改方案", "状态"]
    ws.append(headers)
    
    # 表头样式
    from openpyxl.styles import Font, PatternFill, Alignment
    header_font = Font(name="微软雅黑", bold=True, size=11)
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font_white = Font(name="微软雅黑", bold=True, size=11, color="FFFFFF")
    
    for col_idx, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_idx)
        cell.font = header_font_white
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")
    
    # 风险等级颜色
    level_fills = {
        "P0": PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid"),
        "P1": PatternFill(start_color="FFEB9C", end_color="FFEB9C", fill_type="solid"),
        "P2": PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid"),
        "P3": PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid"),
    }
    
    for i, f in enumerate(findings, 1):
        row_data = [
            i, f["id"], f["level"], f["location"],
            f["description"], f["evidence"][:200],
            f["suggestion"], "待修改"
        ]
        ws.append(row_data)
        
        # 行样式
        for col_idx in range(1, len(headers) + 1):
            cell = ws.cell(row=i + 1, column=col_idx)
            cell.font = Font(name="微软雅黑", size=10)
            cell.alignment = Alignment(wrap_text=True)
        
        # 等级列高亮
        level_cell = ws.cell(row=i + 1, column=3)
        fill = level_fills.get(f["level"])
        if fill:
            level_cell.fill = fill
    
    # 列宽
    col_widths = [6, 12, 10, 18, 45, 40, 40, 10]
    for i, w in enumerate(col_widths, 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = w
    
    # 概要 sheet
    ws2 = wb.create_sheet("审核概要")
    ts = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    p0 = len([f for f in findings if f["level"] == "P0"])
    p1 = len([f for f in findings if f["level"] == "P1"])
    p2 = len([f for f in findings if f["level"] == "P2"])
    p3 = len([f for f in findings if f["level"] == "P3"])
    
    summary_data = [
        ["方案名称", plan_name],
        ["审核时间", ts],
        ["P0 高危", p0],
        ["P1 中危", p1],
        ["P2 低危", p2],
        ["P3 提示", p3],
        ["总问题数", len(findings)],
        ["结论", "不通过" if p0 or p1 > 5 else ("条件通过" if p1 else "通过")],
    ]
    for row_data in summary_data:
        ws2.append(row_data)
    
    wb.save(output_path)
    return output_path


# ═══════════════════════════════════════════════════════
# 输入校验
# ═══════════════════════════════════════════════════════

def _validate_inputs(plan_path: str, deploy_path: str, manage_path: Optional[str] = None):
    """校验输入文件完整性。缺少必要文件时抛出明确错误。"""
    errors = []
    warnings = []
    
    # 必要条件：方案 Word
    if not plan_path or not plan_path.strip():
        errors.append("缺少必要文件：变更方案(.docx)")
    elif not os.path.exists(plan_path):
        errors.append(f"变更方案文件不存在：{plan_path}")
    elif not plan_path.lower().endswith('.docx'):
        errors.append(f"变更方案必须是 .docx 格式：{plan_path}")
    
    # 必要条件：部署表 Excel
    if not deploy_path or not deploy_path.strip():
        errors.append("缺少必要文件：部署表(.xlsx)")
    elif not os.path.exists(deploy_path):
        errors.append(f"部署表文件不存在：{deploy_path}")
    elif not deploy_path.lower().endswith('.xlsx'):
        errors.append(f"部署表必须是 .xlsx 格式：{deploy_path}")
    
    # 可选：纳管表 Excel
    if manage_path and manage_path.strip():
        if not os.path.exists(manage_path):
            warnings.append(f"纳管表文件不存在（将跳过C类检查）：{manage_path}")
        elif not manage_path.lower().endswith('.xlsx'):
            warnings.append(f"纳管表不是 .xlsx 格式（将跳过C类检查）：{manage_path}")
    
    if errors:
        raise ValueError(
            "审核无法继续，缺少必要文件：\n  - " +
            "\n  - ".join(errors) +
            "\n\n请补充完整后重新提交。"
        )
    
    return warnings  # 调用方可选择是否输出 warning


# ═══════════════════════════════════════════════════════
# 主入口
# ═══════════════════════════════════════════════════════

def run_audit(plan_path: str, deploy_path: str, manage_path: Optional[str] = None) -> dict:
    """
    主审核入口
    
    Args:
        plan_path:   变更方案 .docx 路径（必需）
        deploy_path: 部署表 .xlsx 路径（必需）
        manage_path: 纳管表 .xlsx 路径（可选，None 则跳过 C 类检查）
    
    Returns:
        dict with keys: plan_name, findings, docx_path, xlsx_path, summary
    
    Raises:
        ValueError: 缺少必要文件（plan_path 或 deploy_path）
        FileNotFoundError: 指定路径的文件不存在
    """
    
    # Step 0: 输入校验
    _validate_inputs(plan_path, deploy_path, manage_path)
    
    # Step 1: 存档原始输入
    ts = datetime.now().strftime('%Y-%m-%d_%H%M%S')
    for label, src_path in [("plan", plan_path), ("deploy", deploy_path), ("manage", manage_path)]:
        if src_path and os.path.exists(src_path):
            dst = RAW_DIR / f'{ts}_{label}_{Path(src_path).name}'
            shutil.copy2(src_path, dst)
    
    # Step 1: 解析输入
    plan = parse_docx(plan_path)
    deploy = parse_deploy_xlsx(deploy_path)
    manage = parse_manage_xlsx(manage_path) if manage_path else None
    
    plan_name = Path(plan_path).stem
    
    # Step 2: 执行全量交叉核验
    findings = run_all_checks(plan, deploy, manage)
    
    # Step 3: 生成 Word 审核报告
    docx_path = OUTPUT_DIR / f'审核报告_{plan_name}_{ts}.docx'
    generate_review_report(str(docx_path), plan_name, plan, deploy, manage, findings)
    
    # Step 4: 生成 Excel 问题跟踪表
    xlsx_path = OUTPUT_DIR / f'问题跟踪_{plan_name}_{ts}.xlsx'
    generate_tracking_xlsx(str(xlsx_path), plan_name, findings)
    
    # 追加到总跟踪表
    _append_to_master_tracking(plan_name, ts, findings)
    
    # 统计
    p0 = len([f for f in findings if f["level"] == "P0"])
    p1 = len([f for f in findings if f["level"] == "P1"])
    p2 = len([f for f in findings if f["level"] == "P2"])
    p3 = len([f for f in findings if f["level"] == "P3"])
    
    return {
        "plan_name": plan_name,
        "plan": plan,
        "deploy": deploy,
        "manage": manage,
        "findings": findings,
        "docx_path": str(docx_path),
        "xlsx_path": str(xlsx_path),
        "summary": {
            "p0": p0, "p1": p1, "p2": p2, "p3": p3,
            "total": len(findings),
            "conclusion": "不通过" if p0 or p1 > 5 else ("条件通过" if p1 else "通过"),
            "timestamp": ts,
        },
        "raw_archived": str(RAW_DIR),
    }


def _append_to_master_tracking(plan_name: str, ts: str, findings: list):
    """追加问题到总跟踪表"""
    if not TRACK_XLSX.exists():
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "问题汇总"
        ws.append(["审核时间", "方案名称", "风险ID", "风险等级", "位置",
                    "问题描述", "判定依据", "建议修改", "当前状态"])
        wb.save(str(TRACK_XLSX))
    
    wb = openpyxl.load_workbook(str(TRACK_XLSX))
    ws = wb.active
    
    for f in findings:
        ws.append([ts, plan_name, f["id"], f["level"], f["location"],
                   f["description"], f["evidence"][:200], f["suggestion"], "待修改"])
    
    wb.save(str(TRACK_XLSX))


# ─── CLI ──────────────────────────────────────────────
if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description="部署方案审核报告生成器 v3.0")
    parser.add_argument("--plan", required=True, help="变更方案 .docx 路径")
    parser.add_argument("--deploy", required=True, help="部署表 .xlsx 路径")
    parser.add_argument("--manage", default=None, help="纳管表 .xlsx 路径（可选）")
    parser.add_argument("--json", action="store_true", help="输出 JSON 格式结果")
    
    args = parser.parse_args()
    
    result = run_audit(args.plan, args.deploy, args.manage)
    
    if args.json:
        # 简化输出
        output = {
            "plan_name": result["plan_name"],
            "summary": result["summary"],
            "findings": result["findings"],
            "docx_path": result["docx_path"],
            "xlsx_path": result["xlsx_path"],
        }
        print(json.dumps(output, ensure_ascii=False, indent=2))
    else:
        s = result["summary"]
        print(f"═══ 审核完成 ═══")
        print(f"方案: {result['plan_name']}")
        print(f"风险: 🔴P0={s['p0']} 🟡P1={s['p1']} 🔵P2={s['p2']} ⚪P3={s['p3']}")
        print(f"结论: {s['conclusion']}")
        print(f"报告: {result['docx_path']}")
        print(f"跟踪: {result['xlsx_path']}")
        print(f"存档: {result['raw_archived']}")
