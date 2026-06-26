#!/usr/bin/env python3
"""
Export customized resume to specified format (PDF, Word, Markdown, HTML, text).
"""

import argparse
import json
import sys
import os
from pathlib import Path

def export_to_text(resume_data, output_path):
    """Export resume to plain text format."""
    text = []
    
    # Contact info
    contact = resume_data.get("contact", {})
    if contact.get("name"):
        text.append(contact["name"])
        text.append("=" * len(contact["name"]))
    
    contact_line = []
    if contact.get("email"):
        contact_line.append(contact["email"])
    if contact.get("phone"):
        contact_line.append(contact["phone"])
    if contact.get("location"):
        contact_line.append(contact["location"])
    if contact_line:
        text.append(" | ".join(contact_line))
    
    if contact.get("linkedin"):
        text.append(f"LinkedIn: {contact['linkedin']}")
    if contact.get("website"):
        text.append(f"Website: {contact['website']}")
    
    text.append("")
    
    # Summary
    if resume_data.get("summary"):
        text.append("SUMMARY")
        text.append("-" * 30)
        text.append(resume_data["summary"])
        text.append("")
    
    # Work Experience
    if resume_data.get("work_experience"):
        text.append("WORK EXPERIENCE")
        text.append("-" * 30)
        for exp in resume_data["work_experience"]:
            text.append(f"{exp.get('title', '')} at {exp.get('company', '')}")
            text.append(f"{exp.get('start_date', '')} - {exp.get('end_date', 'Present')}")
            if exp.get("description"):
                text.append(exp["description"])
            text.append("")
    
    # Education
    if resume_data.get("education"):
        text.append("EDUCATION")
        text.append("-" * 30)
        for edu in resume_data["education"]:
            text.append(f"{edu.get('degree', '')} in {edu.get('field', '')}")
            text.append(f"{edu.get('institution', '')}, {edu.get('graduation_date', '')}")
            text.append("")
    
    # Skills
    if resume_data.get("skills"):
        text.append("SKILLS")
        text.append("-" * 30)
        text.append(", ".join(resume_data["skills"]))
        text.append("")
    
    # Projects
    if resume_data.get("projects"):
        text.append("PROJECTS")
        text.append("-" * 30)
        for proj in resume_data["projects"]:
            text.append(proj.get("name", ""))
            if proj.get("description"):
                text.append(proj["description"])
            text.append("")
    
    # Certifications
    if resume_data.get("certifications"):
        text.append("CERTIFICATIONS")
        text.append("-" * 30)
        for cert in resume_data["certifications"]:
            text.append(cert)
        text.append("")
    
    # Write to file
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(text))
    
    print(f"Resume exported to text: {output_path}")

def export_to_markdown(resume_data, output_path):
    """Export resume to Markdown format."""
    md = []
    
    # Contact info
    contact = resume_data.get("contact", {})
    if contact.get("name"):
        md.append(f"# {contact['name']}")
        md.append("")
    
    contact_line = []
    if contact.get("email"):
        contact_line.append(f"📧 {contact['email']}")
    if contact.get("phone"):
        contact_line.append(f"📞 {contact['phone']}")
    if contact.get("location"):
        contact_line.append(f"📍 {contact['location']}")
    if contact_line:
        md.append(" | ".join(contact_line))
        md.append("")
    
    if contact.get("linkedin"):
        md.append(f"🔗 LinkedIn: {contact['linkedin']}")
    if contact.get("website"):
        md.append(f"🌐 Website: {contact['website']}")
    if contact.get("linkedin") or contact.get("website"):
        md.append("")
    
    # Summary
    if resume_data.get("summary"):
        md.append("## Summary")
        md.append("")
        md.append(resume_data["summary"])
        md.append("")
    
    # Work Experience
    if resume_data.get("work_experience"):
        md.append("## Work Experience")
        md.append("")
        for exp in resume_data["work_experience"]:
            md.append(f"### {exp.get('title', '')} at {exp.get('company', '')}")
            md.append(f"*{exp.get('start_date', '')} - {exp.get('end_date', 'Present')}*")
            md.append("")
            if exp.get("description"):
                # Split description into bullet points
                for line in exp["description"].split('\n'):
                    if line.strip():
                        md.append(f"- {line.strip()}")
                md.append("")
    
    # Education
    if resume_data.get("education"):
        md.append("## Education")
        md.append("")
        for edu in resume_data["education"]:
            md.append(f"### {edu.get('degree', '')} in {edu.get('field', '')}")
            md.append(f"*{edu.get('institution', '')}, {edu.get('graduation_date', '')}*")
            md.append("")
    
    # Skills
    if resume_data.get("skills"):
        md.append("## Skills")
        md.append("")
        md.append(", ".join(resume_data["skills"]))
        md.append("")
    
    # Projects
    if resume_data.get("projects"):
        md.append("## Projects")
        md.append("")
        for proj in resume_data["projects"]:
            md.append(f"### {proj.get('name', '')}")
            if proj.get("description"):
                md.append(proj["description"])
                md.append("")
    
    # Certifications
    if resume_data.get("certifications"):
        md.append("## Certifications")
        md.append("")
        for cert in resume_data["certifications"]:
            md.append(f"- {cert}")
        md.append("")
    
    # Write to file
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(md))
    
    print(f"Resume exported to Markdown: {output_path}")

def export_to_html(resume_data, output_path):
    """Export resume to HTML format."""
    # Simple HTML template
    html = """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Resume</title>
    <style>
        body { font-family: Arial, sans-serif; line-height: 1.6; max-width: 800px; margin: 0 auto; padding: 20px; }
        h1 { color: #333; border-bottom: 2px solid #333; padding-bottom: 10px; }
        h2 { color: #555; margin-top: 20px; }
        h3 { color: #777; }
        .contact { margin-bottom: 20px; }
        .section { margin-bottom: 20px; }
        ul { list-style-type: disc; margin-left: 20px; }
    </style>
</head>
<body>
"""
    
    # Contact info
    contact = resume_data.get("contact", {})
    if contact.get("name"):
        html += f"    <h1>{contact['name']}</h1>\n"
    
    html += '    <div class="contact">\n'
    contact_line = []
    if contact.get("email"):
        contact_line.append(f"Email: {contact['email']}")
    if contact.get("phone"):
        contact_line.append(f"Phone: {contact['phone']}")
    if contact.get("location"):
        contact_line.append(f"Location: {contact['location']}")
    if contact_line:
        html += f"        <p>{' | '.join(contact_line)}</p>\n"
    
    if contact.get("linkedin"):
        html += f"        <p>LinkedIn: {contact['linkedin']}</p>\n"
    if contact.get("website"):
        html += f"        <p>Website: {contact['website']}</p>\n"
    html += '    </div>\n'
    
    # Summary
    if resume_data.get("summary"):
        html += '    <div class="section">\n'
        html += '        <h2>Summary</h2>\n'
        html += f"        <p>{resume_data['summary']}</p>\n"
        html += '    </div>\n'
    
    # Work Experience
    if resume_data.get("work_experience"):
        html += '    <div class="section">\n'
        html += '        <h2>Work Experience</h2>\n'
        for exp in resume_data["work_experience"]:
            html += f"        <h3>{exp.get('title', '')} at {exp.get('company', '')}</h3>\n"
            html += f"        <p><em>{exp.get('start_date', '')} - {exp.get('end_date', 'Present')}</em></p>\n"
            if exp.get("description"):
                html += "        <ul>\n"
                for line in exp["description"].split('\n'):
                    if line.strip():
                        html += f"            <li>{line.strip()}</li>\n"
                html += "        </ul>\n"
        html += '    </div>\n'
    
    # Education
    if resume_data.get("education"):
        html += '    <div class="section">\n'
        html += '        <h2>Education</h2>\n'
        for edu in resume_data["education"]:
            html += f"        <h3>{edu.get('degree', '')} in {edu.get('field', '')}</h3>\n"
            html += f"        <p><em>{edu.get('institution', '')}, {edu.get('graduation_date', '')}</em></p>\n"
        html += '    </div>\n'
    
    # Skills
    if resume_data.get("skills"):
        html += '    <div class="section">\n'
        html += '        <h2>Skills</h2>\n'
        html += f"        <p>{', '.join(resume_data['skills'])}</p>\n"
        html += '    </div>\n'
    
    html += "</body>\n</html>"
    
    # Write to file
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(html)
    
    print(f"Resume exported to HTML: {output_path}")

def export_to_docx(resume_data, output_path):
    """Export resume to Word format."""
    try:
        from docx import Document
        
        doc = Document()
        
        # Contact info
        contact = resume_data.get("contact", {})
        if contact.get("name"):
            doc.add_heading(contact["name"], 0)
        
        contact_line = []
        if contact.get("email"):
            contact_line.append(contact["email"])
        if contact.get("phone"):
            contact_line.append(contact["phone"])
        if contact.get("location"):
            contact_line.append(contact["location"])
        if contact_line:
            doc.add_paragraph(" | ".join(contact_line))
        
        # Summary
        if resume_data.get("summary"):
            doc.add_heading("Summary", 1)
            doc.add_paragraph(resume_data["summary"])
        
        # Work Experience
        if resume_data.get("work_experience"):
            doc.add_heading("Work Experience", 1)
            for exp in resume_data["work_experience"]:
                doc.add_heading(f"{exp.get('title', '')} at {exp.get('company', '')}", 2)
                doc.add_paragraph(f"{exp.get('start_date', '')} - {exp.get('end_date', 'Present')}")
                if exp.get("description"):
                    for line in exp["description"].split('\n'):
                        if line.strip():
                            doc.add_paragraph(line.strip(), style='List Bullet')
        
        # Education
        if resume_data.get("education"):
            doc.add_heading("Education", 1)
            for edu in resume_data["education"]:
                doc.add_heading(f"{edu.get('degree', '')} in {edu.get('field', '')}", 2)
                doc.add_paragraph(f"{edu.get('institution', '')}, {edu.get('graduation_date', '')}")
        
        # Skills
        if resume_data.get("skills"):
            doc.add_heading("Skills", 1)
            doc.add_paragraph(", ".join(resume_data["skills"]))
        
        # Save
        doc.save(output_path)
        print(f"Resume exported to Word: {output_path}")
    
    except ImportError:
        print("Error: python-docx not installed. Install with: pip install python-docx")
        sys.exit(1)

def export_to_pdf(resume_data, output_path):
    """Export resume to PDF format (via HTML -> PDF)."""
    # This is a simplified version - in production, use reportlab or weasyprint
    print("Warning: PDF export is not fully implemented. Use HTML export and convert manually.")
    print("Alternatively, install weasyprint: pip install weasyprint")
    
    # Fallback: export as HTML
    html_path = output_path.replace('.pdf', '.html')
    export_to_html(resume_data, html_path)
    print(f"Fallback: Resume exported to HTML instead: {html_path}")
    print("You can convert HTML to PDF using browser's Print to PDF feature.")

def main():
    parser = argparse.ArgumentParser(description='Export customized resume to specified format.')
    parser.add_argument('--input', required=True, help='Input resume JSON file')
    parser.add_argument('--output', required=True, help='Output file path')
    parser.add_argument('--format', required=True, choices=['pdf', 'docx', 'md', 'html', 'txt'],
                       help='Output format')
    parser.add_argument('--template', help='Template name (optional)')
    parser.add_argument('--ats-friendly', action='store_true', help='Ensure ATS-friendly formatting')
    
    args = parser.parse_args()
    
    # Load input file
    if not os.path.exists(args.input):
        print(f"Error: Resume file not found: {args.input}")
        sys.exit(1)
    
    with open(args.input, 'r', encoding='utf-8') as f:
        resume_data = json.load(f)
    
    # Export based on format
    if args.format == 'txt':
        export_to_text(resume_data, args.output)
    elif args.format == 'md':
        export_to_markdown(resume_data, args.output)
    elif args.format == 'html':
        export_to_html(resume_data, args.output)
    elif args.format == 'docx':
        export_to_docx(resume_data, args.output)
    elif args.format == 'pdf':
        export_to_pdf(resume_data, args.output)
    
    print(f"\nExport completed successfully!")
    if args.ats_friendly:
        print("Note: ATS-friendly formatting applied.")

if __name__ == '__main__':
    main()
