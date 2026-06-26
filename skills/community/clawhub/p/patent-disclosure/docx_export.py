#!/usr/bin/env python3
"""将交底书 Markdown 转为 Word (.docx)。"""

from __future__ import annotations

import importlib.util
import io
import re
from pathlib import Path
from typing import Any, List

_SKILL_DIR = Path(__file__).resolve().parent
_deps_mod: Any = None

DOCX_MIME = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
MD_MIME = "text/markdown; charset=utf-8"


class ExportNotAvailable(Exception):
    """Word 导出依赖不可用。"""


def _load_deps() -> Any:
    global _deps_mod
    if _deps_mod is not None:
        return _deps_mod
    path = _SKILL_DIR / "export_deps.py"
    spec = importlib.util.spec_from_file_location(
        "skill_export_deps_patent_disclosure", path
    )
    if spec is None or spec.loader is None:
        raise ImportError(f"无法加载 {path}")
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    _deps_mod = mod
    return mod


def _parse_table_row(line: str) -> List[str]:
    parts = [c.strip() for c in line.strip().strip("|").split("|")]
    return parts


def _is_table_separator(line: str) -> bool:
    s = line.strip().strip("|").replace(" ", "")
    return bool(s) and set(s) <= {"|", "-", ":"}


def markdown_to_docx_bytes(markdown: str, document_title: str = "技术交底书") -> bytes:
    deps = _load_deps()
    ok, err = deps.ensure_docx()
    if not ok:
        raise ExportNotAvailable(err or "python-docx 不可用")

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = Document()
    title = (document_title or "技术交底书").strip()
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    lines = markdown.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            table_lines: List[str] = []
            while i < len(lines):
                row = lines[i].strip()
                if not row.startswith("|"):
                    break
                table_lines.append(row)
                i += 1
            data_rows = [
                _parse_table_row(r)
                for r in table_lines
                if r.strip() and not _is_table_separator(r)
            ]
            if data_rows:
                cols = max(len(r) for r in data_rows)
                table = doc.add_table(rows=len(data_rows), cols=cols)
                table.style = "Table Grid"
                for ri, row in enumerate(data_rows):
                    for ci in range(cols):
                        cell_text = row[ci] if ci < len(row) else ""
                        table.rows[ri].cells[ci].text = cell_text
                doc.add_paragraph()
            continue

        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=4)
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
            i += 1
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
            i += 1
            continue

        if stripped.startswith(("- ", "* ")):
            doc.add_paragraph(stripped[2:], style="List Bullet")
            i += 1
            continue

        m = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if m:
            doc.add_paragraph(m.group(2), style="List Number")
            i += 1
            continue

        p = doc.add_paragraph()
        run = p.add_run(stripped)
        run.font.size = Pt(11)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
