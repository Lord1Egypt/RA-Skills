#!/usr/bin/env python3
"""
Analyze current document state to see what's missing
"""

import os
import sys
from docx import Document

def analyze_word_document(doc_path):
    """Analyze Word document content"""
    if not os.path.exists(doc_path):
        print(f"❌ Document not found: {doc_path}")
        return None
    
    try:
        doc = Document(doc_path)
        print(f"📄 Analyzing: {doc_path}")
        print(f"   Size: {os.path.getsize(doc_path):,} bytes")
        print(f"   Paragraphs: {len(doc.paragraphs)}")
        
        # Look for chapter markers
        chapters_found = []
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text.startswith("فصل") or "فصل" in text:
                chapters_found.append((i, text))
        
        print(f"   Chapter headings found: {len(chapters_found)}")
        for idx, chapter in chapters_found:
            print(f"     - {chapter}")
        
        return chapters_found
        
    except Exception as e:
        print(f"❌ Error analyzing document: {e}")
        return None

def check_text_files():
    """Check available text files"""
    print("\n📋 Checking text files:")
    chapters = {}
    
    for chap_num in range(1, 8):
        filename = f"chapter{chap_num}_clean.txt"
        if os.path.exists(filename):
            size = os.path.getsize(filename)
            with open(filename, 'r', encoding='utf-8') as f:
                first_line = f.readline().strip()
            chapters[chap_num] = {
                'size': size,
                'first_line': first_line[:50] + "..." if len(first_line) > 50 else first_line
            }
            print(f"✅ Chapter {chap_num}: {size:,} bytes - {first_line[:50]}...")
        else:
            print(f"❌ Chapter {chap_num}: MISSING")
    
    return chapters

def check_all_documents():
    """Check all Word documents"""
    print("🔍 DOCUMENT ANALYSIS REPORT")
    print("=" * 50)
    
    # Check all .docx files
    doc_files = []
    for f in os.listdir('.'):
        if f.endswith('.docx') and not f.startswith('~$'):
            doc_files.append(f)
    
    print(f"Found {len(doc_files)} Word documents:")
    
    chapters_by_doc = {}
    for doc_file in sorted(doc_files):
        print(f"\n📊 {doc_file}:")
        chapters = analyze_word_document(doc_file)
        if chapters:
            chapters_by_doc[doc_file] = [chap[1] for chap in chapters]
    
    # Check text files
    text_chapters = check_text_files()
    
    # Generate summary
    print("\n" + "=" * 50)
    print("📋 SUMMARY")
    print("=" * 50)
    
    print("\n📚 Available chapters in text files:")
    for chap_num in sorted(text_chapters.keys()):
        info = text_chapters[chap_num]
        print(f"   Chapter {chap_num}: {info['size']:,} bytes")
    
    print("\n📄 Chapters in each document:")
    for doc_file, chapters in chapters_by_doc.items():
        print(f"\n   {doc_file}:")
        for chap in chapters:
            print(f"     - {chap}")
    
    # Identify missing chapters
    print("\n🔍 IDENTIFIED ISSUES:")
    
    # Check ORIGINAL document
    if 'BSc-project-report-ORIGINAL.docx' in chapters_by_doc:
        original_chapters = chapters_by_doc['BSc-project-report-ORIGINAL.docx']
        print(f"1. ORIGINAL has {len(original_chapters)} chapters")
        
        # Expected chapters 1-4 in original
        expected_in_original = 4
        if len(original_chapters) < expected_in_original:
            print(f"   ❌ Missing chapters in ORIGINAL: Expected {expected_in_original}, found {len(original_chapters)}")
    
    # Check text files availability
    missing_text_files = []
    for chap_num in range(1, 8):
        if chap_num not in text_chapters:
            missing_text_files.append(chap_num)
    
    if missing_text_files:
        print(f"2. Missing text files: Chapters {missing_text_files}")
    else:
        print("2. ✅ All text files available (chapters 1-7)")
    
    print("\n🚀 RECOMMENDED ACTION:")
    print("1. Use python-docx to merge ORIGINAL (chapters 1-4) with text files 5-7")
    print("2. Apply proper formatting")
    print("3. Verify all 7 chapters present")

def main():
    try:
        check_all_documents()
    except KeyboardInterrupt:
        print("\n⏹️ Analysis interrupted")
    except Exception as e:
        print(f"❌ Error: {e}")

if __name__ == "__main__":
    main()