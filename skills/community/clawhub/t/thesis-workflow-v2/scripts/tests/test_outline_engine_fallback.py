#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
test_outline_engine_fallback.py - outline_parser 引擎切换 B→A 单向降级测试

v2.0.7 新增(#1 outline_parser 引擎切换任务)
覆盖 F1-F5 决策 + 统一入口 + 单向降级语义。

决策回顾:
  F1 弹窗   = 否: 不调 warnings.warn(用户静默)
  F2 audit   = 是: 降级事件写 state.audit_log
  F3 重试    = 1 次: MinerU 失败 1 次立即降级
  F4 作用域  = 进程级: _fallback_used 是模块全局
  F5 跨 paper = 共享: 跨 paper 共享 _fallback_used

降级语义: B (MinerU) → A (heuristic),A 失败不回 B(单向)
"""

import sys
import os
import warnings
import io
import contextlib
from unittest.mock import patch

import pytest

import outline_parser as op_module
from outline_parser import (
    extract_outline_from_docx,
    extract_outline_from_docx_with_heuristic,
    extract_outline_from_docx_via_mineru,
    _is_mineru_available,
    _log_fallback_to_audit,
    _preprocess_paragraphs,
    _strip_markdown_bold,
    reset_fallback_state,
)


# ============================================================
# Fixture: 构造临时 docx(A 路径解析用)
# ============================================================

@pytest.fixture
def tmp_docx_with_outline(tmp_path):
    """构造一个含'目录'段落的最小 docx(2 章)"""
    import docx as docx_lib
    doc = docx_lib.Document()
    doc.add_paragraph("南京大学")
    doc.add_paragraph("MBA 学位论文开题报告")
    doc.add_paragraph("姓名: 测试")
    doc.add_paragraph("论文题目: 测试论文")
    doc.add_paragraph("研究方向: 测试方向")
    doc.add_paragraph(" ")
    doc.add_paragraph("目录")
    doc.add_paragraph("第1章 绪论")
    doc.add_paragraph("1.1 研究背景")
    doc.add_paragraph("1.1.1 行业背景")
    doc.add_paragraph("1.1.2 研究意义")
    doc.add_paragraph("1.2 研究内容")
    doc.add_paragraph("第2章 理论基础")
    doc.add_paragraph("2.1 理论概述")
    doc.add_paragraph("参考文献")
    out = tmp_path / "test_outline.docx"
    doc.save(str(out))
    return str(out)


@pytest.fixture
def tmp_docx_split_chapters(tmp_path):
    """构造被 Word 拆段的章节(测试 _preprocess_paragraphs)"""
    import docx as docx_lib
    doc = docx_lib.Document()
    doc.add_paragraph("目录")
    doc.add_paragraph("第1章")
    doc.add_paragraph("绪论")
    doc.add_paragraph("1.1")
    doc.add_paragraph("研究背景")
    doc.add_paragraph("第2章")
    doc.add_paragraph("理论基础")
    doc.add_paragraph("参考文献")
    out = tmp_path / "split.docx"
    doc.save(str(out))
    return str(out)


@pytest.fixture(autouse=True)
def _reset_fallback_state_each_test():
    """每个测试前后 reset 模块状态(避免相互污染)"""
    reset_fallback_state()
    yield
    reset_fallback_state()


# ============================================================
# 基础: 模块状态 + reset
# ============================================================

def test_module_state_initialized():
    """v2.0.7 模块级状态变量初始值正确"""
    reset_fallback_state()
    assert op_module._mineru_check_done is False
    assert op_module._mineru_available is False
    assert op_module._fallback_used is False


def test_reset_fallback_state_idempotent():
    """reset 多次安全幂等"""
    op_module._fallback_used = True
    op_module._mineru_check_done = True
    op_module._mineru_available = True

    for _ in range(3):
        reset_fallback_state()
        assert op_module._fallback_used is False
        assert op_module._mineru_check_done is False
        assert op_module._mineru_available is False


def test_is_mineru_available_uses_cache():
    """F4 进程级: _is_mineru_available 检测一次后缓存"""
    call_count = {"n": 0}
    real_which = __import__("shutil").which

    def fake_which(name):
        call_count["n"] += 1
        return real_which(name)

    with patch("outline_parser.shutil.which", side_effect=fake_which):
        # 第 1 次调用: 真实检测
        r1 = _is_mineru_available()
        # 第 2 次调用: 应走缓存,不再调 which
        r2 = _is_mineru_available()
        r3 = _is_mineru_available()

    assert r1 == r2 == r3
    assert call_count["n"] == 1, f"缓存失效: which 被调用 {call_count['n']} 次"


def test_is_mineru_available_rediscover_after_reset():
    """reset 后 _is_mineru_available 会重新检测(不再用缓存)"""
    op_module._mineru_check_done = True
    op_module._mineru_available = True

    reset_fallback_state()

    with patch("outline_parser.shutil.which", return_value=None):
        result = _is_mineru_available()

    assert result is False
    assert op_module._mineru_check_done is True  # 已重新检测
    assert op_module._mineru_available is False


# ============================================================
# A 路径: 独立可用
# ============================================================

def test_a_path_extracts_outline(tmp_docx_with_outline):
    """A 路径独立可用: MinerU 不可用时直接走 A"""
    with patch.object(op_module, "_is_mineru_available", return_value=False):
        tree, issues = extract_outline_from_docx(tmp_docx_with_outline)

    assert len(tree) == 2, f"A 路径应解析 2 章,实际 {len(tree)}"
    assert tree[0]["title"] == "绪论"
    assert tree[1]["title"] == "理论基础"
    # 走 A 不应触发降级
    assert op_module._fallback_used is False


def test_a_path_handles_split_paragraphs(tmp_docx_split_chapters):
    """A 路径 + _preprocess_paragraphs 合并被拆段"""
    with patch.object(op_module, "_is_mineru_available", return_value=False):
        tree, issues = extract_outline_from_docx(tmp_docx_split_chapters)

    # 应能正确解析 2 章(尽管 Word 拆段)
    assert len(tree) == 2, f"拆段未合并: {tree}"
    assert tree[0]["title"] == "绪论"
    assert tree[1]["title"] == "理论基础"


def test_preprocess_paragraphs_rules():
    """_preprocess_paragraphs 4 条合并规则全覆盖"""
    paragraphs = [
        (0, "Normal", "第1章"),
        (1, "Normal", "绪论"),
        (2, "Normal", "1.1"),
        (3, "Normal", "研究背景"),
        (4, "Normal", "1"),
        (5, "Normal", ".1 数字开头"),
        (6, "Normal", ".1.1"),
        (7, "Normal", "三级"),
    ]
    merged = _preprocess_paragraphs(paragraphs)
    texts = [p[2] for p in merged]
    assert "第1章 绪论" in texts, f"规则 3 未触发: {texts}"
    assert "1.1 研究背景" in texts, f"规则 2 未触发: {texts}"
    assert "1.1 数字开头" in texts, f"规则 1 未触发: {texts}"
    assert ".1.1三级" in texts or ".1.1 三级" in texts, f"规则 4 未触发: {texts}"


def test_a_path_returns_no_anchor_error(tmp_path):
    """A 路径: 没有大纲锚点时返回错误"""
    import docx as docx_lib
    doc = docx_lib.Document()
    doc.add_paragraph("没有目录锚点")
    out = tmp_path / "no_anchor.docx"
    doc.save(str(out))

    with patch.object(op_module, "_is_mineru_available", return_value=False):
        tree, issues = extract_outline_from_docx(str(out))

    assert tree == []
    assert any("no_anchor" in i.get("type", "") for i in issues)


# ============================================================
# F1 = 否: 不弹窗
# ============================================================

def test_f1_no_warnings_on_fallback(tmp_docx_with_outline):
    """F1=否: MinerU 失败降级时, 不调 warnings.warn"""
    with warnings.catch_warnings(record=True) as w:
        warnings.simplefilter("always")
        with patch.object(op_module, "_is_mineru_available", return_value=True), \
             patch.object(op_module, "extract_outline_from_docx_via_mineru",
                          side_effect=RuntimeError("simulated")), \
             patch.object(op_module, "_log_fallback_to_audit"):
            tree, issues = extract_outline_from_docx(tmp_docx_with_outline)

    user_warnings = [x for x in w if issubclass(x.category, UserWarning)]
    assert user_warnings == [], f"F1 违反: 触发了 {len(user_warnings)} 个 UserWarning"


# ============================================================
# F2 = 是: 降级写 audit
# ============================================================

def test_f2_calls_audit_on_fallback(tmp_docx_with_outline):
    """F2=是: MinerU 失败时, _log_fallback_to_audit 被调用"""
    audit_calls = []

    def fake_audit(from_engine, to_engine, reason, docx_path):
        audit_calls.append({
            "from_engine": from_engine,
            "to_engine": to_engine,
            "reason": reason,
            "docx_path": docx_path,
        })

    with patch.object(op_module, "_is_mineru_available", return_value=True), \
         patch.object(op_module, "extract_outline_from_docx_via_mineru",
                      side_effect=RuntimeError("simulated failure")), \
         patch.object(op_module, "_log_fallback_to_audit", side_effect=fake_audit):
        tree, issues = extract_outline_from_docx(tmp_docx_with_outline)

    assert len(audit_calls) == 1, f"F2 违反: audit 调用 {len(audit_calls)} 次(应 1 次)"
    call = audit_calls[0]
    assert call["from_engine"] == "mineru"
    assert call["to_engine"] == "heuristic"
    assert "simulated failure" in call["reason"]
    assert call["docx_path"] == tmp_docx_with_outline


def test_f2_no_audit_when_b_succeeds(tmp_docx_with_outline):
    """F2 反向: MinerU 成功时, 不调 audit"""
    audit_calls = []

    with patch.object(op_module, "_is_mineru_available", return_value=True), \
         patch.object(op_module, "extract_outline_from_docx_via_mineru",
                      return_value=([{"title": "mock"}], [])), \
         patch.object(op_module, "_log_fallback_to_audit",
                      side_effect=lambda *a, **kw: audit_calls.append(True)):
        tree, issues = extract_outline_from_docx(tmp_docx_with_outline)

    assert audit_calls == []


def test_f2_real_audit_writes_stdout():
    """F2 真实运行: _log_fallback_to_audit 写 stdout 含 [AUDIT 标记"""
    buf = io.StringIO()
    with contextlib.redirect_stdout(buf):
        _log_fallback_to_audit(
            from_engine="mineru",
            to_engine="heuristic",
            reason="test reason",
            docx_path="/tmp/not_in_workspace/test.docx",  # 路径不在 workspace → 跳过 state 写盘
        )
    out = buf.getvalue()
    assert "[AUDIT" in out
    assert "mineru" in out
    assert "heuristic" in out
    assert "test reason" in out


# ============================================================
# F3 = 1 次: 单次重试
# ============================================================

def test_f3_single_retry(tmp_docx_with_outline):
    """F3=1: MinerU 只尝试 1 次, 失败立即降级"""
    call_count = {"n": 0}

    def fake_mineru(docx_path, language="ch", timeout=120):
        call_count["n"] += 1
        raise RuntimeError("simulated")

    with patch.object(op_module, "_is_mineru_available", return_value=True), \
         patch.object(op_module, "extract_outline_from_docx_via_mineru", side_effect=fake_mineru), \
         patch.object(op_module, "_log_fallback_to_audit"):
        tree, issues = extract_outline_from_docx(tmp_docx_with_outline)

    assert call_count["n"] == 1, f"F3 违反: MinerU 被调用 {call_count['n']} 次(应 1 次)"


# ============================================================
# F4 = 进程级: _fallback_used 全局
# ============================================================

def test_f4_fallback_locks_subsequent_calls():
    """F4=进程级: 降级后 _fallback_used=True, 后续任何调用都走 A"""
    op_module._fallback_used = True

    mineru_called = {"n": 0}

    def fake_mineru(*args, **kwargs):
        mineru_called["n"] += 1
        return [{"title": "should not reach"}], []

    with patch.object(op_module, "_is_mineru_available", return_value=True), \
         patch.object(op_module, "extract_outline_from_docx_via_mineru", side_effect=fake_mineru):
        for _ in range(3):
            extract_outline_from_docx("/tmp/fake.docx")

    assert mineru_called["n"] == 0


# ============================================================
# F5 = 跨 paper 共享
# ============================================================

def test_f5_cross_paper_shared():
    """F5=共享: 降级后跨 paper 都不再试 MinerU"""
    op_module._fallback_used = True

    mineru_called = {"n": 0}

    def fake_mineru(*args, **kwargs):
        mineru_called["n"] += 1
        return [{"title": "x"}], []

    papers = ["paper_A", "paper_B", "paper_C"]
    with patch.object(op_module, "_is_mineru_available", return_value=True), \
         patch.object(op_module, "extract_outline_from_docx_via_mineru", side_effect=fake_mineru):
        for p in papers:
            extract_outline_from_docx(f"/tmp/{p}/test.docx")

    assert mineru_called["n"] == 0


# ============================================================
# 单向降级: B→A, A 失败不回 B
# ============================================================

def test_b_to_a_unidirectional_no_revert():
    """降级语义: A 路径失败不回 B 试 MinerU"""
    with patch.object(op_module, "_is_mineru_available", return_value=False):
        mineru_called = {"n": 0}

        def fake_mineru(*args, **kwargs):
            mineru_called["n"] += 1
            return [{"title": "x"}], []

        with patch.object(op_module, "extract_outline_from_docx_via_mineru", side_effect=fake_mineru):
            # 故意传不存在路径, A 路径应失败但不回 B
            tree, issues = extract_outline_from_docx("/tmp/nonexistent.docx")

    assert mineru_called["n"] == 0, "违反单向降级: A 失败后回 B 试 MinerU"
    assert tree == []


# ============================================================
# B 路径: extract_outline_from_docx_via_mineru 自身
# ============================================================

def test_b_path_raises_when_mineru_unavailable(tmp_docx_with_outline):
    """B 路径自身: mineru 不可用时直接抛 RuntimeError"""
    with patch.object(op_module, "_is_mineru_available", return_value=False):
        with pytest.raises(RuntimeError, match="mineru-open-api 未安装"):
            extract_outline_from_docx_via_mineru(tmp_docx_with_outline)


def test_b_path_invokes_subprocess(tmp_path):
    """B 路径自身: 调用 mineru-open-api CLI subprocess"""
    import shutil as sh
    fake_md = tmp_path / "out.md"
    # 模拟 mineru 输出的 md 文本(已被 _strip_markdown_bold 清理后会被解析)
    fake_md.write_text("第1章 绪论\n1.1 研究背景\n第2章 理论基础\n2.1 理论概述", encoding="utf-8")

    fake_result = type("R", (), {
        "returncode": 0,
        "stdout": "ok",
        "stderr": "",
    })()

    def fake_run(cmd, **kwargs):
        # cmd 形式: ["mineru-open-api", "flash-extract", docx, "-o", tmpdir, "--language", "ch"]
        tmpdir_arg = None
        for i, a in enumerate(cmd):
            if a == "-o" and i + 1 < len(cmd):
                tmpdir_arg = cmd[i + 1]
                break
        if tmpdir_arg:
            sh.copy(str(fake_md), os.path.join(tmpdir_arg, "out.md"))
        return fake_result

    with patch.object(op_module, "_is_mineru_available", return_value=True), \
         patch("subprocess.run", side_effect=fake_run):
        tree, issues = extract_outline_from_docx_via_mineru("/tmp/test.docx")

    assert len(tree) >= 1
    assert tree[0]["title"] == "绪论"


def test_b_path_handles_nonzero_return(tmp_path):
    """B 路径: mineru 返回非 0 → RuntimeError 含 stderr"""
    fake_result = type("R", (), {
        "returncode": 1,
        "stdout": "",
        "stderr": "mineru crashed",
    })()

    with patch.object(op_module, "_is_mineru_available", return_value=True), \
         patch("subprocess.run", return_value=fake_result):
        with pytest.raises(RuntimeError, match="code=1"):
            extract_outline_from_docx_via_mineru("/tmp/test.docx")


def test_b_path_handles_no_md_output():
    """B 路径: mineru 成功但没生成 md → RuntimeError"""
    import subprocess
    fake_result = type("R", (), {
        "returncode": 0,
        "stdout": "",
        "stderr": "",
    })()

    with patch.object(op_module, "_is_mineru_available", return_value=True), \
         patch("subprocess.run", return_value=fake_result):
        with pytest.raises(RuntimeError, match="没生成 md"):
            extract_outline_from_docx_via_mineru("/tmp/test.docx")


# ============================================================
# 集成: 完整降级链路
# ============================================================

def test_full_fallback_chain(tmp_docx_with_outline):
    """集成: MinerU 不可用→走 A; MinerU 可用+成功→B; MinerU 失败→A 且锁"""
    # 第 1 次: MinerU 不可用 → A
    with patch.object(op_module, "_is_mineru_available", return_value=False):
        tree1, _ = extract_outline_from_docx(tmp_docx_with_outline)
    assert len(tree1) == 2
    assert op_module._fallback_used is False

    # 第 2 次: MinerU 可用+成功 → B
    with patch.object(op_module, "_is_mineru_available", return_value=True), \
         patch.object(op_module, "extract_outline_from_docx_via_mineru",
                      return_value=([{"title": "B版"}], [])):
        tree2, _ = extract_outline_from_docx(tmp_docx_with_outline)
    assert tree2 == [{"title": "B版"}]
    assert op_module._fallback_used is False

    # 第 3 次: MinerU 失败 → A 且 _fallback_used=True
    with patch.object(op_module, "_is_mineru_available", return_value=True), \
         patch.object(op_module, "extract_outline_from_docx_via_mineru",
                      side_effect=RuntimeError("now fails")), \
         patch.object(op_module, "_log_fallback_to_audit"):
        tree3, _ = extract_outline_from_docx(tmp_docx_with_outline)
    assert op_module._fallback_used is True
    assert len(tree3) == 2

    # 第 4 次: MinerU 恢复但 _fallback_used=True → 不再试
    mineru_called = {"n": 0}
    with patch.object(op_module, "_is_mineru_available", return_value=True), \
         patch.object(op_module, "extract_outline_from_docx_via_mineru",
                      side_effect=lambda *a, **kw: (mineru_called.update(n=mineru_called["n"]+ 1) or [{"title": "x"}], [])):
        tree4, _ = extract_outline_from_docx(tmp_docx_with_outline)
    assert mineru_called["n"] == 0
    assert len(tree4) == 2


# ============================================================
# 工具函数: _strip_markdown_bold
# ============================================================

def test_strip_markdown_bold():
    """_strip_markdown_bold 清理 markdown 标记"""
    assert _strip_markdown_bold("## **第1章 绪论**") == "第1章 绪论"
    assert _strip_markdown_bold("### 1.1 研究背景") == "1.1 研究背景"
    assert _strip_markdown_bold("**粗体**") == "粗体"
    assert _strip_markdown_bold("普通文本") == "普通文本"
    assert _strip_markdown_bold("  # 标题  ") == "标题"


# ============================================================
# 向后兼容: 现有调用方无破坏
# ============================================================

def test_existing_callers_signature_unchanged(tmp_docx_with_outline):
    """向后兼容: extract_outline_from_docx 签名 + 返回值类型不变"""
    with patch.object(op_module, "_is_mineru_available", return_value=False):
        result = extract_outline_from_docx(tmp_docx_with_outline)

    # 返回 (List[Dict], List[Dict]) 不变
    assert isinstance(result, tuple)
    assert len(result) == 2
    tree, issues = result
    assert isinstance(tree, list)
    assert isinstance(issues, list)


def test_existing_text_path_still_works():
    """向后兼容: extract_outline_from_text 仍可用"""
    from outline_parser import extract_outline_from_text
    text = """
第1章 绪论
1.1 研究背景
1.1.1 行业背景
第2章 理论基础
2.1 理论概述
"""
    tree, issues = extract_outline_from_text(text)
    assert len(tree) == 2
    assert tree[0]["title"] == "绪论"


def test_existing_outline_parse_still_works(tmp_docx_with_outline):
    """向后兼容: outline_parse 公共 API 仍可用"""
    from outline_parser import outline_parse

    with patch.object(op_module, "_is_mineru_available", return_value=False):
        result = outline_parse(tmp_docx_with_outline)

    assert result.get("ok") is True
    assert result["summary"]["l1"] == 2
