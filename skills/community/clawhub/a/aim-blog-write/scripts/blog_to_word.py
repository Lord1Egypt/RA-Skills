#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""博客生成 → Word 文档（含图片下载）

流程：
  1. 调用 seo_agent.sh 生成博客 JSON（SSE → 结构化）
  2. 下载所有图片到 output/<task>/images/
  3. 把 markdown 正文渲染成 .docx，图片就地嵌入；多余图片放最后画廊
"""

from __future__ import annotations

import argparse
import json
import os
import re
import subprocess
import sys
import time
from pathlib import Path
from urllib.parse import urlparse

import requests
from docx import Document
from docx.shared import Inches, RGBColor

SCRIPT_DIR = Path(__file__).resolve().parent
SKILL_ROOT = SCRIPT_DIR.parent


# ====================== 文本清洗 ======================

_THINK_BLOCK_RE = re.compile(r"<think\b[^>]*>.*?</think>", re.DOTALL | re.IGNORECASE)
_THINK_STRAY_RE = re.compile(r"</?think\b[^>]*>", re.IGNORECASE)


def strip_think(md: str) -> str:
    """剥离模型推理泄漏的 <think>...</think> 块及孤立的开/闭标签。"""
    md = _THINK_BLOCK_RE.sub("", md)
    md = _THINK_STRAY_RE.sub("", md)
    return md.lstrip()


# ====================== 调用 seo_agent.sh ======================

def run_seo_agent(payload: dict, timeout: int = 900) -> dict:
    """调用 seo_agent.sh 并解析 stdout 中的 JSON。"""
    seo = SCRIPT_DIR / "seo_agent.sh"
    if not seo.is_file():
        raise FileNotFoundError(f"seo_agent.sh 不存在: {seo}")

    proc = subprocess.run(
        ["bash", str(seo), json.dumps(payload, ensure_ascii=False)],
        capture_output=True, text=True, timeout=timeout,
    )

    out = proc.stdout.strip()
    if not out:
        raise RuntimeError(
            f"seo_agent.sh 无 stdout 输出。stderr={proc.stderr[:500]}"
        )

    # 1) 优先尝试整段当作一个 JSON
    try:
        return json.loads(out)
    except json.JSONDecodeError:
        pass

    # 2) 退而求其次：扫描每一行，找最后一个能 json.loads 的
    last = None
    for line in out.splitlines()[::-1]:
        line = line.strip()
        if line.startswith("{") and line.endswith("}"):
            try:
                last = json.loads(line)
                break
            except json.JSONDecodeError:
                continue
    if last is None:
        raise RuntimeError(
            f"seo_agent.sh 输出非 JSON。stdout 头部: {out[:300]} | stderr: {proc.stderr[:300]}"
        )
    return last


# ====================== 图片下载 ======================

_EXT_FROM_CT = {
    "image/png": ".png",
    "image/jpeg": ".jpg",
    "image/jpg": ".jpg",
    "image/gif": ".gif",
    "image/webp": ".webp",
    "image/bmp": ".bmp",
}
_VALID_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp"}


def _coerce_image_entry(entry) -> tuple[str, str]:
    """把后端返回的图片项规范化为 (url, caption)。
    后端可能返回字符串 URL，也可能返回 {image_url, rephraser_result, aspect_ratio} 字典。"""
    if isinstance(entry, str):
        return entry, ""
    if isinstance(entry, dict):
        url = entry.get("image_url") or entry.get("url") or ""
        caption = entry.get("rephraser_result") or entry.get("alt") or ""
        return url, caption
    return "", ""


def download_image(entry, dest_dir: Path, idx: int) -> tuple[Path, str] | None:
    url, caption = _coerce_image_entry(entry)
    if not url:
        print(f"[image] 第 {idx} 项无 url，跳过: {entry!r:.120}", file=sys.stderr)
        return None

    try:
        r = requests.get(
            url, timeout=60, stream=False,
            proxies={"http": None, "https": None},
        )
    except Exception as e:
        print(f"[image] 下载失败 {url}: {e}", file=sys.stderr)
        return None

    if r.status_code != 200:
        print(f"[image] {url} → HTTP {r.status_code}, 跳过", file=sys.stderr)
        return None

    # 选扩展名：URL path → Content-Type → 默认 .jpg
    ext = os.path.splitext(urlparse(url).path)[1].lower()
    if ext not in _VALID_EXTS:
        ct = r.headers.get("Content-Type", "").split(";")[0].strip().lower()
        ext = _EXT_FROM_CT.get(ct, ".jpg")

    dest = dest_dir / f"image_{idx:02d}{ext}"
    try:
        dest.write_bytes(r.content)
    except OSError as e:
        print(f"[image] 写盘失败 {dest}: {e}", file=sys.stderr)
        return None

    print(f"[image] {idx:02d} → {dest.name} ({len(r.content)} bytes)", file=sys.stderr)
    return dest, caption


# ====================== Markdown → DOCX ======================

_INLINE_PATTERN = re.compile(
    r"(\*\*[^*]+\*\*|__[^_]+__|\*[^*]+\*|_[^_]+_|`[^`]+`|\[[^\]]+\]\([^)]+\))"
)


def render_inline(paragraph, text: str) -> None:
    """渲染行内 markdown：**bold** *italic* `code` [text](url)。
    图片标记 ![alt](url) 已在外层剥离。"""
    pos = 0
    for m in _INLINE_PATTERN.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith("**") or tok.startswith("__"):
            run = paragraph.add_run(tok[2:-2])
            run.bold = True
        elif tok.startswith("`"):
            run = paragraph.add_run(tok[1:-1])
            run.font.name = "Courier New"
        elif tok.startswith("["):
            link_m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", tok)
            if link_m:
                run = paragraph.add_run(link_m.group(1))
                run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
                run.underline = True
            else:
                paragraph.add_run(tok)
        elif tok.startswith("*") or tok.startswith("_"):
            run = paragraph.add_run(tok[1:-1])
            run.italic = True
        else:
            paragraph.add_run(tok)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


_IMAGE_LINE_RE = re.compile(r"^\s*!\[[^\]]*\]\(([^)]+)\)\s*$")
_INLINE_IMAGE_RE = re.compile(r"!\[[^\]]*\]\(([^)]+)\)")
_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET_RE = re.compile(r"^\s*[-*+]\s+(.*)$")
_NUMBER_RE = re.compile(r"^\s*\d+\.\s+(.*)$")
_SPECIAL_LINE_RE = re.compile(r"^(#{1,6}\s|>|\s*[-*+]\s|\s*\d+\.\s|```|!\[)")


def _add_picture_safe(doc: Document, img_path: Path, caption: str = "") -> bool:
    try:
        doc.add_picture(str(img_path), width=Inches(5.5))
    except Exception as e:
        print(f"[docx] 嵌入图片失败 {img_path.name}: {e}", file=sys.stderr)
        return False
    if caption:
        p = doc.add_paragraph()
        run = p.add_run(caption)
        run.italic = True
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    return True


def md_to_docx(md: str, doc: Document, image_files: list[tuple[Path, str]]) -> int:
    """渲染 markdown，按出现顺序消耗 image_files，返回已嵌入的图片数。
    image_files 元素为 (本地路径, caption)。"""
    img_iter = iter(image_files)
    used = 0

    def next_image() -> tuple[Path, str] | None:
        nonlocal used
        try:
            item = next(img_iter)
            used += 1
            return item
        except StopIteration:
            return None

    in_code = False
    code_buf: list[str] = []

    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip()

        # 代码块栅栏
        if stripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buf))
                run.font.name = "Courier New"
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # 整行就是一张图
        m_img = _IMAGE_LINE_RE.match(stripped)
        if m_img:
            item = next_image()
            if item is not None:
                _add_picture_safe(doc, item[0], item[1])
            i += 1
            continue

        # 标题
        m = _HEADING_RE.match(stripped)
        if m:
            level = min(len(m.group(1)), 4)
            doc.add_heading(m.group(2).strip(), level=level)
            i += 1
            continue

        # 项目符号
        m = _BULLET_RE.match(stripped)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            render_inline(p, m.group(1))
            i += 1
            continue

        # 编号列表
        m = _NUMBER_RE.match(stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            render_inline(p, m.group(1))
            i += 1
            continue

        # 引用
        if stripped.startswith(">"):
            text = stripped[1:].lstrip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.italic = True
            i += 1
            continue

        # 空行 → 段落分隔
        if not stripped:
            i += 1
            continue

        # 段落：合并连续的非空、非特殊行
        para_lines = [stripped]
        j = i + 1
        while j < len(lines):
            nxt = lines[j].rstrip()
            if not nxt:
                break
            if _SPECIAL_LINE_RE.match(nxt):
                break
            para_lines.append(nxt)
            j += 1
        para_text = " ".join(para_lines)

        # 段落里如果嵌了 ![](url)，就把段落切片，把图片插在中间
        last = 0
        wrote_anything = False
        for m_i in _INLINE_IMAGE_RE.finditer(para_text):
            before = para_text[last:m_i.start()].strip()
            if before:
                p = doc.add_paragraph()
                render_inline(p, before)
                wrote_anything = True
            item = next_image()
            if item is not None:
                _add_picture_safe(doc, item[0], item[1])
                wrote_anything = True
            last = m_i.end()
        tail = para_text[last:].strip()
        if tail:
            p = doc.add_paragraph()
            render_inline(p, tail)
            wrote_anything = True
        if not wrote_anything:
            # 极端兜底：整段被吃干净，至少留个空行
            doc.add_paragraph("")
        i = j

    # 没用上的图片当作画廊放在末尾
    remaining = list(img_iter)
    if remaining:
        doc.add_heading("配图", level=1)
        for path, caption in remaining:
            _add_picture_safe(doc, path, caption)
            used += 1

    return used


# ====================== 主流程 ======================

def build_docx(result: dict, task_name: str, out_dir: Path) -> dict:
    if not result.get("success"):
        raise RuntimeError(f"seo_agent 返回失败: {result.get('msg', '未知错误')}")

    data = result.get("data", {}) or {}
    md = strip_think(data.get("blog_content") or "")
    keywords = data.get("keywords") or {}
    images = data.get("images") or []

    out_dir.mkdir(parents=True, exist_ok=True)
    img_dir = out_dir / "images"
    img_dir.mkdir(exist_ok=True)

    # 1) 下载图片
    image_files: list[tuple[Path, str]] = []
    for idx, entry in enumerate(images, 1):
        result_item = download_image(entry, img_dir, idx)
        if result_item:
            image_files.append(result_item)

    # 2) 构建 docx
    doc = Document()
    title_match = re.search(r"^#\s+(.*)$", md, re.MULTILINE)
    title = title_match.group(1).strip() if title_match else task_name
    doc.add_heading(title, level=0)

    # 关键词摘要
    if isinstance(keywords, dict) and keywords:
        doc.add_heading("关键词", level=1)
        for k, v in keywords.items():
            p = doc.add_paragraph()
            r1 = p.add_run(f"{k}: ")
            r1.bold = True
            p.add_run(str(v))
    elif isinstance(keywords, list) and keywords:
        doc.add_heading("关键词", level=1)
        for v in keywords:
            doc.add_paragraph(str(v), style="List Bullet")

    doc.add_heading("正文", level=1)
    embedded = md_to_docx(md, doc, image_files)

    # 把第一行 #title 从正文里去掉的视觉处理：md_to_docx 已把它再渲染了一次为 H1
    # 这里不再做删除，保留也无害（用户能在 doc 里直接编辑）

    out_file = out_dir / f"{task_name}.docx"
    doc.save(str(out_file))

    # 顺手存一份原始 JSON，便于排查
    (out_dir / "raw.json").write_text(
        json.dumps(result, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    return {
        "success": True,
        "docx": str(out_file),
        "images_dir": str(img_dir),
        "image_count_downloaded": len(image_files),
        "image_count_embedded": embedded,
        "out_dir": str(out_dir),
    }


def safe_task_name(s: str) -> str:
    cleaned = re.sub(r"[^\w一-龥]+", "_", s).strip("_")
    return cleaned[:40] or "blog"


def main() -> None:
    p = argparse.ArgumentParser(description="生成 SEO 博客并输出 Word 文稿")
    p.add_argument("--theme", required=True, help="博客主题")
    p.add_argument("--industry", required=True, help="行业（多个逗号分隔）")
    p.add_argument("--language", required=True, help="语言代码，如 en / zh")
    p.add_argument("--task-name", default="", help="任务名（默认从 theme 截取）")
    p.add_argument("--out-dir", default="", help="输出目录（默认 output/<task>_<ts>）")
    p.add_argument("--timeout", type=int, default=900, help="seo_agent.sh 调用超时（秒）")
    args = p.parse_args()

    task_name = args.task_name or safe_task_name(args.theme)
    out_dir = Path(args.out_dir) if args.out_dir else (
        SKILL_ROOT / "output" / f"{task_name}_{int(time.time())}"
    )

    payload = {
        "theme": args.theme,
        "industry": args.industry,
        "language": args.language,
    }

    print(f"[1/2] 调用 seo_agent.sh ...", file=sys.stderr)
    try:
        result = run_seo_agent(payload, timeout=args.timeout)
    except Exception as exc:
        print(json.dumps(
            {"success": False, "stage": "seo_agent", "error": str(exc)},
            ensure_ascii=False, indent=2,
        ))
        sys.exit(1)

    print(f"[2/2] 生成 Word ...", file=sys.stderr)
    try:
        out = build_docx(result, task_name, out_dir)
    except Exception as exc:
        print(json.dumps(
            {"success": False, "stage": "build_docx", "error": str(exc)},
            ensure_ascii=False, indent=2,
        ))
        sys.exit(1)

    print(json.dumps(out, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
