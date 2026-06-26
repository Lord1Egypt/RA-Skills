#!/usr/bin/env python3
"""
Markdown 转 Word 工具
====================

功能：
1. 将 Markdown 格式的薪酬报告转换为 Word 文档
2. 保留标题、列表、表格等格式
3. 支持自定义样式

使用方法：
    python3 markdown_to_word.py -i report.md -o report.docx
"""

import argparse
import re
import sys

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("❌ python-docx 未安装，请运行：pip install python-docx")
    sys.exit(1)


class MarkdownToWord:
    """Markdown 转 Word 转换器"""

    # 样式配置
    STYLES = {
        'title': {'size': 22, 'bold': True, 'color': '1F3864', 'align': 'center'},
        'h1': {'size': 18, 'bold': True, 'color': '1F3864'},
        'h2': {'size': 16, 'bold': True, 'color': '2E75B6'},
        'h3': {'size': 14, 'bold': True, 'color': '2E75B6'},
        'h4': {'size': 12, 'bold': True, 'color': '2E75B6'},
        'body': {'size': 11, 'bold': False, 'color': '333333'},
        'table_header': {'size': 10, 'bold': True, 'color': 'FFFFFF'},
        'table_cell': {'size': 10, 'bold': False, 'color': '333333'}
    }

    def __init__(self, input_file, output_file):
        self.input_file = input_file
        self.output_file = output_file
        self.doc = Document()

        # 设置默认字体
        style = self.doc.styles['Normal']
        font = style.font
        font.name = '微软雅黑'
        font.size = Pt(11)
        font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.space_before = Pt(0)

        # 设置中文字体
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    def _set_cell_shading(self, cell, color):
        """设置单元格背景色"""
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(
            qn('w:shd'),
            {
                qn('w:fill'): color,
                qn('w:val'): 'clear'
            }
        )
        shading.append(shading_elm)

    def _add_formatted_text(self, paragraph, text, style_name='body'):
        """添加格式化文本"""
        style = self.STYLES.get(style_name, self.STYLES['body'])

        # 处理粗体和斜体标记
        parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)

        for part in parts:
            run = paragraph.add_run(part.replace('**', '').replace('*', '').replace('`', ''))
            run.font.size = Pt(style['size'])
            run.font.color.rgb = RGBColor(
                int(style['color'][:2], 16),
                int(style['color'][2:4], 16),
                int(style['color'][4:], 16)
            )

            if '**' in part or style.get('bold'):
                run.bold = True
            if '*' in part:
                run.italic = True
            if '`' in part:
                run.font.name = 'Consolas'

        if 'align' in style:
            if style['align'] == 'center':
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def convert(self):
        """执行转换"""
        with open(self.input_file, 'r', encoding='utf-8') as f:
            content = f.read()

        lines = content.split('\n')
        in_table = False
        table_data = []
        in_list = False

        for line in lines:
            # 跳过空行（但保持段落间距）
            if not line.strip():
                if in_table:
                    # 表格结束，渲染表格
                    self._render_table(table_data)
                    table_data = []
                    in_table = False
                if in_list:
                    in_list = False
                continue

            # 检测表格
            if '|' in line and line.strip().startswith('|'):
                if not in_table:
                    in_table = True
                row = [cell.strip() for cell in line.split('|') if cell.strip()]
                table_data.append(row)
                continue
            elif in_table:
                # 表格结束
                self._render_table(table_data)
                table_data = []
                in_table = False

            # 标题
            if line.startswith('# '):
                p = self.doc.add_paragraph()
                self._add_formatted_text(p, line[2:], 'title')
                p.space_after = Pt(12)
            elif line.startswith('## '):
                p = self.doc.add_heading(level=1)
                self._add_formatted_text(p, line[3:], 'h1')
            elif line.startswith('### '):
                p = self.doc.add_heading(level=2)
                self._add_formatted_text(p, line[4:], 'h2')
            elif line.startswith('#### '):
                p = self.doc.add_heading(level=3)
                self._add_formatted_text(p, line[5:], 'h3')

            # 列表
            elif line.startswith('- ') or line.startswith('* '):
                p = self.doc.add_paragraph(style='List Bullet')
                self._add_formatted_text(p, line[2:], 'body')
                in_list = True

            elif re.match(r'^\d+\.\s', line):
                p = self.doc.add_paragraph(style='List Number')
                self._add_formatted_text(p, re.sub(r'^\d+\.\s', '', line), 'body')
                in_list = True

            # 引用
            elif line.startswith('> '):
                p = self.doc.add_paragraph()
                run = p.add_run(line[2:])
                run.italic = True
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

            # 普通段落
            else:
                p = self.doc.add_paragraph()
                self._add_formatted_text(p, line, 'body')

        # 处理文件末尾的表格
        if in_table and table_data:
            self._render_table(table_data)

        # 保存
        self.doc.save(self.output_file)
        print(f"✅ Word 文档已保存：{self.output_file}")

    def _render_table(self, table_data):
        """渲染表格"""
        if not table_data or len(table_data) < 2:
            return

        # 创建表格
        rows = len(table_data)
        cols = max(len(row) for row in table_data)

        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, row_data in enumerate(table_data):
            for j, cell_text in enumerate(row_data):
                if j >= cols:
                    break
                cell = table.cell(i, j)
                cell.text = cell_text

                # 设置字体
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        if i == 0:  # 表头
                            run.font.size = Pt(10)
                            run.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            self._set_cell_shading(cell, '4472C4')
                        else:
                            run.font.size = Pt(10)
                            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                            if i % 2 == 0:
                                self._set_cell_shading(cell, 'D9E2F3')


def main():
    parser = argparse.ArgumentParser(description='Markdown 转 Word 工具')
    parser.add_argument('-i', '--input', required=True, help='输入 Markdown 文件')
    parser.add_argument('-o', '--output', required=True, help='输出 Word 文件')

    args = parser.parse_args()

    converter = MarkdownToWord(args.input, args.output)
    converter.convert()


if __name__ == '__main__':
    main()
