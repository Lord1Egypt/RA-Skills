#!/usr/bin/env python3
"""Grade eval outputs for kai-docx-generator."""

import json
import os
import re
import zipfile
import sys

WORKSPACE = os.path.dirname(os.path.abspath(__file__))


def check_output_exists(eval_dir):
    """Check if .docx file was generated."""
    files = [f for f in os.listdir(eval_dir) if f.endswith(".docx")]
    return bool(files), f"Found: {files}" if files else "No .docx files found"


def check_valid_docx(eval_dir):
    """Check if output is a valid .docx (ZIP with required files)."""
    files = [f for f in os.listdir(eval_dir) if f.endswith(".docx")]
    if not files:
        return False, "No .docx file"
    path = os.path.join(eval_dir, files[0])
    if not zipfile.is_zipfile(path):
        return False, "Not a valid ZIP file"
    with zipfile.ZipFile(path) as z:
        names = z.namelist()
        required = ["word/document.xml", "[Content_Types].xml"]
        missing = [r for r in required if r not in names]
        if missing:
            return False, f"Missing: {missing}"
    return True, "Valid .docx structure"


def check_core_properties_title(eval_dir):
    """Check title in core properties."""
    from docx import Document
    files = [f for f in os.listdir(eval_dir) if f.endswith(".docx")]
    if not files:
        return False, "No .docx file"
    doc = Document(os.path.join(eval_dir, files[0]))
    title = doc.core_properties.title
    return bool(title), f"Title: {title or '(empty)'}"


def check_core_properties_author(eval_dir):
    """Check author in core properties."""
    from docx import Document
    files = [f for f in os.listdir(eval_dir) if f.endswith(".docx")]
    if not files:
        return False, "No .docx file"
    doc = Document(os.path.join(eval_dir, files[0]))
    author = doc.core_properties.author
    return bool(author), f"Author: {author or '(empty)'}"


def check_header_present(eval_dir):
    """Check header text exists."""
    from docx import Document
    files = [f for f in os.listdir(eval_dir) if f.endswith(".docx")]
    if not files:
        return False, "No .docx file"
    doc = Document(os.path.join(eval_dir, files[0]))
    for section in doc.sections:
        header_text = "".join(p.text for p in section.header.paragraphs).strip()
        if header_text:
            return True, f"Header: {header_text}"
    return False, "No header text found"


def check_toc_field_present(eval_dir):
    """Check TOC field code exists."""
    files = [f for f in os.listdir(eval_dir) if f.endswith(".docx")]
    if not files:
        return False, "No .docx file"
    path = os.path.join(eval_dir, files[0])
    with zipfile.ZipFile(path) as z:
        if "word/document.xml" in z.namelist():
            xml = z.read("word/document.xml").decode("utf-8")
            has_toc = "TOC" in xml or "toc" in xml.lower()
            return has_toc, "TOC field found" if has_toc else "No TOC field"
    return False, "Cannot read document.xml"


def check_callout_rendered(eval_dir):
    """Check blockquote/callout rendered."""
    files = [f for f in os.listdir(eval_dir) if f.endswith(".docx")]
    if not files:
        return False, "No .docx file"
    path = os.path.join(eval_dir, files[0])
    with zipfile.ZipFile(path) as z:
        if "word/document.xml" in z.namelist():
            xml = z.read("word/document.xml").decode("utf-8")
            # Callouts are rendered with special shading/background
            has_shading = "w:shd" in xml
            return has_shading, "Callout shading found" if has_shading else "No callout shading"
    return False, "Cannot read document.xml"


def check_lists_rendered(eval_dir):
    """Check lists are rendered."""
    from docx import Document
    files = [f for f in os.listdir(eval_dir) if f.endswith(".docx")]
    if not files:
        return False, "No .docx file"
    doc = Document(os.path.join(eval_dir, files[0]))
    list_count = sum(1 for p in doc.paragraphs if p.style.name and ("List" in p.style.name or "list" in p.style.name.lower()))
    return list_count > 0, f"List paragraphs: {list_count}"


def check_contract_font(eval_dir):
    """Check contract style uses FangSong."""
    files = [f for f in os.listdir(eval_dir) if f.endswith(".docx")]
    if not files:
        return False, "No .docx file"
    path = os.path.join(eval_dir, files[0])
    with zipfile.ZipFile(path) as z:
        names = z.namelist()
        found_in_doc = False
        found_in_styles = False
        if "word/document.xml" in names:
            doc_xml = z.read("word/document.xml").decode("utf-8")
            found_in_doc = "FangSong" in doc_xml or "仿宋" in doc_xml
        if "word/styles.xml" in names:
            styles_xml = z.read("word/styles.xml").decode("utf-8")
            found_in_styles = "FangSong" in styles_xml or "仿宋" in styles_xml
        has_font = found_in_doc or found_in_styles
        loc = "document.xml" if found_in_doc else ("styles.xml" if found_in_styles else "neither")
        return has_font, f"FangSong found in {loc}" if has_font else "No FangSong font reference"
    return False, "Cannot read document"


def check_bold_preserved(eval_dir):
    """Check bold formatting preserved."""
    files = [f for f in os.listdir(eval_dir) if f.endswith(".docx")]
    if not files:
        return False, "No .docx file"
    path = os.path.join(eval_dir, files[0])
    with zipfile.ZipFile(path) as z:
        if "word/document.xml" in z.namelist():
            xml = z.read("word/document.xml").decode("utf-8")
            has_bold = "w:b" in xml
            return has_bold, "Bold formatting found" if has_bold else "No bold formatting"
    return False, "Cannot read document.xml"


def check_headings_present(eval_dir):
    """Check headings rendered."""
    from docx import Document
    files = [f for f in os.listdir(eval_dir) if f.endswith(".docx")]
    if not files:
        return False, "No .docx file"
    doc = Document(os.path.join(eval_dir, files[0]))
    headings = [p for p in doc.paragraphs if p.style.name and "Heading" in p.style.name]
    return len(headings) > 0, f"Heading count: {len(headings)}"


def check_all_placeholders_replaced(eval_dir):
    """Check all template placeholders replaced (no {{...}} remain)."""
    files = [f for f in os.listdir(eval_dir) if f.endswith(".docx")]
    if not files:
        return False, "No .docx file"
    path = os.path.join(eval_dir, files[0])
    with zipfile.ZipFile(path) as z:
        if "word/document.xml" in z.namelist():
            xml = z.read("word/document.xml").decode("utf-8")
            remaining = re.findall(r"\{\{[^}]+\}\}", xml)
            return len(remaining) == 0, f"Unreplaced: {remaining}" if remaining else "All placeholders replaced"
    return False, "Cannot read document.xml"


def check_data_values_present(eval_dir):
    """Check data values appear in document text."""
    from docx import Document
    files = [f for f in os.listdir(eval_dir) if f.endswith(".docx")]
    if not files:
        return False, "No .docx file"
    doc = Document(os.path.join(eval_dir, files[0]))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    # Determine expected values based on which eval this is
    eval_name = os.path.basename(os.path.dirname(eval_dir))
    if eval_name == "eval-3":
        expected = ["技术服务合同", "北京科技", "上海软件"]
    elif eval_name == "eval-7":
        expected = ["年度维保", "城市物业", "设备维护"]
    else:
        expected = []
    found = [v for v in expected if v in full_text]
    return len(found) == len(expected), f"Found {len(found)}/{len(expected)} values: {found}"


def check_dispatch_red_h1(eval_dir):
    """Check dispatch H1 uses red color in styles."""
    files = [f for f in os.listdir(eval_dir) if f.endswith(".docx")]
    if not files:
        return False, "No .docx file"
    path = os.path.join(eval_dir, files[0])
    with zipfile.ZipFile(path) as z:
        if "word/styles.xml" in z.namelist():
            styles = z.read("word/styles.xml").decode("utf-8")
            # Red color #FF0000
            has_red = "FF0000" in styles or "ff0000" in styles.lower()
            return has_red, "Red (#FF0000) found in styles" if has_red else "No red color in styles"
    return False, "Cannot read styles.xml"


def check_strict_mode_exits_nonzero(eval_dir):
    """Check that --strict mode exits with code 1 for unfilled placeholders."""
    import subprocess
    import tempfile
    project_root = "/Users/song/projects/kai-docx-generator"
    data = '{"合同标题":"测试","甲方":"A"}'
    with tempfile.NamedTemporaryFile(suffix=".json", mode="w", delete=False, encoding="utf-8") as f:
        f.write(data)
        data_path = f.name
    try:
        result = subprocess.run(
            ["python", "scripts/fill_template.py", "templates/contract.docx",
             "--data", data_path, "--output", "/tmp/eval8_strict_test.docx", "--strict"],
            capture_output=True, text=True,
            cwd=project_root
        )
        return result.returncode == 1, f"Exit code: {result.returncode}"
    finally:
        os.unlink(data_path)


# Assertion registry
CHECKS = {
    "output_file_exists": check_output_exists,
    "valid_docx": check_valid_docx,
    "core_properties_title": check_core_properties_title,
    "core_properties_author": check_core_properties_author,
    "header_present": check_header_present,
    "toc_field_present": check_toc_field_present,
    "callout_rendered": check_callout_rendered,
    "lists_rendered": check_lists_rendered,
    "contract_font_used": check_contract_font,
    "bold_text_preserved": check_bold_preserved,
    "headings_present": check_headings_present,
    "all_placeholders_replaced": check_all_placeholders_replaced,
    "no_unreplaced_placeholders": check_all_placeholders_replaced,
    "data_values_present": check_data_values_present,
    "dispatch_red_h1": check_dispatch_red_h1,
    "strict_mode_exits_nonzero": check_strict_mode_exits_nonzero,
}


def grade_eval(eval_dir, config_name):
    """Run all assertions for an eval, save grading.json."""
    meta_path = os.path.join(eval_dir, "eval_metadata.json")
    if not os.path.exists(meta_path):
        return None

    with open(meta_path) as f:
        meta = json.load(f)

    config_dir = os.path.join(eval_dir, config_name)
    if not os.path.exists(config_dir):
        return None

    results = []
    for assertion in meta.get("assertions", []):
        name = assertion["name"]
        check_fn = CHECKS.get(name)
        if check_fn:
            passed, evidence = check_fn(config_dir)
            results.append({
                "text": assertion["description"],
                "passed": passed,
                "evidence": evidence,
            })
        else:
            results.append({
                "text": assertion["description"],
                "passed": False,
                "evidence": f"Unknown assertion: {name}",
            })

    grading = {"results": results}
    grading_path = os.path.join(config_dir, "grading.json")
    with open(grading_path, "w", encoding="utf-8") as f:
        json.dump(grading, f, indent=2, ensure_ascii=False)

    return grading


if __name__ == "__main__":
    iteration_dir = os.path.join(WORKSPACE, "iteration-1")
    if not os.path.exists(iteration_dir):
        print("Usage: python grade_evals.py")
        sys.exit(1)

    for eval_name in sorted(os.listdir(iteration_dir)):
        if not eval_name.startswith("eval-"):
            continue
        eval_dir = os.path.join(iteration_dir, eval_name)

        for config in ["with_skill", "old_skill"]:
            config_dir = os.path.join(eval_dir, config)
            if os.path.exists(config_dir):
                result = grade_eval(eval_dir, config)
                if result:
                    passed = sum(1 for r in result["results"] if r["passed"])
                    total = len(result["results"])
                    print(f"{eval_name}/{config}: {passed}/{total} passed")
                else:
                    print(f"{eval_name}/{config}: no metadata")
            else:
                print(f"{eval_name}/{config}: dir not found")
