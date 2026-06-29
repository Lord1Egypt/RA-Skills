from __future__ import annotations

from pathlib import Path

COMMON_MODULE_VERSION = "paperkb-v3.0"


def extract_text(path: str, max_chars: int = 60000) -> str:
    p = Path(path)
    ext = p.suffix.lower()
    if ext in {".txt", ".md", ".markdown"}:
        text = p.read_text(encoding="utf-8", errors="ignore")
    elif ext == ".docx":
        import docx
        doc = docx.Document(str(p))
        parts = [para.text for para in doc.paragraphs if para.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n".join(parts)
    elif ext == ".xlsx":
        import openpyxl
        wb = openpyxl.load_workbook(str(p), read_only=True, data_only=True)
        parts = []
        for ws in wb.worksheets:
            parts.append(f"# 工作表：{ws.title}")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None]
                if cells:
                    parts.append(" | ".join(cells))
        wb.close()
        text = "\n".join(parts)
    elif ext == ".xls":
        import xlrd
        book = xlrd.open_workbook(str(p))
        parts = []
        for sheet in book.sheets():
            parts.append(f"# 工作表：{sheet.name}")
            for row_idx in range(sheet.nrows):
                cells = [str(sheet.cell_value(row_idx, col_idx)) for col_idx in range(sheet.ncols) if sheet.cell_value(row_idx, col_idx) not in {"", None}]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n".join(parts)
    elif ext == ".pdf":
        import fitz
        pdf = fitz.open(str(p))
        text = "\n".join(page.get_text() for page in pdf)
        pdf.close()
    else:
        text = p.read_text(encoding="utf-8", errors="ignore")
    return (text or "")[:max_chars]
