#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
河北青山鼎信新能源科技有限公司技术尽调报告生成脚本
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from datetime import datetime

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    tcPr = cell._tc.get_or_add_tcPr()
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    tcPr.append(shading_elm)

def add_table_with_style(doc, rows, cols, data, header_color="4472C4"):
    """创建表格并应用样式"""
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = str(cell_text)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.runs[0]
            run.font.size = Pt(9)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            # 表头样式
            if i == 0:
                set_cell_shading(cell, header_color)
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
    
    return table

def set_table_width(table, width=8800):
    """设置表格宽度"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{width}" w:type="dxa"/>')
        tblPr.append(tblW)
    else:
        tblW.set(qn('w:w'), str(width))
        tblW.set(qn('w:type'), 'dxa')

def add_heading(doc, text, level):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.runs[0].font.size = Pt(16)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)
    elif level == 2:
        heading.runs[0].font.size = Pt(14)
        heading.paragraph_format.space_before = Pt(10)
        heading.paragraph_format.space_after = Pt(4)
    elif level == 3:
        heading.runs[0].font.size = Pt(12)
        heading.paragraph_format.space_before = Pt(8)
        heading.paragraph_format.space_after = Pt(3)
    return heading

def add_paragraph(doc, text):
    """添加段落"""
    para = doc.add_paragraph(text)
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.5
    for run in para.runs:
        run.font.size = Pt(10.5)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return para

def add_bullet_list(doc, items):
    """添加项目符号列表"""
    for item in items:
        para = doc.add_paragraph(item, style='List Bullet')
        para.paragraph_format.space_before = Pt(1)
        para.paragraph_format.space_after = Pt(1)
        for run in para.runs:
            run.font.size = Pt(10.5)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def create_report():
    doc = Document()
    
    # 设置文档默认字体
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(10.5)
    
    # ========== 封面 ==========
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(120)
    title_run = title.add_run('技术尽调报告')
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.name = '黑体'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(24)
    sub_run = subtitle.add_run('河北青山鼎信新能源科技有限公司')
    sub_run.font.size = Pt(20)
    sub_run.font.bold = True
    sub_run.font.name = '黑体'
    sub_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    # 日期
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.paragraph_format.space_before = Pt(400)
    date_run = date_para.add_run(f'报告日期：{datetime.now().strftime("%Y年%m月%d日")}')
    date_run.font.size = Pt(14)
    date_run.font.name = '宋体'
    date_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 保密声明
    conf_para = doc.add_paragraph()
    conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conf_para.paragraph_format.space_before = Pt(200)
    conf_run = conf_para.add_run('【保密文件】')
    conf_run.font.size = Pt(12)
    conf_run.font.color.rgb = RGBColor(192, 0, 0)
    conf_run.font.bold = True
    
    doc.add_page_break()
    
    # ========== 目录占位 ==========
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_run = toc_title.add_run('目  录')
    toc_run.font.size = Pt(18)
    toc_run.font.bold = True
    toc_run.font.name = '黑体'
    toc_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    toc_items = [
        '第一章 执行摘要',
        '第二章 公司概况',
        '第三章 第一层：材料完整性检查',
        '第四章 第二层：技术可行性分析',
        '第五章 第三层：商业逻辑验证',
        '第六章 第四层：风险识别与评级',
        '第七章 红旗信号清单',
        '第八章 综合风险评级',
        '第九章 投资建议',
        '附录'
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.size = Pt(12)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    doc.add_page_break()
    
    # ========== 第一章 执行摘要 ==========
    add_heading(doc, '第一章 执行摘要', 1)
    add_paragraph(doc, '本报告对河北青山鼎信新能源科技有限公司（以下简称"青山鼎信"或"公司"）进行了全面的技术尽调分析。公司成立于2019年8月，是河北省领先的售电企业，2024年11月海澜电力有限公司入股67%成为控股股东。本报告采用四层尽调金字塔框架，从材料完整性、技术可行性、商业逻辑和风险识别四个维度进行全面评估。')
    
    # 核心发现表格
    add_heading(doc, '1.1 核心发现概览', 2)
    data = [
        ['维度', '关键发现', '风险等级'],
        ['股东结构', '海澜电力67%控股+保定如亿33%（实控人李肥子）', '🟠中高'],
        ['员工规模', '参保仅10人，支撑1400家客户、70亿度交易量', '🔴高'],
        ['技术能力', '专利0项，仅1项商标，科创等级"入门"', '🔴高'],
        ['变更频率', '2024年密集变更（增资5倍、法人、股东、经营范围）', '🟠中高'],
        ['业务规模', '冀南60亿+冀北10亿千瓦时，头部客户中标', '🟢低'],
        ['信用评级', '启信分618，R6等级，存在一定风险', '🟡中']
    ]
    table = add_table_with_style(doc, len(data), 3, data)
    set_table_width(table)
    
    add_paragraph(doc, '综合评估结论：青山鼎信是一家具有河北售电市场基础业务的轻资产公司，在海澜电力控股后获得了品牌背书和资源支持，但存在人员规模与技术能力不匹配、股东结构复杂等风险点，建议持续关注海澜电力的实际整合程度和业务协同效果。')
    
    # ========== 第二章 公司概况 ==========
    add_heading(doc, '第二章 公司概况', 1)
    add_paragraph(doc, '本章节从公司基本信息、股权结构、融资历程和发展里程碑四个维度全面呈现青山鼎信的企业画像。')
    
    add_heading(doc, '2.1 基本信息', 2)
    data = [
        ['项目', '内容'],
        ['公司名称', '河北青山鼎信新能源科技有限公司'],
        ['成立日期', '2019年8月21日'],
        ['注册资本', '2001万元人民币'],
        ['实缴资本', '暂无公开信息'],
        ['法定代表人', '王一蔚（2024年11月变更，原为李肥子）'],
        ['统一社会信用代码', '91130101MA0E00R52T'],
        ['注册地址', '河北省石家庄市裕华区翟营南大街41号财库国际商务港1602室'],
        ['参保人数', '10人（2024年）'],
        ['企业类型', '其他有限责任公司（2024年9月变更）'],
        ['行业分类', '电力、热力、燃气及水生产和供应业'],
        ['企业经营状态', '存续'],
        ['联系方式', '0311-85266688']
    ]
    table = add_table_with_style(doc, len(data), 2, data)
    set_table_width(table)
    
    add_heading(doc, '2.2 股权结构（2024年11月变更后）', 2)
    data = [
        ['序号', '股东名称', '持股比例', '认缴出资额（万元）', '首次持股日期', '股东背景'],
        ['1', '海澜电力有限公司', '67%', '1340.67', '2024-11-13', '千亿级海澜集团旗下，注册资本2.5亿'],
        ['2', '保定如亿能源发展有限责任公司', '33%', '660.33', '2024-08-30', '实控人李肥子，注册资本仅3万元']
    ]
    table = add_table_with_style(doc, len(data), 6, data)
    set_table_width(table)
    
    add_paragraph(doc, '⚠️ 重要提示：保定如亿作为33%股东，注册资本从50万元减至3万元（2025年9月），参保人数0人，疑似壳公司，其持股动机和实际控制关系需进一步核实。')
    
    add_heading(doc, '2.3 股东背景详情', 2)
    add_paragraph(doc, '青山鼎信共有两名股东：海澜电力有限公司（67%）和保定如亿能源发展有限责任公司（33%）。海澜电力是千亿级企业海澜集团旗下公司，具备强大的资金和技术实力；保定如亿则疑似壳公司，其股东背景和持股动机需重点关注。')
    
    add_heading(doc, '2.3.1 海澜电力有限公司', 3)
    data = [
        ['项目', '内容'],
        ['企业性质', '有限责任公司（自然人投资或控股）'],
        ['成立日期', '2010年6月9日'],
        ['注册资本', '2.5亿元人民币（实缴）'],
        ['实缴资本', '2.5亿元人民币'],
        ['法定代表人', '陈炯'],
        ['注册地址', '江苏省江阴市滨江东路2号17楼'],
        ['参保人数', '133人'],
        ['股东结构', '海澜智云科技有限公司80% + 江阴德誉电力科技有限公司20%'],
        ['主营业务', '智慧售电、电力工程安装运维、虚拟电厂、综合能源服务'],
        ['覆盖区域', '江苏、浙江、上海、山东、河南、安徽、山西等省市'],
        ['年交易电量', '连续5年全国交易电量达500亿千瓦时'],
        ['行业地位', '民营独立售电公司头部企业'],
        ['企业资质', '电力工程施工总承包贰级、承装承修承试叁级、电力工程设计乙级']
    ]
    table = add_table_with_style(doc, len(data), 2, data)
    set_table_width(table)
    
    add_paragraph(doc, '海澜电力系海澜集团旗下企业，海澜集团是营收超1400亿元的千亿级民营集团，连续14年蝉联中国民营企业500强。集团涵盖服饰零售、文体旅游、智能低碳、商业管理等产业板块。')
    
    add_heading(doc, '2.3.2 保定如亿能源发展有限责任公司', 3)
    data = [
        ['项目', '内容'],
        ['企业性质', '有限责任公司（自然人投资或控股）'],
        ['成立日期', '2023年9月19日'],
        ['原始注册资本', '50万元'],
        ['现注册资本', '3万元（2025年9月减资）'],
        ['实缴资本', '暂无公开信息'],
        ['法定代表人', '李肥子'],
        ['股东结构', '李肥子70% + 曹建楼30%'],
        ['参保人数', '0人'],
        ['主营业务', '新兴能源技术研发、技术服务、技术推广'],
        ['对外投资', '仅青山鼎信一家'],
        ['企业地址', '河北省保定市高开区朝阳北大街2238号汇博上谷大观C座622室']
    ]
    table = add_table_with_style(doc, len(data), 2, data)
    set_table_width(table)
    
    add_paragraph(doc, '⚠️ 风险提示：保定如亿成立于2023年9月，注册资本仅3万元，参保人数为0，无实际经营迹象，疑似壳公司或代持。其在青山鼎信的33%持股比例与海澜电力的67%形成制衡，需关注是否存在关联关系或代持安排。')
    
    add_heading(doc, '2.4 融资历程与股权演变', 2)
    data = [
        ['时间', '事件', '股权变化', '备注'],
        ['2019-08-21', '公司成立', '曹淑英100%独资，注册资本400万元', '自然人独资企业'],
        ['2024-08-30', '首次引入股东', '曹淑英+保定如亿', '类型变更为其他有限责任公司'],
        ['2024-09-03', '注册资本增资', '400万→2001万元（+1601万）', '保定如亿认缴1601万元'],
        ['2024-09-03', '企业类型变更', '自然人独资→其他有限责任公司', '曹淑英股权被稀释'],
        ['2024-11-01', '曹淑英退出', '保定如亿承接曹淑英股权', '曹淑英完全退出'],
        ['2024-11-13', '海澜电力入股', '海澜电力67%控股', '正式完成工商变更登记'],
        ['2024-11-16', '法定代表人变更', '李肥子→王一蔚', '海澜电力派驻'],
        ['2024-11-16', '经营范围大幅扩展', '新增储能、光伏、物联网、电信业务等', '从售电扩展至综合能源']
    ]
    table = add_table_with_style(doc, len(data), 4, data)
    set_table_width(table)
    
    add_heading(doc, '2.5 发展里程碑', 2)
    add_bullet_list(doc, [
        '2019年8月：公司成立，定位为新能源技术服务商',
        '2019年11月：被列为河北新增2家售电公司之一',
        '2023年12月：中标邢台区域直购电项目',
        '2024年1月：注册首个商标"青山鼎信"',
        '2024年9月：注册资本增资5倍至2001万元，引入保定如亿',
        '2024年11月：海澜电力67%控股，正式完成工商变更',
        '2024年11月：经营范围从售电扩展至储能、光伏、建设工程等综合能源服务',
        '2024年12月：连续中标4个购售电项目（冀东、金隅等大客户）',
        '2025年4月：公司被认定为科技型中小企业'
    ])
    
    add_heading(doc, '2.6 经营范围（2024年11月最新）', 2)
    add_paragraph(doc, '一般项目：新兴能源技术研发；储能技术服务；合同能源管理；节能管理服务；物联网技术服务；物联网应用服务；信息系统集成服务；技术服务、技术开发、技术咨询、技术交流、技术转让、技术推广；市政设施管理；灯具销售；工程管理服务；专业设计服务；工业工程设计服务；光伏发电设备租赁；充电控制设备租赁；设备监理服务；企业管理咨询；贸易经纪；社会经济咨询服务；数据处理和存储支持服务；制冷、空调设备销售；通用设备修理；专用设备修理；电气设备修理。')
    add_paragraph(doc, '许可项目：供电业务；发电业务、输电业务、供（配）电业务；第一类增值电信业务；第二类增值电信业务；建设工程设计；建设工程施工；输电、供电、受电电力设施的安装、维修和试验；特种设备安装改造修理。')
    
    # ========== 第三章 第一层：材料完整性检查 ==========
    add_heading(doc, '第三章 第一层：材料完整性检查', 1)
    add_paragraph(doc, '本章节从团队背景、专利/资质、备案信息、负面信息四个维度核查青山鼎信的材料完整性和真实性。')
    
    add_heading(doc, '3.1 团队背景核查', 2)
    data = [
        ['姓名', '职位', '背景信息', '关联企业', '核查结论'],
        ['王一蔚', '法定代表人、执行董事', '2024年11月由海澜电力派驻', '海澜电力关联企业2家', '🟢可信'],
        ['陈炯', '总经理（海澜电力）', '海澜电力法定代表人、总经理', '海澜智云科技', '🟢可信'],
        ['李肥子', '原法定代表人（已退出）', '保定如亿实控人，70%持股', '保定如亿', '🟡待核实'],
        ['曹建楼', '监事（保定如亿）', '保定如亿30%股东', '保定如亿', '🟡待核实']
    ]
    table = add_table_with_style(doc, len(data), 5, data)
    set_table_width(table)
    
    add_paragraph(doc, '核心管理层中，王一蔚、陈炯来自海澜电力，具有电力行业背景。但公司仅10名参保员工，如何支撑1400家客户的运营服务存疑。李肥子虽退出管理层，但通过保定如亿仍持股33%，需关注其实际影响力。')
    
    add_heading(doc, '3.2 专利/资质核查', 2)
    data = [
        ['类型', '数量', '详情', '核查结论'],
        ['专利', '0项', '无任何专利申请记录', '🔴存疑'],
        ['商标', '1项', '"青山鼎信"（2024-01-22注册，国际分类39类-运输贮藏）', '🟢已确认'],
        ['著作权', '0项', '无软件著作权登记', '🔴存疑'],
        ['电力业务资质', '待核实', '需核查供电业务许可证等', '⚠️待核实'],
        ['建筑业企业资质', '待核实', '2024年11月新增建设工程施工范围', '⚠️待核实']
    ]
    table = add_table_with_style(doc, len(data), 4, data)
    set_table_width(table)
    
    add_paragraph(doc, '⚠️ 重大发现：公司科创等级为"入门"，全国行业排名后99%，0专利、0著作权，仅有1项商标。科技属性薄弱，与"新能源科技"公司定位不符。')
    
    add_heading(doc, '3.3 备案信息核查', 2)
    data = [
        ['备案类型', '状态', '核查渠道', '结论'],
        ['工商登记', '✅ 已备案', '国家企业信用信息公示系统', '🟢可信'],
        ['税务登记', '✅ 增值税一般纳税人', '水滴信用', '🟢可信'],
        ['社保缴纳', '✅ 10人参保', '企查查', '🟢可信'],
        ['电力交易资格', '✅ 河北电力交易中心注册', '河北日报报道', '🟢可信'],
        ['科技型中小企业', '✅ 已认定（2025年4月）', '企查查', '🟢可信']
    ]
    table = add_table_with_style(doc, len(data), 4, data)
    set_table_width(table)
    
    add_heading(doc, '3.4 负面信息核查', 2)
    data = [
        ['负面类型', '核查结果', '严重程度', '风险评估'],
        ['法律诉讼', '暂无公开诉讼记录', '无', '🟢低风险'],
        ['被执行人', '无', '无', '🟢低风险'],
        ['经营异常', '无', '无', '🟢低风险'],
        ['行政处罚', '无', '无', '🟢低风险'],
        ['严重违法失信', '无', '无', '🟢低风险'],
        ['客户投诉/纠纷', '暂无公开信息', '未知', '⚠️待核实']
    ]
    table = add_table_with_style(doc, len(data), 4, data)
    set_table_width(table)
    
    add_paragraph(doc, '公司暂无公开负面记录，信用状况基本良好。但需注意，公司参保人数仅10人，服务1400家客户，人效比极高，需确认是否存在灵活用工或外包服务模式。')
    
    # ========== 第四章 第二层：技术可行性分析 ==========
    add_heading(doc, '第四章 第二层：技术可行性分析', 1)
    add_paragraph(doc, '本章节从售电技术能力、储能技术、虚拟电厂、综合能源管理四个维度评估青山鼎信的技术可行性。鉴于海澜电力已控股67%，部分技术能力应视为可整合资源。')
    
    add_heading(doc, '4.1 售电技术能力评估', 2)
    data = [
        ['能力维度', '青山鼎信现状', '海澜电力能力', '综合评估'],
        ['交易系统', '基础售电系统', '智慧能源运维管理平台，24小时监测', '🟠依赖海澜'],
        ['客户规模', '1400家客户', '签约客户均衡分布于各行业', '🟢已验证'],
        ['年交易量', '冀南60亿+冀北10亿千瓦时', '连续5年500亿千瓦时全国领先', '🟢行业头部'],
        ['偏差考核', '需核实', '帮助客户实现售电零风险零考核', '🟠依赖海澜'],
        ['绿电绿证', '有开展', '年交易绿电500亿千瓦时', '🟠依赖海澜']
    ]
    table = add_table_with_style(doc, len(data), 4, data)
    set_table_width(table)
    
    add_paragraph(doc, '青山鼎信在河北市场具有明确的客户基础和交易量，但售电技术系统大概率依赖海澜电力的平台支持。建议核查公司是否具备独立的电力交易系统和技术团队。')
    
    add_heading(doc, '4.2 储能技术能力评估', 2)
    data = [
        ['维度', '评估', '依据'],
        ['技术积累', '🔴无', '0专利，无储能相关软著'],
        ['项目经验', '⚠️待核实', '需核查是否有储能项目落地'],
        ['海澜支持', '🟢有', '海澜电力具有储能技术服务资质'],
        ['市场定位', '🟡战略布局', '2024年11月新增储能技术服务范围']
    ]
    table = add_table_with_style(doc, len(data), 3, data)
    set_table_width(table)
    
    add_paragraph(doc, '储能业务属于战略性布局，目前无自主技术能力。海澜电力具备储能项目经验，可作为技术支持方。公司在储能领域的发展将高度依赖海澜电力的技术输出。')
    
    add_heading(doc, '4.3 虚拟电厂能力评估', 2)
    data = [
        ['维度', '青山鼎信', '海澜电力', '行业对比'],
        ['虚拟电厂资质', '⚠️待核实', '✅ 有', '头部企业均有布局'],
        ['聚合资源规模', '⚠️待核实', '✅ 有', '行业规模达GW级'],
        ['市场参与', '⚠️待核实', '✅ 有', '冀北已有示范项目'],
        ['技术平台', '⚠️待核实', '✅ 有', '需持续投入研发']
    ]
    table = add_table_with_style(doc, len(data), 4, data)
    set_table_width(table)
    
    add_paragraph(doc, '虚拟电厂行业2024年市场规模已突破200亿元，预计2030年达千亿级，年复合增长率超30%。海澜电力已布局虚拟电厂业务，但青山鼎信自身暂无明确能力。河北省冀北电网虚拟电厂已进入市场型阶段，公司可借势发展。')
    
    add_heading(doc, '4.4 综合能源管理能力评估', 2)
    data = [
        ['能力领域', '现状', '海澜支持', '发展潜力'],
        ['能耗管理', '⚠️待核实', '✅ 有', '🟡一般'],
        ['节能改造', '⚠️待核实', '✅ 有', '🟡一般'],
        ['电力工程', '⚠️待核实', '✅ 电力总包二级', '🟠依赖海澜'],
        ['运维服务', '⚠️待核实', '✅ 有', '🟠依赖海澜']
    ]
    table = add_table_with_style(doc, len(data), 4, data)
    set_table_width(table)
    
    add_paragraph(doc, '综合能源管理是售电公司的升级方向。海澜电力具备电力工程施工总承包贰级资质、承装承修承试叁级资质，可提供"一站式"服务。青山鼎信在综合能源领域的能力将主要依赖海澜电力的资质和团队支持。')
    
    # ========== 第五章 第三层：商业逻辑验证 ==========
    add_heading(doc, '第五章 第三层：商业逻辑验证', 1)
    add_paragraph(doc, '本章节从商业模式画布、竞争格局、市场规模、估值分析四个维度验证青山鼎信的商业逻辑。')
    
    add_heading(doc, '5.1 商业模式画布', 2)
    data = [
        ['模块', '内容'],
        ['价值主张', '为河北工业用户提供智慧售电、绿电绿证、综合能源管理等服务'],
        ['客户细分', '用电客户1400家，以冀南冀北工商业用户为主，中标冀东水泥、金隅等大客户'],
        ['渠道通路', '河北电力交易中心、招投标、大客户直销'],
        ['客户关系', '长期代理购电关系，托管交易、聚合交易'],
        ['收入来源', '售电价差收入+绿电绿证交易+综合能源服务费'],
        ['核心资源', '售电资质、客户资源、海澜电力品牌背书'],
        ['关键活动', '电力交易、需求响应、客户服务、项目中标'],
        ['关键合作', '海澜电力（67%控股）、发电企业、河北电力交易中心'],
        ['成本结构', '购电成本、人力成本、平台运维成本、偏差考核成本']
    ]
    table = add_table_with_style(doc, len(data), 2, data)
    set_table_width(table)
    
    add_paragraph(doc, '商业模式分析：青山鼎信的售电业务模式属于行业主流的"价差盈利"模式，面临"1439号文"后工商业用户全面入市、价格竞争白热化的行业压力。单一价差模式难以维系，需向"售电+"综合服务转型。')
    
    add_heading(doc, '5.2 竞争格局分析', 2)
    add_paragraph(doc, '青山鼎信面临的竞争格局呈现"内外双重压力"的特点：')
    
    add_heading(doc, '5.2.1 全国售电市场竞争格局', 3)
    add_paragraph(doc, '中国售电市场呈现"数量上民营为主，市场份额上国资主导"的倒挂格局：')
    data = [
        ['主体类型', '数量占比', '市场份额', '代表企业'],
        ['电网系售电公司', '约15%', '45%', '国网新源、南网能源'],
        ['发电系售电公司', '约15%', '30%', '华能能源销售、大唐售电'],
        ['地方国资售电公司', '约10%', '10%', '浙能售电、粤电售电'],
        ['民营/外资售电公司', '约60%', '15%', '海澜电力、壳牌能源']
    ]
    table = add_table_with_style(doc, len(data), 4, data)
    set_table_width(table)
    
    add_paragraph(doc, '青山鼎信属于民营售电公司，在河北市场面临电网系和发电系售电公司的激烈竞争。海澜电力控股后，可借助其全国交易量和品牌背书提升竞争力。')
    
    add_heading(doc, '5.2.2 河北售电市场格局', 3)
    data = [
        ['维度', '数据'],
        ['售电公司数量', '223家（2024年3月），位居全国第三'],
        ['2024年市场交易量', '发电侧1205.69亿千瓦时，用户侧916亿千瓦时'],
        ['结算均价', '发电侧431.09元/兆瓦时，用户侧413.38元/兆瓦时'],
        ['青山鼎信份额', '冀南约60亿千瓦时+冀北约10亿千瓦时，约占6-7%']
    ]
    table = add_table_with_style(doc, len(data), 2, data)
    set_table_width(table)
    
    add_heading(doc, '5.3 市场规模分析', 2)
    data = [
        ['市场维度', '全国数据', '河北省数据'],
        ['2024年市场化交易电量', '6.2万亿千瓦时（占全社会用电量63%）', '1137亿千瓦时（2023年）'],
        ['2024年跨省跨区交易电量', '1.4万亿千瓦时', '-'],
        ['用户侧零售市场', '3.6万亿千瓦时（占市场化电量58%）', '-'],
        ['绿证绿电交易', '4460亿千瓦时（同比增长364%）', '97.8万张绿证（2023年）'],
        ['新能源装机', '14.5亿千瓦（占总装机43%）', '清洁能源上网435.99亿千瓦时']
    ]
    table = add_table_with_style(doc, len(data), 3, data)
    set_table_width(table)
    
    add_paragraph(doc, '电力市场空间广阔，河北省2024年市场化交易电量超千亿千瓦时，为青山鼎信提供了充足的市场容量。随着全国统一电力市场加速建设，市场化交易比例将持续提升。')
    
    add_heading(doc, '5.4 估值分析', 2)
    data = [
        ['估值因素', '分析', '参考值'],
        ['注册资本', '2001万元', '门槛较低'],
        ['年交易电量', '70亿千瓦时', '行业头部水平'],
        ['客户数量', '1400家', '规模较大'],
        ['股东背景', '海澜电力67%控股', '品牌背书强'],
        ['净利润', '暂无公开数据', '⚠️需核实'],
        ['营收规模', '暂无公开数据', '⚠️需核实'],
        ['行业估值', '售电公司一般按年交易电量×一定系数估值', '约0.01-0.03元/度']
    ]
    table = add_table_with_style(doc, len(data), 3, data)
    set_table_width(table)
    
    add_paragraph(doc, '由于缺乏公开财务数据，青山鼎信的实际估值难以确定。建议要求提供近三年财务报表，结合利润率和现金流进行估值分析。以年交易量70亿千瓦时计算，若按0.01元/度系数估值，估值约7000万元。')
    
    # ========== 第六章 第四层：风险识别与评级 ==========
    add_heading(doc, '第六章 第四层：风险识别与评级', 1)
    add_paragraph(doc, '本章节从技术风险、市场风险、团队风险、财务风险、法律风险五个维度识别青山鼎信的潜在风险，并构建风险矩阵。')
    
    add_heading(doc, '6.1 技术风险', 2)
    data = [
        ['风险点', '具体描述', '概率', '影响', '等级'],
        ['技术能力薄弱', '0专利、0软著，科创等级入门', '极高', '高', '🔴高'],
        ['核心技术依赖海澜', '售电系统、储能、虚拟电厂均依赖海澜支持', '极高', '中', '🟠中高'],
        ['系统安全风险', '电力交易系统稳定性要求高', '低', '高', '🟡中']
    ]
    table = add_table_with_style(doc, len(data), 5, data)
    set_table_width(table)
    
    add_heading(doc, '6.2 市场风险', 2)
    data = [
        ['风险点', '具体描述', '概率', '影响', '等级'],
        ['竞争加剧', '223家售电公司竞争，电网系/发电系主导市场', '极高', '高', '🔴高'],
        ['价差收窄', '1439号文后工商业全面入市，价格竞争白热化', '极高', '高', '🔴高'],
        ['偏差考核', '若负荷预测不准可能产生高额考核成本', '中', '中', '🟠中高'],
        ['客户流失', '大客户可能自建售电或被竞争对手抢夺', '中', '中', '🟡中']
    ]
    table = add_table_with_style(doc, len(data), 5, data)
    set_table_width(table)
    
    add_heading(doc, '6.3 团队风险', 2)
    data = [
        ['风险点', '具体描述', '概率', '影响', '等级'],
        ['人员规模不足', '10人支撑1400家客户、70亿度交易', '极高', '高', '🔴高'],
        ['团队能力存疑', '科创等级入门，无专利技术', '高', '中', '🟠中高'],
        ['核心人员流失', '关键岗位人员变动可能影响运营', '中', '中', '🟡中']
    ]
    table = add_table_with_style(doc, len(data), 5, data)
    set_table_width(table)
    
    add_heading(doc, '6.4 财务风险', 2)
    data = [
        ['风险点', '具体描述', '概率', '影响', '等级'],
        ['盈利能力不明', '暂无公开财务数据', '高', '高', '🟠中高'],
        ['现金流风险', '售电行业垫资普遍，现金流压力大', '中', '中', '🟡中'],
        ['关联交易风险', '与海澜电力的关联交易定价待核实', '中', '中', '🟡中']
    ]
    table = add_table_with_style(doc, len(data), 5, data)
    set_table_width(table)
    
    add_heading(doc, '6.5 法律风险', 2)
    data = [
        ['风险点', '具体描述', '概率', '影响', '等级'],
        ['保定如亿壳公司风险', '注册资本3万、0人参保的33%股东', '高', '高', '🟠中高'],
        ['股权代持风险', '李肥子通过保定如亿持股33%，实控关系待核实', '中', '高', '🟠中高'],
        ['资质合规风险', '新增电信业务、建设工程等资质待核实', '中', '中', '🟡中'],
        ['合同风险', '客户合同条款、代理授权合规性', '低', '中', '🟡中']
    ]
    table = add_table_with_style(doc, len(data), 5, data)
    set_table_width(table)
    
    # ========== 第七章 红旗信号清单 ==========
    add_heading(doc, '第七章 红旗信号清单', 1)
    add_paragraph(doc, '以下为本次尽调中发现的重大风险信号，需要重点关注或进一步核实：')
    
    add_heading(doc, '7.1 严重红旗（🔴需立即关注）', 2)
    add_bullet_list(doc, [
        '【团队能力】参保仅10人，却服务1400家客户、年交易量70亿千瓦时，人效比极度异常，需核实是否有外包/派遣模式或人员隐藏在母公司',
        '【科技创新】0专利、0软著、科创等级"入门"，与"新能源科技"公司定位严重不符',
        '【股权结构】保定如亿注册资本从50万减至3万元、0人参保，33%持股比例存在壳公司嫌疑',
        '【实控关系】李肥子虽退出管理层，但通过保定如亿持股33%，实际控制关系待核实'
    ])
    
    add_heading(doc, '7.2 重要红旗（🟠需重点关注）', 2)
    add_bullet_list(doc, [
        '【变更频繁】2024年密集发生增资5倍、法人变更、股东变更、经营范围大扩等重大变更，变更动机需分析',
        '【业务依赖】售电技术能力大概率依赖海澜电力，独立性存疑，整合深度待核实',
        '【财务不明】公司暂无公开财务数据，盈利能力和现金流状况无法评估',
        '【竞争激烈】河北223家售电公司，民营公司份额仅15%，生存压力大'
    ])
    
    add_heading(doc, '7.3 一般红旗（🟡需持续关注）', 2)
    add_bullet_list(doc, [
        '【行业趋势】"价差售电"模式难以为继，需向"售电+"综合服务转型',
        '【客户集中度】中标冀东、金隅等大客户，需关注客户集中度风险',
        '【关联交易】与海澜电力的购售电交易定价合理性待核实',
        '【政策变化】电力体制改革政策变化可能影响业务模式'
    ])
    
    # ========== 第八章 综合风险评级 ==========
    add_heading(doc, '第八章 综合风险评级', 1)
    add_paragraph(doc, '本章节基于四层尽调发现，对青山鼎信进行综合风险评级。')
    
    add_heading(doc, '8.1 风险评级维度', 2)
    data = [
        ['评级维度', '评分', '满分', '得分率', '风险等级', '说明'],
        ['材料完整性', '55', '100', '55%', '🟠中', '股东结构复杂，资质待核实'],
        ['技术可行性', '45', '100', '45%', '🟠中高', '核心技术依赖海澜，自主能力弱'],
        ['商业逻辑', '70', '100', '70%', '🟡中低', '商业模式可行，市场空间大'],
        ['风险识别', '50', '100', '50%', '🟠中高', '多项重大风险信号']
    ]
    table = add_table_with_style(doc, len(data), 6, data)
    set_table_width(table)
    
    add_heading(doc, '8.2 综合风险矩阵', 2)
    data = [
        ['风险类型', '严重程度', '发生概率', '风险值(RN)', '综合评级'],
        ['团队能力不足', '高(5)', '极高(5)', '25', '🔴极高'],
        ['股权结构复杂', '高(5)', '高(4)', '20', '🔴高'],
        ['技术能力薄弱', '高(5)', '极高(5)', '25', '🔴极高'],
        ['市场竞争激烈', '高(5)', '极高(5)', '25', '🔴极高'],
        ['价差收窄压力', '高(5)', '极高(5)', '25', '🔴极高'],
        ['盈利能力不明', '高(5)', '高(4)', '20', '🔴高'],
        ['财务风险', '中(3)', '中(3)', '9', '🟠中高'],
        ['法律合规', '中(3)', '中(3)', '9', '🟠中高']
    ]
    table = add_table_with_style(doc, len(data), 5, data)
    set_table_width(table)
    
    add_heading(doc, '8.3 最终评级结论', 2)
    data = [
        ['评级维度', '结论', '依据'],
        ['综合风险评级', '🟠中高风险（BB级）', '多项重大风险信号，需深入尽调'],
        ['投资建议', '⚠️ 有条件投资', '建议完成以下尽调后决策：'],
        ['核心关注点', '1. 10人团队支撑业务真实性', '灵活用工/外包情况'],
        ['核心关注点', '2. 与海澜电力的整合深度', '独立运营还是完全依赖'],
        ['核心关注点', '3. 保定如亿实际控制关系', '是否存在代持安排'],
        ['核心关注点', '4. 近三年财务报表', '盈利能力、现金流核实'],
        ['核心关注点', '5. 主要客户合同条款', '服务期限、违约金等']
    ]
    table = add_table_with_style(doc, len(data), 3, data)
    set_table_width(table)
    
    # ========== 第九章 投资建议 ==========
    add_heading(doc, '第九章 投资建议', 1)
    add_paragraph(doc, '基于本次技术尽调的四层评估框架，提出如下投资建议：')
    
    add_heading(doc, '9.1 投资建议', 2)
    data = [
        ['建议类型', '内容', '优先级'],
        ['🟢 推进条件', '要求公司提供近三年财务报表、10人团队劳动关系说明、与海澜电力的整合协议', 'P0'],
        ['🟢 推进条件', '核实保定如亿的实际控制关系和代持安排', 'P0'],
        ['🟢 推进条件', '核查电力业务许可证、建筑业企业资质等核心资质', 'P1'],
        ['🟢 推进条件', '访谈核心团队成员，评估实际运营能力', 'P1'],
        ['🟡 关注事项', '评估"售电+"转型进展和综合能源服务能力建设', 'P2'],
        ['🟡 关注事项', '核实中标项目的合同金额和服务条款', 'P2'],
        ['🟡 关注事项', '了解与海澜电力的关联交易定价机制', 'P2']
    ]
    table = add_table_with_style(doc, len(data), 3, data)
    set_table_width(table)
    
    add_heading(doc, '9.2 风险缓释建议', 2)
    add_bullet_list(doc, [
        '【股权方面】建议要求保定如亿如实披露股权代持情况，或要求李肥子出具不谋求控制权的承诺函',
        '【团队方面】要求提供完整员工名册和社保记录，核实劳动关系真实性',
        '【技术方面】要求海澜电力出具书面支持承诺，明确技术输出范围和期限',
        '【财务方面】要求提供审计报告，重点关注关联交易定价和利润转移情况',
        '【合同方面】核查主要客户合同，确认服务期限和违约金条款'
    ])
    
    add_heading(doc, '9.3 投资决策建议', 2)
    data = [
        ['场景', '条件', '建议'],
        ['A', '完成P0尽调，核心风险已核实或已缓释', '✅ 可进入投资决策'],
        ['B', '发现股权代持或实控关系不清晰', '❌ 建议终止或重新谈判'],
        ['C', '10人团队存在外包/派遣，非真实劳动关系', '⚠️ 重新评估公司价值和风险'],
        ['D', '财务数据盈利能力不足或现金流持续为负', '⚠️ 谨慎决策或调整估值']
    ]
    table = add_table_with_style(doc, len(data), 3, data)
    set_table_width(table)
    
    # ========== 附录 ==========
    add_heading(doc, '附录', 1)
    add_paragraph(doc, '本附录包含信息核查清单、签字栏、数据来源和免责声明等内容。')
    
    add_heading(doc, '附录A：信息核查清单', 2)
    data = [
        ['核查项', '状态', '说明'],
        ['工商信息', '✅ 已核查', '国家企业信用信息公示系统'],
        ['股东信息', '✅ 已核查', '企查查、启信宝'],
        ['专利信息', '✅ 已核查', '国家知识产权局'],
        ['商标信息', '✅ 已核查', '国家知识产权局'],
        ['软件著作权', '✅ 已核查', '中国版权保护中心'],
        ['电力交易资质', '⚠️ 待核实', '需获取相关证书'],
        ['建筑业企业资质', '⚠️ 待核实', '需获取相关证书'],
        ['财务报表', '⚠️ 待获取', '需公司提供审计报告'],
        ['客户合同', '⚠️ 待获取', '需公司提供样本合同'],
        ['员工名册', '⚠️ 待获取', '需公司提供']
    ]
    table = add_table_with_style(doc, len(data), 3, data)
    set_table_width(table)
    
    add_heading(doc, '附录B：签字栏', 2)
    data = [
        ['角色', '姓名', '签字', '日期'],
        ['项目负责人', '（待填写）', '____________', '2026年__月__日'],
        ['尽调分析师', '（待填写）', '____________', '2026年__月__日'],
        ['法务审核', '（待填写）', '____________', '2026年__月__日'],
        ['财务审核', '（待填写）', '____________', '2026年__月__日'],
        ['风控审核', '（待填写）', '____________', '2026年__月__日']
    ]
    table = add_table_with_style(doc, len(data), 4, data)
    set_table_width(table)
    
    add_heading(doc, '附录C：数据来源', 2)
    add_bullet_list(doc, [
        '企查查（https://www.qcc.com）- 工商信息查询',
        '启信宝（https://www.qixin.com）- 企业信用信息',
        '水滴信用（https://shuidi.cn）- 企业信息聚合',
        '国家企业信用信息公示系统 - 工商登记信息',
        '河北电力交易中心 - 电力交易数据',
        '国家能源局 - 政策文件',
        '中国电力企业联合会 - 行业数据',
        '海澜电力官网（https://www.heilanpower.com）- 股东信息',
        '海澜集团官网（http://www.heilan.com.cn）- 集团信息',
        '河北新闻网 - 新闻报道'
    ])
    
    add_heading(doc, '附录D：免责声明', 2)
    add_paragraph(doc, '本报告仅基于公开可获取的信息编制，不构成任何投资建议。报告中的分析、评级和建议仅供参考，不对报告使用人或第三方因使用本报告而产生的任何损失承担责任。投资决策应由投资方在充分尽调和独立判断的基础上做出。本报告禁止未经授权的传播和引用。')
    
    # 保存文档
    output_path = '河北青山鼎信技术尽调报告.docx'
    doc.save(output_path)
    print(f"✅ 报告已生成：{output_path}")
    return output_path

if __name__ == "__main__":
    create_report()
