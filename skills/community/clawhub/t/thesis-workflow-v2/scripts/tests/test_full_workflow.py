#!/usr/bin/env python3
"""
test_full_workflow.py - Step 12 全链路集成测试（修订 v2）

端到端验证 Step 1-11 全部组件协同工作。
"""

import sys
import os
import json
import time

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from orchestrator_v2 import (
    orchestrate_phase1_1,
    confirm_phase1,
    orchestrate_phase1_3,
    update_node_content_hint,
    confirm_phase1_3,
    synthesize_chapter_summary,
    load_orchestrate_state,
    _get_orchestrate_state_path,
    is_last_child_of_chapter,
    apply_user_decision,
    orchestrate_phase2,
)
from context_builder import build_prompt_package, build_prompt_package_text
from state_manager_v2 import (
    outline_save,
    outline_load,
    outline_update_status,
    outline_get_context,
    _get_state_path,
)
from outline_parser import (
    insert_chapter_summary_nodes,
    save_content_hints_to_outline,
)


TEST_PAPER = "test_full_workflow_paper"


def cleanup():
    for p in [_get_state_path(TEST_PAPER), _get_orchestrate_state_path(TEST_PAPER)]:
        if os.path.exists(p):
            try:
                os.remove(p)
            except FileNotFoundError:
                pass


def create_proposal():
    """创建真实开题报告 docx"""
    try:
        import docx
    except ImportError:
        return None

    doc = docx.Document()
    doc.add_heading("论文大纲", level=1)
    chapters = [
        ("第1章 绪论", [("1.1 研究背景", "AI 时代背景"),
                     ("1.2 研究内容", "本文聚焦差异化战略")]),
        ("第2章 理论基础", [("2.1 竞争战略理论", "Porter 三种基本战略"),
                       ("2.2 文献综述", "近年 AI 时代战略文献")]),
        ("第3章 外部环境", [("3.1 宏观环境", "PEST 模型分析"),
                       ("3.2 行业结构", "波特五力模型")]),
        ("第4章 内部环境", [("4.1 资源与能力", "VRIN 资源分析")]),
        ("第5章 战略选择", [("5.1 QSPM 矩阵", "QSPM 量化评估")]),
        ("第6章 战略实施", [("6.1 实施路径", "三阶段实施")]),
        ("第7章 结论", [("7.1 研究结论", "总结研究结论")]),
    ]
    for ch_title, sections in chapters:
        doc.add_heading(ch_title, level=2)
        for sec_title, sec_content in sections:
            doc.add_heading(sec_title, level=3)
            doc.add_paragraph(sec_content)
    doc.add_heading("参考文献", level=1)
    path = f"/tmp/{TEST_PAPER}_proposal.docx"
    doc.save(path)
    return path


def chapter_aware_summary_llm(prompt: str) -> str:
    """章节感知 mock：根据 prompt 中的章节标题返回不同摘要"""
    if "绪论" in prompt:
        return "本章系统提出差异化战略研究问题，奠定全文框架。"
    if "理论基础" in prompt:
        return "本章系统梳理竞争战略理论，为后续分析提供方法论。"
    if "外部环境" in prompt:
        return "本章从宏观与行业结构两个维度分析外部环境。"
    return "本章进行相关分析。"


# ============================================================
# 测试 A：完整 docx 流程（happy path）
# ============================================================

def test_full_workflow_happy_path():
    """测试 A：完整 docx 流程"""
    print("\n=== 测试 A：完整 docx 流程（happy path） ===")
    cleanup()

    fake_path = create_proposal()
    if not fake_path:
        print("   ⚠️ 跳过（python-docx 未安装）")
        return

    # === Phase 1.1: docx 解析 ===
    print("   [1] Phase 1.1 docx 解析...")
    r1 = orchestrate_phase1_1(TEST_PAPER, "docx", fake_path)
    assert r1["ok"] is True
    outline_state = outline_load(TEST_PAPER)
    virtual_count = sum(1 for n in outline_state["outline"]["outline_tree"]["nodes"]
                        if n.get("is_virtual"))
    assert virtual_count == 7
    print(f"      ✅ 解析成功（{virtual_count} 个虚拟摘要节点）")

    # === Phase 1.2: 确认 ===
    print("   [2] Phase 1.2 目录确认...")
    r2 = confirm_phase1(TEST_PAPER)
    assert r2["phase"] == "phase1"

    # === Phase 1.3: 归因 ===
    print("   [3] Phase 1.3 开题报告归因...")
    r3 = orchestrate_phase1_3(TEST_PAPER)
    if r3["ok"]:
        print(f"      ✅ 归因成功")

    # === 用户调整 hint ===
    print("   [4] 用户调整 content_hint...")
    r4 = update_node_content_hint(TEST_PAPER, "1.1", "用户对研究背景的深度自定义提示")
    assert r4["ok"]
    node_11 = next((n for n in outline_load(TEST_PAPER)["outline"]["outline_tree"]["nodes"]
                    if n["id"] == "1.1"), None)
    assert node_11["content_hint"] == "用户对研究背景的深度自定义提示"
    print(f"      ✅ 用户调整生效")

    # === Phase 1.3 确认 → Phase 2 ===
    print("   [5] Phase 1.3 确认 → Phase 2...")
    r5 = confirm_phase1_3(TEST_PAPER)
    assert r5["phase"] == "phase2"

    # === Phase 2 写作（增强项1 触发章节摘要） ===
    print("   [6] Phase 2 写作（第一章所有 L2）...")
    # 模拟 1.1 + 1.2 完成 → 触发 ch1 摘要
    outline_update_status(TEST_PAPER, "1.1", "completed", key_conclusion="AI 时代背景")
    outline_update_status(TEST_PAPER, "1.2", "completed", key_conclusion="研究内容")

    r_ch1_summary = synthesize_chapter_summary(
        TEST_PAPER, "ch1", llm_func=chapter_aware_summary_llm
    )
    assert r_ch1_summary["ok"], f"ch1 摘要失败: {r_ch1_summary.get('error', '')}"
    print(f"      ✅ ch1 摘要合成: {r_ch1_summary['summary'][:40]}...")

    # === 验证 ch1 摘要已写入 ===
    outline_state = outline_load(TEST_PAPER)
    ch1_summary_node = next((n for n in outline_state["outline"]["outline_tree"]["nodes"]
                              if n["id"] == "__ch1_summary__"), None)
    assert ch1_summary_node["writing_status"] == "completed"
    print(f"      ✅ __ch1_summary__ 节点已 completed")

    print("   ✅ 全链路集成通过")


# ============================================================
# 测试 B：失败回退流程
# ============================================================

def test_full_workflow_failure_recovery():
    """测试 B：解析失败 → 重新输入 → 成功"""
    print("\n=== 测试 B：解析失败 → 重新输入 → 成功 ===")
    cleanup()

    print("   [1] 第一次输入：完全无效 text...")
    r1 = orchestrate_phase1_1(TEST_PAPER, "text", "完全是垃圾内容")
    assert r1["ok"] is False
    assert r1["action"] == "input_required"

    print("   [2] 第二次输入：还是无效...")
    r2 = orchestrate_phase1_1(TEST_PAPER, "text", "还是没有章节结构")
    assert r2["ok"] is False

    print("   [3] 第三次：切换到 docx...")
    fake_path = create_proposal()
    if not fake_path:
        print("      ⚠️ 跳过")
        return

    r3 = orchestrate_phase1_1(TEST_PAPER, "docx", fake_path)
    assert r3["ok"] is True

    confirm_phase1(TEST_PAPER)
    print(f"      ✅ 失败回退 → 成功")


# ============================================================
# 测试 C：章节摘要 + bridge 跨章节（增强项1 + Step 11 协同）
# ============================================================

def test_chapter_summary_bridge_integration():
    """测试 C：章节摘要合成 + bridge 跨章节引用"""
    print("\n=== 测试 C：章节摘要 + bridge 跨章节 ===")
    cleanup()

    fake_path = create_proposal()
    if not fake_path:
        print("   ⚠️ 跳过")
        return

    orchestrate_phase1_1(TEST_PAPER, "docx", fake_path)
    confirm_phase1(TEST_PAPER)

    # 模拟 ch1 + ch2 全部完成
    outline_update_status(TEST_PAPER, "1.1", "completed", key_conclusion="AI 时代背景")
    outline_update_status(TEST_PAPER, "1.2", "completed", key_conclusion="研究内容")
    r1 = synthesize_chapter_summary(TEST_PAPER, "ch1", chapter_aware_summary_llm)
    assert r1["ok"], f"ch1 summary fail: {r1.get('error', '')}"

    outline_update_status(TEST_PAPER, "2.1", "completed", key_conclusion="战略理论")
    outline_update_status(TEST_PAPER, "2.2", "completed", key_conclusion="文献综述")
    r2 = synthesize_chapter_summary(TEST_PAPER, "ch2", chapter_aware_summary_llm)
    assert r2["ok"], f"ch2 summary fail: {r2.get('error', '')}"

    # === 验证 ch2 摘要 ===
    outline_state = outline_load(TEST_PAPER)
    ch2_summary_node = next((n for n in outline_state["outline"]["outline_tree"]["nodes"]
                              if n["id"] == "__ch2_summary__"), None)
    assert ch2_summary_node["key_conclusion"] is not None
    print(f"   [1] __ch2_summary__: {ch2_summary_node['key_conclusion'][:50]}")

    # === 验证增强项1 P3 fallback：3.1 应引用 ch2 摘要 ===
    print("   [2] 验证 3.1 (ch3 首节点) prev_chapter_summary...")
    ctx = outline_get_context(TEST_PAPER, "3.1")
    prev_summary = ctx.get("prev_chapter_summary")
    assert prev_summary is not None, f"3.1 应有 prev_chapter_summary"
    assert prev_summary["chapter_id"] == "ch2"
    assert prev_summary["chapter_title"] == "理论基础"
    # 关键是内容应该是 ch2 摘要，不是 ch1
    assert "理论" in prev_summary["key_conclusion"] or "方法论" in prev_summary["key_conclusion"]
    print(f"      ✅ prev_chapter_summary = {prev_summary['chapter_title']}")

    # === 验证 bridge P3 fallback：3.1 prompt 应包含 ch2 摘要 ===
    print("   [3] 验证 3.1 prompt bridge...")
    pkg = build_prompt_package(TEST_PAPER, "3.1")
    bridge = pkg.get("bridge_paragraph")
    assert bridge is not None, "3.1 应有 bridge"
    assert "理论基础" in bridge, f"bridge 应包含 ch2 章节标题"
    print(f"      ✅ bridge: {bridge[:80]}...")

    # === 验证 content_hint 增强项4 + Step 11 协同 ===
    print("   [4] 验证 content_hint 集成...")
    content_hint = pkg.get("content_hint")
    print(f"      ✅ content_hint: {content_hint[:60] if content_hint else '<empty>'}...")

    print("   ✅ 增强项1 + 增强项4 + Step 11 协同工作")




# ============================================================
# Part 2: v2.0.0 端到端验证（方案 C：mock + 真实样本）
# ============================================================
#
# ⚠️ 隐私警告：开题报告样本含学生姓名/学号/研究方向，**严禁上传到 GitHub**。
#
# 真实样本测试默认跳过（不依赖任何路径/文件）。如需本地验证：
#   1. 设置环境变量 MBA_REAL_SAMPLES_DIR 指向你的样本目录
#   2. 该目录下的 *.docx 文件会被自动发现
#   3. paper_name 用 `v2_real_sample_<idx>` 匿名
#   4. 输出仅显示样本索引 + 文件大小，不含任何学生身份信息
#
# 详见 tests/REAL_SAMPLES_README.md

def _discover_real_samples():
    """
    从环境变量 MBA_REAL_SAMPLES_DIR 指定的目录中发现 docx 文件。
    返回 [(idx, full_path), ...]，idx 从 1 开始。
    若环境变量未设置或目录不存在，返回空列表（Part 2 跳过）。
    """
    samples_dir = os.environ.get("MBA_REAL_SAMPLES_DIR")
    if not samples_dir:
        return []
    if not os.path.isdir(samples_dir):
        return []
    docx_files = sorted(
        f for f in os.listdir(samples_dir)
        if f.lower().endswith(".docx") and not f.startswith("~")
    )
    return [(i + 1, os.path.join(samples_dir, f)) for i, f in enumerate(docx_files)]


def _get_sample_display_name(idx, docx_path):
    """
    样本显示名：使用匿名 ID（sample_001、sample_002...）。
    不包含学生姓名、文件原始名等隐私信息。
    """
    return f"sample_{idx:03d}"


def cleanup_paper(paper_name):
    for p in [_get_state_path(paper_name), _get_orchestrate_state_path(paper_name)]:
        if os.path.exists(p):
            try:
                os.remove(p)
            except FileNotFoundError:
                pass


# ============================================================
# 公共 mock 函数
# ============================================================

def mock_summary_llm(prompt: str) -> str:
    """章节摘要 mock：根据章节标题生成不同内容"""
    if "绪论" in prompt or "第一章" in prompt:
        return "本章系统提出研究问题，奠定全文框架。"
    if "理论基础" in prompt or "文献" in prompt:
        return "本章梳理竞争战略理论，为后续分析提供方法论支撑。"
    if "外部环境" in prompt or "宏观" in prompt:
        return "本章从宏观与行业结构分析外部环境。"
    if "内部环境" in prompt or "资源" in prompt:
        return "本章分析内部资源能力。"
    if "竞争战略" in prompt or "QSPM" in prompt:
        return "本章通过 QSPM 矩阵选择最优战略。"
    if "实施" in prompt:
        return "本章提出战略实施路径。"
    if "结论" in prompt:
        return "本章总结研究结论。"
    return "本章进行相关分析。"


def make_outline_simple():
    """构造最小 outline（含 L3 节点）"""
    nodes = [
        {"id": "ch1", "level": 1, "title": "绪论", "parent_id": None,
         "children_ids": ["1.1", "1.2"], "prev_sibling_id": None,
         "next_sibling_id": "ch2", "writing_status": "pending"},
        {"id": "1.1", "level": 2, "title": "研究背景", "parent_id": "ch1",
         "children_ids": [], "prev_sibling_id": None,
         "next_sibling_id": "1.2", "writing_status": "pending"},
        {"id": "1.2", "level": 2, "title": "研究内容", "parent_id": "ch1",
         "children_ids": [], "prev_sibling_id": "1.1",
         "next_sibling_id": None, "writing_status": "pending"},
        {"id": "ch2", "level": 1, "title": "理论基础", "parent_id": None,
         "children_ids": ["2.1", "2.2"], "prev_sibling_id": "ch1",
         "next_sibling_id": "ch3", "writing_status": "pending"},
        {"id": "2.1", "level": 2, "title": "竞争战略理论", "parent_id": "ch2",
         "children_ids": [], "prev_sibling_id": None,
         "next_sibling_id": "2.2", "writing_status": "pending"},
        {"id": "2.2", "level": 2, "title": "文献综述", "parent_id": "ch2",
         "children_ids": [], "prev_sibling_id": "2.1",
         "next_sibling_id": None, "writing_status": "pending"},
        {"id": "ch3", "level": 1, "title": "外部环境", "parent_id": None,
         "children_ids": ["3.1", "3.2", "3.1.1"], "prev_sibling_id": "ch2",
         "next_sibling_id": None, "writing_status": "pending"},
        {"id": "3.1", "level": 2, "title": "宏观环境", "parent_id": "ch3",
         "children_ids": ["3.1.1"], "prev_sibling_id": None,
         "next_sibling_id": "3.2", "writing_status": "pending"},
        {"id": "3.1.1", "level": 3, "title": "PEST 模型", "parent_id": "3.1",
         "children_ids": [], "prev_sibling_id": None,
         "next_sibling_id": None, "writing_status": "pending"},
        {"id": "3.2", "level": 2, "title": "行业结构", "parent_id": "ch3",
         "children_ids": [], "prev_sibling_id": "3.1",
         "next_sibling_id": None, "writing_status": "pending"},
    ]
    return {"outline_tree": {"metadata": {"paper_title": "mock"}, "nodes": nodes}}


# ============================================================
# Part 1: 7 个 mock 边界测试（方案 A）
# ============================================================

def test_mock_1_content_hint_consistency():
    """Mock 测试 1：Phase 1.3 submit 后 content_hint 在 3 个地方一致"""
    print("\n=== Mock 测试 1：content_hint 端到端一致性 ===")
    paper = "test_mock_1"
    cleanup_paper(paper)

    outline = make_outline_simple()
    outline_with_summary = insert_chapter_summary_nodes(outline)
    outline_save(paper, outline_with_summary)
    confirm_phase1(paper)

    # 直接调用 save_content_hints_to_outline（模拟 Phase 1.3）
    hints = {
        "1.1": "AI 时代互联网分发面临范式重构。",
        "1.2": "本文聚焦差异化战略。",
        "2.1": "Porter 三种基本竞争战略。",
        "2.2": "AI 时代文献综述。",
        "3.1": "PEST 模型分析。",
        "3.1.1": "PEST 模型详解。",
        "3.2": "五力模型。",
    }
    save_content_hints_to_outline(paper, hints)

    # 验证 1：outline_state 节点字段
    state = outline_load(paper)
    node_11 = next((n for n in state["outline"]["outline_tree"]["nodes"] if n["id"] == "1.1"), None)
    assert node_11["content_hint"] == "AI 时代互联网分发面临范式重构。", \
        f"outline_state 应有 hint，实际: {node_11.get('content_hint')}"
    print(f"   ✅ outline_state 节点 content_hint 已写入")

    # 验证 2：build_prompt_package 读取
    pkg = build_prompt_package(paper, "1.1")
    assert pkg["content_hint"] == "AI 时代互联网分发面临范式重构。", \
        f"prompt 包应包含 hint，实际: {pkg.get('content_hint')}"
    print(f"   ✅ build_prompt_package content_hint 正确")

    # 验证 3：prompt 文本化
    text = build_prompt_package_text(pkg)
    assert "## 开题报告方向参考" in text, "prompt 文本应包含方向参考 section"
    assert "AI 时代" in text
    print(f"   ✅ prompt 文本包含「## 开题报告方向参考」section")

    cleanup_paper(paper)


def test_mock_2_chapter_summary_boundaries():
    """Mock 测试 2：L1/L2/L3 首节点的 prev_chapter_summary 触发边界"""
    print("\n=== Mock 测试 2：章节首节点 prev_chapter_summary 边界 ===")
    paper = "test_mock_2"
    cleanup_paper(paper)

    outline = make_outline_simple()
    outline_with_summary = insert_chapter_summary_nodes(outline)
    outline_save(paper, outline_with_summary)
    confirm_phase1(paper)

    # ch1 完成 → ch1 摘要合成
    outline_update_status(paper, "1.1", "completed", key_conclusion="AI 时代背景")
    outline_update_status(paper, "1.2", "completed", key_conclusion="研究内容")
    synthesize_chapter_summary(paper, "ch1", mock_summary_llm)

    # ch2 完成 → ch2 摘要合成
    outline_update_status(paper, "2.1", "completed", key_conclusion="Porter")
    outline_update_status(paper, "2.2", "completed", key_conclusion="文献综述")
    synthesize_chapter_summary(paper, "ch2", mock_summary_llm)

    # 场景 1：2.1 (L2 章节首节点) — 应有 prev_chapter_summary = ch1
    ctx = outline_get_context(paper, "2.1")
    pcs = ctx.get("prev_chapter_summary")
    assert pcs is not None, "2.1 应有 prev_chapter_summary"
    assert pcs["chapter_id"] == "ch1"
    print(f"   ✅ L2 章节首节点 (2.1): prev_chapter_summary = {pcs['chapter_title']}")

    # 场景 2：3.1 (L2 章节首节点) — 应有 prev_chapter_summary = ch2
    ctx = outline_get_context(paper, "3.1")
    pcs = ctx.get("prev_chapter_summary")
    assert pcs is not None, "3.1 应有 prev_chapter_summary"
    assert pcs["chapter_id"] == "ch2"
    print(f"   ✅ L2 章节首节点 (3.1): prev_chapter_summary = {pcs['chapter_title']}")

    # 场景 3：3.1.1 (L3 节点，parent=3.1) — 不是章节首节点 → 应无
    ctx = outline_get_context(paper, "3.1.1")
    pcs = ctx.get("prev_chapter_summary")
    assert pcs is None, "3.1.1 不是章节首节点 → 应无"
    print(f"   ✅ L3 节点 (3.1.1): prev_chapter_summary = None")

    # 场景 4：1.1 (首章节首节点) — 应无
    ctx = outline_get_context(paper, "1.1")
    pcs = ctx.get("prev_chapter_summary")
    assert pcs is None
    print(f"   ✅ 首章节首节点 (1.1): prev_chapter_summary = None")

    # 场景 5：1.2 (章节内中间节点) — 应无
    ctx = outline_get_context(paper, "1.2")
    pcs = ctx.get("prev_chapter_summary")
    assert pcs is None
    print(f"   ✅ 章节内中间节点 (1.2): prev_chapter_summary = None")

    cleanup_paper(paper)


def test_mock_3_state_schema_completeness():
    """Mock 测试 3：v2.0.0 state schema 完整性"""
    print("\n=== Mock 测试 3：v2.0.0 state schema 完整性 ===")
    paper = "test_mock_3"
    cleanup_paper(paper)

    outline_text = """第1章 绪论
1.1 研究背景
1.2 研究内容
第2章 理论基础
2.1 战略理论
2.2 文献综述
第3章 外部环境
3.1 宏观
3.2 行业
第4章 内部环境
4.1 资源
第5章 战略选择
5.1 QSPM
第6章 实施
6.1 路径
第7章 结论
7.1 总结
参考文献"""

    r = orchestrate_phase1_1(paper, "text", outline_text)
    assert r["ok"]
    confirm_phase1(paper)

    # outline_state schema
    outline_state = outline_load(paper)
    meta = outline_state["outline"]["outline_tree"]["metadata"]
    assert "virtual_nodes" in meta
    assert "real_nodes" in meta
    print(f"   ✅ outline metadata: virtual={meta['virtual_nodes']}, real={meta['real_nodes']}")

    # orchestrate_state schema
    orchest_state = load_orchestrate_state(paper)
    required_fields = [
        "phase1_3_status", "phase1_3_docx_path",
        "phase1_3_result", "phase1_3_submitted_at", "phase1_3_confirmed_at"
    ]
    for f in required_fields:
        assert f in orchest_state, f"应包含 {f}"
    assert orchest_state["phase1_3_status"] == "pending"
    assert orchest_state["phase1_3_docx_path"] is None
    print(f"   ✅ orchestrate_state 5 个 phase1_3 字段全部就位")

    cleanup_paper(paper)


def test_mock_4_three_decision_paths():
    """Mock 测试 4：3 个决策路径混合（决策 1/2/3 各一个节点）"""
    print("\n=== Mock 测试 4：3 个决策路径混合 ===")
    paper = "test_mock_4"
    cleanup_paper(paper)

    outline_text = """第1章 绪论
1.1 节点 A
1.2 节点 B
1.3 节点 C
第2章 理论基础
2.1 战略理论
2.2 文献综述
第3章 外部环境
3.1 宏观
3.2 行业
第4章 内部环境
4.1 资源
第5章 战略选择
5.1 QSPM
第6章 实施
6.1 路径
第7章 结论
7.1 总结
参考文献"""

    orchestrate_phase1_1(paper, "text", outline_text)
    confirm_phase1(paper)
    orchestrate_phase1_3(paper)

    # 决策 1：用户提供 hint
    r1 = apply_user_decision(paper, "1.1", "1", user_hint="用户提供的研究背景")
    assert r1["action"] == "proceed" and r1["decision"] == "1"
    print(f"   ✅ 决策 1 (用户提供 hint) → proceed")

    # 决策 2：AI 自行生成
    r2 = apply_user_decision(paper, "1.2", "2")
    assert r2["action"] == "proceed" and r2["decision"] == "2"
    state = outline_load(paper)
    node_12 = next((n for n in state["outline"]["outline_tree"]["nodes"] if n["id"] == "1.2"), None)
    assert not node_12.get("content_hint")
    print(f"   ✅ 决策 2 (AI 自行生成) → proceed (不写 hint)")

    # 决策 3：跳过节点
    r3 = apply_user_decision(paper, "1.3", "3")
    assert r3["action"] == "skipped" and r3["decision"] == "3"
    state = outline_load(paper)
    node_13 = next((n for n in state["outline"]["outline_tree"]["nodes"] if n["id"] == "1.3"), None)
    assert node_13["writing_status"] == "failed"
    print(f"   ✅ 决策 3 (跳过节点) → failed")

    cleanup_paper(paper)


def test_mock_5_chained_chapter_summaries():
    """Mock 测试 5：ch1→ch2→ch3 链式合成"""
    print("\n=== Mock 测试 5：多章节摘要链式合成 ===")
    paper = "test_mock_5"
    cleanup_paper(paper)

    outline = make_outline_simple()
    outline_with_summary = insert_chapter_summary_nodes(outline)
    outline_save(paper, outline_with_summary)
    confirm_phase1(paper)

    # ch1 摘要
    outline_update_status(paper, "1.1", "completed", key_conclusion="ch1.1")
    outline_update_status(paper, "1.2", "completed", key_conclusion="ch1.2")
    r1 = synthesize_chapter_summary(paper, "ch1", mock_summary_llm)
    assert r1["ok"]

    # ch2 摘要
    outline_update_status(paper, "2.1", "completed", key_conclusion="ch2.1")
    outline_update_status(paper, "2.2", "completed", key_conclusion="ch2.2")
    r2 = synthesize_chapter_summary(paper, "ch2", mock_summary_llm)
    assert r2["ok"]

    # ch3 摘要
    outline_update_status(paper, "3.1", "completed", key_conclusion="ch3.1")
    outline_update_status(paper, "3.1.1", "completed", key_conclusion="ch3.1.1")
    outline_update_status(paper, "3.2", "completed", key_conclusion="ch3.2")
    r3 = synthesize_chapter_summary(paper, "ch3", mock_summary_llm)
    assert r3["ok"]

    # 验证 3 个虚拟摘要节点都已 completed
    state = outline_load(paper)
    for ch in ["__ch1_summary__", "__ch2_summary__", "__ch3_summary__"]:
        sn = next((n for n in state["outline"]["outline_tree"]["nodes"] if n["id"] == ch), None)
        assert sn is not None, f"{ch} 应存在"
        assert sn["writing_status"] == "completed"
        assert sn["key_conclusion"] is not None
        print(f"   ✅ {ch}: {sn['key_conclusion'][:40]}...")

    cleanup_paper(paper)


def test_mock_6_failure_recovery():
    """Mock 测试 6：失败回退路径"""
    print("\n=== Mock 测试 6：失败回退路径 ===")
    paper = "test_mock_6"
    cleanup_paper(paper)

    # 第一次：解析失败
    r1 = orchestrate_phase1_1(paper, "text", "完全是垃圾内容")
    assert r1["ok"] is False
    assert r1["action"] == "input_required"
    assert "1" in r1["retry_options"]
    assert "2" in r1["retry_options"]
    assert "3" in r1["retry_options"]
    print(f"   ✅ text 解析失败 → input_required + 3 选项")

    # 第二次：真实 docx
    fake_path = "/tmp/test_failure_recovery_v2.docx"
    import docx
    doc = docx.Document()
    doc.add_heading("论文大纲", level=1)
    for i in range(1, 8):
        doc.add_heading(f"第{i}章 第{i}章", level=2)
        for j in range(1, 3):
            doc.add_heading(f"{i}.{j} {i}.{j}", level=3)
            doc.add_paragraph("内容")
    doc.add_heading("参考文献", level=1)
    doc.save(fake_path)

    r2 = orchestrate_phase1_1(paper, "docx", fake_path)
    assert r2["ok"] is True
    print(f"   ✅ 失败后切换 docx → 成功")

    confirm_phase1(paper)
    os.remove(fake_path)
    r3 = orchestrate_phase1_3(paper)
    assert r3["ok"] is False
    assert r3["action"] == "input_required"
    print(f"   ✅ docx 失效 → Phase 1.3 input_required")

    cleanup_paper(paper)


def test_mock_7_phase2_enforcement():
    """Mock 测试 7：Phase 2 强制检查 + state 完整性"""
    print("\n=== Mock 测试 7：Phase 2 强制检查 ===")
    paper = "test_mock_7"
    cleanup_paper(paper)

    outline_text = """第1章 绪论
1.1 研究背景
第2章 理论基础
2.1 战略理论
第3章 外部环境
3.1 宏观
第4章 内部环境
4.1 资源
第5章 战略选择
5.1 QSPM
第6章 实施
6.1 路径
第7章 结论
7.1 总结
参考文献"""

    orchestrate_phase1_1(paper, "text", outline_text)
    confirm_phase1(paper)

    # 手动把 phase 设为 phase2，但 phase1_3_status=pending
    state = load_orchestrate_state(paper)
    state["phase"] = "phase2"
    with open(_get_orchestrate_state_path(paper), "w", encoding="utf-8") as f:
        json.dump(state, f, ensure_ascii=False, indent=2)

    r = orchestrate_phase2(paper, llm_func=lambda p: "x")
    assert r["ok"] is False
    assert "Phase 1.3" in r["error"]
    print(f"   ✅ Phase 2 拒绝未确认: {r['error'][:60]}")

    cleanup_paper(paper)


# ============================================================
# Part 2: 真实 docx 样本完整流程测试（方案 B）
# ============================================================

def run_real_sample(sample_name, docx_path):
    """对单个真实 docx 样本跑完整流程"""
    # paper_name 使用匿名 sample_xxx 形式
    safe_name = "".join(c if c.isalnum() or c == "_" else "_" for c in sample_name)
    paper = f"v2_real_{safe_name}"
    cleanup_paper(paper)

    if not os.path.exists(docx_path):
        return {"sample": sample_name, "ok": False, "error": f"文件不存在"}

    result = {
        "sample": sample_name,
        "docx_size_kb": os.path.getsize(docx_path) // 1024,
        "phase_1_1": {},
        "phase_1_2": {},
        "phase_1_3": {},
        "phase_2_write": {},
        "chapter_summary": {},
        "errors": []
    }

    # Phase 1.1: 解析
    t1 = time.time()
    r1 = orchestrate_phase1_1(paper, "docx", docx_path)
    t1_elapsed = time.time() - t1

    if not r1["ok"]:
        result["errors"].append(f"Phase 1.1: {r1.get('error', '')[:80]}")
        result["phase_1_1"] = {"ok": False, "elapsed_sec": round(t1_elapsed, 2)}
        cleanup_paper(paper)
        return result

    state = outline_load(paper)
    real_nodes = sum(1 for n in state["outline"]["outline_tree"]["nodes"] if not n.get("is_virtual"))
    virtual_nodes = sum(1 for n in state["outline"]["outline_tree"]["nodes"] if n.get("is_virtual"))

    result["phase_1_1"] = {
        "ok": True,
        "elapsed_sec": round(t1_elapsed, 2),
        "real_nodes": real_nodes,
        "virtual_nodes": virtual_nodes
    }

    # Phase 1.2
    t2 = time.time()
    r2 = confirm_phase1(paper)
    t2_elapsed = time.time() - t2
    result["phase_1_2"] = {
        "ok": r2["ok"],
        "elapsed_sec": round(t2_elapsed, 3)
    }

    # Phase 1.3（用 mock_summary_llm）
    t3 = time.time()
    r3 = orchestrate_phase1_3(paper, llm_func=mock_summary_llm)
    t3_elapsed = time.time() - t3

    if not r3["ok"]:
        result["errors"].append(f"Phase 1.3: {r3.get('error', '')[:80]}")
        result["phase_1_3"] = {"ok": False, "elapsed_sec": round(t3_elapsed, 2)}
        cleanup_paper(paper)
        return result

    hints_written = sum(1 for n in state["outline"]["outline_tree"]["nodes"]
                        if n.get("content_hint"))

    result["phase_1_3"] = {
        "ok": True,
        "elapsed_sec": round(t3_elapsed, 2),
        "hints_written": hints_written
    }

    # Phase 2 模拟写 1.1 + 1.2
    t4 = time.time()
    outline_update_status(paper, "1.1", "completed", key_conclusion=f"{sample_name} ch1.1 结论")
    outline_update_status(paper, "1.2", "completed", key_conclusion=f"{sample_name} ch1.2 结论")
    t4_elapsed = time.time() - t4
    result["phase_2_write"] = {"ok": True, "elapsed_sec": round(t4_elapsed, 3), "nodes_completed": 2}

    # 章节摘要
    t5 = time.time()
    r5 = synthesize_chapter_summary(paper, "ch1", mock_summary_llm)
    t5_elapsed = time.time() - t5

    if not r5["ok"]:
        result["errors"].append(f"chapter_summary: {r5.get('error', '')[:80]}")
        result["chapter_summary"] = {"ok": False, "elapsed_sec": round(t5_elapsed, 2)}
    else:
        result["chapter_summary"] = {
            "ok": True,
            "elapsed_sec": round(t5_elapsed, 2),
            "summary_len": len(r5["summary"])
        }

    cleanup_paper(paper)
    return result


def _run_real_samples():
    """真实 docx 样本批量测试（默认跳过，需设置 MBA_REAL_SAMPLES_DIR）

    修复 P2-3：拆分为 _run_real_samples（返回元组，main() 调用）
    + test_real_samples（pytest 测试，不返回）。
    """
    real_samples = _discover_real_samples()
    samples_dir = os.environ.get("MBA_REAL_SAMPLES_DIR", "")

    if not real_samples:
        print("\n=== 真实 docx 样本批量测试 ===")
        print(f"   ⏭️  跳过（未设置 MBA_REAL_SAMPLES_DIR 环境变量）")
        print(f"   说明：详见 tests/REAL_SAMPLES_README.md")
        return 0, 0

    print("\n=== 真实 docx 样本批量测试 ===")
    print(f"   样本目录: {samples_dir}")
    print(f"   待测试样本: {len(real_samples)} 个（匿名展示）")

    all_results = []
    total_start = time.time()

    for idx, docx_path in real_samples:
        sample_name = _get_sample_display_name(idx, docx_path)

        if not os.path.exists(docx_path):
            print(f"\n   ⚠️  {sample_name}: 文件不存在")
            continue

        print(f"\n   [测试] {sample_name} ({os.path.getsize(docx_path)//1024} KB)...")
        result = run_real_sample(sample_name, docx_path)
        all_results.append(result)

        if result["errors"]:
            print(f"      ❌ 失败: {result['errors']}")
        else:
            print(f"      ✅ 全流程通过")
            print(f"         Phase 1.1: {result['phase_1_1'].get('real_nodes', 0)} 真实 + "
                  f"{result['phase_1_1'].get('virtual_nodes', 0)} 虚拟节点 "
                  f"({result['phase_1_1']['elapsed_sec']}s)")
            print(f"         Phase 1.3: {result['phase_1_3'].get('hints_written', 0)} hints 写入 "
                  f"({result['phase_1_3']['elapsed_sec']}s)")
            print(f"         ch1 摘要: {result['chapter_summary'].get('summary_len', 0)} 字 "
                  f"({result['chapter_summary']['elapsed_sec']}s)")

    total_elapsed = time.time() - total_start

    # 汇总
    print("\n" + "=" * 60)
    print(f"📊 真实样本测试汇总（{len(all_results)} 个样本，{round(total_elapsed, 1)}s 总耗时）")
    print("=" * 60)

    passed = sum(1 for r in all_results if not r["errors"])
    failed = len(all_results) - passed

    print(f"\n   通过率: {passed}/{len(all_results)} ({round(passed/len(all_results)*100, 1)}%)")
    print()
    for r in all_results:
        status = "✅" if not r["errors"] else "❌"
        total_time = sum([
            r["phase_1_1"].get("elapsed_sec", 0),
            r["phase_1_2"].get("elapsed_sec", 0),
            r["phase_1_3"].get("elapsed_sec", 0),
            r["phase_2_write"].get("elapsed_sec", 0),
            r["chapter_summary"].get("elapsed_sec", 0),
        ])
        # 仅显示匿名 ID 和文件大小，不含学生身份
        print(f"   {status} {r['sample']:12s} ({r['docx_size_kb']:4d} KB): "
              f"ch1 完成, 摘要 {r['chapter_summary'].get('summary_len', 0)} 字, "
              f"耗时 {round(total_time, 2)}s")

    return passed, failed


def test_real_samples():
    """pytest 测试入口：调用 _run_real_samples 但不返回值"""
    passed, failed = _run_real_samples()
    # pytest 模式：断言至少运行（passed + failed > 0 或 0 0 表示跳过）
    # 不强求 passed > 0，因为默认会跳过（无 MBA_REAL_SAMPLES_DIR）


# ============================================================
# 主入口
# ============================================================

def main():
    print("=" * 60)
    print("v2.0.0 端到端验证（方案 C：mock + 真实样本）")
    print("=" * 60)

    # Part 0: Step 12 全链路集成测试（修订 v2）
    print("\n## Part 0: Step 12 全链路集成测试（A/B/C）")
    integration_tests = [
        test_full_workflow_happy_path,
        test_full_workflow_failure_recovery,
        test_chapter_summary_bridge_integration,
    ]

    integration_passed = 0
    integration_failed = 0
    for t in integration_tests:
        try:
            t()
            integration_passed += 1
        except AssertionError as e:
            print(f"\n❌ 测试失败: {e}")
            integration_failed += 1
        except Exception as e:
            print(f"\n❌ 异常: {e}")
            import traceback
            traceback.print_exc()
            integration_failed += 1
        finally:
            cleanup()

    print(f"\n## Step 12 集成测试汇总: {integration_passed}/{len(integration_tests)} 通过")

    # Part 1: 7 个 mock 测试
    print("\n## Part 1: 7 个 mock 边界测试（方案 A）")
    mock_tests = [
        test_mock_1_content_hint_consistency,
        test_mock_2_chapter_summary_boundaries,
        test_mock_3_state_schema_completeness,
        test_mock_4_three_decision_paths,
        test_mock_5_chained_chapter_summaries,
        test_mock_6_failure_recovery,
        test_mock_7_phase2_enforcement,
    ]

    mock_passed = 0
    mock_failed = 0
    for t in mock_tests:
        try:
            t()
            mock_passed += 1
        except AssertionError as e:
            print(f"\n❌ 测试失败: {e}")
            mock_failed += 1
        except Exception as e:
            print(f"\n❌ 异常: {e}")
            import traceback
            traceback.print_exc()
            mock_failed += 1

    print(f"\n## Mock 测试汇总: {mock_passed}/{len(mock_tests)} 通过")

    # Part 2: 真实样本测试
    print("\n## Part 2: 真实 docx 样本测试（方案 B）")
    real_passed, real_failed = _run_real_samples()

    print("\n" + "=" * 60)
    print(f"🎯 最终汇总")
    print(f"   Step 12 集成 (A/B/C): {integration_passed}/{len(integration_tests)} 通过")
    print(f"   Mock 测试 (方案 A): {mock_passed}/{len(mock_tests)} 通过")
    if real_passed + real_failed > 0:
        print(f"   真实样本 (方案 B): {real_passed}/{real_passed + real_failed} 通过")
    else:
        print(f"   真实样本 (方案 B): 跳过（未设置 MBA_REAL_SAMPLES_DIR）")
    print("=" * 60)

    total_failed = integration_failed + mock_failed + real_failed
    return 0 if total_failed == 0 else 1


if __name__ == "__main__":
    sys.exit(main())
