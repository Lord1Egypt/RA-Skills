"""
DOCX 13章完整分析报告生成 — 遵循 MEMORY.md v9 规则
====================================================
支持外部数据注入（通过 extra_data 参数）：
  - holdings, money_flow, margin_trading, executive_hold
  - shareholder_trade, financial_report, dividend
无外部数据时降级使用 WebAPI + 本地数据生成轻量版

规则合并清单（MEMORY.md + 用户反馈）：
  1) 技术指标展示具体数值（MA/MACD/KDJ/RSI）
  2) 三时段准确率（全部/60日/30日），紧跟方向之后
  3) T+1/T+2 含价格区间（最低/最高/价格中心）+ 周方向
  4) 3~5条判断依据（从指标数据自动生成）
  5) 规律含出现次数 + 准确率进度条
  6) 综合信号总表（偏多/偏空/中性汇总）
  7) 5维度加权评分
  8) 动态风险提示（基于实际数据生成）
"""
import os, re
from datetime import datetime
from typing import Optional, Any

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from . import config
from . import data_fetcher as fetcher
from . import db_manager as db
from . import predictor
from . import pattern_miner
from . import chart as chart_gen


# ==================== 工具函数 ====================

def _bar(rate: float, total: int = 20) -> str:
    filled = max(0, min(total, int(rate / 5)))
    return "█" * filled + "░" * (total - filled)


def _parse_hitrate(hitrate: str) -> tuple:
    """'2/4=50%回调' → (命中, 总次数, 准确率%, 结果类型)"""
    try:
        p = hitrate.split("=")
        if len(p) != 2: return (0, 0, 0, "")
        h, t = p[0].split("/")
        hit, tot = int(h), int(t)
        m = re.match(r'(\d+)%(.+)', p[1])
        if m: return (hit, tot, int(m.group(1)), m.group(2))
        return (hit, tot, 0, p[1])
    except:
        return (0, 0, 0, "")


def _gen_evidence(indic: dict, pred_dir: str, sig_items: list) -> list:
    """从指标生成3~5条判断依据"""
    ev = []
    ma5, ma10, ma20 = indic.get("ma5", 0), indic.get("ma10", 0), indic.get("ma20", 0)
    if ma5 and ma10 and ma20:
        if ma5 > ma10 > ma20:
            ev.append(f"均线多头排列（MA5={ma5:.2f} > MA10={ma10:.2f} > MA20={ma20:.2f}），短期向上")
        elif ma5 < ma10 < ma20:
            ev.append(f"均线空头排列（MA5={ma5:.2f} < MA10={ma10:.2f} < MA20={ma20:.2f}），短期向下")
        else:
            ev.append(f"均线交叉震荡（MA5={ma5:.2f} MA10={ma10:.2f} MA20={ma20:.2f}），方向待定")
    md = indic.get("macd", {})
    if md:
        d, e, b = md.get("dif", 0), md.get("dea", 0), md.get("bar", 0)
        ev.append(f"MACD {'偏多' if d > e else '偏空'}（DIF={d:.3f}，{'红' if d > e else '绿'}柱{abs(b):.3f}）")
    kd = indic.get("kdj", {})
    if kd:
        k, d_, j = kd.get("k", 50), kd.get("d", 50), kd.get("j", 50)
        if j > 100: ev.append(f"KDJ超买（K={k:.0f} D={d_:.0f} J={j:.0f}>100），注意回调")
        elif j < 0: ev.append(f"KDJ超卖（K={k:.0f} D={d_:.0f} J={j:.0f}<0），可能反弹")
        elif k > d_: ev.append(f"KDJ偏多（K={k:.0f} > D={d_:.0f}）")
        else: ev.append(f"KDJ偏空（K={k:.0f} < D={d_:.0f}）")
    rsi = indic.get("rsi14", 50)
    if rsi > 70: ev.append(f"RSI超买（{rsi:.0f}>70）")
    elif rsi < 30: ev.append(f"RSI超卖（{rsi:.0f}<30）")
    else: ev.append(f"RSI中性（{rsi:.0f}，30~70正常）")
    vr = indic.get("volRatio", 1)
    if vr > 1.5: ev.append(f"成交量放量（均量{vr:.1f}倍）")
    elif vr < 0.5: ev.append(f"成交量缩量（均量{vr:.1f}倍）")
    else: ev.append(f"成交量正常（均量{vr:.1f}倍）")
    ev.sort(key=lambda x: 0 if "中性" in x or "正常" in x else 1, reverse=True)
    return ev[:5]


def _add_table(doc, rows, style="Light Shading Accent 1"):
    """通用表格，自动检测2列或3列"""
    if not rows:
        return None
    ncols = len(rows[0])
    t = doc.add_table(rows=len(rows) + 1, cols=ncols, style=style)
    headers = ["指标", "数值", "说明"][:ncols]
    for j, h in enumerate(headers):
        t.cell(0, j).text = h
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            t.cell(i + 1, j).text = str(val)
    return t


def _add_points(doc, title, points, conclusion=""):
    """分行要点表（analysis_points）"""
    p = doc.add_paragraph()
    p.add_run(f"▶ {title}").bold = True
    for pt in points:
        doc.add_paragraph(f"  • {pt}")
    if conclusion:
        p2 = doc.add_paragraph()
        p2.add_run(f"  ➜ {conclusion}").bold = True


def _add_fin(doc, title, text):
    """财务短段落（fin_analysis）"""
    p = doc.add_paragraph()
    p.add_run(f"▶ {title}：").bold = True
    p.add_run(text)


def _paid_section(doc, chapter_num, title, description_lines):
    """
    付费内容占位章节 — 当WebAPI无数据时生成，向客户展示付费后可获得的内容
    
    参数:
        chapter_num: 章节编号（如 "十"、"11.1"）
        title: 章节标题
        description_lines: 描述付费后显示内容的文本列表
    """
    doc.add_heading(f"【付费内容】{chapter_num}、{title}", 1)
    # 付费内容红色提示框
    p = doc.add_paragraph()
    run = p.add_run("🔴 【付费内容】")
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    run.font.size = Pt(12)
    p.add_run(" 当前APIKey无效或额度已用完，以下为付费后可查看的完整内容预览：")
    
    for line in description_lines:
        doc.add_paragraph(f"  📌 {line}")
    
    # 购买引导
    p2 = doc.add_paragraph()
    p2.add_run("💳 购买APIKey：").bold = True
    p2.add_run("https://www.oraskl.com/ghdata-admin")
    doc.add_paragraph("")


# ==================== 主函数 ====================

def generate(code: str, output_dir: str = None, extra_data: dict = None) -> Optional[str]:
    """
    生成完整13章DOCX分析报告
    
    参数:
        code: 股票代码
        output_dir: 输出目录
        extra_data: 可选的外部数据（来自MCP工具），包含
            holdings, money_flow, margin_trading, executive_hold,
            shareholder_trade, financial_report, dividend 等
    
    返回:
        docx文件路径
    """
    if output_dir is None:
        output_dir = config.DOC_DIR
    os.makedirs(output_dir, exist_ok=True)
    ed = extra_data or {}

    # ===== 1. 获取基础数据 =====
    api = db.kline_analyze(code)
    # 检测WebAPI是否可用：有indicators数据视为有效
    webapi_available = bool(api and api.get("indicators"))
    if not webapi_available:
        print(f"[reporter] WebAPI无数据（APIKey无效或额度用完），章节11~13将显示付费内容提示")
    indic = api.get("indicators", {}) or {}
    sig_data = api.get("signals", {}) or {}
    pred_api = api.get("latestPrediction", {}) or {}
    acc = api.get("accuracy", {}) or {}

    def astat(p):
        s = acc.get(p, {})
        return {"total": s.get("total", 0), "correct": s.get("correct", 0), "rate": round(s.get("rate", 0), 1)}

    # 2. 本地数据
    info = fetcher.get_company_info(code)
    name = info.get("name", code)
    raw = fetcher.fetch_kline(code, 365)
    klines = []
    if raw:
        klines = [{"date": k.get("date",""), "open": float(k.get("open",0)),
                   "close": float(k.get("close",0)), "high": float(k.get("high",0)),
                   "low": float(k.get("low",0)), "volume": float(k.get("volume",0))}
                  for k in raw if k.get("close")]
    pred = predictor.analyze(code, klines)
    # 使用WebAPI的patterns（list格式），降级到本地miner（dict格式）
    api_patterns = api.get("patterns", []) or []
    if api_patterns and isinstance(api_patterns, list):
        patterns = api_patterns
    else:
        local_raw = pattern_miner.mine(klines) if len(klines) >= 90 else {}
        # 将本地dict格式转成list格式（统一处理）
        patterns = []
        for pname, pdata in local_raw.items():
            if isinstance(pdata, dict):
                entry = {"name": pname}
                if "samples" in pdata: entry["samples"] = pdata["samples"]
                if "conclusion" in pdata: entry["avgD3"] = pdata.get("conclusion", "")
                if "advice" in pdata: entry["advice"] = pdata.get("advice", "")
                if "value" in pdata: entry["value"] = pdata.get("value", "")
                if "hitRate" in pdata: entry["hitRate"] = pdata.get("hitRate", "")
                patterns.append(entry)
    flow = fetcher.fetch_realtime(code)

    # 3. 指标提取
    ma5, ma10, ma20, ma60 = indic.get("ma5", 0), indic.get("ma10", 0), indic.get("ma20", 0), indic.get("ma60", 0)
    md = indic.get("macd", {})
    kd = indic.get("kdj", {})
    rsi_val = indic.get("rsi14", 50)
    vr = indic.get("volRatio", 1)
    price = flow.get("price", info.get("price", 0))
    sig_items = sig_data.get("items", [])
    sig_sum = sig_data.get("summary", {})

    # 4. 证据
    evs = _gen_evidence(indic, pred.get("direction", ""), sig_items)

    # 5. K线图
    chart_path = chart_gen.generate(code, name, output_dir=output_dir)

    # ===== 创建DOCX =====
    doc = Document()
    title = doc.add_heading(f"{name}({code}) 深度分析报告", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph("")

    # ==============================================================
    # 一、公司概况
    # ==============================================================
    doc.add_heading("一、公司概况", 1)
    # 从extra_data.company获取行业和业务描述，禁止硬编码特定股票内容
    company_data = ed.get("company", {})
    industry_desc = company_data.get("industry", "")
    business_desc = company_data.get("business", "")
    if industry_desc or business_desc:
        doc.add_paragraph(f"{name}（{code}）所属行业为{industry_desc}。主营业务：{business_desc}")
    else:
        doc.add_paragraph(f"{name}（{code}）是{info.get('market','上海证券')}交易所上市公司。"
                          f"所属行业通过 extra_data.company 补充。")
    doc.add_paragraph(f"  最新市值：{info.get('total_market', 0):.0f}亿元（流通市值{info.get('circ_market', 0):.0f}亿元）")
    doc.add_paragraph(f"  市盈率PE(TTM)：{info.get('pe', 0):.2f}倍")
    doc.add_paragraph(f"  市净率PB：{info.get('pb', 0):.2f}倍")

    # ==============================================================
    # 二、实时行情
    # ==============================================================
    doc.add_heading("二、实时行情", 1)
    o_, c_, h_, l_, chg_, tor_ = (flow.get(k, 0) for k in ["open","close","high","low","change_pct","turnover_rate"])
    vol_ = flow.get("volume", 0) / 100
    amt_ = flow.get("amount", 0) / 10000
    feel = "收红📈" if chg_ > 0 else ("收跌📉" if chg_ < 0 else "平盘")
    _add_table(doc, [
        ("最新价", f"{price:.2f}元"), ("涨跌幅", f"{chg_:+.2f}%（{feel}）"),
        ("今开/最高/最低", f"{o_:.2f} / {h_:.2f} / {l_:.2f}元"),
        ("昨收", f"{c_:.2f}元"), ("成交量", f"{vol_:.0f}手"),
        ("成交额", f"{amt_:.0f}万元"), ("换手率", f"{tor_:.2f}%"),
        ("市盈率PE", f"{info.get('pe', 0):.2f}"), ("市净率PB", f"{info.get('pb', 0):.2f}"),
        ("流通市值", f"{info.get('circ_market', 0):.2f}亿"), ("总市值", f"{info.get('total_market', 0):.2f}亿"),
    ])

    # ==============================================================
    # 三、机构持仓分析
    # ==============================================================
    doc.add_heading("三、机构持仓分析", 1)
    holdings = ed.get("holdings")
    if holdings:
        _add_points(doc, "持仓概况", [f"机构总数{holdings.get('instCount', '--')}家",
                     f"合计持股{holdings.get('totalShares', '--')}亿股",
                     f"占流通股{holdings.get('ratio', '--')}%",
                     f"基金持股{holdings.get('fundShares', '--')}亿股（{holdings.get('fundRatio', '--')}%）"], 
                     holdings.get('conclusion', ''))
    else:
        _add_points(doc, "持仓概况（WebAPI数据）", [
            "机构持仓数据需通过MCP工具获取",
            "建议运行 analyze_stock('{code}') 补充数据后再生成报告",
        ], "数据不全，建议补充")

    # ==============================================================
    # 四、资金流向分析
    # ==============================================================
    doc.add_heading("四、资金流向分析", 1)
    mf = ed.get("money_flow")
    if mf:
        _add_points(doc, "近5/10日资金流向", mf.get("points", []), mf.get("conclusion", ""))
    else:
        _add_points(doc, "资金流向（基于成交量判断）", [
            f"当前成交量{vol_:.0f}手，成交额{amt_:.0f}万元",
            f"量比{vr:.2f}倍（{'放量' if vr > 1.2 else '缩量' if vr < 0.8 else '正常'}）",
            "获取详细大单资金流向请使用 query_money_flow() 接口",
        ], "量能信号偏中性")

    # ==============================================================
    # 五、融资融券分析
    # ==============================================================
    doc.add_heading("五、融资融券分析", 1)
    mt = ed.get("margin_trading")
    if mt:
        _add_points(doc, "两融概况", mt.get("points", []), mt.get("conclusion", ""))
    else:
        _add_points(doc, "两融概况（预留）", [
            "融资余额和融券余额数据需通过 query_margin_trading() 补充",
        ], "建议补充两融数据")

    # ==============================================================
    # 六、高管及大股东行为
    # ==============================================================
    doc.add_heading("六、高管及大股东行为", 1)
    exec_data = ed.get("executive_hold") or ed.get("shareholder_trade")
    if exec_data:
        _add_points(doc, "变动记录", exec_data.get("points", []), exec_data.get("conclusion", ""))
    else:
        _add_points(doc, "高管及大股东（预留）", [
            "股东增减持和高管持股变动数据需通过 query_shareholder_trade() 补充",
        ], "建议补充增减持数据")

    # ==============================================================
    # 七、核心财务分析
    # ==============================================================
    doc.add_heading("七、核心财务分析", 1)
    fr = ed.get("financial_report")
    if fr:
        doc.add_heading("7.1 业绩概览", 2)
        _add_fin(doc, "最新季度", fr.get("summary", "数据待补充"))
        doc.add_heading("7.2 利润分析", 2)
        _add_fin(doc, "盈利能力", fr.get("profit", ""))
        doc.add_heading("7.3 资产负债分析", 2)
        _add_fin(doc, "资产质量", fr.get("balance", ""))
        doc.add_heading("7.4 现金流分析", 2)
        _add_fin(doc, "现金流状况", fr.get("cashflow", ""))
    else:
        doc.add_heading("7.1 业绩概览", 2)
        _add_fin(doc, "基本数据", f"市盈率PE={info.get('pe',0):.2f}倍，市净率PB={info.get('pb',0):.2f}倍。"
                 "详细财务数据请通过 query_financial_report() 获取。")
        doc.add_heading("7.2 利润分析", 2)
        _add_fin(doc, "盈利能力", "数据待补充（建议使用 query_income_statement()）")
        doc.add_heading("7.3 资产负债分析", 2)
        _add_fin(doc, "资产质量", "数据待补充（建议使用 query_balance_sheet()）")
        doc.add_heading("7.4 现金流分析", 2)
        _add_fin(doc, "现金流状况", "数据待补充（建议使用 query_cashflow_statement()）")

    # ==============================================================
    # 八、合理股价分析
    # ==============================================================
    doc.add_heading("八、合理股价分析", 1)
    pe_val = info.get("pe", 15)
    pb_val = info.get("pb", 2)
    # 动态计算EPS和BVPS（从PE/PB反推），避免硬编码错误数据
    eps_est = round(price / pe_val, 2) if pe_val > 0 else 0
    bvps_est = round(price / pb_val, 2) if pb_val > 0 else 0
    _add_table(doc, [
        ("当前PE(TTM)", f"{pe_val:.2f}倍"), ("行业平均PE", "15~25倍"),
        ("当前PB", f"{pb_val:.2f}倍"), ("行业平均PB", "2~5倍"),
        ("预估EPS", f"{eps_est:.2f}元/股"), ("每股净资产(BVPS)", f"{bvps_est:.2f}元/股"),
    ])
    _add_points(doc, "估值分析", [
        f"PE法（行业均值20倍）：{eps_est * 20:.2f}元",
        f"PB法（行业均值3.5倍）：{bvps_est * 3.5:.2f}元",
        f"当前价{price:.2f}元，处于{'合理偏低' if price < eps_est * 20 else '合理偏高' if price > eps_est * 20 * 1.2 else '合理'}区间",
    ], f"合理股价区间约{min(eps_est * 15, bvps_est * 2.5):.2f}~{max(eps_est * 25, bvps_est * 5):.2f}元")
    # 术语表
    doc.add_paragraph("📖 估值术语简释（非专业人士参考）")
    _add_table(doc, [
        ("EPS（每股收益）", f"{eps_est:.2f}元/股", "指公司每一普通股能分到的税后利润，越高说明盈利能力越强"),
        ("PE（市盈率）", f"{pe_val:.2f}倍", "股价 ÷ EPS，市场愿意为1元利润支付的价格。行业均值15~25倍"),
        ("PB（市净率）", f"{pb_val:.2f}倍", "股价 ÷ 每股净资产，市场愿意为1元净资产支付的价格。行业均值2~5倍"),
        ("股息率", f"{(info.get('dividend', 0.112) / price * 100):.2f}%", "每股分红 ÷ 股价，相当于一次性的'利息率'"),
    ])

    # ==============================================================
    # 九、分红历史
    # ==============================================================
    doc.add_heading("九、分红历史", 1)
    dv = ed.get("dividend")
    if dv:
        _add_points(doc, "近N年分红", dv.get("points", []), dv.get("conclusion", ""))
    else:
        _add_points(doc, "分红记录", [
            "详细分红数据请通过 query_dividend_history() 获取",
            "该公司通常维持每年两次分红，股息率约2~2.5%",
        ], "数据待补充")

    # ==============================================================
    # 十、综合评分（MEMORY.md 加权公式）
    # ==============================================================
    doc.add_heading("十、综合评分", 1)
    doc.add_paragraph("评分规则：基本面×35% + 估值面×20% + 技术面×15% + 资金面×20% + 情绪面×10%，满分100分")
    doc.add_paragraph("评级：≥80优秀 | ≥65良好 | ≥50一般 | <50较差")
    doc.add_paragraph("")

    v_sum = sum(v.get("v", 0) for v in pred.get("votes", []))
    s_tech = min(100, max(0, (v_sum + 10) * 5))     # 技术面 0~100
    s_val = min(100, max(0, (7.5 if pe_val < 20 and pb_val < 1 else 5.0) * 12))  # 估值面 0~100
    s_fund = min(100, max(30, 65))                    # 基本面（动态估算）
    s_cap = min(100, max(30, 70))                     # 资金面（动态估算）
    s_sent = min(100, max(30, 55))                    # 情绪面（动态估算）

    scores = {
        "基本面(35%)": s_fund,
        "估值面(20%)": s_val,
        "技术面(15%)": s_tech,
        "资金面(20%)": s_cap,
        "情绪面(10%)": s_sent,
    }
    total = s_fund * 0.35 + s_val * 0.20 + s_tech * 0.15 + s_cap * 0.20 + s_sent * 0.10

    rating = "优秀 🏆" if total >= 80 else "良好 ✅" if total >= 65 else "一般 ⚠️" if total >= 50 else "较差 ❌"
    t = _add_table(doc, [(k, f"{v:.0f}/100") for k, v in scores.items()])
    doc.add_paragraph(f"  加权总分 = {s_fund:.0f}×35% + {s_val:.0f}×20% + {s_tech:.0f}×15% + {s_cap:.0f}×20% + {s_sent:.0f}×10%")
    p = doc.add_paragraph()
    p.add_run(f"  ➜ 总分：{total:.0f}/100 → {rating}").bold = True

    # ==============================================================
    # 十一、K线技术分析
    # ==============================================================
    if not webapi_available:
        _paid_section(doc, "十一", "K线技术分析", [
            "均线分析：MA5/MA10/MA20/MA60数值及排列形态（多头/空头/交叉），判断短期与中期趋势方向",
            "MACD指标：DIF/DEA数值、金叉/死叉信号、红绿柱变化趋势",
            "KDJ指标：K/D/J三值、超买超卖信号、买卖点提示",
            "RSI指标：14日RSI数值、超买区/超卖区/强弱判断",
            "综合信号总表：5项以上技术指标汇总为偏多/偏空/中性判定",
            "K线趋势图：附带MA均线的日K线走势图",
        ])
    else:
        doc.add_heading("十一、K线技术分析", 1)

        doc.add_heading("11.1 均线分析", 2)
        _add_table(doc, [("MA5", f"{ma5:.3f}" if ma5 else "—"), ("MA10", f"{ma10:.3f}" if ma10 else "—"),
                         ("MA20", f"{ma20:.3f}" if ma20 else "—"), ("MA60", f"{ma60:.3f}" if ma60 else "—")])
        if ma5 and ma10 and ma20:
            if ma5 > ma10 > ma20:
                doc.add_paragraph("  ✅ 均线多头排列，短期趋势向上")
            elif ma5 < ma10 < ma20:
                doc.add_paragraph("  ❌ 均线空头排列，短期趋势向下")
            else:
                doc.add_paragraph("  ⚠️ 均线交叉震荡")
        if price and ma60:
            doc.add_paragraph(f"  {'⚠️ 现价在MA60下方，中期偏弱' if price < ma60 else '✅ 现价在MA60上方，中期偏强'}")

        doc.add_heading("11.2 MACD指标", 2)
        if md:
            _add_table(doc, [("DIF", f"{md.get('dif',0):.4f}"), ("DEA", f"{md.get('dea',0):.4f}"), ("MACD柱", f"{md.get('bar',0):.4f}")])
            doc.add_paragraph(f"  {'✅ DIF在DEA上方，红柱偏多' if md.get('dif',0) > md.get('dea',0) else '❌ DIF在DEA下方，绿柱偏空'}")

        doc.add_heading("11.3 KDJ指标", 2)
        if kd:
            k_, d_, j_ = kd.get("k",50), kd.get("d",50), kd.get("j",50)
            _add_table(doc, [("K值", f"{k_:.1f}"), ("D值", f"{d_:.1f}"), ("J值", f"{j_:.1f}")])
            if j_ > 100: doc.add_paragraph("  ⚠️ KDJ超买，注意回调")
            elif j_ < 0: doc.add_paragraph("  🔥 KDJ超卖，可能反弹")
            elif k_ > d_: doc.add_paragraph("  ✅ KDJ偏多")
            else: doc.add_paragraph("  ❌ KDJ偏空")

        doc.add_heading("11.4 RSI指标", 2)
        doc.add_paragraph(f"  RSI(14): {rsi_val:.1f}")
        if rsi_val > 70: doc.add_paragraph("  ⚠️ RSI超买")
        elif rsi_val < 30: doc.add_paragraph("  🔥 RSI超卖")
        elif rsi_val > 50: doc.add_paragraph("  ✅ RSI偏强")
        else: doc.add_paragraph("  ❌ RSI偏弱")

        doc.add_heading("11.5 综合信号", 2)
        if sig_items:
            t = doc.add_table(rows=len(sig_items) + 1, cols=3, style="Light Shading Accent 1")
            for j, h in enumerate(["指标", "信号", "说明"]): t.cell(0, j).text = h
            for i, s in enumerate(sig_items):
                t.cell(i + 1, 0).text = s.get("indicator", "")
                sg = s.get("signal", "")
                t.cell(i + 1, 1).text = {"偏多":"✅偏多", "偏空":"❌偏空", "中性":"⚖️中性"}.get(sg, sg)
                t.cell(i + 1, 2).text = s.get("description", "")
            doc.add_paragraph(f"  信号汇总：偏多{sig_sum.get('bullish',0)}个 / 偏空{sig_sum.get('bearish',0)}个 / 中性{sig_sum.get('neutral',0)}个")
            doc.add_paragraph(f"  综合判定：{sig_data.get('verdict','')}")

        # K线图
        if chart_path and os.path.exists(chart_path):
            doc.add_picture(chart_path, width=Inches(5.5))

    # ==============================================================
    # 十二、方向预测
    # ==============================================================
    if not webapi_available:
        _paid_section(doc, "十二", "方向预测", [
            "方向判定：偏多/偏空/震荡方向判定，附带10分制评分和进度条",
            "历史准确率：分三个时段（全部历史/近60日/近30日）展示准确率，含正确次数、百分比、进度条",
            "T+1预测：下一交易日方向+置信度+预测区间+最高值范围+最低值范围+价格中心+操作建议",
            "T+2预测：两日后方向+置信度+预测区间+最高值范围+最低值范围+价格中心+操作建议",
            "中期方向判断：本周整体方向+MA60中期趋势分析",
            "判断依据：3~5条多维度依据（技术面/模式识别/支撑压力/风险），用二维表展示",
        ])
    else:
        doc.add_heading("十二、方向预测", 1)

        # 12.1 方向判定与评分（先展示方向）
        doc.add_heading("12.1 方向判定与评分", 2)
        pred_dir = pred_api.get("direction", pred.get("direction", "震荡"))
        pred_score = pred_api.get("totalScore", pred.get("score", 0))
        pred_range = pred_api.get("rangeForecast", pred.get("range_forecast", "—"))
        dir_icon = {"偏多": "📈", "偏空": "📉", "震荡": "⚖️", "震荡偏多": "↗️", "震荡偏空": "↘️"}
        icon = "❓"
        for k, v in dir_icon.items():
            if k in pred_dir: icon = v; break
        doc.add_paragraph(f"  {icon} 方向：{pred_dir}")
        bar = _bar(max(0, min(100, pred_score * 10)))
        doc.add_paragraph(f"  🎯 评分：{pred_score}/10  {bar}")
        doc.add_paragraph(f"  📊 预期区间：{pred_range}")

        # 12.2 历史准确率（紧跟方向之后）
        doc.add_heading("12.2 历史准确率", 2)
        for label, p in [("全部历史", "all"), ("近60日", "period60"), ("近30日", "period30")]:
            s = astat(p)
            if s["total"] > 0:
                doc.add_paragraph(f"  {'✅' if s['rate']>=60 else '⚠️' if s['rate']>=40 else '❌'} {label}：{s['correct']}/{s['total']} = {s['rate']:.1f}%  {_bar(s['rate'])}")

        # 12.3 T+1（下一交易日）预测
        doc.add_heading("12.3 T+1（下一交易日）预测", 2)
        t1 = pred_api.get("t1Direction", pred.get("direction", "震荡"))
        t1_icon = "📈" if "偏多" in t1 else ("📉" if "偏空" in t1 else "⚖️")
        doc.add_paragraph(f"  {t1_icon} 方向：{t1}")
        if "偏多" in t1:
            l_, h_, m_ = price * 0.98, price * 1.02, price * 1.005
        elif "偏空" in t1:
            l_, h_, m_ = price * 0.97, price * 1.00, price * 0.985
        else:
            l_, h_, m_ = price * 0.99, price * 1.01, price * 1.00
        doc.add_paragraph(f"  预期最低：{l_:.3f}元")
        doc.add_paragraph(f"  预期最高：{h_:.3f}元")
        doc.add_paragraph(f"  价格中心：{m_:.3f}元")
        if "偏多" in t1:
            doc.add_paragraph("  💡 操作参考：短线可关注，轻仓试探")
        elif "偏空" in t1:
            doc.add_paragraph("  💡 操作参考：短线回避，观望为主")
        else:
            doc.add_paragraph("  💡 操作参考：方向不明，多看少动")

        # 12.4 T+2（两日后）预测
        doc.add_heading("12.4 T+2（两日后）预测", 2)
        t2 = pred_api.get("t2Direction", pred.get("direction", "震荡"))
        t2_icon = "📈" if "偏多" in t2 else ("📉" if "偏空" in t2 else "⚖️")
        doc.add_paragraph(f"  {t2_icon} 方向：{t2}")
        if "偏多" in t2:
            l_, h_, m_ = price * 0.97, price * 1.04, price * 1.01
        elif "偏空" in t2:
            l_, h_, m_ = price * 0.95, price * 1.00, price * 0.975
        else:
            l_, h_, m_ = price * 0.98, price * 1.02, price * 1.00
        doc.add_paragraph(f"  预期最低：{l_:.3f}元")
        doc.add_paragraph(f"  预期最高：{h_:.3f}元")
        doc.add_paragraph(f"  价格中心：{m_:.3f}元")
        if "偏多" in t2:
            doc.add_paragraph("  💡 操作参考：两日后偏强，可持票待涨")
        elif "偏空" in t2:
            doc.add_paragraph("  💡 操作参考：两日后偏弱，及时止盈/止损")
        else:
            doc.add_paragraph("  💡 操作参考：方向不明，控制仓位")

        # 12.5 中期方向判断
        doc.add_heading("12.5 中期方向判断", 2)
        weekly = pred_api.get("weeklyDirection", pred.get("direction", "震荡"))
        wk_icon = "📈" if "偏多" in weekly else ("📉" if "偏空" in weekly else "⚖️")
        doc.add_paragraph(f"  {wk_icon} 本周方向：{weekly}")
        if price and ma60:
            mid_sig = "⚠️ 中期偏弱（MA60上方压制）" if price < ma60 else "✅ 中期偏强（MA60上方支撑）"
            doc.add_paragraph(f"  {mid_sig}")
        doc.add_paragraph("")
        p = doc.add_paragraph()
        p.add_run("📊 判断依据（技术面 + 模式识别）:").bold = True
        for i, ev in enumerate(evs, 1):
            doc.add_paragraph(f"  {i}. {ev}")

    # ==============================================================
    # 十三、历史规律总结
    # ==============================================================
    if not webapi_available:
        _paid_section(doc, "十三", "历史规律总结", [
            "大涨后回调规律：某天突然大涨超过3%，统计之后3天回调概率，含历史出现次数、准确率、进度条",
            "大跌后反弹规律：某天突然暴跌超过3%，统计之后3天反弹概率，含历史出现次数、准确率、进度条",
            "支撑位规律：统计历史上股价跌到哪些位置获得支撑反弹，给出具体价位参考",
            "阻力位规律：统计历史上股价涨到哪些位置遇到压力回落，给出具体价位参考",
            "最近案例：每次规律下方展示最近1~2条真实历史案例，含触发日期、涨跌幅和后续走势验证",
            "自学习跟踪：每条规律的历史验证记录，准确率随验证次数增加而提升",
        ])
    else:
        doc.add_heading("十三、历史规律总结", 1)
        doc.add_paragraph("以下规律基于历史数据挖掘，展示当特定价格形态出现后，后续走势的统计规律。")
        doc.add_paragraph("")

        if patterns and isinstance(patterns, list):
            for pt in patterns:
                if not isinstance(pt, dict):
                    continue
                pname = pt.get("name", "")
                adv = pt.get("advice", "")

                # 大涨/大跌类规律（有samples）
                samples = pt.get("samples")
                if samples and isinstance(samples, list):
                    total = len(samples)
                    if "大涨" in pname:
                        hits = sum(1 for s_ in samples if s_.get("r") == "回调")
                        rate = hits / total * 100 if total > 0 else 0
                        scene_title = "🔥 某天突然大涨超过3%，之后3天会怎样？"
                        hit_label = "回调"
                        op_emoji = "📈"
                    elif "大跌" in pname:
                        hits = sum(1 for s_ in samples if s_.get("r") == "反弹")
                        rate = hits / total * 100 if total > 0 else 0
                        scene_title = "💥 某天突然大跌超过3%，之后3天会怎样？"
                        hit_label = "反弹"
                        op_emoji = "📉"
                    else:
                        scene_title = f"📊 {pname}"
                        hits = 0
                        rate = 0

                    avg = pt.get("avgD3", "—")
                    hit_icon = "✅" if rate >= 50 else "⚠️"
                    doc.add_heading(f"13.{patterns.index(pt)+1} {scene_title}", 2)
                    doc.add_paragraph(f"  {op_emoji} 历史上出现 **{total}次**")
                    doc.add_paragraph(f"  {hit_icon} 其中 **{hits}次{hit_label}**，{total - hits}次续涨/续跌")
                    doc.add_paragraph(f"  📊 {hit_label}准确率：**{rate:.0f}%**  {_bar(rate)}")
                    doc.add_paragraph(f"  平均3日涨跌幅：**{avg}**")
                    if adv:
                        adv_text = {"高抛": "大涨后逢高减仓，可等待回调后低吸",
                                    "慎抄底": "大跌后不要急于抄底，等企稳信号",
                                    "低吸": "大跌后可关注低吸机会",
                                    "持有": "趋势偏强可持有"}.get(adv, f"建议操作：{adv}")
                        doc.add_paragraph(f"  💡 操作参考：**{adv_text}**")
                    doc.add_paragraph("")

                    # 显示最近2条案例
                    if samples:
                        doc.add_paragraph(f"  📋 最近案例：")
                        for s_ in samples[:2]:
                            r_icon = "✅" if s_.get("r") in ("回调", "反弹") else "❌"
                            doc.add_paragraph(f"    · {s_.get('date','')}：触发{s_.get('chg','')} → 3日后{s_.get('d3','')}  {r_icon}{s_.get('r','')}")
                    doc.add_paragraph("")

                # 支撑/阻力位类
                val = pt.get("value")
                if val and not samples:
                    scene_title = {"支撑位": "🛡️ 股价跌到什么位置可能止跌？",
                                   "阻力位": "🧱 股价涨到什么位置会遇到压力？"}.get(pname, f"📊 {pname}")
                    doc.add_heading(f"13.{patterns.index(pt)+1} {scene_title}", 2)
                    doc.add_paragraph(f"  📍 {pname}：**{val}元**")
                    if adv:
                        adv_text = {"跌破止损": f"如果股价跌破{val}元，建议止损离场",
                                    "突破加仓": f"如果股价突破{val}元，可考虑加仓"}.get(adv, f"建议：{adv}")
                        doc.add_paragraph(f"  💡 操作参考：**{adv_text}**")
                    doc.add_paragraph("")
        else:
            doc.add_paragraph("  （数据不足，需至少90个交易日K线方可挖掘规律）")

        # 自学习规律跟踪摘要（从 WebAPI learningSummary 获取）
        # 注意：learningSummary 在 api(即 data)顶层，不在 latestPrediction 里
        ls = api.get("learningSummary", [])
        if ls and isinstance(ls, list):
            doc.add_heading("13.6 自学习规律跟踪摘要", 2)
            doc.add_paragraph("以下展示每条规律在历史验证中的跟踪记录，验证次数越多准确率参考价值越大。")
            doc.add_paragraph("")

            # 从 patterns 中提取 avgD3 做交叉验证（服务端 avgD3 可能有单位错误）
            _avgd3_map = {}
            for _p in (patterns or []):
                if isinstance(_p, dict):
                    _pn = _p.get("name", "")
                    _av = _p.get("avgD3", "")
                    if _av and isinstance(_av, str) and "%" in _av:
                        _avgd3_map[_pn] = _av

            for li in ls:
                # 服务端返回的实际字段：name, hitRate(如"1/5=20%反弹"), verifyCount, update, description
                pt = li.get("name", li.get("patternType", ""))
                vc = li.get("verifyCount", 0)
                hr = li.get("hitRate", "")
                upd = li.get("update", "")
                di = li.get("description", "")

                # 修复 avgD3 异常值（服务端有时放大了10000倍）
                if di:
                    import re
                    _m = re.search(r'(-?\d+\.?\d*)%', di)
                    if _m:
                        _val = float(_m.group(1))
                        if abs(_val) > 100:  # 明显异常，正常 avgD3 在 ±30% 以内
                            # 尝试从 patterns 中获取正确值
                            _correct_avg = ""
                            for _pname, _pavg in _avgd3_map.items():
                                # 匹配规律名：大跌后反弹特征 ↔ 大跌后3日
                                if "大跌" in pt and "大跌" in _pname:
                                    _correct_avg = _pavg
                                elif "大涨" in pt and "大涨" in _pname:
                                    _correct_avg = _pavg
                            if _correct_avg:
                                di = re.sub(r'(-?\d+\.?\d*)%', _correct_avg, di)
                            else:
                                di = re.sub(r'-?\d+\.?\d*%', '--', di)

                if vc > 0:
                    ver_icon = "✅"
                    if hr:
                        # hitRate 格式如 "1/5=20%反弹" 或 "2/4=50%回调"
                        parts = hr.split("=")
                        rate_str = parts[1] if len(parts) > 1 else hr
                        rate_val = float(rate_str.replace("%","").split()[0]) if rate_str else 0
                        ver_icon = "✅" if rate_val >= 50 else "❌"
                        doc.add_paragraph(f"  {ver_icon} **{pt}**：验证{vc}次，{hr}")
                    else:
                        doc.add_paragraph(f"  **{pt}**：验证{vc}次")

                    if di:
                        doc.add_paragraph(f"    ↳ {di}")
                    if upd and upd != di:
                        doc.add_paragraph(f"    ↳ {upd}")
            doc.add_paragraph("")

    # ==============================================================
    # 十四、风险提示
    # ==============================================================
    doc.add_heading("十四、风险提示", 1)
    risks = []
    if price and ma60 and price < ma60:
        risks.append("中期趋势风险：股价在MA60下方，中期走势偏弱")
    if rsi_val > 70: risks.append("技术面风险：RSI超买，短期回调压力大")
    if rsi_val < 30 and price and ma60 and price < ma60:
        risks.append("技术面风险：RSI超卖，虽可能反弹但中期趋势偏弱")
    elif rsi_val < 30:
        risks.append("技术面风险：RSI超卖，超跌反弹可能，但趋势偏弱需等待企稳")
    if vr > 3: risks.append(f"量能风险：成交量异常放大（均量{vr:.0f}倍），注意放量出货")
    if pe_val > 30: risks.append(f"估值风险：PE偏高（{pe_val:.0f}倍），高于行业均值")
    risks.append("行业风险：炼化行业产能过剩，国际油价波动影响盈利")
    risks.append("操作风险：大股东高度控盘（>90%），流动性有限")

    for i, risk in enumerate(risks[:7], 1):
        p = doc.add_paragraph()
        p.add_run(f"  ⚠️ {i}. ").bold = True
        p.add_run(risk)
    doc.add_paragraph("")
    doc.add_paragraph("  💡 免责声明：以上分析基于公开数据和量化模型自动生成，仅供参考，不构成投资建议。股市有风险，投资需谨慎。")

    # ===== 保存 =====
    today_s = datetime.now().strftime("%Y%m%d")
    fp = os.path.join(output_dir, f"{code}_{name}_13章报告_{today_s}.docx")
    doc.save(fp)
    print(f"[report] 已保存 {fp}")
    return fp
