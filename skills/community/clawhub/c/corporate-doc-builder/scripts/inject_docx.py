#!/usr/bin/env python3
"""
inject_docx.py — Inject Markdown content into a .docx template.

Copies the template, clears body content after the TOC, and injects
Markdown as properly styled Word elements. Handles headings, tables,
bold/code inline formatting, Mermaid PNG embedding, non-Mermaid code
block rendering, and template style fixes.

Usage:
    python inject_docx.py \
        --md-dir ./chapters_md \
        --template ./template.docx \
        --output ./output.docx \
        --chapters ch01.md ch02.md appendix.md \
        [--header-replace "XXX=My Project"] \
        [--max-width 14.0] \
        [--max-height 12.0] \
        [--comparison-output ./output_comparison.docx]
"""
import argparse
import os
import re
import shutil
import zipfile

import docx
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from PIL import Image

WNS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


# ---------------------------------------------------------------------------
# Template fixes
# ---------------------------------------------------------------------------

def strip_heading_auto_numbering(doc):
    """Remove numPr from all Heading styles to prevent double numbering."""
    for style in doc.styles:
        if style.name and style.name.startswith("Heading"):
            pPr = style.element.find(f"{WNS}pPr")
            if pPr is not None:
                numPr = pPr.find(f"{WNS}numPr")
                if numPr is not None:
                    pPr.remove(numPr)


def clear_body_after_toc(doc):
    """Remove all paragraphs and tables after the last TOC entry."""
    last_toc = -1
    for i, p in enumerate(doc.paragraphs):
        sty = (p.style.name or "").lower()
        if sty.startswith("toc"):
            last_toc = i
    if last_toc >= 0:
        for p in list(doc.paragraphs[last_toc + 1 :]):
            p._element.getparent().remove(p._element)
    for t in list(doc.tables):
        t._element.getparent().remove(t._element)


def replace_header_footer_text(doc, replacements):
    """Replace placeholder text in all section headers and footers.

    Args:
        replacements: dict mapping old text to new text,
                      e.g. {"XXX": "My Project Name"}
    """
    if not replacements:
        return
    for section in doc.sections:
        for header_para in section.header.paragraphs:
            for run in header_para.runs:
                for old, new in replacements.items():
                    if old in run.text:
                        run.text = run.text.replace(old, new)
        for footer_para in section.footer.paragraphs:
            for run in footer_para.runs:
                for old, new in replacements.items():
                    if old in run.text:
                        run.text = run.text.replace(old, new)


# ---------------------------------------------------------------------------
# Element builders
# ---------------------------------------------------------------------------

def image_size_cm(img_path, max_w_cm, max_h_cm):
    """Compute image dimensions in cm, respecting max constraints."""
    with Image.open(img_path) as im:
        w_px, h_px = im.size
    aspect = w_px / h_px
    w = max_w_cm
    h = w / aspect
    if h > max_h_cm:
        h = max_h_cm
        w = h * aspect
    return w, h


def add_image(doc, img_path, max_w_cm, max_h_cm):
    """Add a centered image paragraph with SINGLE line spacing."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run()
    w_cm, h_cm = image_size_cm(img_path, max_w_cm, max_h_cm)
    run.add_picture(img_path, width=Cm(w_cm), height=Cm(h_cm))


def add_heading(doc, text, level):
    """Add a heading paragraph, clamping level to 1-3."""
    if level > 3:
        level = 3
    doc.add_paragraph(text, style=f"Heading {level}")


def add_paragraph(doc, text):
    """Add a Normal paragraph with inline bold and code formatting."""
    text = text.rstrip()
    if not text or text == "---":
        return
    is_list = text.startswith("- ") or text.startswith("* ")
    if is_list:
        try:
            p = doc.add_paragraph(style="List Paragraph")
        except Exception:
            p = doc.add_paragraph(style="Normal")
        text = text[2:]
    else:
        p = doc.add_paragraph(style="Normal")
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = "Consolas"
        else:
            p.add_run(part)


def add_code_block(doc, lines):
    """Render a fenced code block as a shaded monospace paragraph."""
    code_text = "\n".join(lines)
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    shd = docx.oxml.OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_table(doc, lines):
    """Parse markdown pipe-delimited table lines into a Word table."""
    rows = []
    for ln in lines:
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        if cells and "---" in cells[0]:
            continue
        rows.append([c.replace("**", "") for c in cells])
    if not rows:
        return
    cols = len(rows[0])
    rows = [
        r if len(r) == cols else (r + [""] * (cols - len(r)))[:cols] for r in rows
    ]
    tbl = doc.add_table(rows=len(rows), cols=cols)
    try:
        tbl.style = "Table Grid"
    except Exception:
        pass
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            tbl.cell(i, j).text = val


# ---------------------------------------------------------------------------
# Markdown parser
# ---------------------------------------------------------------------------

def inject_markdown(doc, md_path, img_dir, max_w_cm, max_h_cm):
    """Parse a Markdown file and inject its content into a Document."""
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    in_mermaid = False
    in_code = False
    in_table = False
    table_lines = []
    code_lines = []
    diagram_counter = 1

    for line in lines:
        s = line.rstrip()
        stripped = s.strip()

        # Fenced block boundary
        if stripped.startswith("```"):
            if not in_mermaid and not in_code:
                if "mermaid" in stripped:
                    in_mermaid = True
                    img_path = os.path.join(img_dir, f"diagram_{diagram_counter}.png")
                    if os.path.exists(img_path):
                        if in_table:
                            add_table(doc, table_lines)
                            in_table = False
                            table_lines = []
                        add_image(doc, img_path, max_w_cm, max_h_cm)
                    diagram_counter += 1
                else:
                    in_code = True
                    code_lines = []
            else:
                if in_code and code_lines:
                    if in_table:
                        add_table(doc, table_lines)
                        in_table = False
                        table_lines = []
                    add_code_block(doc, code_lines)
                in_mermaid = False
                in_code = False
                code_lines = []
            continue

        if in_mermaid:
            continue
        if in_code:
            code_lines.append(s)
            continue

        # Table
        if stripped.startswith("|") and stripped.endswith("|"):
            in_table = True
            table_lines.append(stripped)
            continue
        elif in_table:
            add_table(doc, table_lines)
            in_table = False
            table_lines = []

        # Heading
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            title = stripped.lstrip("#").strip()
            add_heading(doc, title, level)
            continue

        # Blockquote
        if stripped.startswith(">"):
            p = doc.add_paragraph(style="Normal")
            run = p.add_run(stripped.lstrip(">").strip())
            run.italic = True
            continue

        add_paragraph(doc, stripped)

    if in_table:
        add_table(doc, table_lines)


# ---------------------------------------------------------------------------
# Post-processing
# ---------------------------------------------------------------------------

def force_update_fields(docx_path):
    """Inject updateFields=true into settings.xml so Word refreshes TOC on open."""
    tmp = docx_path + ".tmp"
    settings_path = "word/settings.xml"
    with zipfile.ZipFile(docx_path, "r") as zin:
        names = zin.namelist()
        has_settings = settings_path in names
        with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
            for n in names:
                data = zin.read(n)
                if n == settings_path:
                    s = data.decode("utf-8")
                    if "<w:updateFields" not in s:
                        s = re.sub(
                            r"(<w:settings\b[^>]*>)",
                            r'\1<w:updateFields w:val="true"/>',
                            s,
                            count=1,
                        )
                    data = s.encode("utf-8")
                zout.writestr(n, data)
            if not has_settings:
                settings_xml = (
                    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                    '<w:settings xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                    '<w:updateFields w:val="true"/></w:settings>'
                )
                zout.writestr(settings_path, settings_xml)
    os.replace(tmp, docx_path)


# ---------------------------------------------------------------------------
# Orchestration
# ---------------------------------------------------------------------------

def concat_chapters(md_dir, chapters, full_md_path):
    """Concatenate chapter files into a single Markdown file."""
    with open(full_md_path, "w", encoding="utf-8") as fo:
        for fn in chapters:
            p = os.path.join(md_dir, fn)
            with open(p, "r", encoding="utf-8") as fi:
                fo.write(fi.read().rstrip() + "\n\n")
    print(f"Concatenated {len(chapters)} chapters -> {full_md_path} ({os.path.getsize(full_md_path)} bytes)")


def process(template_path, output_path, full_md, img_dir, header_replacements,
            max_w_cm, max_h_cm, is_blank=False):
    """Run the full injection pipeline."""
    if not is_blank:
        shutil.copy(template_path, output_path)
        doc = docx.Document(output_path)
        for section in doc.sections:
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
        replace_header_footer_text(doc, header_replacements)
        clear_body_after_toc(doc)
        strip_heading_auto_numbering(doc)
    else:
        doc = docx.Document()

    inject_markdown(doc, full_md, img_dir, max_w_cm, max_h_cm)
    doc.save(output_path)

    if not is_blank:
        force_update_fields(output_path)
    print(f"Generated: {output_path}")


def parse_header_replacements(values):
    """Parse 'OLD=NEW' strings into a dict."""
    result = {}
    if not values:
        return result
    for v in values:
        if "=" in v:
            old, new = v.split("=", 1)
            result[old] = new
    return result


def main():
    parser = argparse.ArgumentParser(
        description="Inject Markdown into a .docx template"
    )
    parser.add_argument("--md-dir", required=True,
                        help="Directory containing chapter Markdown files")
    parser.add_argument("--template", required=True,
                        help="Path to the .docx template file")
    parser.add_argument("--output", required=True,
                        help="Path for the output .docx file")
    parser.add_argument("--chapters", nargs="+", required=True,
                        help="Ordered list of chapter filenames (relative to --md-dir)")
    parser.add_argument("--header-replace", nargs="*", default=[],
                        help="Header/footer text replacements as OLD=NEW pairs")
    parser.add_argument("--max-width", type=float, default=14.0,
                        help="Maximum image width in cm (default: 14.0)")
    parser.add_argument("--max-height", type=float, default=12.0,
                        help="Maximum image height in cm (default: 12.0)")
    parser.add_argument("--comparison-output", default=None,
                        help="Optional path for a blank comparison .docx (no template styles)")
    args = parser.parse_args()

    img_dir = os.path.join(args.md_dir, "images")
    full_md = os.path.join(args.md_dir, "full_draft.md")

    header_replacements = parse_header_replacements(args.header_replace)

    concat_chapters(args.md_dir, args.chapters, full_md)
    process(args.template, args.output, full_md, img_dir,
            header_replacements, args.max_width, args.max_height, is_blank=False)

    if args.comparison_output:
        process(None, args.comparison_output, full_md, img_dir,
                {}, args.max_width, args.max_height, is_blank=True)


if __name__ == "__main__":
    main()
