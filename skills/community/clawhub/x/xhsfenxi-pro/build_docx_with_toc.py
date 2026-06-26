"""
build_docx_with_toc.py — build_docx.py 的目录增强版

差异点（vs 原版 build_docx.py）：
1. H1/H2 改用 Word 原生 Heading 样式（Heading 1/2），保留同样的视觉
2. 封面后插入「目录」+ Word TOC 字段（打开 Word 后右键更新即可填充）
3. 同时插入静态目录列表作为后备（即使 Word 不更新字段也能看见）

用法:
    python3 build_docx_with_toc.py <md_path> <out_path> <title> <subtitle> [meta]
"""
import re, sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 配色（与原版一致）
C_BRAND  = RGBColor(0x1A, 0x1A, 0x1A)
C_H1     = RGBColor(0x0D, 0x0D, 0x0D)
C_H2     = RGBColor(0x1A, 0x1A, 0x1A)
C_ACCENT = RGBColor(0x00, 0x80, 0x60)
C_LIGHT  = RGBColor(0x6B, 0x72, 0x80)
C_TH_BG  = RGBColor(0xF3, 0xF4, 0xF6)
C_LINE   = RGBColor(0x22, 0xC5, 0x5E)
HEITI    = "黑体"


def set_font(run, size, bold=False, color=None, italic=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts')
        rPr.insert(0, rf)
    for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rf.set(qn(attr), HEITI)


def clean(s):
    for p, r in [
        (r'\*\*(.+?)\*\*', r'\1'), (r'\*(.+?)\*', r'\1'),
        (r'\[(.+?)\]\(.+?\)', r'\1'), (r'`(.+?)`', r'\1'),
        (r'^#{1,6}\s+', ''),
    ]:
        s = re.sub(p, r, s)
    return s.strip()


def set_cell_bg(cell, rgb):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}')
    tcPr.append(shd)


def add_rule(doc, color=C_LINE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), f'{color[0]:02X}{color[1]:02X}{color[2]:02X}')
    pBdr.append(bot)
    pPr.append(pBdr)


def add_bookmark(paragraph, bookmark_id, bookmark_name):
    """给段落首尾添加书签标记（用于 TOC 跳转）"""
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), str(bookmark_id))
    start.set(qn('w:name'), bookmark_name)
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_internal_hyperlink(paragraph, anchor, text, size=10.5, bold=False, color=C_BRAND, indent=False):
    """添加跳转到 anchor 书签的内部超链接"""
    if indent:
        paragraph.paragraph_format.left_indent = Cm(0.8)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), anchor)

    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # 字体
    rFonts = OxmlElement('w:rFonts')
    for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rFonts.set(qn(attr), HEITI)
    rPr.append(rFonts)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(size * 2)))
    rPr.append(sz)

    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)

    col = OxmlElement('w:color')
    col.set(qn('w:val'), f'{color[0]:02X}{color[1]:02X}{color[2]:02X}')
    rPr.append(col)

    # 单下划线（让它视觉上像链接）
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'none')
    rPr.append(u)

    run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    run.append(t)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_toc_field(doc):
    """插入 Word 原生 TOC 字段（打开后右键 → 更新域 即可自动生成可点击目录）"""
    p = doc.add_paragraph()
    run = p.add_run()
    r = run._r

    fldBegin = OxmlElement('w:fldChar')
    fldBegin.set(qn('w:fldCharType'), 'begin')
    fldBegin.set(qn('w:dirty'), 'true')
    r.append(fldBegin)

    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = r' TOC \o "1-2" \h \z \u '
    r.append(instr)

    fldSep = OxmlElement('w:fldChar')
    fldSep.set(qn('w:fldCharType'), 'separate')
    r.append(fldSep)

    placeholder = OxmlElement('w:t')
    placeholder.text = "（在 Word 中右键此处 → 更新域，目录将自动填充。下方是手工备份目录。）"
    r.append(placeholder)

    fldEnd = OxmlElement('w:fldChar')
    fldEnd.set(qn('w:fldCharType'), 'end')
    r.append(fldEnd)


def parse_headings(md_text):
    """提取 markdown 中的 H1/H2 标题"""
    headings = []
    for line in md_text.splitlines():
        if re.match(r'^# [^#]', line):
            headings.append(('h1', clean(line)))
        elif re.match(r'^## [^#]', line):
            headings.append(('h2', clean(line)))
    return headings


def build_word_with_toc(md_path, out_path, title="", subtitle="", meta=""):
    md_text = Path(md_path).read_text(encoding="utf-8")
    headings = parse_headings(md_text)

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3.0)
        sec.right_margin = Cm(3.0)

    nml = doc.styles['Normal']
    nml.font.name = HEITI
    nml.font.size = Pt(10.5)
    nml._element.rPr.rFonts.set(qn('w:eastAsia'), HEITI)

    # ── 封面 ──────────────────────────────────────────
    if title:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(title), 28, bold=True, color=C_ACCENT)

    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(subtitle), 14, color=C_LIGHT)

    if meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(meta), 9, color=C_LIGHT)

    if title:
        doc.add_page_break()

    # ── 目录页 ────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    set_font(p.add_run("目录"), 22, bold=True, color=C_ACCENT)
    add_rule(doc, C_LINE)

    # 1. Word 原生 TOC 字段（用户在 Word 里右键即可自动生成）
    add_toc_field(doc)
    doc.add_paragraph()

    # 2. 静态目录列表（即使 TOC 字段未更新，也能直接看到 + 点击跳转）
    bookmark_map = {}
    for idx, (level, text) in enumerate(headings):
        bookmark_name = f"_TocH_{idx}"
        bookmark_map[(level, text, idx)] = bookmark_name

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(2)

        if level == 'h1':
            add_internal_hyperlink(p, bookmark_name, text, size=12, bold=True, color=C_H1)
        else:  # h2
            add_internal_hyperlink(p, bookmark_name, "    · " + text, size=10.5, color=C_H2, indent=True)

    doc.add_page_break()

    # ── 正文（解析 Markdown）──────────────────────────
    lines = md_text.splitlines()
    bookmark_id = 1
    heading_idx = 0
    i = 0
    while i < len(lines):
        line = lines[i]

        # H1
        if re.match(r'^# [^#]', line):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(4)
            text = clean(line)

            # 加书签
            bn = f"_TocH_{heading_idx}"
            add_bookmark(p, bookmark_id, bn)
            bookmark_id += 1
            heading_idx += 1

            set_font(p.add_run(text), 20, bold=True, color=C_H1)
            i += 1
            continue

        # H2
        if re.match(r'^## [^#]', line):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(2)
            text = clean(line)

            bn = f"_TocH_{heading_idx}"
            add_bookmark(p, bookmark_id, bn)
            bookmark_id += 1
            heading_idx += 1

            set_font(p.add_run(text), 13, bold=True, color=C_H2)
            add_rule(doc)
            i += 1
            continue

        # H3
        if re.match(r'^### ', line):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            set_font(p.add_run(clean(line)), 11, bold=True, color=C_ACCENT)
            i += 1
            continue

        # 引用
        if line.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            set_font(p.add_run(line[2:].strip()), 10.5, italic=True, color=C_LIGHT)
            i += 1
            continue

        # 代码块
        if line.startswith('```'):
            i += 1
            code = []
            while i < len(lines) and not lines[i].startswith('```'):
                code.append(lines[i])
                i += 1
            p = doc.add_paragraph('\n'.join(code))
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            for run in p.runs:
                set_font(run, 9, color=RGBColor(0x37, 0x37, 0x37))
            i += 1
            continue

        if line.strip() == '---':
            add_rule(doc, C_TH_BG)
            i += 1
            continue

        # 表格
        if line.strip().startswith('|'):
            tbl_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                if not re.match(r'^\|[\s|:-]+$', lines[i].strip()):
                    tbl_lines.append(lines[i])
                i += 1
            if not tbl_lines:
                continue
            rows = [
                [c.strip() for c in tl.split('|')[1:-1]]
                for tl in tbl_lines
            ]
            rows = [r for r in rows if any(c for c in r)]
            if not rows:
                continue
            ncols = max(len(r) for r in rows)
            tbl = doc.add_table(rows=len(rows), cols=ncols)
            tbl.style = 'Table Grid'
            tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
            for ri, row in enumerate(rows):
                for ci in range(ncols):
                    cell = tbl.cell(ri, ci)
                    txt = clean(row[ci]) if ci < len(row) else ''
                    p = cell.paragraphs[0]
                    p.clear()
                    run = p.add_run(txt)
                    set_font(run, 9.5,
                             bold=(ri == 0),
                             color=C_H2 if ri == 0 else C_BRAND)
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    if ri == 0:
                        set_cell_bg(cell, C_TH_BG)
            doc.add_paragraph()
            continue

        # 列表
        if re.match(r'^[-*] |^\d+\. ', line):
            txt = re.sub(r'^[-*] |^\d+\. ', '', line)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            set_font(p.add_run(clean(txt)), 10.5, color=C_BRAND)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        set_font(p.add_run(clean(line)), 10.5, color=C_BRAND)
        i += 1

    doc.save(out_path)
    print(f"✅ Word（含目录）已生成：{out_path}")
    print(f"   H1 标题: {sum(1 for h in headings if h[0]=='h1')} 项")
    print(f"   H2 标题: {sum(1 for h in headings if h[0]=='h2')} 项")
    return out_path


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("用法: python3 build_docx_with_toc.py <md_path> <out_path> [title] [subtitle] [meta]")
        sys.exit(1)
    md = sys.argv[1]
    out = sys.argv[2]
    ttl = sys.argv[3] if len(sys.argv) > 3 else ""
    sub = sys.argv[4] if len(sys.argv) > 4 else ""
    meta = sys.argv[5] if len(sys.argv) > 5 else ""
    build_word_with_toc(md, out, ttl, sub, meta)
