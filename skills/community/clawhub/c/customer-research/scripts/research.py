#!/usr/bin/env python3
"""
客户调研脚本 - 围绕 AI 领导发言、招标记录、数据相关、基本情况
输出：Word 调研报告 + 破冰 PPT（3-5 页客户视角）

兼容：Hermes Agent, OpenClaw, 或其他支持 web 工具的 agent framework
"""

import os
import sys
import json
from datetime import datetime
from pathlib import Path

# ---------------------------------------------------------------------------
# 环境检测与工具适配
# ---------------------------------------------------------------------------

def get_workspace_dir():
    """获取工作目录，兼容不同 agent framework"""
    # 优先使用环境变量
    if os.getenv("HERMES_HOME"):
        return Path(os.getenv("HERMES_HOME"))
    elif os.getenv("OPENCLAW_HOME"):
        return Path(os.getenv("OPENCLAW_HOME")) / "workspace"
    else:
        # 默认当前用户主目录
        return Path.home() / ".hermes"

def get_output_dir():
    """获取输出目录"""
    workspace = get_workspace_dir()
    output_dir = workspace / "customer_research"
    output_dir.mkdir(parents=True, exist_ok=True)
    return output_dir

def search_web(query, count=5):
    """
    使用 web_fetch 搜索 - 兼容 Hermes 和 OpenClaw
    
    在 Hermes 环境下，调用方应使用 web_search/web_fetch 工具
    在 OpenClaw 环境下，调用方应使用 web-fetch 命令
    
    这里返回搜索指令，由 agent framework 执行
    """
    # 返回搜索指令，由上层 agent 执行
    # 这样避免了 subprocess 调用
    return {
        "action": "web_search",
        "query": query,
        "count": count,
        "note": "此搜索指令由 agent framework 执行，非本脚本直接调用"
    }

# ---------------------------------------------------------------------------
# 调研逻辑
# ---------------------------------------------------------------------------

def get_search_dimensions(customer_name):
    """生成搜索维度"""
    return [
        {
            "name": "AI 相关领导发言",
            "keywords": [
                f"{customer_name} AI 讲话",
                f"{customer_name} 人工智能 指示",
                f"{customer_name} 数字化 转型 领导"
            ]
        },
        {
            "name": "AI 相关招标记录",
            "keywords": [
                f"{customer_name} AI 招标",
                f"{customer_name} 人工智能 采购",
                f"{customer_name} 智能 中标"
            ]
        },
        {
            "name": "数据相关内容",
            "keywords": [
                f"{customer_name} 数据治理",
                f"{customer_name} 数据中台",
                f"{customer_name} 数据分析"
            ]
        },
        {
            "name": "基本情况与最新动态",
            "keywords": [
                f"{customer_name} 官网",
                f"{customer_name} 简介",
                f"{customer_name} 最新动态"
            ]
        }
    ]

def research_customer(customer_name, search_results=None):
    """
    执行客户调研
    
    Args:
        customer_name: 客户名称
        search_results: 可选，agent framework 执行搜索后的实际结果字典
                       格式：{"AI 相关领导发言": ["结果 1", "结果 2", ...], ...}
    
    Returns:
        结果字典，包含实际搜索内容或搜索指令
    """
    dimensions = get_search_dimensions(customer_name)
    
    # 如果已提供实际搜索结果，直接使用
    if search_results and isinstance(search_results, dict):
        results = {}
        for dim in dimensions:
            dim_name = dim["name"]
            results[dim_name] = search_results.get(dim_name, ["公开渠道未检索到相关信息"])
        return results
    
    # 否则返回搜索指令（由 agent framework 执行后再次调用）
    results = {}
    for dim in dimensions:
        results[dim["name"]] = []
        for keyword in dim["keywords"]:
            search_cmd = search_web(keyword)
            results[dim["name"]].append({
                "keyword": keyword,
                "search_instruction": search_cmd,
                "status": "pending",
                "note": "⚠️ 此搜索指令必须由 agent framework 执行后，再次调用本函数传入实际结果"
            })
    return results

# ---------------------------------------------------------------------------
# 报告生成
# ---------------------------------------------------------------------------

def generate_word_report(customer_name, results, output_path):
    """创建 Word 调研报告"""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        # 如果 python-docx 未安装，生成 Markdown 格式
        return generate_word_report_markdown(customer_name, results, output_path)
    
    doc = Document()
    
    # 标题
    title = doc.add_heading(f'客户调研报告：{customer_name}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 基本信息
    doc.add_paragraph(f'**调研时间：** {datetime.now().strftime("%Y 年 %m 月 %d 日")}')
    doc.add_paragraph(f'**调研人：** AI Assistant')
    doc.add_paragraph()
    
    # 执行摘要
    doc.add_heading('📌 执行摘要', level=1)
    doc.add_paragraph(f'本次调研围绕 {customer_name} 的 AI 相关领导发言、招标记录、数据相关内容及基本情况展开。')
    doc.add_paragraph()
    
    # 各维度详情
    sections = [
        ("1️⃣ AI 相关领导发言", "AI 相关领导发言"),
        ("2️⃣ AI 相关招标记录", "AI 相关招标记录"),
        ("3️⃣ 数据相关内容", "数据相关内容"),
        ("4️⃣ 基本情况与最新动态", "基本情况与最新动态")
    ]
    
    for emoji_name, key_name in sections:
        doc.add_heading(emoji_name, level=2)
        items = results.get(key_name, [])
        if items:
            for i, item in enumerate(items, 1):
                if isinstance(item, dict):
                    # 如果是待执行的搜索指令
                    p = doc.add_paragraph(style='List Number')
                    p.add_run(f'关键词：{item.get("keyword", "N/A")}（待搜索）')
                else:
                    # 如果是已填充的文本
                    p = doc.add_paragraph(style='List Number')
                    clean_text = str(item)[:300].replace('\n', ' ').strip()
                    p.add_run(f'{clean_text}...')
        else:
            doc.add_paragraph('暂无相关公开信息', style='Intense Quote')
        doc.add_paragraph()
    
    # 销售机会分析
    doc.add_heading('🎯 销售机会分析', level=1)
    doc.add_paragraph('可能的痛点：', style='Heading 2')
    pain_points = [
        '多院区数据孤岛问题',
        '业务多元化需要综合医疗信息系统',
        '数字化转型压力',
        '三级医院升级需要配套信息化系统'
    ]
    for point in pain_points:
        doc.add_paragraph(point, style='List Bullet')
    
    doc.add_paragraph('决策链推测：', style='Heading 2')
    doc.add_paragraph('• EB（经济决策人）：待确认')
    doc.add_paragraph('• Technical Buyer：信息科主任/CTO（待拜访确认）')
    doc.add_paragraph('• Champion：待识别')
    
    doc.add_paragraph('切入建议：', style='Heading 2')
    suggestions = [
        '首访目标：信息科/IT 部门负责人',
        '切入点：多院区数据治理、智能问数、3D 可视化',
        'MEDDIC 行动：确认预算、决策流程、培养 Champion'
    ]
    for sug in suggestions:
        doc.add_paragraph(sug, style='List Bullet')
    
    # 信息来源
    doc.add_heading('📎 信息来源', level=1)
    doc.add_paragraph('本报告基于公开网络搜索整理，主要来源包括：')
    doc.add_paragraph('• 百度百科', style='List Bullet')
    doc.add_paragraph('• Bing 搜索引擎', style='List Bullet')
    doc.add_paragraph('• 公开新闻报道', style='List Bullet')
    
    # 保存文档
    doc.save(output_path)
    print(f"✅ Word 报告已生成：{output_path}")
    return output_path

def generate_word_report_markdown(customer_name, results, output_path):
    """生成 Markdown 格式报告（当 python-docx 未安装时）"""
    content = f"""# 客户调研报告：{customer_name}

**调研时间：** {datetime.now().strftime("%Y 年 %m 月 %d 日")}
**调研人：** AI Assistant

---

## 📌 执行摘要

本次调研围绕 {customer_name} 的 AI 相关领导发言、招标记录、数据相关内容及基本情况展开。

---

## 1️⃣ AI 相关领导发言

"""
    for item in results.get("AI 相关领导发言", []):
        if isinstance(item, dict):
            content += f"- 关键词：{item.get('keyword', 'N/A')}（待搜索）\n"
        else:
            content += f"- {item}\n"
    
    content += """
---

## 2️⃣ AI 相关招标记录

"""
    for item in results.get("AI 相关招标记录", []):
        if isinstance(item, dict):
            content += f"- 关键词：{item.get('keyword', 'N/A')}（待搜索）\n"
        else:
            content += f"- {item}\n"
    
    content += """
---

## 3️⃣ 数据相关内容

"""
    for item in results.get("数据相关内容", []):
        if isinstance(item, dict):
            content += f"- 关键词：{item.get('keyword', 'N/A')}（待搜索）\n"
        else:
            content += f"- {item}\n"
    
    content += """
---

## 4️⃣ 基本情况与最新动态

"""
    for item in results.get("基本情况与最新动态", []):
        if isinstance(item, dict):
            content += f"- 关键词：{item.get('keyword', 'N/A')}（待搜索）\n"
        else:
            content += f"- {item}\n"
    
    content += """
---

## 🎯 销售机会分析

**可能的痛点：**
- 多院区数据孤岛问题
- 业务多元化需要综合医疗信息系统
- 数字化转型压力

**决策链推测：**
- EB（经济决策人）：待确认
- Technical Buyer：信息科主任/CTO（待拜访确认）
- Champion：待识别

**切入建议：**
- 首访目标：信息科/IT 部门负责人
- 切入点：多院区数据治理、智能问数、3D 可视化
- MEDDIC 行动：确认预算、决策流程、培养 Champion

---

## 📎 信息来源

本报告基于公开网络搜索整理。
"""
    
    # 使用安全的文件写入
    output_path = Path(output_path).with_suffix('.md')
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(content)
    
    print(f"✅ Markdown 报告已生成：{output_path}")
    return str(output_path)

def generate_research_results_json(customer_name, search_results, output_path):
    """将调研结果保存为 JSON 文件，供 PPT 生成脚本使用。
    
    ⚠️ ASI09 合规：此函数是调研结果的标准输出格式。
    PPT 生成脚本必须读取此 JSON 文件来生成内容，
    禁止跳过此步骤直接生成通用模板。
    
    Args:
        customer_name: 客户名称
        search_results: dict，调研搜索结果
        output_path: 输出 JSON 文件路径
    
    Returns:
        str: 保存的 JSON 文件路径
    """
    research_data = {
        "customer_name": customer_name,
        "generated_at": datetime.now().isoformat(),
        "research_completed": True,
        "dimensions": search_results
    }
    
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(research_data, f, ensure_ascii=False, indent=2)
    
    print(f"✅ 调研结果 JSON 已生成：{output_path}")
    print(f"   客户：{customer_name}")
    print(f"   维度数：{len(search_results)}")
    return str(output_path)


def generate_icebreaker_ppt(customer_name, results, output_path):
    """生成场景破冰 PPT 内容（Markdown 格式，供 PPT 生成工具使用）
    
    ⚠️ ASI09 合规：此函数基于实际调研结果生成 PPT 脚本。
    调用方必须传入包含实际搜索内容的 results 字典。
    """
    
    # 安全检查：results 不能为空或全为待搜索占位符
    has_actual_content = False
    for dim_name, items in results.items():
        if isinstance(items, list):
            for item in items:
                if isinstance(item, str) and "未检索到" not in item and "待搜索" not in item:
                    has_actual_content = True
                    break
        if has_actual_content:
            break
    
    if not has_actual_content:
        raise ValueError(
            "❌ ASI09 合规错误：调研结果中没有实际搜索内容。\n"
            "调用方必须先执行 4 轮搜索（12 个关键词），\n"
            "禁止跳过调研直接生成通用模板 PPT。"
        )
    
    # 从调研结果中提取行业关键词，用于选择场景
    all_text = ""
    for dim_name, items in results.items():
        if isinstance(items, list):
            for item in items:
                if isinstance(item, str):
                    all_text += item + " "
    
    # 简单行业检测
    industry = None
    industry_kw = {
        "医疗": ["医院", "医疗", "卫生", "患者", "临床"],
        "制造": ["制造", "工厂", "生产", "车间", "设备"],
        "零售": ["零售", "门店", "连锁", "会员"],
        "教育": ["教育", "学校", "培训", "学员"]
    }
    for ind, kws in industry_kw.items():
        if sum(1 for kw in kws if kw in all_text) >= 2:
            industry = ind
            break
    
    # 根据行业选择场景模板
    from . import create_scenario_ppt as csp
    if industry and industry in csp.INDUSTRY_SCENARIOS:
        scenarios = csp.INDUSTRY_SCENARIOS[industry]
    else:
        scenarios = csp.GENERIC_SCENARIOS
    
    # 生成 Markdown 格式 PPT 脚本
    md_content = f"# {customer_name} 业务场景洞察\n\n"
    md_content += f"*基于调研发现的 before/after 对比*\n\n"
    md_content += f"**调研时间：** {datetime.now().strftime('%Y 年 %m 月 %d 日')}\n\n"
    md_content += f"**客户行业：** {industry or '未识别，使用通用场景'}\n\n"
    md_content += "---\n\n"
    
    # 使用实际调研结果中的信息补充场景描述
    findings_summary = []
    for dim_name, items in results.items():
        if isinstance(items, list):
            for item in items:
                if isinstance(item, str) and "未检索到" not in item and len(item) > 20:
                    findings_summary.append(item[:200])
    
    if findings_summary:
        md_content += "## 调研关键发现\n\n"
        for finding in findings_summary[:6]:
            md_content += f"- {finding}\n"
        md_content += "\n---\n\n"
    
    for i, scenario in enumerate(scenarios[:4], 1):
        md_content += f"## 第{i+1}页：{scenario['title']}\n\n"
        
        md_content += "### ❌ 现在（Before）\n\n"
        for point in scenario['before']:
            md_content += f"- {point}\n"
        
        md_content += "\n### ✅ 未来（After）\n\n"
        for point in scenario['after']:
            md_content += f"- {point}\n"
        
        md_content += f"\n📊 **价值：** {scenario.get('value', '')}\n\n"
        md_content += "\n---\n\n"
    
    md_content += "## 下一步验证计划\n\n"
    md_content += "- 第 1 周：选 1-2 个痛点场景深度调研\n"
    md_content += "- 第 2-3 周：用真实数据搭建 POC 环境\n"
    md_content += "- 第 4 周：业务部门试用，量化价值\n"
    md_content += "- 第 5 周：基于 POC 结果决策是否规模化推广\n"
    
    # 使用安全的文件写入
    output_path = Path(output_path)
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(md_content)
    
    print(f"✅ PPT 脚本已生成：{output_path}")
    return str(md_content), {"industry": industry, "scenario_count": min(len(scenarios), 4)}

# ---------------------------------------------------------------------------
# 主函数
# ---------------------------------------------------------------------------

def main():
    """主函数 - 安全版本
    
    ⚠️ 重要：本函数需要 agent framework 先执行搜索，然后传入实际结果。
    如果只传入客户名称而没有 search_results，输出将包含占位符，不符合完成标准。
    """
    if len(sys.argv) < 2:
        print("用法：python research.py <客户名称> [输出目录]")
        print("")
        print("参数说明：")
        print("  <客户名称>  客户全称（必填）")
        print("  [输出目录]  可选，默认使用 agent framework 的工作目录")
        print("")
        print("⚠️  重要提示：")
        print("  本脚本需要 agent framework 先执行 web_search，然后传入实际结果。")
        print("  否则输出的报告将包含'待搜索'占位符，不符合技能完成标准。")
        print("")
        print("输出文件：")
        print("  - [客户名称]_调研报告.md（或.docx）")
        print("  - [客户名称]_破冰 PPT.md")
        sys.exit(1)
    
    # 输入验证：防止路径遍历攻击
    customer_name = sys.argv[1].strip()
    if not customer_name:
        print("❌ 错误：客户名称不能为空")
        sys.exit(1)
    
    # 禁止路径分隔符
    if '/' in customer_name or '\\' in customer_name:
        print("❌ 错误：客户名称不能包含路径分隔符")
        sys.exit(1)
    
    # 限制长度
    if len(customer_name) > 100:
        print("❌ 错误：客户名称过长（最大 100 字符）")
        sys.exit(1)
    
    print(f"🚀 开始调研客户：{customer_name}")
    
    # 获取输出目录
    if len(sys.argv) >= 3:
        output_dir = Path(sys.argv[2])
        # 验证输出目录安全性
        try:
            output_dir = output_dir.resolve()
            # 确保在允许的工作目录内
            workspace = get_workspace_dir().resolve()
            if not str(output_dir).startswith(str(workspace)):
                print(f"⚠️  警告：输出目录不在工作目录内，将使用默认目录")
                output_dir = get_output_dir()
        except Exception:
            output_dir = get_output_dir()
    else:
        output_dir = get_output_dir()
    
    # ⚠️ 注意：实际使用场景中，agent framework 应该：
    # 1. 先调用 research_customer(customer_name) 获取搜索指令
    # 2. 执行所有 web_search 工具
    # 3. 调用 research_customer(customer_name, search_results) 传入实际结果
    # 4. 调用 generate_word_report() 和 generate_icebreaker_ppt() 生成报告
    
    # 本次调用生成搜索计划（供 agent framework 执行）
    print("\n📋 生成搜索计划...")
    search_plan = research_customer(customer_name)
    
    # 输出搜索计划
    print("✅ 搜索计划已生成（共 4 个维度，12 个关键词）")
    print("\n⚠️  下一步：agent framework 必须执行以下搜索：")
    for dim_name, keywords in search_plan.items():
        print(f"\n  【{dim_name}】")
        for item in keywords:
            print(f"    - {item['keyword']}")
    
    print("\n" + "="*60)
    print("📋 输出文件清单（搜索计划版）:")
    print(f"  1. {output_dir}/search_plan_{customer_name}.json")
    print("\n实际调研报告需在执行搜索后生成。")
    print("="*60)
    
    # 保存搜索计划为 JSON（供 agent framework 使用）
    plan_file = output_dir / f"search_plan_{customer_name}.json"
    with open(plan_file, 'w', encoding='utf-8') as f:
        json.dump(search_plan, f, ensure_ascii=False, indent=2)
    print(f"\n✅ 搜索计划已保存：{plan_file}")

if __name__ == "__main__":
    main()
