#!/usr/bin/env python3
"""
XLSX/DOCX Measured Data Extraction and Verification Script v3
- Supports parameterized column mapping
- Supports auto-detection of headers
- Supports more tolerance formats and judgement symbols
"""

from __future__ import annotations

import argparse
import os
import sys

import openpyxl

from inspection_utils import (
    classify_judge_text,
    clean_text,
    contains_numeric_spec,
    evaluate_value_against_tolerance,
    extract_rows_from_sheet,
    maybe_parse_float,
    parse_tolerance,
    summarize_precision,
)


def fmt_path(abspath, base_dir):
    try:
        rel = os.path.relpath(abspath, base_dir)
        return rel.replace("\\", "/")
    except ValueError:
        return abspath


def build_cli_mapping(args):
    return {
        "seq_col": args.seq_col,
        "item_col": args.item_col,
        "std_col": args.std_col,
        "method_col": args.method_col,
        "data_start": args.data_start,
        "data_cols": args.data_cols,
        "judge_col": args.judge_col,
        "header_row": args.header_row,
        "data_start_row": args.data_start_row,
    }


def analyze_row(row):
    tolerance = parse_tolerance(row.get("std"))
    judge_kind, normalized_judge = classify_judge_text(row.get("judge"))
    numeric_values = []
    text_values = []
    for value in row.get("values", []):
        number = maybe_parse_float(value)
        if number is not None:
            numeric_values.append(number)
        elif clean_text(value):
            text_values.append(clean_text(value))

    result = {
        "seq": row.get("seq"),
        "item": row.get("item"),
        "sheet": row.get("sheet"),
        "row_index": row.get("row_index"),
        "std": row.get("std"),
        "method": row.get("method"),
        "judge": row.get("judge"),
        "judge_kind": judge_kind,
        "normalized_judge": normalized_judge,
        "numeric_values": numeric_values,
        "text_values": text_values,
        "tolerance": tolerance,
        "match": True,
        "expected": "",
        "issues": [],
    }

    if tolerance is None:
        if judge_kind in ("nonstandard_pass", "other"):
            result["issues"].append("Judgement column uses non-standard terminology")
            result["match"] = False
        return result

    if not numeric_values:
        if contains_numeric_spec(row.get("std", "")) and text_values:
            result["issues"].append("Standard is numeric tolerance, but data column provides no numeric value")
            result["match"] = False
        return result

    failed_values = []
    margins = []
    for value in numeric_values:
        in_range, margin = evaluate_value_against_tolerance(value, tolerance)
        if not in_range:
            failed_values.append(value)
        if margin is not None:
            margins.append(margin)

    all_in_tol = len(failed_values) == 0
    result["expected"] = "Pass" if all_in_tol else "Fail"
    result["failed_values"] = failed_values
    result["min_margin"] = min(margins) if margins else None

    if all_in_tol:
        if judge_kind == "pass":
            return result
        if judge_kind == "fail":
            result["issues"].append("Measured value within tolerance, but judgement is Fail")
        elif judge_kind == "missing":
            result["issues"].append("Measured value within tolerance, but judgement column is empty")
        else:
            result["issues"].append("Measured value within tolerance, but judgement column uses non-standard pass term")
        result["match"] = False
        return result

    if judge_kind == "fail":
        return result

    if judge_kind == "pass":
        result["issues"].append("High risk: measured value out of tolerance, but judgement is Pass")
    elif judge_kind == "missing":
        result["issues"].append("Measured value out of tolerance, but judgement column is empty")
    else:
        result["issues"].append("Measured value out of tolerance, but judgement column uses non-standard fail term")
    result["match"] = False
    return result


def process_xlsx(path, w, cli_mapping):
    name = os.path.basename(path)
    workbook = openpyxl.load_workbook(path, data_only=True)
    file_rows = []
    sheet_layouts = []

    try:
        for sheet_name in workbook.sheetnames:
            ws = workbook[sheet_name]
            extracted = extract_rows_from_sheet(ws, sheet_name, cli_mapping)
            rows = extracted["rows"]
            if not rows:
                continue
            sheet_layouts.append((sheet_name, extracted["layout"], len(rows)))
            file_rows.extend(rows)
    finally:
        workbook.close()

    w("### {0}".format(name))
    if not file_rows:
        w("No valid data rows detected. Try specifying column mapping manually.")
        w()
        return {"type": "unknown", "n": 0, "issues": 0}

    w("Extracted **{0}** valid data records".format(len(file_rows)))
    w()
    w("**Detected column mapping:**")
    for sheet_name, layout, row_count in sheet_layouts:
        w(
            "- `{0}`: header_row={1}, data_start_row={2}, seq={3}, item={4}, std={5}, method={6}, data={7}-{8}, judge={9}, rows={10}".format(
                sheet_name,
                layout["header_row"],
                layout["data_start_row"],
                layout["seq_col"],
                layout["item_col"],
                layout["std_col"],
                layout["method_col"],
                layout["data_start"],
                layout["data_start"] + layout["data_cols"] - 1,
                layout["judge_col"],
                row_count,
            )
        )
    w()

    analyzed_rows = [analyze_row(row) for row in file_rows]
    checked_rows = [row for row in analyzed_rows if row.get("tolerance") or row.get("numeric_values")]
    issues = [row for row in analyzed_rows if row.get("issues")]
    high_risk = [row for row in issues if any("High risk" in issue for issue in row["issues"])]

    w("- Unique inspection items: {0}".format(len(set(row["item"] for row in file_rows if row["item"]))))
    w("- Verified data rows: {0}".format(len(checked_rows)))
    w("- Judgement anomalies: {0}".format(len(issues)))
    if high_risk:
        w("- High-risk items: {0}".format(len(high_risk)))
    w()

    if issues:
        w("**Anomaly details:**")
        for row in issues[:50]:
            joined = "; ".join(row["issues"])
            w(
                "- `{0}` sheet=`{1}` row={2} item=`{3}` std=`{4}` judge=`{5}` -> {6}".format(
                    row["seq"],
                    row["sheet"],
                    row["row_index"],
                    row["item"][:30],
                    row["std"][:30],
                    row["judge"][:12],
                    joined,
                )
            )
        w()
    else:
        w("All identified data rows have consistent judgement results with measured values")
        w()

    precision_summary = summarize_precision(file_rows)
    if precision_summary:
        w("**Measurement precision distribution:**")
        for method, summary in sorted(precision_summary.items()):
            distribution = ", ".join(
                "{0} decimals:{1}".format(decimals, count)
                for decimals, count in sorted(summary["distribution"].items())
            )
            w(
                "- `{0}`: dominant precision={1} decimals, distribution={2}".format(
                    method[:20], summary["dominant_decimals"], distribution
                )
            )
        w()

    max_values = max(len(row["values"]) for row in file_rows) if file_rows else 0
    header_cells = ["Seq", "Sheet", "Item", "Standard", "Method"]
    header_cells.extend(["Value{0}".format(index + 1) for index in range(max_values)])
    header_cells.append("Judgement")
    w("**Detailed data:**")
    w()
    w("| " + " | ".join(header_cells) + " |")
    w("|" + "|".join(["------"] * len(header_cells)) + "|")
    for row in file_rows[:200]:
        cells = [
            str(row["seq"]),
            row["sheet"],
            row["item"][:20],
            row["std"][:18],
            row["method"][:10],
        ]
        for value in row["values"]:
            cells.append(clean_text(value)[:12])
        if len(row["values"]) < max_values:
            cells.extend([""] * (max_values - len(row["values"])))
        cells.append(row["judge"][:10])
        w("| " + " | ".join(cells) + " |")
    w()
    return {"type": "xlsx", "n": len(checked_rows), "issues": len(issues)}


def process_docx(path, w):
    name = os.path.basename(path)
    results = {"pass": 0, "fail": 0, "issues": 0}
    try:
        from docx import Document

        doc = Document(path)
        if not doc.tables:
            w("### {0}".format(name))
            w("Document contains no tables")
            w()
            return results

        total_pass = 0
        total_fail = 0
        total_issues = 0
        w("### {0}".format(name))
        w("Document contains {0} tables".format(len(doc.tables)))
        w()
        for table_index, table in enumerate(doc.tables):
            if len(table.rows) < 2:
                continue
            headers = [clean_text(cell.text)[:12] for cell in table.rows[0].cells]
            judge_idx = -1
            for idx, header in enumerate(headers):
                if "Judge" in header or "OK" in header.upper():
                    judge_idx = idx
                    break
            if judge_idx < 0:
                continue

            table_issues = 0
            for row in table.rows[1:]:
                cells = [clean_text(cell.text) for cell in row.cells]
                if not cells or not cells[0]:
                    continue
                judge_text = cells[judge_idx] if judge_idx < len(cells) else ""
                judge_kind, _ = classify_judge_text(judge_text)
                if judge_kind == "pass":
                    total_pass += 1
                else:
                    total_fail += 1
                if judge_kind in ("missing", "other", "nonstandard_pass"):
                    table_issues += 1
                    total_issues += 1
            w("- Table #{0}: {1} rows x {2} cols, judgement anomalies/non-standard {3}".format(
                table_index + 1, len(table.rows), len(table.columns), table_issues
            ))
        w()
        w("**Summary:** Pass {0} / Total {1}".format(total_pass, total_pass + total_fail))
        w("All passed" if total_fail == 0 and total_issues == 0 else "Review judgement anomalies or non-standard terms")
        w()
        results["pass"] = total_pass
        results["fail"] = total_fail
        results["issues"] = total_issues
    except Exception as exc:
        w("### {0}".format(name))
        w("DOCX parsing failed: {0}".format(exc))
        w()
    return results


def collect_files(args):
    if args.file:
        file_path = os.path.abspath(args.file)
        if not os.path.isfile(file_path):
            print("[Error] File not found: {0}".format(file_path), file=sys.stderr)
            sys.exit(1)
        base_dir = (
            os.path.dirname(os.path.dirname(file_path))
            if os.path.basename(os.path.dirname(file_path)) == "output"
            else os.path.dirname(file_path)
        )
        return base_dir, [file_path]

    if not args.dir:
        print("[Error] Please specify --dir or --file", file=sys.stderr)
        sys.exit(1)

    base_dir = os.path.abspath(args.dir)
    output_dir = os.path.join(base_dir, "output")
    if not os.path.isdir(output_dir):
        print("[Error] output/ not found: {0}".format(output_dir), file=sys.stderr)
        sys.exit(1)
    files_to_process = sorted(
        os.path.join(output_dir, name)
        for name in os.listdir(output_dir)
        if name.lower().endswith((".xlsx", ".xls", ".docx", ".doc")) and not name.startswith("~")
    )
    return base_dir, files_to_process


def main():
    parser = argparse.ArgumentParser(description="Dimension Inspection Report -> Data Verification")
    parser.add_argument("--dir", type=str, default="", help="Process entire task folder (find XLSX/DOCX from DIR/output/)")
    parser.add_argument("--file", type=str, default="", help="Process single file (XLSX or DOCX)")
    parser.add_argument("--output", type=str, default="", help="Report output path (default: same directory as input)")
    parser.add_argument("--seq-col", type=int, default=None, help="Sequence column")
    parser.add_argument("--item-col", type=int, default=None, help="Inspection item column")
    parser.add_argument("--std-col", type=int, default=None, help="Standard column")
    parser.add_argument("--method-col", type=int, default=None, help="Method column")
    parser.add_argument("--data-start", type=int, default=None, help="Data start column")
    parser.add_argument("--data-cols", type=int, default=None, help="Number of data columns")
    parser.add_argument("--judge-col", type=int, default=None, help="Judgement column")
    parser.add_argument("--header-row", type=int, default=None, help="Header row number")
    parser.add_argument("--data-start-row", type=int, default=None, help="Data start row number")
    args = parser.parse_args()

    base_dir, files_to_process = collect_files(args)
    report_path = os.path.abspath(args.output) if args.output else os.path.join(base_dir, "DataVerificationReport.md")
    cli_mapping = build_cli_mapping(args)
    report_lines = []

    def w(text=""):
        report_lines.append(text)

    w("# Dimension Inspection Report — Data Verification Report")
    w()
    w("**Working directory:** `{0}`".format(fmt_path(base_dir, base_dir)))
    w("**Notes:** Supports auto-header detection; column mapping may be overridden via CLI.")
    w("**Review rules:** OK/Pass/(check mark)/(circle) = Pass, NOK/Fail/NG = Fail, PASS/Pass(in Chinese)/Correct = non-standard.")
    w()

    summary = []
    total_n = 0
    for file_path in files_to_process:
        ext = os.path.splitext(file_path)[1].lower()
        w("---")
        if ext in (".xlsx", ".xls"):
            result = process_xlsx(file_path, w, cli_mapping)
            n = result.get("n", 0)
            issues = result.get("issues", 0)
        elif ext in (".docx", ".doc"):
            result = process_docx(file_path, w)
            n = result.get("pass", 0) + result.get("fail", 0)
            issues = result.get("issues", 0)
        else:
            w("### {0}".format(os.path.basename(file_path)))
            w("Unsupported file format")
            w()
            n = 0
            issues = 0
        total_n += n
        summary.append((os.path.basename(file_path), n, issues))

    w("---")
    w("## Summary")
    w()
    w("| File | Valid Data Rows | Judgement Consistent |")
    w("|------|:---------:|:--------:|")
    total_issues = 0
    for name, n, issues in summary:
        if issues == 0 and n > 0:
            w("| {0} | {1} | All consistent |".format(name[:30], n))
        elif issues > 0:
            w("| {0} | {1} | {2} items require manual review |".format(name[:30], n, issues))
        else:
            w("| {0} | {1} | No valid data |".format(name[:30], n))
        total_issues += issues
    w()
    w("**Total valid records:** {0}".format(total_n))
    w("**Issues found:** No judgement anomalies" if total_issues == 0 else "**Anomalies:** {0} items require manual review".format(total_issues))
    w()
    w("---")
    w("*Auto-generated by script -- `scripts/extract_verify_data.py`*")

    with open(report_path, "w", encoding="utf-8") as handle:
        handle.write("\n".join(report_lines))

    print("[Done] Data verification report generated: {0}".format(report_path))
    print()
    print("=" * 70)
    print("[File Locations] (Copy and paste into File Explorer)")
    print("=" * 70)
    print("  [Report] Verification report: {0}".format(report_path))
    for file_path in files_to_process:
        print("  [Source] Source file:        {0}".format(file_path))
    print("  [Dir]    Working directory:  {0}".format(base_dir))
    print()
    print("[Tip] Select the path above -> Ctrl+C -> paste into File Explorer address bar")
    print()


if __name__ == "__main__":
    main()
