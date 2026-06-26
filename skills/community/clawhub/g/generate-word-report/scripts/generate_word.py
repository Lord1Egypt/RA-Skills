# -*- coding: utf-8 -*-
"""
生成 Word 日报（无模板依赖版）
直接用 python-docx 构建文档结构，不依赖 M1.docx 模板。

文档结构：
  1. 隔夜环球市场总结（纯文字，无数字）
  2. 市场表现方面
     2.1 美国（美股+市场结构+个股合并 / 美国经济情况 / 美债 / 美元指数 / 美国就业市场情况）
     2.2 中国及中国香港（A股 / 港股）
     2.3 黄金
  3. 政策方面（美国 / 欧洲 / 中国）
  4. 科技方面
  5. 其他市场（欧洲股市 / 欧元区经济 / 亚太市场）
  6. 今日关注经济数据
"""
import sys
sys.path.insert(0, r"C:\Users\qu669\.openclaw\workspace-yoyo")
sys.stdout.reconfigure(encoding='utf-8')
import os, json, datetime, logging, re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import config

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[logging.FileHandler(config.LOG_FILE, encoding="utf-8"), logging.StreamHandler(sys.stdout)])
log = logging.getLogger(__name__)

# ─────────────────────────────────────────────
# 样式配置
# ─────────────────────────────────────────────

COLOR_TITLE    = RGBColor(0x1F, 0x49, 0x7D)
COLOR_H1       = RGBColor(0x1F, 0x49, 0x7D)
COLOR_H2       = RGBColor(0x2E, 0x74, 0xB5)
COLOR_H3       = RGBColor(0x00, 0x00, 0x00)
COLOR_BODY     = RGBColor(0x00, 0x00, 0x00)
COLOR_GRAY     = RGBColor(0x60, 0x60, 0x60)
COLOR_ACCENT   = RGBColor(0xC0, 0x00, 0x00)

FONT_NAME_CN   = "微软雅黑"
FONT_NAME_EN   = "Arial"

# ─────────────────────────────────────────────
# 辅助函数
# ─────────────────────────────────────────────

def gp(key, mdata, default=None):
    """获取价格 - 支持别名key自动回退"""
    aliases = {
        'COMEX黄金期货':      ['COMEX黄金', '黄金期货'],
        '现货黄金(XAUUSD)':  ['现货黄金', 'XAUUSD'],
        '美元指数(DXY)':      ['美元指数', 'DXY'],
        '10年期美债收益率':  ['10年期美债', '美债收益率'],
        '2年期美债收益率':   ['2年期美债'],
    }
    for k in [key] + aliases.get(key, []):
        # 直接从顶层查
        d = mdata.get(k, {})
        if isinstance(d, dict) and d:
            v = d.get('price')
            if v is not None:
                return v
        # 从市场表现子字典查
        mp = mdata.get('市场表现', {})
        for region in mp.values():
            if isinstance(region, dict):
                d = region.get(k, {})
                if isinstance(d, dict) and d:
                    v = d.get('price')
                    if v is not None:
                        return v
    return default

def gc(key, mdata, default=None):
    """获取涨跌幅% - 支持别名key自动回退"""
    aliases = {
        'COMEX黄金期货':      ['COMEX黄金', '黄金期货'],
        '现货黄金(XAUUSD)':  ['现货黄金', 'XAUUSD'],
        '美元指数(DXY)':      ['美元指数', 'DXY'],
    }
    for k in [key] + aliases.get(key, []):
        # 直接从顶层查
        d = mdata.get(k, {})
        if isinstance(d, dict) and d:
            v = d.get('change')
            if v is not None:
                return v
        # 从市场表现子字典查
        mp = mdata.get('市场表现', {})
        for region in mp.values():
            if isinstance(region, dict):
                d = region.get(k, {})
                if isinstance(d, dict) and d:
                    v = d.get('change')
                    if v is not None:
                        return v
    return default

def clean_text(text):
    """清理连续标点：连续两个句号、连续逗号、逗号+句号、句号+逗号等"""
    if not text:
        return text
    for p in ['。', '，', '、', '；', '：']:
        text = re.sub(p + '{2,}', p, text)
    text = re.sub(r'，。', '。', text)
    text = re.sub(r'。，', '。', text)
    text = text.strip().rstrip('，、;:')
    return text

def fp(v):
    """格式化价格数字 - 统一保留两位小数"""
    if v is None: return "N/A"
    if abs(v) >= 10000: return f"{v:,.2f}"
    return f"{v:,.2f}"

def fpc(v):
    """涨跌% 带符号"""
    if v is None or v == 0: return "N/A"
    return f"{'+' if v > 0 else ''}{v:.2f}%"

def fpc_abs(v):
    """涨跌% 绝对值（0显示0.00%，不显示N/A）"""
    if v is None: return "N/A"
    return f"{abs(v):.2f}%"

def chn(v):
    """涨跌 → 汉字（0=持稳）"""
    if v is None: return "涨跌N/A"
    if v == 0: return "持稳"
    return "涨" if v > 0 else "跌"

def chn_full(v):
    """涨跌 → 完整汉字（0=持稳）"""
    if v is None: return "涨跌N/A"
    if v == 0: return "持稳"
    return "上涨" if v > 0 else "下跌"

def chn_bp(v):
    """涨跌 → 基点数（带单位）"""
    if v is None or v == 0: return "N/A"
    return f"{abs(v):.2f}个基点"

# ─────────────────────────────────────────────
# 文档构建工具
# ─────────────────────────────────────────────

class DocBuilder:
    def __init__(self, doc: Document):
        self.doc = doc
        self._set_document_styles()

    def _set_document_styles(self):
        style = self.doc.styles['Normal']
        style.font.name = FONT_NAME_CN
        style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME_CN)
        style.font.size = Pt(10.5)

    def add_heading(self, text, level=1):
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        if level == 0:
            run.font.name = FONT_NAME_CN
            run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME_CN)
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = COLOR_TITLE
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif level == 1:
            run.font.name = FONT_NAME_CN
            run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME_CN)
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = COLOR_H1
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(4)
        elif level == 2:
            run.font.name = FONT_NAME_CN
            run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME_CN)
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = COLOR_H2
            para.paragraph_format.space_before = Pt(8)
            para.paragraph_format.space_after = Pt(2)
        elif level == 3:
            run.font.name = FONT_NAME_CN
            run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME_CN)
            run.font.size = Pt(10.5)
            run.font.bold = True
            run.font.color.rgb = COLOR_H3
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after = Pt(0)
        return para

    def add_body(self, text, indent=False):
        text = clean_text(text)
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = FONT_NAME_CN
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME_CN)
        run.font.size = Pt(10.5)
        run.font.color.rgb = COLOR_BODY
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(4)
        if indent:
            para.paragraph_format.left_indent = Inches(0.3)
        return para

    def add_subtitle(self, text):
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = FONT_NAME_CN
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME_CN)
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_GRAY
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_after = Pt(8)
        return para

    def add_divider(self):
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)
        return para

    def add_bullet(self, text, indent_level=0):
        para = self.doc.add_paragraph(style='List Bullet')
        run = para.add_run(text)
        run.font.name = FONT_NAME_CN
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME_CN)
        run.font.size = Pt(10.5)
        run.font.color.rgb = COLOR_BODY
        para.paragraph_format.space_before = Pt(1)
        para.paragraph_format.space_after = Pt(1)
        base_indent = 0.3
        para.paragraph_format.left_indent = Inches(base_indent + indent_level * 0.25)
        return para

# ─────────────────────────────────────────────
# 内容生成函数
# ─────────────────────────────────────────────

def build_global_summary(mdata):
    """模块1：隔夜环球市场总结 - 直接使用market_data.json中的完整环球市场速览"""
    summary_data = mdata.get('环球市场速览', {})
    paragraphs = summary_data.get('段落列表', [])
    
    if paragraphs:
        # 使用完整的环球市场速览，每句话作为独立段落
        return "\n".join(paragraphs)
    
    # 回退逻辑：如果没有速览数据，使用原有合成方式
    dji_chg    = gc('道琼斯工业平均指数', mdata)
    spx_chg    = gc('标普500指数', mdata)
    nasdaq_chg = gc('纳斯达克综合指数', mdata)
    sh_chg     = gc('上证指数', mdata)
    hsi_chg    = gc('恒生指数', mdata)
    gold_chg   = gc('COMEX黄金期货', mdata) or gc('现货黄金(XAUUSD)', mdata)
    yield10_chg = gc('10年期美债收益率', mdata)
    usdidx_chg  = gc('美元指数(DXY)', mdata)
    eu_dax_chg  = gc('德国DAX 30', mdata)
    asi_nik_chg = gc('日经225指数', mdata)

    news      = mdata.get('新闻', {})
    us_news   = news.get('美国', [])
    china_news = news.get('中国', [])

    if (dji_chg or 0) > 0:
        main = "美国股市在逢低买盘支撑下小幅走强，"
    else:
        main = "美国股市受地缘政治不确定性拖累震荡走低，"

    news_snippets = []
    for n in us_news[:2]:
        t = n.get('title', '') or n.get('内容', '') or ''
        if t and len(t) > 10:
            news_snippets.append(t.replace('。', '').replace('，', ''))
    for n in china_news[:1]:
        t = n.get('title', '') or n.get('内容', '') or ''
        if t and len(t) > 10:
            news_snippets.append(t.replace('。', '').replace('，', ''))

    snippet_text = ""
    if news_snippets:
        snippet_text = "，" + "，".join(news_snippets[:2]) + "等因素亦加剧市场波动。"

    text_summary = (f"{main}投资者对全球贸易前景及中东局势的担忧挥之不去，"
                    f"避险情绪与风险偏好交替主导市场。{snippet_text}"
                    f"整体来看，全球市场在宏观不确定性笼罩下延续震荡格局。")

    assets = []
    if dji_chg is not None:
        assets.append(f"道指{fpc(dji_chg)}")
    if spx_chg is not None:
        assets.append(f"标普500{fpc(spx_chg)}")
    if nasdaq_chg is not None:
        assets.append(f"纳指{fpc(nasdaq_chg)}")
    if sh_chg is not None:
        assets.append(f"上证指数{fpc(sh_chg)}")
    if hsi_chg is not None:
        assets.append(f"恒生指数{fpc(hsi_chg)}")
    if eu_dax_chg is not None:
        assets.append(f"德国DAX{fpc(eu_dax_chg)}")
    if asi_nik_chg is not None:
        assets.append(f"日经225{fpc(asi_nik_chg)}")
    if gold_chg is not None:
        assets.append(f"黄金{fpc(gold_chg)}")
    if yield10_chg is not None:
        assets.append(f"10Y美债收益率{fpc(yield10_chg)}")
    if usdidx_chg is not None:
        assets.append(f"美元指数{fpc(usdidx_chg)}")

    asset_summary = ""
    if assets:
        asset_summary = "\n\n【主要资产表现】" + "、".join(assets) + "。"

    return text_summary + asset_summary

def build_market_performance_us(mdata):
    """模块2·美国：返回5段文字
    [0] 美股（含市场结构+重要个股板块）
    [1] 美国经济情况
    [2] 美债
    [3] 美元指数
    [4] 美国就业市场情况
    """
    dji     = (gp('道琼斯工业平均指数', mdata), gc('道琼斯工业平均指数', mdata))
    spx     = (gp('标普500指数', mdata), gc('标普500指数', mdata))
    nasdaq  = (gp('纳斯达克综合指数', mdata), gc('纳斯达克综合指数', mdata))
    yield10 = (gp('10年期美债收益率', mdata), gc('10年期美债收益率', mdata))
    yield2  = (gp('2年期美债收益率', mdata), gc('2年期美债收益率', mdata))
    usdidx  = (gp('美元指数(DXY)', mdata), gc('美元指数(DXY)', mdata))
    usd_cny = gp('USD/CNY', mdata)
    eur_usd = gp('EUR/USD', mdata)
    econ    = mdata.get('经济数据', {})
    news_map = mdata.get('新闻', {})
    us_news  = news_map.get('美国', [])
    tech_news = news_map.get('科技', [])

    # ── 美股（含市场结构+个股板块合并） ──
    dji_dir    = 1 if dji[1]   and dji[1]   > 0 else (-1 if dji[1]   and dji[1]   < 0 else 0)
    spx_dir    = 1 if spx[1]   and spx[1]   > 0 else (-1 if spx[1]   and spx[1]   < 0 else 0)
    nasdaq_dir = 1 if nasdaq[1] and nasdaq[1] > 0 else (-1 if nasdaq[1] and nasdaq[1] < 0 else 0)
    dirs = [dji_dir, spx_dir, nasdaq_dir]
    if all(d == 1  for d in dirs): us_chg_str = "集体收涨"
    elif all(d == -1 for d in dirs): us_chg_str = "集体收跌"
    elif nasdaq_dir == 0 and dji_dir == spx_dir:
        us_chg_str = "集体收涨" if dji_dir == 1 else "集体收跌"
    else: us_chg_str = "涨跌不一"

    dji_str    = f"道指{chn(dji[1])}{fpc_abs(dji[1])}报{fp(dji[0])}点"
    spx_str    = f"标普500指数{chn(spx[1])}{fpc_abs(spx[1])}报{fp(spx[0])}点"
    nasdaq_str = f"纳指{chn(nasdaq[1])}{fpc_abs(nasdaq[1])}报{fp(nasdaq[0])}点"

    if nasdaq[1] and nasdaq[1] > 0:
        tech_driver = "科技股普遍反弹，人工智能相关板块重新获得资金关注，整体成长风格有所回暖"
        energy_driver = "能源股随油价波动有所震荡，但地缘风险溢价仍在"
    elif nasdaq[1] and nasdaq[1] < 0:
        tech_driver = "科技成长股普遍承压，估值压力叠加利率预期令高弹性个股面临阶段性调整"
        energy_driver = "能源股受地缘风险推动逆势走强，资金向防御性板块迁移"
    else:
        tech_driver = "光通信、光模块等概念仍是市场主线之一，部分个股出现高位强势震荡"
        energy_driver = "地缘风险持续发酵带动油气股维持相对强势"

    political_note = ("特朗普政府贸易政策的不确定性持续扰动市场情绪，"
                      "企业在资本开支决策上趋于谨慎，但一季度财报整体表现好于预期，为市场提供底部支撑。")

    news_parts = []
    for item in (us_news + tech_news)[:6]:
        title = item.get('title', '') or item.get('内容', '') or ''
        if not title or len(title) < 10: continue
        title = title.strip().rstrip('。.')
        news_parts.append(title)
        if len(news_parts) >= 4: break

    news_text = ""
    if news_parts:
        lines = ["当日美股市场重要动态包括："]
        for j, t in enumerate(news_parts, 1):
            lines.append(f"{j}）{t}。")
        news_text = "".join(lines)
    else:
        news_text = ("当日美股市场未出现重大个股或板块新闻，整体交投受宏观因素主导。"
                     "市场短期仍将以消息面驱动为主，建议关注美股财报季期间个股基本面变化。")

    para1 = (f"美国三大股指{us_chg_str}，{dji_str}，{spx_str}，{nasdaq_str}。"
             f"受美伊局势变化及油价波动影响，盘中震荡幅度较大；"
             f"科技股承压，能源股则受地缘风险推动有所表现。从市场结构来看，{tech_driver}；"
             f"{energy_driver}。{political_note}{news_text}")

    # ── 美国经济情况 ──
    us_unemp_val = None
    us_econ = econ.get('美国', {})
    if isinstance(us_econ, dict):
        unemp = us_econ.get('失业率', {})
        if isinstance(unemp, dict):
            us_unemp_val = unemp.get('数值', 'N/A')
        if us_unemp_val is None:
            us_unemp_val = 'N/A'
    else:
        us_unemp_val = 'N/A'

    para2 = (f"美国近期失业率为{us_unemp_val}；"
             f"就业市场边际走弱压力有所上升，企业裁员计划增加，"
             f"但整体劳动力市场仍具韧性，消费支出在经历年初放缓后出现企稳迹象，"
             f"美联储在评估关税影响前预计维持利率观望。")

    # ── 美债 ──
    # 处理 0 变化：当变动为 0 或 None 时，不显示涨跌方向，只显示收益率数值
    def _bond_dir(v):
        if v is None or v == 0:
            return None
        return "上涨" if v > 0 else "下跌"

    y10_dir = _bond_dir(yield10[1])
    y10_bp  = f"{abs(yield10[1]):.2f}" if yield10[1] is not None else "N/A"
    y2_dir  = _bond_dir(yield2[1])
    y2_bp   = f"{abs(yield2[1]):.2f}" if yield2[1] is not None else "N/A"
    y10_val = f"{yield10[0]:.2f}" if yield10[0] else "N/A"
    y2_val  = f"{yield2[0]:.2f}" if yield2[0] else "N/A"
    same_dir = (yield10[1] is not None and yield2[1] is not None and
                bool((yield10[1] > 0) == (yield2[1] > 0)))
    bond_str = "同步走" if same_dir else "分化走势"
    y10_part = f"{y10_dir}{y10_bp}至{y10_val}%" if y10_dir else f"维持在{y10_val}%"
    y2_part  = f"{y2_dir}{y2_bp}至{y2_val}%" if y2_dir else f"维持在{y2_val}%"
    para3 = (f"美债收益率呈现{bond_str}，"
             f"10年期美债收益率{y10_part}，"
             f"2年期美债收益率{y2_part}。"
             f"美伊局势消息面反复令避险资金在股债之间频繁切换，短期美债波动加大。")

    # ── 美元指数 ──
    if usdidx[1] is not None:
        dxy_str = f"美元指数{chn(usdidx[1])}{fpc_abs(usdidx[1])}报{fp(usdidx[0])}"
    else:
        dxy_str = "美元指数变动待更新"
    fed_candidate = "凯文·沃什"
    para4 = (f"{dxy_str}，USD/CNY报{usd_cny or 'N/A'}，EUR/USD报{eur_usd or 'N/A'}。"
             f"避险需求回升叠加美联储主席候选人{fed_candidate}释放偏鹰派信号，令美元短期获得支撑。"
             f"美联储短期维持利率不变的预期持续强化，年内降息预期进一步降温。")

    # 美国就业市场数据段落（5段专业叙事）
    emp = mdata.get('美国就业市场', {})

    def _fw(v):
        if isinstance(v, (int, float)):
            return f"{v/10000:.1f}"
        return str(v) if v != '暂未披露' else '暂未披露'

    ic  = emp.get('当周初请失业金', {})
    ic4 = emp.get('初请四周均值', {})
    cc  = emp.get('续请失业金', {})
    ch  = emp.get('Challenger裁员', {})
    nfp = emp.get('非农前瞻预测', {})

    cur_ic   = _fw(ic.get('当期值', 0))
    prev_ic  = _fw(ic.get('前值', 0))
    dt_ic    = ic.get('期间', '')
    prev_dt_ic = ic.get('前值期间', '')

    cur_ic4  = _fw(ic4.get('当期值', 0))
    prev_ic4 = _fw(ic4.get('前值', 0))
    chg_ic4  = ic4.get('边际变化', 'N/A')
    if chg_ic4 == '暂未披露':
        chg_ic4_str = '暂未披露'
    elif isinstance(chg_ic4, (int, float)):
        chg_ic4_str = str(int(chg_ic4))
    else:
        chg_ic4_str = str(chg_ic4)

    cur_cc   = _fw(cc.get('当期值', 0))
    prev_cc  = _fw(cc.get('前值', 0))
    dt_cc    = cc.get('期间', '')
    prev_dt_cc = cc.get('前值期间', '')

    val_ch_raw = ch.get('当期值', None)
    val_ch = f"{val_ch_raw:.1f}" if isinstance(val_ch_raw, float) else str(val_ch_raw) if val_ch_raw else '暂未披露'
    prev_ch = str(ch.get('前值', '暂未披露'))
    reasons = list(ch.get('核心诱因', {}).keys())
    reason_text = '、'.join(reasons) if reasons else '待披露'

    payroll_nfp = nfp.get('新增就业预期', 'N/A')
    period_nfp  = nfp.get('期间', '')

    # Paragraph 1: 逐条罗列
    para_p1 = (f"当周初请失业金：当期值{cur_ic}万人（前值期间{dt_ic}），前值{prev_ic}万人"
               f"（前值期间{prev_dt_ic}）；初请四周均值：当期值{cur_ic4}万人，边际变化{chg_ic4_str}；"
               f"续请失业金：当期值{cur_cc}万人（前值期间{dt_cc}），前值{prev_cc}万人"
               f"（前值期间{prev_dt_cc}）；Challenger企业裁员：当期值{val_ch}万人（前值{prev_ch}），"
               f"核心诱因包括：{reason_text}；非农前瞻：市场预期新增就业"
               f"{payroll_nfp}万人（{period_nfp}）。")

    # Paragraph 2: 整体分化格局
    if isinstance(chg_ic4, (int, float)) and chg_ic4 < 0:
        ic4_trend = "回落"
    elif isinstance(chg_ic4, (int, float)) and chg_ic4 > 0:
        ic4_trend = "攀升"
    else:
        ic4_trend = "变化"
    para_p2 = (f"当周初请失业金录得{cur_ic}万人，环比下降{prev_ic}万人，初请四周均值同步{ic4_trend}"
               f"至{cur_ic4}万人，边际减少{chg_ic4_str}人。续请失业金降至{cur_cc}万人，前值期间续请人数"
               f"录得{prev_cc}万人。整体而言，当期就业数据呈现初请与续请双降格局，初请端边际改善"
               f"相对明显，然续请绝对量仍处相对高位，映射就业市场并未出现显著恶化，但下行压力有所累积。")

    # Paragraph 3: 裁员结构
    para_p3 = (f"Challenger企业当期裁员规模录得{val_ch}万人（前值{prev_ch}），已连续处于历史"
               f"相对较高区间。分项结构显示，{reason_text}构成当期裁员核心驱动因素，反映企业在"
               f"贸易政策不确定性背景下的保守雇佣策略。行业层面，制造业、零售业及部分科技领域"
               f"均面临不同程度的人员优化压力。")

    # Paragraph 4: 初请续请差异与宏观联动
    para_p4 = (f"从初请与续请数据背离来看，初请端边际改善速度快于续请端绝对量回落幅度，表明"
               f"单周裁员波动有所收敛，但存量待领取失业金人数尚未形成趋势性下行，就业市场韧性"
               f"仍受企业招聘端滞后调整支撑。宏观层面，就业数据分化格局对大类资产形成差异化指引："
               f"劳动力市场边际走弱信号压低美债实际利率预期，对黄金价格形成支撑；同时，薪资增长"
               f"压力缓解亦降低美联储被动收紧概率，风险资产估值端压力有所减轻。")

    # Paragraph 5: 核心特征与非农前瞻
    nfp_narrate = (f"当月非农前瞻数据录得市场预期新增就业{payroll_nfp}万人（{period_nfp}），"
                   f"关注服务业新增就业能否对制造业景气下行形成对冲，以及薪资增速是否维持韧性。"
                   ) if payroll_nfp != 'N/A' else (
                   f"当月非农前瞻数据暂未披露，密切关注后续发布情况。"
                   )
    para_p5 = (f"综合高频数据研判，当前美国劳动力市场核心运行特征如下：招聘需求边际放缓、"
               f"企业裁员计划有所上升、就业稳定性指标出现松动迹象。"
               f"{nfp_narrate}"
               f"当前就业市场整体仍处周期性调整通道，趋势性拐点的确认仍需后续数据验证。")

    # Combine all into single para5 (for return structure compatibility)
    para5 = para_p1 + " " + para_p2 + " " + para_p3 + " " + para_p4 + " " + para_p5

    return [para1, para3, para4, para5]

def build_market_performance_china_hk(mdata):
    """模块2·中国及中国香港"""
    sh     = (gp('上证指数', mdata), gc('上证指数', mdata))
    sz     = (gp('深证成指', mdata), gc('深证成指', mdata))
    cy     = (gp('创业板指', mdata), gc('创业板指', mdata))
    csi    = (gp('沪深300', mdata), gc('沪深300', mdata))
    hsi    = (gp('恒生指数', mdata), gc('恒生指数', mdata))
    hstech = (gp('恒生科技指数', mdata), gc('恒生科技指数', mdata))
    # 恒生中国企业指数（数据源暂无）
    hscei  = (None, None)

    sh_str  = f"上证指数{chn_full(sh[1])}{fpc_abs(sh[1]) if sh[1] is not None else 'N/A'}报{fp(sh[0])}点"
    sz_str  = f"深证成指{chn_full(sz[1])}{fpc_abs(sz[1]) if sz[1] is not None else 'N/A'}"
    cy_str  = f"创业板指{chn_full(cy[1])}{fpc_abs(cy[1]) if cy[1] is not None else 'N/A'}"
    csi_str = f"沪深300{chn_full(csi[1])}{fpc_abs(csi[1]) if csi[1] is not None else 'N/A'}"
    para_a  = (f"A股呈现结构性分化格局，{sh_str}，{sz_str}，{cy_str}，{csi_str}。"
               f"光纤光模块板块表现亮眼，政策宽松预期对部分周期股形成支撑；"
               f"科技股经历前期回调后出现分化，资金在基本面扎实的龙头与题材炒作之间权衡。")

    hsi_str     = f"恒生指数{chn_full(hsi[1])}{fpc_abs(hsi[1]) if hsi[1] is not None else 'N/A'}报{fp(hsi[0])}点"
    hstech_str  = f"恒生科技指数{chn_full(hstech[1])}{fpc_abs(hstech[1]) if hstech[1] is not None else 'N/A'}"
    para_hk     = (f"港股震荡承压，{hsi_str}，{hstech_str}。"
                   f"半导体、黄金股及AI应用类个股跌幅居前，美联储鹰派信号压制全球成长股估值；"
                   f"油气股及部分医药股逆势走强，南向资金维持净买入态势。")

    return [para_a, para_hk]

def build_market_performance_gold(mdata):
    """模块2·黄金"""
    comex_gold = gp('COMEX黄金期货', mdata)
    comex_chg  = gc('COMEX黄金期货', mdata)
    spot_gold  = gp('现货黄金(XAUUSD)', mdata)
    spot_chg   = gc('现货黄金(XAUUSD)', mdata)

    # 任一存在即可（现货和期货至少有一个有数据）
    gold_price = comex_gold if comex_gold is not None else spot_gold
    gold_chg   = comex_chg  if comex_chg  is not None else spot_chg
    if gold_price is None: gold_price = 0
    if gold_chg   is None: gold_chg   = 0

    return (f"现货黄金{chn_full(gold_chg)}{fpc_abs(gold_chg)}至每盎司{fp(gold_price)}美元，"
            f"COMEX黄金期货参考价格{fp(gold_price) if gold_price else 'N/A'}美元/盎司。"
            f"美伊局势未如预期进一步升级，避险买盘有所消退；美联储官员接连释放偏鹰表态，"
            f"实际利率预期回升令黄金持有成本上升，短期金价或维持高位震荡格局，"
            f"若地缘风险再度升温则可能重启升势。")

def _str(v, fallback="暂无相关数据。"):
    """从政策列表中提取内容，优先取内容字段（内容损坏时用标题），并清理媒体来源字样"""
    if isinstance(v, list):
        if not v:
            return fallback
        item = v[0]
        content = item.get('内容', '') or item.get('title', '') or ''
        # 清理内容中的网页干扰元素（新浪APP提示、HTML残余、表格碎片等）
        content = _clean_policy_content(content)
        # 内容正常（无严重乱码）则使用；乱码程度高则回退到标题
        if content and len(content) > 20 and not _is_garbled(content):
            return content
        # fallback to 标题 (strip media source suffix pattern like _xxx_xxx)
        title = item.get('标题', '') or item.get('title', '') or ''
        if title and len(title) > 10:
            # strip trailing media source patterns like "_财经头条__新浪财经"
            cleaned = re.sub(r'_{1,2}[^_]+_{1,2}[^_]+$', '', title).strip()
            return cleaned + '。'
        return fallback
    if isinstance(v, str) and v.strip():
        return v.strip()
    return fallback


def _clean_policy_content(text):
    """"清理政策内容中的网页干扰元素，逐行处理返回清洁文本"""
    if not text:
        return text
    lines = text.split('\n')
    cleaned = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if re.search(r'^下载新浪财经APP', stripped):
            continue
        if re.match(r'^\|.*\|$', stripped):
            continue
        if re.match(r'^\|.*\d{2}-\d{2}', stripped):
            continue
        if re.match(r'^[#\u3000\s]+$', stripped):
            continue
        line = re.sub(r'^(Stockfisher|stockfisher)[\s\u3000]*', '', line, flags=re.IGNORECASE)
        line = re.sub(r'^\u3000+', '', line)
        line = re.sub(r'^#+\s*', '', line)
        if line.strip():
            cleaned.append(line)
    result = '\n'.join(cleaned)
    result = re.sub(r'\n{3,}', '\n\n', result)
    return result.strip()


def _is_garbled(text):
    """判断文本是否疑似乱码（大量全角空格 或 连续控制字符）"""
    if not text or len(text) < 10:
        return False
    wide_space = text.count('\u3000')
    # 全角空格过多（>5且占比>5%）→ 乱码
    if wide_space > 5 and wide_space > len(text) * 0.05:
        return True
    import re
    # 含连续不可打印ASCII控制字符 → 乱码
    if re.search(r'[\x00-\x1f]{3,}', text):
        return True
    return False

def build_policy(mdata):
    """第五章：政策方面——按区域拆分归类，不混写"""
    policy = mdata.get('政策动态', {})

    # 区域关键词
    US_KWS = ['美联储', '美国联准', 'Fed', '特朗普', '拜登', '美国国会', '美国政府', '美国商务部', '美国贸易', '美方', '华盛顿', '美国白宫']
    EU_KWS = ['欧洲央行', '欧央行', 'ECB', '欧盟', '德意志', '法国', '英国央行', '意大利', '欧元区', '欧洲经济', '欧洲议会']
    CN_KWS = ['中国', '央行', '人民银行', '银保监会', '证监会', '财政部', '发改委', '商务部', '习近平', '李强', '国务院', '港交所', '沪深', 'A股', '内陆']

    def get_text(item):
        """Extract content field only; do NOT fall back to title (title causes duplication)."""
        content = item.get('内容', '') or item.get('content', '')
        if content and len(content) > 15:
            c = _clean_policy_content(content)
            if c and len(c) > 15 and not _is_garbled(c):
                return c
        return ''

    def split_sentences(text):
        """Split text into sentences, protecting percentages and numbers."""
        if not text:
            return []
        import re
        # Split on Chinese sentence punctuation, but NOT on . inside numbers (95.2%)
        parts = re.split(r'(?<=[。；？！?!])(?!\s*[a-zA-Z0-9%])', text)
        result = []
        for p in parts:
            p = p.strip()
            if not p or len(p) < 6:
                continue
            # If starts with number/symbol (broken mid-sentence), merge backward
            if result and re.match(r'^[a-z0-9%\u3000]', p):
                prev = result.pop()
                merged = prev.rstrip('。；？！') + '。' + p
                result.append(merged)
            else:
                result.append(p)
        return result

    def assign_region(sentence):
        """Assign a sentence to US/EU/CN based on keywords. Returns None if unclear."""
        s = sentence

        # ── 特殊case优先排除（日本不是中国，即使含"央行"也排除）─────────
        if '日本' in s or '日央行' in s:
            return None  # 日本政策暂不在 US/EU/CN 三类中，跳过

        us_score = sum(1 for kw in US_KWS if kw in s)
        eu_score = sum(1 for kw in EU_KWS if kw in s)
        cn_score = sum(1 for kw in CN_KWS if kw in s)

        # 如果是混合新闻（多个区域同时出现），优先按主体归类
        # 美联储/欧洲央行/中国央行 明确归类
        if any(kw in s for kw in ['美联储', 'Fed', '欧央行', '欧洲央行', '人民银行', '中国央行']):
            if '美联储' in s or 'Fed' in s:
                return 'US'
            if '欧央行' in s or '欧洲央行' in s or 'ECB' in s:
                return 'EU'
            if '人民银行' in s or '中国央行' in s:
                return 'CN'

        # 按分数
        scores = {'US': us_score, 'EU': eu_score, 'CN': cn_score}
        max_score = max(scores.values())
        if max_score == 0:
            return None
        # 取最高分区域；分数相同则优先 US > EU > CN
        for region in ['US', 'EU', 'CN']:
            if scores[region] == max_score:
                return region
        return None

    def accumulate(region_buckets):
        """Convert buckets {region: [sentences]} into trimmed paragraphs.
        Deduplicate by first 30-char prefix to avoid near-duplicates."""
        output = {}
        for region, sents in region_buckets.items():
            seen = set()
            unique = []
            for s in sents:
                key = s[:30]
                if key not in seen:
                    seen.add(key)
                    unique.append(s)
            trimmed = unique[:4]
            if trimmed:
                result = '。'.join(trimmed)
                if not result.endswith('。'):
                    result += '。'
                output[region] = result
            else:
                output[region] = None
        return output

    # 按区域收集所有新闻句
    region_buckets = {'US': [], 'EU': [], 'CN': []}

    for region_key, news_list in policy.items():
        if not isinstance(news_list, list):
            continue
        for item in news_list:
            text = get_text(item)
            if not text:
                continue
            sentences = split_sentences(text)
            for sent in sentences:
                region = assign_region(sent)
                if region and region in region_buckets:
                    region_buckets[region].append(sent)

    # 合并同类
    merged = accumulate(region_buckets)

    # 输出顺序：US, EU, CN
    FALLBACK = {
        'US': '美国政策动态：无新增重要信息。',
        'EU': '欧洲政策动态：无新增重要信息。',
        'CN': '中国政策动态：无新增重要信息。',
    }
    return [
        merged.get('US') or FALLBACK['US'],
        merged.get('EU') or FALLBACK['EU'],
        merged.get('CN') or FALLBACK['CN'],
    ]

def build_tech(mdata):
    """
    第四章：科技方面
    来源：market_data['企业动态']（Step2采集后存为'企业动态'，代码已兼容查找'科技企业动态'兜底）
    兼容两种数据格式：
      1. 新格式（Step2采集）: {'美股科技': [...], 'A股科技': [...], '港股科技': [...], '欧日韩科技': [...]}
      2. 旧格式（其他来源）: {'美国': [...], '中国': [...]} (直接按地区)
    筛选：巨头并购、资本运作、境外IPO等核心事件
    ...
    """
    import re

    # 兼容两种数据格式
    tech = mdata.get('企业动态', {}) or mdata.get('科技企业动态', {})
    all_items = []
    for region_items in tech.values():
        if isinstance(region_items, list):
            all_items.extend(region_items)

    # ── 核心事件关键词 ──
    CORE_KWS = [
        '并购', '收购', '合并', '要约收购', '私有化',
        '上市', 'IPO', '境外上市', '港股上市', '美股上市', '纳斯达克', '纽交所',
        '回购', '配股', '减持', '增持', '定向增发',
        '拆分', '分拆', '募资', '融资', '参股', '合资', '战略合作',
    ]
    # ── 排除词（公司名含这些→跳过） ──
    SKIP_COMPANY = [
        '新能源', '汽车', '小订', '投资者', '二手车', '货运',
        '全球资本峰会', '直播预告', '财报日历', '直击业绩会',
        'AI快讯', 'IPO热点小时报', '投资日志', '股票回购一览',
        '主力资金', '研报', '行业报告', '评级', '目标价',
        '每日经济新闻', '新浪财经', '搜狐财经', '证券之星',
        'Class A', 'Class B', 'A股股票',
        # ── 导航/栏目碎片词（公司名含这些说明是噪声采集） ──
        '要闻', '金融', '评论', '产经', '创投', '滚动',
        '新股', '基金', '港美股',
    ]
    # ── 排除词（事件含这些→跳过）
    # 注：避免单字符匹配误伤，如'涨'/'跌'在"股价上涨"等正常新闻中也出现
    SKIP_EVENT = [
        '新能源', '汽车', '小订', '盘前', '收盘',
        '分析师', '预计', '看好', '有望', '目标价', '评级',
        '机构', '券商',
        '2023年', '2024年', '2025年',
    ]

    # ── 中国科技企业白名单 ──
    CHINA_TECH_NAMES = [
        '阿里', '腾讯', '字节', '百度', '京东', '美团', '小米', '华为',
        'OPPO', 'vivo', '比亚迪', '宁德时代', '中芯', '中芯国际',
        '科大讯飞', '商汤', '寒武纪', '海康威视', '金龙', '搜狐',
        '网易', '新浪', '知乎', '哔哩', 'b站', 'B站', '阿里云',
        '字节跳动', '云深处', '望圆科技', '航锦科技', '泰坦科技',
        '百胜中国', '港股', 'A股', '中概',
    ]

    # ── 事件正文噪声检测（导航碎片/标签行） ──
    NAV_FRAGMENT_PATTERNS = [
        r'^[\u4e00-\u9fff]{1,4}\s+[\u4e00-\u9fff]{1,4}\s+[\u4e00-\u9fff]{1,4}\s+[\u4e00-\u9fff]{1,4}(\s+[\u4e00-\u9fff]{1,4})+$',
        r'^(?:京|津|冀|晋|蒙|辽|吉|黑|沪|苏|浙|皖|闽|赣|鲁|豫|鄂|湘|粤|桂|琼|渝|川|黔|滇|藏|陕|甘|青|宁|新){5,}.*$',
        r'^(?:要闻|金融|评论|产经|创投|滚动|新股|基金|股票|期货|外汇|债券|理财)[\s\u4e00-\u9fff]{20,}$',
    ]
    NAV_FRAGMENT_COMPANY_KEYWORDS = [
        '要闻', '金融', '评论', '产经', '创投', '滚动',
        '新股', '基金', '港美股',
    ]

    def is_nav_fragment(text):
        """检测是否为导航/标签栏碎片文本"""
        if not text or len(text.strip()) < 15:
            return True
        for pat in NAV_FRAGMENT_PATTERNS:
            if re.match(pat, text.strip()):
                return True
        words = text.strip().split()
        if len(words) >= 6:
            short_words = [w for w in words if len(w) <= 4 and re.match(r'^[\u4e00-\u9fff]+$', w)]
            if len(short_words) >= 5:
                if not re.search(r'\d', text) and not any(text.rstrip().endswith(p) for p in '。？！?'):
                    return True
        return False

    def is_core_event(item):
        """宽松版本：接受核心资本事件 或 中国科技公司新闻"""
        event = item.get('事件', '')
        company = item.get('公司', '')
        event_first_30 = event[:30]

        # 排除词
        if any(kw in company for kw in SKIP_COMPANY):
            return False
        if any(kw in event for kw in SKIP_EVENT):
            return False
        # ── 导航碎片检测（正文是噪声） ──
        if is_nav_fragment(event):
            return False
        # 事件前30字含4+个导航栏目词 → 正文被污染，跳过
        nav_word_count = sum(1 for kw in NAV_FRAGMENT_COMPANY_KEYWORDS if kw in event_first_30)
        if nav_word_count >= 4:
            return False

        tech_names = [
            '阿里', '腾讯', '字节', '百度', '京东', '美团', '小米', '华为',
            'OPPO', 'vivo', '比亚迪', '宁德时代', '中芯', '中芯国际',
            '科大讯飞', '商汤', '寒武纪', '海康威视', '金龙', '搜狐',
            '网易', '新浪', '知乎', '哔哩', 'b站', 'B站', '阿里云',
            '字节跳动', '云深处', '望圆科技', '航锦科技', '泰坦科技',
            '百胜中国', '港股', 'A股', '中概', '英伟达', 'nvidia', 'AMD',
            '苹果', 'Apple', '微软', 'Google', 'Meta', '特斯拉', 'ASML',
        ]

        company_is_tech = any(name in company or name in event for name in tech_names)
        combined = event + company
        if company_is_tech:
            return True
        if any(kw in combined for kw in CORE_KWS):
            return True
        return False

    def is_china_tech(item):
        company = item.get('公司', '')
        event = item.get('事件', '')
        if any(name in company for name in CHINA_TECH_NAMES):
            return True
        if any(name in event for name in CHINA_TECH_NAMES):
            return True
        if re.search(r'[\u4e00-\u9fff]', company) and 4 <= len(company) <= 20:
            return True
        if re.search(r'\d{5}\.HK', company):
            return True
        return False

    def is_real_company(company):
        if not company or len(company.strip()) < 3:
            return False
        bad_patterns = [
            '快讯', '日报', '时报', '投资日志', '热点', '回购一览',
            'AI快讯', '全球', '直播', '研报', '行业', 'Class A', 'Class B', 'A股股票',
        ]
        if any(p in company for p in bad_patterns):
            return False
        if len(company) > 50:  # 放宽中文标题长度限制
            return False
        return True

    def clean_event(text):
        if not text:
            return ''
        text = re.sub(r'\(文章[。.][^)]*\)', '', text)
        text = re.sub(r'\d{4}[-/]\d{2}[-/]\d{2}\s*[·\-–—]\S.*$', '', text)
        text = re.sub(r'[·\-]\S{2,}(?:官方)?(?:网易号|搜狐号|百家号)?$', '', text)
        text = re.sub(r'\s+', ' ', text).strip()
        text = re.sub(r'_{1,2}[^_]+$', '', text)
        return text

    # 过滤 + 去重
    candidates = [it for it in all_items if is_core_event(it)]
    seen_keys = set()
    deduped = []
    for item in candidates:
        key = clean_event(item.get('事件', ''))[:50]
        if key not in seen_keys:
            seen_keys.add(key)
            deduped.append(item)

    def _end_with_punct(text):
        """确保段落以中文句号结尾（无则补加）"""
        if not text:
            return text
        if text.endswith('。') or text.endswith('！') or text.endswith('？'):
            return text
        return text + '。'

    def _smart_truncate(text, max_len=350):
        """
        智能截断：
          1. 跳过标题（冒号「：」前内容），从冒号后第一句开始
          2. 跳过引子句（如"这两天/据CNBC多方证实"等）
          3. 取第一个实质性句子（在 max_len 内），以句号结尾
        """
        if not text:
            return ''
        LEAD_IN_PATTERNS = [
            r'^这两天[,，]?\s*',
            r'^据(CNBC|彭博社|华尔街|路透社)[多方]*[,，]?\s*',
            r'^值得关注的是[,，]?\s*',
            r'^公开资料显示?[,，]?\s*',
            r'^据\d{4}[-/]\d{2}[-/]\d{2}[,，]?\s*',
        ]

        # 1. 如果有「：」，跳过标题部分，从冒号后内容开始
        colon_idx = text.find('：')
        if colon_idx > 3 and colon_idx < 60:
            text = text[colon_idx + 1:].strip()

        # 2. 跳过开头的引子句（到第一个句号）
        temp = text
        for pat in LEAD_IN_PATTERNS:
            temp = re.sub(pat, '', temp)
        # 如果去掉引子后句子变短了（说明有引子被去掉），说明首个句号是引子句尾
        # 需要找到第二个句号
        if temp != text:
            # 找第一个句号位置（在原始 text 中）
            first_dot = text.find('。')
            second_dot = text.find('。', first_dot + 1) if first_dot >= 0 else -1
            if second_dot > first_dot and second_dot < max_len:
                text = text[second_dot + 1:].strip()

        # 3. 在 max_len 内找完整句子
        m = re.search(r'[。；？！?!]', text[:max_len])
        if m:
            end_pos = m.start() + 1
            result = text[:end_pos].strip()
            if result:
                return _end_with_punct(result)

        # 4. 无完整句：找逗号断点
        segment = text[:max_len]
        cp = segment.rfind('，')
        dp = segment.rfind('、')
        cut = max(cp, dp)
        if cut > max_len * 0.4:
            result = segment[:cut].strip()
            return _end_with_punct(result)

        # 5. 截断至 max_len 并补句号
        result = segment.strip().rstrip(',，、;；:.。')
        return _end_with_punct(result) if result else ''

    def build_paragraph(item):
        event = clean_event(item.get('事件', ''))
        company = item.get('公司', '').strip()
        if not event:
            return ''
        # 智能截断：保留完整句子，不截断在中间
        truncated = _smart_truncate(event, max_len=200)
        if is_real_company(company) and 3 <= len(company) <= 25:
            # 公司名作为前缀，注意不要与事件开头重复
            if truncated and not truncated.startswith(company[:4]) and not truncated.startswith(company[:2]):
                text = f"{company}：{truncated}"
            else:
                text = truncated
        else:
            text = truncated
        return text

    # 构建段落（最多3个，优先中国科技企业）
    paragraphs = []
    used_keys = set()

    # 1. 中国科技企业
    china_items = [it for it in deduped
                   if is_china_tech(it) and is_real_company(it.get('公司', ''))]
    for item in china_items:
        if len(paragraphs) >= 3:
            break
        key = clean_event(item.get('事件', ''))[:50]
        if key in used_keys:
            continue
        used_keys.add(key)
        para = build_paragraph(item)
        if para:
            paragraphs.append(para)

    # 2. 其余真实公司（补足至3条）
    for item in deduped:
        if len(paragraphs) >= 3:
            break
        key = clean_event(item.get('事件', ''))[:50]
        if key in used_keys:
            continue
        company = item.get('公司', '').strip()
        if not is_real_company(company):
            continue
        used_keys.add(key)
        para = build_paragraph(item)
        if para:
            paragraphs.append(para)

    return paragraphs if paragraphs else ['暂无科技企业重大动态。']

def build_other_markets(mdata):
    """第五章：其他市场"""
    eu_dax    = (gp('德国DAX 30', mdata), gc('德国DAX 30', mdata))
    eu_cac    = (gp('法国CAC 40', mdata), gc('法国CAC 40', mdata))
    eu_ftse   = (gp('英国富时100', mdata), gc('英国富时100', mdata))
    asi_nik   = (gp('日经225指数', mdata), gc('日经225指数', mdata))
    asi_kos   = (gp('韩国综合指数', mdata), gc('韩国综合指数', mdata))
    asi_asx   = (gp('澳洲S&P/ASX 200', mdata), gc('澳洲S&P/ASX 200', mdata))
    econ      = mdata.get('经济数据', {})

    # === Europe: merge stock indices + economic data ===
    if eu_dax[1] is not None:
        if eu_cac[1] is not None and eu_ftse[1] is not None:
            all_pos_eu = (eu_dax[1] > 0) and (eu_cac[1] > 0) and (eu_ftse[1] > 0)
            all_neg_eu = (eu_dax[1] < 0) and (eu_cac[1] < 0) and (eu_ftse[1] < 0)
            if all_pos_eu:
                eu_joint = "集体上涨"
            elif all_neg_eu:
                eu_joint = "集体下跌"
            else:
                eu_joint = "涨跌分化"
            eu1 = ("欧洲三大股指" + eu_joint + "：" +
                   "德国DAX指数" + chn(eu_dax[1]) + fpc_abs(eu_dax[1]) +
                   ("：" + fp(eu_dax[0]) + "点 " if eu_dax[0] else "") +
                   "法国CAC40指数" + chn(eu_cac[1]) + fpc_abs(eu_cac[1]) +
                   ("：" + fp(eu_cac[0]) + "点 " if eu_cac[0] else "") +
                   "英国富时100指数" + chn(eu_ftse[1]) + fpc_abs(eu_ftse[1]) +
                   ("：" + fp(eu_ftse[0]) + "点" if eu_ftse[0] else ""))
        elif eu_cac[1] is not None:
            eu_joint = chn_full(eu_dax[1])
            eu1 = ("德国DAX指数" + chn(eu_dax[1]) + fpc_abs(eu_dax[1]) +
                   ("：" + fp(eu_dax[0]) + "点 " if eu_dax[0] else "") +
                   "法国CAC40指数" + chn(eu_cac[1]) + fpc_abs(eu_cac[1]) +
                   ("：" + fp(eu_cac[0]) + "点" if eu_cac[0] else ""))
        else:
            eu1 = ("德国DAX指数" + chn(eu_dax[1]) + fpc_abs(eu_dax[1]) +
                   ("：" + fp(eu_dax[0]) + "点" if eu_dax[0] else ""))
    else:
        eu1 = "欧洲主要股指数据待更新。"

    # Economic data
    eu_pmi   = econ.get('欧元区PMI', {})
    eu_cpi   = econ.get('欧元区CPI', {})
    eu_unemp = econ.get('欧元区失业率', {})
    parts = []
    if eu_pmi.get('最新值'):
        v = eu_pmi.get('最新值', 'N/A')
        parts.append("欧元区4月制造业PMI为" + str(v) + "，" +
                     ("萎缩不振" if float(v) < 50 else "维持扩张") + "，"
                     "服务业PMI为" + str(eu_pmi.get('服务业', 'N/A')) + "。")
    if eu_cpi.get('最新值'):
        v = eu_cpi.get('最新值', 'N/A')
        parts.append("欧元区4月CPI初值为" + str(v) + "，"
                     "核心CPI为" + str(eu_cpi.get('核心CPI', 'N/A')) + "，"
                     "通胀压力" + str(eu_cpi.get('备注', '仍处观察')) + "。")
    if eu_unemp.get('最新值'):
        v = eu_unemp.get('最新值', 'N/A')
        parts.append("欧元区失业率为" + str(v) + "，" +
                     str(eu_unemp.get('备注', '德国和法国就业市场韧性好坏参半')) + "。")

    eu2 = "".join(parts)[:250] if parts else ("欧元区经济前景仍受贸易摩擦和能源供给扰动影响，企业盈利增速趋缓。")
    eu_merged = eu1.rstrip("。") + "。" + eu2 + "。"

    # === Asia: dynamic analysis ===
    if asi_nik[1] is not None:
        all_pos_asi = all(v[1] > 0 for v in [asi_nik, asi_kos, asi_asx] if v[1] is not None)
        all_neg_asi = all(v[1] < 0 for v in [asi_nik, asi_kos, asi_asx] if v[1] is not None)
        if all_pos_asi:
            asia_sum = "亚洲主要股指集体上涨，"
        elif all_neg_asi:
            asia_sum = "亚洲主要股指集体下跌，"
        else:
            asia_sum = "亚洲主要股指涨跌分化，"
        asia_parts = []
        if asi_kos[1] is not None:
            asia_parts.append("韩国综合指数" + chn(asi_kos[1]) + fpc_abs(asi_kos[1]) +
                             ("：" + fp(asi_kos[0]) + "点 " if asi_kos[0] else ""))
        if asi_nik[1] is not None:
            asia_parts.append("日经225指数" + chn(asi_nik[1]) + fpc_abs(asi_nik[1]) +
                             ("：" + fp(asi_nik[0]) + "点 " if asi_nik[0] else ""))
        if asi_asx[1] is not None:
            asia_parts.append("澳洲ASX200指数" + chn(asi_asx[1]) + fpc_abs(asi_asx[1]) +
                             ("：" + fp(asi_asx[0]) + "点" if asi_asx[0] else ""))
        asia_str = "，".join(asia_parts) if asia_parts else "亚洲主要股指数据待更新，"

        if all_pos_asi:
            asia_analysis = "亚洲市场风险偏好回升，科技股和周期股走强提振亚太平稳。"
        elif all_neg_asi:
            asia_analysis = "全球避险情绪升温，外资流出亚太市场。"
        else:
            asia_analysis = "亚太市场走势分化，外资流出抑制亚太市场。"

        asia1 = asia_sum + asia_str + "。" + asia_analysis
    else:
        asia1 = "亚洲市场数据待更新。"

    return [eu_merged, asia1]


def build_economic_calendar(mdata):
    """模块6：今日经济数据"""
    cal = mdata.get('今日经济数据', [])
    if cal:
        items = [item.strip() for item in cal if item.strip()]
        if items:
            return "；".join(items) + "。"
    return ("今日重点关注：09:45 中国4月财新服务业/综合PMI；"
            "17:00 欧元区4月制造业PMI终值；"
            "17:00 欧元区3月失业率；"
            "20:30 美国上周初请失业金人数；"
            "22:00 美国3月成屋销售总数。")

# ─────────────────────────────────────────────
# 主流程
# ─────────────────────────────────────────────

log.info(f"📄 生成 Word 日报（无模板版）... 报告日期: {config.REPORT_DATE}")

if not os.path.exists(config.MARKET_DATA_FILE):
    log.error(f"❌ 数据不存在: {config.MARKET_DATA_FILE}"); sys.exit(1)

with open(config.MARKET_DATA_FILE, 'r', encoding='utf-8') as f:
    mdata = json.load(f)

doc = Document()
builder = DocBuilder(doc)

# 文档标题
builder.add_heading(f"金融市场日报", level=0)
builder.add_subtitle(f"报告日期：{config.REPORT_DATE}　　数据截至：{config.DATA_DATE}")

# 模块1：隔夜环球市场总结
builder.add_heading("一、隔夜环球市场总结", level=1)
builder.add_body(build_global_summary(mdata))

# 模块2：市场表现方面
builder.add_heading("二、市场表现方面", level=1)

# 2.1 美国（5个子章节）
builder.add_heading("2.1  美国", level=2)
us_paras = build_market_performance_us(mdata)
builder.add_heading("美股", level=3)
builder.add_body(us_paras[0])
builder.add_heading("美债", level=3)
builder.add_body(us_paras[1])
builder.add_heading("美元指数", level=3)
builder.add_body(us_paras[2])
builder.add_heading("美国就业市场情况", level=3)
builder.add_body(us_paras[3])

# 2.2 中国及中国香港
builder.add_heading("2.2  中国及中国香港", level=2)
cn_hk_paras = build_market_performance_china_hk(mdata)
builder.add_heading("A股", level=3)
builder.add_body(cn_hk_paras[0])
builder.add_heading("港股", level=3)
builder.add_body(cn_hk_paras[1])

# 2.3 黄金
builder.add_heading("2.3  黄金", level=2)
builder.add_body(build_market_performance_gold(mdata))

# 模块3：政策方面
builder.add_heading("三、政策方面", level=1)
policy_paras = build_policy(mdata)
builder.add_heading("美国政策", level=3)
builder.add_body(policy_paras[0])
builder.add_heading("欧洲政策", level=3)
builder.add_body(policy_paras[1])
builder.add_heading("中国政策", level=3)
builder.add_body(policy_paras[2])

# 模块4：科技方面
builder.add_heading("四、科技方面", level=1)
tech_paras = build_tech(mdata)
for para in tech_paras:
    builder.add_body(para)

# 模块5：其他市场
builder.add_heading("五、其他市场", level=1)
other_paras = build_other_markets(mdata)
builder.add_heading("欧洲市场", level=3)
builder.add_body(other_paras[0])
builder.add_heading("亚太市场", level=3)
builder.add_body(other_paras[1])

# 模块6：今日关注经济数据
builder.add_heading("六、今日关注经济数据", level=1)
builder.add_body(build_economic_calendar(mdata))

# ── 附录：休市情况 ──────────────────────────────────────────────
def build_holiday_notice(mdata):
    """
    休市情况段落生成器，严格遵循三类规范：
      1. 完整休市：当日适逢节日，全球多地市场全日休市
      2. 半日/提前休市：海外市场提前收市，国内正常
      3. 正常交易日：不输出任何内容
    数据来源：market_data['休市情况']
    """
    hol = mdata.get('休市情况', {})
    if not hol:
        return None  # 无数据 → 正常交易日，不输出

    # 解析各市场状态
    def status(key):
        v = hol.get(key, {})
        if isinstance(v, dict):
            return v.get('状态', '待确认')
        return '待确认'

    us_stk   = status('美国股市')
    eu_stk   = status('欧洲股市')
    cn_stk   = status('中国股市')
    hk_stk   = status('港股')
    au_stk   = status('澳大利亚股市')
    cme_metal = status('CME贵金属/美油')
    ice_brent = status('ICE布油')
    cn_fut_night = status('国内期货夜盘')

    def is_full_closed(key_status):
        return '休市' in key_status or '休市' in str(key_status)

    def is_half_day(key_status):
        return '半日' in key_status or '提前' in key_status or '收盘' in key_status

    # ── 场景判断 ──
    # 完整休市：至少有欧美或港股全日休市
    full_closed = (
        is_full_closed(us_stk) or is_full_closed(eu_stk) or is_full_closed(hk_stk)
    )
    # 半日/提前休市：有任一市场提前收市，但无全日休市
    half_day = (
        is_half_day(us_stk) or is_half_day(eu_stk) or is_half_day(hk_stk) or
        is_half_day(au_stk) or is_half_day(cme_metal) or is_half_day(ice_brent)
    )
    if not full_closed and not half_day:
        return None  # 正常交易日，不输出

    if full_closed:
        # ── 场景1：完整休市 ──
        parts = []
        # 开篇
        parts.append("当日适逢节假日，全球多地金融市场休市。")
        # 欧美股市
        markets = []
        if is_full_closed(us_stk):
            markets.append("美国股市")
        if is_full_closed(eu_stk):
            markets.append("欧洲股市")
        if is_full_closed(hk_stk):
            markets.append("中国香港股市")
        if is_full_closed(au_stk):
            markets.append("澳大利亚股市")
        if markets:
            parts.append("其中，" + "、".join(markets) + "全天暂停交易。")
        # CME贵金属 + 美油（合并为一个表述）
        cme_parts = []
        if is_full_closed(cme_metal):
            cme_parts.append("芝商所（CME）贵金属及美油期货")
        if cme_parts:
            parts.append("此外，" + "、".join(cme_parts) + "暂停交易。")
        # ICE布油
        if is_full_closed(ice_brent):
            parts.append("洲际交易所（ICE）布伦特原油期货同步休市。")
        # 国内期货夜盘
        if is_full_closed(cn_fut_night):
            parts.append("国内方面，上海黄金交易所（上金所）、上海期货交易所（上期所）、郑州商品交易所（郑商所）、大连商品交易所（大商所）暂停夜盘交易。")
        return "".join(parts)

    elif half_day:
        # ── 场景2：半日/提前休市 ──
        parts = []
        parts.append("当日海外部分市场受节假日影响，交易时段有所调整。")
        # 美股
        if is_half_day(us_stk):
            parts.append("美国股市提前收盘，交易时段缩短；")
        # 欧洲
        eu_adj = []
        if is_half_day(eu_stk):
            eu_adj.append("欧洲股市")
        if is_half_day(au_stk):
            eu_adj.append("澳大利亚股市")
        if eu_adj:
            parts.append("".join(eu_adj) + "同步调整交易时间；")
        # 港股
        if is_half_day(hk_stk):
            parts.append("中国香港股市提前收市；")
        # CME/ICE
        cme_adj = []
        if is_half_day(cme_metal):
            cme_adj.append("芝商所贵金属及美油期货交易时段调整")
        if is_half_day(ice_brent):
            cme_adj.append("洲际交易所布油期货交易时段调整")
        if cme_adj:
            parts.append("、".join(cme_adj) + "；")
        # 国内
        parts.append("国内方面，上海黄金交易所、上海期货交易所、郑州商品交易所、大连商品交易所及中国各金融市场全日正常交易，夜盘正常运行。")
        text = "".join(parts)
        # 清理末尾标点
        text = text.rstrip(';；') + '。'
        return text

    return None

holiday_para = build_holiday_notice(mdata)
if holiday_para:
    builder.add_heading("七、全球市场休市安排", level=1)
    builder.add_body(holiday_para)

log.info(f"   道指: {gp('道琼斯工业平均指数', mdata)} | 10Y: {gp('10年期美债收益率', mdata)}% | COMEX黄金: ${gp('COMEX黄金期货', mdata)}")

# 保存
out = os.path.join(config.OUTPUT_DIR,
                   f"金融市场日报_{config.TODAY.strftime('%Y%m%d')}.docx")
for attempt in range(3):
    try:
        doc.save(out)
        log.info(f"\n✅ Word 报告已保存: {out}")
        break
    except PermissionError:
        alt_out = os.path.join(config.OUTPUT_DIR,
            f"金融市场日报_{config.TODAY.strftime('%Y%m%d')}_v{attempt+2}.docx")
        try:
            doc.save(alt_out)
            log.info(f"\n✅ Word 报告已保存（文件被锁，已另存为）: {alt_out}")
            out = alt_out
            break
        except PermissionError:
            if attempt == 2: raise
            import time; time.sleep(1)

log.info("   Word 文档生成完成")
