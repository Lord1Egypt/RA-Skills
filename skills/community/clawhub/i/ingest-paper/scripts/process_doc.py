# -*- coding: utf-8 -*-
"""
process_doc.py — 多格式文档文字提取（全类型版，替代旧的 process_pdf.py）

按文件后缀自动选择提取方式：
    .pdf            → pymupdf
    .docx           → python-docx
    .xlsx / .xls    → openpyxl
    .txt / .md      → 直接读取
    其他            → 尝试按纯文本读取

用法：
    python3 process_doc.py --path /tmp/paperkb/xxx.pdf
    python3 process_doc.py --path /tmp/paperkb/notes.md

纯文字输入（用户直接打字的会议纪要等）不经过本脚本——
SKILL.md 指示 OpenClaw 直接把消息文本写成 .txt 再传入，或直接进入分析步骤。

输出单行 JSON：
    {"success": true, "text_path": "...", "chars": 12345,
     "truncated": false, "head": "前600字", "fmt": "pdf"}
"""
from __future__ import annotations

import argparse
import json
import os
import sys
from pathlib import Path

from kb_common import ensure_tmp

MAX_CHARS = int(os.environ.get("PAPERKB_MAX_CHARS", "60000"))


def _out(payload: dict) -> None:
    print(json.dumps(payload, ensure_ascii=False))


def _fail(error: str, message: str) -> None:
    _out({"success": False, "error": error, "message": message})
    sys.exit(0)


def _extract_pdf(path: Path) -> str:
    try:
        import fitz
    except ImportError:
        _fail("pymupdf_missing", "未安装 pymupdf，请运行 setup.sh 安装依赖。")
    doc = fitz.open(str(path))
    parts = [page.get_text() for page in doc]
    doc.close()
    return "\n".join(parts)


def _extract_docx(path: Path) -> str:
    try:
        import docx  # python-docx
    except ImportError:
        _fail("python_docx_missing", "未安装 python-docx，请运行 setup.sh 安装依赖。")
    d = docx.Document(str(path))
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    # 同时提取表格内容
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_xlsx(path: Path) -> str:
    try:
        import openpyxl
    except ImportError:
        _fail("openpyxl_missing", "未安装 openpyxl，请运行 setup.sh 安装依赖。")
    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"# 工作表：{ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                parts.append(" | ".join(cells))
    wb.close()
    return "\n".join(parts)


def _extract_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--path", required=True)
    args = parser.parse_args()

    p = Path(args.path)
    if not p.exists():
        _fail("file_not_found",
              f"找不到文件：{p}。"
              "如果是 arxiv 论文，可按 /tmp/paperkb/arxiv_<arxiv_id>.pdf 重建路径重试。")

    ext = p.suffix.lower()
    try:
        if ext == ".pdf":
            text, fmt = _extract_pdf(p), "pdf"
        elif ext == ".docx":
            text, fmt = _extract_docx(p), "docx"
        elif ext in (".xlsx", ".xls"):
            text, fmt = _extract_xlsx(p), "xlsx"
        elif ext in (".txt", ".md", ".markdown"):
            text, fmt = _extract_text(p), ext.lstrip(".")
        else:
            # 兜底：尝试当纯文本读
            text, fmt = _extract_text(p), "text"
    except SystemExit:
        raise
    except Exception as exc:  # noqa: BLE001
        _fail("extract_failed", f"文件解析失败（{ext}）：{exc}")

    text = (text or "").strip()
    if not text:
        _fail("empty_text",
              f"从文件中提取不到文字（{ext}）。若是纯扫描 PDF，当前版本不支持 OCR；"
              "若是会议纪要等，可直接把文字内容发给我。")

    truncated = len(text) > MAX_CHARS
    if truncated:
        text = text[:MAX_CHARS]

    ensure_tmp()
    txt_path = p.with_suffix(".extracted.txt")
    txt_path.write_text(text, encoding="utf-8")

    _out({
        "success": True,
        "text_path": str(txt_path),
        "chars": len(text),
        "truncated": truncated,
        "head": text[:600],
        "fmt": fmt,
    })


if __name__ == "__main__":
    main()
