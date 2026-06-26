"""
Official Doc Formatter - 公文格式转换核心模块
基于 GB/T 9704-2012 标准，将 Markdown 转换为党政机关公文格式

v1.1.0 升级内容：
  - 引入 markdown-it-py 替换逐行解析，支持多行段落、嵌套列表、表格、图片、链接
  - # 标题智能判断：单个视为大标题（居中不加序号），多个视为一级标题（加序号）
  - 首行缩进修正为 640 twips（精确对应国标2个三号汉字宽度）
  - 新增表格、图片、超链接、代码块支持
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from markdown_it import MarkdownIt
import re
import os
import base64
import tempfile
import urllib.request
import urllib.error

__version__ = "1.1.0"

# ── 国标常量 ──
FIRST_LINE_INDENT_TWIPS = 640  # 首行缩进 640 twips = 2个三号汉字宽度
LINE_SPACING_PT = 26           # 固定行距 26pt（用户有意设置）

# 中文数字映射
CHINESE_NUMS = [
    '一', '二', '三', '四', '五', '六', '七', '八', '九', '十',
    '十一', '十二', '十三', '十四', '十五', '十六', '十七', '十八', '十九', '二十'
]


# ══════════════════════════════════════════════════════════════
#  工具函数
# ══════════════════════════════════════════════════════════════

def set_line_spacing(paragraph, spacing_pt=LINE_SPACING_PT):
    """设置固定行间距"""
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    spacing.set(qn('w:line'), str(int(spacing_pt * 20)))
    spacing.set(qn('w:lineRule'), 'exact')
    pPr.append(spacing)


def set_first_line_indent(paragraph, twips=FIRST_LINE_INDENT_TWIPS):
    """设置首行缩进（640 twips = 国标2个三号汉字）"""
    pPr = paragraph._p.get_or_add_pPr()
    indent = OxmlElement('w:ind')
    indent.set(qn('w:firstLine'), str(twips))
    pPr.append(indent)


def set_font(run, font_name, font_size, bold=False, color=None, underline=False):
    """设置字体样式"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    if underline:
        run.font.underline = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def remove_all_spaces(text):
    """删除所有空格"""
    return text.replace(' ', '')


def convert_quotes(text):
    """转换英文引号为中文引号"""
    text = re.sub(r'"([^"]+)"', '\u201c\\1\u201d', text)
    text = re.sub(r"''([^']+)'", '\u2018\\1\u2019', text)
    text = re.sub(r"'([^']+)'", '\u2018\\1\u2019', text)
    return text


def clean_text(text):
    """清理文本：删除空格 + 转换引号"""
    text = remove_all_spaces(text)
    text = convert_quotes(text)
    return text


def add_page_number(doc):
    """添加页码 — 居中显示，样式：— 1 —"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run1 = paragraph.add_run('— ')
        set_font(run1, '宋体', 14)

        # PAGE 域
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run2 = paragraph.add_run()
        run2._element.append(fldChar1)

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE '
        run3 = paragraph.add_run()
        run3._element.append(instrText)

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run4 = paragraph.add_run()
        run4._element.append(fldChar2)

        run5 = paragraph.add_run(' —')
        set_font(run5, '宋体', 14)


def download_image(url, timeout=10):
    """下载图片到临时文件，返回本地路径；失败返回 None"""
    try:
        if url.startswith('data:image'):
            # base64 内嵌图片
            match = re.match(r'data:image/(\w+);base64,(.*)', url)
            if match:
                ext = match.group(1)
                data = base64.b64decode(match.group(2))
                tmp = tempfile.NamedTemporaryFile(suffix=f'.{ext}', delete=False)
                tmp.write(data)
                tmp.close()
                return tmp.name
        else:
            req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
            with urllib.request.urlopen(req, timeout=timeout) as resp:
                content_type = resp.headers.get('Content-Type', '')
                ext = '.png'
                if 'jpeg' in content_type or 'jpg' in content_type:
                    ext = '.jpg'
                elif 'gif' in content_type:
                    ext = '.gif'
                elif 'webp' in content_type:
                    ext = '.webp'
                tmp = tempfile.NamedTemporaryFile(suffix=ext, delete=False)
                tmp.write(resp.read())
                tmp.close()
                return tmp.name
    except Exception:
        pass
    return None


# ══════════════════════════════════════════════════════════════
#  内联内容渲染（处理加粗、斜体、链接等 inline token）
# ══════════════════════════════════════════════════════════════

def render_inline_content(paragraph, inline_tokens, font_name='仿宋_GB2312', font_size=16):
    """渲染 inline tokens 到段落中，支持加粗、斜体、链接、代码"""
    for token in inline_tokens:
        ttype = token.type
        if ttype == 'text':
            text = clean_text(token.content)
            if text:
                run = paragraph.add_run(text)
                set_font(run, font_name, font_size)
        elif ttype == 'softbreak':
            # 软换行，忽略
            pass
        elif ttype == 'hardbreak':
            run = paragraph.add_run('\n')
            set_font(run, font_name, font_size)
        elif ttype == 'strong_open':
            pass  # 由子节点处理
        elif ttype == 'strong_close':
            pass
        elif ttype == 'em_open':
            pass
        elif ttype == 'em_close':
            pass
        elif ttype == 'link_open':
            pass
        elif ttype == 'link_close':
            pass
        elif ttype == 'code_inline':
            text = clean_text(token.content)
            if text:
                run = paragraph.add_run(text)
                set_font(run, '仿宋_GB2312', font_size)
        elif ttype == 'image':
            # inline image 在段落中跳过，由 block 级 image 处理
            pass
        else:
            # 未知类型，尝试输出文本
            if hasattr(token, 'content') and token.content:
                text = clean_text(token.content)
                if text:
                    run = paragraph.add_run(text)
                    set_font(run, font_name, font_size)


def render_inline_with_formatting(paragraph, tokens, font_name='仿宋_GB2312', font_size=16):
    """
    渲染 inline tokens，支持加粗/斜体/链接的格式切换。
    使用状态机跟踪当前是否在 strong/em/link 内部。
    """
    in_strong = False
    in_em = False
    in_link = False
    link_url = ''

    for token in tokens:
        ttype = token.type

        if ttype == 'strong_open':
            in_strong = True
        elif ttype == 'strong_close':
            in_strong = False
        elif ttype == 'em_open':
            in_em = True
        elif ttype == 'em_close':
            in_em = False
        elif ttype == 'link_open':
            in_link = True
            link_url = token.attrGet('href') or ''
        elif ttype == 'link_close':
            in_link = False
            link_url = ''
        elif ttype == 'text':
            text = clean_text(token.content)
            if text:
                run = paragraph.add_run(text)
                # 根据加粗/斜体切换字体
                if in_strong:
                    set_font(run, '黑体', font_size)
                elif in_em:
                    set_font(run, '楷体_GB2312', font_size)
                else:
                    set_font(run, font_name, font_size)
                # 链接加蓝色下划线
                if in_link:
                    set_font(run, run.font.name or font_name, font_size,
                             color=RGBColor(0x05, 0x63, 0xC1), underline=True)
        elif ttype == 'softbreak':
            pass
        elif ttype == 'hardbreak':
            run = paragraph.add_run('\n')
            set_font(run, font_name, font_size)
        elif ttype == 'code_inline':
            text = clean_text(token.content)
            if text:
                run = paragraph.add_run(text)
                set_font(run, '仿宋_GB2312', font_size)
        elif ttype == 'image':
            pass  # 由 block 级处理
        else:
            if hasattr(token, 'content') and token.content:
                text = clean_text(token.content)
                if text:
                    run = paragraph.add_run(text)
                    set_font(run, font_name, font_size)


# ══════════════════════════════════════════════════════════════
#  核心转换引擎
# ══════════════════════════════════════════════════════════════

def md_to_docx(md_content, output_path):
    """
    将 Markdown 内容转换为符合 GB/T 9704-2012 标准的 Word 文档

    参数:
        md_content (str): Markdown 格式的文本内容
        output_path (str): 输出 Word 文件路径（.docx）

    返回:
        bool: 转换是否成功
    """
    try:
        doc = Document()

        # ── 页面设置 ──
        section = doc.sections[0]
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)

        # ── 解析 Markdown ──
        md = MarkdownIt("commonmark", {"html": False}).enable("table")
        tokens = md.parse(md_content)

        # ── 第一遍扫描：统计 # 标题数量 ──
        h1_count = 0
        for token in tokens:
            if token.type == 'heading_open' and token.tag == 'h1':
                h1_count += 1

        # ── 第二遍渲染 ──
        level1_counter = 0  # ## 对应的一级标题计数
        level2_counter = 0  # ### 对应的二级标题计数
        level3_counter = 0  # #### 对应的三级标题计数

        i = 0
        while i < len(tokens):
            token = tokens[i]
            ttype = token.type

            # ── 标题处理 ──
            if ttype == 'heading_open':
                tag = token.tag  # h1 / h2 / h3 / h4
                # 下一个 token 是 heading_inline（内容）
                i += 1
                content_token = tokens[i]
                heading_text = content_token.content.strip()
                heading_children = content_token.children or []

                if tag == 'h1':
                    if h1_count == 1:
                        # 单个 # → 大标题：居中，小标宋二号，不加序号
                        paragraph = doc.add_paragraph()
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        set_line_spacing(paragraph)
                        render_inline_with_formatting(paragraph, heading_children,
                                                      '方正小标宋简体', 22)
                    else:
                        # 多个 # → 一级标题：加序号，黑体三号
                        level1_counter += 1
                        level2_counter = 0
                        level3_counter = 0
                        prefix = f"{CHINESE_NUMS[level1_counter - 1]}、" if level1_counter <= len(CHINESE_NUMS) else f"{level1_counter}、"
                        paragraph = doc.add_paragraph()
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        set_line_spacing(paragraph)
                        set_first_line_indent(paragraph)
                        run = paragraph.add_run(prefix)
                        set_font(run, '黑体', 16)
                        render_inline_with_formatting(paragraph, heading_children, '黑体', 16)

                elif tag == 'h2':
                    level2_counter += 1
                    level3_counter = 0
                    prefix = f"{CHINESE_NUMS[level2_counter - 1]}、" if level2_counter <= len(CHINESE_NUMS) else f"{level2_counter}、"
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    set_line_spacing(paragraph)
                    set_first_line_indent(paragraph)
                    run = paragraph.add_run(prefix)
                    set_font(run, '黑体', 16)
                    render_inline_with_formatting(paragraph, heading_children, '黑体', 16)

                elif tag == 'h3':
                    level3_counter += 1
                    prefix = f"（{CHINESE_NUMS[level3_counter - 1]}）" if level3_counter <= len(CHINESE_NUMS) else f"（{level3_counter}）"
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    set_line_spacing(paragraph)
                    set_first_line_indent(paragraph)
                    run = paragraph.add_run(prefix)
                    set_font(run, '楷体_GB2312', 16)
                    render_inline_with_formatting(paragraph, heading_children, '楷体_GB2312', 16)

                elif tag == 'h4':
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    set_line_spacing(paragraph)
                    set_first_line_indent(paragraph)
                    render_inline_with_formatting(paragraph, heading_children, '仿宋_GB2312', 16)

                # 跳过 heading_close
                i += 1

            # ── 段落处理（含图片段落识别）──
            elif ttype == 'paragraph_open':
                i += 1
                content_token = tokens[i]
                if content_token.type == 'inline':
                    children = content_token.children or []
                    # 检查是否是纯图片段落（只有一个 image token）
                    is_image_only = (len(children) == 1 and children[0].type == 'image')
                    if is_image_only:
                        img_token = children[0]
                        img_url = img_token.attrGet('src') or ''
                        img_alt = img_token.attrGet('alt') or ''
                        _handle_standalone_image(doc, img_url, img_alt)
                    else:
                        paragraph = doc.add_paragraph()
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        set_line_spacing(paragraph)
                        set_first_line_indent(paragraph)
                        render_inline_with_formatting(paragraph, children, '仿宋_GB2312', 16)
                # 跳过 paragraph_close
                i += 1

            # ── 有序列表 ──
            elif ttype == 'ordered_list_open':
                # 收集整个列表
                list_items = _collect_list_items(tokens, i)
                for idx, (item_children, item_nesting) in enumerate(list_items):
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    set_line_spacing(paragraph)
                    indent_twips = FIRST_LINE_INDENT_TWIPS + item_nesting * 320
                    set_first_line_indent(paragraph, indent_twips)
                    prefix = f"{idx + 1}."
                    run = paragraph.add_run(prefix)
                    set_font(run, '仿宋_GB2312', 16)
                    render_inline_with_formatting(paragraph, item_children, '仿宋_GB2312', 16)
                i = _skip_to_close(tokens, i, 'ordered_list_close')

            # ── 无序列表 ──
            elif ttype == 'bullet_list_open':
                list_items = _collect_list_items(tokens, i)
                for idx, (item_children, item_nesting) in enumerate(list_items):
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    set_line_spacing(paragraph)
                    indent_twips = FIRST_LINE_INDENT_TWIPS + item_nesting * 320
                    set_first_line_indent(paragraph, indent_twips)
                    run = paragraph.add_run('•')
                    set_font(run, '仿宋_GB2312', 16)
                    render_inline_with_formatting(paragraph, item_children, '仿宋_GB2312', 16)
                i = _skip_to_close(tokens, i, 'bullet_list_close')

            # ── 表格处理 ──
            elif ttype == 'table_open':
                table_data = _collect_table(tokens, i)
                if table_data:
                    headers = table_data[0]
                    rows = table_data[1:]
                    num_cols = len(headers)
                    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    table.style = 'Table Grid'

                    # 表头
                    for col_idx, header_text in enumerate(headers):
                        cell = table.rows[0].cells[col_idx]
                        cell.text = ''
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        set_line_spacing(p)
                        run = p.add_run(clean_text(header_text))
                        set_font(run, '黑体', 16)

                    # 数据行
                    for row_idx, row_data in enumerate(rows):
                        for col_idx, cell_text in enumerate(row_data):
                            if col_idx < num_cols:
                                cell = table.rows[row_idx + 1].cells[col_idx]
                                cell.text = ''
                                p = cell.paragraphs[0]
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                set_line_spacing(p)
                                run = p.add_run(clean_text(cell_text))
                                set_font(run, '仿宋_GB2312', 16)

                i = _skip_to_close(tokens, i, 'table_close')

            # ── 代码块 ──
            elif ttype == 'fence':
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                set_line_spacing(paragraph)
                set_first_line_indent(paragraph)
                code_text = token.content.strip()
                run = paragraph.add_run(code_text)
                set_font(run, '仿宋_GB2312', 14)

            # ── 水平线 ──
            elif ttype == 'hr':
                pass  # 忽略

            # ── 引用块 ──
            elif ttype == 'blockquote_open':
                i = _skip_to_close(tokens, i, 'blockquote_close')

            # ── 其他 token 跳过 ──
            else:
                pass

            i += 1

        # ── 添加页码 ──
        add_page_number(doc)

        # ── 保存 ──
        doc.save(output_path)
        return True

    except Exception as e:
        print(f"转换失败: {str(e)}")
        import traceback
        traceback.print_exc()
        return False


# ══════════════════════════════════════════════════════════════
#  辅助函数：token 集合
# ══════════════════════════════════════════════════════════════

def _collect_list_items(tokens, start_idx):
    """
    从 list_open 位置开始，收集所有 list_item 的内容和嵌套层级。
    返回 [(children_tokens, nesting_level), ...]
    """
    items = []
    i = start_idx + 1
    nesting = 0

    while i < len(tokens):
        t = tokens[i]
        if t.type in ('ordered_list_close', 'bullet_list_close'):
            # 检查是否是最外层的 close
            break
        elif t.type == 'list_item_open':
            i += 1
            # 收集 item 内容
            item_children = []
            while i < len(tokens) and tokens[i].type != 'list_item_close':
                if tokens[i].type == 'inline':
                    item_children = tokens[i].children or []
                elif tokens[i].type in ('ordered_list_open', 'bullet_list_open'):
                    # 嵌套列表，递归收集
                    sub_items = _collect_list_items(tokens, i)
                    for sub_children, sub_nesting in sub_items:
                        items.append((sub_children, nesting + 1))
                    i = _skip_to_close(tokens, i,
                                       'ordered_list_close' if tokens[i].type == 'ordered_list_open'
                                       else 'bullet_list_close')
                i += 1
            if item_children:
                items.append((item_children, nesting))
        else:
            i += 1

    return items


def _skip_to_close(tokens, start_idx, close_type):
    """跳转到指定类型的 close token，返回其索引"""
    depth = 1
    i = start_idx + 1
    open_type = close_type.replace('_close', '_open')
    while i < len(tokens) and depth > 0:
        if tokens[i].type == open_type:
            depth += 1
        elif tokens[i].type == close_type:
            depth -= 1
            if depth == 0:
                return i
        i += 1
    return i


def _collect_table(tokens, start_idx):
    """
    从 table_open 开始，收集表格数据。
    返回二维列表 [[cell_text, ...], ...]，第一行为表头。
    """
    rows = []
    current_row = []
    i = start_idx + 1

    while i < len(tokens):
        t = tokens[i]
        if t.type == 'table_close':
            break
        elif t.type == 'tr_open':
            current_row = []
        elif t.type == 'tr_close':
            if current_row:
                rows.append(current_row)
            current_row = []
        elif t.type == 'td_open' or t.type == 'th_open':
            i += 1
            cell_text = ''
            while i < len(tokens) and tokens[i].type not in ('td_close', 'th_close'):
                if tokens[i].type == 'inline':
                    cell_text = tokens[i].content.strip()
                i += 1
            current_row.append(cell_text)
        i += 1

    return rows


# ══════════════════════════════════════════════════════════════
#  独立图片段落处理（在主循环外补充）
# ══════════════════════════════════════════════════════════════

def _handle_standalone_image(doc, url, alt_text=''):
    """处理独立图片段落：下载图片并嵌入 Word，居中显示"""
    local_path = download_image(url)
    if local_path and os.path.exists(local_path):
        try:
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_line_spacing(paragraph)
            run = paragraph.add_run()
            run.add_picture(local_path, width=Inches(5.0))

            # 图注
            if alt_text:
                caption = doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_line_spacing(caption)
                run = caption.add_run(clean_text(alt_text))
                set_font(run, '仿宋_GB2312', 12)
        except Exception:
            # 图片嵌入失败，用文字替代
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_line_spacing(paragraph)
            run = paragraph.add_run(f'[图片: {alt_text or url}]')
            set_font(run, '仿宋_GB2312', 16)
        finally:
            try:
                os.unlink(local_path)
            except Exception:
                pass
    else:
        # 下载失败，用文字替代
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_line_spacing(paragraph)
        run = paragraph.add_run(f'[图片: {alt_text or url}]')
        set_font(run, '仿宋_GB2312', 16)
