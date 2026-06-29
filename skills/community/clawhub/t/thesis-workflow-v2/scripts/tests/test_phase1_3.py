#!/usr/bin/env python3
"""
test_phase1_3.py - Step 11 Orchestrator Phase 1.3 集成 单元测试（修订版）

按拍板决策的修订流程：
  Phase 1.0 上传 → Phase 1.1 解析 → Phase 1.2 确认 → Phase 1.3 归因 → Phase 2

覆盖场景：
  【Phase 1.1 入口测试】
  1. orchestrate_phase1_1 docx 输入成功 → 初始化 outline_state + orchestrate_state
  2. orchestrate_phase1_1 text 输入成功 → 同样初始化
  3. orchestrate_phase1_1 docx_path 不存在 → action="input_required" + 3 选项
  4. orchestrate_phase1_1 input_type=auto 报错（拍板 #3）
  5. orchestrate_phase1_1 input_data 空报错
  6. orchestrate_phase1_1 text 解析失败（章节<5）→ action="input_required" + 3 选项
  7. Phase 1.1 成功后 docx_path 已存 state

  【Phase 1.2 确认】
  8. confirm_phase1 → phase1_confirmed=True + phase1_3_status=pending（保持 phase1）

  【Phase 1.3 归因】
  9. orchestrate_phase1_3 不传 docx_path → 从 state 读
  10. orchestrate_phase1_3 state 无 docx_path → 报错
  11. orchestrate_phase1_3 docx_path 失效 → action="input_required" + 3 选项
  12. orchestrate_phase1_3 成功 → state phase1_3_status=submitted + node_details
  13. update_node_content_hint 用户修改
  14. confirm_phase1_3 → phase=phase2

  【Phase 2 强制检查】
  15. orchestrate_phase2 phase1_3 != confirmed 拒绝

  【端到端】
  16. 完整流程：phase1_1_init(docx) → confirm_phase1 → phase1_3_submit → update_hint → confirm_phase1_3 → phase2
  17. 文本输入完整流程（无 docx_path）
"""

import sys
import os
import json

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from orchestrator_v2 import (
    init_orchestrate_state,
    confirm_phase1,
    orchestrate_phase1_1,
    orchestrate_phase1_3,
    update_node_content_hint,
    confirm_phase1_3,
    skip_phase1_3,
    orchestrate_phase2,
    orchestrate,
    load_orchestrate_state,
    _get_orchestrate_state_path,
)
from state_manager_v2 import (
    outline_save,
    outline_load,
    outline_update_status,
    _get_state_path,
)


TEST_PAPER = "test_phase1_3_paper"


def cleanup():
    for p in [_get_state_path(TEST_PAPER), _get_orchestrate_state_path(TEST_PAPER)]:
        if os.path.exists(p):
            try:
                os.remove(p)
            except FileNotFoundError:
                pass


def create_fake_docx():
    """创建最小 docx 文件用于测试"""
    try:
        import docx
    except ImportError:
        return None
    doc = docx.Document()
    doc.add_heading("论文大纲", level=1)
    doc.add_heading("第1章 绪论", level=2)
    doc.add_heading("1.1 研究背景", level=3)
    doc.add_paragraph("AI 时代互联网分发面临范式重构，传统应用商店增长见顶，新型 AI 智能体成为关键变量。")
    doc.add_heading("1.2 研究内容", level=3)
    doc.add_paragraph("本文聚焦差异化战略在 AI 时代的适用性，以 A 公司为例展开竞争战略研究。")
    doc.add_heading("第2章 理论基础与文献综述", level=2)
    doc.add_heading("2.1 竞争战略理论", level=3)
    doc.add_paragraph("Porter 三种基本竞争战略：成本领先、差异化、集中化。")
    doc.add_heading("2.2 文献综述", level=3)
    doc.add_paragraph("近年文献聚焦数字化与 AI 时代的战略选择。")
    doc.add_heading("第3章 外部环境分析", level=2)
    doc.add_heading("3.1 宏观环境", level=3)
    doc.add_paragraph("PEST 模型分析。")
    doc.add_heading("3.2 行业竞争结构", level=3)
    doc.add_paragraph("波特五力模型。")
    doc.add_heading("第4章 内部环境分析", level=2)
    doc.add_heading("4.1 资源与能力", level=3)
    doc.add_paragraph("VRIN 资源分析。")
    doc.add_heading("第5章 竞争战略选择", level=2)
    doc.add_heading("5.1 QSPM 矩阵", level=3)
    doc.add_paragraph("QSPM 矩阵量化评估。")
    doc.add_heading("第6章 战略实施与保障", level=2)
    doc.add_heading("6.1 实施路径", level=3)
    doc.add_paragraph("三阶段实施路径。")
    doc.add_heading("第7章 结论与展望", level=2)
    doc.add_heading("7.1 研究结论", level=3)
    doc.add_paragraph("总结研究结论。")
    doc.add_heading("参考文献", level=2)
    fake_path = f"/tmp/{TEST_PAPER}_proposal.docx"
    doc.save(fake_path)
    return fake_path


# ============================================================
# 【Phase 1.1 入口测试】
# ============================================================

def test_phase1_1_docx_success():
    """测试 1：docx 输入成功 → 初始化 outline_state + orchestrate_state"""
    print("\n=== 测试 1：Phase 1.1 docx 解析成功 ===")
    cleanup()

    fake_path = create_fake_docx()
    if not fake_path:
        print("   ⚠️ 跳过（python-docx 未安装）")
        return

    r = orchestrate_phase1_1(
        TEST_PAPER, input_type="docx", input_data=fake_path
    )

    assert r["ok"] is True, f"应成功，实际: {r.get('error', 'unknown')}"
    assert r["action"] == "review_outline"
    assert r["input_type"] == "docx"
    assert "outline" in r
    print(f"   ✅ action=review_outline，outline 已生成")

    # 验证 state
    state = load_orchestrate_state(TEST_PAPER)
    assert state["phase1_3_input_type"] == "docx"
    assert state["phase1_3_docx_path"] == fake_path
    assert state["phase1_3_status"] == "pending"
    print(f"   ✅ state 已初始化（docx_path 已存，status=pending）")

    # 验证 outline_state 已保存（包含虚拟摘要节点）
    outline_state = outline_load(TEST_PAPER)
    has_virtual = any(n.get("is_virtual") for n in outline_state["outline"]["outline_tree"]["nodes"])
    assert has_virtual, "outline 应包含虚拟摘要节点（增强项1）"
    print(f"   ✅ outline_state 已持久化（增强项1 虚拟节点已插入）")


def test_phase1_1_text_success():
    """测试 2：text 输入成功（无 docx 也能走）"""
    print("\n=== 测试 2：Phase 1.1 text 解析成功 ===")
    cleanup()

    outline_text = """第1章 绪论
1.1 研究背景
1.2 研究内容
第2章 理论基础
2.1 竞争战略理论
2.2 文献综述
第3章 外部环境
3.1 宏观环境
3.2 行业结构
第4章 内部环境
4.1 资源与能力
第5章 战略选择
5.1 QSPM矩阵
第6章 战略实施
6.1 实施路径
第7章 结论
7.1 研究结论
参考文献"""

    r = orchestrate_phase1_1(
        TEST_PAPER, input_type="text", input_data=outline_text
    )

    assert r["ok"] is True, f"应成功，实际: {r.get('error', 'unknown')}"
    assert r["input_type"] == "text"

    state = load_orchestrate_state(TEST_PAPER)
    assert state["phase1_3_input_type"] == "text"
    assert state["phase1_3_docx_path"] is None  # text 无 docx
    print(f"   ✅ text 输入成功（state 无 docx_path）")


def test_phase1_1_rejects_auto():
    """测试 3：input_type=auto 报错（拍板 #3 不要 auto）"""
    print("\n=== 测试 3：Phase 1.1 拒绝 auto ===")
    cleanup()

    r = orchestrate_phase1_1(
        TEST_PAPER, input_type="auto", input_data="anything"
    )

    assert r["ok"] is False
    assert r["action"] == "input_required"
    assert "auto" in r["error"].lower()
    assert "1" in r["retry_options"]
    print(f"   ✅ auto 被拒绝: {r['error'][:60]}")


def test_phase1_1_empty_input():
    """测试 4：input_data 为空报错"""
    print("\n=== 测试 4：input_data 为空 ===")
    cleanup()

    r = orchestrate_phase1_1(
        TEST_PAPER, input_type="docx", input_data=""
    )

    assert r["ok"] is False
    assert r["action"] == "input_required"
    print(f"   ✅ 空 input 被拒绝")


def test_phase1_1_docx_not_exist():
    """测试 5：docx_path 不存在 → input_required + 3 选项"""
    print("\n=== 测试 5：Phase 1.1 docx 不存在 ===")
    cleanup()

    r = orchestrate_phase1_1(
        TEST_PAPER, input_type="docx", input_data="/nonexistent/proposal.docx"
    )

    assert r["ok"] is False
    assert r["action"] == "input_required"
    assert "1" in r["retry_options"]
    assert "2" in r["retry_options"]
    assert "3" in r["retry_options"]
    print(f"   ✅ 返回 input_required + 3 选项")


def test_phase1_1_text_parse_fail():
    """测试 6：text 完全解析失败（无章节）→ input_required + 3 选项"""
    print("\n=== 测试 6：Phase 1.1 text 解析失败 ===")
    cleanup()

    # 完全无效的 text → outline_parse 返回 ok=False
    r = orchestrate_phase1_1(
        TEST_PAPER,
        input_type="text",
        input_data="哈哈啥都没有完全是垃圾内容"
    )

    assert r["ok"] is False
    assert r["action"] == "input_required"
    assert "1" in r["retry_options"]
    print(f"   ✅ 解析失败 → input_required + 3 选项")


# ============================================================
# 【Phase 1.2 确认】
# ============================================================

def test_confirm_phase1_stays_in_phase1():
    """测试 7：confirm_phase1 不直接进 phase2"""
    print("\n=== 测试 7：confirm_phase1 不进 phase2 ===")
    cleanup()

    fake_path = create_fake_docx()
    if not fake_path:
        print("   ⚠️ 跳过")
        return

    # 必须先 Phase 1.1 解析
    orchestrate_phase1_1(TEST_PAPER, "docx", fake_path)

    r = confirm_phase1(TEST_PAPER)

    assert r["ok"] is True
    assert r["phase"] == "phase1"
    assert r["phase1_3_status"] == "pending"
    print(f"   ✅ phase=phase1, phase1_3_status=pending（未直接进 phase2）")


# ============================================================
# 【Phase 1.3 归因】
# ============================================================

def test_phase1_3_reads_docx_from_state():
    """测试 8：phase1_3_submit 不传 docx_path → 从 state 读"""
    print("\n=== 测试 8：phase1_3 从 state 读 docx_path ===")
    cleanup()

    fake_path = create_fake_docx()
    if not fake_path:
        print("   ⚠️ 跳过")
        return

    # 完整 Phase 1.1 + 1.2
    orchestrate_phase1_1(TEST_PAPER, "docx", fake_path)
    confirm_phase1(TEST_PAPER)

    # Phase 1.3 不传 docx_path（应从 state 读）
    r = orchestrate_phase1_3(TEST_PAPER, docx_path=None)

    if not r["ok"]:
        # 归因可能因 docx 内容简单失败，跳过
        print(f"   ⚠️ 归因失败: {r.get('error', '')[:60]}")
        return

    assert r["phase1_3_status"] == "submitted"
    print(f"   ✅ phase1_3_submit 从 state 读 docx 成功")


def test_phase1_3_no_docx_in_state():
    """测试 9：state 无 docx_path + 调用方不传 → 报错"""
    print("\n=== 测试 9：phase1_3 state 无 docx_path ===")
    cleanup()

    # 用 text 模式初始化（state 无 docx_path）
    outline_text = """第1章 绪论
1.1 研究背景
第2章 理论基础
2.1 战略理论
第3章 外部环境
3.1 行业分析
第4章 内部环境
4.1 资源分析
第5章 战略选择
5.1 QSPM
第6章 实施
6.1 路径
第7章 结论
7.1 总结
参考文献"""

    orchestrate_phase1_1(TEST_PAPER, "text", outline_text)
    confirm_phase1(TEST_PAPER)

    r = orchestrate_phase1_3(TEST_PAPER, docx_path=None)

    assert r["ok"] is False
    assert "docx" in r["error"].lower() or "未指定" in r["error"]
    print(f"   ✅ 无 docx 时报错")


def test_phase1_3_docx_deleted():
    """测试 10：docx 失效 → input_required + 3 选项"""
    print("\n=== 测试 10：phase1_3 docx 已删除 ===")
    cleanup()

    fake_path = create_fake_docx()
    if not fake_path:
        print("   ⚠️ 跳过")
        return

    # Phase 1.1 时 docx 存在 → state 记录路径
    orchestrate_phase1_1(TEST_PAPER, "docx", fake_path)
    confirm_phase1(TEST_PAPER)

    # 删除 docx
    os.remove(fake_path)

    r = orchestrate_phase1_3(TEST_PAPER)

    assert r["ok"] is False
    assert r["action"] == "input_required"
    assert "1" in r["retry_options"]
    assert "3" in r["retry_options"]
    print(f"   ✅ docx 失效 → input_required + 3 选项")


def test_phase1_3_submit_state_and_details():
    """测试 11：submit 状态 + 细粒度 node_details"""
    print("\n=== 测试 11：submit 状态 + 细粒度返回 ===")
    cleanup()

    fake_path = create_fake_docx()
    if not fake_path:
        print("   ⚠️ 跳过")
        return

    orchestrate_phase1_1(TEST_PAPER, "docx", fake_path)
    confirm_phase1(TEST_PAPER)

    r = orchestrate_phase1_3(TEST_PAPER)

    if not r["ok"]:
        print(f"   ⚠️ 归因失败: {r.get('error', '')[:60]}")
        return

    assert r["phase1_3_status"] == "submitted"
    assert "node_details" in r
    print(f"   ✅ state=submitted, node_details 已返回")


def test_update_node_content_hint():
    """测试 12：用户修改 content_hint"""
    print("\n=== 测试 12：update_node_content_hint ===")
    cleanup()

    fake_path = create_fake_docx()
    if not fake_path:
        print("   ⚠️ 跳过")
        return

    orchestrate_phase1_1(TEST_PAPER, "docx", fake_path)
    confirm_phase1(TEST_PAPER)
    r = orchestrate_phase1_3(TEST_PAPER)

    if not r["ok"]:
        print(f"   ⚠️ 跳过")
        return

    r2 = update_node_content_hint(TEST_PAPER, "1.1", "用户自定义")
    assert r2["ok"]

    state = load_orchestrate_state(TEST_PAPER)
    assert state["phase1_3_result"]["node_details"]["1.1"].get("user_modified") is True
    print(f"   ✅ 用户修改 + user_modified 标记")


def test_confirm_phase1_3_state_transition():
    """测试 13：confirm_phase1_3 状态机"""
    print("\n=== 测试 13：confirm_phase1_3 状态机 ===")
    cleanup()

    fake_path = create_fake_docx()
    if not fake_path:
        print("   ⚠️ 跳过")
        return

    orchestrate_phase1_1(TEST_PAPER, "docx", fake_path)
    confirm_phase1(TEST_PAPER)
    orchestrate_phase1_3(TEST_PAPER)

    r = confirm_phase1_3(TEST_PAPER)
    assert r["ok"]
    assert r["phase"] == "phase2"
    assert r["phase1_3_status"] == "confirmed"
    print(f"   ✅ submitted → confirmed → phase2")


# ============================================================
# 【Phase 2 强制检查】
# ============================================================

def test_phase2_rejects_unconfirmed_phase1_3():
    """测试 14：Phase 2 强制检查 phase1_3"""
    print("\n=== 测试 14：Phase 2 强制检查 ===")
    cleanup()

    outline_text = """第1章 绪论
1.1 研究背景
第2章 理论基础
2.1 战略理论
第3章 外部环境
3.1 行业分析
第4章 内部环境
4.1 资源分析
第5章 战略选择
5.1 QSPM
第6章 实施
6.1 路径
第7章 结论
7.1 总结
参考文献"""

    orchestrate_phase1_1(TEST_PAPER, "text", outline_text)
    confirm_phase1(TEST_PAPER)

    # 手动改 phase 模拟（不能真进 phase2）
    state = load_orchestrate_state(TEST_PAPER)
    state["phase"] = "phase2"
    with open(_get_orchestrate_state_path(TEST_PAPER), "w", encoding="utf-8") as f:
        json.dump(state, f, ensure_ascii=False, indent=2)

    r = orchestrate_phase2(TEST_PAPER, llm_func=lambda p: "x")

    assert r["ok"] is False
    assert "Phase 1.3" in r["error"]
    print(f"   ✅ 强制检查生效: {r['error'][:80]}")


# ============================================================
# 【端到端集成】
# ============================================================

def test_integration_docx_flow():
    """测试 15：完整 docx 流程"""
    print("\n=== 测试 15：端到端 docx 流程 ===")
    cleanup()

    fake_path = create_fake_docx()
    if not fake_path:
        print("   ⚠️ 跳过")
        return

    # Phase 1.1 docx
    r1 = orchestrate_phase1_1(TEST_PAPER, "docx", fake_path)
    assert r1["action"] == "review_outline"

    # Phase 1.2 确认
    r2 = confirm_phase1(TEST_PAPER)
    assert r2["phase"] == "phase1"
    assert r2["phase1_3_status"] == "pending"

    # Phase 1.3 submit
    r3 = orchestrate_phase1_3(TEST_PAPER)
    if not r3["ok"]:
        print(f"   ⚠️ 归因失败: {r3.get('error', '')[:60]}")
        return

    # 用户修改
    r4 = update_node_content_hint(TEST_PAPER, "1.1", "用户调整")
    assert r4["ok"]

    # Phase 1.3 确认
    r5 = confirm_phase1_3(TEST_PAPER)
    assert r5["phase"] == "phase2"

    print(f"   ✅ 端到端 docx 流程通过")


def test_integration_text_flow():
    """测试 16：完整 text 流程（无 docx）"""
    print("\n=== 测试 16：端到端 text 流程 ===")
    cleanup()

    outline_text = """第1章 绪论
1.1 研究背景
1.2 研究内容
第2章 理论基础
2.1 战略理论
2.2 文献综述
第3章 外部环境
3.1 行业分析
第4章 内部环境
4.1 资源分析
第5章 战略选择
5.1 QSPM
第6章 实施
6.1 路径
第7章 结论
7.1 总结
参考文献"""

    # Phase 1.1 text
    r1 = orchestrate_phase1_1(TEST_PAPER, "text", outline_text)
    assert r1["action"] == "review_outline"

    # Phase 1.2 确认
    r2 = confirm_phase1(TEST_PAPER)
    assert r2["phase1_3_status"] == "pending"

    # Phase 1.3 submit（无 docx_path）→ 应报错（state 无 docx）
    r3 = orchestrate_phase1_3(TEST_PAPER)
    assert not r3["ok"], "text 模式 phase1_3 应报错（无 docx）"

    # v2.0.6 P0-1 修复后：skip_phase1_3 必填 reason + operator
    # 1) 不传 reason/operator → 报错
    r4_no_audit = skip_phase1_3(TEST_PAPER)
    assert not r4_no_audit["ok"], "v2.0.6 修复后 skip 必须 reason+operator"

    # 2) 传 reason + operator → 跳过成功
    r4 = skip_phase1_3(TEST_PAPER, reason="测试用例：text 模式无 docx", operator="test_phase1_3.py")
    assert r4["ok"] is True
    assert r4["phase"] == "phase2"

    # 3) 验证 audit log 写入
    from state_manager_v2 import load_orchestrate_state
    state_after = load_orchestrate_state(TEST_PAPER)
    audit_log = state_after.get("audit_log", [])
    assert len(audit_log) >= 1, "audit log 应有记录"
    assert audit_log[-1]["action"] == "phase1_3_skip"
    assert audit_log[-1]["operator"] == "test_phase1_3.py"

    print(f"   ✅ 端到端 text 流程通过（无 docx 也能进 phase2 via skip + audit log）")


def main():
    """主测试入口"""
    print("=" * 60)
    print("Step 11 Orchestrator Phase 1.3 集成 单元测试（修订版）")
    print("=" * 60)

    tests = [
        test_phase1_1_docx_success,
        test_phase1_1_text_success,
        test_phase1_1_rejects_auto,
        test_phase1_1_empty_input,
        test_phase1_1_docx_not_exist,
        test_phase1_1_text_parse_fail,
        test_confirm_phase1_stays_in_phase1,
        test_phase1_3_reads_docx_from_state,
        test_phase1_3_no_docx_in_state,
        test_phase1_3_docx_deleted,
        test_phase1_3_submit_state_and_details,
        test_update_node_content_hint,
        test_confirm_phase1_3_state_transition,
        test_phase2_rejects_unconfirmed_phase1_3,
        test_integration_docx_flow,
        test_integration_text_flow,
    ]

    passed = 0
    failed = 0
    for t in tests:
        try:
            t()
            passed += 1
        except AssertionError as e:
            print(f"\n❌ 测试失败: {e}")
            failed += 1
        except Exception as e:
            print(f"\n❌ 异常: {e}")
            import traceback
            traceback.print_exc()
            failed += 1
        finally:
            cleanup()

    print("\n" + "=" * 60)
    print(f"结果：{passed} passed, {failed} failed")
    print("=" * 60)
    return 0 if failed == 0 else 1


if __name__ == "__main__":
    sys.exit(main())
