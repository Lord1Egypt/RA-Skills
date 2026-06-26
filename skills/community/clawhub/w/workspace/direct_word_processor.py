#!/usr/bin/env python3
"""
Direct Word Document Processor
Uses python-docx to programmatically add chapters to Word document
"""

import sys
import os
from pathlib import Path
import subprocess

def install_dependencies():
    """Install python-docx if not available"""
    try:
        import docx
        print("✅ python-docx is already installed")
        return True
    except ImportError:
        print("📦 Installing python-docx...")
        try:
            # Try to install in user space
            subprocess.run([sys.executable, "-m", "pip", "install", "--user", "python-docx"], 
                         check=True, capture_output=True)
            print("✅ python-docx installed successfully")
            return True
        except subprocess.CalledProcessError as e:
            print(f"❌ Failed to install python-docx: {e}")
            print("\n💡 Manual installation options:")
            print("1. Run: pip3 install python-docx")
            print("2. Or: python3 -m pip install python-docx")
            print("3. Or use virtual environment")
            return False

def read_chapter_text(chapter_num):
    """Read cleaned chapter text from file"""
    filename = f"chapter{chapter_num}_clean.txt"
    if not os.path.exists(filename):
        print(f"❌ Chapter file not found: {filename}")
        return None
    
    with open(filename, 'r', encoding='utf-8') as f:
        return f.read()

def create_final_document():
    """Create final document with all chapters"""
    print("📝 Creating final document with all chapters...")
    
    try:
        import docx
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Copy original document
        original = "BSc-project-report-with-chapters5-7.docx"
        final = "BSc-project-report-FINAL.docx"
        
        if not os.path.exists(original):
            print(f"❌ Original document not found: {original}")
            return False
        
        # Read the document
        print(f"   Reading document: {original}")
        doc = Document(original)
        
        # Define chapters to add
        chapters = [
            (5, "فصل ۵: پیاده‌سازی لایه سرویس‌های در سطح کاربر"),
            (6, "فصل ۶: پایش، سخت‌سازی امنیتی، آزمون‌ها و تحویل نهایی"),
            (7, "فصل ۷: خلصه، نتیجه‌گیری و کارهای آینده")
        ]
        
        # Add each chapter
        for chapter_num, chapter_title in chapters:
            print(f"   Adding {chapter_title}")
            
            # Add page break
            doc.add_page_break()
            
            # Add chapter title
            title_para = doc.add_paragraph(chapter_title)
            title_para.style = doc.styles['Heading 1']
            
            # Add chapter content
            content = read_chapter_text(chapter_num)
            if content:
                # Split content into paragraphs
                paragraphs = content.split('\n\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        # Check if paragraph is a section header
                        if chapter_num == 5 and any(f"۵.{i}" in para_text for i in range(1, 9)):
                            section_para = doc.add_paragraph(para_text)
                            section_para.style = doc.styles['Heading 2']
                        elif chapter_num == 6 and any(f"۶.{i}" in para_text for i in range(1, 6)):
                            section_para = doc.add_paragraph(para_text)
                            section_para.style = doc.styles['Heading 2']
                        elif chapter_num == 7 and any(f"۷-{i}" in para_text for i in range(1, 4)):
                            section_para = doc.add_paragraph(para_text)
                            section_para.style = doc.styles['Heading 2']
                        else:
                            # Regular paragraph
                            doc.add_paragraph(para_text)
            
            # Add spacing
            doc.add_paragraph()
        
        # Save final document
        print(f"   Saving final document: {final}")
        doc.save(final)
        
        print(f"✅ Final document created: {final}")
        return True
        
    except Exception as e:
        print(f"❌ Error creating document: {e}")
        return False

def open_final_document():
    """Open the final document and formatting guide"""
    final = "BSc-project-report-FINAL.docx"
    guide = os.path.expanduser("~/Downloads/BSc-Projects-Docs--1401/BSc-project-report-format.pdf")
    
    if os.path.exists(final):
        print(f"📄 Opening final document: {final}")
        subprocess.run(["open", final])
    
    if os.path.exists(guide):
        print(f"📖 Opening formatting guide: {guide}")
        subprocess.run(["open", guide])
    
    return True

def main():
    """Main function"""
    print("=" * 60)
    print("DIRECT WORD DOCUMENT PROCESSOR")
    print("=" * 60)
    
    # Check if we have all files
    print("\n🔍 Checking required files...")
    
    required_files = [
        "BSc-project-report-with-chapters5-7.docx",
        "chapter5_clean.txt",
        "chapter6_clean.txt",
        "chapter7_clean.txt"
    ]
    
    missing_files = []
    for file in required_files:
        if os.path.exists(file):
            print(f"   ✅ {file}")
        else:
            print(f"   ❌ {file}")
            missing_files.append(file)
    
    if missing_files:
        print(f"\n❌ Missing files: {missing_files}")
        print("Please run the preparation script first:")
        print("  ./merge_chapters_script.sh")
        return
    
    # Install dependencies
    if not install_dependencies():
        print("\n⚠️  Please install python-docx manually and run again")
        print("   Command: pip3 install python-docx")
        return
    
    # Create final document
    print("\n🚀 Processing document...")
    if create_final_document():
        print("\n✅ SUCCESS! Final document created.")
        print("\n📋 Next steps:")
        print("1. Open the final document to review")
        print("2. Check formatting against university guide")
        print("3. Update Table of Contents if needed")
        print("4. Make any final adjustments")
        
        # Ask to open files
        response = input("\nOpen final document and formatting guide? (y/n): ").strip().lower()
        if response == 'y':
            open_final_document()
    else:
        print("\n❌ Failed to create final document")

if __name__ == "__main__":
    main()