#!/usr/bin/env python3
"""Render final reports for the evidence-enhanced paper analysis workflow.

Given structured extraction JSON and verification JSON, create Markdown, HTML,
and DOCX reports. This mirrors the Dify Scheme A aggregation node while adding
local file outputs that are easy to inspect on Ubuntu Desktop.
"""
from __future__ import annotations

import argparse
import html
import json
import os
import re
import sys
import zipfile
from pathlib import Path
from typing import Any


def parse_json_text(text: str) -> dict[str, Any]:
    if not text:
        return {}
    s = text.strip()
    s = re.sub(r"^```json\s*", "", s)
    s = re.sub(r"^```", "", s)
    s = re.sub(r"```$", "", s)
    s = s.strip()
    try:
        data = json.loads(s)
        return data if isinstance(data, dict) else {"raw_text": data}
    except Exception:
        match = re.search(r"\{[\s\S]*\}", s)
        if match:
            try:
                data = json.loads(match.group(0))
                return data if isinstance(data, dict) else {"raw_text": data}
            except Exception:
                return {"raw_text": s}
        return {"raw_text": s}


def is_english(language: str) -> bool:
    return (language or "").strip() == "英文"


def empty_text(language: str) -> str:
    return "None" if is_english(language) else "无"


def clean_item(x: Any) -> str:
    if isinstance(x, (dict, list)):
        return json.dumps(x, ensure_ascii=False)
    return str(x).strip()


def list_md(items: Any, language: str) -> str:
    fallback = empty_text(language)
    if not items:
        return fallback
    if isinstance(items, list):
        lines = [f"- {clean_item(x)}" for x in items if clean_item(x)]
        return "\n".join(lines) if lines else fallback
    s = clean_item(items)
    return s if s else fallback


def results_md(items: Any, language: str) -> str:
    fallback = empty_text(language)
    if not items or not isinstance(items, list):
        return fallback
    lines = []
    en = is_english(language)
    for it in items:
        if isinstance(it, dict):
            if en:
                lines.append(f"- Dataset: {it.get('dataset','')}; Metric: {it.get('metric','')}; Value: {it.get('value','')}; Baseline: {it.get('baseline','')}; Improvement: {it.get('improvement','')}")
            else:
                lines.append(f"- 数据集：{it.get('dataset','')}；指标：{it.get('metric','')}；结果：{it.get('value','')}；对比基线：{it.get('baseline','')}；提升：{it.get('improvement','')}")
        else:
            s = clean_item(it)
            if s:
                lines.append(f"- {s}")
    return "\n".join(lines) if lines else fallback


def evidence_md(items: Any, language: str) -> str:
    fallback = empty_text(language)
    if not items or not isinstance(items, list):
        return fallback
    en = is_english(language)
    lines = []
    for it in items[:8]:
        if isinstance(it, dict):
            if en:
                lines.append(f"- Field: {it.get('field','')} | Claim: {it.get('claim','')} | Evidence: {it.get('evidence','')}")
            else:
                lines.append(f"- 字段：{it.get('field','')}｜结论：{it.get('claim','')}｜证据：{it.get('evidence','')}")
        else:
            s = clean_item(it)
            if s:
                lines.append(f"- {s}")
    return "\n".join(lines) if lines else fallback


def verify_md(items: Any, language: str) -> str:
    fallback = empty_text(language)
    if not items or not isinstance(items, list):
        return fallback
    en = is_english(language)
    lines = []
    for it in items:
        if isinstance(it, dict):
            if en:
                lines.append(f"- Claim: {it.get('claim','')} | Status: {it.get('status','')} | Evidence: {it.get('evidence','')}")
            else:
                lines.append(f"- {it.get('claim','')}｜状态：{it.get('status','')}｜证据：{it.get('evidence','')}")
        else:
            s = clean_item(it)
            if s:
                lines.append(f"- {s}")
    return "\n".join(lines) if lines else fallback


def issues_md(items: Any, language: str) -> str:
    fallback = empty_text(language)
    if not items or not isinstance(items, list):
        return fallback
    en = is_english(language)
    lines = []
    for it in items:
        if isinstance(it, dict):
            if en:
                lines.append(f"- Field: {it.get('field','')} | Problem: {it.get('problem','')} | Severity: {it.get('severity','')}")
            else:
                lines.append(f"- 字段：{it.get('field','')}｜问题：{it.get('problem','')}｜严重度：{it.get('severity','')}")
        else:
            s = clean_item(it)
            if s:
                lines.append(f"- {s}")
    return "\n".join(lines) if lines else fallback


def build_markdown(structured_json: str, verification_json: str, language: str) -> str:
    data = parse_json_text(structured_json)
    verify = parse_json_text(verification_json)
    en = is_english(language)
    title = data.get("title", "Untitled" if en else "未识别题目")

    if en:
        return f"""# Paper Analysis

## Title
{title}

## Task and Problem
- Task: {data.get('task', '')}
- Background: {data.get('background', '')}
- Problem Statement: {data.get('problem_statement', '')}

## Method Overview
- Method Name: {data.get('method_name', '')}
- Method Core: {data.get('method_core', '')}

## Experimental Elements

### Datasets
{list_md(data.get('datasets', []), language)}

### Baselines
{list_md(data.get('baselines', []), language)}

### Metrics
{list_md(data.get('metrics', []), language)}

## Main Results
{results_md(data.get('main_results', []), language)}

## Contributions
{list_md(data.get('contributions', []), language)}

## Ablations and Limitations

### Ablations
{list_md(data.get('ablations', []), language)}

### Limitations
{list_md(data.get('limitations', []), language)}

## Evidence Spans
{evidence_md(data.get('evidence_spans', []), language)}

## Consistency Check
- Overall Score: {verify.get('overall_score', '')}
- Hallucination Risk: {verify.get('hallucination_risk', '')}
- Final Verdict: {verify.get('final_verdict', '')}

### Verified Claims
{verify_md(verify.get('verified_claims', []), language)}

### Issues
{issues_md(verify.get('issues', []), language)}
"""
    return f"""# 论文分析结果

## 论文题目
{title}

## 任务与问题
- 研究任务：{data.get('task', '')}
- 背景：{data.get('background', '')}
- 问题定义：{data.get('problem_statement', '')}

## 方法概述
- 方法名称：{data.get('method_name', '')}
- 方法核心：{data.get('method_core', '')}

## 实验要素

### 数据集
{list_md(data.get('datasets', []), language)}

### 基线方法
{list_md(data.get('baselines', []), language)}

### 评价指标
{list_md(data.get('metrics', []), language)}

## 主要结果
{results_md(data.get('main_results', []), language)}

## 贡献提炼
{list_md(data.get('contributions', []), language)}

## 消融与局限性

### 消融实验
{list_md(data.get('ablations', []), language)}

### 局限性
{list_md(data.get('limitations', []), language)}

## 证据片段
{evidence_md(data.get('evidence_spans', []), language)}

## 一致性校验
- 总评分：{verify.get('overall_score', '')}
- 幻觉风险：{verify.get('hallucination_risk', '')}
- 最终结论：{verify.get('final_verdict', '')}

### 已核验结论
{verify_md(verify.get('verified_claims', []), language)}

### 发现的问题
{issues_md(verify.get('issues', []), language)}
"""


def markdown_to_html(markdown: str, title: str = "Paper Analysis") -> str:
    try:
        import markdown as markdown_lib  # type: ignore
        body = markdown_lib.markdown(markdown, extensions=["tables", "fenced_code", "sane_lists"])
    except Exception:
        body_lines: list[str] = []
        in_ul = False
        for raw in markdown.splitlines():
            line = raw.rstrip()
            if line.startswith("### "):
                if in_ul:
                    body_lines.append("</ul>"); in_ul = False
                body_lines.append(f"<h3>{html.escape(line[4:])}</h3>")
            elif line.startswith("## "):
                if in_ul:
                    body_lines.append("</ul>"); in_ul = False
                body_lines.append(f"<h2>{html.escape(line[3:])}</h2>")
            elif line.startswith("# "):
                if in_ul:
                    body_lines.append("</ul>"); in_ul = False
                body_lines.append(f"<h1>{html.escape(line[2:])}</h1>")
            elif line.startswith("- "):
                if not in_ul:
                    body_lines.append("<ul>"); in_ul = True
                body_lines.append(f"<li>{html.escape(line[2:])}</li>")
            elif not line.strip():
                if in_ul:
                    body_lines.append("</ul>"); in_ul = False
            else:
                if in_ul:
                    body_lines.append("</ul>"); in_ul = False
                body_lines.append(f"<p>{html.escape(line)}</p>")
        if in_ul:
            body_lines.append("</ul>")
        body = "\n".join(body_lines)
    return f"""<!doctype html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<title>{html.escape(title)}</title>
<style>
body {{ max-width: 920px; margin: 40px auto; padding: 0 24px; font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif; line-height: 1.65; color: #1f2937; }}
h1, h2, h3 {{ color: #111827; line-height: 1.25; }}
h1 {{ border-bottom: 2px solid #e5e7eb; padding-bottom: 12px; }}
h2 {{ margin-top: 32px; border-bottom: 1px solid #e5e7eb; padding-bottom: 6px; }}
li {{ margin: 6px 0; }}
code, pre {{ background: #f3f4f6; }}
</style>
</head>
<body>
{body}
</body>
</html>
"""


def write_docx_with_python_docx(markdown: str, output: Path) -> bool:
    try:
        from docx import Document  # type: ignore
    except Exception:
        return False
    doc = Document()
    for raw in markdown.splitlines():
        line = raw.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(line)
    doc.save(output)
    return True


def write_minimal_docx(markdown: str, output: Path) -> None:
    def xml_escape(s: str) -> str:
        return html.escape(s, quote=False)
    paras = []
    for raw in markdown.splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("#"):
            line = line.lstrip("#").strip()
        elif line.startswith("- "):
            line = "• " + line[2:].strip()
        paras.append(f"<w:p><w:r><w:t>{xml_escape(line)}</w:t></w:r></w:p>")
    document_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:body>%s<w:sectPr><w:pgSz w:w="12240" w:h="15840"/><w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440"/></w:sectPr></w:body></w:document>""" % "".join(paras)
    content_types = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"><Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/><Default Extension="xml" ContentType="application/xml"/><Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/></Types>"""
    rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/></Relationships>"""
    output.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", content_types)
        zf.writestr("_rels/.rels", rels)
        zf.writestr("word/document.xml", document_xml)


def write_docx(markdown: str, output: Path) -> None:
    output.parent.mkdir(parents=True, exist_ok=True)
    if not write_docx_with_python_docx(markdown, output):
        write_minimal_docx(markdown, output)


def render_one(structured_path: Path, verification_path: Path, language: str, output_md: Path, output_html: Path, output_docx: Path) -> dict[str, str]:
    structured = structured_path.read_text(encoding="utf-8", errors="ignore")
    verification = verification_path.read_text(encoding="utf-8", errors="ignore")
    markdown = build_markdown(structured, verification, language)
    output_md.parent.mkdir(parents=True, exist_ok=True)
    output_md.write_text(markdown, encoding="utf-8")
    output_html.write_text(markdown_to_html(markdown), encoding="utf-8")
    write_docx(markdown, output_docx)
    return {"markdown": str(output_md), "html": str(output_html), "docx": str(output_docx)}


def render_manifest(manifest_path: Path) -> list[dict[str, str]]:
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    outputs: list[dict[str, str]] = []
    language = manifest.get("language", "中文")
    for paper in manifest.get("papers", []):
        result = render_one(
            Path(paper["structured_json_file"]),
            Path(paper["verification_json_file"]),
            paper.get("language") or language,
            Path(paper["final_markdown_file"]),
            Path(paper["final_html_file"]),
            Path(paper["final_docx_file"]),
        )
        outputs.append({"paper_id": paper.get("id", ""), **result})
    return outputs


def main() -> int:
    parser = argparse.ArgumentParser(description="Render paper analysis reports from generated JSON files.")
    parser.add_argument("--manifest", help="Manifest created by prepare_papers.py. If set, per-paper paths are read from the manifest.")
    parser.add_argument("--structured-json", help="Path to structured_result.json for single-paper mode.")
    parser.add_argument("--verification-json", help="Path to verification_result.json for single-paper mode.")
    parser.add_argument("--language", default="中文", choices=["中文", "英文"])
    parser.add_argument("--output-md", help="Output Markdown path for single-paper mode.")
    parser.add_argument("--output-html", help="Output rendered HTML path for single-paper mode.")
    parser.add_argument("--output-docx", help="Output DOCX path for single-paper mode.")
    args = parser.parse_args()

    if args.manifest:
        outputs = render_manifest(Path(args.manifest).expanduser())
        print(json.dumps({"outputs": outputs}, ensure_ascii=False, indent=2))
        return 0

    required = [args.structured_json, args.verification_json, args.output_md, args.output_html, args.output_docx]
    if not all(required):
        parser.error("Either --manifest or all single-paper paths are required.")
    result = render_one(
        Path(args.structured_json).expanduser(),
        Path(args.verification_json).expanduser(),
        args.language,
        Path(args.output_md).expanduser(),
        Path(args.output_html).expanduser(),
        Path(args.output_docx).expanduser(),
    )
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
