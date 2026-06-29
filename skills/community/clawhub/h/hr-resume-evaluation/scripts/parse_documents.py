from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass
from pathlib import Path


SUPPORTED_RESUME_SUFFIXES = {".md", ".markdown", ".txt", ".pdf", ".docx"}
TEXT_SUFFIXES = {".md", ".markdown", ".txt"}


@dataclass(frozen=True)
class ParsedDocument:
    path: Path
    file_name: str
    candidate_id: str
    text: str
    ok: bool
    error: str | None = None


def stable_candidate_id(path: Path) -> str:
    stem = Path(path).stem.strip().lower()
    slug = re.sub(r"[^A-Za-z0-9]+", "-", stem).strip("-")
    if stem.isascii():
        return slug or "candidate"

    digest = hashlib.sha256(stem.encode("utf-8")).hexdigest()[:8]
    return f"{slug}-{digest}" if slug else f"candidate-{digest}"


def discover_resumes(folder: Path) -> list[Path]:
    return sorted(
        path
        for path in Path(folder).iterdir()
        if path.is_file() and path.suffix.lower() in SUPPORTED_RESUME_SUFFIXES
    )


def parse_document(path: Path) -> ParsedDocument:
    path = Path(path)
    candidate_id = stable_candidate_id(path)
    suffix = path.suffix.lower()

    if suffix not in SUPPORTED_RESUME_SUFFIXES:
        return ParsedDocument(path, path.name, candidate_id, "", False, "unsupported_extension")

    try:
        if suffix in TEXT_SUFFIXES:
            text = path.read_text(encoding="utf-8-sig")
        elif suffix == ".pdf":
            text = _read_pdf(path)
        else:
            text = _read_docx(path)
    except Exception as exc:
        return ParsedDocument(path, path.name, candidate_id, "", False, f"parse_error: {exc}")

    normalized_text = _normalize_text(text)
    if not normalized_text:
        return ParsedDocument(path, path.name, candidate_id, "", False, "empty_text")

    return ParsedDocument(path, path.name, candidate_id, normalized_text, True)


def _read_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(path)
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _read_docx(path: Path) -> str:
    from docx import Document

    document = Document(path)
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def _normalize_text(text: str) -> str:
    lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    return "\n".join(line.rstrip() for line in lines).strip()
