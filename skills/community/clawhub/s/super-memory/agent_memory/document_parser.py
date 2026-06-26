from __future__ import annotations

import os
import re
import logging
from typing import List, Dict, Optional
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)


@dataclass
class DocumentSection:
    title: str
    level: int
    content: str
    page_num: int = 0
    char_offset: int = 0
    position: int = 0
    metadata: dict = field(default_factory=dict)


@dataclass
class ParsedDocument:
    title: str
    source_path: str
    source_type: str
    sections: List[DocumentSection]
    total_chars: int = 0
    metadata: dict = field(default_factory=dict)


class DocumentParser:

    SUPPORTED_FORMATS = {'.pdf', '.md', '.markdown', '.docx', '.doc', '.txt', '.rst', '.html', '.htm'}

    _MD_HEADING_RE = re.compile(r'^(#{1,6})\s+(.+)$', re.MULTILINE)
    _HTML_HEADING_RE = re.compile(r'<h([1-6])[^>]*>(.*?)</h[1-6]>', re.IGNORECASE | re.DOTALL)
    _HTML_PARA_RE = re.compile(r'<p[^>]*>(.*?)</p>', re.IGNORECASE | re.DOTALL)
    _HTML_TAG_RE = re.compile(r'<[^>]+>')
    _HTML_ENTITY_RE = re.compile(r'&(?:nbsp|lt|gt|amp|quot);')
    _CHINESE_CHAPTER_RE = re.compile(r'^第[一二三四五六七八九十百千\d]+[章节篇部]')
    _CHINESE_NUM_RE = re.compile(r'^[一二三四五六七八九十]+[、.]')
    _PAREN_CHINESE_RE = re.compile(r'^[（(][一二三四五六七八九十]+[）)]')
    _NUMBERED_SECTION_RE = re.compile(r'^(\d+)(\.\d+)*[.、\s]')

    _HTML_ENTITY_MAP = {
        '&nbsp;': ' ', '&lt;': '<', '&gt;': '>', '&amp;': '&', '&quot;': '"',
    }

    def parse(self, source: str, source_type: str = None) -> ParsedDocument:
        fmt = source_type or self.detect_format(source)

        if fmt == 'markdown':
            text = self._read_text(source)
            return self._parse_markdown(text, source)
        elif fmt == 'html':
            text = self._read_text(source)
            return self._parse_html(text, source)
        elif fmt == 'pdf':
            return self._parse_pdf(source)
        elif fmt == 'docx':
            return self._parse_docx(source)
        elif fmt in ('text', 'rst'):
            text = self._read_text(source)
            return self._parse_text(text, source)
        else:
            text = self._read_text(source)
            return self._parse_text(text, source)

    def _parse_markdown(self, text: str, source_path: str = "") -> ParsedDocument:
        lines = text.split('\n')
        sections: List[DocumentSection] = []
        current_title = ""
        current_level = 0
        current_content_lines: List[str] = []
        position = 0
        char_offset = 0

        doc_title = os.path.splitext(os.path.basename(source_path))[0] if source_path else ""

        for line in lines:
            m = self._MD_HEADING_RE.match(line)
            if m:
                if current_title or current_content_lines:
                    content = '\n'.join(current_content_lines).strip()
                    if content or current_title:
                        sections.append(DocumentSection(
                            title=current_title,
                            level=current_level if current_level else 0,
                            content=content,
                            char_offset=char_offset,
                            position=position,
                        ))
                        position += 1
                        char_offset += len(content) + len(current_title) + 1

                current_level = len(m.group(1))
                current_title = m.group(2).strip()
                current_content_lines = []

                if not doc_title and current_level == 1:
                    doc_title = current_title
            else:
                current_content_lines.append(line)

        if current_title or current_content_lines:
            content = '\n'.join(current_content_lines).strip()
            sections.append(DocumentSection(
                title=current_title,
                level=current_level if current_level else 0,
                content=content,
                char_offset=char_offset,
                position=position,
            ))

        self._merge_short_sections(sections)

        total_chars = sum(len(s.content) + len(s.title) for s in sections)

        return ParsedDocument(
            title=doc_title,
            source_path=source_path,
            source_type='markdown',
            sections=sections,
            total_chars=total_chars,
        )

    def _parse_text(self, text: str, source_path: str = "") -> ParsedDocument:
        lines = text.split('\n')
        sections: List[DocumentSection] = []
        position = 0
        char_offset = 0

        doc_title = os.path.splitext(os.path.basename(source_path))[0] if source_path else ""

        current_title = ""
        current_level = 0
        current_content_lines: List[str] = []

        for line in lines:
            stripped = line.rstrip()

            if not stripped.strip():
                if current_content_lines:
                    content = '\n'.join(current_content_lines).strip()
                    if content:
                        sections.append(DocumentSection(
                            title=current_title,
                            level=current_level,
                            content=content,
                            char_offset=char_offset,
                            position=position,
                        ))
                        position += 1
                        char_offset += len(content) + len(current_title) + 1
                    current_title = ""
                    current_level = 0
                    current_content_lines = []
                continue

            cn_level, cn_title = self._detect_chinese_chapter(stripped)
            num_level, num_title = self._detect_numbered_section(stripped)

            detected_level = 0
            detected_title = ""

            if cn_level > 0:
                detected_level = cn_level
                detected_title = cn_title
            elif num_level > 0:
                detected_level = num_level
                detected_title = num_title
            elif stripped.startswith('  ') or stripped.startswith('\t'):
                current_content_lines.append(stripped)
                continue

            if detected_level > 0:
                if current_content_lines:
                    content = '\n'.join(current_content_lines).strip()
                    if content:
                        sections.append(DocumentSection(
                            title=current_title,
                            level=current_level,
                            content=content,
                            char_offset=char_offset,
                            position=position,
                        ))
                        position += 1
                        char_offset += len(content) + len(current_title) + 1

                current_title = detected_title
                current_level = detected_level
                current_content_lines = []

                if not doc_title and detected_level == 1:
                    doc_title = detected_title
            else:
                current_content_lines.append(stripped)

        if current_content_lines:
            content = '\n'.join(current_content_lines).strip()
            if content:
                sections.append(DocumentSection(
                    title=current_title,
                    level=current_level,
                    content=content,
                    char_offset=char_offset,
                    position=position,
                ))

        self._merge_short_sections(sections)

        total_chars = sum(len(s.content) + len(s.title) for s in sections)

        return ParsedDocument(
            title=doc_title,
            source_path=source_path,
            source_type='text',
            sections=sections,
            total_chars=total_chars,
        )

    def _parse_html(self, text: str, source_path: str = "") -> ParsedDocument:
        sections: List[DocumentSection] = []
        position = 0
        char_offset = 0

        doc_title = os.path.splitext(os.path.basename(source_path))[0] if source_path else ""

        title_match = re.search(r'<title[^>]*>(.*?)</title>', text, re.IGNORECASE | re.DOTALL)
        if title_match:
            doc_title = self._strip_html_tags(title_match.group(1)).strip()

        for m in self._HTML_HEADING_RE.finditer(text):
            level = int(m.group(1))
            raw_title = self._strip_html_tags(m.group(2)).strip()
            if not raw_title:
                continue
            sections.append(DocumentSection(
                title=raw_title,
                level=level,
                content="",
                char_offset=m.start(),
                position=position,
            ))
            position += 1
            char_offset += len(raw_title) + 1

            if not doc_title and level == 1:
                doc_title = raw_title

        for m in self._HTML_PARA_RE.finditer(text):
            raw_text = self._strip_html_tags(m.group(1)).strip()
            if not raw_text:
                continue
            sections.append(DocumentSection(
                title="",
                level=0,
                content=raw_text,
                char_offset=m.start(),
                position=position,
            ))
            position += 1
            char_offset += len(raw_text) + 1

        sections.sort(key=lambda s: s.char_offset)

        for i, sec in enumerate(sections):
            sec.position = i

        self._merge_short_sections(sections)

        total_chars = sum(len(s.content) + len(s.title) for s in sections)

        return ParsedDocument(
            title=doc_title,
            source_path=source_path,
            source_type='html',
            sections=sections,
            total_chars=total_chars,
        )

    def _parse_pdf(self, file_path: str) -> ParsedDocument:
        try:
            import PyPDF2
        except ImportError:
            logger.warning("PyPDF2 not available, falling back to text parsing")
            text = self._read_text(file_path)
            doc = self._parse_text(text, file_path)
            doc.source_type = 'pdf'
            return doc

        sections: List[DocumentSection] = []
        position = 0
        char_offset = 0
        doc_title = os.path.splitext(os.path.basename(file_path))[0]

        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page_num, page in enumerate(reader.pages, 1):
                    page_text = page.extract_text() or ""
                    if not page_text.strip():
                        continue

                    page_lines = page_text.split('\n')
                    for line in page_lines:
                        stripped = line.strip()
                        if not stripped:
                            continue

                        cn_level, cn_title = self._detect_chinese_chapter(stripped)
                        num_level, num_title = self._detect_numbered_section(stripped)

                        if cn_level > 0:
                            sections.append(DocumentSection(
                                title=cn_title,
                                level=cn_level,
                                content="",
                                page_num=page_num,
                                char_offset=char_offset,
                                position=position,
                            ))
                            position += 1
                            char_offset += len(cn_title) + 1
                        elif num_level > 0:
                            sections.append(DocumentSection(
                                title=num_title,
                                level=num_level,
                                content="",
                                page_num=page_num,
                                char_offset=char_offset,
                                position=position,
                            ))
                            position += 1
                            char_offset += len(num_title) + 1
                        else:
                            sections.append(DocumentSection(
                                title="",
                                level=0,
                                content=stripped,
                                page_num=page_num,
                                char_offset=char_offset,
                                position=position,
                            ))
                            position += 1
                            char_offset += len(stripped) + 1

                    if not doc_title and page_num == 1 and sections:
                        first = sections[0]
                        if first.title:
                            doc_title = first.title
        except Exception as e:
            logger.warning("PDF parsing failed: %s, falling back to text", e)
            text = self._read_text(file_path)
            doc = self._parse_text(text, file_path)
            doc.source_type = 'pdf'
            return doc

        self._merge_short_sections(sections)

        total_chars = sum(len(s.content) + len(s.title) for s in sections)

        return ParsedDocument(
            title=doc_title,
            source_path=file_path,
            source_type='pdf',
            sections=sections,
            total_chars=total_chars,
            metadata={"pages": max((s.page_num for s in sections), default=0)},
        )

    def _parse_docx(self, file_path: str) -> ParsedDocument:
        try:
            import docx as python_docx
        except ImportError:
            logger.warning("python-docx not available, falling back to text parsing")
            text = self._read_text(file_path)
            doc = self._parse_text(text, file_path)
            doc.source_type = 'docx'
            return doc

        sections: List[DocumentSection] = []
        position = 0
        char_offset = 0
        doc_title = os.path.splitext(os.path.basename(file_path))[0]

        try:
            document = python_docx.Document(file_path)
            for para in document.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                style_name = (para.style.name or "").lower() if para.style else ""
                level = 0

                if 'heading 1' in style_name or 'heading1' in style_name:
                    level = 1
                elif 'heading 2' in style_name or 'heading2' in style_name:
                    level = 2
                elif 'heading 3' in style_name or 'heading3' in style_name:
                    level = 3
                elif 'heading 4' in style_name or 'heading4' in style_name:
                    level = 4
                elif 'heading 5' in style_name or 'heading5' in style_name:
                    level = 5
                elif 'heading 6' in style_name or 'heading6' in style_name:
                    level = 6
                elif 'title' in style_name:
                    level = 1
                    doc_title = text
                else:
                    cn_level, cn_title = self._detect_chinese_chapter(text)
                    num_level, num_title = self._detect_numbered_section(text)
                    if cn_level > 0:
                        level = cn_level
                    elif num_level > 0:
                        level = num_level

                if level > 0:
                    sections.append(DocumentSection(
                        title=text,
                        level=level,
                        content="",
                        char_offset=char_offset,
                        position=position,
                    ))
                else:
                    sections.append(DocumentSection(
                        title="",
                        level=0,
                        content=text,
                        char_offset=char_offset,
                        position=position,
                    ))

                position += 1
                char_offset += len(text) + 1
        except Exception as e:
            logger.warning("DOCX parsing failed: %s, falling back to text", e)
            text = self._read_text(file_path)
            doc = self._parse_text(text, file_path)
            doc.source_type = 'docx'
            return doc

        self._merge_short_sections(sections)

        total_chars = sum(len(s.content) + len(s.title) for s in sections)

        return ParsedDocument(
            title=doc_title,
            source_path=file_path,
            source_type='docx',
            sections=sections,
            total_chars=total_chars,
        )

    def _detect_chinese_chapter(self, line: str) -> tuple[int, str]:
        stripped = line.strip()
        if not stripped:
            return (0, "")

        m = self._CHINESE_CHAPTER_RE.match(stripped)
        if m:
            level = 1
            if '节' in m.group(0):
                level = 2
            elif '篇' in m.group(0) or '部' in m.group(0):
                level = 0
            return (level, stripped)

        m = self._PAREN_CHINESE_RE.match(stripped)
        if m:
            return (3, stripped)

        m = self._CHINESE_NUM_RE.match(stripped)
        if m:
            return (2, stripped)

        return (0, "")

    def _detect_numbered_section(self, line: str) -> tuple[int, str]:
        stripped = line.strip()
        if not stripped:
            return (0, "")

        m = self._NUMBERED_SECTION_RE.match(stripped)
        if not m:
            return (0, "")

        dots = stripped[:m.end()].count('.')
        level = dots + 1
        return (level, stripped)

    @staticmethod
    def detect_format(source: str) -> str:
        ext = os.path.splitext(source)[1].lower()
        format_map = {
            '.pdf': 'pdf',
            '.md': 'markdown',
            '.markdown': 'markdown',
            '.docx': 'docx',
            '.doc': 'docx',
            '.txt': 'text',
            '.rst': 'rst',
            '.html': 'html',
            '.htm': 'html',
        }
        return format_map.get(ext, 'text')

    def _read_text(self, file_path: str) -> str:
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception as e:
            logger.warning("Failed to read file %s: %s", file_path, e)
            return ""

    def _strip_html_tags(self, text: str) -> str:
        text = self._HTML_TAG_RE.sub(' ', text)
        for entity, char in self._HTML_ENTITY_MAP.items():
            text = text.replace(entity, char)
        text = re.sub(r'\s+', ' ', text)
        return text.strip()

    def _merge_short_sections(self, sections: List[DocumentSection]) -> None:
        if len(sections) <= 1:
            return

        merged: List[DocumentSection] = [sections[0]]

        for i in range(1, len(sections)):
            prev = merged[-1]
            curr = sections[i]

            is_short = (len(curr.content) + len(curr.title)) < 10 and curr.level == 0
            is_title_only = curr.level > 0 and not curr.content

            if is_short:
                if curr.title:
                    if prev.content:
                        prev.content += '\n' + curr.title
                    else:
                        prev.title = prev.title or curr.title
                        prev.level = prev.level or curr.level
                if curr.content:
                    if prev.content:
                        prev.content += '\n' + curr.content
                    else:
                        prev.content = curr.content
                prev.char_offset = min(prev.char_offset, curr.char_offset)
            else:
                merged.append(curr)

        sections.clear()
        sections.extend(merged)

        for i, sec in enumerate(sections):
            sec.position = i
