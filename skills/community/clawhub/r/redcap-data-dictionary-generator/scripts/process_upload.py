#!/usr/bin/env python3
"""
REDCap 数据字典生成器 - 处理上传的 Word 文档

Usage:
    python process_upload.py <file_path> [--output <output_csv>]
"""

import csv
import io
import sys
import argparse
from pathlib import Path

try:
    from markitdown import MarkItDown
    MARKITDOWN_AVAILABLE = True
except ImportError:
    MARKITDOWN_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


REDCAP_HEADERS = [
    "Variable / Field Name",
    "Form Name",
    "Section Header",
    "Field Type",
    "Field Label",
    "Choices, Calculations, OR Slider Labels",
    "Field Note",
    "Text Validation Type OR Show Slider Number",
    "Text Validation Min",
    "Text Validation Max",
    "Identifier?",
    "Branching Logic (Show field only if...)",
    "Required Field?",
    "Custom Alignment",
    "Question Number (surveys only)",
    "Matrix Group Name"
]


def parse_docx_smarter(docx_path):
    """
    智能解析 DOCX 文件，识别表单和字段结构
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx 未安装")
    
    doc = Document(docx_path)
    fields = []
    current_form = "general"
    current_section = ""
    form_name = "general"
    
    def normalize_form_name(text):
        """标准化表单名"""
        import re
        # 移除非字母数字字符
        text = re.sub(r'[^a-zA-Z0-9\u4e00-\u9fa5]', '', text)
        # 转小写，下划线连接
        text = re.sub(r'\s+', '_', text.lower())
        return text or "general"
    
    def normalize_var_name(text):
        """标准化变量名"""
        import re
        # 只保留字母、数字、下划线
        text = re.sub(r'[^a-zA-Z0-9]', '', text)
        return text.lower()
    
    # 遍历所有段落
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        style_name = para.style.name if para.style else ""
        
        # 检测表单标题（一级标题）
        if style_name in ['Heading 1', 'Heading 1 Char', 'Title'] or _is_form_pattern(text):
            form_name = normalize_form_name(text)
            current_form = form_name
            current_section = ""
            continue
        
        # 检测分节标题（二级标题）
        if style_name in ['Heading 2', 'Heading 2 Char'] or _is_section_pattern(text):
            current_section = text
            continue
        
        # 解析字段定义
        field = _extract_field_from_text(text, current_form, current_section, normalize_var_name)
        if field:
            fields.append(field)
    
    # 遍历所有表格
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) >= 2 and cells[0]:
                field_name = normalize_var_name(cells[0])
                if field_name and len(field_name) > 1:
                    field_label = cells[1] if len(cells) > 1 else field_name
                    field_type = cells[2].lower() if len(cells) > 2 else 'text'
                    choices = cells[3] if len(cells) > 3 else ''
                    
                    field = {
                        'Variable / Field Name': field_name,
                        'Form Name': current_form,
                        'Section Header': current_section,
                        'Field Type': _normalize_field_type(field_type),
                        'Field Label': field_label,
                        'Choices, Calculations, OR Slider Labels': _normalize_choices(choices),
                        'Field Note': '',
                        'Text Validation Type OR Show Slider Number': '',
                        'Text Validation Min': '',
                        'Text Validation Max': '',
                        'Identifier?': 'y' if _is_identifier(field_name) else '',
                        'Branching Logic (Show field only if...)': '',
                        'Required Field?': '',
                        'Custom Alignment': '',
                        'Question Number (surveys only)': '',
                        'Matrix Group Name': ''
                    }
                    fields.append(field)
    
    return fields


def _is_form_pattern(text):
    """判断是否为表单标题模式"""
    patterns = ['表单', 'Form', '模块', 'Module', 'CRF', 'Visit', '访视', '评估']
    return any(p in text for p in patterns) and len(text) < 100


def _is_section_pattern(text):
    """判断是否为分节标题模式"""
    patterns = ['信息', 'Information', '数据', 'Data', '测量', 'Measurement', 
                 '检查', 'Exam', '实验室', 'Lab', '病史', 'History', '体格', 'Physical']
    return any(p in text for p in patterns) and len(text) < 100


def _extract_field_from_text(text, form_name, section, normalize_var_name):
    """从文本中提取字段信息"""
    import re
    
    # 跳过标题行
    if text.startswith('#') or text.endswith(':'):
        return None
    
    # 尝试多种分隔符
    delimiters = ['：', ':', ' - ', '-', '\t']
    parts = None
    
    for d in delimiters:
        if d in text:
            parts = text.split(d, 1)
            break
    
    if not parts or len(parts) != 2:
        return None
    
    var_name = parts[0].strip()
    rest = parts[1].strip()
    
    # 验证变量名
    if not re.match(r'^[a-zA-Z\u4e00-\u9fa5][a-zA-Z0-9_\u4e00-\u9fa5]*$', var_name):
        return None
    
    # 标准化变量名
    var_name = normalize_var_name(var_name)
    if not var_name or len(var_name) < 2:
        return None
    
    # 判断字段类型
    field_type = 'text'
    choices = ''
    
    # 从标签中提取选项信息
    if '[' in rest and ']' in rest:
        # 可能包含选项
        option_match = re.search(r'\[(.*?)\]', rest)
        if option_match:
            options_str = option_match.group(1)
            if '|' in options_str or ',' in options_str:
                choices = _normalize_choices(options_str)
                rest = rest.replace(f'[{options_str}]', '').strip()
    
    # 从变量名推断类型
    if 'date' in var_name.lower():
        field_type = 'text'
    elif '_yn' in var_name.lower() or '_yes' in var_name.lower():
        choices = '0, No | 1, Yes'
        field_type = 'dropdown'
    
    return {
        'Variable / Field Name': var_name,
        'Form Name': form_name,
        'Section Header': section,
        'Field Type': field_type,
        'Field Label': rest,
        'Choices, Calculations, OR Slider Labels': choices,
        'Field Note': '',
        'Text Validation Type OR Show Slider Number': '',
        'Text Validation Min': '',
        'Text Validation Max': '',
        'Identifier?': 'y' if _is_identifier(var_name) else '',
        'Branching Logic (Show field only if...)': '',
        'Required Field?': '',
        'Custom Alignment': '',
        'Question Number (surveys only)': '',
        'Matrix Group Name': ''
    }


def _normalize_field_type(field_type):
    """标准化字段类型"""
    type_map = {
        'dropdown': 'dropdown',
        'drop': 'dropdown',
        'select': 'dropdown',
        'radio': 'radio',
        'checkbox': 'checkbox',
        'check': 'checkbox',
        'text': 'text',
        'string': 'text',
        'notes': 'notes',
        'note': 'notes',
        'textarea': 'notes',
        'calc': 'calc',
        'calculated': 'calc',
        'calculation': 'calc',
        'file': 'file',
        'attachment': 'file',
        'date': 'text',
        'datetime': 'text',
        'number': 'number',
        'integer': 'integer',
        'numeric': 'number'
    }
    return type_map.get(field_type.lower().strip(), 'text')


def _normalize_choices(choices_str):
    """标准化选项格式"""
    if not choices_str:
        return ''
    
    # 替换各种分隔符
    choices_str = choices_str.replace(',', ' | ')
    choices_str = choices_str.replace('；', ' | ')
    choices_str = choices_str.replace(';', ' | ')
    
    # 清理多余的空格
    import re
    choices_str = re.sub(r'\s+', ' ', choices_str)
    
    # 确保每个选项格式为 "编码, 标签"
    parts = choices_str.split('|')
    normalized_parts = []
    
    for part in parts:
        part = part.strip()
        if not part:
            continue
        
        # 如果没有逗号，尝试添加默认编码
        if ',' not in part:
            # 尝试提取数字前缀作为编码
            match = re.match(r'^(\d+)\s*(.+)', part)
            if match:
                code, label = match.groups()
                part = f"{code}, {label.strip()}"
            else:
                # 使用索引作为编码
                idx = len(normalized_parts)
                part = f"{idx}, {part}"
        
        normalized_parts.append(part)
    
    return ' | '.join(normalized_parts)


def _is_identifier(var_name):
    """判断是否为隐私标识字段"""
    identifier_keywords = [
        'name', '姓名', 'first', 'last', 'lastname', 'firstname',
        'address', '地址', 'phone', '电话', 'tel', 'mobile',
        'email', '邮件', 'id_card', '身份证', 'ssn', 'social'
    ]
    return any(k in var_name.lower() for k in identifier_keywords)


def generate_csv(fields, output_path=None):
    """生成 CSV 内容"""
    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=REDCAP_HEADERS)
    writer.writeheader()
    
    for field in fields:
        writer.writerow(field)
    
    csv_content = output.getvalue()
    
    if output_path:
        with open(output_path, 'w', encoding='utf-8-sig', newline='') as f:
            f.write(csv_content)
        return output_path
    
    return csv_content


def main():
    parser = argparse.ArgumentParser(description='处理上传的 Word 文档生成 REDCap 数据字典')
    parser.add_argument('file_path', help='Word 文件路径')
    parser.add_argument('-o', '--output', help='输出 CSV 文件路径')
    
    args = parser.parse_args()
    
    file_path = Path(args.file_path)
    
    if not file_path.exists():
        print(f"错误: 文件不存在 - {args.file_path}")
        sys.exit(1)
    
    print(f"正在解析文件: {file_path}")
    
    try:
        fields = parse_docx_smarter(str(file_path))
        print(f"解析完成，共发现 {len(fields)} 个字段")
        
        output_path = args.output or str(file_path.with_suffix('.csv'))
        result_path = generate_csv(fields, output_path)
        
        print(f"数据字典已生成: {result_path}")
        
        # 输出 CSV 内容预览
        csv_content = generate_csv(fields)
        print("\n=== CSV 内容预览 ===")
        print(csv_content[:2000])  # 显示前2000字符
        
    except Exception as e:
        print(f"处理失败: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()