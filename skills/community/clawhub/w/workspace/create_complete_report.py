#!/usr/bin/env python3
"""
Create complete BSc report with all 7 chapters and proper formatting
"""

import os
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE

def read_text_file(filename):
    """Read text file with proper encoding"""
    try:
        with open(filename, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        print(f"❌ Error reading {filename}: {e}")
        return None

def apply_formatting(doc):
    """Apply university formatting guidelines"""
    print("🎨 Applying formatting guidelines...")
    
    # Set document properties
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)  # ~2.54cm
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Create or modify styles
    styles = doc.styles
    
    # Normal style (body text)
    if 'Normal' not in styles:
        style = styles.add_style('Normal', WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = styles['Normal']
    
    font = style.font
    font.name = 'B Nazanin'
    font.size = Pt(14)
    
    # Heading 1 style
    if 'Heading 1' not in styles:
        style = styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = styles['Heading 1']
    
    font = style.font
    font.name = 'B Nazanin'
    font.size = Pt(16)
    font.bold = True
    
    # Heading 2 style
    if 'Heading 2' not in styles:
        style = styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = styles['Heading 2']
    
    font = style.font
    font.name = 'B Nazanin'
    font.size = Pt(15)
    font.bold = True
    
    # Apply to all paragraphs
    for para in doc.paragraphs:
        # Justify all text paragraphs
        if para.text.strip() and not para.text.startswith('فصل'):
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Set line spacing
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        # Apply appropriate style
        if para.text.startswith('فصل'):
            para.style = doc.styles['Heading 1']
        elif any(marker in para.text for marker in ['۱.', '۲.', '۳.', '۴.', '۵.', '۶.', '۷.']):
            para.style = doc.styles['Heading 2']
        else:
            para.style = doc.styles['Normal']
    
    return doc

def create_complete_report():
    """Create complete report with all 7 chapters"""
    
    print("📝 CREATING COMPLETE BSc PROJECT REPORT")
    print("="*50)
    
    # Step 1: Load ORIGINAL document (chapters 1-3)
    print("\n📄 Step 1: Loading ORIGINAL document (chapters 1-3)...")
    if not os.path.exists("BSc-project-report-ORIGINAL.docx"):
        print("❌ ORIGINAL document not found!")
        return False
    
    try:
        doc = Document("BSc-project-report-ORIGINAL.docx")
        print(f"✅ Loaded ORIGINAL: {len(doc.paragraphs)} paragraphs")
    except Exception as e:
        print(f"❌ Error loading ORIGINAL: {e}")
        return False
    
    # Step 2: Add chapters 4-7 from text files
    print("\n📄 Step 2: Adding chapters 4-7 from text files...")
    
    chapters_to_add = [
        (4, "فصل ۴: ساخت کلاستر Kubernetes", "chapter4_clean.txt"),
        (5, "فصل ۵: پیاده‌سازی لایه سرویس‌های در سطح کاربر", "chapter5_clean.txt"),
        (6, "فصل ۶: پایش، سخت‌سازی امنیتی، آزمون‌ها و تحویل نهایی", "chapter6_clean.txt"),
        (7, "فصل ۷: خلصه، نتیجه‌گیری و کارهای آینده", "chapter7_clean.txt")
    ]
    
    for chap_num, title, filename in chapters_to_add:
        print(f"  📖 Adding {title}...")
        
        # Add page break before new chapter
        doc.add_page_break()
        
        # Add chapter title
        title_para = doc.add_paragraph(title)
        title_para.style = 'Heading 1'
        
        # Read and add chapter content
        content = read_text_file(filename)
        if content:
            # Split content into paragraphs
            lines = content.strip().split('\n')
            for line in lines:
                if line.strip():  # Skip empty lines
                    para = doc.add_paragraph(line.strip())
                    para.style = 'Normal'
            
            print(f"    ✅ Added {len(lines)} paragraphs")
        else:
            print(f"    ⚠️  No content for chapter {chap_num}")
    
    # Step 3: Apply formatting
    print("\n🎨 Step 3: Applying university formatting...")
    doc = apply_formatting(doc)
    
    # Step 4: Save final document
    print("\n💾 Step 4: Saving final document...")
    output_file = "BSc-project-report-COMPLETE-FORMATTED.docx"
    doc.save(output_file)
    
    # Verification
    print("\n🔍 Step 5: Verification...")
    final_doc = Document(output_file)
    
    # Count chapters
    chapter_headings = []
    for para in final_doc.paragraphs:
        if para.text.startswith('فصل'):
            chapter_headings.append(para.text[:50])
    
    print(f"✅ Final document saved: {output_file}")
    print(f"📏 File size: {os.path.getsize(output_file):,} bytes")
    print(f"📄 Total paragraphs: {len(final_doc.paragraphs)}")
    print(f"📚 Chapters found: {len(chapter_headings)}")
    
    print("\n📋 Chapter headings:")
    for i, heading in enumerate(chapter_headings[:10]):
        print(f"  {i+1}. {heading}")
    
    if len(chapter_headings) >= 7:
        print("\n🎉 SUCCESS: Complete report with all 7 chapters created!")
        print("\n📋 FORMATTING APPLIED:")
        print("  • Font: B Nazanin")
        print("  • Body text: 14pt")
        print("  • Headings: 16pt")
        print("  • Text: Justified")
        print("  • Line spacing: 1.5")
        print("  • Margins: Standard academic")
    else:
        print(f"\n⚠️  WARNING: Only {len(chapter_headings)} chapters found")
    
    return True

def main():
    print("🔧 BSc Project Report Creator")
    print("="*50)
    
    # Check required files
    print("\n📋 Checking required files:")
    
    required = [
        "BSc-project-report-ORIGINAL.docx",
        "chapter4_clean.txt",
        "chapter5_clean.txt", 
        "chapter6_clean.txt",
        "chapter7_clean.txt"
    ]
    
    all_ok = True
    for file in required:
        if os.path.exists(file):
            size = os.path.getsize(file)
            print(f"✅ {file}: {size:,} bytes")
        else:
            print(f"❌ {file}: MISSING")
            all_ok = False
    
    if not all_ok:
        print("\n❌ Missing required files. Cannot proceed.")
        return
    
    print("\n" + "="*50)
    print("🚀 Starting report creation...")
    print("="*50)
    
    success = create_complete_report()
    
    if success:
        print("\n" + "="*50)
        print("✅ COMPLETION INSTRUCTIONS:")
        print("="*50)
        print("\n1. Open 'BSc-project-report-COMPLETE-FORMATTED.docx' in Word")
        print("2. Verify all 7 chapters are present")
        print("3. Check formatting matches guidelines")
        print("4. Update Table of Contents (References → Update Table)")
        print("5. Save final version")
        print("\n⏱️  Estimated time: 5-10 minutes")
    else:
        print("\n❌ Failed to create complete report")

if __name__ == "__main__":
    main()