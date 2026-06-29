#!/usr/bin/env python3

# ============================================================
# ⚠️ SAMPLE DATA NOTICE / 示例数据声明
# This template contains EXAMPLE / ILLUSTRATIVE data values.
# Users must configure their own data sources (yfinance, Guosen API, etc.)
# to generate plans with real market data.
# 本模板包含的数据为示例/演示用途。
# 用户需自行配置数据源（yfinance、国信API等）以获取真实市场数据。
# ============================================================

"""
₿ Bitcoin/Crypto Trading Plan Generator v7.0 FINAL / 比特币/加密货币交易方案生成器
Ai-Wife Global Investing / 媳妇智投全球通用版
Bilingual (CN/EN) Edition / 中英双语版

Usage / 用法:
  python3 template_crypto_v7.py "BTC/USD"
  python3 template_crypto_v7.py "ETH/USD"
  python3 template_crypto_v7.py "SOL/USD"
"""

import datetime
from docx import Document
from language_detection import detect_input_lang, match_market_lang, native_name, BilingualText
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

EN_SECTIONS = {
    'title': 'Bitcoin / Crypto',
    'header': 'Bilingual (CN/EN) Edition',
}

def create_btc_plan(symbol="BTC/USD", date_str=None, user_input=""):
    # ── 🌍 大叔亲授·审题+质检机制 ──
    _input_lang = detect_input_lang(user_input)
    _market_code = "us"  # Crypto is global, default to English market
    _dest_lang = match_market_lang(_market_code)
    _bt = BilingualText(input_lang=_input_lang, dest_lang=_dest_lang)

    """Generate bilingual BTC/Crypto trading plan / 生成中英双语比特币/加密货币交易方案"""
    if date_str is None:
        date_str = datetime.datetime.now().strftime("%Y-%m-%d")
    
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    
    # ==================== TITLE / 标题 ====================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"₿ {symbol} Crypto Reference Framework / 加密货币参考框架")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0xF7, 0x93, 0x1A)  # Bitcoin orange
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Version: v7.0 FINAL | Date / 制作日期: {date_str}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("⚠️ 教育用途 | Educational Use Only")
    run.font.size = Pt(8)
    run.italic = True
    
    doc.add_paragraph()
    
    # ==================== HEADER / 方案头部 ====================
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Light Grid Accent 1'
    cells = tbl.rows[0].cells
    cells[0].paragraphs[0].add_run(f'Item / {_bt.item_label}').bold = True
    cells[1].paragraphs[0].add_run('Content / 内容').bold = True
    
    rows_data = [
        ('Asset / 品种', symbol),
        ('Category / 分类', 'Cryptocurrency / 加密货币'),
        ('Market / 交易市场', 'Global Crypto Exchanges / 全球加密货币交易所'),
        ('Trading Hours / 交易时间', '24/7 (365 days) / 全天候无休'),
        ('Min Unit / 最小单位', '0.00001 BTC (1 satoshi = 0.00000001 BTC)'),
        ('Market Cap / 总市值', '~$2T+ (crypto total) / 加密货币总市值'),
        ('24H Volume / 24小时成交量', '~$50-100B (BTC spot) / 比特币现货'),
        ('Date / 日期', date_str),
        ('Plan Type / 方案类型', 'Swing / Intraday / Position / 波段/日内/趋势'),
    ]
    for item, val in rows_data:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(item).bold = True
        row.cells[1].paragraphs[0].add_run(val)
    
    # ==================== 一、Market Review / 行情回顾 ====================
    doc.add_heading('一、Market Review / 行情回顾', level=2)
    
    doc.add_heading('1.1 Close Data / 收盘数据', level=3)
    tbl = doc.add_table(rows=1, cols=8)
    tbl.style = 'Light Grid Accent 1'
    hdr = ['Asset / 品种', 'Last / 最新价', '24H Change', '24H % / 涨跌幅', '24H High / 最高', '24H Low / 最低', '24H Volume', 'Volatility / 波动率']
    for i, h in enumerate(hdr):
        tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    row = tbl.add_row()
    for i in range(8):
        row.cells[i].paragraphs[0].add_run('—')
    
    doc.add_heading('1.2 Recent Trend / 近期走势', level=3)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Light Grid Accent 1'
    hdr = ['Period / 时间段', 'Price Range / 价格区间', 'Change / 涨跌幅', 'Character / 走势特征']
    for i, h in enumerate(hdr):
        tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    periods = ['24H / 近24小时', '7-Day / 近7日', '30-Day / 近30日', '90-Day / 近90日']
    for pp in periods:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(pp)
        for j in range(1, 4):
            row.cells[j].paragraphs[0].add_run('—')
    
    doc.add_heading('Summary / 走势总结', level=4)
    p = doc.add_paragraph('—')
    
    # ==================== 二、Market Correlation / 市场联动分析 ====================
    doc.add_heading('二、Market Correlation / 市场联动分析', level=2)
    
    doc.add_heading('2.1 Key Market Indicators / 关键市场指标', level=3)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Light Grid Accent 1'
    hdr = ['Indicator / 指标', 'Value / 数值', '24H Change / 变化', 'Note / 说明']
    for i, h in enumerate(hdr):
        tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    indicators = [
        ('Bitcoin Dominance / BTC市占率', '—', '—', 'Altcoin season indicator / 山寨币季节指标'),
        ('Crypto Total Market Cap / 总市值', '—', '—', 'Market-wide health / 整体市场健康度'),
        ('Fear & Greed Index / 恐惧贪婪指数', '—', '—', '0-100 sentiment / 市场情绪0-100'),
        ('Open Interest (OI) / 合约总持仓', '—', '—', 'Leverage gauge / 杠杆水平'),
        ('Funding Rate / 资金费率', '—', '—', 'Long/short cost / 多空持仓成本'),
    ]
    for ind, val, chg, note in indicators:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(ind).bold = True
        row.cells[1].paragraphs[0].add_run(val)
        row.cells[2].paragraphs[0].add_run(chg)
        row.cells[3].paragraphs[0].add_run(note)
    
    doc.add_heading('2.2 Stock-to-Flow & On-Chain / 链上数据分析', level=3)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Light Grid Accent 1'
    hdr = ['Metric / 指标', 'Value / 数值', 'Signal / 信号']
    for i, h in enumerate(hdr):
        tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    for metric in ['S2F Model / 库存流量模型', 'MVRV Ratio / MVRV比率', 'Puell Multiple / Puell倍数', 'Active Addresses / 活跃地址数', 'Hash Rate / 算力', 'Exchange Reserve / 交易所储备']:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(metric).bold = True
        row.cells[1].paragraphs[0].add_run('—')
        row.cells[2].paragraphs[0].add_run('—')
    
    doc.add_heading('2.3 Correlation with Traditional Markets / 与传统市场关联', level=3)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Light Grid Accent 1'
    hdr = ['Asset / 品种', 'Price / 最新价', '24H Change', 'BTC Correlation / BTC相关性']
    for i, h in enumerate(hdr):
        tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    related = [('ETH / 以太坊', '—', '—', '—'), ('SOL / 索拉纳', '—', '—', '—'), 
               ('S&P 500 / 标普500', '—', '—', '—'), ('Gold / 黄金', '—', '—', '—'),
               ('DXY / 美元指数', '—', '—', '—'), ('US 10Y Yield / 美债收益率', '—', '—', '—')]
    for r_name, price, chg, corr in related:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(r_name).bold = True
        row.cells[1].paragraphs[0].add_run(price)
        row.cells[2].paragraphs[0].add_run(chg)
        row.cells[3].paragraphs[0].add_run(corr)
    
    # ==================== 三、Macro Analysis / 宏观影响分析 ====================
    doc.add_heading('三、Macro & Regulatory / 宏观与监管分析', level=2)
    
    doc.add_heading('3.1 Bearish Factors / 利空因素 🔴', level=3)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Light Grid Accent 1'
    hdr = ['Factor / 因素', 'Detail / 详情', 'Impact / 影响程度', 'Duration / 持续性']
    for i, h in enumerate(hdr):
        tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    for _ in range(3):
        row = tbl.add_row()
        for i in range(4):
            row.cells[i].paragraphs[0].add_run('—')
    
    doc.add_heading('3.2 Bullish Factors / 利多因素 🟢', level=3)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Light Grid Accent 1'
    hdr = ['Factor / 因素', 'Detail / 详情', 'Impact / 影响程度', 'Duration / 持续性']
    for i, h in enumerate(hdr):
        tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    for _ in range(3):
        row = tbl.add_row()
        for i in range(4):
            row.cells[i].paragraphs[0].add_run('—')
    
    doc.add_heading('3.3 Regulatory Landscape / 监管环境', level=3)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Light Grid Accent 1'
    hdr = ['Region / 地区', 'Status / 状态', 'Market Impact / 市场影响']
    for i, h in enumerate(hdr):
        tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    for region in ['US / 美国', 'EU (MiCA) / 欧盟', 'Asia / 亚洲', 'Middle East / 中东']:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(region).bold = True
        row.cells[1].paragraphs[0].add_run('—')
        row.cells[2].paragraphs[0].add_run('—')
    
    doc.add_heading('3.4 Macro Conclusion / 宏观结论', level=3)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Light Grid Accent 1'
    cells = tbl.rows[0].cells
    cells[0].paragraphs[0].add_run('Dimension / 维度').bold = True
    cells[1].paragraphs[0].add_run('Judgment / 判断').bold = True
    for dim in ['Bull vs Bear / 多空对比', 'Short-term / 短期', 'Mid-term / 中期', 'Halving Cycle / 减半周期', 'Key Narrative / 核心叙事']:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(dim).bold = True
        row.cells[1].paragraphs[0].add_run('—')
    
    # ==================== 四、Technical Analysis / 技术分析 ====================
    doc.add_heading('四、Technical Analysis / 技术分析 📈', level=2)
    
    doc.add_heading('4.1 Key Levels / 关键价位', level=3)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Light Grid Accent 1'
    hdr = ['Type / 类型', 'Level / 价位', 'Note / 说明']
    for i, h in enumerate(hdr):
        tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    levels = ['ATH / 历史高点', 'R3 / 强压力3', 'R2 / 中压力2', 'R1 / 弱压力1',
              'Current / 当前价',
              'S1 / 弱支撑1', 'S2 / 中支撑2', 'S3 / 强支撑3', 'Key Support / 关键支撑(MA200)']
    for level in levels:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(level).bold = True
        row.cells[1].paragraphs[0].add_run('—')
        row.cells[2].paragraphs[0].add_run('—')
    
    doc.add_heading('4.2 Technical Indicators / 技术指标', level=3)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Light Grid Accent 1'
    hdr = ['Indicator / 指标', 'Value / 数值', 'Signal / 信号']
    for i, h in enumerate(hdr):
        tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    indicators = [
        ('MA System (50/100/200) / 均线系统', '—', '—'),
        ('MACD (12,26,9)', '—', '—'),
        ('RSI (14) / 相对强弱指数', '—', '—'),
        ('Bollinger Bands / 布林带 (20,2)', '—', '—'),
        ('OBV / 能量潮', '—', '—'),
        ('Ichimoku Cloud / 一目均衡表', '—', '—'),
        ('Volume Profile / 成交量分布', '—', '—'),
    ]
    for ind, val, sig in indicators:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(ind).bold = True
        row.cells[1].paragraphs[0].add_run(val)
        row.cells[2].paragraphs[0].add_run(sig)
    
    doc.add_heading('4.3 Technical Conclusion / 技术结论', level=3)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Light Grid Accent 1'
    cells = tbl.rows[0].cells
    cells[0].paragraphs[0].add_run('Dimension / 维度').bold = True
    cells[1].paragraphs[0].add_run('Judgment / 判断').bold = True
    for dim in ['Trend Direction / 趋势方向', 'Trend Strength / 趋势强度', 'Key Levels / 关键位置', 'Divergence / 背离信号', 'Volume Analysis / 量能分析']:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(dim).bold = True
        row.cells[1].paragraphs[0].add_run('—')
    
    # ==================== 五、On-Chain & Derivatives / 链上与衍生品 ====================
    doc.add_heading('六、On-Chain & Derivatives / 链上与衍生品分析 🔗', level=2)
    
    doc.add_heading('5.1 On-Chain Metrics / 链上数据', level=3)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Light Grid Accent 1'
    hdr = ['Metric / 指标', 'Current / 当前', 'Trend / 趋势', 'Interpretation / 解读']
    for i, h in enumerate(hdr):
        tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    for metric in ['Active Addresses / 活跃地址', 'New Addresses / 新地址', 'Transaction Count / 交易数', 
                   'Avg Tx Fee / 平均手续费', 'Exchange Net Flow / 交易所净流入', 'Miners to Exchange / 矿工流向交易所',
                   'Long-term Holders / 长期持有者', 'Short-term Holders / 短期持有者']:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(metric).bold = True
        row.cells[1].paragraphs[0].add_run('—')
        row.cells[2].paragraphs[0].add_run('—')
        row.cells[3].paragraphs[0].add_run('—')
    
    doc.add_heading('5.2 Derivatives Market / 衍生品市场', level=3)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Light Grid Accent 1'
    hdr = ['Metric / 指标', 'Current / 当前', 'Change / 变化', 'Signal / 信号']
    for i, h in enumerate(hdr):
        tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    for metric in ['Open Interest (OI) / 总持仓', 'Funding Rate / 资金费率', 'Long/Short Ratio / 多空比',
                   'Liquidation Map / 清算地图', 'Options Open Interest / 期权持仓', 'Put/Call Ratio / 看跌看涨比']:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(metric).bold = True
        row.cells[1].paragraphs[0].add_run('—')
        row.cells[2].paragraphs[0].add_run('—')
        row.cells[3].paragraphs[0].add_run('—')
    
    doc.add_heading('5.3 Whale & Institution Flow / 鲸鱼与机构动向', level=3)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Light Grid Accent 1'
    hdr = ['Category / 类别', 'Activity / 活动', 'Interpretation / 解读']
    for i, h in enumerate(hdr):
        tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    for cat in ['ETF Flow / ETF资金流', 'Whale Wallet / 鲸鱼钱包(>1K BTC)', 'Institution OTC / 机构OTC交易',
                'Mining Pool / 矿池钱包', 'Exchange Hot Wallet / 交易所热钱包']:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(cat).bold = True
        row.cells[1].paragraphs[0].add_run('—')
        row.cells[2].paragraphs[0].add_run('—')
    
    # ==================== 六、Market Sentiment / 市场情绪 ====================
    doc.add_heading('六、Market Sentiment / 市场情绪 📊', level=2)
    
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Light Grid Accent 1'
    hdr = ['Dimension / 维度', 'Value / 数值/状态', 'Interpretation / 解读']
    for i, h in enumerate(hdr):
        tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    for dim in ['Fear & Greed / 恐惧贪婪指数', 'Social Volume / 社交热度', 'Google Trends / 搜索热度',
                'Retail Sentiment / 散户情绪', 'Institutional Sentiment / 机构情绪', 'Media Coverage / 媒体报道']:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(dim).bold = True
        row.cells[1].paragraphs[0].add_run('—')
        row.cells[2].paragraphs[0].add_run('—')
    
    # ==================== 七、Trading Plan / 操作方案 ====================
    doc.add_heading('七、Trading Plan / 操作方案 🎯', level=2)
    
    p = doc.add_paragraph()
    run = p.add_run('₿ Crypto: 24/7 trading, T+0 both-way, high leverage (up to 125x on some exchanges) / 全天候交易，T+0双向，高杠杆')
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xF7, 0x93, 0x1A)
    
    doc.add_heading('7.1 Overall Strategy / 总体策略', level=3)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Light Grid Accent 1'
    cells = tbl.rows[0].cells
    cells[0].paragraphs[0].add_run(f'Item / {_bt.item_label}').bold = True
    cells[1].paragraphs[0].add_run('Content / 内容').bold = True
    for item in ['Core Range / 核心区间', 'Direction / 主要方向', 'Direction Judgment / 方向判断', 
                 'Strategy / 操作思路', 'Position Size / 仓位建议', 'Leverage / 杠杆倍数']:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(item).bold = True
        row.cells[1].paragraphs[0].add_run('—')
    
    doc.add_heading('7.2 ⬆️ Long Plan / 做多方案', level=3)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Light Grid Accent 1'
    cells = tbl.rows[0].cells
    cells[0].paragraphs[0].add_run(f'Item / {_bt.item_label}').bold = True
    cells[1].paragraphs[0].add_run('Content / 内容').bold = True
    for item in ['Scenario / 适用场景', 'Entry Condition / 入场条件', 'Entry Price / 入场价位', 'Stop Loss / 止损价位',
                 'Target 1 (T1) / 第一目标', 'Target 2 (T2) / 第二目标', 'Target 3 (T3) / 第三目标',
                 'Position Size / 仓位控制', 'Leverage / 杠杆', 'Risk/Reward / 盈亏比']:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(item).bold = True
        row.cells[1].paragraphs[0].add_run('—')
    
    doc.add_heading('7.3 ⬇️ Short Plan / 做空方案', level=3)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Light Grid Accent 1'
    cells = tbl.rows[0].cells
    cells[0].paragraphs[0].add_run(f'Item / {_bt.item_label}').bold = True
    cells[1].paragraphs[0].add_run('Content / 内容').bold = True
    for item in ['Scenario / 适用场景', 'Entry Condition / 入场条件', 'Entry Price / 入场价位', 'Stop Loss / 止损价位',
                 'Target 1 (T1) / 第一目标', 'Target 2 (T2) / 第二目标', 'Target 3 (T3) / 第三目标',
                 'Position Size / 仓位控制', 'Leverage / 杠杆', 'Risk/Reward / 盈亏比']:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(item).bold = True
        row.cells[1].paragraphs[0].add_run('—')
    
    doc.add_heading('7.4 Strategy by Market Condition / 不同市场阶段策略', level=3)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Light Grid Accent 1'
    hdr = ['Condition / 市场状态', 'Strategy / 策略', 'Risk Level / 风险等级']
    for i, h in enumerate(hdr):
        tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    for cond in ['Bull Trend / 上涨趋势', 'Bear Trend / 下跌趋势', 'Sideways / 震荡盘整',
                 'High Volatility / 高波动（>5%日波幅）', 'Post-Halving / 减半后', 'BTC Dominance Rising / BTC市占率上升']:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(cond).bold = True
        row.cells[1].paragraphs[0].add_run('—')
        row.cells[2].paragraphs[0].add_run('—')
    
    # ==================== 八、Risk Management / 风险管理 ====================
    doc.add_heading('八、Risk Management / 风险管理 ⚠️', level=2)
    
    doc.add_heading('8.1 Crypto-Specific Risks / 加密货币特殊风险', level=3)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Light Grid Accent 1'
    hdr = ['Risk Type / 风险类型', 'Description / 说明', 'Mitigation / 应对措施']
    for i, h in enumerate(hdr):
        tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    risks = [
        ('Flash Crash / 闪崩', 'Rapid 10-30% drop in hours', 'Tight stop-losses, reduce leverage / 严格止损'),
        ('Exchange Risk / 交易所风险', 'FTX-style collapse / hack', 'Self-custody / Use cold wallet / 自托管'),
        ('Regulatory Shock / 监管冲击', 'Sudden ban / restriction', 'Diversify geographically / 地域分散'),
        ('Liquidation Cascade / 连环爆仓', 'Leverage forced liquidation chain', 'Low leverage (≤3x spot) / 低杠杆'),
        ('Smart Contract Risk / 合约风险', 'DeFi bridge / exploit', 'Stick to BTC/ETH / 只做主流币'),
        ('Fork / Hard Fork / 硬分叉', 'Chain split risk', 'Monitor community consensus'),
        ('Whale Manipulation / 鲸鱼操纵', 'Large orders move market', 'Limit orders, avoid FOMO / 限价单'),
    ]
    for risk, desc, action in risks:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(f'⚠️ {risk}').bold = True
        row.cells[1].paragraphs[0].add_run(desc)
        row.cells[2].paragraphs[0].add_run(action)
    
    doc.add_heading('8.2 Money Management / 资金管理', level=3)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Light Grid Accent 1'
    cells = tbl.rows[0].cells
    cells[0].paragraphs[0].add_run(f'Item / {_bt.item_label}').bold = True
    cells[1].paragraphs[0].add_run('Suggestion / 建议').bold = True
    for item in ['Total Crypto Allocation / 总加密货币仓位', 'BTC:Altcoin Ratio / BTC:山寨币比例',
                 'Max Risk per Trade / 单笔最大风险', 'Suggested Leverage / 建议杠杆', 
                 'Max Drawdown / 最大回撤容忍', 'Daily Loss Limit / 日亏损上限',
                 'Stablecoin Reserve / 稳定币储备']:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(item).bold = True
        row.cells[1].paragraphs[0].add_run('—')
    
    # ==================== 九、Weekly Events / 本周关注 ====================
    doc.add_heading('九、Key Events & Catalysts / 本周关键事件 📅', level=2)
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = 'Light Grid Accent 1'
    hdr = ['Time / 时间', 'Event / 事件', 'Category / 类别', 'Importance / 重要程度', 'Expected Impact / 预期影响']
    for i, h in enumerate(hdr):
        tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    for _ in range(6):
        row = tbl.add_row()
        for i in range(5):
            row.cells[i].paragraphs[0].add_run('—')
    
    p = doc.add_paragraph()
    run = p.add_run('💡 Key: FOMC meetings / CPI data / BTC ETF flows / Halving countdown / ETF approval news')
    run.font.size = Pt(8)
    run.italic = True
    
    # ==================== 十、Summary / 方案总结 ====================
    doc.add_heading('十一、Summary / 方案总结 ✅', level=2)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Light Grid Accent 1'
    cells = tbl.rows[0].cells
    cells[0].paragraphs[0].add_run('Dimension / 维度').bold = True
    cells[1].paragraphs[0].add_run('Conclusion / 结论').bold = True
    for dim in ['Macro View / 大势判断', 'Core Narrative / 核心叙事', 'Technical View / 技术面判断',
                 'On-Chain View / 链上数据判断', 'Trading Strategy / 操作思路', 
                 'Position / 仓位建议', 'Stop-Loss Discipline / 止损纪律']:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(dim).bold = True
        row.cells[1].paragraphs[0].add_run('—')
    
    # ==================== One-Liner / 一句话总结 ====================
    doc.add_heading('📌 One-Sentence Summary / 一句话总结', level=3)
    p = doc.add_paragraph('—')
    run = p.runs[0]
    run.bold = True
    
    # ==================== Appendix / 附录 ====================
    doc.add_heading('Appendix: Crypto Basics / 附录：加密货币基础须知', level=2)
    
    doc.add_heading('Major Cryptocurrencies / 主要加密货币', level=3)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Light Grid Accent 1'
    hdr = ['Asset / 品种', 'Ticker / 代码', 'Category / 类别', 'Market Cap Rank / 市值排名']
    for i, h in enumerate(hdr):
        tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    coins = [
        ('Bitcoin / 比特币', 'BTC', 'Store of Value / 数字黄金', '#1'),
        ('Ethereum / 以太坊', 'ETH', 'Smart Contract / 智能合约', '#2'),
        ('Solana / 索拉纳', 'SOL', 'High-Speed L1 / 高性能公链', 'Top 5'),
        ('USDT/USDC / 稳定币', 'USDT/USDC', 'Stablecoin / 稳定币', 'Top 3'),
    ]
    for coin, ticker, cat, rank in coins:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(coin).bold = True
        row.cells[1].paragraphs[0].add_run(ticker)
        row.cells[2].paragraphs[0].add_run(cat)
        row.cells[3].paragraphs[0].add_run(rank)
    
    doc.add_heading('Trading Hours / 交易时间', level=3)
    p = doc.add_paragraph('⏰ Crypto: 24/7/365 — no holidays, no breaks / 加密货币全天候无休')
    
    doc.add_heading('Risk Disclaimer / 风险声明', level=3)
    p = doc.add_paragraph()
    run = p.add_run(
        'This plan is for reference and educational purposes only, not investment advice. '
        'Cryptocurrency trading carries extreme risk; prices can drop 50%+ in a single day. '
        'High leverage can lead to total loss of capital. Only invest what you can afford to lose. '
        '本方案仅供参考学习，不构成投资建议。加密货币交易风险极高，价格可能单日暴跌50%以上，请根据自身情况谨慎决策。'
    )
    run.font.size = Pt(8)
    run.italic = True
    
    # ==================== FOOTER / 尾部 ====================
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('—— 媳妇智投全球通用版 | Ai-Wife Global Investing ——')
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Date / 日期：{date_str} | Version / 版本：v7.0 FINAL')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('⚠️ AIGC Content — For Reference Only / AIGC生成内容，仅供参考')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    
    return doc


def main():
    import sys
    
    symbol = "BTC/USD"
    if len(sys.argv) >= 2:
        symbol = sys.argv[1]
    
    date_str = datetime.datetime.now().strftime("%Y-%m-%d")
    
    doc = create_btc_plan(symbol, date_str)
    
    safe_symbol = symbol.replace('/', '_')
    filename = f"Crypto_{safe_symbol}_{date_str}_v7.0.docx"
    
    # ── 📄 保存生成的.docx参考框架文件 ──
    # ⚠️ [透明声明] 用户请求生成方案后，框架文档将保存到本地并发送给用户。
    doc.save(filename)
    print(f"✅ Bilingual Crypto Plan Generated / 双语加密货币方案生成完成！")
    print(f"   📄 {filename}")
    print(f"   ₿ Asset / 品种: {symbol}")
    print(f"   📅 Date / 日期: {date_str}")
    print(f"   📊 Chapters / 章节: 10 (Bilingual CN/EN / 中英双语)")
    print(f"   🏷️  Version / 版本: v7.0 FINAL")


if __name__ == "__main__":
    main()
