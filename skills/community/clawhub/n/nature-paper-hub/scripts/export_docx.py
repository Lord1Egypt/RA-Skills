#!/usr/bin/env python3
"""
export_docx.py — Generate a Nature-portfolio-style Word manuscript from a JSON draft.

Usage:
    python3 export_docx.py --input ~/Downloads/nature-paper-draft.json \
                           --output ~/Downloads/nature-paper-20260517.docx
"""

import argparse
import json
import os
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_font(run, name="Times New Roman", size_pt=12, bold=False, italic=False,
             color=None):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_paragraph_format(para, space_before=0, space_after=0,
                          line_spacing=None, first_line_indent=None,
                          left_indent=None, alignment=None):
    pf = para.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if line_spacing is not None:
        from docx.shared import Pt as _Pt
        pf.line_spacing = _Pt(line_spacing)
    if first_line_indent is not None:
        pf.first_line_indent = Pt(first_line_indent)
    if left_indent is not None:
        pf.left_indent = Pt(left_indent)
    if alignment is not None:
        para.alignment = alignment


def add_page_numbers(doc):
    """Insert a bottom-centre page number field into the default footer."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.clear()

    run = para.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    run._r.append(fld)

    run2 = para.add_run()
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    run2._r.append(instrText)

    run3 = para.add_run()
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run3._r.append(fld2)

    for run in (run, run2, run3):
        set_font(run, size_pt=10)


def add_section_heading(doc, text, level="section"):
    """Add a bold section heading (or italic-bold subsection heading)."""
    para = doc.add_paragraph()
    set_paragraph_format(para, space_before=12, space_after=4, line_spacing=24)
    run = para.add_run(text)
    if level == "section":
        set_font(run, size_pt=12, bold=True)
    else:  # subsection
        set_font(run, size_pt=11, bold=True, italic=True)
    return para


def add_body_paragraph(doc, text, first=False):
    """Add a double-spaced body paragraph."""
    para = doc.add_paragraph()
    indent = 0 if first else 0           # Nature style: no first-line indent
    set_paragraph_format(para, space_before=0, space_after=0,
                          line_spacing=24, first_line_indent=indent)
    run = para.add_run(text)
    set_font(run, size_pt=12)
    return para


# ---------------------------------------------------------------------------
# Main builder
# ---------------------------------------------------------------------------

def build_docx(data: dict, output_path: Path):
    doc = Document()

    # ── Page setup ──────────────────────────────────────────────────────────
    section = doc.sections[0]
    margin = Cm(2.54)
    section.top_margin = margin
    section.bottom_margin = margin
    section.left_margin = margin
    section.right_margin = margin

    # ── Page numbers ────────────────────────────────────────────────────────
    add_page_numbers(doc)

    # ── Header note (line-numbers reminder) ─────────────────────────────────
    note_para = doc.add_paragraph()
    set_paragraph_format(note_para, space_before=0, space_after=6, line_spacing=12)
    note_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_run = note_para.add_run(
        "[Note for submission: enable continuous line numbers in your journal portal "
        "before uploading this manuscript.]"
    )
    set_font(note_run, size_pt=9, italic=True, color=(128, 128, 128))

    # ── Journal tag ─────────────────────────────────────────────────────────
    journal = data.get("journal", "Nature Portfolio")
    j_para = doc.add_paragraph()
    set_paragraph_format(j_para, space_before=0, space_after=4, line_spacing=12)
    j_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    j_run = j_para.add_run(journal.upper())
    set_font(j_run, size_pt=9, bold=True, color=(80, 80, 80))

    # ── Title ────────────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    set_paragraph_format(title_para, space_before=12, space_after=8, line_spacing=24)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(data.get("title", "Manuscript Title"))
    set_font(title_run, size_pt=14, bold=True)

    # ── Authors ──────────────────────────────────────────────────────────────
    authors = data.get("authors", [])
    auth_para = doc.add_paragraph()
    set_paragraph_format(auth_para, space_before=0, space_after=4, line_spacing=24)
    auth_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    auth_run = auth_para.add_run(", ".join(authors))
    set_font(auth_run, size_pt=11)

    # ── Affiliations ─────────────────────────────────────────────────────────
    for i, aff in enumerate(data.get("affiliations", []), 1):
        aff_para = doc.add_paragraph()
        set_paragraph_format(aff_para, space_before=0, space_after=2, line_spacing=14)
        aff_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label = f"{i}. " if len(data.get("affiliations", [])) > 1 else ""
        aff_run = aff_para.add_run(f"{label}{aff}")
        set_font(aff_run, size_pt=9, italic=True)

    # ── Corresponding email ──────────────────────────────────────────────────
    email = data.get("corresponding_email", "")
    if email:
        em_para = doc.add_paragraph()
        set_paragraph_format(em_para, space_before=2, space_after=8, line_spacing=14)
        em_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        em_run = em_para.add_run(f"Correspondence: {email}")
        set_font(em_run, size_pt=9, italic=True)

    doc.add_paragraph()  # visual spacer

    # ── Abstract ─────────────────────────────────────────────────────────────
    abs_para = doc.add_paragraph()
    set_paragraph_format(abs_para, space_before=0, space_after=2, line_spacing=14)
    abs_label = abs_para.add_run("Abstract  ")
    set_font(abs_label, size_pt=12, bold=True)
    abs_body = abs_para.add_run(data.get("abstract", ""))
    set_font(abs_body, size_pt=12)

    # ── Keywords ─────────────────────────────────────────────────────────────
    keywords = data.get("keywords", [])
    if keywords:
        kw_para = doc.add_paragraph()
        set_paragraph_format(kw_para, space_before=4, space_after=8, line_spacing=14)
        kw_label = kw_para.add_run("Keywords: ")
        set_font(kw_label, size_pt=10, bold=True)
        kw_body = kw_para.add_run("; ".join(keywords))
        set_font(kw_body, size_pt=10)

    # ── Introduction ─────────────────────────────────────────────────────────
    add_section_heading(doc, "Introduction")
    intro = data.get("introduction", "")
    if intro:
        for i, para_text in enumerate(intro.split("\n\n")):
            para_text = para_text.strip()
            if para_text:
                add_body_paragraph(doc, para_text, first=(i == 0))

    # ── Results ──────────────────────────────────────────────────────────────
    add_section_heading(doc, "Results")
    for subsec in data.get("results", []):
        add_section_heading(doc, subsec.get("title", ""), level="subsection")
        content = subsec.get("content", "")
        for i, para_text in enumerate(content.split("\n\n")):
            para_text = para_text.strip()
            if para_text:
                add_body_paragraph(doc, para_text, first=(i == 0))

    # ── Discussion ────────────────────────────────────────────────────────────
    add_section_heading(doc, "Discussion")
    discussion = data.get("discussion", "")
    if discussion:
        for i, para_text in enumerate(discussion.split("\n\n")):
            para_text = para_text.strip()
            if para_text:
                add_body_paragraph(doc, para_text, first=(i == 0))

    # ── Methods ───────────────────────────────────────────────────────────────
    add_section_heading(doc, "Methods")
    for subsec in data.get("methods", []):
        add_section_heading(doc, subsec.get("title", ""), level="subsection")
        content = subsec.get("content", "")
        for i, para_text in enumerate(content.split("\n\n")):
            para_text = para_text.strip()
            if para_text:
                add_body_paragraph(doc, para_text, first=(i == 0))

    # ── Data Availability ────────────────────────────────────────────────────
    da = data.get("data_availability", "")
    if da:
        add_section_heading(doc, "Data Availability")
        add_body_paragraph(doc, da)

    # ── Author Contributions ─────────────────────────────────────────────────
    ac = data.get("author_contributions", "")
    if ac:
        add_section_heading(doc, "Author Contributions")
        add_body_paragraph(doc, ac)

    # ── Competing Interests ───────────────────────────────────────────────────
    ci = data.get("competing_interests", "")
    if ci:
        add_section_heading(doc, "Competing Interests")
        add_body_paragraph(doc, ci)

    # ── Acknowledgements ──────────────────────────────────────────────────────
    ack = data.get("acknowledgements", "")
    if ack:
        add_section_heading(doc, "Acknowledgements")
        add_body_paragraph(doc, ack)

    # ── References ────────────────────────────────────────────────────────────
    refs = data.get("references", [])
    if refs:
        add_section_heading(doc, "References")
        for ref_text in refs:
            ref_para = doc.add_paragraph()
            # Hanging indent: left_indent=18pt, first_line_indent=-18pt
            pf = ref_para.paragraph_format
            pf.left_indent = Pt(18)
            pf.first_line_indent = Pt(-18)
            pf.space_before = Pt(2)
            pf.space_after = Pt(2)
            pf.line_spacing = Pt(12)
            ref_run = ref_para.add_run(ref_text)
            set_font(ref_run, size_pt=10)

    # ── Figure Legends ────────────────────────────────────────────────────────
    fig_legends = data.get("figure_legends", [])
    if fig_legends:
        add_section_heading(doc, "Figure Legends")
        for legend in fig_legends:
            leg_para = doc.add_paragraph()
            set_paragraph_format(leg_para, space_before=4, space_after=4, line_spacing=14)

            # Bold "Figure X |" prefix  (split on first " " after the pipe)
            if "|" in legend:
                prefix, rest = legend.split("|", 1)
                prefix_run = leg_para.add_run(prefix.rstrip() + " |")
                set_font(prefix_run, size_pt=10, bold=True)
                body_run = leg_para.add_run(rest)
                set_font(body_run, size_pt=10)
            else:
                run = leg_para.add_run(legend)
                set_font(run, size_pt=10)

    # ── Save ──────────────────────────────────────────────────────────────────
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------

def parse_args():
    today = datetime.today().strftime("%Y%m%d")
    parser = argparse.ArgumentParser(
        description="Generate a Nature-portfolio-style Word manuscript from a JSON draft."
    )
    parser.add_argument(
        "--input",
        default=str(Path("~/Downloads/nature-paper-draft.json").expanduser()),
        help="Path to the input JSON file.",
    )
    parser.add_argument(
        "--output",
        default=str(Path(f"~/Downloads/nature-paper-{today}.docx").expanduser()),
        help="Path for the output .docx file.",
    )
    return parser.parse_args()


def main():
    args = parse_args()
    input_path = Path(args.input).expanduser().resolve()
    output_path = Path(args.output).expanduser().resolve()

    if not input_path.exists():
        print(f"❌ Input file not found: {input_path}", file=sys.stderr)
        sys.exit(1)

    with input_path.open("r", encoding="utf-8") as fh:
        data = json.load(fh)

    build_docx(data, output_path)
    print(f"✅ Export successful → {output_path}")


if __name__ == "__main__":
    main()
