#!/usr/bin/env python3
"""
Doc Weaver — transform Markdown into polished Word (.docx) and PDF documents.

Uses python-docx for native .docx generation with real formatting:
  - Multi-level heading numbering with Word heading styles
  - Clickable table of contents (TOC field)
  - Tables with borders, styled headers, and alternating row shading
  - Code blocks with monospace font + grey background
  - Blockquotes with indentation + left border
  - Headers (title) and footers (page X of Y + date)
  - Cover page support (when template has cover=True)

PDF output via pandoc + weasyprint (intermediate Markdown -> PDF).
Requires: python-docx, pandoc, weasyprint
"""

import re
import sys
import shutil
import subprocess
import tempfile
import importlib.util
from pathlib import Path
from datetime import datetime

# --- python-docx imports ---
try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# --- Template definitions ---
TEMPLATES = {
    "prd": {
        "name": "Product Requirements Document",
        "cover": True,
        "fonts": {"heading": "Arial", "body": "Calibri"},
        "colors": {
            "primary": "#1a73e8", "accent": "#e8f0fe",
            "heading_text": "#174ea6", "body_text": "#202124",
        },
    },
    "report": {
        "name": "Report",
        "cover": True,
        "fonts": {"heading": "Georgia", "body": "Calibri"},
        "colors": {
            "primary": "#333333", "accent": "#f5f5f5",
            "heading_text": "#1a1a1a", "body_text": "#333333",
        },
    },
    "academic": {
        "name": "Academic Paper",
        "cover": True,
        "fonts": {"heading": "Times New Roman", "body": "Times New Roman"},
        "colors": {
            "primary": "#000000", "accent": "#ffffff",
            "heading_text": "#000000", "body_text": "#000000",
        },
    },
    "manual": {
        "name": "User Manual",
        "cover": True,
        "fonts": {"heading": "Helvetica", "body": "Helvetica"},
        "colors": {
            "primary": "#005a9e", "accent": "#e6f2ff",
            "heading_text": "#003d6b", "body_text": "#333333",
        },
    },
    "contract": {
        "name": "Contract",
        "cover": True,
        "fonts": {"heading": "Times New Roman", "body": "Times New Roman"},
        "colors": {
            "primary": "#000000", "accent": "#ffffff",
            "heading_text": "#000000", "body_text": "#000000",
        },
    },
    "proposal": {
        "name": "Business Proposal",
        "cover": True,
        "fonts": {"heading": "Helvetica", "body": "Arial"},
        "colors": {
            "primary": "#2d5f8a", "accent": "#eaf2f8",
            "heading_text": "#1a3a5c", "body_text": "#2c3e50",
        },
    },
    "resume": {
        "name": "Resume / CV",
        "cover": False,
        "fonts": {"heading": "Calibri", "body": "Calibri"},
        "colors": {
            "primary": "#2c3e50", "accent": "#ecf0f1",
            "heading_text": "#1a252f", "body_text": "#2c3e50",
        },
    },
    "newsletter": {
        "name": "Email Newsletter",
        "cover": True,
        "fonts": {"heading": "Georgia", "body": "Georgia"},
        "colors": {
            "primary": "#c0392b", "accent": "#fdecea",
            "heading_text": "#922b21", "body_text": "#333333",
        },
    },
    "meeting-minutes": {
        "name": "Meeting Minutes",
        "cover": True,
        "fonts": {"heading": "Arial", "body": "Arial"},
        "colors": {
            "primary": "#27ae60", "accent": "#eafaf1",
            "heading_text": "#1e8449", "body_text": "#333333",
        },
    },
    "whitepaper": {
        "name": "Technical Whitepaper",
        "cover": True,
        "fonts": {"heading": "Times New Roman", "body": "Times New Roman"},
        "colors": {
            "primary": "#1a1a2e", "accent": "#f0f0f5",
            "heading_text": "#0d0d1a", "body_text": "#2d2d3f",
        },
    },
}

DEFAULT_TEMPLATE = "report"

# --- Document type auto-detection ---
def detect_doc_type(text: str) -> str:
    """Auto-detect document type from content keywords."""
    keywords = {
        "prd": ["prd", "product requirement", "feature", "stakeholder", "user story",
                 "product spec", "requirements document"],
        "report": ["report", "analysis", "findings", "summary", "conclusion",
                    "quarterly", "annual report"],
        "academic": ["abstract", "introduction", "methodology", "literature", "reference",
                      "bibliography", "thesis", "dissertation", "research"],
        "manual": ["manual", "guide", "instructions", "steps", "how to",
                    "getting started", "tutorial", "setup"],
        "contract": ["contract", "agreement", "party", "clause", "terms",
                      "hereby", "hereinafter", "indemnification", "warranty"],
        "proposal": ["proposal", "proposed solution", "scope of work", "deliverables",
                      "budget", "timeline", "executive summary", "business case"],
        "resume": ["resume", "curriculum vitae", "work experience", "education",
                    "skills", "certification", "contact information"],
        "newsletter": ["newsletter", "issue #", "weekly digest", "subscriber",
                        "unsubscribe", "featured article", "editor's note"],
        "meeting-minutes": ["meeting minutes", "attendees", "agenda", "action items",
                             "next meeting", "minutes of", "discussion", "motion"],
        "whitepaper": ["whitepaper", "white paper", "technical overview", "architecture",
                        "benchmark", "use case", "industry analysis", "solution architecture"],
    }
    text_lower = text.lower()
    scores = {}
    for doc_type, kw_list in keywords.items():
        scores[doc_type] = sum(1 for kw in kw_list if kw in text_lower)
    best = max(scores, key=scores.get)
    return best if scores[best] > 0 else "report"


# --- Markdown parser ---
def parse_markdown(text: str) -> list:
    """Parse Markdown text into a flat list of structured element dicts."""
    elements = []
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        if not line.strip():
            i += 1
            continue

        # Code blocks (fenced)
        if line.strip().startswith("```"):
            lang = line.strip()[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1
            elements.append({"type": "code_block", "lang": lang, "code": "\n".join(code_lines)})
            continue

        # Blockquotes
        if line.strip().startswith(">"):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote_lines.append(lines[i].strip()[1:].strip())
                i += 1
            elements.append({"type": "blockquote", "lines": quote_lines})
            continue

        # Tables
        if "|" in line and line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and "|" in lines[i]:
                table_lines.append(lines[i])
                i += 1
            if len(table_lines) >= 2:
                elements.append({"type": "table", "data": parse_table(table_lines)})
            continue

        # Headings
        m = re.match(r'^(#{1,6})\s+(.+)$', line)
        if m:
            level = len(m.group(1))
            text_content = m.group(2).strip()
            elements.append({"type": "heading", "level": level, "text": text_content})
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^[-*_]{3,}\s*$', line):
            elements.append({"type": "hr"})
            i += 1
            continue

        # Unordered lists
        if re.match(r'^\s*[-*+]\s+', line):
            list_items = []
            indent_level = len(line) - len(line.lstrip())
            while i < len(lines) and re.match(r'^\s*[-*+]\s+', lines[i]):
                item_text = re.sub(r'^\s*[-*+]\s+', '', lines[i])
                list_items.append({"text": item_text, "indent": indent_level})
                i += 1
            elements.append({"type": "list", "ordered": False, "items": list_items})
            continue

        # Ordered lists
        if re.match(r'^\s*\d+[.)]\s+', line):
            list_items = []
            while i < len(lines) and re.match(r'^\s*\d+[.)]\s+', lines[i]):
                item_text = re.sub(r'^\s*\d+[.)]\s+', '', lines[i])
                list_items.append({"text": item_text})
                i += 1
            elements.append({"type": "list", "ordered": True, "items": list_items})
            continue

        # Paragraph
        para_lines = []
        while i < len(lines) and lines[i].strip() and \
                not re.match(r'^(#{1,6}\s|```|>|\|)', lines[i]) and \
                not re.match(r'^\s*[-*+]\s+', lines[i]) and \
                not re.match(r'^\s*\d+[.)]\s+', lines[i]):
            para_lines.append(lines[i])
            i += 1
        elements.append({"type": "paragraph", "text": " ".join(para_lines)})
        continue

    return elements


def parse_table(lines: list) -> list:
    """Parse a Markdown table into header rows and data rows."""
    if len(lines) < 2:
        return []
    header = [c.strip() for c in lines[0].strip().strip("|").split("|")]
    start = 2 if len(lines) > 2 and re.match(r'^[\|\s\-:]+$', lines[1]) else 1
    rows = []
    for line in lines[start:]:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if cells and any(c for c in cells):
            rows.append(cells)
    return [header] + rows


# --- Heading numbering ---
def auto_number_headings(elements: list) -> list:
    """Add hierarchical numbering (1, 1.1, 1.1.1) to heading elements in place."""
    counters = [0, 0, 0, 0, 0, 0]
    for el in elements:
        if el["type"] != "heading":
            continue
        level = el["level"]
        counters[level - 1] += 1
        for i in range(level, 6):
            counters[i] = 0
        num_parts = [str(counters[i]) for i in range(level) if counters[i] > 0]
        el["number"] = ".".join(num_parts)
        el["display"] = f"{el['number']} {el['text']}"
    return elements

# --- Color helpers ---
def hex_to_rgb(hex_color: str):
    hex_color = hex_color.lstrip("#")
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))


# --- Cell shading helper ---
def set_cell_shading(cell, color_hex: str):
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


# --- Main .docx generator ---
def generate_docx(elements: list, template_name: str, output_path: Path) -> Path:
    """Generate a real .docx file with professional formatting."""
    if not HAS_DOCX:
        raise RuntimeError("python-docx is required. Install with: pip install python-docx")

    template = TEMPLATES.get(template_name, TEMPLATES[DEFAULT_TEMPLATE])
    doc = Document()

    # Page setup (A4)
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # Style configuration
    style = doc.styles['Normal']
    style.font.name = template["fonts"]["body"]
    style.font.size = Pt(10.5)
    style.font.color.rgb = RGBColor(*hex_to_rgb(template["colors"]["body_text"]))
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    heading_font = template["fonts"]["heading"]
    heading_color = hex_to_rgb(template["colors"]["heading_text"])
    for i in range(1, 4):
        h_style = doc.styles[f'Heading {i}']
        h_style.font.name = heading_font
        h_style.font.color.rgb = RGBColor(*heading_color)
        h_style.font.bold = True
        sizes = {1: Pt(18), 2: Pt(14), 3: Pt(12)}
        befores = {1: Pt(24), 2: Pt(18), 3: Pt(12)}
        afters = {1: Pt(12), 2: Pt(8), 3: Pt(6)}
        h_style.font.size = sizes[i]
        h_style.paragraph_format.space_before = befores[i]
        h_style.paragraph_format.space_after = afters[i]

    # Cover page
    heading_elements = [el for el in elements if el["type"] == "heading"]
    has_cover = template.get("cover", False) and len(heading_elements) > 0

    if has_cover:
        for _ in range(6):
            doc.add_paragraph("")
        title_text = heading_elements[0].get("display", heading_elements[0]["text"])
        cover_title = doc.add_paragraph()
        cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cover_title.add_run(title_text)
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.name = heading_font
        run.font.color.rgb = RGBColor(*hex_to_rgb(template["colors"]["primary"]))

        cover_sub = doc.add_paragraph()
        cover_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        today_str = datetime.now().strftime("%B %d, %Y")
        sub_run = cover_sub.add_run(f"{template['name']}\n{today_str}")
        sub_run.font.size = Pt(14)
        sub_run.font.name = template["fonts"]["body"]
        sub_run.font.color.rgb = RGBColor(*hex_to_rgb(template["colors"]["heading_text"]))
        doc.add_page_break()

    # Table of Contents
    toc_headings = heading_elements
    if has_cover:
        toc_headings = heading_elements[1:] if len(heading_elements) > 1 else []

    if toc_headings:
        doc.add_heading("Table of Contents", level=1)
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fldChar1)
        run2 = paragraph.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
        run2._r.append(instrText)
        run3 = paragraph.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
        run3._r.append(fldChar2)
        run4 = paragraph.add_run("[ Right-click and select 'Update Field' to populate TOC ]")
        run4.font.color.rgb = RGBColor(128, 128, 128)
        run4.font.size = Pt(9)
        run5 = paragraph.add_run()
        fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run5._r.append(fldChar3)
        doc.add_page_break()

    # Render content (skip cover heading)
    skip_first_heading = has_cover
    for el in elements:
        if el["type"] == "heading":
            if skip_first_heading:
                skip_first_heading = False
                continue
            level = min(el["level"], 3)
            display_text = el.get("display", el["text"])
            doc.add_heading(display_text, level=level)

        elif el["type"] == "paragraph":
            para = doc.add_paragraph()
            _add_formatted_text(para, el["text"], template)

        elif el["type"] == "list":
            _render_list(doc, el, template)

        elif el["type"] == "table":
            _render_table(doc, el["data"], template)

        elif el["type"] == "code_block":
            _render_code_block(doc, el, template)

        elif el["type"] == "blockquote":
            _render_blockquote(doc, el, template)

        elif el["type"] == "hr":
            para = doc.add_paragraph()
            run = para.add_run("\u2500" * 60)
            run.font.color.rgb = RGBColor(200, 200, 200)
            run.font.size = Pt(8)

    # Headers & Footers
    header = section.header
    header.is_linked_to_previous = False
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header_para.add_run()
    title_for_header = heading_elements[0].get("text", "Document") if heading_elements else "Document"
    header_run.text = title_for_header
    header_run.font.size = Pt(8)
    header_run.font.color.rgb = RGBColor(128, 128, 128)
    header_run.font.name = template["fonts"]["body"]

    footer = section.footer
    footer.is_linked_to_previous = False
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run_page = footer_para.add_run("Page ")
    run_page.font.size = Pt(8)
    run_page.font.color.rgb = RGBColor(128, 128, 128)

    # PAGE field
    fldChar_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    footer_para.add_run()._r.append(fldChar_begin)
    instrText_page = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    footer_para.add_run()._r.append(instrText_page)
    fldChar_sep = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    footer_para.add_run()._r.append(fldChar_sep)
    run_num = footer_para.add_run("1")
    run_num.font.size = Pt(8)
    fldChar_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    footer_para.add_run()._r.append(fldChar_end)

    run_of = footer_para.add_run(" of ")
    run_of.font.size = Pt(8)
    run_of.font.color.rgb = RGBColor(128, 128, 128)

    # NUMPAGES field
    fldChar_begin2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    footer_para.add_run()._r.append(fldChar_begin2)
    instrText_numpages = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> NUMPAGES </w:instrText>')
    footer_para.add_run()._r.append(instrText_numpages)
    fldChar_sep2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    footer_para.add_run()._r.append(fldChar_sep2)
    run_total = footer_para.add_run("1")
    run_total.font.size = Pt(8)
    fldChar_end2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    footer_para.add_run()._r.append(fldChar_end2)

    run_space = footer_para.add_run("  |  ")
    run_space.font.size = Pt(8)
    run_date = footer_para.add_run(datetime.now().strftime("%Y-%m-%d"))
    run_date.font.size = Pt(8)
    run_date.font.color.rgb = RGBColor(128, 128, 128)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path

# --- Inline text formatting ---
def _add_formatted_text(paragraph, text: str, template: dict):
    """Add paragraph text with basic inline formatting (**bold**, *italic*, `code`)."""
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    last_end = 0
    for m in pattern.finditer(text):
        if m.start() > last_end:
            run = paragraph.add_run(text[last_end:m.start()])
            run.font.name = template["fonts"]["body"]
        if m.group(2):
            run = paragraph.add_run(m.group(2))
            run.bold = True
            run.font.name = template["fonts"]["body"]
        elif m.group(3):
            run = paragraph.add_run(m.group(3))
            run.italic = True
            run.font.name = template["fonts"]["body"]
        elif m.group(4):
            run = paragraph.add_run(m.group(4))
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F0" w:val="clear"/>')
            run._r.get_or_add_rPr().append(shading)
        last_end = m.end()
    if last_end < len(text):
        run = paragraph.add_run(text[last_end:])
        run.font.name = template["fonts"]["body"]


def _render_list(doc, el: dict, template: dict):
    """Render bullet or numbered list."""
    for item in el.get("items", []):
        style_name = 'List Bullet' if not el.get("ordered") else 'List Number'
        para = doc.add_paragraph(style=style_name)
        para.clear()
        _add_formatted_text(para, item["text"], template)


def _render_table(doc, data: list, template: dict):
    """Render a table with header styling, borders, and alternating row shading."""
    if not data:
        return
    doc.add_paragraph("")
    rows = len(data)
    cols = max(len(row) for row in data) if data else 1
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            if j >= cols:
                break
            cell = table.cell(i, j)
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(str(cell_text))
            run.font.name = template["fonts"]["body"]
            run.font.size = Pt(9.5)

            if i == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                set_cell_shading(cell, template["colors"]["primary"].lstrip("#"))
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                if i % 2 == 0:
                    set_cell_shading(cell, template["colors"]["accent"].lstrip("#"))
    doc.add_paragraph("")


def _render_code_block(doc, el: dict, template: dict):
    """Render a code block with monospace font and grey background."""
    if el.get("lang"):
        label_para = doc.add_paragraph()
        label_run = label_para.add_run(f"[ {el['lang']} ]")
        label_run.font.size = Pt(8)
        label_run.font.color.rgb = RGBColor(128, 128, 128)
        label_run.italic = True

    code_para = doc.add_paragraph()
    code_para.paragraph_format.left_indent = Cm(0.5)
    code_para.paragraph_format.space_before = Pt(4)
    code_para.paragraph_format.space_after = Pt(4)

    pPr = code_para._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
    pPr.append(shading)

    for idx, line in enumerate(el["code"].splitlines()):
        if idx > 0:
            code_para.add_run("\n")
        run = code_para.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph("")


def _render_blockquote(doc, el: dict, template: dict):
    """Render a blockquote with indent and left grey border."""
    text = " ".join(el.get("lines", []))
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(1.0)
    para.paragraph_format.right_indent = Cm(1.0)

    pPr = para._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="12" w:space="8" w:color="999999"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    run = para.add_run(text)
    run.font.name = template["fonts"]["body"]
    run.font.size = Pt(9.5)
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


# --- PDF generation via pandoc + weasyprint ---
def md_to_pdf_via_pandoc(md_text: str, template_name: str, output_path: Path) -> Path:
    """Convert Markdown text to PDF using pandoc + weasyprint."""
    template = TEMPLATES.get(template_name, TEMPLATES[DEFAULT_TEMPLATE])

    with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False) as f:
        md_path = f.name
        f.write(md_text)

    try:
        primary_color = template["colors"]["primary"]
        heading_font = template["fonts"]["heading"]
        body_font = template["fonts"]["body"]

        css_content = f"""@page {{
    margin: 2.5cm;
    @bottom-center {{
        content: counter(page) " of " counter(pages);
        font-size: 8pt;
        color: #888;
    }}
}}
body {{
    font-family: "{body_font}", sans-serif;
    font-size: 10.5pt;
    color: #333;
    line-height: 1.5;
}}
h1 {{ font-family: "{heading_font}", sans-serif; color: {primary_color}; font-size: 20pt; }}
h2 {{ font-family: "{heading_font}", sans-serif; color: {primary_color}; font-size: 15pt; }}
h3 {{ font-family: "{heading_font}", sans-serif; color: {primary_color}; font-size: 12pt; }}
code {{ font-family: "Courier New", monospace; background-color: #f5f5f5; padding: 1px 4px; }}
pre {{ background-color: #f5f5f5; padding: 10px; border-radius: 4px; font-size: 9pt; }}
blockquote {{
    border-left: 3px solid #999;
    margin-left: 0;
    padding-left: 1em;
    color: #555;
    font-style: italic;
}}
table {{ border-collapse: collapse; width: 100%; }}
th {{ background-color: {primary_color}; color: white; padding: 6px 10px; }}
td {{ border: 1px solid #ddd; padding: 4px 8px; }}
tr:nth-child(even) {{ background-color: #f9f9f9; }}
"""

        css_path = md_path.replace('.md', '.css')
        with open(css_path, 'w') as cf:
            cf.write(css_content)

        result = subprocess.run(
            [
                "pandoc", md_path,
                "-o", str(output_path),
                "--from", "markdown",
                "--pdf-engine", "weasyprint",
                "--css", css_path,
                "--metadata", f"title={template['name']}",
                "--metadata", f"date={datetime.now().strftime('%Y-%m-%d')}",
            ],
            capture_output=True, text=True, timeout=60,
        )

        Path(css_path).unlink(missing_ok=True)

        if result.returncode != 0:
            raise RuntimeError(f"pandoc PDF generation failed: {result.stderr}")

        return output_path

    finally:
        Path(md_path).unlink(missing_ok=True)

# --- Preview render (Markdown with CSS annotations) ---
def render_preview(elements: list, template_name: str) -> str:
    """Render a styled Markdown preview with template annotations."""
    template = TEMPLATES.get(template_name, TEMPLATES[DEFAULT_TEMPLATE])
    lines = []
    lines.append(f"<!-- Doc Weaver -- Template: {template['name']} -->")
    lines.append(f"<style>")
    lines.append(f"  h1,h2,h3,h4,h5,h6 {{ font-family: {template['fonts']['heading']}; }}")
    lines.append(f"  body {{ font-family: {template['fonts']['body']}; }}")
    lines.append(f"  a {{ color: {template['colors']['primary']}; }}")
    lines.append(f"</style>")
    lines.append("")

    heading_els = [el for el in elements if el["type"] == "heading"]
    has_cover = template.get("cover", False) and len(heading_els) > 0
    cover_skipped = False

    if has_cover:
        title = heading_els[0].get("display", heading_els[0]["text"])
        lines.append(f"# {title}")
        lines.append("")
        lines.append(f"*{template['name']} -- {datetime.now().strftime('%B %d, %Y')}*")
        lines.append("")
        lines.append("---")
        lines.append("")
        cover_skipped = True

    toc_items = [
        el for el in elements
        if el["type"] == "heading" and not (cover_skipped and el is heading_els[0])
    ]
    if toc_items:
        lines.append("## Table of Contents")
        lines.append("")
        for h in toc_items:
            indent = "  " * (h["level"] - 1)
            lines.append(f"{indent}- {h.get('display', h['text'])}")
        lines.append("")
        lines.append("---")
        lines.append("")

    for el in elements:
        if el["type"] == "heading":
            if cover_skipped and el is heading_els[0]:
                continue
            level = el["level"]
            lines.append(f"{'#' * level} {el.get('display', el['text'])}")
            lines.append("")

        elif el["type"] == "code_block":
            lines.append(f"```{el.get('lang', '')}")
            lines.append(el["code"])
            lines.append("```")
            lines.append("")

        elif el["type"] == "table":
            data = el["data"]
            if not data:
                continue
            lines.append("| " + " | ".join(data[0]) + " |")
            lines.append("| " + " | ".join(["---"] * len(data[0])) + " |")
            for row in data[1:]:
                lines.append("| " + " | ".join(row) + " |")
            lines.append("")

        elif el["type"] == "blockquote":
            for qline in el.get("lines", []):
                lines.append(f"> {qline}")
            lines.append("")

        elif el["type"] == "paragraph":
            lines.append(el["text"])
            lines.append("")

        elif el["type"] == "list":
            for item in el.get("items", []):
                prefix = "1." if el.get("ordered") else "-"
                lines.append(f"{prefix} {item['text']}")
            lines.append("")

        elif el["type"] == "hr":
            lines.append("---")
            lines.append("")

    return "\n".join(lines)


# --- CLI ---
def list_templates() -> list:
    return [{"name": k, "label": v["name"], "cover": v["cover"],
             "heading_font": v["fonts"]["heading"],
             "body_font": v["fonts"]["body"],
             "primary": v["colors"]["primary"]}
            for k, v in TEMPLATES.items()]


def check_runtime() -> list:
    """Return dependency checks for the installed conversion toolchain."""
    checks = [
        {
            "name": "python-docx",
            "required_for": ".docx output",
            "ok": HAS_DOCX,
            "hint": "pip install python-docx",
        },
        {
            "name": "pandoc",
            "required_for": ".pdf output",
            "ok": shutil.which("pandoc") is not None,
            "hint": "brew install pandoc",
        },
        {
            "name": "weasyprint",
            "required_for": ".pdf output",
            "ok": importlib.util.find_spec("weasyprint") is not None,
            "hint": "pip install weasyprint",
        },
    ]
    return checks


def print_doctor_report() -> int:
    """Print a concise readiness report and return an exit code."""
    checks = check_runtime()
    print("Doc Weaver Runtime Check")
    print("========================")
    for check in checks:
        status = "OK" if check["ok"] else "MISSING"
        print(f"[{status}] {check['name']:<12} required for {check['required_for']}")
        if not check["ok"]:
            print(f"        install: {check['hint']}")

    can_docx = next(c for c in checks if c["name"] == "python-docx")["ok"]
    can_pdf = all(c["ok"] for c in checks)
    print("")
    print(f".docx generation: {'available' if can_docx else 'blocked'}")
    print(f".pdf generation:  {'available' if can_pdf else 'blocked'}")
    if can_docx and not can_pdf:
        print("PDF dependencies are optional; Word output will still work.")
    return 0 if can_docx else 1


def main():
    import argparse
    parser = argparse.ArgumentParser(
        description="Doc Weaver -- Markdown to polished Word (.docx) and PDF documents"
    )
    parser.add_argument("--input", "-i", help="Input Markdown file")
    parser.add_argument("--template", "-t",
                        choices=list(TEMPLATES.keys()) + ["auto"],
                        default="auto")
    parser.add_argument("--output", "-o", help="Output file (.docx or .pdf)")
    parser.add_argument("--show-templates", action="store_true",
                        help="List available templates")
    parser.add_argument("--preview", action="store_true",
                        help="Preview as styled Markdown (no file conversion)")
    parser.add_argument("--doctor", action="store_true",
                        help="Check local dependencies for Word/PDF generation")
    args = parser.parse_args()

    if args.doctor:
        return print_doctor_report()

    if args.show_templates:
        print(f"\n{'Key':<20} {'Name':<30} {'Cover':<8} {'Heading Font':<25} {'Body Font':<20} {'Primary'}")
        print("-" * 120)
        for t in list_templates():
            cover_str = "cover" if t["cover"] else "no cover"
            print(f"  {t['name']:<18}  {t['label']:<28}  {cover_str:<6}  "
                  f"{t['heading_font']:<23}  {t['body_font']:<18}  {t['primary']}")
        print(f"\n  Total: {len(TEMPLATES)} templates\n")
        return 0

    # Read input
    if args.input:
        text = Path(args.input).read_text()
    else:
        # Self-test sample PRD
        text = """# My Project PRD

## Overview

This is a product requirements document for the new **analytics dashboard**.

## Goals

- Launch MVP by **Q3 2026**
- Support 10K concurrent users
- Achieve 99.9% uptime SLA

## Features

### Core Features

1. Login and authentication
2. Real-time data dashboard
3. Custom report builder
4. CSV/PDF export

### Premium Features

- White-label branding
- API access
- Priority support

## Technical Architecture

The system uses a **microservices architecture** with React frontend and Python backend.

```python
def get_dashboard(user_id):
    \"\"\"Fetch user dashboard data.\"\"\"
    return Dashboard.query.filter_by(user_id=user_id).first()
```

## API Endpoints

| Endpoint   | Method | Auth Required | Description          |
|------------|--------|---------------|----------------------|
| /api/login | POST   | No            | User authentication  |
| /api/data  | GET    | Yes           | Fetch dashboard data |
| /api/export| POST   | Yes           | Export reports       |

> **Important:** All API endpoints must enforce rate limiting to prevent abuse.

## Timeline

- **Alpha**: August 2026
- **Beta**: October 2026
- **GA**: December 2026

## Appendix

See [API Documentation](https://docs.example.com) for details.
"""
        if args.preview:
            print("[doc-weaver] No input provided, self-testing with sample PRD.\\n")

    # Detect template
    if args.template == "auto":
        template_name = detect_doc_type(text)
        print(f"[doc-weaver] Auto-detected document type: {template_name} "
              f"({TEMPLATES[template_name]['name']})")
    else:
        template_name = args.template

    # Parse
    elements = parse_markdown(text)
    elements = auto_number_headings(elements)

    # Preview mode
    if args.preview or (not args.preview and not args.output):
        if not args.input and not args.preview:
            print("[doc-weaver] No input file or --output specified. Preview mode:\\n")
        rendered = render_preview(elements, template_name)
        print(rendered)
        return 0

    # Determine output format
    out_path = Path(args.output)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    ext = out_path.suffix.lower()

    if ext == ".docx":
        generate_docx(elements, template_name, out_path)
        print(f"[OK] Word document saved: {out_path}")
        return 0

    elif ext == ".pdf":
        docx_path = out_path.with_suffix(".docx")
        generate_docx(elements, template_name, docx_path)
        print(f"[OK] Word document saved: {docx_path}")

        md_for_pdf = render_preview(elements, template_name)
        md_for_pdf = re.sub(r'<!--.*?-->', '', md_for_pdf, flags=re.DOTALL)
        md_for_pdf = re.sub(r'<style>.*?</style>', '', md_for_pdf, flags=re.DOTALL)

        try:
            md_to_pdf_via_pandoc(md_for_pdf, template_name, out_path)
            print(f"[OK] PDF saved: {out_path}")
        except Exception as e:
            print(f"[WARN] PDF generation failed: {e}")
            print(f"       Word document is available at: {docx_path}")
            return 1
        return 0

    else:
        docx_path = out_path if ext == ".docx" else out_path.with_suffix(".docx")
        generate_docx(elements, template_name, docx_path)
        print(f"[OK] Word document saved: {docx_path}")
        return 0


if __name__ == "__main__":
    sys.exit(main())
