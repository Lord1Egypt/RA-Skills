# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 创建文档
doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(10)

# 标题
heading = doc.add_heading('阿福 Skills 汇总文档', level=0)
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
heading_run = heading.runs[0]
heading_run.font.name = 'Microsoft YaHei'
heading_run.font.size = Pt(16)
heading_run.font.bold = True

# 文档信息
info = doc.add_paragraph()
info.add_run('文档维护：阿福（AI 助理） | 更新时间：2026-03-07 00:25 | 版本：v1.0')

doc.add_paragraph('_' * 80)

# Skills 总览
doc.add_heading('📋 Skills 总览', level=1)
doc.add_paragraph('当前共有 2 个 Skills：')

# 表格
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Skill 名称'
hdr_cells[1].text = '一句话概述'
hdr_cells[2].text = '输出格式'

# 加粗表头
for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

# 添加数据
skills_data = [
    ('豆包专家点评', '自动处理豆包会话，生成专家级点评 HTML 网页', 'HTML + 语音'),
    ('项目知识专家', '以行业专家视角构建项目知识库', '飞书文档 + 批注')
]

for name, desc, output in skills_data:
    row_cells = table.add_row().cells
    row_cells[0].text = name
    row_cells[1].text = desc
    row_cells[2].text = output

doc.add_paragraph()

# Skill 1
doc.add_heading('🎯 Skill 1: 豆包专家点评系统', level=1)
doc.add_paragraph('一句话概述：自动处理豆包会话内容，生成包含知识架构图、深度洞察和行动建议的专家级点评 HTML 网页。')

doc.add_heading('🔄 完整流程（6 步）', level=2)

steps1 = [
    ('步骤 1：保存原始内容', 'doubao-sessions/日期 - 序号.md'),
    ('步骤 2：更新 worklog.txt', '添加今日记录'),
    ('步骤 3：生成专家点评 HTML', '6 章节标准（专家评分/核心观点/深度洞察/知识架构/对比分析/行动建议）'),
    ('步骤 4：Chrome 自动打开', '独立窗口显示'),
    ('步骤 5：TTS 语音生成', 'MP3 到 Temp 目录'),
    ('步骤 6：飞书发送 + 自动播放', 'filePath 参数 + Start-Process')
]

for step, desc in steps1:
    p = doc.add_paragraph(style='List Number')
    p.add_run(step).bold = True
    doc.add_paragraph(desc, style='List Bullet')

# Skill 2
doc.add_heading('🎯 Skill 2: 项目知识库构建专家', level=1)
doc.add_paragraph('一句话概述：基于用户提供的信息，以行业专家视角构建项目知识库，输出结构化飞书文档。')

doc.add_heading('🔄 完整流程（8 步迭代）', level=2)

steps2 = [
    '意图识别 - 无需触发词，理解真实意图',
    '专家视角分析 - Critical Thinking + 行业对标',
    '补全架构 - 业务/技术/组织三大架构',
    '飞书文档生成（v1.0）- 3 部分结构，分块写入',
    '发送批注清单 - 5 种类型（❓疑问/⚠️风险/💡建议/📝待补充/✅确认）',
    '等待用户回复 - 解析每条批注回复',
    '修订文档（v1.1）- 更新内容 + 版本号 + 修订记录',
    '循环迭代 - 直到所有批注处理完成'
]

for step in steps2:
    doc.add_paragraph(step, style='List Number')

doc.add_heading('📊 版本号规范', level=2)
doc.add_paragraph('格式：v{主版本}.{次版本}.{修订号}', style='List Bullet')
doc.add_paragraph('v1.0.0 - 初稿', style='List Bullet')
doc.add_paragraph('v1.1.0 - 小修订', style='List Bullet')
doc.add_paragraph('v2.0.0 - 重大更新', style='List Bullet')

# 用户偏好
doc.add_heading('🎯 用户偏好（通用）', level=1)

preferences = [
    '视觉学习者 - 结构化图表',
    '自动播放 - TTS 后自动播放',
    '金字塔原理 - 结论先行',
    'Critical Thinking - 主动思考',
    'MECE 法则 - 不重不漏'
]

for pref in preferences:
    doc.add_paragraph(pref, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('_' * 80)
doc.add_paragraph('文档同步：本文档已同步到飞书和 Word 双格式')

# 保存文档
save_path = r'C:\Users\Xiabi\.openclaw\workspace\skills\阿福 Skills 汇总文档.docx'
doc.save(save_path)
print(f'Word document saved to: {save_path}')
print(f'File size: {len(open(save_path, "rb").read())} bytes')
