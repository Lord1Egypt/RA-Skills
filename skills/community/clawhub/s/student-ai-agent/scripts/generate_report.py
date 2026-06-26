"""
报告生成器 - 生成 .docx 格式的学术报告
使用 python-docx 库
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import json
from pathlib import Path
from datetime import datetime


def create_report(title, student_name, course_name, sections, figures=None, code_file=None, output_path=None):
    """
    生成一份完整的学术报告 (.docx)
    
    Args:
        title: 报告标题
        student_name: 学生姓名
        course_name: 课程名称
        sections: list of dicts, each with 'heading' and 'content'
        figures: list of figure paths to include
        code_file: path to code file for appendix
        output_path: output file path
    """
    doc = Document()
    
    # ---- Page Setup ----
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    
    # ---- Title Page ----
    # Add some spacing before title
    for _ in range(4):
        doc.add_paragraph()
    
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.font.size = Pt(24)
    run.bold = True
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # Student info
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info_para.add_run(f"\n{student_name}\n{course_name}\n{datetime.now().strftime('%B %d, %Y')}")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    doc.add_page_break()
    
    # ---- Table of Contents placeholder ----
    toc_heading = doc.add_heading('Table of Contents', level=1)
    doc.add_paragraph('[Auto-generate in Word: References > Table of Contents]')
    doc.add_page_break()
    
    # ---- Main Sections ----
    for section in sections:
        heading_level = section.get('level', 1)
        doc.add_heading(section['heading'], level=heading_level)
        
        content = section.get('content', '')
        if isinstance(content, list):
            for item in content:
                para = doc.add_paragraph(item, style='List Bullet')
                para.paragraph_format.space_after = Pt(4)
        else:
            for paragraph_text in content.split('\n'):
                if paragraph_text.strip():
                    para = doc.add_paragraph(paragraph_text)
                    para.paragraph_format.space_after = Pt(6)
                    for run in para.runs:
                        run.font.size = Pt(11)
                        run.font.name = 'Times New Roman'
    
    # ---- Figures ----
    if figures:
        doc.add_heading('Figures', level=1)
        for i, fig_path in enumerate(figures, 1):
            fig_path = Path(fig_path)
            if fig_path.exists():
                doc.add_picture(str(fig_path), width=Inches(5.5))
                caption = doc.add_paragraph(f'Figure {i}: {fig_path.stem}')
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.runs[0].font.size = Pt(10)
                caption.runs[0].italic = True
                doc.add_paragraph()
    
    # ---- Code Appendix ----
    if code_file:
        code_path = Path(code_file)
        if code_path.exists():
            doc.add_page_break()
            doc.add_heading('Appendix: Source Code', level=1)
            code_text = code_path.read_text(encoding='utf-8')
            para = doc.add_paragraph()
            run = para.add_run(code_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(8)
    
    # ---- Save ----
    if output_path is None:
        output_path = Path('output/report/report.docx')
    
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"  ✅ 报告已生成: {output_path}")
    return output_path


def create_report_from_json(json_path, output_path=None):
    """从JSON结构文件生成报告"""
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    return create_report(
        title=data.get('title', 'Untitled Report'),
        student_name=data.get('student_name', 'Student'),
        course_name=data.get('course_name', 'Course'),
        sections=data.get('sections', []),
        figures=data.get('figures', []),
        code_file=data.get('code_file'),
        output_path=output_path
    )


if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        create_report_from_json(sys.argv[1])
    else:
        # Demo
        create_report(
            title="Demo Report",
            student_name="Student Name",
            course_name="CS101",
            sections=[
                {"heading": "Introduction", "content": "This is a demo report."},
                {"heading": "Methodology", "content": "We used method X."},
                {"heading": "Results", "content": "Results show Y."},
                {"heading": "Conclusion", "content": "In conclusion, Z."},
            ]
        )
