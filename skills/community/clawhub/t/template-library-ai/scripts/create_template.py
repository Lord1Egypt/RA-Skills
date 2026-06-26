#!/usr/bin/env python3
"""
根据用户描述从零创建 .docx 模板。
支持：纸张/边距/页眉页脚/封面布局/字体字号/对齐/图片/表格/{{}}占位符

用法：
    python scripts/create_template.py --output ./output.docx < params.json

参数文件格式 (JSON):
{
  "page": {
    "size": "A4",           // A4, B5, 或 [w_mm, h_mm]
    "margin_top_cm": 2.54,
    "margin_bottom_cm": 2.54,
    "margin_left_cm": 3.17,
    "margin_right_cm": 3.17
  },
  "header": "XX大学毕业论文",
  "footer_page_number": true,
  "cover": [
    {"type": "image", "path": "C:/logo.png", "width_cm": 3, "align": "left"},
    {"type": "text", "content": "毕业论文", "font": "黑体", "size_pt": 22, "bold": true, "align": "center"},
    {"type": "text", "content": "{{论文题目}}", "font": "宋体", "size_pt": 16, "align": "center"},
    {"type": "text", "content": "{{姓名}}", "font": "宋体", "size_pt": 14, "align": "center"},
    {"type": "text", "content": "{{学号}}", "font": "宋体", "size_pt": 14, "align": "center"}
  ],
  "body": {
    "font": "宋体", "size_pt": 12, "line_spacing": 1.5,
    "first_line_indent_cm": 0.74, "alignment": "justify"
  },
  "headings": [
    {"level": 1, "font": "黑体", "size_pt": 16, "bold": true, "space_before_pt": 12},
    {"level": 2, "font": "黑体", "size_pt": 14, "bold": true, "space_before_pt": 6}
  ],
  "placeholders": [
    {"text": "{{内容:研究背景}}", "style": "body"},
    {"text": "{{内容:研究方法}}", "style": "body"},
    {"text": "{{TABLE:实验数据}}", "style": "body"},
    {"text": "{{参考文献}}", "style": "body"}
  ]
}
"""

import json, os, sys
from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_font(cell, font_name, font_size):
    """Set east-asian font for table cells."""
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.name = font_name
            r.font.size = font_size
            r._r.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def make_page_number(paragraph):
    """Insert auto page number field into a footer paragraph."""
    run = paragraph.add_run()
    fldChar = run._r.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    run._r.append(fldChar)
    run2 = paragraph.add_run()
    instr = run2._r.makeelement(qn('w:instrText'), {})
    instr.text = ' PAGE '
    run2._r.append(instr)
    run3 = paragraph.add_run()
    run3._r.append(run3._r.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'}))


def add_run_with_font(paragraph, text, font_name='宋体', size_pt=12,
                      bold=False, color=None, east_asian=None):
    """Add a run with Chinese font support."""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    ea = east_asian or font_name
    run._r.rPr.rFonts.set(qn('w:eastAsia'), ea)
    return run


def create_template(params):
    """Main function: create .docx from parameters dict."""
    doc = Document()
    s = doc.sections[0]

    # === PAGE SETUP ===
    if params.get('page', {}).get('size') == 'B5':
        s.page_width = Cm(17.6)
        s.page_height = Cm(25)
    else:  # Default A4
        s.page_width = Cm(21)
        s.page_height = Cm(29.7)

    p = params.get('page', {})
    s.top_margin = Cm(p.get('margin_top_cm', 2.54))
    s.bottom_margin = Cm(p.get('margin_bottom_cm', 2.54))
    s.left_margin = Cm(p.get('margin_left_cm', 3.17))
    s.right_margin = Cm(p.get('margin_right_cm', 3.17))

    # === HEADER ===
    if params.get('header'):
        h = s.header.paragraphs[0]
        h.text = params['header']
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in h.runs:
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(128, 128, 128)

    # === FOOTER ===
    if params.get('footer_page_number', True):
        f = s.footer.paragraphs[0]
        f.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run_with_font(f, '- ', size_pt=9, font_name='宋体')
        make_page_number(f)
        add_run_with_font(f, ' -', size_pt=9, font_name='宋体')

    # === COVER ===
    for item in params.get('cover', []):
        if item['type'] == 'image':
            img_path = item.get('path')
            if img_path and os.path.exists(img_path):
                p = doc.add_paragraph()
                align_map = {'left': WD_ALIGN_PARAGRAPH.LEFT,
                             'center': WD_ALIGN_PARAGRAPH.CENTER,
                             'right': WD_ALIGN_PARAGRAPH.RIGHT}
                p.alignment = align_map.get(item.get('align', 'center'))
                width = Cm(item.get('width_cm', 4))
                p.add_run().add_picture(img_path, width=width)
            else:
                # Image not found, insert placeholder
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_run_with_font(p, '{{校徽}}', font_name='宋体', size_pt=10)

        elif item['type'] == 'text':
            p = doc.add_paragraph()
            align_map = {'left': WD_ALIGN_PARAGRAPH.LEFT,
                         'center': WD_ALIGN_PARAGRAPH.CENTER,
                         'right': WD_ALIGN_PARAGRAPH.RIGHT}
            p.alignment = align_map.get(item.get('align', 'center'), WD_ALIGN_PARAGRAPH.CENTER)
            add_run_with_font(p, item['content'],
                              font_name=item.get('font', '宋体'),
                              size_pt=item.get('size_pt', 12),
                              bold=item.get('bold', False))

    # Page break after cover
    doc.add_page_break()

    # === TABLE OF CONTENTS placeholder ===
    p_toc = doc.add_paragraph()
    p_toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_with_font(p_toc, '目  录', font_name='黑体', size_pt=16, bold=True)

    # Insert TOC field
    p_toc2 = doc.add_paragraph()
    run = p_toc2.add_run()
    fldChar = run._r.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    run._r.append(fldChar)
    run2 = p_toc2.add_run()
    instr = run2._r.makeelement(qn('w:instrText'), {})
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._r.append(instr)
    run3 = p_toc2.add_run()
    run3._r.append(run3._r.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'separate'}))
    p_toc2.add_run('（更新域以生成目录）')
    run4 = p_toc2.add_run()
    run4._r.append(run4._r.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'}))

    doc.add_page_break()

    # === HEADING STYLES ===
    heading_styles = {}
    for h_def in params.get('headings', []):
        level = h_def.get('level', 1)
        hs = doc.styles[f'Heading {level}']
        hs.font.name = h_def.get('font', '黑体')
        hs.font.size = Pt(h_def.get('size_pt', 16))
        hs.font.bold = h_def.get('bold', True)
        hs._element.rPr.rFonts.set(qn('w:eastAsia'), h_def.get('font', '黑体'))
        if h_def.get('space_before_pt'):
            hs.paragraph_format.space_before = Pt(h_def['space_before_pt'])
        heading_styles[level] = hs

    # === BODY STYLE ===
    body_cfg = params.get('body', {})
    normal = doc.styles['Normal']
    normal.font.name = body_cfg.get('font', '宋体')
    normal.font.size = Pt(body_cfg.get('size_pt', 12))
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), body_cfg.get('font', '宋体'))
    normal.paragraph_format.line_spacing = body_cfg.get('line_spacing', 1.5)
    normal.paragraph_format.first_line_indent = Cm(body_cfg.get('first_line_indent_cm', 0))

    # === PLACEHOLDER SECTIONS ===
    placeholders = params.get('placeholders', [])
    if placeholders:
        add_run_with_font(doc.add_paragraph(), '{{绪论}}', font_name='黑体', size_pt=16, bold=True)
        add_run_with_font(doc.add_paragraph(), '{{内容:研究背景}}', size_pt=12)

        # Table placeholder
        add_run_with_font(doc.add_paragraph(), '{{TABLE:实验数据}}', size_pt=12)
        # Empty table structure
        tbl = doc.add_table(rows=2, cols=3)
        tbl.style = 'Table Grid'
        for i in range(2):
            for j in range(3):
                tbl.cell(i, j).text = f'  '

        add_run_with_font(doc.add_paragraph(), '{{内容:分析与讨论}}', size_pt=12)

        # More section placeholders
        for ph in placeholders:
            style_def = ph.get('style', 'body')
            if style_def == 'heading':
                add_run_with_font(doc.add_paragraph(), ph['text'],
                                  font_name='黑体', size_pt=16, bold=True)
            else:
                add_run_with_font(doc.add_paragraph(), ph['text'], size_pt=12)

    # === REFERENCES ===
    doc.add_page_break()
    add_run_with_font(doc.add_paragraph(), '参考文献', font_name='黑体', size_pt=16, bold=True)
    add_run_with_font(doc.add_paragraph(), '{{参考文献}}', size_pt=12)

    return doc


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print('Usage: python create_template.py <params.json> <output.docx>')
        print('  or:  python create_template.py --stdin <output.docx> (read JSON from stdin)')
        sys.exit(1)

    if sys.argv[1] == '--stdin':
        params = json.loads(sys.stdin.read())
        output_path = sys.argv[2]
    else:
        with open(sys.argv[1], 'r', encoding='utf-8') as f:
            params = json.load(f)
        output_path = sys.argv[2]

    doc = create_template(params)
    doc.save(output_path)
    print(f'OK {output_path}')
