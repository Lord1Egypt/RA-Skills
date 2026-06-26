#!/usr/bin/env python3
"""
batch-format-converter 核心转换引擎
支持：CSV↔Excel↔JSON↔PDF/Markdown/DOCX/HTML/PNG/TXT
"""

import os
import sys
import json
import base64
import logging
import urllib.request
import urllib.error
from pathlib import Path
from typing import Optional, List, Dict, Any, Tuple

# ─── 91Skillhub Token Verification ─────────────────────────────────────────────
#
# 重要：已迁移到 skills-developer/shared/token_validator.py
# 本地保留此模块以确保独立运行能力
#
VERIFY_URL = "https://geo-api.yk-global.com/validate"  # 修正：正确的验证接口

# Tier 映射（与 yk global 前缀规范一致）
TIER_MAP = {
    "FREE":  "FREE",
    "BSC":   "BASIC",
    "STD":   "STANDARD",
    "PRO":   "PRO",
    "ENT":   "ENTERPRISE",
    "MAX":   "MAX",
}

# 已知 91Skillhub 前缀
VALID_PREFIXES = {
    "GEO", "PROFIT", "INV", "DATA", "MON",
    "PDF", "BANK", "CONTRACT", "EMAIL", "CONV",
    "RPT", "SENTIMENT",
}


def _get_cached(key: str) -> dict:
    """读取本地缓存（5分钟TTL）"""
    import time
    cache_dir = Path.home() / ".batch_converter_cache"
    cache_dir.mkdir(parents=True, exist_ok=True)
    cache_file = cache_dir / f"{key[:8].replace('/', '_')}.json"
    if not cache_file.exists():
        return None
    try:
        with open(cache_file) as f:
            data = json.load(f)
        if time.time() - data.get("_ts", 0) > 300:
            return None
        return data
    except Exception:
        return None


def _set_cached(key: str, data: dict) -> None:
    """写入本地缓存"""
    import time
    cache_dir = Path.home() / ".batch_converter_cache"
    cache_dir.mkdir(parents=True, exist_ok=True)
    cache_file = cache_dir / f"{key[:8].replace('/', '_')}.json"
    try:
        data["_ts"] = time.time()
        with open(cache_file, "w") as f:
            json.dump(data, f)
    except Exception:
        pass


def _map_prefix_to_tier(api_key: str) -> str:
    """根据 API key 前缀推断套餐级别。"""
    if not api_key:
        return "FREE"
    upper = api_key.upper()
    if "ENT" in upper:
        return "ENTERPRISE"
    if "MAX" in upper:
        return "MAX"
    if "PRO" in upper:
        return "PRO"
    if "STD" in upper:
        return "STANDARD"
    if "BSC" in upper:
        return "BASIC"
    if "FREE" in upper:
        return "FREE"
    return "FREE"


def verify_token(api_key: str) -> dict:
    """
    验证 API key via geo-api.yk-global.com。

    降级策略：
    1. 无 key → FREE
    2. key 不属于 91Skillhub 体系 → FREE
    3. 网络错误 → FREE（不阻断使用）
    4. 验证失败 → FREE
    5. 验证成功 → 对应 tier

    缓存：结果缓存 5 分钟。
    """
    if not api_key:
        return {"valid": False, "tier": "FREE", "error": "No API key"}

    # 快速判断：不在已知前缀列表 = 外部 key，跳过验证
    prefix = api_key.split("-")[0].upper() if "-" in api_key else api_key[:4].upper()
    if prefix not in VALID_PREFIXES:
        return {"valid": False, "tier": "FREE", "error": "Not a 91Skillhub key"}

    # 缓存查询
    cached = _get_cached(api_key)
    if cached:
        return cached

    try:
        req = urllib.request.Request(
            VERIFY_URL,
            method="POST",
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
            },
            data=b"{}",  # POST body required
        )
        with urllib.request.urlopen(req, timeout=10) as resp:
            data = json.loads(resp.read().decode("utf-8"))
            # 修正：使用 valid 字段
            if data.get("valid", False):
                tier = _map_prefix_to_tier(api_key)
                result = {"valid": True, "tier": tier, "prefix": data.get("prefix", ""),
                           "plan_id": data.get("plan_id"), "quota_remaining": data.get("quota_remaining")}
            else:
                result = {"valid": False, "tier": "FREE",
                           "error": data.get("error", "Invalid or expired key")}
            _set_cached(api_key, result)
            return result
    except urllib.error.HTTPError as e:
        try:
            err_body = json.loads(e.read().decode("utf-8"))
            err_msg = err_body.get("error", f"HTTP {e.code}")
        except Exception:
            err_msg = f"HTTP {e.code}"
        result = {"valid": False, "tier": "FREE", "error": err_msg}
        _set_cached(api_key, result)
        return result
    except Exception as e:
        # 网络异常 → FREE，不阻断使用
        return {"valid": False, "tier": "FREE", "error": f"Network error: {e}"}


# ─── 依赖检测 ────────────────────────────────────────────────
_deps = {}
try:
    import pandas as pd
    _deps["pandas"] = True
except ImportError:
    _deps["pandas"] = False

try:
    import fitz  # PyMuPDF
    _deps["pymupdf"] = True
except ImportError:
    _deps["pymupdf"] = False

try:
    import pdfplumber
    _deps["pdfplumber"] = True
except ImportError:
    _deps["pdfplumber"] = False

try:
    import docx
    _deps["python-docx"] = True
except ImportError:
    _deps["python-docx"] = False

try:
    from PIL import Image
    _deps["pillow"] = True
except ImportError:
    _deps["pillow"] = False

try:
    import pandoc
    _deps["pandoc"] = True
except ImportError:
    _deps["pandoc"] = False


logging.basicConfig(level=logging.INFO, format="%(levelname)s | %(message)s")
log = logging.getLogger("converter")


# ─── 编码检测 ────────────────────────────────────────────────
def detect_encoding(file_path: str) -> str:
    """自动检测文件编码，依次尝试 UTF-8 → GBK → ISO-8859-1"""
    for enc in ["utf-8", "gbk", "gb2312", "iso-8859-1"]:
        try:
            with open(file_path, "r", encoding=enc) as f:
                f.read()
            return enc
        except (UnicodeDecodeError, UnicodeError):
            continue
    return "utf-8"


# ─── 表格转换 ────────────────────────────────────────────────
def csv_to_excel(csv_path: str, out_path: str) -> str:
    enc = detect_encoding(csv_path)
    df = pd.read_csv(csv_path, encoding=enc)
    df.to_excel(out_path, index=False, engine="openpyxl")
    return out_path


def csv_to_json(csv_path: str, out_path: str) -> str:
    enc = detect_encoding(csv_path)
    df = pd.read_csv(csv_path, encoding=enc)
    df.to_json(out_path, orient="records", force_ascii=False, indent=2)
    return out_path


def excel_to_csv(xlsx_path: str, out_path: str) -> str:
    df = pd.read_excel(xlsx_path, engine="openpyxl")
    df.to_csv(out_path, index=False, encoding="utf-8-sig")
    return out_path


def excel_to_json(xlsx_path: str, out_path: str) -> str:
    xl = pd.ExcelFile(xlsx_path, engine="openpyxl")
    sheets = {sheet: xl.parse(sheet).to_dict(orient="records") for sheet in xl.sheet_names}
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(sheets, f, ensure_ascii=False, indent=2)
    return out_path


def excel_to_png(xlsx_path: str, out_path: str) -> str:
    """将 Excel 表格渲染为 PNG 图片"""
    df = pd.read_excel(xlsx_path, engine="openpyxl")
    # 简单表格渲染为图片
    from PIL import Image, ImageDraw, ImageFont
    col_widths = [max(len(str(v)) for v in df[col].astype(str)) + 2 for col in df.columns]
    row_height = 30
    header_height = 35
    total_width = sum(col_widths) * 9 + 20
    total_height = (len(df) + 1) * row_height + header_height + 20

    img = Image.new("RGB", (total_width, total_height), color="white")
    draw = ImageDraw.Draw(img)

    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 12)
        header_font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 13)
    except Exception:
        font = ImageFont.load_default()
        header_font = font

    x = 10
    for i, col in enumerate(df.columns):
        draw.text((x, 10), str(col), fill="1a3c5e", font=header_font)
        x += col_widths[i] * 9

    y = header_height
    for _, row in df.iterrows():
        x = 10
        for i, val in enumerate(row):
            draw.text((x, y), str(val)[:50], fill="black", font=font)
            x += col_widths[i] * 9
        y += row_height

    img.save(out_path)
    return out_path


def json_to_excel(json_path: str, out_path: str) -> str:
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    if isinstance(data, dict):
        # 多 sheet 场景
        with pd.ExcelWriter(out_path, engine="openpyxl") as writer:
            for sheet_name, records in data.items():
                if isinstance(records, list):
                    df = pd.DataFrame(records)
                else:
                    df = pd.DataFrame([records])
                safe_name = str(sheet_name)[:31]
                df.to_excel(writer, sheet_name=safe_name, index=False)
    else:
        df = pd.DataFrame(data)
        df.to_excel(out_path, index=False, engine="openpyxl")
    return out_path


def txt_to_csv(txt_path: str, out_path: str) -> str:
    enc = detect_encoding(txt_path)
    with open(txt_path, "r", encoding=enc) as f:
        lines = f.readlines()
    # 简单按逗号分割，无表头
    rows = [line.strip().split(",") for line in lines if line.strip()]
    if rows:
        df = pd.DataFrame(rows)
        df.to_csv(out_path, index=False, header=False, encoding="utf-8-sig")
    return out_path


# ─── 文档转换 ────────────────────────────────────────────────
def markdown_to_docx(md_path: str, out_path: str) -> str:
    try:
        import pandoc
        doc = pandoc.read(open(md_path, encoding="utf-8").read())
        pandoc.write(doc, out_path, format="docx")
    except Exception:
        # fallback: python-docx 直接创建简单 docx
        from docx import Document
        with open(md_path, "r", encoding="utf-8") as f:
            content = f.read()
        doc = Document()
        for para in content.split("\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
        doc.save(out_path)
    return out_path


def markdown_to_html(md_path: str, out_path: str) -> str:
    try:
        import pandoc
        doc = pandoc.read(open(md_path, encoding="utf-8").read())
        pandoc.write(doc, out_path, format="html5")
    except Exception:
        with open(md_path, "r", encoding="utf-8") as f:
            content = f.read().replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        html = f"<!DOCTYPE html><html><head><meta charset='utf-8'><title>Converted</title></head><body><pre>{content}</pre></body></html>"
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(html)
    return out_path


def docx_to_markdown(docx_path: str, out_path: str) -> str:
    try:
        import pandoc
        doc = pandoc.read(open(docx_path, "rb").read(), format="docx")
        pandoc.write(doc, out_path, format="markdown")
    except Exception:
        from docx import Document
        doc = Document(docx_path)
        lines = []
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text.strip() + "\n")
        with open(out_path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))
    return out_path


def html_to_markdown(html_path: str, out_path: str) -> str:
    try:
        import pandoc
        doc = pandoc.read(open(html_path, encoding="utf-8").read(), format="html")
        pandoc.write(doc, out_path, format="markdown")
    except Exception:
        with open(html_path, "r", encoding="utf-8") as f:
            content = f.read()
        import re
        text = re.sub(r"<[^>]+>", "", content)
        text = re.sub(r"\s+", " ", text).strip()
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(text)
    return out_path


# ─── PDF / 图片转换 ──────────────────────────────────────────
def pdf_to_png(pdf_path: str, out_dir: str, dpi: int = 150) -> List[str]:
    if not _deps["pymupdf"]:
        raise ImportError("PyMuPDF (fitz) 未安装，无法处理 PDF")
    import fitz
    doc = fitz.open(pdf_path)
    os.makedirs(out_dir, exist_ok=True)
    paths = []
    for i, page in enumerate(doc):
        mat = fitz.Matrix(dpi / 72, dpi / 72)
        pix = page.get_pixmap(matrix=mat)
        out = os.path.join(out_dir, f"page_{i+1:03d}.png")
        pix.save(out)
        paths.append(out)
    return paths


def pdf_to_jpg(pdf_path: str, out_dir: str, dpi: int = 150) -> List[str]:
    if not _deps["pymupdf"]:
        raise ImportError("PyMuPDF (fitz) 未安装，无法处理 PDF")
    import fitz
    doc = fitz.open(pdf_path)
    os.makedirs(out_dir, exist_ok=True)
    paths = []
    for i, page in enumerate(doc):
        mat = fitz.Matrix(dpi / 72, dpi / 72)
        pix = page.get_pixmap(matrix=mat)
        out = os.path.join(out_dir, f"page_{i+1:03d}.jpg")
        pix.save(out)
        paths.append(out)
    return paths


def image_to_pdf(image_path: str, out_path: str) -> str:
    if not _deps["pillow"]:
        raise ImportError("Pillow 未安装，无法处理图片")
    from PIL import Image
    img = Image.open(image_path)
    if img.mode != "RGB":
        img = img.convert("RGB")
    img.save(out_path, "PDF", resolution=150)
    return out_path


# ─── 批量处理 ────────────────────────────────────────────────
def batch_convert(
    file_paths: List[str],
    target_format: str,
    output_dir: str,
    ai_instruction: Optional[str] = None,
    plan: str = "FREE",
) -> Dict[str, Any]:
    """
    批量转换主入口
    plan: FREE | STD | PRO | MAX
    """
    os.makedirs(output_dir, exist_ok=True)

    results = {"success": [], "failed": [], "skipped": []}
    converted_count = 0

    for fp in file_paths:
        fname = Path(fp).name
        stem = Path(fp).stem
        ext = target_format.lower().strip(".")

        try:
            out_path = os.path.join(output_dir, f"{stem}.{ext}")

            src_ext = Path(fp).suffix.lower().strip(".")

            # ── 单个文件转换路由 ──────────────────────
            if src_ext == "csv":
                if ext in ("xlsx", "xls"):
                    csv_to_excel(fp, out_path)
                elif ext == "json":
                    csv_to_json(fp, out_path)
                else:
                    results["skipped"].append({"file": fname, "reason": f"CSV不支持转为 {target_format}"})
                    continue

            elif src_ext in ("xlsx", "xls"):
                if ext == "json":
                    excel_to_json(fp, out_path)
                elif ext == "csv":
                    excel_to_csv(fp, out_path)
                elif ext in ("png", "jpg", "jpeg"):
                    excel_to_png(fp, out_path)
                else:
                    results["skipped"].append({"file": fname, "reason": f"Excel不支持转为 {target_format}"})
                    continue

            elif src_ext == "json":
                if ext in ("xlsx", "xls"):
                    json_to_excel(fp, out_path)
                else:
                    results["skipped"].append({"file": fname, "reason": f"JSON不支持转为 {target_format}"})
                    continue

            elif src_ext == "md":
                if ext == "docx":
                    markdown_to_docx(fp, out_path)
                elif ext == "html":
                    markdown_to_html(fp, out_path)
                else:
                    results["skipped"].append({"file": fname, "reason": f"Markdown不支持转为 {target_format}"})
                    continue

            elif src_ext == "docx":
                if ext == "md":
                    docx_to_markdown(fp, out_path)
                else:
                    results["skipped"].append({"file": fname, "reason": f"DOCX不支持转为 {target_format}"})
                    continue

            elif src_ext == "html":
                if ext == "md":
                    html_to_markdown(fp, out_path)
                else:
                    results["skipped"].append({"file": fname, "reason": f"HTML不支持转为 {target_format}"})
                    continue

            elif src_ext == "pdf":
                if ext == "png":
                    imgs = pdf_to_png(fp, output_dir)
                    results["success"].append({"file": fname, "outputs": imgs})
                    converted_count += 1
                    continue
                elif ext in ("jpg", "jpeg"):
                    imgs = pdf_to_jpg(fp, output_dir)
                    results["success"].append({"file": fname, "outputs": imgs})
                    converted_count += 1
                    continue
                else:
                    results["skipped"].append({"file": fname, "reason": f"PDF不支持转为 {target_format}"})
                    continue

            elif src_ext in ("png", "jpg", "jpeg"):
                if ext == "pdf":
                    image_to_pdf(fp, out_path)
                else:
                    results["skipped"].append({"file": fname, "reason": f"图片不支持转为 {target_format}"})
                    continue

            elif src_ext == "txt":
                if ext == "csv":
                    txt_to_csv(fp, out_path)
                else:
                    results["skipped"].append({"file": fname, "reason": f"TXT不支持转为 {target_format}"})
                    continue

            else:
                results["skipped"].append({"file": fname, "reason": f"不支持的源格式: {src_ext}"})
                continue

            results["success"].append({"file": fname, "output": out_path})
            converted_count += 1

        except Exception as e:
            log.error(f"转换失败 {fname}: {e}")
            results["failed"].append({"file": fname, "error": str(e)})

    results["converted_count"] = converted_count
    return results


# ─── 套餐额度检查 ─────────────────────────────────────────────
def check_quota(
    plan: str,
    total_converted_count: int,
    remaining: int,
) -> Tuple[bool, str]:
    """检查是否还能继续转换，返回 (can_proceed, message)"""
    if plan == "FREE":
        if remaining <= 0:
            return False, "⚠️ 免费额度已用尽（共10个文件）。升级到标准版（¥9.9/月，每天100个）继续使用。"
        return True, f"免费额度剩余 {remaining} 个文件"
    elif plan in ("STD", "PRO", "MAX"):
        return True, "套餐额度充足"
    return False, "未识别的套餐"


# ─── 飞书卡片构建 ─────────────────────────────────────────────
def build_result_card(results: Dict[str, Any], plan: str) -> Dict[str, Any]:
    succ = len(results["success"])
    fail = len(results["failed"])
    skip = len(results["skipped"])

    succ_files = [r["file"] for r in results["success"]]
    fail_files = [f"{r['file']}: {r['error']}" for r in results["failed"]]

    title = "✅ 转换完成" if fail == 0 else "⚠️ 部分完成"
    color = "green" if fail == 0 else "orange"

    card = {
        "msg_type": "interactive",
        "card": {
            "header": {
                "title": {"tag": "plain_text", "text": f"📦 格式批量互转 | {title}"},
                "template": color
            },
            "elements": [
                {"tag": "markdown", "content": f"**成功:** {succ} 个  |  **失败:** {fail} 个  |  **跳过:** {skip} 个"},
                {"tag": "hr"},
            ]
        }
    }

    if succ_files:
        file_list = "\n".join([f"- {f}" for f in succ_files])
        card["card"]["elements"].append(
            {"tag": "markdown", "content": f"**成功文件:**\n{file_list}"}
        )
    if fail_files:
        err_list = "\n".join([f"- {f}" for f in fail_files])
        card["card"]["elements"].append(
            {"tag": "markdown", "content": f"**失败文件:**\n{err_list}"}
        )
    if plan == "FREE":
        card["card"]["elements"].append(
            {"tag": "note", "elements": [
                {"tag": "plain_text", "text": "💡 免费版共10个文件绝对计数，用完请升级。"}
            ]}
        )

    return card


# ─── CLI / 入口 ──────────────────────────────────────────────
def main():
    import argparse
    parser = argparse.ArgumentParser(description="batch-format-converter")
    parser.add_argument("--files", nargs="+", required=True, help="源文件路径列表")
    parser.add_argument("--format", required=True, help="目标格式，如 xlsx, json, pdf ...")
    parser.add_argument("--output-dir", default="/tmp/converter_output", help="输出目录")
    parser.add_argument("--plan", default="FREE", choices=["FREE", "STD", "PRO", "MAX"],
                        help="Subscription tier (default: FREE)")
    parser.add_argument("--api-key", default="",
                        help="91Skillhub API key for automatic tier verification")
    parser.add_argument("--total-count", type=int, default=0, help="免费版累计已转换数")
    parser.add_argument("--ai-instruction", default=None, help="AI自定义指令（Pro/Max）")
    args = parser.parse_args()

    # ── Tier 推断优先级 ──────────────────────────────────────────
    # 1. 优先用 --api-key 显式传入的 key 进行验证
    # 2. 其次用 OPENAI_API_KEY 环境变量
    # 3. 最终降级到 --plan CLI 参数（手动指定）
    import os
    api_key = args.api_key or os.environ.get("OPENAI_API_KEY", "")
    if api_key:
        verify_result = verify_token(api_key)
        if verify_result["valid"]:
            plan = verify_result["tier"]
            print(f"[INFO] Token verified: valid=True, tier={plan}, prefix={verify_result.get('prefix','')}")
        else:
            plan = "FREE"
            print(f"[WARN] Token invalid ({verify_result['error']}), using FREE tier")
    else:
        plan = args.plan
        print(f"[INFO] No API key, using CLI plan={plan}")

    # 额度检查
    remaining = 10 - args.total_count
    can_proceed, msg = check_quota(plan, args.total_count, remaining)
    print(f"[quota] {msg}")

    if not can_proceed:
        print("[ERROR] 额度不足，终止转换")
        sys.exit(1)

    print(f"[start] 转换 {len(args.files)} 个文件，目标格式: {args.format}")
    results = batch_convert(
        args.files,
        args.format,
        args.output_dir,
        ai_instruction=args.ai_instruction,
        plan=plan,
    )
    print(f"[done] 成功 {len(results['success'])}，失败 {len(results['failed'])}，跳过 {len(results['skipped'])}")
    print("[results]", json.dumps(results, ensure_ascii=False))


if __name__ == "__main__":
    main()
