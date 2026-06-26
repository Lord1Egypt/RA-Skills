# -*- coding: utf-8 -*-
"""Verify a tex2docx output document."""
import sys, re
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

if len(sys.argv) < 2:
    print("Usage: python verify.py <docx_path>")
    sys.exit(1)

doc = Document(sys.argv[1])
ns_m = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
ns_a = 'http://schemas.openxmlformats.org/drawingml/2006/main'

print(f"=== Verification Report ===")
print(f"Paragraphs: {len(doc.paragraphs)}")
print(f"Tables: {len(doc.tables)}")

# Equations
omml = sum(1 for p in doc.paragraphs if p._element.findall(f'.//{{{ns_m}}}oMath'))
print(f"OMML equations: {omml}")

# Images
img = sum(1 for p in doc.paragraphs if p._element.findall(f'.//{{{ns_a}}}blip'))
print(f"Images embedded: {img}")

# Captions
caps = sum(1 for p in doc.paragraphs if p.text.strip().startswith('Fig.'))
print(f"Figure captions: {caps}")

# Tables
for ti, t in enumerate(doc.tables):
    h = [c.text.strip()[:15] for c in t.rows[0].cells]
    print(f"  Table {ti+1}: {len(t.rows)}r x {len(t.columns)}c, header={h}")

# References
refs = sum(1 for p in doc.paragraphs if re.match(r'^\[b\d+\]', p.text.strip()))
print(f"References: {refs}")
print(f"REFERENCES title: {'YES' if any('REFERENCES' in p.text for p in doc.paragraphs) else 'NO'}")

# Content checks
full = ' '.join([p.text for p in doc.paragraphs])
for kw in ['CNN-BiLSTM', 'Abstract', 'Methodology', 'Conclusion']:
    print(f"  {'OK' if kw in full else 'MISS'}: {kw}")

# LaTeX residue
for bad in ['htbp', '\\begin{', '\\end{', '\\toprule']:
    if bad in full:
        print(f"  LaTeX residue: {bad}")

import os
print(f"\nFile size: {os.path.getsize(sys.argv[1])} bytes")
print("=== Done ===")
