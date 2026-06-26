#!/usr/bin/env python3
"""
单体配方投料表生成器
根据联合分析结果（formulation.json）自动生成 Word 投料表。

用法：
    python gen_monomer_formulation_table.py <整合结果目录> [--output <输出路径.docx>]

参数：
    整合结果目录   包含"*_按物种整合"子目录的联合分析结果目录
    --output       输出 Word 文件路径（默认: 整合结果目录/物种单体配方投料表.docx）

依赖：
    pip install python-docx

单体-氨基酸类别对照（固定）：
    Hydrophobic → BA（丙烯酸正丁酯，128.17 g/mol, ρ=0.90 g/mL）
    Nucleophilic → HEA（丙烯酸羟乙酯，116.12 g/mol, ρ=1.10 g/mL）
    Aromatic → PEA（丙烯酸苯氧乙酯，192.21 g/mol, ρ=1.10 g/mL）
    Amide → AAm（丙烯酰胺 0.18 g/mL溶液，71.08 g/mol）
    Acidic → CBEA（丙烯酸羧乙酯，144.13 g/mol, ρ=1.10 g/mL）
    Cationic → ATAC（80%质量分数水溶液，折算纯品193.67 g/mol, ρ≈1.15 g/mL）

固定参数（2026-05-11 更新）：
    单体总浓度: 1.0 M
    最终总体积: 1.5 mL（DMSO 补至该体积）
    交联剂: 不使用
    引发剂 2-oxoglutaric acid: 3.75 mg, MW=146.10（固体，记录质量）
    最终溶液中引发剂浓度约 17 mM
"""

import json, os, sys, argparse
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

# ========== 常量 ==========
MW = {
    "BA":   128.17,   # Hydrophobic
    "HEA":  116.12,   # Nucleophilic
    "PEA":  192.21,   # Aromatic
    "AAm":   71.08,   # Amide (固体，按质量)
    "CBEA": 144.13,   # Acidic
    "ATAC": 193.67,   # Cationic (80%水溶液)
}

# 液体单体密度 (g/mL) - 用于质量→体积转换
DENSITY = {
    "BA":   0.90,     # 丙烯酸正丁酯
    "HEA":  1.10,     # 丙烯酸羟乙酯
    "PEA":  1.10,     # 丙烯酸苯氧乙酯 (估算，与HEA相近)
    "CBEA": 1.10,     # 丙烯酸羧乙酯 (估算)
    "ATAC": 1.15,     # 80% ATAC水溶液 (水溶液密度)
}

CAT_TO_MONO = {
    "Hydrophobic": "BA",
    "Nucleophilic": "HEA",
    "Aromatic":     "PEA",
    "Amide":        "AAm",
    "Acidic":       "CBEA",
    "Cationic":     "ATAC",
}

# ========== 新参数（2026-05-11 更新）==========
MONOMER_CONC_M = 1.0      # 单体总浓度 (M)
FINAL_VOLUME_ML = 1.5     # 最终总体积 (mL)

# 根据浓度和体积计算总摩尔数
TOTAL_MOL = MONOMER_CONC_M * (FINAL_VOLUME_ML / 1000)  # 0.0015 mol

# AAm 现在按液体处理，浓度 0.18 g/mL → 密度 0.18 g/mL
AAm_DENSITY = 0.18  # g/mL

# ATAC 质量分数 (80% 水溶液)
ATAC_WF = 0.80

# ========== 交联剂与引发剂 ==========
USE_CROSSLINKER = False   # 不使用交联剂
INITIATOR_MASS_MG = 3.75  # 引发剂质量 (mg)
INITIATOR_MW = 146.10

# 引发剂摩尔数 = 质量 / MW (mmol)
INITIATOR_MMOL = INITIATOR_MASS_MG / INITIATOR_MW * 1000


def find_species_dirs(base_dir):
    """自动发现 *_按物种整合 子目录及其中的物种，或直接从整合结果目录读取"""
    results = []
    base_dir = Path(base_dir)
    
    # 情况1: 包含 *_按物种整合 子目录（旧结构）
    for entry in sorted(base_dir.iterdir()):
        if not entry.is_dir() or "按物种整合" not in entry.name:
            continue
        cat_name = entry.name.replace("_按物种整合", "")
        for species_dir in sorted(entry.iterdir()):
            if not species_dir.is_dir():
                continue
            fpath = species_dir / "氨基酸对分析" / "formulation.json"
            if not fpath.is_file():
                continue
            with open(fpath) as f:
                data = json.load(f)
            phi = data.get("category_phi", {})
            results.append({
                "category": cat_name,
                "species": species_dir.name,
                "phi": {cat: float(phi.get(cat, 0)) for cat in CAT_TO_MONO},
            })
    
    # 情况2: 直接是整合结果目录（新结构：物种名/氨基酸对分析/formulation.json）
    if not results:
        for species_dir in sorted(base_dir.iterdir()):
            if not species_dir.is_dir() or species_dir.name.startswith('_'):
                continue
            fpath = species_dir / "氨基酸对分析" / "formulation.json"
            if not fpath.is_file():
                continue
            with open(fpath) as f:
                data = json.load(f)
            phi = data.get("category_phi", {})
            results.append({
                "category": "整合结果",
                "species": species_dir.name,
                "phi": {cat: float(phi.get(cat, 0)) for cat in CAT_TO_MONO},
            })
    
    return results


def calc_monomer_amounts(phi, total_mol):
    """
    根据φ值百分比和总摩尔数计算各单体投料量
    返回: dict {单体名: (数值, 单位)}
    所有单体（包括AAm）→ 体积 (μL)
    """
    amounts = {}
    for cat, mono in CAT_TO_MONO.items():
        pct = phi.get(cat, 0)
        mol_i = total_mol * pct / 100.0
        mass_mg = mol_i * MW[mono] * 1000  # mg
        
        if mono == "AAm":
            volume_uL = mass_mg / AAm_DENSITY  # μL
            amounts[mono] = (volume_uL, "μL")
        elif mono == "ATAC":
            solution_mass_mg = mass_mg / ATAC_WF
            volume_uL = solution_mass_mg / DENSITY[mono]  # μL
            amounts[mono] = (volume_uL, "μL")
        else:
            volume_uL = mass_mg / DENSITY[mono]  # μL
            amounts[mono] = (volume_uL, "μL")
    return amounts


def generate_docx(all_species, outpath):
    total_mol = TOTAL_MOL
    initiator_mg = INITIATOR_MASS_MG
    initiator_mmol = INITIATOR_MMOL

    # 交联剂（可能不使用）
    if USE_CROSSLINKER:
        crosslinker_mg = CROSSLINKER_MMOL * CROSSLINKER_MW
        crosslinker_uL = crosslinker_mg / CROSSLINKER_DENSITY
    else:
        crosslinker_mg = 0
        crosslinker_uL = 0

    doc = Document()

    # 横向页面
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        new_w, new_h = section.page_height, section.page_width
        section.page_width, section.page_height = new_w, new_h
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # 标题
    title = doc.add_heading('物种单体配方投料表', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 说明
    info = doc.add_paragraph()
    crosslinker_info = f'交联剂: {crosslinker_mg:.2f} mg = {crosslinker_uL:.2f} μL | ' if USE_CROSSLINKER else '交联剂: 无 | '
    info_run = info.add_run(
        f'单体总浓度: {MONOMER_CONC_M} M | 最终体积: {FINAL_VOLUME_ML} mL | '
        f'{crosslinker_info}'
        f'引发剂: {initiator_mg:.2f} mg ({initiator_mmol:.4f} mmol) | DMSO: 补至 {FINAL_VOLUME_ML} mL'
    )
    info_run.font.size = Pt(9)

    # 单体对照表
    doc.add_paragraph()
    doc.add_heading('单体-氨基酸类别对照', level=2)
    ref_table = doc.add_table(rows=7, cols=6, style='Table Grid')
    ref_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    ref_headers = ['氨基酸类别', '单体缩写', '单体全称', '分子式', 'M (g/mol)', '密度/浓度']
    for i, h in enumerate(ref_headers):
        cell = ref_table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    ref_data = [
        ['Hydrophobic (疏水性)', 'BA', 'Butyl Acrylate', 'C₇H₁₂O₂', '128.17', '0.90 g/mL'],
        ['Nucleophilic (亲核性)', 'HEA', '2-Hydroxyethyl Acrylate', 'C₅H₈O₃', '116.12', '1.10 g/mL'],
        ['Aromatic (芳香性)', 'PEA', '2-Phenoxyethyl Acrylate', 'C₁₁H₁₂O₃', '192.21', '1.10 g/mL'],
        ['Amide (酰胺类)', 'AAm', 'Acrylamide (0.18 g/mL)', 'C₃H₅NO', '71.08', '0.18 g/mL'],
        ['Acidic (酸性)', 'CBEA', '2-Carboxyethyl Acrylate', 'C₆H₈O₄', '144.13', '1.10 g/mL'],
        ['Cationic (阳离子性)', 'ATAC (80% w.f.)', 'Acryloyloxyethyl Trimethyl Ammonium Chloride (80% aq.)', 'C₈H₁₆ClNO₂', '193.67', '1.15 g/mL'],
    ]
    for ri, rd in enumerate(ref_data):
        for ci, val in enumerate(rd):
            cell = ref_table.rows[ri + 1].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(8)

    doc.add_paragraph()

    # 按分类生成表格
    headers = ['序号', '物种名', 'BA\n(μL)', 'HEA\n(μL)', 'PEA\n(μL)',
               'AAm\n(μL)', 'CBEA\n(μL)', 'ATAC\n(μL)',
               '引发剂\n(mg)', 'DMSO\n(μL)']

    cat_order = []
    seen = set()
    for sp in all_species:
        if sp["category"] not in seen:
            cat_order.append(sp["category"])
            seen.add(sp["category"])

    global_idx = 0
    for cat_key in cat_order:
        species_in_cat = [s for s in all_species if s["category"] == cat_key]
        if not species_in_cat:
            continue

        doc.add_heading(f'{cat_key}（{len(species_in_cat)}物种）', level=2)
        table = doc.add_table(rows=len(species_in_cat) + 1, cols=len(headers), style='Table Grid')
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 表头
        for ci, h in enumerate(headers):
            cell = table.rows[0].cells[ci]
            cell.text = h
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(8)
            shading = cell._element.get_or_add_tcPr()
            elm = shading.makeelement(qn('w:shd'), {qn('w:fill'): '4472C4', qn('w:val'): 'clear'})
            shading.append(elm)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)

        # 数据行
        for ri, sp in enumerate(species_in_cat):
            global_idx += 1
            amounts = calc_monomer_amounts(sp["phi"], TOTAL_MOL)
            row = table.rows[ri + 1]
            def fmt(val, unit):
                return f'{int(round(val))}' if unit == 'μL' else f'{val:.1f}'
            
            total_liquid_uL = (
                amounts["BA"][0] + amounts["HEA"][0] + amounts["PEA"][0] +
                amounts["AAm"][0] + amounts["CBEA"][0] + amounts["ATAC"][0]
            )
            if USE_CROSSLINKER:
                total_liquid_uL += crosslinker_uL
            
            dmso_uL = FINAL_VOLUME_ML * 1000 - total_liquid_uL
            if dmso_uL < 0:
                dmso_uL = 0
            
            values = [
                str(global_idx), sp["species"],
                fmt(*amounts["BA"]), fmt(*amounts["HEA"]), fmt(*amounts["PEA"]),
                fmt(*amounts["AAm"]), fmt(*amounts["CBEA"]), fmt(*amounts["ATAC"]),
                f'{initiator_mg:.3f}', f'{int(round(dmso_uL))}',
            ]
            for ci, val in enumerate(values):
                cell = row.cells[ci]
                cell.text = val
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci == 1 else WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.size = Pt(8)
                        if ci == 1:
                            run.italic = True
            if ri % 2 == 1:
                for cell in row.cells:
                    shading = cell._element.get_or_add_tcPr()
                    elm = shading.makeelement(qn('w:shd'), {qn('w:fill'): 'D9E2F3', qn('w:val'): 'clear'})
                    shading.append(elm)

        doc.add_paragraph()

    doc.save(outpath)
    print(f"✅ 已生成: {outpath}（{global_idx} 个物种）")
    return global_idx


def main():
    parser = argparse.ArgumentParser(description='生成物种单体配方投料表 Word 文档（单体总浓度 1.0 M，终体积 1.5 mL）')
    parser.add_argument('result_dir', help='整合分析结果目录（含 *_按物种整合 子目录）')
    parser.add_argument('--output', '-o', help='输出 docx 路径（默认: 结果目录/物种单体配方投料表.docx）')
    args = parser.parse_args()

    if not os.path.isdir(args.result_dir):
        print(f"❌ 目录不存在: {args.result_dir}")
        sys.exit(1)

    outpath = args.output or os.path.join(args.result_dir, '物种单体配方投料表.docx')

    all_species = find_species_dirs(args.result_dir)
    if not all_species:
        print(f"❌ 未找到任何物种配方数据（formulation.json）")
        sys.exit(1)

    print(f"📊 共发现 {len(all_species)} 个物种配方")
    generate_docx(all_species, outpath)


if __name__ == '__main__':
    main()
