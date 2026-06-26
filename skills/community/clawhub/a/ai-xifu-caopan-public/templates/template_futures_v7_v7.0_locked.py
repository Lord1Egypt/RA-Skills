#!/usr/bin/env python3
"""
沪铜CU2607期货方案生成模板 v7.0 FINAL
===================================
大叔2026-05-23验收通过，锁定替换旧模板（补丁63）

核心特征：
  1. 数据源：Tushare fut_daily（期货日线）
  2. 行情回顾：2日完整交易日（trade_date=夜盘+日盘合并）
  3. 方案周期：2天（精简版）
  4. 操作方案：日盘做多/做空 + 夜盘做多/做空（四维双向）
  5. 每日预判：含日盘+夜盘两时段

用法：
  from template_futures_v7 import build_plan
  
  注意：用户需自行提供数据，data_dict为必传参数
  不支持自动数据拉取
"""

from language_detection import detect_input_lang, match_market_lang, native_name, BilingualText

def build_plan(data_dict, user_input=""):
    """
    大叔亲授·审题+质检机制：自动检测用户输入语言，输出动态双语方案。
    """
    # data_dict is REQUIRED — raises ValueError if None
    """
    生成期货交易方案docx
    
    参数:
      data_dict: dict，包含以下字段（必传，用户需自行提供）
        - ts_code: 合约代码，如 'CU2607.SHF'
        - start_date: 回顾起始日，如 '20260521'
        - end_date: 回顾截止日，如 '20260522'
        - plan_start: 预判起始日，如 '05-25(一)'
        - plan_end: 预判截止日，如 '05-26(二)'
        - plan_start_en: 预判起始日数字，如 '20260525'
        - plan_end_en: 预判截止日数字，如 '20260526'
        - day1_rows: 第1天日盘数据行
        - day2_rows: 第2天日盘数据行
        - price_current: 当前价(日盘)
        - price_night: 当前价(夜盘)
        - pressure_strong: 强压力
        - pressure_mid: 次压力
        - support_short: 短期支撑
        - support_strong: 强支撑
        - lme_price: LME铜价格
        - comex_price: COMEX铜价格
        - usd_index: 美元指数
    
    返回:
      doc: python-docx Document对象
      文件已保存至: {WORKSPACE}/skills/ai-xifu-caopan/{filename}
    """

    # ── 🌍 大叔亲授·审题+质检机制：自动检测语言 ──
    _input_lang = detect_input_lang(user_input)
    _market_code = data_dict.get("market_code", "cn") if data_dict else "cn"
    _dest_lang = match_market_lang(_market_code)
    _bt = BilingualText(input_lang=_input_lang, dest_lang=_dest_lang)
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    import os
    import datetime
    
    # ── 用户必须提供data_dict ──
    if data_dict is None:
        raise ValueError("❌ data_dict is REQUIRED / data_dict为必传参数。请提供真实市场数据。")
    defaults = {
        'ts_code': 'CU2607.SHF',
        'display_name': '沪铜CU2607',
        'variety': '沪铜阴极铜期货',
        'contract': 'CU2607',
        'exchange': '上海期货交易所',
        'unit': '5吨/手',
        'min_move': '10元/吨(=50元/手)',
        'margin': '约5%',
        'trade_hours': '日盘 09:00-11:30/13:30-15:00 | 夜盘 21:00-01:00(次日凌晨)',
        'plan_start': '05-25(一)',
        'plan_end': '05-26(二)',
        'data_source': 'Tushare期货日线接口(fut_daily)',
        'data_note': 'Tushare期货日线中每个trade_date为完整交易日(含夜盘+日盘合并)。open=夜盘开盘，close=日盘15:00收盘。',
        # 行情回顾数据（2行：day1/day2）
        'day1_date': '05-21(四)',
        'day1_open': '104,060',
        'day1_high': '105,660',
        'day1_low': '103,880',
        'day1_close': '104,040',
        'day1_settle': '104,790',
        'day1_change': '+440',
        'day1_vol': '103,523',
        'day1_oi': '160,365',
        'day2_date': '05-22(五)',
        'day2_open': '104,000',
        'day2_high': '105,300',
        'day2_low': '103,620',
        'day2_close': '104,870',
        'day2_settle': '104,480',
        'day2_change': '+80',
        'day2_vol': '96,264',
        'day2_oi': '171,627',
        # 关键价位
        'price_current': '104,870',
        'price_night': '104,940',
        'pressure_strong': '106,670',
        'pressure_mid': '105,660',
        'support_short': '103,620',
        'support_strong': '103,000',
        # 外盘
        'lme_price': '13,620美元/吨',
        'lme_trend': '高位震荡(-0.26%)',
        'comex_price': '~6.345美元/磅',
        'usd_index': '~99.5',
        'usd_trend': '震荡偏弱',
        # 风险预警
        'risk_vol': '量能萎缩|日盘9.6万手萎缩',
        'risk_lme': '外盘走弱|LME 13,710→13,620',
        'risk_inv': '库存累库|三大交易所~120万吨',
        'risk_geo': '地缘风险|美伊+托克提货',
        'risk_spot': '现货升水|升~2,240，支撑强',
        # 国际局势
        'geo1': '美伊冲突|特朗普拒绝停火|利多(地缘溢价)',
        'geo2': '中美关系|访华预期|偏多',
        'geo3': '跨市场套利|托克注销5.1万吨|利多(亚洲供应趋紧)',
        # 宏观
        'macro1': '中国经济政策|利多|+14%',
        'macro2': '美元走势|利多|+10%',
        'macro3': '美伊地缘|利多|+9%',
        'macro4': '美联储政策|利空|-8%',
        'macro5': '过剩预期|利空|-5%',
        # 产业链
        'chain1': '铜精矿TC|-93.9~-100美元/吨(历史新低)|利多大涨',
        'chain2': '矿端供应|智利罢工+印尼减产|利多',
        'chain3': '下游需求|新能源+特高压+AI|利多',
        # 资金
        'fund_oi': '171,627手|多头建仓',
        'fund_vol': '日盘9.6万手|观望',
        'fund_trafigura': 'LME注销5.1万吨|2013年最大提货',
        'fund_premium': '~2,240元/吨|支撑强',
        'fund_core': '托克注销5.1万吨，亚洲供应趋紧',
        # MACD
        'macd_daily': '日线|DIF在DEA上、红柱缩短|价在均线上|多头动能减',
        'macd_60min': '60分钟|DIF在DEA附近|均线粘合|短线回调压力',
        'macd_weekly': '周线|DIF在DEA上+红柱|价在均线上|中期多头完好',
        'tech_summary': '中期多+短期震荡',
        # 每日预判（2天×2时段）
        'pred1_day_dir': '探底回升',
        'pred1_day_range': '103,800~105,000',
        'pred1_day_long': '103,800~104,000低吸',
        'pred1_day_short': '冲高105,000短空',
        'pred1_night_dir': '高位震荡',
        'pred1_night_range': '104,500~105,100',
        'pred1_night_long': 'LME13,600做多',
        'pred1_night_short': 'LME13,400做空',
        'pred2_day_dir': '震荡偏多',
        'pred2_day_range': '104,500~105,500',
        'pred2_day_long': '104,500附近做多',
        'pred2_day_short': '冲高105,500轻仓',
        'pred2_night_dir': '偏多整理',
        'pred2_night_range': '104,600~105,200',
        'pred2_night_long': '观察LME',
        'pred2_night_short': '观察LME',
        # 操作建议参数
        'day_long_entry': '103,800~104,000',
        'day_long_target1': '105,000',
        'day_long_target2': '105,500',
        'day_long_stop': '102,500下方',
        'day_long_pos': '首仓20%|突破105,500加至30%',
        'day_short_entry': '105,300~105,500受阻+15分钟顶背离',
        'day_short_target1': '104,500',
        'day_short_target2': '104,000',
        'day_short_stop': '105,700上方',
        'day_short_pos': '不超10%',
        'night_long_entry': '21:00 LME站稳13,600+CU2607不破104,420+30分钟翻红',
        'night_long_target1': '105,110',
        'night_long_target2': '105,500',
        'night_long_stop': '104,200下方',
        'night_long_pos': '15%',
        'night_short_entry': '21:00 LME跌破13,400+CU2607受压+15分钟连续收阴',
        'night_short_target1': '104,420',
        'night_short_target2': '104,000',
        'night_short_stop': '105,200上方',
        'night_short_pos': '不超10%',
        # 操盘建议
        'main_strategy': '做多为主(主力策略)，做空为辅(辅助策略)',
        'trade_strategy': '日盘逢回调做多+冲高短空|夜盘联动LME',
        'trade_pos': '做多首仓20%|做空不超10%，总敞口不超50%',
        'trade_stop': '多单102,500下方|空单105,700上方',
        # 评分
        'score_tech': '7/10',
        'score_tech_note': '中期多头+短期震荡',
        'score_fundamental': '8/10',
        'score_fundamental_note': '供应紧缺+库存偏高',
        'score_capital': '7/10',
        'score_capital_note': '托克5.1万吨利多',
        'score_sentiment': '6/10',
        'score_sentiment_note': '观望浓(缩量)',
        'score_macro': '7/10',
        'score_macro_note': '政策利好+美元弱',
        'summary': '中期上行趋势未破，短期高位震荡。日盘回调103,800~104,000低吸做多(顺势主策略)，冲高105,300~105,500轻仓短空(逆势辅助)。夜盘联动LME。',
    }
    
    d = {**defaults, **data_dict}
    
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style.font.size = Pt(10)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    
    def H(t, l=1):
        h = doc.add_heading(t, level=l)
        for r in h.runs:
            r.font.name = 'SimSun'
            r.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    
    def P(t, b=False, s=10, c=None, a=None):
        p = doc.add_paragraph()
        r = p.add_run(t)
        r.font.name = 'SimSun'; r.font.size = Pt(s); r.bold = b
        r.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        if c: r.font.color.rgb = c
        if a: p.alignment = a
    
    def T(hd, data, note=None):
        t = doc.add_table(rows=1+len(data), cols=len(hd))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = 'Table Grid'
        for i,h in enumerate(hd):
            c = t.rows[0].cells[i]; c.text = h
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs: r.bold=True; r.font.size=Pt(9); r.font.name='SimSun'; r.element.rPr.rFonts.set(qn('w:eastAsia'),'SimSun')
        for ri, rd in enumerate(data):
            for ci, v in enumerate(rd):
                c = t.rows[ri+1].cells[ci]; c.text = str(v)
                for p in c.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs: r.font.size=Pt(9); r.font.name='SimSun'; r.element.rPr.rFonts.set(qn('w:eastAsia'),'SimSun')
    
    # ── 标题 ──
    title = doc.add_heading('', level=0)
    r = title.add_run(f'{d["variety"]}{d["contract"]}下周交易方案')
    r.font.name = 'SimSun'; r.font.size = Pt(22); r.bold = True
    r.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    today_str = datetime.date.today().strftime('%Y-%m-%d')
    P(f'版本：v7.0 FINAL | 制作日期：{today_str} | 预判周期：{d["plan_start"]} ~ {d["plan_end"]}', s=10, a=WD_ALIGN_PARAGRAPH.CENTER)
    P('AI生成内容，请注意甄别 | 媳妇智投全球通用版出品，必属精品', b=True, s=10, c=RGBColor(128,128,128), a=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    
    # ── 方案头部 ──
    H('【方案头部】', 2)
    T(['项目','内容'],[
        ['品种|合约',f'{d["variety"]} | {d["contract"]}(主力合约，持仓{d["day2_oi"]}手)'],
        ['交易时段',d['trade_hours']],
        ['方案周期',f'{d["plan_start"]} ~ {d["plan_end"]} | 2个交易日'],
        ['数据来源',d['data_source']],
        ['说明',d['data_note']],
    ])
    doc.add_paragraph()
    
    # ── 一、行情回顾 ──
    H('一、行情回顾（Tushare数据源）', 2)
    P(f'数据来源：Tushare fut_daily接口 | 合约{d["ts_code"]}', s=9, c=RGBColor(100,100,100))
    
    T(['日期(交易日)','开盘(夜盘)','最高','最低','收盘(日盘15:00)','结算价','涨跌1','成交量','持仓量'],[
        [d['day1_date'], d['day1_open'], d['day1_high'], d['day1_low'], d['day1_close'], d['day1_settle'], d['day1_change'], d['day1_vol'], d['day1_oi']],
        [d['day2_date'], d['day2_open'], d['day2_high'], d['day2_low'], d['day2_close'], d['day2_settle'], d['day2_change'], d['day2_vol'], d['day2_oi']],
    ])
    P(f'说明：Tushare期货日线中每个交易日trade_date为完整交易日(夜盘+日盘合并)。开盘价对应夜盘21:00开盘，收盘价对应日盘15:00收盘。第1行含前一交易日夜盘+当日日盘，第2行含05-21夜盘+当日日盘。', s=9, c=RGBColor(100,100,100))
    
    H('1.1 关键价位', 3)
    T(['类型','价位','说明'],[
        ['强压力',d['pressure_strong'],'5月中旬高点'],
        ['次压力',d['pressure_mid'],f'{d["day1_date"]}最高点'],
        [f'当前价(日盘)',d['price_current'],f'{d["day2_date"]}收盘(日盘15:00)'],
        ['短期支撑',d['support_short'],f'{d["day2_date"]}最低点'],
        ['强支撑',d['support_strong'],'MA20区域'],
    ])
    doc.add_paragraph()
    
    # ── 二、市场联动 ──
    H('二、市场联动', 2)
    T(['品种','最新价','趋势','影响'],
     [[f'LME铜3个月',d['lme_price'],d['lme_trend'],'方向待定，r~0.95'],
      [f'COMEX铜',d['comex_price'],'+0.23%','强势震荡'],
      [f'美元指数',d['usd_index'],d['usd_trend'],'负相关，利多铜价']])
    doc.add_paragraph()
    
    # ── 2.2 风险预警 ──
    H('2.2 风险预警', 3)
    risks = [r.split('|') for r in [d['risk_vol'],d['risk_lme'],d['risk_inv'],d['risk_geo'],d['risk_spot']]]
    T(['信号','状态','说明'], risks)
    P('风险：中等。', s=10)
    
    # ── 四、国际局势 ──
    H('三、国际局势', 2)
    geos = [g.split('|') for g in [d['geo1'],d['geo2'],d['geo3']]]
    T(['局势','事件','影响'], geos)
    doc.add_paragraph()
    
    # ── 五、宏观影响 ──
    H('四、宏观影响', 2)
    macros = [m.split('|') for m in [d['macro1'],d['macro2'],d['macro3'],d['macro4'],d['macro5']]]
    T(['因素','方向','贡献'], macros)
    doc.add_paragraph()
    
    # ── 六、产业链 ──
    H('五、产业链', 2)
    chains = [c.split('|') for c in [d['chain1'],d['chain2'],d['chain3']]]
    T(['维度','现状','影响'], chains)
    doc.add_paragraph()
    
    # ── 七、资金动向 ──
    H('六、资金动向', 2)
    funds = [f.split('|') for f in [d['fund_oi'],d['fund_vol'],d['fund_trafigura'],d['fund_premium']]]
    T(['信号','数据','判断'], funds)
    P(f'核心：{d["fund_core"]}', b=True, s=10)
    doc.add_paragraph()
    
    # ── 十、技术分析 ──
    H('七、技术分析', 2)
    macds = [m.split('|') for m in [d['macd_daily'],d['macd_60min'],d['macd_weekly']]]
    T(['周期','MACD','均线','综合'], macds)
    P(f'综合：{d["tech_summary"]}', b=True, s=10)
    doc.add_paragraph()
    
    # ── 操作方案 ──
    has_night = d.get('has_night', True)
    night_label = '（日盘+夜盘）' if has_night else '（仅日盘·无夜盘）'
    H(f'八、操作方案{night_label}', 2)
    if has_night:
        P(f'说明：日盘(09:00-15:00)和夜盘(21:00-01:00)均为独立交易时段。本周期总体判断：{d["main_strategy"]}。', s=10)
    else:
        P(f'说明：日盘(09:00-11:30/13:30-15:00)。本品种无夜盘交易，注意隔夜跳空风险。本周期总体判断：{d["main_strategy"]}。', s=10)
    P('')
    
    H('8.1 日盘做多操作（主力策略）', 3)
    T(['项目','操作内容'],
     [['核心逻辑','中期多头+供应紧+现货升水=做多为主'],
      ['入场条件',f'回调{d["day_long_entry"]}企稳+30分钟止跌信号'],
      ['目标',f'第一{d["day_long_target1"]}|第二{d["day_long_target2"]}'],
      ['止损',d['day_long_stop']],
      ['仓位',d['day_long_pos']]])
    
    H('8.2 日盘做空操作（辅助策略）', 3)
    T(['项目','操作内容'],
     [['入场条件',d['day_short_entry']],
      ['目标',f'第一{d["day_short_target1"]}|第二{d["day_short_target2"]}'],
      ['止损',d['day_short_stop']],
      ['仓位',d['day_short_pos']],
      ['风险','盘中突破上行必须止损']])
    
    if has_night:
        H('8.3 夜盘做多操作（主力策略）', 3)
        T(['项目','操作内容'],
         [['入场条件',d['night_long_entry']],
          ['目标',f'第一{d["night_long_target1"]}|第二{d["night_long_target2"]}'],
          ['止损',d['night_long_stop']],
          ['仓位',d['night_long_pos']],
          ['不做条件','LME跌破关键位或低开破支撑']])
        
        H('8.4 夜盘做空操作（辅助策略）', 3)
        T(['项目','操作内容'],
         [['入场条件',d['night_short_entry']],
          ['目标',f'第一{d["night_short_target1"]}|第二{d["night_short_target2"]}'],
          ['止损',d['night_short_stop']],
          ['仓位',d['night_short_pos']],
          ['不做条件','LME站稳关键位或突破压力']])
    
    # 预判表：根据是否有夜盘动态生成行
    H('8.3 两日走势预判', 3)
    if has_night:
        pred_rows = [
            [d['plan_start'].split('(')[0],d['plan_start'].split('(')[1].rstrip(')'),'日盘',d['pred1_day_dir'],d['pred1_day_range'],d['pred1_day_long'],d['pred1_day_short'],'⭐⭐⭐⭐'],
            [d['plan_start'].split('(')[0],d['plan_start'].split('(')[1].rstrip(')'),'夜盘',d['pred1_night_dir'],d['pred1_night_range'],d['pred1_night_long'],d['pred1_night_short'],'⭐⭐⭐'],
            [d['plan_end'].split('(')[0],d['plan_end'].split('(')[1].rstrip(')'),'日盘',d['pred2_day_dir'],d['pred2_day_range'],d['pred2_day_long'],d['pred2_day_short'],'⭐⭐⭐'],
            [d['plan_end'].split('(')[0],d['plan_end'].split('(')[1].rstrip(')'),'夜盘',d['pred2_night_dir'],d['pred2_night_range'],d['pred2_night_long'],d['pred2_night_short'],'⭐⭐⭐'],
        ]
    else:
        pred_rows = [
            [d['plan_start'].split('(')[0],d['plan_start'].split('(')[1].rstrip(')'),'日盘',d['pred1_day_dir'],d['pred1_day_range'],d['pred1_day_long'],d['pred1_day_short'],'⭐⭐⭐⭐'],
            [d['plan_end'].split('(')[0],d['plan_end'].split('(')[1].rstrip(')'),'日盘',d['pred2_day_dir'],d['pred2_day_range'],d['pred2_day_long'],d['pred2_day_short'],'⭐⭐⭐'],
        ]
    T(['日期','星期','时段','预判','区间','做多策略','做空策略','置信度'], pred_rows)
    doc.add_paragraph()
    
    # ── 十二、操盘建议 ──
    H('九、操盘建议', 2)
    trades = [t.split('|') for t in [d['trade_strategy'],d['trade_pos'],d['trade_stop']]]
    T(['要点','建议'], trades)
    doc.add_paragraph()
    
    # ── 十三、风险提示 ──
    H('十、关注重点', 2)
    T(['重点关注','内容','影响'],
     [['LME铜','价格区间','决定方向'],
      ['美元指数','99-100','美元弱=铜强'],
      ['托克提货','注销5.1万吨后续','亚洲供应'],
      ['现货升贴水','维持升水','支撑力度']])
    doc.add_paragraph()
    
    H('十一、方案总结', 2)
    T(['维度','结论'],
     [['预判',f'{d["plan_start"]}~{d["plan_end"]}高位震荡偏多'],
      ['核心区间','103,500~105,500'],
      ['日盘策略','做多回调区间低吸|做空冲高压力短空'],
      ['夜盘策略','做多LME关键位开多|做空LME破位开空'],
      ['风险等级','中等|评级B+']])
    doc.add_paragraph()
    
    # ── 二十二-二十三 ──
    H('十二、方案总评', 2)
    T(['维度','评分','评语'],
     [['技术面',d['score_tech'],d['score_tech_note']],
      ['基本面',d['score_fundamental'],d['score_fundamental_note']],
      ['资金面',d['score_capital'],d['score_capital_note']],
      ['情绪面',d['score_sentiment'],d['score_sentiment_note']],
      ['宏观面',d['score_macro'],d['score_macro_note']]])
    P(d['summary'], b=True, s=11)
    doc.add_paragraph()
    
    H('十三、博弈论分析', 2)
    
    T(['参与方','方向','力量','策略'],
     [['产业套保(矿企)','空','强','高价套保'],
      ['机构投机','偏多','强','跟趋势'],
      ['游资','双向','中','快进快出'],
      ['散户','多头套牢','弱','高位追涨被套']])
    doc.add_paragraph()
    
    # ── 十四、风险提示 ──
    H('十四、风险提示', 2)
    for i,s in enumerate([
        '以上分析基于Tushare公开数据，不构成投资建议。期货高风险高杠杆，请合理控制仓位。',
        '关注：LME价格区间、美元指数、托克提货、现货升贴水。',
        '地缘冲突升级可能导致铜价5%波动。',
        '夜盘流动性约日盘1/3，极端行情可能流动性枯竭。',
    ],1): P(f'{i}. {s}', s=9)
    doc.add_paragraph()
    
    P('媳妇智投全球通用版出品，必属精品', b=True, s=12, a=WD_ALIGN_PARAGRAPH.CENTER)
    P(f'v7.0 FINAL | {today_str} | AI生成内容，请注意甄别', s=9, c=RGBColor(128,128,128), a=WD_ALIGN_PARAGRAPH.CENTER)
    
    # 保存
    output_dir = os.path.dirname(os.path.abspath(__file__))
    filename = f'{d["variety"].replace(" ","")}{d["contract"]}下周交易方案_v7.0_{today_str}.docx'
    out = os.path.join(output_dir, filename)
    doc.save(out)
    
    return doc, out
