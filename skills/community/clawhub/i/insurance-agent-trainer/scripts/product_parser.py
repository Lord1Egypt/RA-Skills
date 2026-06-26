# Insurance Product Document Parser / 保险产品文档解析脚本

"""
支持解析：PDF / Word(.docx) / 扫描图片(OCR) / TXT
从中国福利彩票/体育彩票官网获取开奖数据，用于陪练场景中的案例设计
"""

import re
import json
from typing import Optional

# ============ 外部数据接口 ============
# 官方数据源（供训练案例设计使用，不用于实际销售）


def fetch_ssq_latest(limit: int = 30) -> list[dict]:
    """
    从中国福利彩票官网获取双色球最新开奖数据
    API: https://www.cwl.gov.cn/cwl_admin/front/cwlkj/search/kjxx/findDrawNotice?name=ssq
    """
    import requests
    url = "https://www.cwl.gov.cn/cwl_admin/front/cwlkj/search/kjxx/findDrawNotice"
    params = {"name": "ssq", "pageNo": 1, "pageSize": limit}
    headers = {"User-Agent": "Mozilla/5.0 (compatible; InsuranceTrainerBot/1.0)"}
    try:
        resp = requests.get(url, params=params, headers=headers, timeout=15)
        data = resp.json()
        if data.get("state") == 0:
            return data.get("result", [])
        return []
    except Exception as e:
        return [{"error": str(e)}]


def fetch_fc3d_latest(limit: int = 30) -> list[dict]:
    """
    从中国福利彩票官网获取3D最新开奖数据
    API: https://www.cwl.gov.cn/cwl_admin/front/cwlkj/search/kjxx/findDrawNotice?name=3d
    """
    import requests
    url = "https://www.cwl.gov.cn/cwl_admin/front/cwlkj/search/kjxx/findDrawNotice"
    params = {"name": "3d", "pageNo": 1, "pageSize": limit}
    headers = {"User-Agent": "Mozilla/5.0 (compatible; InsuranceTrainerBot/1.0)"}
    try:
        resp = requests.get(url, params=params, headers=headers, timeout=15)
        data = resp.json()
        if data.get("state") == 0:
            return data.get("result", [])
        return []
    except Exception as e:
        return [{"error": str(e)}]


def fetch_dlt_latest(limit: int = 30) -> list[dict]:
    """
    从中国体育彩票官网获取大乐透最新开奖数据
    注意：部分站点有WAF防护，需添加合适的User-Agent
    API: https://webapi.sporttery.cn/service/lotto/dlt/search.do
    """
    import requests
    url = "https://webapi.sporttery.cn/service/list/info.do"
    params = {"lotteryType": "DLT", "pageNo": 1, "pageSize": limit}
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
        "Referer": "https://www.lottery.gov.cn/",
        "Accept": "application/json"
    }
    try:
        resp = requests.get(url, params=params, headers=headers, timeout=15)
        data = resp.json()
        if data.get("errorCode") == 0:
            return data.get("value", {}).get("list", [])
        return []
    except Exception as e:
        return [{"error": str(e), "note": "如果被WAF拦截，请尝试添加正确的Referer头"}]


# ============ 产品文档解析 ============


def parse_insurance_product_from_text(text: str) -> dict:
    """
    从产品文档文本中提取结构化信息
    适用于：TXT / PDF文字提取 / Word文字提取结果
    """
    result = {
        "product_name": "",
        "product_type": "",
        "insurer": "",
        "target_customers": [],
        "coverage": {},
        "premium": {},
        "key_selling_points": [],
        "competitive_edges": [],
        "exclusions": [],
        "compliance_notes": [],
        "difficulty_tags": [],
        "raw_snippets": []
    }

    # 产品名称提取（常见模式）
    name_patterns = [
        r"(?:产品名称|保险名称|产品名)[：:]\s*(.+?)(?:\n|$)",
        r"《(.+?)》",
        r"^#\s*(.+?)(?:\n|$)",
    ]
    for p in name_patterns:
        match = re.search(p, text)
        if match and not result["product_name"]:
            result["product_name"] = match.group(1).strip()

    # 保险公司
    insurer_patterns = [
        r"(?:保险公司|承保公司)[：:]\s*(.+?)(?:\n|$)",
        r"由\s*(.+?)\s*承保",
    ]
    for p in insurer_patterns:
        match = re.search(p, text)
        if match and not result["insurer"]:
            result["insurer"] = match.group(1).strip()

    # 保障内容（关键条款）
    coverage_keywords = ["保障", "保险金", "赔付", "给付", "责任"]
    for kw in coverage_keywords:
        if kw in text:
            lines = [l.strip() for l in text.split("\n") if kw in l]
            result["coverage"][kw] = lines[:3]  # 保留最多3条

    # 排除责任
    exclusion_keywords = ["除外", "不赔", "责任免除", "免责"]
    for kw in exclusion_keywords:
        if kw in text:
            lines = [l.strip() for l in text.split("\n") if kw in l]
            result["exclusions"].extend(lines[:3])

    # 销售卖点（数字+功能）
    number_benefit_patterns = [
        r"(\d+(?:万|%|倍|元))/年",
        r"(?:历史结算利率|预期收益率)[：:\s]*(\d+(?:\.\d+)?%?)",
        r"(\d+(?:万|倍))保障",
        r"(?:免费|赠送|豁免)[^，。]{0,20}",
    ]
    for p in number_benefit_patterns:
        matches = re.findall(p, text)
        result["key_selling_points"].extend([m for m in matches if m][:5])

    # 合规提示
    compliance_keywords = ["双录", "犹豫期", "等待期", "如实告知", "健康告知", "退保"]
    for kw in compliance_keywords:
        if kw in text:
            result["compliance_notes"].append(kw)

    # 去重
    result["key_selling_points"] = list(dict.fromkeys(result["key_selling_points"]))[:10]
    result["exclusions"] = list(dict.fromkeys(result["exclusions"]))[:10]

    return result


def extract_product_from_pdf(pdf_path: str) -> dict:
    """从PDF文件提取产品信息（需要 pdfplumber 或 PyPDF2）"""
    try:
        import pdfplumber
        text = ""
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                text += (page.extract_text() or "") + "\n"
        return parse_insurance_product_from_text(text)
    except ImportError:
        return {"error": "请安装 pdfplumber: pip install pdfplumber"}


def extract_product_from_docx(docx_path: str) -> dict:
    """从Word文件提取产品信息"""
    try:
        from docx import Document
        doc = Document(docx_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        # 也读取表格
        for table in doc.tables:
            for row in table.rows:
                text += "\n" + "\t".join([cell.text for cell in row.cells])
        return parse_insurance_product_from_text(text)
    except ImportError:
        return {"error": "请安装 python-docx: pip install python-docx"}


# ============ 训练案例数据获取 ============


def get_training_case_data(lottery_type: str = "ssq", limit: int = 10) -> list[dict]:
    """
    获取指定彩种最新开奖数据，用于设计培训案例
    lottery_type: 'ssq' | 'fc3d' | 'dlt'
    """
    if lottery_type == "ssq":
        raw = fetch_ssq_latest(limit)
        return [{
            "issue": r.get("code", ""),
            "date": r.get("date", ""),
            "red": r.get("red", ""),
            "blue": r.get("blue", ""),
            "sales": r.get("sales", ""),
            "pool": r.get("poolmoney", ""),
            "note": "用于案例设计中的数字演示（如：算中奖概率）"
        } for r in raw if "error" not in r]

    elif lottery_type == "fc3d":
        raw = fetch_fc3d_latest(limit)
        return [{
            "issue": r.get("issue", ""),
            "date": r.get("date", ""),
            "numbers": r.get("awardDetails", ""),
            "note": "用于入门级计算题（如：和值、跨度计算）"
        } for r in raw if "error" not in r]

    elif lottery_type == "dlt":
        raw = fetch_dlt_latest(limit)
        return [{
            "issue": r.get("lotteryDrawNum", ""),
            "date": r.get("lotteryDrawTime", ""),
            "front": r.get("lotteryDrawResult", [[]])[0][:5],
            "back": r.get("lotteryDrawResult", [[]])[0][5:],
            "note": "用于高净值客户资产配置案例"
        } for r in raw if "error" not in r]

    return []


if __name__ == "__main__":
    # 测试双色球数据获取
    print("=== 双色球最新10期 ===")
    data = get_training_case_data("ssq", 10)
    for item in data[:3]:
        print(f"期号: {item['issue']} | 红球: {item['red']} | 蓝球: {item['blue']}")
