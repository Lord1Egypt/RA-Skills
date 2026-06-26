#!/usr/bin/env python3
"""高校招生通知 → Word文档转换脚本"""
import sys
sys.dont_write_bytecode = True

import json
import os
import sys
import datetime
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

SKILL_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTPUT_DIR = '/tmp'

# 导入监控脚本的查询功能
sys.path.insert(0, os.path.join(SKILL_DIR, 'scripts'))
from monitor import load_schools, get_all_school_notices, format_results_for_display


def create_word_report(results):
    """根据查询结果生成 Word 文档"""
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # ====== 封面标题 ======
    title = doc.add_heading('🎓 高校招生通知监控报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x1A, 0x5C, 0x8A)
    
    # 报告信息
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    now = datetime.datetime.now()
    info_para.add_run(f'生成时间：{now.strftime("%Y-%m-%d %H:%M")}').font.size = Pt(11)
    
    # 分隔线
    doc.add_paragraph('—' * 40).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    total_schools = 0
    total_notices = 0
    total_errors = 0
    
    for school in results:
        total_schools += 1
        
        # 学校标题
        h = doc.add_heading(f'🏫 {school["name"]}', level=1)
        for run in h.runs:
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(0x1A, 0x5C, 0x8A)
        
        for site in school['sites']:
            # 网站信息
            p = doc.add_paragraph()
            p.add_run(f'📢 {site["name"]}').bold = True
            p.add_run(f'   来源：{site["url"]}')
            p.paragraph_format.space_after = Pt(4)
            
            notices = site['notices']
            if notices:
                err_type = notices[0].get('error_type')
                if err_type:
                    # 错误情况
                    err_p = doc.add_paragraph()
                    err_p.add_run(f'  ❌ {notices[0]["title"]}').font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                    err_p.add_run(f'\n  {notices[0]["error_hint"]}')
                    total_errors += 1
                else:
                    # 正常通知 - 用表格展示
                    table = doc.add_table(rows=1, cols=3)
                    table.style = 'Light Shading Accent 1'
                    
                    # 表头
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = '#'
                    hdr_cells[1].text = '标题'
                    hdr_cells[2].text = '日期'
                    
                    # 设置表头宽度
                    hdr_cells[0].width = Cm(1.5)
                    hdr_cells[1].width = Cm(12)
                    hdr_cells[2].width = Cm(3)
                    
                    for i, notice in enumerate(notices, 1):
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(i)
                        row_cells[1].text = notice['title']
                        row_cells[2].text = notice.get('date', '')
                        
                        # 设置字体大小
                        for cell in row_cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(10)
                    
                    total_notices += len(notices)
                    
                    # 添加链接提示
                    link_para = doc.add_paragraph()
                    link_para.paragraph_format.space_before = Pt(2)
                    link_para.add_run(f'  共 {len(notices)} 条通知，点击标题可查看详情').font.size = Pt(9)
            else:
                doc.add_paragraph('  ⚠️ 暂无通知')
        
        doc.add_paragraph()  # 学校间距
    
    # ====== 页脚统计 ======
    doc.add_paragraph('—' * 40).alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary.add_run(f'📈 统计：{total_schools} 所学校 · {total_notices} 条通知').bold = True
    if total_errors > 0:
        summary.add_run(f'\n⚠️ {total_errors} 个网站访问异常')
    
    # 页脚
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.add_run(f'由高校招生监控 Skill 自动生成').font.size = Pt(9)
    
    return doc


def main():
    # 解析参数
    school_names = None
    for arg in sys.argv[1:]:
        if arg.startswith('--schools='):
            school_names = [s.strip() for s in arg.split('=', 1)[1].split(',')]
        elif arg == '--schools' and len(sys.argv) > sys.argv.index(arg) + 1:
            idx = sys.argv.index(arg)
            school_names = [s.strip() for s in sys.argv[idx + 1].split(',')]
        elif arg == '--help' or arg == '-h':
            print("""📄 高校招生通知 → Word文档

用法:
  python3 scripts/to_word.py --schools 深圳大学,武汉大学

参数:
  --schools <学校名1,学校名2>   指定要查询的学校
  --help / -h                   显示帮助信息
""")
            return
    
    if not school_names:
        print("⚠️ 请指定要查询的学校，例如：")
        print("   python3 scripts/to_word.py --schools 深圳大学,武汉大学")
        return
    
    # 查询通知
    results = get_all_school_notices(school_names)
    if not results:
        return
    
    # 生成 Word
    doc = create_word_report(results)
    
    # 保存文件
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    now = datetime.datetime.now()
    filename = f"招生监控报告_{now.strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    
    print(f"\n📄 报告已保存到: {filepath}")


if __name__ == "__main__":
    main()
