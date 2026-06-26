"""
整合报告生成器 v3
=========================
基于 v2 的增量优化版本，新增：
- Phase 0 参考资料管理（reference_material.txt）
- 术语表前置生成（从参考资料中提取）
- 大纲快照机制（plan_outline_snapshot.md）
- 批量版本快照（snapshot_batch_*.md）
- 单章原地修改工具（inline_edit）
- 全局配置（config.json）

自包含设计：单章转换逻辑直接内嵌，不依赖外部脚本
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import glob, os, re, subprocess, sys, json as json_module, shutil, hashlib
from datetime import datetime
from concurrent.futures import ProcessPoolExecutor, as_completed
from typing import Dict, List, Tuple, Optional, Any

# ============ 全局配置 ============
CHAPTERS_DIR = 'F:/agent/chapters'
PLAN_FILE = CHAPTERS_DIR + '/plan.json'
PROGRESS_FILE = CHAPTERS_DIR + '/progress.json'
GLOSSARY_FILE = CHAPTERS_DIR + '/glossary.json'
REFERENCE_FILE = CHAPTERS_DIR + '/reference_material.txt'
OUTLINE_SNAPSHOT = CHAPTERS_DIR + '/plan_outline_snapshot.md'
CONFIG_FILE = CHAPTERS_DIR + '/config.json'
FINAL_DOC = 'F:/agent/整合报告.docx'
CHARS_PER_PAGE = 950
HASH_FILE = CHAPTERS_DIR + '/content_hashes.json'  # 增量更新：章节内容hash清单
MERMAID_TEMP = CHAPTERS_DIR + '/mermaid_temp'     # Mermaid渲染临时目录
# Playwright Chromium 配置（mmdc 专用）
MERMAID_PUPPETEER_CONFIG = CHAPTERS_DIR + '/mermaid_temp/puppeteer_config.json'


# ============ 增量更新：内容Hash ============

def compute_content_hash(content: str) -> str:
    """计算内容MD5（排除空白符差异）"""
    normalized = re.sub(r'\s+', '', content.strip())
    return hashlib.md5(normalized.encode('utf-8')).hexdigest()

def load_hashes() -> Dict[str, str]:
    if os.path.exists(HASH_FILE):
        try:
            with open(HASH_FILE, 'r', encoding='utf-8') as f:
                return json_module.load(f)
        except Exception:
            pass
    return {}

def save_hashes(hashes: Dict[str, str]):
    with open(HASH_FILE, 'w', encoding='utf-8') as f:
        json_module.dump(hashes, f, ensure_ascii=False, indent=2)

def get_changed_chapters(chapters_data: List[Tuple], hashes: Dict[str, str]) -> List[Tuple]:
    """返回实际发生变化的章节列表（增量更新依据）"""
    changed = []
    for item in chapters_data:
        seq = item[0]
        content = item[3]
        new_hash = compute_content_hash(content)
        if hashes.get(seq) != new_hash:
            changed.append(item)
    return changed


# ============ Mermaid 图表渲染 ============

def ensure_mermaid_deps():
    """检查并返回mermaid CLI调用命令（None表示不可用）"""
    local_cli = r'E:\lonb\LobsterAI\node_modules\@mermaid-js\mermaid-cli\src\cli.js'
    # 规范化并去重 ..
    local_cli = os.path.normpath(local_cli)
    candidates = [
        ('local', [local_cli, '--version']),
        ('local_node', ['node', local_cli, '--version']),
        ('mmdc', ['mmdc', '--version']),
        ('npx_mmdc', ['npx', 'mmdc', '--version']),
    ]
    for name, cmd in candidates:
        try:
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=10)
            if result.returncode == 0:
                if name == 'local' or name == 'local_node':
                    return 'E:\\lonb\\LobsterAI\\node_modules\\@mermaid-js\\mermaid-cli\\src\\cli.js'  # mmdc 完整路径（render时会用node调用）
        except Exception:
            continue
    return None

MERMAID_CLI = ensure_mermaid_deps()

def render_mermaid_image(code: str, out_path: str, cli: str = None) -> bool:
    """
    调用 mermaid CLI 将代码块渲染为PNG
    cli: 'mmdc' | 'npx mermaid' 等
    Returns: 是否成功
    """
    if cli is None:
        cli = MERMAID_CLI
    if cli is None:
        return False

    os.makedirs(os.path.dirname(out_path), exist_ok=True)

    # 写临时文件
    import tempfile
    tmp_input = os.path.join(CHAPTERS_DIR, '_mermaid_tmp.mmd')
    with open(tmp_input, 'w', encoding='utf-8') as f:
        f.write(code)

    try:
        # mmdc 是 .js 文件时需用 node 调用
        if cli.endswith('.js'):
            cmd = ['node', cli, '-i', tmp_input, '-o', out_path]
        else:
            cmd = cli.split() + ['-i', tmp_input, '-o', out_path]
        # 注入 Playwright Chromium 配置
        if os.path.exists(MERMAID_PUPPETEER_CONFIG):
            cmd += ['-p', MERMAID_PUPPETEER_CONFIG]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        return os.path.exists(out_path)
    except Exception:
        return False
    finally:
        if os.path.exists(tmp_input):
            os.remove(tmp_input)

def process_mermaid_blocks(content: str) -> Tuple[str, List[str]]:
    """
    检测并渲染 content 中的 mermaid 图表代码块。
    返回: (processed_content, list_of_rendered_image_paths)
    渲染失败时：保留原始代码块，附加【图表渲染失败，请手动替换】提示
    """
    rendered_images = []
    mermaid_blocks = list(re.finditer(r'```mermaid\n(.*?)```', content, re.DOTALL))

    if not mermaid_blocks:
        return content, []

    processed = content
    for m in reversed(mermaid_blocks):  # 逆序处理，从后往前替换
        code = m.group(1).strip()
        # 生成唯一文件名
        block_idx = len(mermaid_blocks) - 1 - mermaid_blocks[::-1].index(m)
        img_name = f'mermaid_{block_idx:03d}.png'
        img_path = os.path.join(MERMAID_TEMP, img_name)

        success = False
        if MERMAID_CLI:
            success = render_mermaid_image(code, img_path, MERMAID_CLI)

        if success:
            rendered_images.append(img_path)
            replacement = f'\n[Mermaid图表已渲染，见附件: {img_name}]\n'
        else:
            replacement = (
                f'\n```mermaid\n{code}\n```\n\n'
                f'<!-- ⚠️ Mermaid图表（渲染工具mmdc未安装或渲染失败，'
                f'请在支持Mermaid的编辑器中查看，或手动替换为图片） -->\n'
            )
        # 用替换文稿重建内容
        processed = processed[:m.start()] + replacement + processed[m.end():]

    return processed, rendered_images


# ============ Word TOC 字段生成 ============

NSMAP = 'xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas" ' \
        'xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" ' \
        'xmlns:o="urn:schemas-microsoft-com:office:office" ' \
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" ' \
        'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" ' \
        'xmlns:v="urn:schemas-microsoft-com:vml" ' \
        'xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing" ' \
        'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" ' \
        'xmlns:w10="urn:schemas-microsoft-com:office:word" ' \
        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" ' \
        'xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" ' \
        'xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup" ' \
        'xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk" ' \
        'xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml" ' \
        'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"'


def _make_bookmark_start(bookmark_id: int, bookmark_name: str) -> OxmlElement:
    el = OxmlElement('w:bookmarkStart')
    el.set(qn('w:id'), str(bookmark_id))
    el.set(qn('w:name'), bookmark_name)
    return el

def _make_bookmark_end(bookmark_id: int) -> OxmlElement:
    el = OxmlElement('w:bookmarkEnd')
    el.set(qn('w:id'), str(bookmark_id))
    return el


def add_toc_entry(doc, seq: str, title: str, page_num: int, toc_type: str = 'chapter'):
    """
    生成真实的Word TOC条目（使用 FORMTEXT + PAGEREF 字段）。
    seq: 章节序号，如"一"或"第一章"
    toc_type: 'summary'（执行摘要）| 'chapter'（章节）
    """
    bm_id = 100 + hash(title) % 1000

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(22)
    p.paragraph_format.space_after = Pt(4)

    if toc_type == 'summary':
        # 执行摘要：纯文本，无超链接
        p.paragraph_format.first_line_indent = Cm(-0.74)
        r = p.add_run(seq + '　' + title)
        r.font.size = Pt(12)
        cjk(r, '宋体')
        return

    # ---- 章节TOC条目：带超链接 + Tab + 页码字段 ----
    prefix = seq + '　'
    p.paragraph_format.first_line_indent = Cm(-0.74)

    # 前缀文本
    r_prefix = p.add_run(prefix)
    r_prefix.font.size = Pt(12)
    cjk(r_prefix, '宋体')

    # 超链接（链接到本章书签）
    bookmark_name = f'_Toc_{bm_id}'
    run = p.add_run()
    run.font.size = Pt(12)
    cjk(run, '宋体')

    # 插入 FORMTEXT 字段（显示标题）
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld_char_begin)

    instr_text = OxmlElement('w:instrText')
    instr_text.text = f' FORMTEXT '
    run._r.append(instr_text)

    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_char_end)

    # 插入 Tab + PAGEREF 字段（显示页码）
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    p._p.append(tab)

    tab_char = OxmlElement('w:tabChar')
    tab_char.set(qn('w:val'), 'right')
    p._p.append(tab_char)

    run_page = p.add_run()
    run_page.font.size = Pt(12)
    cjk(run_page, '宋体')

    # PAGEREF 字段
    fld_char_begin2 = OxmlElement('w:fldChar')
    fld_char_begin2.set(qn('w:fldCharType'), 'begin')
    run_page._r.append(fld_char_begin2)

    instr_text2 = OxmlElement('w:instrText')
    instr_text2.text = f' PAGEREF {bookmark_name} \\h '
    run_page._r.append(instr_text2)

    fld_char_end2 = OxmlElement('w:fldChar')
    fld_char_end2.set(qn('w:fldCharType'), 'end')
    run_page._r.append(fld_char_end2)

    # 添加书签（供 PAGEREF 引用）
    p._p.insert(0, _make_bookmark_start(bm_id, bookmark_name))
    p._p.append(_make_bookmark_end(bm_id))

    return bm_id, bookmark_name


# ============ 配置读写 ============

def load_config() -> Dict[str, Any]:
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
                return json_module.load(f)
        except Exception:
            pass
    return {"project_name": "", "topic": "", "audience": "", "doc_type": "可行性研究报告", "style": "专业严谨", "custom_constraints": ""}


def save_config(cfg: Dict[str, Any]):
    with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
        json_module.dump(cfg, f, ensure_ascii=False, indent=2)


def load_plan() -> Dict[str, Any]:
    if os.path.exists(PLAN_FILE):
        try:
            with open(PLAN_FILE, 'r', encoding='utf-8') as f:
                return json_module.load(f)
        except Exception:
            pass
    return make_default_plan()


def make_default_plan() -> Dict[str, Any]:
    return {"project_name": "", "chapters": []}


def save_plan(plan: Dict[str, Any]):
    with open(PLAN_FILE, 'w', encoding='utf-8') as f:
        json_module.dump(plan, f, ensure_ascii=False, indent=2)


# ============ 参考资料管理 ============

def load_reference() -> str:
    """加载参考资料"""
    if os.path.exists(REFERENCE_FILE):
        try:
            with open(REFERENCE_FILE, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception:
            pass
    return ""


def save_reference(text: str):
    """保存参考资料"""
    with open(REFERENCE_FILE, 'w', encoding='utf-8') as f:
        f.write(text)
    print(f"[REF] 参考资料已保存，共 {len(text)} 字符")


def extract_terms_from_reference(text: str, max_terms=80) -> List[Dict[str, Any]]:
    """
    从参考资料中提取术语（专业词汇提取）
    策略：提取重复出现2次以上的中文词组（>=4字），过滤停用词
    """
    if not text:
        return []

    stopwords = {
        '以及', '包括', '可以', '通过', '根据', '按照', '为了', '由于', '其中',
        '其他', '相关', '以上', '以下', '对于', '并且', '或者', '等等',
        '本项目', '本公司', '本系统', '本章', '本节', '本文', '本案',
        '进行', '完成', '实现', '提供', '使用', '管理', '系统', '建设',
        '方案', '项目', '数据', '平台', '技术', '功能', '模块'
    }

    # 提取中文词组
    pattern = re.compile(r'[\u4e00-\u9fff]{4,}')
    candidates = pattern.findall(text)

    # 统计频次
    freq: Dict[str, int] = {}
    for w in candidates:
        if w not in stopwords and len(w) >= 4:
            freq[w] = freq.get(w, 0) + 1

    # 过滤：出现>=2次
    filtered = {w: c for w, c in freq.items() if c >= 2}
    sorted_terms = sorted(filtered.items(), key=lambda x: -x[1])[:max_terms]

    return [{"term": t, "count": c, "source": "reference"} for t, c in sorted_terms]


def build_reference_summary(text: str, max_chars=3000) -> str:
    """构建参考资料摘要（供子Agent使用）"""
    if not text:
        return ""
    # 取前max_chars
    summary = text[:max_chars]
    if len(text) > max_chars:
        summary += f"\n\n[...参考资料共 {len(text)} 字符，此处省略中间部分...]\n\n" + text[-1000:]
    return summary


# ============ 字体辅助 ============

def cjk(run, name):
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)


# ============ Markdown → docx 表格辅助函数 ============

def _clean_inline(text):
    """清除行内markdown符号"""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    return text.strip()

def _is_table_line(line):
    stripped = line.strip()
    return stripped.startswith('|') and stripped.endswith('|')

def _is_separator_line(line):
    """判断是否为 markdown 表格分隔行（如 |----|----|）"""
    stripped = line.strip().strip('|')
    return bool(re.match(r'^[\s\-:.|]+$', stripped))

def _parse_md_table(rows):
    """将markdown表格行列表解析为二维字符串数组"""
    result = []
    for line in rows:
        stripped = line.strip().strip('|')
        cols = stripped.split('|')
        result.append([_clean_inline(c.strip()) for c in cols])
    return result

def _add_table_to_doc(doc, rows):
    """将解析后的表格写入docx"""
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    col_count = max(col_count, 1)
    tbl = doc.add_table(rows=len(rows), cols=col_count)
    tbl.style = 'Table Grid'
    for r_idx, row_data in enumerate(rows):
        cells = tbl.rows[r_idx].cells
        actual = len(cells)
        for c_idx in range(actual):
            text = row_data[c_idx] if c_idx < len(row_data) else ''
            cells[c_idx].text = text
            for para in cells[c_idx].paragraphs:
                for run in para.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(10)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after = Pt(2)

def _flush_table(doc, pending_table):
    """将收集的表格行写入doc，然后清空缓冲区"""
    if pending_table:
        _add_table_to_doc(doc, pending_table)
        pending_table.clear()

def _write_para(doc, line, font='宋体', size=12, bold=False,
                first_indent=Cm(0.74), before=Pt(2), after=Pt(6),
                alignment=None, is_heading=False):
    """写入正文段落（统一封装，方便多处复用）"""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.first_line_indent = first_indent
    p.paragraph_format.line_spacing = Pt(22)
    p.paragraph_format.space_before = before
    p.paragraph_format.space_after = after
    r = p.add_run(_clean_inline(line))
    r.font.size = Pt(size)
    r.font.bold = bold
    cjk(r, font)
    return p

# ============ Markdown → docx ============

def md_to_paragraphs(doc, content, add_page_break=True):
    """将markdown内容写入docx，正确处理表格和Mermaid图表"""
    # --- Mermaid 预处理 ---
    processed_content, rendered_images = process_mermaid_blocks(content)
    # 建立 mermaid 代码 → 图片路径 的映射
    mermaid_img_map = {}
    if rendered_images:
        for img_path in rendered_images:
            img_name = os.path.basename(img_path)
            # 从处理后的内容中提取 mermaid 代码作为 key（用标记）
            for m in re.finditer(r'\[Mermaid图表已渲染，见附件:\s*(\S+)\]', processed_content):
                fname = m.group(1)
                if fname == img_name:
                    mermaid_img_map[m.group(0)] = img_path
                    break

    lines = processed_content.split('\n')
    i = 0
    pending_table = []
    mermaid_img_iter = iter(rendered_images) if rendered_images else iter([])

    while i < len(lines):
        line = lines[i].rstrip()
        i += 1

        # --- Mermaid 已渲染图片插入 ---
        if '[Mermaid图表已渲染，见附件:' in line:
            img_path = next(mermaid_img_iter, None)
            if img_path and os.path.exists(img_path):
                _flush_table(doc, pending_table)
                try:
                    p = doc.add_paragraph()
                    run = p.add_run()
                    run.add_picture(img_path, width=Inches(5.5))
                except Exception as e:
                    # 图片插入失败降级为文字提示
                    p = doc.add_paragraph()
                    r = p.add_run(line + ' [图片渲染失败]')
                    r.font.size = Pt(10); cjk(r, '宋体')
            continue

        if not line.strip():
            _flush_table(doc, pending_table)
            continue

        if line.startswith('# '):
            _flush_table(doc, pending_table)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(10)
            r = p.add_run(_clean_inline(line[2:]))
            r.font.size = Pt(18); r.font.bold = True; cjk(r, '黑体')
            continue

        if line.startswith('## '):
            _flush_table(doc, pending_table)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(6)
            r = p.add_run(_clean_inline(line[3:]))
            r.font.size = Pt(14); r.font.bold = True; cjk(r, '楷体')
            continue

        if line.startswith('### '):
            _flush_table(doc, pending_table)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
            r = p.add_run(_clean_inline(line[4:]))
            r.font.size = Pt(12); r.font.bold = True; cjk(r, '仿宋')
            continue

        if line.startswith('#### '):
            _flush_table(doc, pending_table)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(3)
            r = p.add_run(_clean_inline(line[5:]))
            r.font.size = Pt(11); r.font.bold = True; cjk(r, '仿宋')
            continue

        # markdown 表格行
        if _is_table_line(line):
            if not _is_separator_line(line):
                pending_table.append(line)
            continue

        # 非表格行 → flush 缓存表格，写入正文段落
        _flush_table(doc, pending_table)
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = Pt(22)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(_clean_inline(line))
        r.font.size = Pt(12); cjk(r, '宋体')

    # 处理末尾可能残留的表格
    _flush_table(doc, pending_table)
    if add_page_break:
        doc.add_page_break()


# ============ 章节解析（错误隔离）============

def safe_parse_chapter(fpath: str) -> Optional[Tuple]:
    fname = os.path.basename(fpath).replace('.txt', '')
    seq = fname.split('-')[0]
    try:
        with open(fpath, 'r', encoding='utf-8') as fp:
            content = fp.read()
    except Exception as e:
        print(f"[ERROR] 读取失败 {fname}: {e}")
        return None
    h2_entries = [l[3:].strip() for l in content.split('\n') if l.strip().startswith('## ')]
    title = fname
    for line in content.split('\n'):
        line = line.strip()
        if line.startswith('# '):
            title = line[2:].strip()
            break
    return (seq, fname, title, content, h2_entries)


def parse_chapters(txt_files: List[str]) -> List[Tuple]:
    seen_seq = set()
    chapters, errors = [], []
    for f in txt_files:
        seq = os.path.basename(f).replace('.txt', '').split('-')[0]
        if seq in seen_seq:
            continue
        result = safe_parse_chapter(f)
        if result is None:
            errors.append(os.path.basename(f)); continue
        seen_seq.add(seq); chapters.append(result)
    if errors:
        print(f"[WARN] 以下章节解析失败（已跳过）: {errors}")
    return chapters


# ============ 字符统计 ============

def count_chars(text: str) -> int:
    return len([c for c in text if c.strip()])


# ============ Glossary 生成（前置版）============

def generate_glossary(txt_files: List[str] = None, ref_text: str = "", max_terms=80) -> Dict[str, Any]:
    """从参考资料和章节内容中生成术语表"""
    all_terms: Dict[str, int] = {}

    # 从参考资料提取
    if ref_text:
        ref_terms = extract_terms_from_reference(ref_text, max_terms)
        for item in ref_terms:
            all_terms[item['term']] = all_terms.get(item['term'], 0) + item['count']

    # 从章节内容提取
    if txt_files:
        stopwords = {'以及', '包括', '可以', '通过', '根据', '按照', '为了', '由于', '其中', '其他', '相关', '以上', '以下', '对于', '并且', '或者', '等等', '本项目', '本公司', '本系统'}
        pattern = re.compile(r'[\u4e00-\u9fff]{4,}')
        for f in txt_files:
            try:
                with open(f, 'r', encoding='utf-8') as fp:
                    content = fp.read()
                for w in pattern.findall(content):
                    if w not in stopwords and len(w) >= 4:
                        all_terms[w] = all_terms.get(w, 0) + 1
            except Exception:
                continue

    # 按频次排序取前max_terms
    sorted_terms = sorted(all_terms.items(), key=lambda x: -x[1])[:max_terms]
    glossary = {
        "generated_at": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        "total_ref_chars": len(ref_text),
        "terms": [{"term": t, "count": c} for t, c in sorted_terms]
    }
    with open(GLOSSARY_FILE, 'w', encoding='utf-8') as f:
        json_module.dump(glossary, f, ensure_ascii=False, indent=2)
    print(f"[GLOSSARY] 术语表已生成: {GLOSSARY_FILE}（共 {len(sorted_terms)} 个术语）")
    return glossary


def load_glossary() -> Dict[str, Any]:
    if os.path.exists(GLOSSARY_FILE):
        try:
            with open(GLOSSARY_FILE, 'r', encoding='utf-8') as f:
                return json_module.load(f)
        except Exception:
            pass
    return {"terms": []}


def glossary_to_prompt_text(glossary: Dict[str, Any], max_terms=30) -> str:
    """将术语表转为子Agent可读的提示文本"""
    terms = glossary.get('terms', [])
    if not terms:
        return "（术语表暂无数据，完成 Batch A 后自动生成）"
    display = terms[:max_terms]
    lines = [f"- {t['term']}（出现{t['count']}次）" for t in display]
    suffix = f"\n（共 {len(terms)} 个术语，仅展示前 {max_terms} 个）" if len(terms) > max_terms else ""
    return '\n'.join(lines) + suffix


# ============ 大纲快照 ============

def save_outline_snapshot(plan: Dict[str, Any]):
    """保存规划师输出的大纲快照"""
    lines = [f"# 文档大纲快照（{datetime.now().strftime('%Y-%m-%d %H:%M')}）"]
    project_name = plan.get('project_name', '未知项目')
    lines.append(f"\n项目：{project_name}\n")
    for ch in plan.get('chapters', []):
        lines.append(f"第{ch.get('seq','?')}章 | {ch.get('title','')} | Batch {ch.get('batch','')} | 约{ch.get('word_count',0)}字 | 依赖:{ch.get('dependencies',[])}")
    content = '\n'.join(lines)
    with open(OUTLINE_SNAPSHOT, 'w', encoding='utf-8') as f:
        f.write(content)
    print(f"[SNAPSHOT] 大纲快照已保存: {OUTLINE_SNAPSHOT}")


# ============ 批量版本快照 ============

def save_batch_snapshot(batch_label: str, batch_chapters: List[Tuple]):
    """保存每批完成后的章节内容快照"""
    snapshot_file = f"{CHAPTERS_DIR}/snapshot_{batch_label}_{datetime.now().strftime('%Y%m%d_%H%M')}.md"
    lines = [f"# {batch_label} 快照（{datetime.now().strftime('%Y-%m-%d %H:%M')}）"]
    for seq, fname, title, content, _ in batch_chapters:
        lines.append(f"\n---\n## 第{seq}章 {title}\n")
        # 只保存前200字预览
        preview = content[:300].replace('\n', ' ').strip()
        lines.append(f"[预览] {preview}...")
    with open(snapshot_file, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))
    print(f"[SNAPSHOT] 批次快照已保存: {snapshot_file}")


# ============ 跨章一致性审查 ============

def extract_quantities(text: str) -> Dict[str, str]:
    qty = {}
    pattern = re.compile(r'(\d+(?:\.\d+)?)\s*(万元|万元/年|万元\/年|人|人\/日|台|套|个|次|年|月|天|%)')
    for m in pattern.finditer(text):
        key = f"{m.group(1)}{m.group(2)}"
        qty[key] = m.group(0)
    return qty


def check_cross_chapter_consistency(chapters_data: List[Tuple]) -> List[Dict]:
    issues = []
    all_qty = [(seq, fname, extract_quantities(content)) for seq, fname, title, content, _ in chapters_data]
    for i in range(len(all_qty) - 1):
        seq_a, fname_a, qty_a = all_qty[i]
        seq_b, fname_b, qty_b = all_qty[i + 1]
        shared = set(qty_a.keys()) & set(qty_b.keys())
        for key in shared:
            ma = re.match(r'^(\d+(?:\.\d+)?)', key)
            mb = re.match(r'^(\d+(?:\.\d+)?)', key)
            if ma and mb:
                try:
                    if float(ma.group(1)) != float(mb.group(1)):
                        issues.append({
                            "seq_a": seq_a, "seq_b": seq_b,
                            "item": key,
                            "value_a": qty_a[key], "value_b": qty_b[key]
                        })
                except ValueError:
                    continue
    return issues


# ============ 执行摘要 ============

def _build_summary(chapters_data, max_chars=800):
    lines, total = [], 0
    for seq, fname, title, content, h2_list in chapters_data:
        para_lines = []
        for line in content.split('\n'):
            line = line.strip()
            if not line or line.startswith('# ') or line.startswith('## ') or line.startswith('### '):
                continue
            para_lines.append(line)
            if len(para_lines) >= 3:
                break
        if not para_lines:
            continue
        para_text = ''.join(para_lines[:2])
        if total + len(para_text) > max_chars:
            remaining = max_chars - total
            if remaining > 50:
                lines.append(para_text[:remaining] + '…')
            break
        lines.append(para_text)
        total += len(para_text)
    return lines or ['本报告对项目建设进行了全面可行性分析。']


# ============ 最终文档生成 ============

def generate_final_doc(chapters_data, page_estimates, output_path=FINAL_DOC, incremental=True):
    """
    生成整合报告 docx。
    incremental=True（默认）：对比 hash，仅章节内容变化才重写该章；
                              无变化时跳过重写，直接复用已有章节 docx。
    """
    plan = load_plan()

    # --- 增量更新：检查哪些章节发生了变化 ---
    changed_chapters = chapters_data
    if incremental:
        hashes = load_hashes()
        changed_chapters = get_changed_chapters(chapters_data, hashes)
        unchanged = [item for item in chapters_data if item not in changed_chapters]
        changed_seqs = {item[0] for item in changed_chapters}
        if unchanged and not changed_chapters:
            print(f"[INCREMENTAL] 所有 {len(chapters_data)} 章内容未变化，跳过重写")
            return None
        elif unchanged:
            print(f"[INCREMENTAL] {len(unchanged)} 章未变化，{len(changed_chapters)} 章需重写: {changed_seqs}")

    doc = Document()
    s = doc.sections[0]
    s.page_height = Inches(11.69); s.page_width = Inches(8.27)
    s.top_margin = Inches(1.0); s.bottom_margin = Inches(1.0)
    s.left_margin = Inches(1.18); s.right_margin = Inches(1.18)

    # 封面
    for _ in range(6): doc.add_paragraph()
    for txt, size, bold, font in [
        (plan.get('org_name', '编制单位'), Pt(26), True, '黑体'),
        (plan.get('project_name', '项目名称'), Pt(32), True, '黑体'),
    ]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt); r.font.size = size; r.font.bold = bold; cjk(r, font)
    for _ in range(3): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(plan.get('doc_type', '可行性研究报告'))
    r.font.size = Pt(22); cjk(r, '楷体')
    for _ in range(8): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    unit = plan.get('编制单位', '编制单位')
    build_time = plan.get('编制时间', datetime.now().strftime('%Y年%m月'))
    r = p.add_run(f'编制单位：{unit}\n编制时间：{build_time}')
    r.font.size = Pt(14); cjk(r, '宋体')
    doc.add_page_break()

    # 执行摘要
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(10)
    r = p.add_run('执行摘要'); r.font.size = Pt(18); r.font.bold = True; cjk(r, '黑体')
    for pt in _build_summary(changed_chapters if changed_chapters else chapters_data):
        p2 = doc.add_paragraph()
        p2.paragraph_format.first_line_indent = Cm(0.74)
        p2.paragraph_format.line_spacing = Pt(22)
        p2.paragraph_format.space_after = Pt(6)
        r2 = p2.add_run(pt); r2.font.size = Pt(12); cjk(r2, '宋体')
    doc.add_page_break()

    # 目录（使用真实 Word TOC 字段）
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(10)
    r = p.add_run('目  录'); r.font.size = Pt(18); r.font.bold = True; cjk(r, '黑体')
    add_toc_entry(doc, '一', '执行摘要', 1, toc_type='summary')
    seen = set()
    for seq, fname, title, content, h2_list in (changed_chapters if changed_chapters else chapters_data):
        if seq in seen: continue
        seen.add(seq)
        start = page_estimates.get(seq, (1, 0, 1))[0]
        if not seq.isdigit(): continue
        add_toc_entry(doc, f'第{int(seq)}章', title, start, toc_type='chapter')
    doc.add_page_break()

    # 各章节
    target = changed_chapters if changed_chapters else chapters_data
    for seq, fname, title, content, h2_list in target:
        md_to_paragraphs(doc, content, add_page_break=True)

    # --- 更新 hash 清单（增量记录）---
    if incremental:
        new_hashes = {}
        for item in (changed_chapters if changed_chapters else chapters_data):
            seq, content = item[0], item[3]
            new_hashes[seq] = compute_content_hash(content)
        # 合并未变化章节的旧 hash
        old_hashes = load_hashes()
        old_hashes.update(new_hashes)
        save_hashes(old_hashes)

    doc.save(output_path)
    print(f"[DONE] 整合报告已保存: {output_path}")
    return output_path


# ============ 整合报告主流程 ============

def generate_with_accurate_toc(txt_dir=CHAPTERS_DIR, final_doc=FINAL_DOC):
    txt_files = sorted(glob.glob(f'{txt_dir}/*.txt'))
    if not txt_files:
        print(f"[ERROR] 未找到章节文件: {txt_dir}/*.txt"); return None

    chapters_data = parse_chapters(txt_files)
    if not chapters_data:
        print("[ERROR] 所有章节解析均失败"); return None
    print(f"[PARSE] 解析 {len(chapters_data)} 个章节")

    # 更新术语表（含参考资料）
    ref_text = load_reference()
    generate_glossary(txt_files, ref_text=ref_text)

    # 跨章一致性审查
    issues = check_cross_chapter_consistency(chapters_data)
    if issues:
        print(f"[CONSISTENCY] 发现 {len(issues)} 个潜在不一致:")
        for iss in issues:
            print(f"  - {iss['message']}")
    else:
        print("[CONSISTENCY] 跨章一致性检查通过 (OK)")

    # 估算页码
    pe = {}
    cur = 7
    for seq, fname, title, content, h2_list in chapters_data:
        cc = count_chars(content)
        ep = max(1, (cc + CHARS_PER_PAGE - 1) // CHARS_PER_PAGE)
        pe[seq] = (cur, cc, ep); cur += ep

    print("[BUILD] 生成整合报告...")
    generate_final_doc(chapters_data, pe, output_path=final_doc)

    # 纯文本版
    md_path = final_doc.replace('.docx', '-纯文本.md')
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write('\n\n---\n\n'.join(c for _, _, _, c, _ in chapters_data))
    print(f"[MD] 纯文本版已保存: {md_path}")

    return final_doc


# ============ 单章 docx 转换 ============

def convert_single_chapter_inline(txt_path, docx_path):
    """将txt章节文件转换为docx，正确解析markdown表格"""
    try:
        doc = Document()
        s = doc.sections[0]
        s.page_height = Inches(11.69); s.page_width = Inches(8.27)
        s.top_margin = Inches(1.0); s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.18); s.right_margin = Inches(1.18)

        with open(txt_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()

        pending_table = []

        for line in lines:
            line = line.rstrip()

            if not line.strip():
                _flush_table(doc, pending_table)
                continue

            if line.startswith('# '):
                _flush_table(doc, pending_table)
                p = doc.add_paragraph(); p.alignment = 1
                p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(10)
                r = p.add_run(_clean_inline(line[2:])); r.font.size = Pt(18); r.font.bold = True; cjk(r, '黑体')
                continue

            if line.startswith('## '):
                _flush_table(doc, pending_table)
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(6)
                r = p.add_run(_clean_inline(line[3:])); r.font.size = Pt(14); r.font.bold = True; cjk(r, '楷体')
                continue

            if line.startswith('### '):
                _flush_table(doc, pending_table)
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
                r = p.add_run(_clean_inline(line[4:])); r.font.size = Pt(12); r.font.bold = True; cjk(r, '仿宋')
                continue

            if line.startswith('#### '):
                _flush_table(doc, pending_table)
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(3)
                r = p.add_run(_clean_inline(line[5:])); r.font.size = Pt(11); r.font.bold = True; cjk(r, '仿宋')
                continue

            # markdown 表格行
            if _is_table_line(line):
                if not _is_separator_line(line):
                    pending_table.append(line)
                continue

            # 普通正文段落
            _flush_table(doc, pending_table)
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0.74)
            p.paragraph_format.line_spacing = Pt(22)
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(_clean_inline(line)); r.font.size = Pt(12); cjk(r, '宋体')

        # 处理末尾残留表格
        _flush_table(doc, pending_table)
        doc.save(docx_path)
        return docx_path
    except Exception as e:
        print(f"[ERROR] 转换失败 {txt_path}: {e}")
        raise


def _convert_worker(args) -> Tuple[str, bool, str]:
    txt_path, docx_path = args
    try:
        convert_single_chapter_inline(txt_path, docx_path)
        return (docx_path, True, '')
    except Exception as e:
        return (txt_path, False, str(e))


# ============ 批量并行转换 ============

def batch_convert_txt_to_docx(txt_dir=CHAPTERS_DIR, max_concurrent=8, progress_file=PROGRESS_FILE, incremental=True):
    """
    批量将 txt 章节转换为 docx。
    incremental=True（默认）：对比内容hash，仅转换有变化的章节。
    force=False：跳过已存在的docx（默认True）。
    """
    txt_files = sorted(glob.glob(os.path.join(txt_dir, '*.txt')))
    if not txt_files:
        print(f"[ERROR] 未找到 .txt 文件"); return []

    hashes = load_hashes() if incremental else {}
    jobs = []
    for tf in txt_files:
        docx_path = tf.replace('.txt', '.docx')
        content_hash = compute_content_hash(open(tf, 'r', encoding='utf-8').read())
        if incremental and os.path.exists(docx_path):
            if hashes.get(os.path.basename(tf)) == content_hash:
                print(f"  [SKIP] {os.path.basename(tf)} 内容未变化，跳过")
                continue
        jobs.append((tf, docx_path))

    if not jobs:
        print("[INFO] 所有章节已是最新（无变化），跳过转换")
        return []
    print(f"[BATCH] 待转换 {len(jobs)} 个章节，并发上限 {max_concurrent}")
    completed, failed = [], []
    with ProcessPoolExecutor(max_workers=max_concurrent) as executor:
        futures = {executor.submit(_convert_worker, job): job for job in jobs}
        for future in as_completed(futures):
            docx_path, ok, err = future.result()
            if ok:
                # 更新 hash
                txt_path = docx_path.replace('.docx', '.txt')
                if os.path.exists(txt_path):
                    hashes[os.path.basename(txt_path)] = compute_content_hash(
                        open(txt_path, 'r', encoding='utf-8').read()
                    )
                completed.append(docx_path); print(f"  [OK] {os.path.basename(docx_path)}")
            else:
                failed.append((docx_path, err)); print(f"  [FAIL] {os.path.basename(docx_path)}: {err}")

    if incremental and completed:
        save_hashes(hashes)

    print(f"\n[BATCH] {len(completed)}/{len(jobs)} 成功，{len(failed)} 失败")
    return completed


# ============ 进度文件 ============

def load_progress() -> Dict:
    if os.path.exists(PROGRESS_FILE):
        try:
            with open(PROGRESS_FILE, 'r', encoding='utf-8') as f:
                return json_module.load(f)
        except Exception:
            pass
    return {"total": 0, "completed": 0, "batches": [], "current": ""}


# ============ CLI 入口 ============

if __name__ == '__main__':
    if len(sys.argv) >= 2 and sys.argv[1] == '--convert-one':
        if len(sys.argv) != 4:
            print("用法: python integrate_report.py --convert-one <in.txt> <out.docx>"); sys.exit(1)
        convert_single_chapter_inline(sys.argv[2], sys.argv[3])
        print(f"saved: {sys.argv[3]}", flush=True); sys.exit(0)

    elif len(sys.argv) >= 2 and sys.argv[1] == 'convert-batch':
        txt_dir = sys.argv[2] if len(sys.argv) > 2 else CHAPTERS_DIR
        batch_convert_txt_to_docx(txt_dir=txt_dir)

    elif len(sys.argv) >= 2 and sys.argv[1] == 'glossary':
        txt_files = sorted(glob.glob(f'{CHAPTERS_DIR}/*.txt'))
        ref_text = load_reference()
        generate_glossary(txt_files, ref_text=ref_text)

    elif len(sys.argv) >= 2 and sys.argv[1] == 'check':
        txt_files = sorted(glob.glob(f'{CHAPTERS_DIR}/*.txt'))
        chapters_data = parse_chapters(txt_files)
        issues = check_cross_chapter_consistency(chapters_data)
        if not issues:
            print("[OK] 跨章一致性检查通过，无不一致项")
        else:
            for iss in issues:
                print(f"[WARN] {iss['message']}")

    elif len(sys.argv) >= 2 and sys.argv[1] == 'status':
        prog = load_progress()
        print(f"进度: {prog.get('completed',0)}/{prog.get('total','?')}")
        if prog.get('current'): print(f"状态: {prog['current']}")

    elif len(sys.argv) >= 2 and sys.argv[1] == 'ref':
        # 仅查看/更新参考资料
        if len(sys.argv) >= 3:
            action = sys.argv[2]
            if action == 'show':
                ref = load_reference()
                print(f"参考资料: {len(ref)} 字符")
                print(ref[:500] if ref else "(空)")
            elif action == 'clear':
                save_reference("")
                print("参考资料已清空")
        else:
            ref = load_reference()
            print(f"当前参考资料: {len(ref)} 字符")

    else:
        # 默认：生成整合报告
        txt_dir = sys.argv[1] if len(sys.argv) > 1 else CHAPTERS_DIR
        result = generate_with_accurate_toc(txt_dir=txt_dir)
        if result:
            print(f"\n[DONE] 整合报告生成完成: {result}")
