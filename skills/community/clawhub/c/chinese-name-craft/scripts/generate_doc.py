#!/usr/bin/env python3
"""Generate a formatted Word document from naming analysis JSON data.

Usage:
    python generate_doc.py <input.json> [output.docx]

JSON schema: see chinese-name-craft/references/output-template.md
"""

import json
import sys
import os

def check_deps():
    try:
        from docx import Document
        return True
    except ImportError:
        print("Installing python-docx...")
        import subprocess
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
        return True

def generate(json_path, output_path=None):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    with open(json_path, "r", encoding="utf-8-sig") as f:
        data = json.load(f)

    if not output_path:
        father = data.get("father_name", "")
        mother = data.get("mother_name", "")
        suffix = data.get("child_label", "之子")
        output_path = os.path.join(
            os.path.dirname(json_path),
            f"\u300a{father}\u5148\u751f{mother}\u5973\u58eb{suffix}\u300b\u5b9a\u540d\u65b9\u6848.docx"
        )

    doc = Document()

    # --- styles ---
    style = doc.styles["Normal"]
    style.font.name = "SimSun"
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    # --- helpers ---
    def _run(p, text, font_name="SimSun", size=12, bold=False, color=None):
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = font_name
        run.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        if color:
            run.font.color.rgb = color
        return run

    def title(text, size=20, color=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, text, "SimHei", size, bold=True, color=color)
        return p

    def heading(text, level=1):
        h = doc.add_heading(text, level=level)
        for r in h.runs:
            r.font.name = "SimHei"
            r.element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
        return h

    def para(text, bold=False, indent=False, font="SimSun", size=12):
        p = doc.add_paragraph()
        if indent:
            p.paragraph_format.first_line_indent = Pt(24)
        _run(p, text, font, size, bold)
        return p

    def quote(text):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        r = _run(p, text, "KaiTi", 11)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        return p

    def centered(text, font="SimHei", size=14, bold=True, color=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, text, font, size, bold, color)
        return p

    def make_table(rows_data, bold_header=True):
        nrows = len(rows_data)
        ncols = max(len(r) for r in rows_data)
        table = doc.add_table(rows=nrows, cols=ncols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, row in enumerate(rows_data):
            for j, val in enumerate(row):
                cell = table.rows[i].cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(val)
                run.font.size = Pt(11)
                if i == 0 or j == 0:
                    run.bold = True
                    run.font.name = "SimHei"
                    run.element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
                else:
                    run.font.name = "SimSun"
                    run.element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
        return table

    def wuge_table(wg):
        labels = ["\u5929\u683c", "\u4eba\u683c", "\u5730\u683c", "\u5916\u683c", "\u603b\u683c"]
        keys = ["tian", "ren", "di", "wai", "zong"]
        wx_keys = ["tian5", "ren5", "di5", "wai5", "zong5"]
        rows = [
            labels,
            [wg[k] for k in keys],
            [wg[k] for k in wx_keys],
        ]
        return make_table(rows)

    # ==================== DOCUMENT BODY ====================

    # Title
    title_text = data.get("title", "\u5b9a\u540d\u65b9\u6848")
    title(title_text, color=RGBColor(0x8B, 0x00, 0x00))
    doc.add_paragraph()

    # Greeting
    if "greeting" in data:
        para(data["greeting"], indent=True)

    doc.add_paragraph()

    # Section 1: Basic Info
    heading("\u4e00\u3001\u57fa\u672c\u4fe1\u606f", 1)
    for item in data.get("basic_info", []):
        para(item)

    doc.add_paragraph()

    # Section 2: Bazi
    heading("\u4e8c\u3001\u516b\u5b57\u6392\u76d8", 1)
    bazi = data.get("bazi", {})
    bazi_rows = []
    label_map = {"shishen": "\u5341\u795e", "tiangan": "\u5929\u5e72",
                  "dizhi": "\u5730\u652f", "wuxing": "\u4e94\u884c", "canggan": "\u85cf\u5e72"}
    for key, label in label_map.items():
        if key in bazi:
            bazi_rows.append([label] + bazi[key])
    if bazi_rows:
        make_table(bazi_rows)

    doc.add_paragraph()

    # Wuxing stats
    if "wuxing_stats" in data:
        para("\u4e94\u884c\u7edf\u8ba1\uff08\u5929\u5e72\u5730\u652f\u672c\u6c14\uff09\uff1a", bold=True)
        for line in data["wuxing_stats"]:
            para("\u3000\u3000" + line)

    doc.add_paragraph()

    if "need_wuxing" in data:
        centered("\u9700\u8865\u4e94\u884c\uff1a" + data["need_wuxing"],
                 size=14, color=RGBColor(0x8B, 0x00, 0x00))

    doc.add_paragraph()

    # Section 3: Expert Commentary
    heading("\u4e09\u3001\u547d\u5c40\u4e13\u5bb6\u70b9\u8bc4", 1)
    for item in data.get("expert_comment", []):
        para(item, indent=True)

    doc.add_paragraph()

    # Section 4: Names
    heading("\u56db\u3001\u5b9a\u540d\u65b9\u6848", 1)

    for name_data in data.get("names", []):
        heading(name_data["title"], 2)

        # Wuge table
        if "wuge" in name_data:
            wuge_table(name_data["wuge"])
            doc.add_paragraph()

        # Sancai
        if "sancai" in name_data:
            para("\u4e09\u624d\u914d\u7f6e\uff1a" + name_data["sancai"], bold=True)

        # Ziyi
        if "ziyi" in name_data:
            para("\u5b57\u4e49\u539f\u7406\uff1a", bold=True)
            for item in name_data["ziyi"]:
                para(item, indent=True)

        # Jingyi
        if "jingyi" in name_data:
            para("\u7ecf\u610f\u539f\u7406\uff1a", bold=True)
            for item in name_data["jingyi"]:
                para(item, indent=True)

        # Yinlv
        if "yinlv" in name_data:
            para("\u97f3\u5f8b\uff1a" + name_data["yinlv"], bold=True)
        if "yinlv_comment" in name_data:
            para(name_data["yinlv_comment"], indent=True)

        # Shuli
        if "shuli" in name_data:
            para("\u6570\u7406\u539f\u7406\uff1a", bold=True)
            for item in name_data["shuli"]:
                para(item, indent=True)

        # Chuangyi
        if "chuangyi" in name_data:
            para("\u521b\u610f\u5b9a\u4f4d\uff1a", bold=True)
            para(name_data["chuangyi"], indent=True)

        doc.add_paragraph()

    # Section 5: Summary
    heading("\u4e94\u3001\u59d3\u540d\u603b\u8bc4", 1)
    if "summary_title" in data:
        para(data["summary_title"], bold=True)
    for item in data.get("summary", []):
        para(item, indent=True)

    doc.add_paragraph()

    # Disclaimer
    if "disclaimer" in data:
        quote(data["disclaimer"])

    # Save
    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python generate_doc.py <input.json> [output.docx]")
        sys.exit(1)

    check_deps()

    json_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None

    out = generate(json_path, output_path)
    print(f"Document saved: {out}")