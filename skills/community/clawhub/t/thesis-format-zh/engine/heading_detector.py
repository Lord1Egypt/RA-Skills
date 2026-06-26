"""
标题检测与样式化模块

策略：正则匹配 + 字号 + 加粗 三重判断
支持 H1-H4 四级标题，可插拔检测策略
"""

import re
from docx.oxml.ns import qn


# ═══════════════════════════════════════════════
# 匹配模式（可按学校扩展）
# ═══════════════════════════════════════════════

# H1: 章标题
H1_PATTERNS = [
    r'^第[一二三四五六七八九十\d]+章',
    r'^绪\s*论$',
    r'^结\s*(论|语)$',
    r'^参考文献$',
    r'^致\s*谢$',
    r'^附\s*录',
    r'^摘\s*要$',
    r'^Abstract$',
]

# H2: 二级标题
H2_PATTERNS = [
    r'^[\d]+\.[\d]+',                    # 1.1 / 2.3
    r'^[一二三四五六七八九十]+、',       # 一、二、
]

# H3: 三级标题
H3_PATTERNS = [
    r'^[\d]+\.[\d]+\.[\d]+',            # 1.1.1 / 2.3.5
    r'^（[一二三四五六七八九十]+）',      # （一）（二）
]

# H4: 四级标题
H4_PATTERNS = [
    r'^[\d]+\.[\d]+\.[\d]+\.[\d]+',     # 1.1.1.1
]


# ═══════════════════════════════════════════════
# 工具函数
# ═══════════════════════════════════════════════

def _get_para_font_info(para):
    """提取段落字号信息: (max_size_pt, has_bold, has_explicit_bold)"""
    runs = [r for r in para.runs if r.text.strip()]
    if not runs:
        return (0, False, False)
    sizes = [r.font.size / 12700 for r in runs if r.font.size]
    max_size = max(sizes) if sizes else 0
    has_explicit_bold = any(r.bold is True for r in runs)
    has_bold_style = any(r.bold is not False for r in runs)
    return (max_size, has_bold_style, has_explicit_bold)


def _is_heading_candidate(para):
    """判断段落是否是潜在标题（短文本 + 有字号）"""
    text = para.text.strip()
    if not text or len(text) > 80:
        return False
    # 跳过目录条目（含 tab + 页码的）
    if '\t' in text:
        return False
    # 跳过已有 Heading 样式的段落（已处理过）
    if para.style.name.startswith('Heading'):
        return False
    max_size, _, _ = _get_para_font_info(para)
    return max_size > 0


# ═══════════════════════════════════════════════
# 核心检测
# ═══════════════════════════════════════════════

def classify_heading(para):
    """
    根据文本内容和字号判断标题级别。

    返回: (level, confidence)  level为1/2/3/4 或 None, confidence为0-1
    """
    text = para.text.strip()
    max_size, has_bold, explicit_bold = _get_para_font_info(para)

    # H1: 章标题（通常加粗15pt）
    for pat in H1_PATTERNS:
        if re.match(pat, text):
            if max_size >= 14 and explicit_bold:
                return (1, 0.95)
            elif max_size >= 14:
                return (1, 0.8)
            elif max_size >= 12:
                return (1, 0.6)

    # H4: 精确匹配 1.1.1.1（必须在 H3 之前检测）
    for pat in H4_PATTERNS:
        if re.match(pat, text):
            if max_size >= 11:
                return (4, 0.85)

    # H3: 精确匹配 1.1.1 或 （一）
    for pat in H3_PATTERNS:
        if re.match(pat, text):
            if max_size >= 11:
                return (3, 0.85)

    # H2: 匹配 1.1 / 一、
    for pat in H2_PATTERNS:
        if re.match(pat, text):
            if max_size >= 13:
                return (2, 0.9)
            elif max_size >= 11 and explicit_bold:
                return (2, 0.7)

    return (None, 0)


def detect_and_apply(doc, start_idx=None, confidence_threshold=0.6):
    """
    扫描文档，识别手动格式化的标题，替换为 Heading 1/2/3/4 样式。

    Args:
        doc: python-docx Document 对象
        start_idx: 从该段落索引开始检测（None 则自动找正文起始）
        confidence_threshold: 置信度阈值，低于此值跳过

    Returns:
        dict: {"h1": n, "h2": n, "h3": n, "h4": n, "skipped": n}
    """
    if start_idx is None:
        start_idx = find_body_start(doc)
    stats = {"h1": 0, "h2": 0, "h3": 0, "h4": 0, "skipped": 0}

    for i, para in enumerate(doc.paragraphs):
        if i < start_idx:
            continue
        # 跳过目录段落
        if para.style.name.startswith(('toc', 'TOC')):
            continue
        # 跳过已经是标题样式的段落
        if para.style.name.startswith('Heading'):
            continue

        if not _is_heading_candidate(para):
            continue

        level, confidence = classify_heading(para)
        if level is None:
            continue
        if confidence < confidence_threshold:
            stats["skipped"] += 1
            continue

        style_name = f"Heading {level}"
        try:
            para.style = doc.styles[style_name]
            stats[f"h{level}"] += 1
            # 清除手动加粗，让样式控制
            for r in para.runs:
                r.bold = None
        except KeyError:
            stats["skipped"] += 1

    return stats


def find_body_start(doc):
    """找到正文起始段落索引（跳过封面/目录，找第一个真正的章标题）"""
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        # 跳过目录条目（含 tab 页码）
        if '\t' in text:
            continue
        # 优先匹配章标题
        if re.match(r'^第[一二三四五六七八九十\d]+章', text):
            return i
    # 兜底：找无 tab 的摘要
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if '\t' not in text and re.match(r'^摘\s*要$', text):
            return i
    return 0
