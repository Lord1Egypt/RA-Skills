#!/usr/bin/env python3
"""
Word文档转Markdown转换器
使用python-docx读取.docx文件并转换为Markdown格式
"""

import sys
import re
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

# 尝试导入python-docx，如果不可用则使用备用方法
try:
    from docx import Document
    from docx.shared import Pt
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("警告: python-docx未安装，将使用备用方法")

def extract_text_from_docx_simple(docx_path):
    """使用zip和XML解析提取文本（备用方法）"""
    text_content = []
    
    try:
        with zipfile.ZipFile(docx_path, 'r') as z:
            # 读取document.xml
            xml_content = z.read('word/document.xml')
            root = ET.fromstring(xml_content)
            
            # Word文档命名空间
            namespaces = {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            }
            
            # 提取所有段落
            for paragraph in root.findall('.//w:p', namespaces):
                para_text = []
                for text_elem in paragraph.findall('.//w:t', namespaces):
                    if text_elem.text:
                        para_text.append(text_elem.text)
                
                if para_text:
                    text_content.append(''.join(para_text))
                else:
                    text_content.append('')  # 空段落
                    
    except Exception as e:
        print(f"简单提取失败: {e}")
        return None
    
    return '\n'.join(text_content)

def convert_docx_to_markdown(docx_path, output_path=None):
    """将Word文档转换为Markdown"""
    
    if output_path is None:
        output_path = str(Path(docx_path).with_suffix('.md'))
    
    print(f"正在读取: {docx_path}")
    
    if HAS_DOCX:
        try:
            doc = Document(docx_path)
            
            md_lines = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    md_lines.append('')
                    continue
                
                # 检测标题样式
                style_name = para.style.name if para.style else "Normal"
                
                if style_name.startswith('Heading 1') or style_name.startswith('标题 1'):
                    md_lines.append(f"# {text}")
                elif style_name.startswith('Heading 2') or style_name.startswith('标题 2'):
                    md_lines.append(f"## {text}")
                elif style_name.startswith('Heading 3') or style_name.startswith('标题 3'):
                    md_lines.append(f"### {text}")
                elif style_name.startswith('Heading 4') or style_name.startswith('标题 4'):
                    md_lines.append(f"#### {text}")
                else:
                    # 普通段落
                    # 处理加粗和斜体
                    md_text = text
                    md_lines.append(md_text)
            
            # 处理表格
            for table in doc.tables:
                md_lines.append('')
                md_lines.append('| ' + ' | '.join([cell.text for cell in table.rows[0].cells]) + ' |')
                md_lines.append('|' + '|'.join(['---' for _ in table.rows[0].cells]) + '|')
                for row in table.rows[1:]:
                    md_lines.append('| ' + ' | '.join([cell.text for cell in row.cells]) + ' |')
                md_lines.append('')
            
            markdown_content = '\n'.join(md_lines)
            
        except Exception as e:
            print(f"python-docx读取失败: {e}")
            print("切换到备用方法...")
            markdown_content = extract_text_from_docx_simple(docx_path)
    else:
        markdown_content = extract_text_from_docx_simple(docx_path)
    
    if markdown_content:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        print(f"Markdown已保存: {output_path}")
        
        # 统计信息
        char_count = len(markdown_content)
        line_count = len(markdown_content.split('\n'))
        print(f"字符数: {char_count}, 行数: {line_count}")
        
        return output_path
    else:
        print("转换失败")
        return None

if __name__ == "__main__":
    if len(sys.argv) < 2:
        # 默认处理论文初稿v4.0.docx
        docx_file = "/Users/openclaw2026/.qclaw/workspace/workdoc/论文初稿v4.0.docx"
    else:
        docx_file = sys.argv[1]
    
    output_file = "/Users/openclaw2026/.qclaw/workspace/论文初稿v4.0.md"
    convert_docx_to_markdown(docx_file, output_file)
