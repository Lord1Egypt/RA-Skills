#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
将MD格式招标文件分析报告转换为DOCX格式
- 清理LaTeX公式
- 全文微软雅黑
- 无首行缩进
- 页码格式：P4, P9-P15
"""

# 全局字体配置（修改此处即可统一切换字体）
FONT_BODY = '微软雅黑'       # 正文/表格数据字体
FONT_HEADING = '微软雅黑'    # 标题字体

import re
import os
import sys
from collections import Counter
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def clean_latex_formula(text):
    """
    清理LaTeX公式，转换为纯文本
    """
    if not text or not isinstance(text, str):
        return text
    
    # 移除LaTeX公式分隔符 $$
    text = re.sub(r'\$\$(.*?)\$\$', lambda m: _convert_formula(m.group(1)), text, flags=re.DOTALL)
    
    # 移除行内LaTeX公式 $ ... $
    text = re.sub(r'\$(.*?)\$', lambda m: _convert_formula(m.group(1)), text)
    
    return text

def _convert_formula(latex_str):
    """将LaTeX公式片段转换为纯文本"""
    if not latex_str:
        return ''
    
    result = latex_str
    
    # 移除\text{}，保留内部文本
    result = re.sub(r'\\text\{([^}]*)\}', r'\1', result)
    
    # 转换常见数学符号
    result = result.replace('\\times', 'x')
    result = result.replace('\\div', '÷')
    result = result.replace('\\pm', '±')
    result = result.replace('\\geq', '>=')
    result = result.replace('\\leq', '<=')
    result = result.replace('\\neq', '≠')
    result = result.replace('\\approx', '≈')
    result = result.replace('\\infty', '∞')
    result = result.replace('\\alpha', 'α')
    result = result.replace('\\beta', 'β')
    result = result.replace('\\gamma', 'γ')
    result = result.replace('\\delta', 'δ')
    result = result.replace('\\Delta', 'Δ')
    result = result.replace('\\sum', 'Σ')
    result = result.replace('\\prod', 'Π')
    result = result.replace('\\int', '∫')
    
    # 转换分数 \frac{分子}{分母} 为 (分子/分母)
    result = re.sub(r'\\frac\{([^}]*)\}\{([^}]*)\}', r'(\1/\2)', result)
    
    # 转换上下标
    result = re.sub(r'\^(\d+)', r'^\1', result)  # 上标
    result = re.sub(r'_(\d+)', r'_\1', result)   # 下标
    
    # 移除多余的反斜杠
    result = result.replace('\\%', '%')
    result = result.replace('\\(', '(')
    result = result.replace('\\)', ')')
    
    # 移除LaTeX注释
    result = re.sub(r'%.*', '', result)
    
    return result.strip()

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    for edge in ('top', 'start', 'bottom', 'end', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            for key in ['sz', 'val', 'color', 'space', 'shadow']:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), edge_data[key])
            tcBorders.append(element)
    
    tc.append(tcBorders)

def set_cell_margins(cell, top=100, start=100, bottom=100, end=100):
    """设置单元格边距"""
    tc = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    
    margins = {'top': top, 'left': start, 'bottom': bottom, 'right': end}
    for margin, value in margins.items():
        mar = OxmlElement('w:{}'.format(margin))
        mar.set(qn('w:w'), str(value))
        mar.set(qn('w:type'), 'dxa')
        tcMar.append(mar)
    
    tc.append(tcMar)

def set_paragraph_no_indent(paragraph):
    """设置段落无首行缩进"""
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.left_indent = None
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(3)

def apply_font(run, font_name, font_size, bold=False, color=None):
    """
    完整设置 run 的字体，覆盖 ascii/hAnsi/eastAsia/cs 四个属性，
    防止 Word 在不同系统/版本下回退到 MS Gothic 等默认字体。
    """
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

    # 通过 XML 直接设置四个字体属性，确保不被 Word 主题字体覆盖
    rFonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    rFonts.set(qn('w:ascii'),    font_name)
    rFonts.set(qn('w:hAnsi'),   font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'),       font_name)  # 复杂脚本字体，缺失会导致回退 MS Gothic

def set_cell_text(cell, text, font_size=9.5, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    """
    安全地设置单元格文字及样式。
    不使用 cell.text = value（在部分 python-docx 版本中 runs 为空导致样式失效），
    改用 clear() + add_run() 确保 run 一定存在。
    """
    para = cell.paragraphs[0]
    para.clear()  # 清空段落内容（保留段落本身）
    para.alignment = align
    set_paragraph_no_indent(para)

    run = para.add_run(str(text) if text else '')
    apply_font(run, FONT_BODY, font_size, bold=bold, color=color)

def add_styled_table(doc, headers, rows):
    """添加带样式的表格（全文微软雅黑，无首行缩进）"""
    if not rows:
        print('[WARN] 表格无数据行，跳过')
        return None
    
    try:
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 表头样式
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            cell = hdr_cells[i]
            
            # 用 set_cell_text 写入，保证 run 存在且样式生效
            set_cell_text(cell, header, font_size=10.5, bold=True,
                          color=RGBColor(255, 255, 255),
                          align=WD_ALIGN_PARAGRAPH.CENTER)
            
            # 设置表头背景色（深蓝色）
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:val'), 'clear')
            shading_elm.set(qn('w:color'), 'auto')
            shading_elm.set(qn('w:fill'), '2E5496')
            cell._element.get_or_add_tcPr().append(shading_elm)
            
            # 垂直居中
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # 数据行样式
        for i, row_data in enumerate(rows):
            row_cells = table.rows[i + 1].cells
            
            # 确保数据行长度与表头一致
            while len(row_data) < len(headers):
                row_data.append('')
            
            for j, cell_data in enumerate(row_data[:len(headers)]):
                cell = row_cells[j]
                
                # 清理LaTeX公式
                cleaned_data = clean_latex_formula(str(cell_data))
                
                # 用 set_cell_text 写入，保证 run 存在且样式生效
                set_cell_text(cell, cleaned_data, font_size=9.5, bold=False,
                              align=WD_ALIGN_PARAGRAPH.LEFT)
                
                # 垂直居中
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                # 设置单元格边距
                set_cell_margins(cell)
        
        # 设置表格边框
        for row in table.rows:
            for cell in row.cells:
                set_cell_border(
                    cell,
                    top={"sz": "4", "val": "single", "color": "auto"},
                    bottom={"sz": "4", "val": "single", "color": "auto"},
                    left={"sz": "4", "val": "single", "color": "auto"},
                    right={"sz": "4", "val": "single", "color": "auto"},
                )
        
        doc.add_paragraph()  # 表格后空一行
        return table
    
    except Exception as e:
        print('[ERROR] 添加表格失败 - {}'.format(str(e)))
        import traceback
        traceback.print_exc()
        return None

def _is_page_col(header_text):
    """判断某列表头是否为'页码'列"""
    if not header_text:
        return False
    return '页码' in header_text or 'page' in header_text.lower() or '页数' in header_text


def _is_page_value(text):
    """判断某单元格内容是否像页码（P5、P9-P15、P4、P8、P15 等）"""
    if not text:
        return False
    return bool(re.search(r'P\d', text))


def _is_seq_value(text):
    """判断某单元格内容是否像序号（纯数字或 X.Y 格式）"""
    if not text:
        return False
    return bool(re.match(r'^\d+(\.\d+)?$', text.strip())) and len(text.strip()) <= 6


def _align_short_row(cells, target_cols, headers, complete_rows):
    """
    智能对齐短行：在缺失的列位置插入空单元格。

    策略优先级：
    1. 页码列对齐：末列是页码列且末单元格像页码 → 页码放末列，中间补空
    2. 参考行+表头语义对齐：
       a) 基于完整行统计哪些列常为空
       b) 叠加表头语义（"备注"等列空列可能性高）
       c) 结合当前行内容模式（页码值→页码列，序号值→序号列）调整
    3. 兜底：末尾补空
    """
    deficit = target_cols - len(cells)
    if deficit <= 0:
        return cells[:target_cols]

    # --- 策略1：页码列对齐 ---
    if len(headers) > 0 and _is_page_col(headers[-1]) and len(cells) > 0 and _is_page_value(cells[-1]):
        non_page = cells[:-1]
        page_val = cells[-1]
        padding_needed = target_cols - len(non_page) - 1
        normalized = non_page + [''] * padding_needed + [page_val]
        print('[INFO] 短行对齐(页码策略): {}列→{}列, 页码放末列'.format(len(cells), target_cols))
        return normalized

    # --- 策略2：参考行+表头语义对齐 ---
    emptiness = [0.0] * target_cols

    # 2a: 基于表头语义的先验得分（始终计算，作为基础）
    for col_idx in range(min(len(headers), target_cols)):
        h = headers[col_idx].strip()
        if '备注' in h or '说明' in h or '其他' in h:
            emptiness[col_idx] = 0.7   # 高先验：备注列常为空
        elif _is_page_col(h):
            emptiness[col_idx] = 0.02  # 低先验：页码列很少为空
        elif '序号' in h or '编号' in h:
            emptiness[col_idx] = 0.01  # 很低先验：序号列几乎不为空
        else:
            emptiness[col_idx] = 0.2   # 默认先验

    # 2b: 叠加参考行的实际空列频率（有参考行时增强判断）
    if complete_rows:
        for col_idx in range(target_cols):
            empty_count = sum(1 for r in complete_rows if col_idx < len(r) and not r[col_idx].strip())
            freq = empty_count / len(complete_rows)
            # 用加权平均融合：先验 40% + 实际频率 60%
            emptiness[col_idx] = 0.4 * emptiness[col_idx] + 0.6 * freq

    # 2c: 结合当前行内容模式调整得分
    # 如果某单元格明显属于某列，降低该列的空列可能性
    for cell in cells:
        cell_text = cell.strip()
        if _is_page_value(cell_text):
            for col_idx in range(min(len(headers), target_cols)):
                if _is_page_col(headers[col_idx]):
                    emptiness[col_idx] = 0.0
        elif _is_seq_value(cell_text):
            for col_idx in range(min(len(headers), target_cols)):
                if '序号' in headers[col_idx] or '编号' in headers[col_idx]:
                    emptiness[col_idx] = 0.0

    # 选择得分最高的 deficit 个位置作为空列插入点
    # 同分时优先选择靠后的位置（空列如"备注"通常靠后）
    scored_positions = list(enumerate(emptiness))
    scored_positions.sort(key=lambda x: (-x[1], -x[0]))
    insert_positions = set(idx for idx, score in scored_positions[:deficit])

    # 构建对齐后的行
    result = []
    cell_idx = 0
    for col_idx in range(target_cols):
        if col_idx in insert_positions:
            result.append('')
        else:
            if cell_idx < len(cells):
                result.append(cells[cell_idx])
                cell_idx += 1
            else:
                result.append('')

    strategy = '参考行+语义' if complete_rows else '表头语义'
    print('[INFO] 短行对齐({}策略): {}列→{}列, 空列插入位置={}, 得分={}'.format(
        strategy, len(cells), target_cols, sorted(insert_positions),
        [round(emptiness[i], 2) for i in sorted(insert_positions)]))
    return result


def parse_markdown_table(table_lines):
    """解析Markdown表格行，返回headers和rows。

    增强逻辑：
    - 以表头列数为标准列数（header_cols）
    - 如果某行列数 > header_cols，多余列合并到最后一列
    - 如果某行列数 < header_cols：
        a) 页码列对齐：末列是页码列且末单元格像页码 → 页码放末列，中间补空
        b) 参考行对齐：有完整行时，基于完整行分析哪些列常为空，将空列插入最可能的位置
        c) 表头语义对齐：无完整行时，基于表头语义（"备注"等列常为空）启发式插入
        d) 兜底：末尾补空
    - 如果表头列数 < 数据主流列数，自动扩展表头
    """
    if len(table_lines) < 2:
        return None, None

    # 解析表头（第一行）
    header_line = table_lines[0]
    headers = [cell.strip() for cell in header_line.split('|')[1:-1]]

    # 扫描所有数据行
    raw_rows = []
    for row_line in table_lines[2:]:
        if not row_line.strip():
            continue
        cells = [cell.strip() for cell in row_line.split('|')[1:-1]]
        cells = [clean_latex_formula(cell) for cell in cells]
        raw_rows.append(cells)

    if not raw_rows:
        return headers, []

    # 确定目标列数：优先用表头列数，但如果数据主流列数 > 表头列数则扩展
    col_counter = Counter(len(row) for row in raw_rows)
    most_common_data_cols = col_counter.most_common(1)[0][0]
    target_cols = max(len(headers), most_common_data_cols)

    # 扩展表头（如果数据主流列数 > 表头列数）
    if target_cols > len(headers):
        print('[WARN] 表格数据主流列数({}) > 表头列数({})，已自动扩展表头'.format(target_cols, len(headers)))
        headers.extend(['（未命名）{}'.format(i + 1) for i in range(len(headers), target_cols)])

    # 预先识别完整行（列数 == target_cols），用于参考行对齐
    complete_rows = [r for r in raw_rows if len(r) == target_cols]

    # 统一每一行的列数为 target_cols
    rows = []
    for cells in raw_rows:
        if len(cells) > target_cols:
            # 多余列合并到最后一列，用" | "分隔（保留所有内容）
            merged = ' | '.join(cells[target_cols - 1:])
            normalized = cells[:target_cols - 1] + [merged]
            rows.append(normalized)
        elif len(cells) < target_cols:
            # 调用智能对齐函数
            normalized = _align_short_row(cells, target_cols, headers, complete_rows)
            rows.append(normalized)
        else:
            rows.append(cells)

    return headers, rows

def convert_md_to_docx(md_path, docx_path):
    """将Markdown文件转换为DOCX"""
    print('[INFO] 正在读取Markdown文件：{}'.format(md_path))
    
    try:
        with open(md_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
    except Exception as e:
        print('[ERROR] 无法读取文件 {} - {}'.format(md_path, str(e)))
        return None
    
    print('[INFO] 正在创建DOCX文档...')
    doc = Document()
    
    # 设置默认字体（全文微软雅黑）
    style = doc.styles['Normal']
    style.font.name = FONT_BODY
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    style.paragraph_format.first_line_indent = None  # 无首行缩进
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(3)
    # Normal 样式的东亚字体也要设置
    style.element.rPr.get_or_add_rFonts().set(qn('w:eastAsia'), FONT_BODY)
    
    # 按行处理
    i = 0
    table_count = 0
    paragraph_count = 0
    
    while i < len(lines):
        line = lines[i].rstrip('\n')
        
        # 跳过空行
        if not line.strip():
            i += 1
            continue
        
        # 检测表格开始（当前行以|开头，且下一行也以|开头）
        if line.strip().startswith('|') and i + 1 < len(lines) and lines[i + 1].strip().startswith('|'):
            # 收集表格的所有行
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].rstrip('\n'))
                i += 1
            
            # 解析表格
            headers, rows = parse_markdown_table(table_lines)
            
            if headers and rows:
                add_styled_table(doc, headers, rows)
                table_count += 1
                print('[INFO] 已添加表格 {}：{} 行'.format(table_count, len(rows)))
            else:
                print('[WARN] 表格解析失败，跳过')
            
            continue  # 跳过下面的 i += 1
        
        # 检测标题
        elif line.startswith('# '):
            heading = doc.add_heading(line[2:], level=1)
            set_paragraph_no_indent(heading)
            for run in heading.runs:
                apply_font(run, FONT_HEADING, 14, bold=True)
            paragraph_count += 1
        
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:], level=2)
            set_paragraph_no_indent(heading)
            for run in heading.runs:
                apply_font(run, FONT_HEADING, 12, bold=True)
            paragraph_count += 1
        
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:], level=3)
            set_paragraph_no_indent(heading)
            for run in heading.runs:
                apply_font(run, FONT_HEADING, 11, bold=True)
            paragraph_count += 1
        
        # 普通段落
        else:
            p = doc.add_paragraph()
            set_paragraph_no_indent(p)
            run = p.add_run(line)
            apply_font(run, FONT_BODY, 10.5)
            paragraph_count += 1
        
        i += 1
    
    # 保存文档
    print('[INFO] 正在保存DOCX文档：{}'.format(docx_path))
    try:
        doc.save(docx_path)
        print('[OK] DOCX报告已生成：{}'.format(docx_path))
        print('[INFO] 统计：{} 个表格，{} 个段落'.format(table_count, paragraph_count))
        return docx_path
    except Exception as e:
        print('[ERROR] 保存文档失败 - {}'.format(str(e)))
        import traceback
        traceback.print_exc()
        return None

def main():
    """主函数"""
    if len(sys.argv) < 2:
        print('Usage: py md_to_docx.py <Markdown文件路径> [输出DOCX路径]')
        print('Example: py md_to_docx.py "input.md" "output.docx"')
        sys.exit(1)
    
    md_path = sys.argv[1]
    
    if not os.path.exists(md_path):
        print('[ERROR] 找不到文件: {}'.format(md_path))
        sys.exit(1)
    
    # 生成输出文件名
    if len(sys.argv) >= 3:
        docx_path = sys.argv[2]
    else:
        base_name = os.path.splitext(md_path)[0]
        docx_path = base_name + '.docx'
    
    print('=' * 60)
    print('Markdown to DOCX Converter')
    print('=' * 60)
    print('Input: {}'.format(md_path))
    print('Output: {}'.format(docx_path))
    print('-' * 60)
    
    output_file = convert_md_to_docx(md_path, docx_path)
    
    if output_file and os.path.exists(output_file):
        size_kb = os.path.getsize(output_file) / 1024
        print('-' * 60)
        print('[OK] Document generated successfully!')
        print('  Location: {}'.format(output_file))
        print('  File size: {:.1f} KB'.format(size_kb))
        print('=' * 60)
    else:
        print('[FAIL] Document generation failed!')
        sys.exit(1)

if __name__ == '__main__':
    main()
