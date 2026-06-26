#!/usr/bin/env python3
"""
Auto Dimension Report Skill - Core Script v3
Functions:
  PDF -> DOCX (embedded image + path reference dual indexing)
  DOCX/XLSX image extraction + path index generation
  output/_ImageIndex.xlsx summary of all images
"""

import argparse
import json
import os
import re
import shutil
import sys
import zipfile
from io import BytesIO
from pathlib import Path


# -- Garbage text detection ------------------------------------------------
def _is_garbage_text(text: str) -> bool:
    """
    Determine if extracted text is "garbage" (hidden error OCR layer or garbled).
    Two signals — either qualifies as garbage:
    1. Avg line length short (< 20) and median line length < 8
    2. Readable CJK+letter+digit char ratio below 30% (QUALITY_THRESHOLD)
    """
    if not text:
        return True
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    if not lines:
        return True
    # Signal 1: avg line length + median line length
    avg_line_len = sum(len(l) for l in lines) / len(lines)
    median_len = sorted(len(l) for l in lines)[len(lines) // 2]
    if avg_line_len < 20 and median_len < 8:
        return True
    # Signal 2: readable char ratio
    import re as _re
    good = len(_re.findall(r'[\u4e00-\u9fff\uff00-\uffefa-zA-Z0-9]', text))
    if good / len(text) < 0.3:
        return True
    return False

import fitz
from docx import Document as DocxDoc
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

# ============================================================
# Configuration
# ============================================================
SUPPORTED_EXTS = {".pdf", ".docx", ".xlsx", ".xlsm"}
KEEP_AS_IS = {".docx", ".xlsx", ".xlsm"}
EXCLUDE_DIR_FRAGMENTS = ["副本", "复制", "Copy", "output", "image"]
IMAGE_INDEX_FILE = "_ImageIndex.xlsx"


# ============================================================
# Utility functions
# ============================================================

def ensure_dir(path: Path):
    path.mkdir(parents=True, exist_ok=True)


def clean_xml_text(text: str) -> str:
    text = re.sub(r'[\x00\x08\x0b\x0c\x0e-\x1f]', '', text)
    text = re.sub(r'[\x7f-\x9f]', '', text)
    return text


def add_para(doc, text, bold=False, font_size=10.5,
             font_name="Microsoft YaHei", font_fallback="Consolas",
             space_before=None, space_after=None):
    text = clean_xml_text(text)
    if not text or not text.strip():
        return None
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(font_size)
    r.font.name = font_name
    r.bold = bold
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def add_image_ref(doc, rel_path, font_size=8):
    """Add image path reference annotation in DOCX."""
    from docx.shared import RGBColor
    text = clean_xml_text(f"[Image Reference] {rel_path}")
    if not text:
        return
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(font_size)
    r.font.name = "Consolas"
    r.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    return p


def is_source_file(f: Path, base_dir: Path) -> bool:
    if not f.is_file():
        return False
    ext = f.suffix.lower()
    if ext not in SUPPORTED_EXTS:
        return False
    if ext == ".docx" and f.name.lower().endswith(".pdf.docx"):
        return False
    try:
        rel = f.relative_to(base_dir)
        for part in rel.parts:
            for frag in EXCLUDE_DIR_FRAGMENTS:
                if frag.lower() in part.lower():
                    return False
    except ValueError:
        pass
    return True


# ============================================================
# Image extraction
# ============================================================

def extract_from_zip(zip_path, image_subdir, media_prefix):
    """Extract images from ZIP archive, return [(out_name, out_path), ...]."""
    results = []
    try:
        with zipfile.ZipFile(zip_path, "r") as zf:
            media_files = sorted([n for n in zf.namelist() if n.startswith(media_prefix)])
            for idx, mp in enumerate(media_files, 1):
                ext = os.path.splitext(mp)[1] or ".png"
                out_name = f"{zip_path.name}-{idx}{ext}"
                out_path = image_subdir / out_name
                if not out_path.exists():
                    with zf.open(mp) as src, open(out_path, "wb") as dst:
                        shutil.copyfileobj(src, dst)
                results.append((out_name, out_path))
    except Exception as e:
        print(f"    [Warning] Image extraction failed: {e}")
    return results


def extract_pdf_images(doc, pdf_path, image_subdir):
    """Extract embedded images from PDF, return [(img_name, img_path), ...]."""
    results = []
    stem = pdf_path.name
    for pn in range(len(doc)):
        page = doc[pn]
        images = page.get_images(full=True)
        for idx, img in enumerate(images, 1):
            xref = img[0]
            bi = doc.extract_image(xref)
            img_bytes = bi["image"]
            img_ext = bi["ext"]
            img_name = f"{stem}-p{pn+1}-img{idx}.{img_ext}"
            img_path = image_subdir / img_name
            if not img_path.exists():
                with open(img_path, "wb") as f:
                    f.write(img_bytes)
            results.append((img_name, img_path, pn + 1, idx, img_bytes, img_ext))
    return results


# ============================================================
# PDF -> DOCX (embedded image + path reference dual indexing)
# ============================================================

def page_to_blocks(page):
    blocks = []
    # Text
    text_blocks = page.get_text("dict")["blocks"]
    for tb in text_blocks:
        if tb["type"] == 0:
            lines = tb.get("lines", [])
            for line in lines:
                spans = line.get("spans", [])
                line_text = ""
                line_fs = 10.5
                for sp in spans:
                    t = sp.get("text", "").strip()
                    if t:
                        line_text += t
                        line_fs = sp.get("size", line_fs)
                if line_text.strip():
                    blocks.append({
                        "type": "text", "y": line["bbox"][1],
                        "content": line_text.strip(), "font_size": line_fs,
                    })
    # Tables
    try:
        tables = page.find_tables()
        for table in tables:
            bbox = table.bbox
            rows = table.extract()
            if rows and any(any(c and str(c).strip() for c in row) for row in rows):
                blocks.append({"type": "table", "y": bbox[1], "content": rows})
    except Exception:
        pass
    # Images - keep only position info; actual images passed externally
    images = page.get_images(full=True)
    for idx, img in enumerate(images, 1):
        try:
            rects = page.get_image_bbox(img)
            if rects:
                blk = {
                    "type": "image", "img_idx": idx, "page_num": page.number + 1,
                    "xref": img[0],
                }
                if isinstance(rects, list):
                    blk["y"] = rects[0].y0
                    blk["bbox"] = rects[0]
                else:
                    blk["y"] = rects.y0
                    blk["bbox"] = rects
                blocks.append(blk)
        except Exception:
            pass
    blocks.sort(key=lambda x: x["y"])
    return blocks


def _ensure_embed_format(img_bytes, img_ext):
    """Convert image to python-docx supported format (JPEG/PNG/GIF/BMP/TIFF)."""
    supported = {"jpeg", "jpg", "png", "gif", "bmp", "tiff", "tif"}
    ext = img_ext.lower()
    if ext in supported:
        return img_bytes, ext
    # Unsupported format — convert to PNG
    try:
        from PIL import Image as PilImage
        stream = BytesIO(img_bytes)
        pil_img = PilImage.open(stream)
        out = BytesIO()
        pil_img.save(out, format="PNG")
        return out.getvalue(), "png"
    except Exception:
        return img_bytes, img_ext


def convert_pdf_to_docx(pdf_path, docx_path, image_subdir, task_dir):
    """
    PDF -> DOCX: embedded images (for human viewing) + path references (for agent).
    Returns [(img_name, rel_path), ...]
    """
    from PIL import Image as PilImage
    doc = fitz.open(str(pdf_path))
    stem = pdf_path.name
    img_records = []  # record all image paths

    # Pre-extract all images (one-time)
    all_images = []
    for pn in range(len(doc)):
        page = doc[pn]
        for idx, img in enumerate(page.get_images(full=True), 1):
            xref = img[0]
            bi = doc.extract_image(xref)
            raw_bytes = bi["image"]
            raw_ext = bi["ext"]
            # Ensure format is compatible with DOCX embedding
            img_bytes, img_ext = _ensure_embed_format(raw_bytes, raw_ext)
            img_name = f"{stem}-p{pn+1}-img{idx}.{img_ext}"
            img_path = image_subdir / img_name
            if not img_path.exists():
                with open(img_path, "wb") as f:
                    f.write(img_bytes)
            # Relative path: from output/ to image/
            rel_path = str(Path("image") / stem / img_name)
            all_images.append({
                "page": pn + 1, "idx": idx, "name": img_name,
                "path": img_path, "rel": rel_path,
                "bytes": img_bytes,
            })
            print(f"    [Image] {img_name}")

    # Per-page detection: aggregate all text, then run garbage detection
    total_raw_text = sum(len(page.get_text().strip()) for page in doc)
    combined_text = '\n'.join(page.get_text().strip() for page in doc)
    has_text = total_raw_text >= 50 and not _is_garbage_text(combined_text)

    word_doc = DocxDoc()
    section = word_doc.sections[0]
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    style = word_doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)

    if has_text:
        print(f"    [Type] Text-based PDF (text + tables + embedded images + path index)")
        for pn in range(len(doc)):
            if pn > 0:
                word_doc.add_page_break()
            page = doc[pn]
            blocks = page_to_blocks(page)

            add_para(word_doc, f"=== Page {pn+1} ===", bold=True, font_size=11)

            for blk in blocks:
                if blk["type"] == "text":
                    fs = blk.get("font_size", 10.5)
                    add_para(word_doc, blk["content"],
                             bold=fs >= 14, font_size=min(fs, 14))

                elif blk["type"] == "table":
                    rows = blk["content"]
                    if not rows:
                        continue
                    nc = max(len(r) for r in rows)
                    if nc == 0:
                        continue
                    tbl = word_doc.add_table(rows=len(rows), cols=nc)
                    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                    tbl.style = "Table Grid"
                    for ri, row in enumerate(rows):
                        for ci in range(nc):
                            ct = row[ci] if ci < len(row) else ""
                            cell = tbl.cell(ri, ci)
                            cell.text = ""
                            safe = clean_xml_text(str(ct))
                            if safe:
                                rn = cell.paragraphs[0].add_run(safe)
                                rn.font.size = Pt(8)
                                rn.font.name = "Microsoft YaHei"
                                if ri == 0:
                                    rn.bold = True
                    word_doc.add_paragraph()

                elif blk["type"] == "image":
                    # Find matching image
                    pg = blk["page_num"]
                    idx = blk["img_idx"]
                    matches = [im for im in all_images
                               if im["page"] == pg and im["idx"] == idx]
                    if matches:
                        im = matches[0]
                        # 1. Embed image (for human viewing); on failure, keep only path reference
                        try:
                            bbox = blk.get("bbox")
                            if bbox:
                                ratio = min((bbox.x1 - bbox.x0) / page.rect.width * 10.89, 10.3)
                                word_doc.add_picture(BytesIO(im["bytes"]),
                                                     width=Inches(max(ratio, 1.0)))
                            else:
                                word_doc.add_picture(BytesIO(im["bytes"]),
                                                     width=Inches(5.5))
                        except Exception as e:
                            pass
                        # 2. Path reference (for agent)
                        add_image_ref(word_doc, im["rel"])
                        img_records.append(im["rel"])
    else:
        print(f"    [Type] Scanned/image-based PDF (full-page render embedding + path index)")
        from docx.shared import RGBColor
        for pn in range(len(doc)):
            if pn > 0:
                word_doc.add_page_break()
            page = doc[pn]
            pix = page.get_pixmap(matrix=fitz.Matrix(200 / 72, 200 / 72))
            img_bytes = pix.tobytes("png")
            img_name = f"{stem}-p{pn+1}.png"
            img_path = image_subdir / img_name
            if not img_path.exists():
                pix.save(str(img_path))
            # Embed full-page image
            word_doc.add_picture(BytesIO(img_bytes), width=Inches(10.89))
            rel_path = str(Path("image") / stem / img_name)
            add_image_ref(word_doc, rel_path)
            img_records.append(rel_path)

    doc.close()
    word_doc.save(str(docx_path))
    print(f"    [Done] -> {docx_path.name} (embedded image + path index)")
    return img_records


# ============================================================
# Image index generation
# ============================================================

def build_image_index(output_dir, task_dir):
    """Generate output/_ImageIndex.xlsx recording all image relative paths."""
    try:
        from openpyxl import Workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "Image Index"
        ws.append(["Source File", "Image Relative Path", "Image Absolute Path"])

        rows = []
        # Scan all files in output/
        for f in sorted(output_dir.iterdir()):
            if f.is_dir() or f.name.startswith("_") or f.suffix.lower() == ".json":
                continue
            stem = f.name
            # For .pdf.docx, the source file is .pdf
            if f.name.lower().endswith(".pdf.docx"):
                src_name = f.name[:-len(".docx")]  # strip .docx to get original pdf name
            else:
                src_name = f.name
            img_dir = task_dir / "image" / src_name
            if img_dir.exists():
                for img in sorted(img_dir.iterdir()):
                    rel = str(Path("image") / src_name / img.name)
                    rows.append([src_name, rel, str(img.resolve())])

        for r in rows:
            ws.append(r)

        # Column widths
        ws.column_dimensions["A"].width = 50
        ws.column_dimensions["B"].width = 80
        ws.column_dimensions["C"].width = 120

        idx_path = output_dir / IMAGE_INDEX_FILE
        wb.save(str(idx_path))
        print(f"  [Index] Generated {idx_path.name} ({len(rows)} records)")
    except Exception as e:
        print(f"  [Warning] Index file generation failed: {e}")
        return []


# ============================================================
# File processing
# ============================================================

def process_file(src_path, output_dir, image_dir, task_dir):
    stem = src_path.name
    ext = src_path.suffix.lower()
    img_subdir = image_dir / stem
    ensure_dir(img_subdir)

    print(f"\n{'='*60}")
    print(f"Source file: {stem}")
    print(f"{'='*60}")

    img_records = []

    if ext == ".pdf":
        out_path = output_dir / f"{stem}.docx"
        if out_path.exists():
            print(f"  [Skip] Target already exists: {out_path.name}")
            return
        print(f"  [Process] PDF -> DOCX (embedded image + path index)")
        img_records = convert_pdf_to_docx(src_path, out_path, img_subdir, task_dir)

    elif ext in KEEP_AS_IS:
        out_path = output_dir / stem
        if out_path.exists():
            print(f"  [Skip] Target already exists: {stem}")
        else:
            shutil.copy2(src_path, out_path)
            print(f"  [Copy] -> {stem} (retained embedded images)")
        # Extract images
        if ext == ".docx":
            imgs = extract_from_zip(src_path, img_subdir, "word/media/")
        else:
            imgs = extract_from_zip(src_path, img_subdir, "xl/media/")
        print(f"  Extracted embedded images: {len(imgs)} -> image/{stem}/")
        for name, path in imgs:
            rel = str(Path("image") / stem / name)
            img_records.append(rel)
            print(f"    [Image] {rel}")

    print(f"  [Output] -> output/{out_path.name}")
    return img_records


# ============================================================
# Main entry point
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description="Auto Dimension Report Skill v3 - PDF -> DOCX (dual index) & image extraction & index summary"
    )
    parser.add_argument("--dir", "-d", required=True, help="Task folder path")
    parser.add_argument("--skip-index", action="store_true", help="Skip generating _ImageIndex.xlsx")
    args = parser.parse_args()

    task_dir = Path(args.dir)
    if not task_dir.exists():
        print(f"[Error] Directory does not exist: {task_dir}")
        sys.exit(1)

    output_dir = task_dir / "output"
    image_dir = task_dir / "image"
    ensure_dir(output_dir)
    ensure_dir(image_dir)

    print(f"=" * 60)
    print(f"Auto Dimension Report Skill v3 - Extraction & Conversion + Dual Index")
    print(f"=" * 60)
    print(f"Task directory: {task_dir}")
    print(f"  output/ -> Processed files (embedded images + path index)")
    print(f"  image/  -> Images stored in per-source-file subdirectories")
    print(f"  _ImageIndex.xlsx -> Summary index of all images")
    print()

    all_files = sorted(f for f in task_dir.rglob("*") if is_source_file(f, task_dir))
    if not all_files:
        print("[Info] No processable source files found")
        return

    print(f"Found {len(all_files)} source files")
    for f in all_files:
        print(f"  - {f.relative_to(task_dir)}")

    for src in all_files:
        process_file(src, output_dir, image_dir, task_dir)

    # Generate summary index
    if not args.skip_index:
        build_image_index(output_dir, task_dir)

    out_count = len(list(output_dir.iterdir())) if output_dir.exists() else 0
    img_subdirs = [d for d in image_dir.iterdir() if d.is_dir()] if image_dir.exists() else []
    total_img = sum(len(list(d.iterdir())) for d in img_subdirs)

    print(f"\n{'='*60}")
    print(f"All processing complete!")
    print(f"  output/ files: {out_count} (all files include index markers)")
    print(f"  image/ subdirectories: {len(img_subdirs)}")
    print(f"  Total images: {total_img}")
    print(f"{'='*60}")

    # -- Print paths (Windows-friendly for easy copy) --
    print()
    print("[File Locations] (Copy and paste into File Explorer)")
    print(f"  Source folder:  {task_dir}")
    print(f"  output/:       {output_dir}")
    print(f"  image/:        {image_dir}")
    print()
    print("[Tip] Select the path above -> Ctrl+C -> paste into File Explorer address bar")
    print()

if __name__ == "__main__":
    main()
