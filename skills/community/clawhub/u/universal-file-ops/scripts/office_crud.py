#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
office_crud.py — Office 文件增删查改（.docx / .xlsx）

依赖（可选，自动降级）：
  python-docx  ：处理 .docx
  openpyxl    ：处理 .xlsx

如果依赖缺失，read 操作返回二进制提示，create/update 操作报错引导安装。

标准化 IO 接口：同 text_crud.py（JSON 到 stdout）
"""

import argparse
import json
import os
import sys
import base64
from typing import Any, Dict, Optional

sys.path.insert(0, os.path.dirname(__file__))
from utils import (
    read_input,
    success_output,
    error_output,
    backup_file,
    log_operation,
    ensure_data_dirs,
    is_safe_path,
    print_output,
)

# ── 依赖检测 ────────────────────────────────────────────────────────────────

HAS_DOCX = False
HAS_XLSX = False

try:
    import docx
    HAS_DOCX = True
except ImportError:
    pass

try:
    import openpyxl
    HAS_XLSX = True
except ImportError:
    pass


# ── .docx 操作 ─────────────────────────────────────────────────────────────

def docx_read(file_path: str) -> Dict[str, Any]:
    if not HAS_DOCX:
        raise RuntimeError("需要 python-docx：pip install python-docx")
    doc = docx.Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs]
    full_text = "\n".join(paragraphs)
    return {
        "format": "docx",
        "paragraphs_count": len(paragraphs),
        "text": full_text,
        "images_count": len(doc.inline_shapes),
    }


def docx_create(file_path: str, text: str, backup: bool = True) -> Dict[str, Any]:
    if not HAS_DOCX:
        raise RuntimeError("需要 python-docx：pip install python-docx")
    rollback_id = backup_file(file_path) if (backup and os.path.exists(file_path)) else None
    os.makedirs(os.path.dirname(os.path.abspath(file_path)), exist_ok=True)
    doc = docx.Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(file_path)
    return {
        "format": "docx",
        "backup_file": rollback_id,
        "size": os.path.getsize(file_path),
    }


def docx_update(file_path: str, text: str, backup: bool = True) -> Dict[str, Any]:
    if not HAS_DOCX:
        raise RuntimeError("需要 python-docx：pip install python-docx")
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")
    rollback_id = backup_file(file_path) if backup else None
    doc = docx.Document(file_path)
    doc.paragraphs.clear()   # 简单实现：清空后重写
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(file_path)
    return {
        "format": "docx",
        "backup_file": rollback_id,
        "size": os.path.getsize(file_path),
    }


# ── .xlsx 操作 ─────────────────────────────────────────────────────────────

def xlsx_read(file_path: str) -> Dict[str, Any]:
    if not HAS_XLSX:
        raise RuntimeError("需要 openpyxl：pip install openpyxl")
    wb = openpyxl.load_workbook(file_path, data_only=True)
    sheets = {}
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            rows.append(list(row))
        sheets[sheet_name] = {
            "dimensions": ws.dimensions,
            "max_row": ws.max_row,
            "max_column": ws.max_column,
            "rows": rows[:100],   # 最多返回前 100 行
        }
    return {
        "format": "xlsx",
        "sheet_names": wb.sheetnames,
        "sheets": sheets,
    }


def xlsx_create(file_path: str, sheets: Dict = None, backup: bool = True) -> Dict[str, Any]:
    if not HAS_XLSX:
        raise RuntimeError("需要 openpyxl：pip install openpyxl")
    rollback_id = backup_file(file_path) if (backup and os.path.exists(file_path)) else None
    os.makedirs(os.path.dirname(os.path.abspath(file_path)), exist_ok=True)
    wb = openpyxl.Workbook()
    if sheets:
        for sheet_name, rows in sheets.items():
            ws = wb.create_sheet(title=sheet_name)
            for row in rows:
                ws.append(row)
        # 删除默认创建的空 Sheet
        if "Sheet" in wb.sheetnames:
            del wb["Sheet"]
    else:
        ws = wb.active
        ws.title = "Sheet1"
    wb.save(file_path)
    return {
        "format": "xlsx",
        "backup_file": rollback_id,
        "size": os.path.getsize(file_path),
    }


# ── 主分发 ─────────────────────────────────────────────────────────────────

def run(action: str, file_path: str, **kwargs) -> str:
    if not is_safe_path(file_path):
        return error_output(f"路径不合法: {file_path}", file_path)

    ext = os.path.splitext(file_path)[1].lower()
    backup = kwargs.get("backup", True)

    try:
        if ext == ".docx":
            if action == "read":
                result = docx_read(file_path)
            elif action == "create":
                result = docx_create(file_path, kwargs.get("content", ""), backup)
            elif action == "update":
                result = docx_update(file_path, kwargs.get("content", ""), backup)
            elif action == "delete":
                rollback_id = backup_file(file_path) if backup else None
                os.remove(file_path)
                result = {"format": "docx", "deleted": True, "backup_file": rollback_id}
            else:
                return error_output(f"不支持的 action: {action}", file_path)

        elif ext == ".xlsx":
            if action == "read":
                result = xlsx_read(file_path)
            elif action == "create":
                result = xlsx_create(file_path, kwargs.get("sheets"), backup)
            elif action == "update":
                return error_output("xlsx update 请使用 xlsx_create 覆盖写入", file_path)
            elif action == "delete":
                rollback_id = backup_file(file_path) if backup else None
                os.remove(file_path)
                result = {"format": "xlsx", "deleted": True, "backup_file": rollback_id}
            else:
                return error_output(f"不支持的 action: {action}", file_path)
        else:
            return error_output(f"不支持的文件格式: {ext}", file_path)

        rollback_id = result.get("backup_file", None)
        log_operation(action, file_path, True, rollback_id)
        return success_output(action, file_path, result, rollback_id)

    except Exception as e:
        log_operation(action, file_path, False, detail=str(e))
        return error_output(str(e), file_path)


# ── CLI / JSON 入口 ─────────────────────────────────────────────────────────

def parse_args():
    p = argparse.ArgumentParser(description="Office 文件增删查改（.docx / .xlsx）")
    p.add_argument("--action",   choices=["read", "create", "update", "delete"])
    p.add_argument("--file",      help="目标文件路径（.docx 或 .xlsx）")
    p.add_argument("--content",   default="", help="写入文本内容（docx）")
    p.add_argument("--no-backup", action="store_true")
    p.add_argument("--input",     help="JSON 输入文件")
    return p.parse_args()


def main():
    args = parse_args()
    input_data = {}
    if args.input:
        with open(args.input, "r", encoding="utf-8") as f:
            input_data = json.load(f)
    else:
        stdin_data = read_input()
        if stdin_data:
            input_data = stdin_data

    if input_data:
        action   = input_data.get("action")
        file_path = input_data.get("file")
        kwargs    = {k: v for k, v in input_data.items() if k not in ("action", "file")}
    else:
        if not args.action or not args.file:
            print_output(
                error_output("--action 和 --file 为必填参数（或使用 JSON 输入）", None),
                is_error=True,
            )
            sys.exit(1)
        action   = args.action
        file_path = args.file
        kwargs    = {
            "content": args.content,
            "backup": not args.no_backup,
        }

    result_json = run(action, file_path, **kwargs)
    print_output(result_json, is_error=('"success": false' in result_json.lower()))
    sys.exit(0 if '"success": true' in result_json.lower() else 1)


if __name__ == "__main__":
    main()
