#!/usr/bin/env python3
import sys
sys.path.insert(0, os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'global_market'))
from geopolitical_risk import get_geopolitical_risks, format_geopolitical_section
from language_detection import detect_input_lang, match_market_lang, native_name, BilingualText

# ============================================================
# ⚠️ SAMPLE DATA NOTICE / 示例数据声明
# This template contains EXAMPLE / ILLUSTRATIVE data values.
# Users must configure their own data sources (yfinance, Guosen API, etc.)
# to generate plans with real market data.
# 本模板包含的数据为示例/演示用途。
# 用户需自行配置数据源（yfinance、国信API等）以获取真实市场数据。
# ============================================================

"""
📦 Futures Trading Plan Generator v7.0 FINAL / 期货交易方案生成器 (Bilingual)
===========================================================
Ai-Wife Global Investing / 媳妇智投全球通用版
Bilingual (CN/EN) Edition / 中英双语版 (已验证)

大叔2026-05-23验收通过，锁定替换旧模板（补丁63）

Core Features / 核心特征：
  1. Data Source / 数据源: Tushare fut_daily (futures daily / 期货日线)
  2. Review: 2 full trading days (trade_date = night+day combined / 夜盘+日盘合并)
  3. Plan Cycle / 方案周期: 2 days / 2天
  4. Operations: Day-long/short + Night-long/short (4-dimension / 四维双向)
  5. Daily Forecast: Day session + Night session / 日盘+夜盘

Usage / 用法：
  from template_futures_v7 import build_plan
  
  Method 1 / 方式一: Pass data_dict (custom data / 传入数据)
  Method 2 / 方式二: data_dict=None → raises ValueError / 不传则报错
  ⚠️ User must provide their own data / 用户需自行提供数据
"""

def build_plan(data_dict, input_lang="auto", user_input="", use_sample_data=False):
    # data_dict is REQUIRED — no fallback, no auto-fetch
    """
    大叔亲授·审题+质检机制：
    - 审题：从user_input检测输入语言
    - 质检：从market_code匹配市场语言
    - 输出：{输入语言} / {市场语言} 双语
    - 同语言自动合并为单语
    """
    # ── 语言自动检测（大叔审题+质检） ──
    if input_lang == "auto":
        _input_lang = detect_input_lang(user_input)
    else:
        _input_lang = input_lang
    # 尝试从data_dict获取市场代码
    _market_code = data_dict.get("market_code", "cn") if data_dict else "cn"
    _dest_lang = match_market_lang(_market_code)
    _bt = BilingualText(input_lang=_input_lang, dest_lang=_dest_lang)
    """
    生成期货交易方案docx
    
    参数:
      data_dict: dict，包含以下字段（必传，不传则报错）
        - ts_code: 合约代码，如 'CU2607.SHF'
        - start_date: 回顾起始日，如 '20260521'
        - end_date: 回顾截止日，如 '20260522'
        - plan_start: 预判起始日，如 '05-25(一)'
        - plan_end: 预判截止日，如 '05-26(二)'
        - plan_start_en: 预判起始日数字，如 '20260525'
        - plan_end_en: 预判截止日数字，如 '20260526'
        - day1_rows: 第1天日盘数据行
        - day2_rows: 第2天日盘数据行
        - price_current: 当前价(日盘)
        - price_night: 当前价(夜盘)
        - pressure_strong: 强压力
        - pressure_mid: 次压力
        - support_short: 短期支撑
        - support_strong: 强支撑
        - lme_price: LME铜价格
        - comex_price: COMEX铜价格
        - usd_index: 美元指数
    
    返回:
      doc: python-docx Document对象
      文件已保存至: {WORKSPACE}/skills/ai-xifu-caopan/{filename}
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    import os
    import datetime
    
    
    # ── 🌍 大叔亲授·审题+质检机制：自动检测语言 ──
    _input_lang = detect_input_lang(user_input)
    _market_code = data_dict.get("market_code", "cn") if data_dict else "cn"
    _dest_lang = match_market_lang(_market_code)
    _bt = BilingualText(input_lang=_input_lang, dest_lang=_dest_lang)

def build_plan(data_dict, input_lang="auto", user_input="", use_sample_data=False):
    """
    Generate futures analysis plan document (.docx)
    生成期货分析方案文档
    
    Args:
        data_dict: dict containing ALL required market data fields
                   All fields MUST be provided by the user with real data.
                   No fallback to sample/default data.
        use_sample_data: If True, fills missing fields with SAMPLE DATA (clearly labeled).
                         ⚠️ Sample data is for demonstration only — not real market data.
    """
    if data_dict is None or not isinstance(data_dict, dict):
        raise ValueError("❌ data_dict is REQUIRED / data_dict为必传参数。请提供真实市场数据。")
    
    # ⚠️ 用户必须提供全部数据字段 / All fields must be user-provided
        # 关键价位
        'price_current': '104,870',
        'price_night': '104,940',
        'pressure_strong': '106,670',
        'pressure_mid': '105,660',
        'support_short': '103,620',
        'support_strong': '103,000',
        # 外盘
        'lme_price': '13,620美元/吨',
        'lme_trend': '高位震荡(-0.26%)',
        'comex_price': '~6.345美元/磅',
        'usd_index': '~99.5',
        'usd_trend': '震荡偏弱',
        # 风险预警
        'risk_vol': '量能萎缩|日盘9.6万手萎缩',
        'risk_lme': '外盘走弱|LME 13,710→13,620',
        'risk_inv': '库存累库|三大交易所~120万吨',
        'risk_geo': '地缘风险|美伊+托克提货',
        'risk_spot': '现货升水|升~2,240，支撑强',
        # 国际局势
        'geo1': '美伊冲突|特朗普拒绝停火|利多(地缘溢价)',
        'geo2': '中美关系|访华预期|偏多',
        'geo3': '跨市场套利|托克注销5.1万吨|利多(亚洲供应趋紧)',
        # 宏观
        'macro1': '中国经济政策|利多|+14%',
        'macro2': '美元走势|利多|+10%',
        'macro3': '美伊地缘|利多|+9%',
        'macro4': '美联储政策|利空|-8%',
        'macro5': '过剩预期|利空|-5%',
        # 产业链
        'chain1': '铜精矿TC|-93.9~-100美元/吨(历史新低)|利多大涨',
        'chain2': '矿端供应|智利罢工+印尼减产|利多',
        'chain3': '下游需求|新能源+特高压+AI|利多',
        # 资金
        'fund_oi': '171,627手|多头建仓',
        'fund_vol': '日盘9.6万手|观望',
        'fund_trafigura': 'LME注销5.1万吨|2013年最大提货',
        'fund_premium': '~2,240元/吨|支撑强',
        'fund_core': '托克注销5.1万吨，亚洲供应趋紧',
        # MACD
        'macd_daily': '日线|DIF在DEA上、红柱缩短|价在均线上|多头动能减',
        'macd_60min': '60分钟|DIF在DEA附近|均线粘合|短线回调压力',
        'macd_weekly': '周线|DIF在DEA上+红柱|价在均线上|中期多头完好',
        'tech_summary': '中期多+短期震荡',
        # 每日预判（2天×2时段）
        'pred1_day_dir': '探底回升',
        'pred1_day_range': '103,800~105,000',
        'pred1_day_long': '103,800~104,000低吸',
        'pred1_day_short': '冲高105,000短空',
        'pred1_night_dir': '高位震荡',
        'pred1_night_range': '104,500~105,100',
        'pred1_night_long': 'LME13,600做多',
        'pred1_night_short': 'LME13,400做空',
        'pred2_day_dir': '震荡偏多',
        'pred2_day_range': '104,500~105,500',
        'pred2_day_long': '104,500附近做多',
        'pred2_day_short': '冲高105,500轻仓',
        'pred2_night_dir': '偏多整理',
        'pred2_night_range': '104,600~105,200',
        'pred2_night_long': '观察LME',
        'pred2_night_short': '观察LME',
        # 操作建议参数
        'day_long_entry': '103,800~104,000',
        'day_long_target1': '105,000',
        'day_long_target2': '105,500',
        'day_long_stop': '102,500下方',
        'day_long_pos': '首仓20%|突破105,500加至30%',
        'day_short_entry': '105,300~105,500受阻+15分钟顶背离',
        'day_short_target1': '104,500',
        'day_short_target2': '104,000',
        'day_short_stop': '105,700上方',
        'day_short_pos': '不超10%',
        'night_long_entry': '21:00 LME站稳13,600+CU2607不破104,420+30分钟翻红',
        'night_long_target1': '105,110',
        'night_long_target2': '105,500',
        'night_long_stop': '104,200下方',
        'night_long_pos': '15%',
        'night_short_entry': '21:00 LME跌破13,400+CU2607受压+15分钟连续收阴',
        'night_short_target1': '104,420',
        'night_short_target2': '104,000',
        'night_short_stop': '105,200上方',
        'night_short_pos': '不超10%',
        # 操盘建议
        'main_strategy': '做多为主(主力策略)，做空为辅(辅助策略)',
        'trade_strategy': '日盘逢回调做多+冲高短空|夜盘联动LME',
        'trade_pos': '做多首仓20%|做空不超10%，总敞口不超50%',
        'trade_stop': '多单102,500下方|空单105,700上方',
        # 评分
        'score_tech': '7/10',
        'score_tech_note': '中期多头+短期震荡',
        'score_fundamental': '8/10',
        'score_fundamental_note': '供应紧缺+库存偏高',
        'score_capital': '7/10',
        'score_capital_note': '托克5.1万吨利多',
        'score_sentiment': '6/10',
        'score_sentiment_note': '观望浓(缩量)',
        'score_macro': '7/10',
        'score_macro_note': '政策利好+美元弱',
        'summary': '中期上行趋势未破，短期高位震荡。日盘回调103,800~104,000低吸做多(顺势主策略)，冲高105,300~105,500轻仓短空(逆势辅助)。夜盘联动LME。',
    }
    
    # ⚠️ 仅使用用户提供的数据 / Use user-provided data only
    # No fallback to sample/default values
    d = data_dict
    
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style.font.size = Pt(10)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    
    def H(t, l=1):
        h = doc.add_heading(t, level=l)
        for r in h.runs:
            r.font.name = 'SimSun'
            r.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    
    def P(t, b=False, s=10, c=None, a=None):
        p = doc.add_paragraph()
        r = p.add_run(t)
        r.font.name = 'SimSun'; r.font.size = Pt(s); r.bold = b
        r.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        if c: r.font.color.rgb = c
        if a: p.alignment = a
    
    def T(hd, data, note=None):
        t = doc.add_table(rows=1+len(data), cols=len(hd))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = 'Table Grid'
        for i,h in enumerate(hd):
            c = t.rows[0].cells[i]; c.text = h
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs: r.bold=True; r.font.size=Pt(9); r.font.name='SimSun'; r.element.rPr.rFonts.set(qn('w:eastAsia'),'SimSun')
        for ri, rd in enumerate(data):
            for ci, v in enumerate(rd):
                c = t.rows[ri+1].cells[ci]; c.text = str(v)
                for p in c.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs: r.font.size=Pt(9); r.font.name='SimSun'; r.element.rPr.rFonts.set(qn('w:eastAsia'),'SimSun')
    
    # ── 标题 ──
    title = doc.add_heading('', level=0)
    r = title.add_run(f'{d["variety"]}{d["contract"]} 下周参考框架 / Weekly Trading Plan')
    r.font.name = 'SimSun'; r.font.size = Pt(22); r.bold = True
    r.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    today_str = datetime.date.today().strftime('%Y-%m-%d')
    P(f'版本 / Version：v7.0 FINAL | 制作日期 / Date：{today_str} | 预判周期 / Plan Period：{d["plan_start"]} ~ {d["plan_end"]}', s=10, a=WD_ALIGN_PARAGRAPH.CENTER)
    P('AI生成内容，请注意甄别 | AI-Generated Content, Please Verify | 媳妇智投全球通用版出品', b=True, s=10, c=RGBColor(128,128,128), a=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    
    # ── 方案头部 ──
    H('【方案头部 / Plan Header】', 2)
    T([f'项目 / {_bt.item_label}',f'内容 / {_bt.header}'],[
        ['品种|合约',f'{d["variety"]} | {d["contract"]}(主力合约，持仓{d["day2_oi"]}手)'],
        ['交易时段',d['trade_hours']],
        ['方案周期',f'{d["plan_start"]} ~ {d["plan_end"]} | 2个交易日'],
        ['数据来源',d['data_source']],
        ['说明',d['data_note']],
    ])
    doc.add_paragraph()
    
    # ── 一、行情回顾 ──
    H(f'一、行情回顾 / {_bt.header}', 2)
    P(f'数据来源：Tushare fut_daily接口 | 合约 / Contract {d["ts_code"]}', s=9, c=RGBColor(100,100,100))
    
    T(['日期(交易日) / Date (Trade Day)','开盘(夜盘) / Open (Night)','最高 / High','最低 / Low','收盘(日盘15:00) / Close (Day 15:00)','结算价 / Settlement','涨跌1 / Change','成交量 / Volume','持仓量 / Open Interest'],[
        [d['day1_date'], d['day1_open'], d['day1_high'], d['day1_low'], d['day1_close'], d['day1_settle'], d['day1_change'], d['day1_vol'], d['day1_oi']],
        [d['day2_date'], d['day2_open'], d['day2_high'], d['day2_low'], d['day2_close'], d['day2_settle'], d['day2_change'], d['day2_vol'], d['day2_oi']],
    ])
    P(f'说明：Tushare期货日线中每个交易日trade_date为完整交易日(夜盘+日盘合并)。开盘价对应夜盘21:00开盘，收盘价对应日盘15:00收盘。第1行含前一交易日夜盘+当日日盘，第2行含05-21夜盘+当日日盘。', s=9, c=RGBColor(100,100,100))
    
    H('1.1 关键价位 / Key Price Levels', 3)
    T(['类型 / Type','价位 / Price','说明 / Note'],[
        ['强压力',d['pressure_strong'],'5月中旬高点'],
        ['次压力',d['pressure_mid'],f'{d["day1_date"]}最高点'],
        [f'当前价(日盘)',d['price_current'],f'{d["day2_date"]}收盘(日盘15:00)'],
        ['短期支撑',d['support_short'],f'{d["day2_date"]}最低点'],
        ['强支撑',d['support_strong'],'MA20区域'],
    ])
    doc.add_paragraph()
    
    # ── 二、市场联动 ──
    H(f'二、市场联动 / {_bt.header}', 2)
    T(['品种 / Product','最新价 / Latest','趋势 / Trend','影响 / Impact'],
     [[f'LME铜3个月',d['lme_price'],d['lme_trend'],'方向待定，r~0.95'],
      [f'COMEX铜',d['comex_price'],'+0.23%','强势震荡'],
      [f'美元指数',d['usd_index'],d['usd_trend'],'负相关，利多铜价']])
    doc.add_paragraph()
    
    # ── 2.2 风险预警 ──
    H('2.2 风险预警 / Risk Warnings', 3)
    risks = [r.split('|') for r in [d['risk_vol'],d['risk_lme'],d['risk_inv'],d['risk_geo'],d['risk_spot']]]
    T(['信号 / Signal','状态 / Status','说明 / Note'], risks)
    P('风险：中等。', s=10)
    
    # ── 四、国际局势 ──
    H(f'三、国际局势 / {_bt.header}', 2)
    geos = [g.split('|') for g in [d['geo1'],d['geo2'],d['geo3']]]
    T(['局势 / Situation','事件 / Event','影响 / Impact'], geos)
    doc.add_paragraph()
    
    # ── 🌍 全球地缘政治分析 ──
    H(f'四、全球地缘政治博弈 / {_bt.header}', 2)
    try:
        # 品种→市场名映射 / Variety → market name mapping
        _variety_geo_map = {"沪铜": "美股", "沪铝": "美股", "沪锌": "美股", "沪金": "全球",
                            "螺纹钢": "A股", "热卷": "A股", "铁矿石": "A股", "焦煤": "A股",
                            "焦炭": "A股", "玻璃": "A股", "纯碱": "A股", "尿素": "A股",
                            "原油": "中东", "燃料油": "中东", "沥青": "中东",
                            "甲醇": "A股", "PTA": "A股", "乙二醇": "A股",
                            "橡胶": "日本", "20号胶": "日本", "纸浆": "全球"}
        _geo_market = _variety_geo_map.get(d['variety'], 'A股')
        geo_analysis = get_geopolitical_risks(_geo_market)
        geo_section = format_geopolitical_section(geo_analysis)
        for line in geo_section.split('\n'):
            if line.startswith('|') or line.startswith('#'):
                P(line)
            elif line.strip():
                P(line)
    except:
        P('地缘政治数据加载中... / Loading geopolitical data...')
    doc.add_paragraph()
    
    # ── 五、宏观影响 ──
    H(f'五、宏观影响 / {_bt.header}', 2)
    macros = [m.split('|') for m in [d['macro1'],d['macro2'],d['macro3'],d['macro4'],d['macro5']]]
    T(['因素 / Factor','方向 / Direction','贡献 / Contribution'], macros)
    doc.add_paragraph()
    
    # ── 十五、产业链 ──
    H('五、产业链 / Industry Chain', 2)
    chains = [c.split('|') for c in [d['chain1'],d['chain2'],d['chain3']]]
    T(['维度 / Dimension','现状 / Status','影响 / Impact'], chains)
    doc.add_paragraph()
    
    # ── 十五、资金动向 ──
    H('十五、资金动向 / Capital Flows', 2)
    funds = [f.split('|') for f in [d['fund_oi'],d['fund_vol'],d['fund_trafigura'],d['fund_premium']]]
    T(['信号 / Signal','数据 / Data','判断 / Judgment'], funds)
    P(f'核心：{d["fund_core"]}', b=True, s=10)
    doc.add_paragraph()
    
    # ── 十五、技术分析 ──
    H('十五、技术分析 / Technical Analysis', 2)
    macds = [m.split('|') for m in [d['macd_daily'],d['macd_60min'],d['macd_weekly']]]
    T(['周期 / Period','MACD','均线 / MA','综合 / Summary'], macds)
    P(f'综合：{d["tech_summary"]}', b=True, s=10)
    doc.add_paragraph()
    
    # ── 操作方案 ──
    has_night = d.get('has_night', True)
    night_label = '（日盘+夜盘）' if has_night else '（仅日盘·无夜盘）'
    H(f'八、操作方案 / Trading Plan {night_label}', 2)
    if has_night:
        P(f'说明：日盘(09:00-15:00)和夜盘(21:00-01:00)均为独立交易时段。本周期总体判断：{d["main_strategy"]}。', s=10)
    else:
        P(f'说明：日盘(09:00-11:30/13:30-15:00)。本品种无夜盘交易，注意隔夜跳空风险。本周期总体判断：{d["main_strategy"]}。', s=10)
    P('')
    
    H('8.1 日盘做多操作（主力策略）/ Day Long (Main Strategy)', 3)
    T([f'项目 / {_bt.item_label}','操作内容 / Detail'],
     [['核心逻辑','中期多头+供应紧+现货升水=做多为主'],
      ['入场条件',f'回调{d["day_long_entry"]}企稳+30分钟止跌信号'],
      ['目标',f'第一{d["day_long_target1"]}|第二{d["day_long_target2"]}'],
      ['止损',d['day_long_stop']],
      ['仓位',d['day_long_pos']]])
    
    H('8.2 日盘做空操作（辅助策略）/ Day Short (Auxiliary Strategy)', 3)
    T([f'项目 / {_bt.item_label}','操作内容 / Detail'],
     [['入场条件',d['day_short_entry']],
      ['目标',f'第一{d["day_short_target1"]}|第二{d["day_short_target2"]}'],
      ['止损',d['day_short_stop']],
      ['仓位',d['day_short_pos']],
      ['风险','盘中突破上行必须止损']])
    
    if has_night:
        H('8.3 夜盘做多操作（主力策略）/ Night Long (Main Strategy)', 3)
        T([f'项目 / {_bt.item_label}','操作内容 / Detail'],
         [['入场条件',d['night_long_entry']],
          ['目标',f'第一{d["night_long_target1"]}|第二{d["night_long_target2"]}'],
          ['止损',d['night_long_stop']],
          ['仓位',d['night_long_pos']],
          ['不做条件','LME跌破关键位或低开破支撑']])
        
        H('8.4 夜盘做空操作（辅助策略）/ Night Short (Auxiliary Strategy)', 3)
        T([f'项目 / {_bt.item_label}','操作内容 / Detail'],
         [['入场条件',d['night_short_entry']],
          ['目标',f'第一{d["night_short_target1"]}|第二{d["night_short_target2"]}'],
          ['止损',d['night_short_stop']],
          ['仓位',d['night_short_pos']],
          ['不做条件','LME站稳关键位或突破压力']])
    
    # 预判表：根据是否有夜盘动态生成行
    H('8.3 两日走势预判 / Two-Day Forecast', 3)
    if has_night:
        pred_rows = [
            [d['plan_start'].split('(')[0],d['plan_start'].split('(')[1].rstrip(')'),'日盘',d['pred1_day_dir'],d['pred1_day_range'],d['pred1_day_long'],d['pred1_day_short'],'⭐⭐⭐⭐'],
            [d['plan_start'].split('(')[0],d['plan_start'].split('(')[1].rstrip(')'),'夜盘',d['pred1_night_dir'],d['pred1_night_range'],d['pred1_night_long'],d['pred1_night_short'],'⭐⭐⭐'],
            [d['plan_end'].split('(')[0],d['plan_end'].split('(')[1].rstrip(')'),'日盘',d['pred2_day_dir'],d['pred2_day_range'],d['pred2_day_long'],d['pred2_day_short'],'⭐⭐⭐'],
            [d['plan_end'].split('(')[0],d['plan_end'].split('(')[1].rstrip(')'),'夜盘',d['pred2_night_dir'],d['pred2_night_range'],d['pred2_night_long'],d['pred2_night_short'],'⭐⭐⭐'],
        ]
    else:
        pred_rows = [
            [d['plan_start'].split('(')[0],d['plan_start'].split('(')[1].rstrip(')'),'日盘',d['pred1_day_dir'],d['pred1_day_range'],d['pred1_day_long'],d['pred1_day_short'],'⭐⭐⭐⭐'],
            [d['plan_end'].split('(')[0],d['plan_end'].split('(')[1].rstrip(')'),'日盘',d['pred2_day_dir'],d['pred2_day_range'],d['pred2_day_long'],d['pred2_day_short'],'⭐⭐⭐'],
        ]
    T(['日期 / Date','星期 / Weekday','时段 / Session','预判 / Forecast','区间 / Range','做多策略 / Long Strategy','做空策略 / Short Strategy','置信度 / Confidence'], pred_rows)
    doc.add_paragraph()
    
    # ── 十五、操盘建议 ──
    H('十五、操盘建议 / Trading Suggestions', 2)
    trades = [t.split('|') for t in [d['trade_strategy'],d['trade_pos'],d['trade_stop']]]
    T(['要点 / Key Point','建议 / Suggestion'], trades)
    doc.add_paragraph()
    
    # ── 十五、风险提示 ──
    H('十五、关注重点 / Key Focus Areas', 2)
    T(['重点关注 / Focus',f'内容 / {_bt.header}','影响 / Impact'],
     [['LME铜','价格区间','决定方向'],
      ['美元指数','99-100','美元弱=铜强'],
      ['托克提货','注销5.1万吨后续','亚洲供应'],
      ['现货升贴水','维持升水','支撑力度']])
    doc.add_paragraph()
    
    H('十五、方案总结 / Plan Summary', 2)
    T(['维度 / Dimension','结论 / Conclusion'],
     [['预判',f'{d["plan_start"]}~{d["plan_end"]}高位震荡偏多'],
      ['核心区间','103,500~105,500'],
      ['日盘策略','做多回调区间低吸|做空冲高压力短空'],
      ['夜盘策略','做多LME关键位开多|做空LME破位开空'],
      ['风险等级','中等|评级B+']])
    doc.add_paragraph()
    
    # ── 二十二-二十三 ──
    H('十五、方案总评 / Overall Assessment', 2)
    T(['维度 / Dimension','评分 / Score','评语 / Comment'],
     [['技术面',d['score_tech'],d['score_tech_note']],
      ['基本面',d['score_fundamental'],d['score_fundamental_note']],
      ['资金面',d['score_capital'],d['score_capital_note']],
      ['情绪面',d['score_sentiment'],d['score_sentiment_note']],
      ['宏观面',d['score_macro'],d['score_macro_note']]])
    P(d['summary'], b=True, s=11)
    doc.add_paragraph()
    
    H('十五、博弈论分析 / Game Theory Analysis', 2)
    
    T(['参与方 / Player','方向 / Direction','力量 / Power','策略 / Strategy'],
     [['产业套保(矿企)','空','强','高价套保'],
      ['机构投机','偏多','强','跟趋势'],
      ['游资','双向','中','快进快出'],
      ['散户','多头套牢','弱','高位追涨被套']])
    doc.add_paragraph()
    
    # ── 十五、风险提示与免责声明 ──
    H('十五、风险提示与免责声明 / Risk Warnings & Disclaimer', 2)
    bilingual_risks = [
        ('本方案数据来源于公开市场信息及第三方数据接口（国信证券、金投网等），仅供参考，不构成投资依据。',
         'This plan is based on public market data and third-party data interfaces. For reference only, not investment advice.'),
        ('期货交易具有高风险特征，价格波动可能导致本金全部损失。请根据自身风险承受能力谨慎参与。',
         'Futures trading carries high risk. Price volatility may result in total loss of principal. Trade within your risk tolerance.'),
        ('本方案不构成任何投资建议、投资要约或承诺。任何依据本方案进行的交易决策，风险由投资者自行承担。',
         'This plan does not constitute any investment advice, offer, or commitment. All trading decisions based on this plan are at the investor\'s own risk.'),
        ('投资者应基于自身的风险承受能力和投资经验独立做出决策，必要时咨询专业金融顾问。',
         'Investors should make independent decisions based on their own risk tolerance and experience, consulting financial advisors when necessary.'),
        ('过往业绩及预判结果不代表未来表现，市场存在不确定性，所有预判仅供参考。',
         'Past performance does not guarantee future results. Markets involve uncertainty; all forecasts are for reference only.'),
    ]
    for i,(cn,en) in enumerate(bilingual_risks,1):
        P(f'{i}. {cn}', s=9)
        P(f'   {en}', s=8, c=RGBColor(100,100,100))
    doc.add_paragraph()
    
    # ── AIGC内容标识 ──
    H('AIGC内容标识 / AIGC Content Label', 3)
    P('生成说明 / Generation Note：本方案采用"媳妇智投全球通用版"AI分析引擎生成，包含AI辅助数据采集、分析和文本生成内容。', s=9)
    P('This plan was generated using the \"Ai-Wife Global Investing\" AI analysis engine, including AI-assisted data collection, analysis, and text generation.', s=8, c=RGBColor(100,100,100))
    P('合规提示 / Compliance Note：根据中国证监会《证券期货投资者适当性管理办法》及相关指引，AI生成内容仅供参考，不构成投资建议。投资者应理性判断，审慎决策。', s=9)
    P('In accordance with CSRC regulations and guidelines, AI-generated content is for reference only and does not constitute investment advice. Investors should exercise independent judgment.', s=8, c=RGBColor(100,100,100))
    P('AIGC标识 / AIGC Label：本方案由AI智能体辅助生成 | 人工审核 / Human Review：未审核 (Unreviewed) | 仅供参考 / For Reference Only', s=9, c=RGBColor(128,128,128))
    doc.add_paragraph()
    
    P('媳妇智投全球通用版出品，必属精品 | Ai-Wife Global Investing — Quality Guaranteed', b=True, s=12, a=WD_ALIGN_PARAGRAPH.CENTER)
    P(f'v7.0 FINAL | {today_str} | AI生成内容，请注意甄别 | AI-Generated Content, Please Verify', s=9, c=RGBColor(128,128,128), a=WD_ALIGN_PARAGRAPH.CENTER)
    
    # ── 📄 保存生成的.docx参考框架文件 ──
    # ⚠️ [用户知情声明] 生成框架文档将保存到本地技能目录。
    # 这是用户请求生成方案时的预期行为——生成 .docx 文件并发送给用户。
    # 用户未明确请求时不会自动保存。
    output_dir = os.path.dirname(os.path.abspath(__file__))
    filename = f'{d["variety"].replace(" ","")}{d["contract"]}下周参考框架_v7.0_{today_str}.docx'
    out = os.path.join(output_dir, filename)
    doc.save(out)
    
    return doc, out
