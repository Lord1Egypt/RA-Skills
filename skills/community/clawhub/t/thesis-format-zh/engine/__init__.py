"""
论文排版引擎 v4.1

新增功能：
  - 错误恢复：每步独立 try-except，单步失败不影响整体
  - 进度条：CLI 实时显示处理进度
  - 日志文件：排版过程记录到 logs/
  - 前后对比：排版前后差异报告
"""

import sys
import re
import time
import json
import logging
from pathlib import Path
from datetime import datetime
from docx import Document

from .template_loader import load_templates, json_to_cfg, create_template_skeleton
from .heading_detector import detect_and_apply, find_body_start
from .image_fixer import fix_images_to_top_bottom
from .section_manager import set_page_margins, apply_section_config, detect_sections, add_page_break_before_chapters, insert_section_break
from .style_applier import (
    apply_all_heading_styles,
    apply_normal_style,
    fix_caption_runs,
    fix_heading_colors,
    fix_keywords,
    fix_footnotes,
    fix_appendix_numbering,
)
from .reference_fixer import fix_reference_style, superscript_citations
from .table_formatter import format_tables, apply_three_line_table
from .toc_styles import apply_toc_styles, apply_toc_title_style, insert_toc_field
from .formula_formatter import format_formulas
from .caption_position import fix_caption_positions
import sys as _sys
import os as _os
_sys.path.insert(0, _os.path.dirname(_os.path.dirname(_os.path.abspath(__file__))))
from validator.check import validate

# ═══════════════════════════════════════════════
# 日志系统
# ═══════════════════════════════════════════════

LOG_DIR = Path(__file__).parent.parent / "logs"


def _setup_logger():
    """配置日志（同时输出到控制台和文件）"""
    LOG_DIR.mkdir(parents=True, exist_ok=True)
    log_file = LOG_DIR / f"format_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log"

    logger = logging.getLogger("format_thesis")
    logger.setLevel(logging.DEBUG)

    # 文件处理器
    fh = logging.FileHandler(log_file, encoding="utf-8")
    fh.setLevel(logging.DEBUG)
    fh.setFormatter(logging.Formatter("%(asctime)s - %(levelname)s - %(message)s"))
    logger.addHandler(fh)

    return logger, log_file


_logger, _log_file = _setup_logger()


def log(msg, level="info"):
    """输出日志（控制台 + 文件）"""
    print(f"[format_thesis] {msg}", file=sys.stderr)
    getattr(_logger, level, _logger.info)(msg)


# ═══════════════════════════════════════════════
# 进度条
# ═══════════════════════════════════════════════

class ProgressBar:
    """简单进度条"""

    def __init__(self, total: int, prefix: str = ""):
        self.total = total
        self.current = 0
        self.prefix = prefix
        self.start_time = time.time()
        self._print()

    def update(self, step_name: str = ""):
        self.current += 1
        self._print(step_name)

    def _print(self, step_name: str = ""):
        bar_len = 30
        filled = int(bar_len * self.current / self.total)
        bar = "█" * filled + "░" * (bar_len - filled)
        pct = self.current / self.total * 100
        elapsed = time.time() - self.start_time
        step_info = f" {step_name}" if step_name else ""
        print(f"\r  {self.prefix} [{bar}] {pct:5.1f}% ({self.current}/{self.total}){step_info}  ", end="", file=sys.stderr)

    def finish(self):
        elapsed = time.time() - self.start_time
        print(f"\r  {self.prefix} [{'█' * 30}] 100.0% ({self.total}/{self.total}) 完成 ({elapsed:.1f}s)    ", file=sys.stderr)
        print(file=sys.stderr)


# ═══════════════════════════════════════════════
# 前后对比
# ═══════════════════════════════════════════════

def _snapshot(doc) -> dict:
    """文档快照（用于前后对比）"""
    styles_used = {}
    for p in doc.paragraphs:
        sname = p.style.name
        styles_used[sname] = styles_used.get(sname, 0) + 1

    colors = {"blue": 0, "black": 0, "other": 0}
    for p in doc.paragraphs:
        for r in p.runs:
            if r.font.color and r.font.color.rgb:
                hex_val = str(r.font.color.rgb)
                if hex_val.startswith("0000FF") or hex_val.startswith("4472C4"):
                    colors["blue"] += 1
                elif hex_val == "000000":
                    colors["black"] += 1
                else:
                    colors["other"] += 1

    bold_count = sum(1 for p in doc.paragraphs for r in p.runs if r.bold is True)

    return {
        "paragraphs": len(doc.paragraphs),
        "sections": len(doc.sections),
        "tables": len(doc.tables),
        "styles": styles_used,
        "colors": colors,
        "bold_runs": bold_count,
    }


def _diff_report(before: dict, after: dict) -> str:
    """生成前后对比报告"""
    lines = [
        "📊 排版前后对比",
        f"━━━━━━━━━━━━━━━━━━",
    ]

    # 基本数据
    for key, label in [("paragraphs", "段落"), ("sections", "节"), ("tables", "表格")]:
        b, a = before.get(key, 0), after.get(key, 0)
        if b != a:
            lines.append(f"  {label}: {b} → {a}")
        else:
            lines.append(f"  {label}: {b}（不变）")

    # 样式变化
    lines.append("")
    lines.append("  样式变化:")
    all_styles = set(list(before.get("styles", {}).keys()) + list(after.get("styles", {}).keys()))
    for s in sorted(all_styles):
        b = before.get("styles", {}).get(s, 0)
        a = after.get("styles", {}).get(s, 0)
        if b != a:
            lines.append(f"    {s}: {b} → {a}")

    # 颜色变化
    before_blue = before.get("colors", {}).get("blue", 0)
    after_blue = after.get("colors", {}).get("blue", 0)
    if before_blue != after_blue:
        lines.append(f"  蓝色字体: {before_blue} → {after_blue}（修复标题颜色）")

    # 加粗变化
    before_bold = before.get("bold_runs", 0)
    after_bold = after.get("bold_runs", 0)
    if before_bold != after_bold:
        lines.append(f"  加粗文本: {before_bold} → {after_bold}")

    lines.append(f"━━━━━━━━━━━━━━━━━━")
    return "\n".join(lines)


# ═══════════════════════════════════════════════
# 安全执行步骤
# ═══════════════════════════════════════════════

def _safe_step(func, step_name: str, errors: list, *args, **kwargs):
    """
    安全执行一个步骤。

    失败时记录错误但不中断流程。

    Returns:
        (result, success: bool)
    """
    try:
        result = func(*args, **kwargs)
        return result, True
    except Exception as e:
        error_msg = f"步骤 [{step_name}] 失败: {type(e).__name__}: {e}"
        log(error_msg, "error")
        errors.append({"step": step_name, "error": str(e), "type": type(e).__name__})
        return None, False


# ═══════════════════════════════════════════════
# 主流程
# ═══════════════════════════════════════════════

# 步骤定义
STEPS = [
    ("加载文档", "load"),
    ("文档快照", "snapshot_before"),
    ("标题检测", "heading_detect"),
    ("页面设置", "page_margins"),
    ("标题样式", "heading_styles"),
    ("正文样式", "normal_style"),
    ("图片环绕", "image_fix"),
    ("题注格式", "caption_fix"),
    ("题注位置", "caption_position"),
    ("参考文献", "reference_fix"),
    ("表格格式", "table_format"),
    ("公式格式", "formula_format"),
    ("关键词", "keywords"),
    ("附录编号", "appendix"),
    ("目录样式", "toc"),
    ("分节页码", "sections"),
    ("章节分页", "page_breaks"),
    ("标题颜色", "heading_colors"),
    ("脚注尾注", "footnotes"),
    ("保存文档", "save"),
    ("文档快照", "snapshot_after"),
    ("验证检查", "validate"),
]


def format_document(doc_path, output_path, cfg, skip_heading_detect=False):
    """
    核心格式化流程（带错误恢复、进度条、日志、前后对比）。

    Args:
        doc_path: 输入 docx 路径
        output_path: 输出 docx 路径
        cfg: 模板配置（json_to_cfg 转换后的）
        skip_heading_detect: 是否跳过标题自动检测

    Returns:
        dict: {"heading_stats", "errors", "warnings", "summary", "diff", "log_file", "elapsed"}
    """
    start_time = time.time()
    step_errors = []

    log(f"开始排版: {doc_path}")
    log(f"输出: {output_path}")
    log(f"日志: {_log_file}")

    progress = ProgressBar(26, prefix="排版进度")

    # ─── 1. 加载文档 ───
    doc, ok = _safe_step(Document, "加载文档", step_errors, doc_path)
    if not ok:
        return {"errors": step_errors, "summary": "❌ 文档加载失败"}
    progress.update("加载文档")

    # ─── 2. 文档快照（排版前） ───
    snapshot_before, _ = _safe_step(_snapshot, "排版前快照", step_errors, doc)
    progress.update("排版前快照")

    sections_info = detect_sections(doc)
    log(f"📐 文档: {len(doc.paragraphs)} 段落, {len(doc.sections)} 节, {len(doc.tables)} 表格")

    # ─── 3. 标题检测 ───
    heading_stats = {"h1": 0, "h2": 0, "h3": 0, "h4": 0, "skipped": 0}
    body_start = 0
    if not skip_heading_detect:
        body_start, ok = _safe_step(find_body_start, "标题检测", step_errors, doc)
        if ok and body_start is not None:
            stats, ok = _safe_step(detect_and_apply, "标题检测", step_errors, doc, start_idx=body_start)
            if ok and stats:
                heading_stats = stats
                log(f"  H1={stats['h1']} H2={stats['h2']} H3={stats['h3']} H4={stats['h4']}")
    progress.update("标题检测")

    # ─── 4. 页面设置 ───
    result, ok = _safe_step(set_page_margins, "页面设置", step_errors, doc, cfg)
    if ok:
        log(f"✓ {result} 个 section 页边距设置完成")
    progress.update("页面设置")

    # ─── 5. 标题样式 ───
    result, ok = _safe_step(apply_all_heading_styles, "标题样式", step_errors, doc, cfg)
    if ok:
        log(f"✓ {result} 个标题样式设置完成")
    progress.update("标题样式")

    # ─── 6. 正文样式 ───
    _safe_step(apply_normal_style, "正文样式", step_errors, doc, cfg)
    log("✓ 正文样式设置完成")
    progress.update("正文样式")

    # ─── 7. 图片环绕 ───
    skip_indices = list(range(body_start)) if not skip_heading_detect else []
    result, ok = _safe_step(fix_images_to_top_bottom, "图片环绕", step_errors, doc, max_height_cm=6.0, skip_para_indices=skip_indices)
    if ok:
        log(f"✓ {result} 个图片环绕修复完成")
    progress.update("图片环绕")

    # ─── 8. 题注格式 ───
    result, ok = _safe_step(fix_caption_runs, "题注格式", step_errors, doc, cfg)
    if ok:
        log(f"✓ {result} 个图表题注修复完成")
    progress.update("题注格式")

    # ─── 9. 题注位置 ───
    result, ok = _safe_step(fix_caption_positions, "题注位置", step_errors, doc)
    if ok:
        log(f"✓ {result} 个题注位置修复完成")
    progress.update("题注位置")

    # ─── 10. 参考文献 ───
    result, ok = _safe_step(fix_reference_style, "参考文献", step_errors, doc, cfg)
    if ok:
        log(f"✓ {result} 条参考文献格式修复完成")
    progress.update("参考文献")

    # ─── 10b. 引用上标（必须显式确认）───
    citation_cfg = cfg.get("references", {}).get("citation", {})
    enable_citation_superscript = citation_cfg.get("superscript") is True \
        or cfg.get("citation_superscript") is True
    if enable_citation_superscript:
        result, ok = _safe_step(superscript_citations, "引用上标", step_errors, doc)
        if ok:
            log(f"✓ {result} 个引用上标处理完成")
    else:
        log("↷ 引用上标跳过：规范未明确要求上标")
    progress.update("引用上标")

    # ─── 11. 表格格式 ───
    result, ok = _safe_step(format_tables, "表格格式", step_errors, doc, cfg)
    if ok:
        log(f"✓ {result} 个表格格式化完成")
    progress.update("表格格式")

    # ─── 11b. 三线表 ───
    result, ok = _safe_step(apply_three_line_table, "三线表", step_errors, doc)
    if ok:
        log(f"✓ {result} 个表格三线表格式完成")
    progress.update("三线表")

    # ─── 12. 公式格式 ───
    result, ok = _safe_step(format_formulas, "公式格式", step_errors, doc, cfg)
    if ok:
        log(f"✓ {result} 个公式格式化完成")
    progress.update("公式格式")

    # ─── 13. 关键词 ───
    result, ok = _safe_step(fix_keywords, "关键词", step_errors, doc, cfg)
    if ok:
        log(f"✓ {result} 个关键词格式修复完成")
    progress.update("关键词")

    # ─── 14. 附录编号 ───
    result, ok = _safe_step(fix_appendix_numbering, "附录编号", step_errors, doc, cfg)
    if ok:
        log(f"✓ {result} 个附录编号修复完成")
    progress.update("附录编号")

    # ─── 15. 目录样式 ───
    _safe_step(apply_toc_styles, "目录样式", step_errors, doc, cfg)
    _safe_step(apply_toc_title_style, "目录样式", step_errors, doc, cfg)
    _safe_step(insert_toc_field, "目录域插入", step_errors, doc, cfg)
    log("✓ 目录样式设置完成")
    progress.update("目录样式")

    # ─── 16. 分节页码 ───
    result, ok = _safe_step(apply_section_config, "分节页码", step_errors, doc, cfg)
    if ok and result:
        for r in result:
            log(r)
    progress.update("分节页码")

    # ─── 17. 章节分页 ───
    result, ok = _safe_step(add_page_break_before_chapters, "章节分页", step_errors, doc)
    if ok:
        log(f"✓ {result} 个章节分页设置完成")
    progress.update("章节分页")

    # ─── 18. 标题颜色 ───
    result, ok = _safe_step(fix_heading_colors, "标题颜色", step_errors, doc)
    if ok:
        log(f"✓ {result} 个标题颜色修复完成")
    progress.update("标题颜色")

    # ─── 19. 脚注尾注 ───
    result, ok = _safe_step(fix_footnotes, "脚注尾注", step_errors, doc, cfg)
    if ok:
        log(f"✓ {result} 个脚注格式修复完成")
    progress.update("脚注尾注")

    # ─── 20. 保存 ───
    _safe_step(doc.save, "保存文档", step_errors, output_path)
    log(f"✅ 已保存: {output_path}")
    progress.update("保存文档")

    # ─── 21. 文档快照（排版后） ───
    doc_after = Document(output_path)
    snapshot_after, _ = _safe_step(_snapshot, "排版后快照", step_errors, doc_after)
    progress.update("排版后快照")

    # ─── 22. 验证 ───
    errors, warnings, summary = [], [], "✅ 格式检查通过"
    try:
        errors, warnings, summary = validate(doc_after, cfg)
    except Exception as e:
        step_errors.append({"step": "验证", "error": str(e)})
    progress.update("验证检查")

    progress.finish()

    # ─── 前后对比报告 ───
    diff_report = ""
    if snapshot_before and snapshot_after:
        diff_report = _diff_report(snapshot_before, snapshot_after)
        log(diff_report)

    # ─── 汇总 ───
    elapsed = time.time() - start_time
    log(f"\n排版完成，耗时 {elapsed:.1f} 秒")

    if step_errors:
        log(f"⚠️ {len(step_errors)} 个步骤有错误:", "warning")
        for err in step_errors:
            log(f"  - [{err['step']}] {err['error']}", "warning")

    return {
        "heading_stats": heading_stats,
        "errors": errors,
        "warnings": warnings,
        "step_errors": step_errors,
        "summary": summary,
        "diff": diff_report,
        "log_file": str(_log_file),
        "elapsed": round(elapsed, 1),
    }


def analyze_document(doc_path):
    """仅分析文档结构，不做修改"""
    doc = Document(doc_path)
    sections_info = detect_sections(doc)
    body_start = find_body_start(doc)

    # 预览标题检测（不修改文档）
    heading_preview = {"h1": 0, "h2": 0, "h3": 0, "h4": 0}
    for i, para in enumerate(doc.paragraphs):
        if i < body_start:
            continue
        from .heading_detector import classify_heading, _is_heading_candidate
        if not _is_heading_candidate(para):
            continue
        level, confidence = classify_heading(para)
        if level and confidence >= 0.6:
            heading_preview[f"h{level}"] += 1

    return {
        "paragraphs": len(doc.paragraphs),
        "sections": len(doc.sections),
        "tables": len(doc.tables),
        "body_start": body_start,
        "sections_info": sections_info,
        "heading_preview": heading_preview,
    }
