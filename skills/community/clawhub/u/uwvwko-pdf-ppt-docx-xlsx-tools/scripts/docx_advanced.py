"""DOCX 高级操作: 合并 / 提取文本 / 信息

用法:
    合并:    python docx_advanced.py merge file1.docx file2.docx ... [output.docx]
    提取文本: python docx_advanced.py extract_text input.docx [output.txt]
    信息:    python docx_advanced.py info input.docx
"""
from docx import Document
import sys
import os


def merge(files, out_path):
    if len(files) < 2:
        print("错误: 合并至少需要 2 个 DOCX 文件")
        sys.exit(1)
    merged = Document(files[0])
    for f in files[1:]:
        if not os.path.isfile(f):
            print(f"错误: 文件不存在 - {f}")
            sys.exit(1)
        src = Document(f)
        for para in src.paragraphs:
            merged.add_paragraph(para.text, style=para.style)
        for table in src.tables:
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            if rows:
                t = merged.add_table(rows=len(rows), cols=len(rows[0]))
                for r_idx, row_data in enumerate(rows):
                    for c_idx, cell_text in enumerate(row_data):
                        t.cell(r_idx, c_idx).text = cell_text
        merged.add_page_break()
        print(f"Merged: {f}")
    merged.save(out_path)
    print(f"完成: {len(files)} 个文件合并 -> {out_path}")


def extract_text(docx_path, out_path):
    doc = Document(docx_path)
    with open(out_path, "w", encoding="utf-8") as f:
        for para in doc.paragraphs:
            if para.text.strip():
                f.write(para.text + "\n")
        for i, table in enumerate(doc.tables):
            f.write(f"\n--- Table {i + 1} ---\n")
            for row in table.rows:
                f.write(" | ".join(cell.text for cell in row.cells) + "\n")
    print(f"完成: 文本提取 -> {out_path}")


def info(docx_path):
    doc = Document(docx_path)
    core = doc.core_properties
    print(f"文件: {docx_path}")
    print(f"段落数: {len(doc.paragraphs)}")
    print(f"表格数: {len(doc.tables)}")
    print(f"标题: {core.title or 'N/A'}")
    print(f"作者: {core.author or 'N/A'}")
    print(f"主题: {core.subject or 'N/A'}")
    print(f"创建时间: {core.created or 'N/A'}")
    print(f"修改时间: {core.modified or 'N/A'}")


def main():
    if len(sys.argv) < 3:
        print(__doc__)
        sys.exit(1)

    action = sys.argv[1].lower()

    if action == "merge":
        out = "merged.docx"
        files = []
        for arg in sys.argv[2:]:
            if arg.endswith(".docx") and os.path.isfile(arg):
                files.append(arg)
            else:
                out = arg
        if not files:
            print("错误: 未找到有效的 DOCX 文件")
            sys.exit(1)
        merge(files, out)

    elif action == "extract_text":
        docx_path = sys.argv[2]
        out_path = sys.argv[3] if len(sys.argv) > 3 else "output.txt"
        extract_text(docx_path, out_path)

    elif action == "info":
        info(sys.argv[2])

    else:
        print(f"未知操作: {action}")
        print(__doc__)
        sys.exit(1)


if __name__ == "__main__":
    main()