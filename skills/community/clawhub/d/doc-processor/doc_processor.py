#!/usr/bin/env python3
"""
文档处理 Python Skill - 统一处理 Word/Excel/PDF
支持格式:.docx, .xlsx, .pdf, .csv, .txt, .md
"""

import os
import sys
import json
import subprocess
import shutil
from pathlib import Path
from typing import Optional, Dict, Any, List, Union, Callable
from dataclasses import dataclass, asdict, field
from datetime import datetime

# ========== 可选依赖 ==========
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    WORD_AVAILABLE = True
except ImportError:
    WORD_AVAILABLE = False
    Document = None

try:
    import pandas as pd
    from openpyxl import load_workbook
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    pd = None
    load_workbook = None


# ========== v2.7.0 新增模块导入 ==========
import importlib

def _import_module(name):
    """灵活导入模块"""
    try:
        return importlib.import_module(name)
    except ImportError:
        return None

# 导入所有新模块
# v2.7.11: 移除 AI 功能,由 OpenClaw 主程序统一处理
# AI 功能已通过架构重构移除,本 Skill 专注文档处理

logging_config = _import_module('logging_config')
error_handler = _import_module('error_handler')
performance = _import_module('performance')
# ai_service = _import_module('ai_service')  # 已移除 v2.7.11
batch_processor = _import_module('batch_processor')
template_cache = _import_module('template_cache')
excel_style = _import_module('excel_style')
pdf_generator = _import_module('pdf_generator')

# 设置默认值
if logging_config:
    setup_logger = logging_config.setup_logger
    configure_logging = logging_config.configure_logging
else:
    def setup_logger(*args, **kwargs):
        import logging
        return logging.getLogger('doc_processor')
    def configure_logging(*args, **kwargs):
        pass

if error_handler:
    DocProcessorError = error_handler.DocProcessorError
    file_not_found = error_handler.file_not_found
    dependency_missing = error_handler.dependency_missing
else:
    class DocProcessorError(Exception):
        pass
    def file_not_found(*args, **kwargs):
        raise FileNotFoundError(*args)
    def dependency_missing(*args, **kwargs):
        raise ImportError(*args)

if performance:
    PerformanceMonitor = performance.PerformanceMonitor
    get_monitor = performance.get_monitor
else:
    class PerformanceMonitor:
        pass
    def get_monitor():
        return None

if batch_processor:
    BatchResult = batch_processor.BatchResult
    ProgressTracker = batch_processor.ProgressTracker
    SimpleProgress = batch_processor.SimpleProgress
else:
    class BatchResult:
        pass
    class ProgressTracker:
        pass
    class SimpleProgress:
        pass

if template_cache:
    TemplateCache = template_cache.TemplateCache
else:
    class TemplateCache:
        def __init__(self, *args, **kwargs):
            pass
        def stats(self):
            return {}
        def clear(self):
            pass
        def get(self, *args):
            return None

if excel_style:
    StylePreserver = excel_style.StylePreserver
else:
    class StylePreserver:
        pass

if pdf_generator:
    PDFGenerator = pdf_generator.PDFGenerator
    generate_pdf = pdf_generator.generate_pdf
else:
    class PDFGenerator:
        pass
    def generate_pdf(*args, **kwargs):
        pass

# 初始化 logger
logger = setup_logger()


# ========== 数据结构 ==========
@dataclass
class DocInfo:
    """文档元信息"""
    path: str
    type: str  # pdf, word, excel, csv, txt, md
    size_bytes: int
    filename: str = ""
    size_human: str = ""
    pages: Optional[int] = None
    sheets: Optional[List[str]] = None
    rows: Optional[int] = None
    columns: Optional[int] = None
    created_at: Optional[str] = None
    modified_at: Optional[str] = None
    author: Optional[str] = None

    def __post_init__(self):
        if not self.size_human:
            self.size_human = self._format_size(self.size_bytes)

    @staticmethod
    def _format_size(bytes_: int) -> str:
        for unit in ['B', 'KB', 'MB', 'GB']:
            if bytes_ < 1024:
                return f"{bytes_:.1f} {unit}"
            bytes_ /= 1024
        return f"{bytes_:.1f} TB"

    def to_dict(self):
        return asdict(self)


@dataclass
class DocContent:
    """文档内容"""
    type: str
    format: str  # text, structured, table
    data: Any
    metadata: Dict = field(default_factory=dict)

    def to_dict(self):
        return {
            'type': self.type,
            'format': self.format,
            'data': self.data,
            'metadata': self.metadata
        }


# ========== 异常类 ==========
class DependencyError(Exception):
    """依赖缺失错误"""
    pass


class FormatError(Exception):
    """格式不支持错误"""
    pass


# ========== v2.0 新增数据结构 ==========
@dataclass
class StyleDefinition:
    """样式定义 - v2.0 新增"""
    page: Dict = field(default_factory=dict)
    sections: List[Dict] = field(default_factory=list)
    styles: Dict[str, Dict] = field(default_factory=dict)
    constraints: Dict[str, Any] = field(default_factory=dict)


@dataclass
class TemplateInfo:
    """模板信息 - v2.0 新增"""
    path: str
    style: StyleDefinition
    content_type: str = ""
    usage_count: int = 0
    last_used: Optional[str] = None


@dataclass
class GenerationResult:
    """生成结果 - v2.0 新增"""
    output_path: str
    template_used: Optional[TemplateInfo] = None
    style_applied: Optional[StyleDefinition] = None
    stats: Dict = field(default_factory=dict)


# ========== v2.0 新增:样式提取器 ==========
class StyleExtractor:
    """样式提取器 - v2.0 新增"""

    def __init__(self):
        pass

    def extract(self, template_path: Union[str, Path]) -> StyleDefinition:
        """从模板提取样式定义"""
        path = Path(template_path)
        ext = path.suffix.lower()

        if ext == '.docx':
            return self._extract_word_style(path)
        elif ext == '.xlsx':
            return self._extract_excel_style(path)
        else:
            return StyleDefinition()

    def _extract_word_style(self, path: Path) -> StyleDefinition:
        """提取 Word 样式 - v2.1 增强版"""
        if not WORD_AVAILABLE:
            return StyleDefinition()

        from docx import Document
        doc = Document(path)

        sections = []
        styles = {}

        # 提取章节结构
        current_section = None
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                if current_section:
                    sections.append(current_section)

                # v2.1: 提取完整样式信息
                current_section = {
                    'title': para.text,
                    'level': int(para.style.name[-1]) if para.style.name[-1].isdigit() else 1,
                    'style': {
                        'font': para.style.font.name if para.style.font.name else '宋体',
                        'size': para.style.font.size.pt if para.style.font.size else 12,
                        'bold': para.style.font.bold,
                        'color': self._get_color_hex(para.style.font.color),
                    },
                    'format': self._extract_paragraph_format(para),
                    'content_type': self._infer_content_type(para.text),
                    'children': []
                }
            elif current_section and para.text.strip():
                # v2.1: 提取段落格式信息
                current_section['children'].append({
                    'type': 'paragraph',
                    'style': {
                        'font': para.style.font.name if para.style.font.name else '宋体',
                        'size': para.style.font.size.pt if para.style.font.size else 12,
                    },
                    'format': self._extract_paragraph_format(para),
                    'max_length': len(para.text) * 2
                })

        if current_section:
            sections.append(current_section)

        return StyleDefinition(
            page={'size': 'A4', 'orientation': 'portrait'},
            sections=sections,
            styles=styles,
            constraints={}
        )

    def _get_color_hex(self, color) -> str:
        """获取颜色 HEX 值 - v2.1 新增"""
        if color is None:
            return '#000000'
        try:
            if hasattr(color, 'rgb') and color.rgb is not None:
                rgb = color.rgb
                # RGBColor 对象需要转换为整数
                if hasattr(rgb, '__int__'):
                    return f'#{int(rgb):06x}'
                else:
                    return '#000000'
            else:
                return '#000000'
        except Exception:
            return '#000000'

    def _get_pt_value(self, value) -> float:
        """获取磅值 - v2.1 新增"""
        if value is None:
            return 0.0
        try:
            return value.pt if hasattr(value, 'pt') else float(value)
        except (AttributeError, ValueError, TypeError):
            return 0.0

    def _extract_paragraph_format(self, para) -> Dict:
        """提取段落格式 - v2.1 新增"""
        """注意:python-docx 中段落格式通过 paragraph.paragraph_format 访问"""
        try:
            fmt = para.paragraph_format if hasattr(para, 'paragraph_format') else None
            if fmt is None:
                return {
                    'space_after': 0.0,
                    'line_spacing': 1.5,
                    'left_indent': 0.0,
                    'first_line_indent': 0.0
                }

            return {
                'space_after': self._get_pt_value(fmt.space_after),
                'line_spacing': fmt.line_spacing if fmt.line_spacing else 1.5,
                'left_indent': self._get_pt_value(fmt.left_indent),
                'first_line_indent': self._get_pt_value(fmt.first_line_indent)
            }
        except Exception as e:
            return {
                'space_after': 0.0,
                'line_spacing': 1.5,
                'left_indent': 0.0,
                'first_line_indent': 0.0
            }

    def _extract_excel_style(self, path: Path) -> StyleDefinition:
        """提取 Excel 样式"""
        if not EXCEL_AVAILABLE:
            return StyleDefinition()

        xl = load_workbook(path, read_only=True)

        return StyleDefinition(
            page={'size': 'A4'},
            sections=[{'title': sheet, 'level': 1} for sheet in xl.sheetnames],
            styles={},
            constraints={}
        )

    def _infer_content_type(self, title: str) -> str:
        """从标题推断内容类型"""
        title_lower = title.lower()

        if any(k in title_lower for k in ['摘要', '概述', '总结', '重点', 'main', 'summary']):
            return 'summary'
        elif any(k in title_lower for k in ['列表', '任务', '项目', '完成', 'list', 'task']):
            return 'list'
        elif any(k in title_lower for k in ['表格', '数据', '统计', 'table', 'data']):
            return 'table'
        elif any(k in title_lower for k in ['问题', '困难', '挑战', 'problem', 'issue']):
            return 'problem'
        elif any(k in title_lower for k in ['计划', '安排', '下一步', 'plan', 'next']):
            return 'plan'
        else:
            return 'text'


# ========== v2.0 新增:内容适配器 ==========
class ContentAdapter:
    """内容适配器 - v2.7.11 简化版 (移除 AI 功能)"""

    def __init__(self):
        """
        初始化内容适配器

        v2.7.11: 移除 AI 功能,由 OpenClaw 主程序统一处理
        """
        self._summary_cache = {}

    def _adapt_summary(self, content: str, max_length: int = 150) -> str:
        """
        适配为摘要 - v2.7.11 简化版

        v2.7.11: 移除 AI 服务,仅保留简单摘要
        如需 AI 摘要,请使用 OpenClaw 主程序的 LLM 能力
        """
        if not isinstance(content, str):
            content = str(content)

        if len(content) <= max_length:
            return content

        # 简单摘要(按句子分割)
        sentences = content.replace('。', '。\n').split('\n')
        summary = []
        current_length = 0
        for sentence in sentences:
            if current_length + len(sentence) + 1 <= max_length:
                summary.append(sentence)
                current_length += len(sentence) + 1
            else:
                break
        return '。'.join(summary) + '。' if summary else content[:max_length] + '...'

    def adapt_all(self, content: Union[Dict, str], style_def: StyleDefinition) -> Dict:
        """适配所有章节内容"""
        if isinstance(content, str):
            content = self._parse_content(content)

        adapted = {}
        for section in style_def.sections:
            section_title = section.get('title', '')
            raw = content.get(section_title, '')
            adapted[section_title] = self._adapt_section(raw, section)

        return adapted

    def _adapt_section(self, content: Any, section: Dict) -> Any:
        """适配单个章节"""
        content_type = section.get('content_type', 'text')

        if content_type == 'summary':
            return self._adapt_summary(content)
        elif content_type == 'list':
            return self._adapt_list(content)
        elif content_type == 'plan':
            return self._adapt_plan(content)
        else:
            return self._adapt_text(content)

    def _adapt_summary(self, content: str) -> str:
        """适配为摘要 - v2.1 AI 增强版"""
        if not isinstance(content, str):
            content = str(content)

        max_length = 150

        if len(content) <= max_length:
            return content

        # AI 总结(v2.1 新增)
        try:
            prompt = f"""请将以下内容总结为{max_length}字以内的摘要,保留核心要点:

{content[:500]}  # 只取前 500 字避免超长

要求:
1. 保留关键信息(项目名、成果、数据)
2. 语言简洁专业
3. 不超过{max_length}字

摘要:"""

            # 调用 AI 总结(简化实现,实际应调用 LLM)
            # 当前使用简单启发式总结
            sentences = content.split('。')
            summary = []
            current_length = 0

            for sentence in sentences:
                if current_length + len(sentence) + 1 <= max_length:
                    summary.append(sentence)
                    current_length += len(sentence) + 1
                else:
                    break

            if summary:
                return '。'.join(summary) + '。'
            else:
                # 如果第一句就超长,直接截断
                return content[:max_length] + '...'

        except Exception as e:
            # 降级方案:简单截断
            return content[:max_length] + '...'

    def _adapt_list(self, content: Any, min_items: int = 3, max_items: int = 8) -> List[str]:
        """适配为列表 - v2.1 AI 增强版"""
        # 解析为列表项
        if isinstance(content, str):
            items = [line.strip() for line in content.split('\n') if line.strip()]
        elif isinstance(content, list):
            items = content
        else:
            items = [str(content)]

        # 格式统一
        formatted = []
        for item in items:
            if not item.startswith('-') and not item.startswith('•'):
                item = f"• {item}"
            formatted.append(item)

        # AI 补充不足项(v2.1 新增)
        if len(formatted) < min_items:
            try:
                # 分析现有项的主题
                existing_topics = [item.replace('• ', '').split(' - ')[0] for item in formatted]

                # 生成补充项的提示
                prompt = f"""已有列表项:
{chr(10).join(formatted)}

请补充到{min_items}项,保持格式一致:
- 使用动词开头(完成/开发/修复/实现)
- 包含成果描述
- 如果适用,标注状态(已完成/进行中)
- 与现有项主题相关

只需返回补充的项,每行一个,不要编号。"""

                # 简化实现:生成通用补充项
                # v2.7.11: AI 功能已移除
                additional_templates = [
                    "• 参与团队讨论和代码审查",
                    "• 学习相关技术文档",
                    "• 协助解决团队技术问题"
                ]

                needed = min_items - len(formatted)
                for i in range(min(needed, len(additional_templates))):
                    formatted.append(additional_templates[i])

            except Exception:
                pass  # 不补充,保持原样

        return formatted[:max_items]

    def _adapt_plan(self, content: Any) -> List[str]:
        """适配为计划 - v2.1 AI 增强版"""
        # 先作为列表解析
        items = self._adapt_list(content, min_items=2, max_items=5)

        # AI 添加时间估算(v2.1 新增)
        enhanced_items = []
        for item in items:
            # 检查是否已有时间估算
            if '预计' in item or '天' in item or '小时' in item or '周' in item:
                enhanced_items.append(item)
            else:
                try:
                    # 简化实现:根据任务类型估算
                    # v2.7.11: AI 功能已移除
                    time_estimate = self._estimate_time(item)
                    if time_estimate:
                        enhanced_items.append(f"{item} - 预计{time_estimate}")
                    else:
                        enhanced_items.append(item)
                except Exception:
                    enhanced_items.append(item)

        return enhanced_items[:5]

    def _estimate_time(self, task: str) -> str:
        """根据任务描述估算时间(v2.1 简化版)"""
        task_lower = task.lower()

        # 关键词匹配估算
        if any(k in task_lower for k in ['修复', 'bug', '问题']):
            return '1-2 天'
        elif any(k in task_lower for k in ['开发', '实现', '功能']):
            return '3-5 天'
        elif any(k in task_lower for k in ['优化', '改进', '性能']):
            return '2-3 天'
        elif any(k in task_lower for k in ['文档', '说明', 'readme']):
            return '1 天'
        elif any(k in task_lower for k in ['测试', 'unittest', 'coverage']):
            return '1-2 天'
        elif any(k in task_lower for k in ['学习', '研究', '调研']):
            return '2-4 天'
        else:
            return '2-3 天'  # 默认估算

    def _adapt_text(self, content: Any) -> str:
        """适配为文本"""
        return str(content)

    def _parse_content(self, text: str) -> Dict:
        """解析自然语言内容"""
        # 简单解析(v2.0 MVP)
        return {'content': text}


# ========== 文档处理器 ==========
class DocumentProcessor:
    """统一文档处理器 - v2.0 增强版"""

    SUPPORTED_FORMATS = {
        'read': ['.pdf', '.docx', '.xlsx', '.csv', '.txt', '.md'],
        'write': ['.docx', '.xlsx', '.csv', '.txt', '.md'],
        'convert': True,
        'merge': ['.docx', '.xlsx'],
    }

    def __init__(
        self,
        workspace: str = None,
        enable_cache: bool = True,
        enable_perf_monitor: bool = True
    ):
        """
        初始化文档处理器 - v2.7.11 简化版 (移除 AI 功能)

        Args:
            workspace: 工作目录
            enable_cache: 是否启用模板缓存
            enable_perf_monitor: 是否启用性能监控

        v2.7.11 变更:
            - 移除 ai_service_type 参数
            - AI 功能由 OpenClaw 主程序统一处理
            - 本 Skill 专注文档处理
        """
        self.workspace = Path(workspace) if workspace else Path.cwd()
        self.capabilities = self._check_capabilities()

        # v2.0 功能
        self.style_extractor = StyleExtractor()
        self.content_adapter = ContentAdapter()
        self.user_template_dir = self.workspace.parent / "user-templates"
        self._ensure_user_template_dir()

        # v2.5.0 功能:模板缓存
        self.enable_cache = enable_cache
        if enable_cache:
            self.template_cache = TemplateCache(max_size=100, max_memory_mb=512)
        else:
            self.template_cache = None

        # v2.7.0 功能:性能监控
        self.enable_perf_monitor = enable_perf_monitor
        self.perf_monitor = get_monitor() if enable_perf_monitor else None

        # v2.6.0 功能:Excel 样式保持
        self.style_preserver = StylePreserver()

        # v2.6.0 功能:PDF 生成
        self.pdf_generator = None  # 按需创建

        logger.info(f"DocumentProcessor 已初始化,工作目录:{self.workspace}")

    def _ensure_user_template_dir(self):
        """确保用户模板目录存在"""
        self.user_template_dir.mkdir(parents=True, exist_ok=True)
        (self.user_template_dir / "templates").mkdir(exist_ok=True)

    def _check_capabilities(self) -> Dict[str, bool]:
        """检查可用功能"""
        return {
            'pdf': self._check_pdf_tools(),
            'word': WORD_AVAILABLE,
            'excel': EXCEL_AVAILABLE,
        }

    def _check_pdf_tools(self) -> bool:
        """检查 PDF 系统工具"""
        return (
            shutil.which('pdftotext') is not None and
            shutil.which('pdfinfo') is not None
        )

    def get_capabilities(self) -> Dict:
        """返回能力报告"""
        return {
            'capabilities': self.capabilities,
            'supported_formats': self.SUPPORTED_FORMATS,
            'warnings': self._get_warnings()
        }

    def _get_warnings(self) -> List[str]:
        """获取警告信息"""
        warnings = []
        if not self.capabilities['word']:
            warnings.append("Word 支持未启用:pip install python-docx")
        if not self.capabilities['excel']:
            warnings.append("Excel 支持未启用:pip install openpyxl pandas")
        if not self.capabilities['pdf']:
            warnings.append("PDF 支持未启用:sudo apt install poppler-utils")
        return warnings

    # ========== 公共接口 ==========
    def get_info(self, file_path: str) -> DocInfo:
        """获取文档信息"""
        path = self._resolve_path(file_path)
        self._validate_file(path)

        ext = path.suffix.lower()
        stat = path.stat()

        doc_type = self._get_doc_type(ext)

        info = DocInfo(
            path=str(path.absolute()),
            filename=path.name,
            type=doc_type,
            size_bytes=stat.st_size,
            modified_at=datetime.fromtimestamp(stat.st_mtime).isoformat()
        )

        if doc_type == 'pdf':
            self._enrich_pdf_info(path, info)
        elif doc_type == 'word':
            self._enrich_word_info(path, info)
        elif doc_type == 'excel':
            self._enrich_excel_info(path, info)
        elif doc_type == 'csv':
            self._enrich_csv_info(path, info)

        return info

    def read(self, file_path: str, **kwargs) -> DocContent:
        """读取文档内容"""
        path = self._resolve_path(file_path)
        self._validate_file(path)

        ext = path.suffix.lower()
        doc_type = self._get_doc_type(ext)

        if doc_type == 'pdf':
            return self._read_pdf(path, **kwargs)
        elif doc_type == 'word':
            return self._read_word(path, **kwargs)
        elif doc_type == 'excel':
            return self._read_excel(path, **kwargs)
        elif doc_type == 'csv':
            return self._read_csv(path, **kwargs)
        elif doc_type in ['txt', 'md']:
            return self._read_txt(path, **kwargs)
        else:
            raise FormatError(f"不支持读取的格式:{ext}")

    def write(self, file_path: str, content: Any, **kwargs) -> str:
        """写入文档"""
        path = self._resolve_path(file_path)
        ext = path.suffix.lower()
        doc_type = self._get_doc_type(ext)

        path.parent.mkdir(parents=True, exist_ok=True)

        if doc_type == 'word':
            return self._write_word(path, content, **kwargs)
        elif doc_type == 'excel':
            return self._write_excel(path, content, **kwargs)
        elif doc_type == 'csv':
            return self._write_csv(path, content, **kwargs)
        elif doc_type in ['txt', 'md']:
            return self._write_txt(path, content, **kwargs)
        else:
            raise FormatError(f"不支持写入的格式:{ext}")

    def convert(self, src_path: str, dst_path: str) -> str:
        """转换文档格式 - v2.7.5 修复路径解析 bug"""
        content = self.read(src_path)
        # v2.7.5 修复:解析目标路径到 workspace
        dst_path = self._resolve_path(dst_path)
        dst_ext = dst_path.suffix.lower()
        src_ext = Path(src_path).suffix.lower()

        # v2.7.5 新增:确保目标目录存在
        dst_path.parent.mkdir(parents=True, exist_ok=True)

        # v2.7.0 新增:Word/Excel → TXT
        if dst_ext == '.txt':
            if content.type == 'word':
                text = '\n'.join([p['text'] for p in content.data.get('paragraphs', [])])
                return self._write_txt(dst_path, text)
            elif content.type == 'excel':
                # Excel → TXT (第一个 sheet)
                first_sheet = list(content.data.values())[0]
                text = '\n'.join([','.join(map(str, row)) for row in first_sheet.get('rows', [])])
                return self._write_txt(dst_path, text)

        # 特殊处理:Word/Excel → CSV 时提取表格数据
        if dst_ext == '.csv':
            if content.type == 'word':
                tables = content.data.get('tables', [])
                if not tables:
                    raise FormatError("Word 文件中没有表格,无法转换为 CSV")
                # 提取第一个表格的 data 字段 (二维列表)
                table_data = tables[0].get('data', [])
                if not table_data:
                    raise FormatError("Word 表格为空")
                return self._write_csv(dst_path, table_data)
            elif content.type == 'excel':
                # 使用第一个 sheet 的 rows 数据
                first_sheet = list(content.data.values())[0]
                # Excel 读取后结构:{'columns': [...], 'rows': [...], 'count': N}
                csv_data = first_sheet.get('rows', [])
                if not csv_data:
                    raise FormatError("Excel 文件为空")
                return self._write_csv(dst_path, csv_data)

        # 特殊处理:CSV/TXT → Excel 时需要正确转换数据结构
        if dst_ext == '.xlsx' and src_ext in ['.csv', '.txt']:
            # CSV 读取后结构:{'columns': [...], 'rows': [...], 'count': N}
            if isinstance(content.data, dict) and 'columns' in content.data:
                columns = content.data['columns']
                rows = content.data['rows']
                # 重建二维列表(包含表头)
                excel_data = [columns] + rows
                return self._write_excel(dst_path, excel_data, has_header=True)
            elif isinstance(content.data, list):
                # 已经是二维列表
                return self._write_excel(dst_path, content.data, has_header=True)

        return self.write(dst_path, content.data)

    def merge(self, file_paths: List[str], output_path: str) -> str:
        """合并多个文档"""
        if not file_paths:
            raise ValueError("文件列表为空")

        output = self._resolve_path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)

        ext = output.suffix.lower()

        if ext == '.docx':
            return self._merge_word(file_paths, output)
        elif ext == '.xlsx':
            return self._merge_excel(file_paths, output)
        else:
            raise FormatError(f"不支持合并的格式:{ext}")

    def extract(self, file_path: str, **kwargs) -> Dict:
        """提取文档数据"""
        content = self.read(file_path)
        result = content.to_dict()

        if 'sheet' in kwargs:
            if content.type == 'excel':
                result['data'] = content.data.get(kwargs['sheet'], {})

        return result

    # ========== 路径处理 ==========
    def _resolve_path(self, file_path: str) -> Path:
        """解析路径"""
        path = Path(file_path)
        if not path.is_absolute():
            path = self.workspace / path
        return path.resolve()

    def _validate_file(self, path: Path):
        """验证文件"""
        if not path.exists():
            raise FileNotFoundError(f"文件不存在:{path}")
        if not path.is_file():
            raise ValueError(f"不是文件:{path}")

    def _get_doc_type(self, ext: str) -> str:
        """获取文档类型"""
        type_map = {
            '.pdf': 'pdf',
            '.docx': 'word',
            '.doc': 'word_legacy',
            '.xlsx': 'excel',
            '.xls': 'excel_legacy',
            '.csv': 'csv',
            '.txt': 'txt',
            '.md': 'md'
        }
        return type_map.get(ext, 'unknown')

    # ========== PDF 操作 ==========
    def _read_pdf(self, path: Path, layout: bool = False,
                  pages: str = None) -> DocContent:
        """读取 PDF"""
        if not self.capabilities['pdf']:
            raise DependencyError("PDF 工具未安装:sudo apt install poppler-utils")

        cmd = ['pdftotext']
        if layout:
            cmd.append('-layout')
        if pages:
            cmd.extend(['-f', pages.split('-')[0]])
            if '-' in pages:
                cmd.extend(['-l', pages.split('-')[1]])
        cmd.extend([str(path), '-'])

        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
        if result.returncode != 0:
            raise Exception(f"pdftotext 失败:{result.stderr}")

        return DocContent(
            type='pdf',
            format='text',
            data=result.stdout,
            metadata={'pages': self._get_pdf_pages(path)}
        )

    def _get_pdf_pages(self, path: Path) -> int:
        """获取 PDF 页数"""
        try:
            result = subprocess.run(
                ['pdfinfo', str(path)],
                capture_output=True,
                text=True,
                timeout=10
            )
            for line in result.stdout.split('\n'):
                if line.startswith('Pages:'):
                    return int(line.split(':')[1].strip())
        except Exception as e:
            print(f"获取页数失败:{e}", file=sys.stderr)
        return 0

    def _enrich_pdf_info(self, path: Path, info: DocInfo):
        """丰富 PDF 信息"""
        info.pages = self._get_pdf_pages(path)

    # ========== Word 操作 ==========
    def _read_word(self, path: Path, include_tables: bool = True) -> DocContent:
        """读取 Word"""
        if not self.capabilities['word']:
            raise DependencyError("python-docx 未安装:pip install python-docx")

        doc = Document(path)
        paragraphs = []
        tables = []

        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append({
                    'text': para.text,
                    'style': para.style.name if para.style else None
                })

        if include_tables:
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append({
                    'index': i,
                    'rows': len(table.rows),
                    'cols': len(table.columns),
                    'data': table_data
                })

        return DocContent(
            type='word',
            format='structured',
            data={
                'paragraphs': paragraphs,
                'tables': tables
            },
            metadata={
                'author': getattr(doc.core_properties, 'author', None),
                'created': str(getattr(doc.core_properties, 'created', None)),
                'modified': str(getattr(doc.core_properties, 'modified', None))
            }
        )

    def _write_word(self, path: Path, content: Union[str, Dict],
                    template: str = None) -> str:
        """写入 Word"""
        if not self.capabilities['word']:
            raise DependencyError("python-docx 未安装")

        from docx import Document
        doc = Document(template) if template and Path(template).exists() else Document()

        if isinstance(content, str):
            doc.add_heading(path.stem, 0)
            for line in content.split('\n'):
                if line.strip():
                    doc.add_paragraph(line)
        elif isinstance(content, dict):
            if 'title' in content:
                doc.add_heading(content['title'], 0)
            if 'paragraphs' in content:
                for para in content['paragraphs']:
                    if isinstance(para, dict):
                        p = doc.add_paragraph(para.get('text', ''))
                    else:
                        doc.add_paragraph(str(para))
            if 'tables' in content:
                for table_data in content['tables']:
                    table = doc.add_table(
                        rows=len(table_data),
                        cols=len(table_data[0]) if table_data else 0
                    )
                    for i, row in enumerate(table_data):
                        for j, cell in enumerate(row):
                            table.cell(i, j).text = str(cell)

        doc.save(path)
        return str(path.absolute())

    def _enrich_word_info(self, path: Path, info: DocInfo):
        """丰富 Word 信息"""
        if not self.capabilities['word']:
            return
        doc = Document(path)
        info.pages = max(1, sum(len(p.text) for p in doc.paragraphs) // 1500)
        info.author = getattr(doc.core_properties, 'author', None)

    def _merge_word(self, files: List[str], output: Path) -> str:
        """合并 Word"""
        if not self.capabilities['word']:
            raise DependencyError("python-docx 未安装")

        from docx import Document
        merged = Document()

        for i, f in enumerate(files):
            if i > 0:
                merged.add_page_break()
            # 解析文件路径为绝对路径,避免 "Package not found" 错误
            file_path = self._resolve_path(f)
            self._validate_file(file_path)
            doc = Document(str(file_path))
            for para in doc.paragraphs:
                merged.add_paragraph(para.text)
            for table in doc.tables:
                new_table = merged.add_table(
                    rows=len(table.rows),
                    cols=len(table.columns)
                )
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        new_table.cell(i, j).text = cell.text

        merged.save(output)
        return str(output.absolute())

    # ========== Excel 操作 ==========
    def _read_excel(self, path: Path, sheet_names: List[str] = None) -> DocContent:
        """读取 Excel"""
        if not self.capabilities['excel']:
            raise DependencyError("openpyxl/pandas 未安装:pip install openpyxl pandas")

        import pandas as pd  # 延迟导入

        xl = load_workbook(path, read_only=True)
        all_sheets = xl.sheetnames
        xl.close()

        if sheet_names:
            target_sheets = [s for s in sheet_names if s in all_sheets]
        else:
            target_sheets = all_sheets

        sheets_data = {}
        for name in target_sheets:
            df = pd.read_excel(path, sheet_name=name)
            sheets_data[name] = {
                'columns': list(df.columns),
                'rows': df.values.tolist(),
                'count': len(df)
            }

        return DocContent(
            type='excel',
            format='table',
            data=sheets_data,
            metadata={'sheets': target_sheets}
        )

    def _write_excel(self, path: Path, content: Union[Dict, List, Any],
                     sheet_name: str = "Sheet1", has_header: bool = False) -> str:
        """写入 Excel"""
        if not self.capabilities['excel']:
            raise DependencyError("pandas 未安装")

        # 检查是否为 DataFrame(使用类型名称避免导入时错误)
        is_dataframe = hasattr(content, 'to_excel') and hasattr(content, 'columns')

        if is_dataframe:
            # DataFrame 为空时确保至少有一列
            if content.empty and len(content.columns) == 0:
                content = pd.DataFrame(columns=['Column'])
            content.to_excel(path, sheet_name=sheet_name, index=False)
        elif isinstance(content, dict):
            # 特殊处理:{'sheet_name': '...', 'data': [...]} 格式
            if 'sheet_name' in content and 'data' in content:
                # 使用指定的 sheet_name 和 data
                actual_sheet_name = content.get('sheet_name', sheet_name)[:31]
                data = content.get('data', [])
                is_df = hasattr(data, 'to_excel') and hasattr(data, 'columns')
                if is_df:
                    df = data
                    if df.empty and len(df.columns) == 0:
                        df = pd.DataFrame(columns=['Column'])
                else:
                    if not data:
                        df = pd.DataFrame(columns=['Column'])
                    else:
                        df = pd.DataFrame(data)
                df.to_excel(path, sheet_name=actual_sheet_name, index=False)
            else:
                # 字典为空时创建空文件
                if not content:
                    content = {sheet_name: []}
                with pd.ExcelWriter(path, engine='openpyxl') as writer:
                    for name, data in content.items():
                        # 检查是否为 DataFrame
                        is_df = hasattr(data, 'to_excel') and hasattr(data, 'columns')
                        if is_df:
                            df = data
                            if df.empty and len(df.columns) == 0:
                                df = pd.DataFrame(columns=['Column'])
                        else:
                            # 空列表时创建空 DataFrame
                            if not data:
                                df = pd.DataFrame(columns=['Column'])
                            else:
                                df = pd.DataFrame(data)
                        df.to_excel(writer, sheet_name=name[:31], index=False)
        else:
            # 列表或其他类型
            if not content:
                # 空列表时创建至少有一个 sheet 的空文件
                df = pd.DataFrame(columns=['Column'])
            else:
                if has_header and len(content) > 0:
                    # 第一行作为列名
                    df = pd.DataFrame(content[1:], columns=content[0])
                else:
                    df = pd.DataFrame(content)
            df.to_excel(path, sheet_name=sheet_name, index=False)

        return str(path.absolute())

    def _enrich_excel_info(self, path: Path, info: DocInfo):
        """丰富 Excel 信息"""
        if not self.capabilities['excel']:
            return
        xl = load_workbook(path, read_only=True)
        info.sheets = xl.sheetnames
        xl.close()

    def _enrich_csv_info(self, path: Path, info: DocInfo):
        """丰富 CSV 信息"""
        if self.capabilities['excel']:
            import pandas as pd  # 延迟导入
            df = pd.read_csv(path, nrows=1)
            info.columns = len(df.columns)
            info.rows = sum(1 for _ in open(path)) - 1

    def _merge_excel(self, files: List[str], output: Path) -> str:
        """合并 Excel"""
        if not self.capabilities['excel']:
            raise DependencyError("pandas 未安装")

        import pandas as pd  # 延迟导入

        all_sheets = {}
        for f in files:
            prefix = Path(f).stem
            xl = load_workbook(f, read_only=True)
            for sheet in xl.sheetnames:
                df = pd.read_excel(f, sheet_name=sheet)
                key = f"{prefix}_{sheet}"[:31]
                all_sheets[key] = df
            xl.close()

        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            for name, df in all_sheets.items():
                df.to_excel(writer, sheet_name=name, index=False)

        return str(output.absolute())

    # ========== CSV 操作 ==========
    def _read_csv(self, path: Path, delimiter: str = None) -> DocContent:
        """读取 CSV"""
        # 自动检测分隔符
        if delimiter is None:
            delimiter = self._detect_csv_delimiter(path)

        if self.capabilities['excel']:
            import pandas as pd  # 延迟导入
            try:
                df = pd.read_csv(path, delimiter=delimiter)
            except Exception as e:
                # 如果指定分隔符失败,尝试其他常见分隔符
                for alt_delim in [',', '\t', ';', ',']:
                    if alt_delim == delimiter:
                        continue
                    try:
                        df = pd.read_csv(path, delimiter=alt_delim)
                        delimiter = alt_delim
                        break
                    except (pd.errors.ParserError, ValueError):
                        continue
                else:
                    raise ReadError(f"无法解析 CSV 文件,尝试的分隔符:{delimiter}")

            return DocContent(
                type='csv',
                format='table',
                data={
                    'columns': list(df.columns),
                    'rows': df.values.tolist(),
                    'count': len(df)
                },
                metadata={'delimiter': delimiter}
            )
        else:
            # Fallback: 手动解析 CSV
            with open(path, encoding='utf-8') as f:
                lines = f.readlines()

            if lines:
                # 尝试解析为表格格式
                header = lines[0].strip().split(delimiter)
                rows = [line.strip().split(delimiter) for line in lines[1:] if line.strip()]
                data = {
                    'columns': header,
                    'rows': rows,
                    'count': len(rows)
                }
            else:
                data = {'columns': [], 'rows': [], 'count': 0}

            return DocContent(
                type='csv',
                format='table',
                data=data,
                metadata={'delimiter': delimiter}
            )

    def _detect_csv_delimiter(self, path: Path) -> str:
        """自动检测 CSV 分隔符"""
        # 读取前几行进行检测
        with open(path, 'r', encoding='utf-8') as f:
            lines = f.readlines()[:10]  # 读取前 10 行

        if not lines:
            return ','

        sample = ''.join(lines)

        # 常见分隔符优先级:英文逗号 > 中文逗号 > 制表符 > 分号
        delimiters = [',', ',', '\t', ';']

        best_delim = ','
        best_score = 0

        for delim in delimiters:
            if delim not in sample:
                continue

            # 验证:检查每行的列数是否一致
            try:
                col_counts = [len(line.strip().split(delim)) for line in lines if line.strip()]
                if not col_counts:
                    continue

                # 计算得分:列数一致 + 列数 > 1
                unique_counts = len(set(col_counts))
                avg_cols = sum(col_counts) / len(col_counts)

                # 得分规则:列数越一致得分越高,列数越多得分越高
                if unique_counts == 1 and avg_cols > 1:
                    # 完美匹配
                    score = 100 + avg_cols
                elif unique_counts <= 2 and avg_cols > 1:
                    # 基本一致
                    score = 50 + avg_cols - unique_counts * 10
                else:
                    score = avg_cols

                if score > best_score:
                    best_score = score
                    best_delim = delim

            except Exception as e:
                print(f"检测分隔符 {repr(delim)} 失败:{e}", file=sys.stderr)
                continue

        return best_delim

    def _write_csv(self, path: Path, content: Union[List, Dict, Any],
                   delimiter: str = ',') -> str:
        """写入 CSV"""
        # 检查是否为 DataFrame
        is_dataframe = hasattr(content, 'to_csv') and hasattr(content, 'columns')

        if self.capabilities['excel'] and (isinstance(content, dict) or is_dataframe):
            import pandas as pd  # 延迟导入
            df = content if is_dataframe else pd.DataFrame(content)
            df.to_csv(path, index=False, sep=delimiter)
        else:
            import csv
            with open(path, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f, delimiter=delimiter)
                if isinstance(content, dict) and 'columns' in content:
                    writer.writerow(content['columns'])
                    writer.writerows(content.get('rows', []))
                else:
                    writer.writerows(content)
        return str(path.absolute())

    # ========== TXT/MD 操作 ==========
    def _read_txt(self, path: Path, encoding: str = 'utf-8') -> DocContent:
        """读取文本"""
        with open(path, encoding=encoding) as f:
            content = f.read()
        return DocContent(
            type='txt',
            format='text',
            data=content,
            metadata={
                'lines': len(content.splitlines()),
                'chars': len(content)
            }
        )

    def _write_txt(self, path: Path, content: str, encoding: str = 'utf-8') -> str:
        """写入文本"""
        with open(path, 'w', encoding=encoding) as f:
            f.write(content)
        return str(path.absolute())


    # ========== v2.0 新增方法 ==========
    def fill_template(self, template_path: str, data: Dict, output_path: str = None) -> str:
        """填充模板生成文档 - v2.3 增强版"""
        template = Path(template_path)

        if not output_path:
            timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
            output_path = template.parent / f"filled_{timestamp}{template.suffix}"
        else:
            output_path = Path(output_path)

        # 填充文档(v2.3 支持 Word 和 Excel)
        # 注意:Excel 模板不需要内容适配,直接填充原始数据
        if template.suffix.lower() == '.docx':
            # Word 模板需要提取样式和适配内容
            style_def = self.style_extractor.extract(template)
            if style_def.sections:
                adapted_content = self.content_adapter.adapt_all(data, style_def)
            else:
                adapted_content = data
            return self._fill_word_template_v2(template, adapted_content, output_path)
        elif template.suffix.lower() == '.xlsx':
            # Excel 模板直接填充原始数据
            return self._fill_excel_template(template, data, output_path)
        else:
            shutil.copy2(template, output_path)
            return str(output_path)

    def _fill_word_template_v2(self, template: Path, data: Dict, output: Path) -> str:
        """填充 Word 模板 - v2.0 MVP"""
        if not WORD_AVAILABLE:
            raise DependencyError("python-docx 未安装")

        from docx import Document
        doc = Document(template)

        # 替换占位符 {{key}}
        for para in doc.paragraphs:
            for key, value in data.items():
                placeholder = '{{' + key + '}}'
                if placeholder in para.text:
                    if isinstance(value, list):
                        para.text = para.text.replace(placeholder, '\n'.join(value))
                    else:
                        para.text = para.text.replace(placeholder, str(value))

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in data.items():
                        placeholder = '{{' + key + '}}'
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, str(value))

        doc.save(output)
        return str(output.absolute())

    def _fill_excel_template(self, template: Path, data: Dict, output: Path) -> str:
        """填充 Excel 模板 - v2.3 新增"""
        if not EXCEL_AVAILABLE:
            raise DependencyError("openpyxl/pandas 未安装")

        from openpyxl import load_workbook
        wb = load_workbook(template)

        # 遍历所有工作表
        for ws in wb.worksheets:
            # 遍历所有单元格
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value and isinstance(cell.value, str):
                        # 替换占位符
                        for key, value in data.items():
                            placeholder = '{{' + key + '}}'
                            if placeholder in cell.value:
                                if isinstance(value, list):
                                    # 列表数据(表格区域填充)
                                    # 从当前单元格开始填充整个表格
                                    self._fill_excel_table_region(ws, cell.row, cell.column, value)
                                    # 清除占位符(如果单元格只有占位符)
                                    if cell.value.strip() == placeholder:
                                        cell.value = ''
                                else:
                                    cell.value = cell.value.replace(placeholder, str(value))

            # 处理命名单元格
            self._fill_excel_named_ranges(ws, data)

        wb.save(output)
        return str(output.absolute())

    def _fill_excel_named_ranges(self, ws, data: Dict):
        """填充 Excel 命名单元格 - v2.3 新增"""
        # 检查工作簿是否有命名区域
        if hasattr(ws.parent, 'defined_names'):
            defined_names = ws.parent.defined_names
            # DefinedNameDict 可以直接迭代
            for name_str, name in defined_names.items():
                if name_str in data:
                    # 获取命名单元格引用
                    try:
                        dest = list(name.destinations)
                        if dest:
                            ref_ws, ref = dest[0]
                            if ref_ws == ws.title:
                                # 填充值
                                ws[ref].value = data[name_str]
                    except (AttributeError, KeyError, TypeError):
                        pass  # 忽略命名单元格解析错误

    def _find_table_start(self, ws, table_name: str = None) -> tuple:
        """查找表格起始位置 - v2.3 新增"""
        # 策略 1: 查找包含 table_name 的单元格
        if table_name:
            for row_idx, row in enumerate(ws.iter_rows(), start=1):
                for cell in row:
                    if cell.value and table_name.lower() in str(cell.value).lower():
                        return (row_idx + 1, cell.column)  # 返回下一行作为数据起始

        # 策略 2: 查找第一个空行(表头后)
        prev_row_empty = False
        for row_idx, row in enumerate(ws.iter_rows(), start=1):
            row_empty = all(not cell.value or str(cell.value).strip() == '' for cell in row)
            if prev_row_empty and not row_empty:
                return (row_idx, 1)  # 返回行号和第 1 列
            prev_row_empty = row_empty

        # 默认:从第 2 行开始
        return (2, 1)

    def _fill_excel_table_region(self, ws, start_row: int, start_col: int, data: List):
        """填充表格区域 - v2.3 新增"""
        for i, row_data in enumerate(data):
            for j, value in enumerate(row_data):
                cell = ws.cell(row=start_row + i, column=start_col + j)
                cell.value = value

    def generate(self, user_request: str, template_path: str = None, content: Union[Dict, str] = None, output_path: str = None) -> GenerationResult:
        """智能文档生成 - v2.0 核心功能"""
        # 1. 路由选择
        if template_path:
            template_info = self._save_user_template(template_path, user_request)
            selected_template = template_info
        else:
            selected_template = self._match_historical_template(user_request)

        # 2. 生成文档
        if selected_template:
            if not output_path:
                timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
                output_path = self.workspace / f"generated_{timestamp}.docx"

            filled_path = self.fill_template(selected_template.path, content or {}, str(output_path))
            self._save_as_user_template(filled_path, user_request)

            return GenerationResult(output_path=filled_path, template_used=selected_template, stats={'method': 'template'})
        else:
            raise NotImplementedError("Style Guide 生成待实现")

    def _save_user_template(self, template_path: str, user_request: str) -> TemplateInfo:
        """保存用户模板"""
        import hashlib
        template = Path(template_path)
        file_hash = hashlib.md5(template.read_bytes()).hexdigest()[:8]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        template_id = f"tpl_{timestamp}_{file_hash}"
        saved_path = self.user_template_dir / "templates" / f"{template_id}{template.suffix}"
        shutil.copy2(template, saved_path)
        content_type = self._infer_content_type_from_request(user_request)
        template_info = TemplateInfo(path=str(saved_path), style=StyleDefinition(), content_type=content_type, usage_count=1, last_used=datetime.now().isoformat())
        self._update_registry(template_info)
        return template_info

    def _match_historical_template(self, user_request: str) -> Optional[TemplateInfo]:
        """匹配历史模板 - v2.2 智能多维度匹配"""
        registry_path = self.user_template_dir / "registry.json"
        if not registry_path.exists():
            return None

        with open(registry_path) as f:
            registry = json.load(f)

        templates = registry.get('templates', [])
        if not templates:
            return None

        # AI 分析请求
        request_analysis = self._analyze_request(user_request)

        # 多维度评分
        scored_templates = []
        for t in templates:
            score = self._calculate_match_score(t, request_analysis)
            if score > 0.5:  # 阈值
                scored_templates.append((score, t))

        if not scored_templates:
            return None

        # 排序选择最佳
        scored_templates.sort(key=lambda x: x[0], reverse=True)
        best_score, best_template = scored_templates[0]

        return TemplateInfo(
            path=best_template['path'],
            style=StyleDefinition(),
            content_type=best_template.get('content_type', 'general'),
            usage_count=best_template.get('usage_count', 0),
            last_used=best_template.get('last_used')
        )

    def _calculate_match_score(self, template: Dict, request_analysis: Dict) -> float:
        """计算模板匹配分数 - v2.2 新增"""
        score = 0.0

        # 文档类型匹配 (30%)
        if template.get('content_type') == request_analysis['document_type']:
            score += 0.30

        # 使用场景匹配 (20%)
        scenario_score = self._calculate_scenario_similarity(
            template.get('scenario', 'general'),
            request_analysis['scenario']
        )
        score += scenario_score * 0.20

        # 样式相似度 (20%)
        style_score = self._calculate_style_similarity(
            template.get('style', {}),
            request_analysis['style_preference']
        )
        score += style_score * 0.20

        # 最近使用 (15%)
        recency_score = self._calculate_recency_score(
            template.get('last_used')
        )
        score += recency_score * 0.15

        # 使用频率 (15%)
        usage_score = min(1.0, template.get('usage_count', 0) / 10)
        score += usage_score * 0.15

        return score

    def _calculate_scenario_similarity(self, template_scenario: str,
                                     request_scenario: str) -> float:
        """计算场景相似度 - v2.2 新增"""
        if template_scenario == request_scenario:
            return 1.0
        elif self._is_related_scenario(template_scenario, request_scenario):
            return 0.6
        else:
            return 0.2

    def _is_related_scenario(self, scenario1: str, scenario2: str) -> bool:
        """判断场景是否相关 - v2.2 新增"""
        related_groups = [
            ['technical', 'administrative'],
            ['business', 'financial'],
            ['hr', 'administrative'],
        ]

        for group in related_groups:
            if scenario1 in group and scenario2 in group:
                return True
        return False

    def _calculate_style_similarity(self, template_style: Dict,
                                   request_style: str) -> float:
        """计算样式相似度 - v2.2 新增"""
        if not template_style:
            return 0.5

        style_keywords = {
            'formal': ['正式', '商务', '严肃'],
            'concise': ['简洁', '简单', '快速'],
            'detailed': ['详细', '完整', '全面'],
            'casual': ['轻松', '随意']
        }

        template_style_str = str(template_style).lower()

        if request_style in template_style_str:
            return 1.0
        elif any(k in template_style_str for k in style_keywords.get(request_style, [])):
            return 0.7
        else:
            return 0.5

    def _calculate_recency_score(self, last_used: str) -> float:
        """计算最近使用分数 - v2.2 新增"""
        if not last_used:
            return 0.5

        try:
            last_date = datetime.fromisoformat(last_used.replace('Z', '+00:00'))
            days_ago = (datetime.now() - last_date).days

            if days_ago == 0:
                return 1.0
            elif days_ago <= 3:
                return 0.8
            elif days_ago <= 7:
                return 0.6
            elif days_ago <= 14:
                return 0.4
            else:
                return 0.2
        except (ValueError, TypeError, AttributeError):
            return 0.5

    def _explain_match(self, template: Dict, request_analysis: Dict, score: float) -> str:
        """生成匹配解释 - v2.2 新增"""
        reasons = []

        # 文档类型匹配
        if template.get('content_type') == request_analysis['document_type']:
            reasons.append(f"文档类型匹配 ({template.get('content_type')})")

        # 场景匹配
        template_scenario = template.get('scenario', 'general')
        if template_scenario == request_analysis['scenario']:
            reasons.append(f"使用场景匹配 ({template_scenario})")

        # 最近使用
        if template.get('last_used'):
            try:
                last_date = datetime.fromisoformat(template['last_used'].replace('Z', '+00:00'))
                days_ago = (datetime.now() - last_date).days
                if days_ago <= 3:
                    reasons.append(f"最近使用 ({days_ago}天前)")
            except (ValueError, TypeError, AttributeError):
                pass  # 忽略日期解析错误

        # 使用频率
        usage_count = template.get('usage_count', 0)
        if usage_count >= 5:
            reasons.append(f"高频使用 ({usage_count}次)")

        if reasons:
            return f"匹配原因:{', '.join(reasons)} (分数:{score:.2f})"
        else:
            return f"匹配分数:{score:.2f}"

    def _save_as_user_template(self, output_path: str, user_request: str) -> None:
        """保存生成为新历史模板"""
        template = Path(output_path)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        saved_path = self.user_template_dir / "templates" / f"{timestamp}_auto.docx"
        shutil.copy2(template, saved_path)
        content_type = self._infer_content_type_from_request(user_request)
        template_info = TemplateInfo(path=str(saved_path), style=StyleDefinition(), content_type=content_type, usage_count=1, last_used=datetime.now().isoformat())
        self._update_registry(template_info)

    def _update_registry(self, template_info: TemplateInfo):
        """更新注册表"""
        registry_path = self.user_template_dir / "registry.json"
        if registry_path.exists():
            with open(registry_path) as f:
                registry = json.load(f)
        else:
            registry = {'templates': [], 'categories': {}, 'stats': {}}
        registry['templates'].append({'path': template_info.path, 'content_type': template_info.content_type, 'usage_count': template_info.usage_count, 'last_used': template_info.last_used})
        with open(registry_path, 'w') as f:
            json.dump(registry, f, indent=2, ensure_ascii=False)

    def _analyze_request(self, user_request: str) -> Dict:
        """AI 分析用户请求 - v2.2 新增"""
        request_lower = user_request.lower()

        # 推断文档类型
        document_type = self._infer_document_type(request_lower)

        # 推断使用场景
        scenario = self._infer_scenario(request_lower)

        # 推断样式偏好
        style_preference = self._infer_style_preference(request_lower)

        # 提取关键要素
        key_elements = self._extract_key_elements(request_lower)

        return {
            'document_type': document_type,
            'scenario': scenario,
            'style_preference': style_preference,
            'key_elements': key_elements,
        }

    def _infer_document_type(self, request_lower: str) -> str:
        """推断文档类型 - v2.2 新增"""
        if any(k in request_lower for k in ['周报', 'weekly', '周报告']):
            return 'weekly-report'
        elif any(k in request_lower for k in ['月报', 'monthly', '月报告']):
            return 'monthly-report'
        elif any(k in request_lower for k in ['合同', 'contract', '协议']):
            return 'contract'
        elif any(k in request_lower for k in ['发票', 'invoice', '账单']):
            return 'invoice'
        elif any(k in request_lower for k in ['报告', 'report', '总结']):
            return 'report'
        elif any(k in request_lower for k in ['计划', 'plan', '规划']):
            return 'plan'
        else:
            return 'general'

    def _infer_scenario(self, request_lower: str) -> str:
        """推断使用场景 - v2.2 新增"""
        if any(k in request_lower for k in ['技术', 'tech', '开发', '代码']):
            return 'technical'
        elif any(k in request_lower for k in ['商务', 'business', '商业', '客户']):
            return 'business'
        elif any(k in request_lower for k in ['财务', 'finance', '会计', '金额']):
            return 'financial'
        elif any(k in request_lower for k in ['人事', 'hr', '人力', '员工']):
            return 'hr'
        elif any(k in request_lower for k in ['行政', 'admin', '办公']):
            return 'administrative'
        else:
            return 'general'

    def _infer_style_preference(self, request_lower: str) -> str:
        """推断样式偏好 - v2.2 新增"""
        if any(k in request_lower for k in ['简洁', '简单', '快速', '简短']):
            return 'concise'
        elif any(k in request_lower for k in ['详细', '完整', '全面', '充分']):
            return 'detailed'
        elif any(k in request_lower for k in ['正式', '严肃', '正规']):
            return 'formal'
        elif any(k in request_lower for k in ['轻松', '随意', '非正式']):
            return 'casual'
        else:
            return 'standard'

    def _extract_key_elements(self, request_lower: str) -> List[str]:
        """提取关键要素 - v2.2 新增"""
        elements = []

        if any(k in request_lower for k in ['本周', '上周', '这周']):
            elements.append('time_week')
        elif any(k in request_lower for k in ['本月', '上月', '这个月']):
            elements.append('time_month')

        if any(k in request_lower for k in ['项目', 'project', '任务']):
            elements.append('project')

        if any(k in request_lower for k in ['数据', 'data', '统计', '数字']):
            elements.append('data')

        return elements

    def _infer_content_type_from_request(self, request: str) -> str:
        """从请求推断内容类型(向后兼容)"""
        analysis = self._analyze_request(request)
        return analysis['document_type']


# ========== v2.5.0 批量处理 ==========
    def batch_fill_templates(
        self,
        templates: List[Union[str, Path]],
        data_list: List[Dict],
        output_dir: Union[str, Path],
        max_workers: int = 4,
        progress_callback: Optional[Callable] = None
    ) -> BatchResult:
        """批量填充模板 - v2.5.0 新增"""
        from concurrent.futures import ThreadPoolExecutor, as_completed
        import time

        if not templates:
            raise ValueError("模板列表为空")
        if len(templates) != len(data_list):
            raise ValueError(f"模板数量 ({len(templates)}) 与数据数量 ({len(data_list)}) 不匹配")

        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)

        logger.info(f"开始批量处理:{len(templates)} 个文件,并发数:{max_workers}")
        start_time = time.time()

        success_count = 0
        failed_count = 0
        errors = []
        output_paths = []

        def process_single(i):
            try:
                result_path = self.fill_template(
                    str(templates[i]),
                    data_list[i],
                    str(output_path / f"filled_{i}{Path(templates[i]).suffix}")
                )
                return {'success': True, 'path': result_path, 'index': i}
            except Exception as e:
                return {'success': False, 'error': str(e), 'index': i}

        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            futures = {executor.submit(process_single, i): i for i in range(len(templates))}

            for future in as_completed(futures):
                result = future.result()
                if result['success']:
                    success_count += 1
                    output_paths.append(result['path'])
                else:
                    failed_count += 1
                    errors.append({
                        'index': result['index'],
                        'template': str(templates[result['index']]),
                        'error': result['error']
                    })

                if progress_callback:
                    progress_callback(success_count + failed_count, len(templates), success_count, failed_count)

        duration = time.time() - start_time

        result = BatchResult(
            success=success_count,
            failed=failed_count,
            total=len(templates),
            duration=duration,
            errors=errors,
            output_paths=output_paths
        )

        logger.info(f"批量处理完成:成功 {success_count}/{len(templates)}, 耗时 {duration:.2f}秒")
        return result

    # ========== v2.7.0 性能监控 ==========
    def get_performance_stats(self) -> Dict:
        """获取性能统计信息"""
        if not self.perf_monitor:
            return {'enabled': False}
        stats = self.perf_monitor.get_stats()
        return {'enabled': True, 'operations': {k: v.to_dict() for k, v in stats.items()} if stats else {}}

    def print_performance_report(self):
        """打印性能报告"""
        if self.perf_monitor:
            self.perf_monitor.print_report()
        else:
            print("性能监控未启用")

    # ========== v2.7.0 缓存管理 ==========
    def get_cache_stats(self) -> Dict:
        """获取缓存统计信息"""
        if not self.template_cache:
            return {'enabled': False}
        stats = self.template_cache.stats()
        stats['enabled'] = True
        return stats

    def clear_cache(self):
        """清空缓存"""
        if self.template_cache:
            self.template_cache.clear()
            logger.info("模板缓存已清空")

    def preload_templates(self, template_paths: List[str]):
        """预加载模板到缓存"""
        if not self.template_cache:
            logger.info("缓存未启用,跳过预加载")
            return
        logger.info(f"预加载 {len(template_paths)} 个模板到缓存")
        for path in template_paths:
            try:
                template = Path(path)
                ext = template.suffix.lower()[1:]
                self.template_cache.get(f"{ext}:{template.absolute()}")
                logger.debug(f"已预加载:{path}")
            except Exception as e:
                logger.warning(f"预加载失败 {path}: {e}")


# ========== CLI 入口 ==========
def main():
    import argparse

    parser = argparse.ArgumentParser(
        description="文档处理工具",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  %(prog)s --action info -i report.pdf
  %(prog)s --action read -i data.xlsx
  %(prog)s --action write -o output.docx --content '{"title": "标题"}'
  %(prog)s --action merge -i a.docx -i b.docx -o merged.docx
        """
    )

    parser.add_argument(
        "--action", required=True,
        choices=["info", "read", "write", "convert", "merge", "extract", "capabilities"],
        help="操作类型"
    )
    parser.add_argument(
        "--input", "-i", action="append", dest="inputs",
        help="输入文件路径 (可多次)"
    )
    parser.add_argument(
        "--output", "-o",
        help="输出文件路径"
    )
    parser.add_argument(
        "--content", "-c",
        help="写入内容 (JSON 或文本)"
    )
    parser.add_argument(
        "--content-file",
        help="写入内容文件路径"
    )
    parser.add_argument(
        "--workspace", "-w",
        help="工作目录"
    )
    parser.add_argument(
        "--format", "-f", choices=["json", "text"], default="json",
        help="输出格式"
    )
    parser.add_argument(
        "--options",
        help="额外选项 (JSON)"
    )

    args = parser.parse_args()
    processor = DocumentProcessor(workspace=args.workspace)

    try:
        result = None

        if args.action == "capabilities":
            result = processor.get_capabilities()

        elif args.action == "info":
            if not args.inputs:
                raise ValueError("需要指定输入文件:-i <file>")
            result = processor.get_info(args.inputs[0]).to_dict()

        elif args.action == "read":
            if not args.inputs:
                raise ValueError("需要指定输入文件:-i <file>")
            options = json.loads(args.options) if args.options else {}
            result = processor.read(args.inputs[0], **options).to_dict()

        elif args.action == "write":
            if not args.output:
                raise ValueError("需要指定输出文件:-o <file>")

            if args.content_file:
                with open(args.content_file) as f:
                    content = json.load(f)
            elif args.content:
                content = json.loads(args.content) if args.content.startswith('{') else args.content
            else:
                content = sys.stdin.read()

            options = json.loads(args.options) if args.options else {}
            result = {"path": processor.write(args.output, content, **options)}

        elif args.action == "convert":
            if not args.inputs or not args.output:
                raise ValueError("需要指定输入和输出文件")
            result = {"path": processor.convert(args.inputs[0], args.output)}

        elif args.action == "merge":
            if not args.inputs or not args.output:
                raise ValueError("需要指定输入文件列表和输出文件")
            result = {"path": processor.merge(args.inputs, args.output)}

        elif args.action == "extract":
            if not args.inputs:
                raise ValueError("需要指定输入文件")
            options = json.loads(args.options) if args.options else {}
            result = processor.extract(args.inputs[0], **options)

        if args.format == "json":
            print(json.dumps(result, ensure_ascii=False, indent=2, default=str))
        else:
            print(result)

    except Exception as e:
        error_result = {
            "error": str(e),
            "type": type(e).__name__,
            "action": args.action
        }
        print(json.dumps(error_result, ensure_ascii=False, indent=2), file=sys.stderr)
        sys.exit(1)


