#!/usr/bin/env python3
"""
DeckCraft v5.3 — Importer Tests

Tests the source importers (PDF, DOCX, MD) on synthetic documents and
verifies the resulting outline is well-formed and produces PPTX output.

Run:
    python3 tests/test_importers.py
"""
import sys, os, json, tempfile
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from engine.importers import pdf_to_outline, docx_to_outline, text_to_outline, detect_and_import
from engine.importers.base import (
    looks_like_toc, looks_like_stats, looks_like_table,
    looks_like_section_divider, looks_like_cover,
    extract_bullets, extract_table, extract_stats, extract_heading,
)
from engine import DeckEngine
from scripts.generate_ppt import PAGE_HANDLERS


passed = 0
total = 0


def expect(actual, expected, desc):
    global passed, total
    total += 1
    if actual == expected:
        print(f"  ✓ {desc}")
        passed += 1
    else:
        print(f"  ✗ {desc}: expected {expected!r}, got {actual!r}")


def expect_true(cond, desc):
    expect(bool(cond), True, desc)


def expect_raises(fn, exc_type, desc):
    global passed, total
    total += 1
    try:
        fn()
    except exc_type as e:
        print(f"  ✓ {desc}: {type(e).__name__}: {str(e)[:80]}")
        passed += 1
        return
    except Exception as e:
        print(f"  ✗ {desc}: expected {exc_type.__name__}, got {type(e).__name__}: {e}")
        return
    print(f"  ✗ {desc}: expected {exc_type.__name__}, no exception raised")


# ── Heuristics ─────────────────────────────────────────────────────────
print("=" * 60)
print("Heuristics")
print("=" * 60)

expect(looks_like_cover("Q3 Plan\n2026", is_first_page=True), True,
       "Short text + first page → cover")
expect(looks_like_cover("Q3 Plan\n2026", is_first_page=False), False,
       "Short text + not first page → not cover")
expect(looks_like_toc("Agenda\n01  A\n02  B\n03  C"), True,
       "Agenda heading + numbered list → toc")
expect(looks_like_toc("01  A\n02  B\n03  C\n04  D"), True,
       "4 numbered short lines → toc")
expect(looks_like_toc("01. A\n02. B\n03. C"), True,
       "3 numbered lines with dot → toc")
expect(looks_like_toc("- P1\n- P2\n- P3"), False,
       "3 bullets are NOT toc")
expect(looks_like_toc("Industry\n- P1\n- P2"), False,
       "2 bullets + heading → not toc")
expect(looks_like_stats("99%  Uptime\n$2M  Revenue"), True,
       "Stat lines detected")
expect(looks_like_stats("Black Friday push\nThree moments"), False,
       "Plain lines → not stats")
expect(looks_like_table("| A | B |\n|---|---|\n| 1 | 2 |"), True,
       "Pipe table detected")

# extract_bullets
bullets = extract_bullets("- First\n- Second\n  - Nested\n* Third")
expect(len(bullets) >= 2, True, "extract_bullets gets multiple items")
expect("First" in bullets, True, "extract_bullets captures 'First'")

# extract_heading
expect(extract_heading("Big Title\nSubtitle\n- bullet"), "Big Title",
       "extract_heading returns first non-empty line")

# extract_table
table = extract_table("| A | B |\n|---|---|\n| 1 | 2 |\n| 3 | 4 |")
expect(table is not None, True, "extract_table parses pipe table")
if table:
    expect(table["headers"], ["A", "B"], "extract_table gets headers")
    expect(len(table["rows"]), 2, "extract_table gets 2 data rows")

# ── Text importer ──────────────────────────────────────────────────────
print()
print("=" * 60)
print("Text/Markdown importer")
print("=" * 60)

with tempfile.TemporaryDirectory() as tmp:
    md_path = os.path.join(tmp, "test.md")
    with open(md_path, "w") as f:
        f.write("""# My Report

A short summary.

# Section 1: Background

Some context here.

# Section 2: Findings

- Finding one
- Finding two
- Finding three
""")

    outline = text_to_outline(md_path)
    expect(outline["theme"], "business", "outline theme defaults to business")
    expect(outline["canvas"], "16:9", "outline canvas defaults to 16:9")
    expect(len(outline["pages"]) >= 2, True, "text_to_outline produces multiple pages")
    expect(outline["pages"][0]["type"], "cover", "first page is cover")
    expect(any(p["type"] == "closing" for p in outline["pages"]), True,
           "text_to_outline appends closing")

# ── DOCX importer ──────────────────────────────────────────────────────
print()
print("=" * 60)
print("DOCX importer")
print("=" * 60)

try:
    import docx
    with tempfile.TemporaryDirectory() as tmp:
        docx_path = os.path.join(tmp, "test.docx")
        d = docx.Document()
        d.add_heading("My Doc", level=0)
        d.add_paragraph("Author")
        d.add_paragraph("Date")
        d.add_heading("Agenda", level=1)
        d.add_paragraph("01  Topic 1")
        d.add_paragraph("02  Topic 2")
        d.add_paragraph("03  Topic 3")
        d.add_heading("Findings", level=1)
        d.add_paragraph("- Finding A")
        d.add_paragraph("- Finding B")
        d.add_heading("Stats", level=1)
        d.add_paragraph("99%  Rate")
        d.add_paragraph("$2M  Revenue")
        d.save(docx_path)

        outline = docx_to_outline(docx_path)
        expect(len(outline["pages"]) >= 3, True, "docx_to_outline produces pages")
        expect(outline["pages"][0]["type"], "cover", "first page is cover")
        types = [p["type"] for p in outline["pages"]]
        expect("toc" in types, True, "Agenda section detected as toc")
        expect("stat_cards" in types, True, "Stats section detected as stat_cards")
except ImportError:
    print("  ⚠ python-docx not installed, skipping DOCX tests")

# ── PDF importer ───────────────────────────────────────────────────────
print()
print("=" * 60)
print("PDF importer")
print("=" * 60)

try:
    import fitz
    with tempfile.TemporaryDirectory() as tmp:
        pdf_path = os.path.join(tmp, "test.pdf")
        doc = fitz.open()
        for title, body in [
            ("Q3 Plan", "Strategic Roadmap\nMarketing Team\n2026"),
            ("Agenda", "01  Part A\n02  Part B\n03  Part C"),
            ("Findings", "Key takeaways:\n- Point 1\n- Point 2\n- Point 3"),
            ("Numbers", "99%  Rate\n$2M  Revenue\n12  Partners"),
        ]:
            page = doc.new_page()
            page.insert_text((72, 72), title, fontsize=20)
            page.insert_text((72, 100), body, fontsize=12)
        doc.save(pdf_path)
        doc.close()

        outline = pdf_to_outline(pdf_path)
        expect(len(outline["pages"]) >= 3, True, "pdf_to_outline produces pages")
        types = [p["type"] for p in outline["pages"]]
        expect("cover" in types, True, "First page detected as cover")
        expect("toc" in types, True, "Agenda page detected as toc")
        expect("stat_cards" in types, True, "Numbers page detected as stat_cards")
except ImportError:
    print("  ⚠ PyMuPDF not installed, skipping PDF tests")

# ── detect_and_import ─────────────────────────────────────────────────
print()
print("=" * 60)
print("detect_and_import")
print("=" * 60)

with tempfile.TemporaryDirectory() as tmp:
    expect_raises(lambda: detect_and_import("/nonexistent.pdf"),
                  FileNotFoundError, "Missing file → FileNotFoundError")

    md_path = os.path.join(tmp, "x.md")
    with open(md_path, "w") as f:
        f.write("# Title\n\nBody.")
    outline = detect_and_import(md_path)
    expect(outline["pages"][0]["type"], "cover", "MD detected by extension")

    bogus = os.path.join(tmp, "x.xyz")
    with open(bogus, "w") as f:
        f.write("hi")
    expect_raises(lambda: detect_and_import(bogus), ValueError,
                  "Unknown extension → ValueError")

# ── End-to-end: outline → PPTX ────────────────────────────────────────
print()
print("=" * 60)
print("End-to-end: outline → PPTX (gate check)")
print("=" * 60)

with tempfile.TemporaryDirectory() as tmp:
    md_path = os.path.join(tmp, "x.md")
    with open(md_path, "w") as f:
        f.write("""# Test Deck

A simple test.

# Section A

- A point
- Another point

# Section B

99%  Test
$1M  Cost
""")
    outline = detect_and_import(md_path)
    eng = DeckEngine()
    for page in outline["pages"]:
        handler = PAGE_HANDLERS.get(page.get("type"))
        if handler:
            handler(eng, page)
    pptx_path = os.path.join(tmp, "out.pptx")
    eng.save(pptx_path)
    expect(os.path.exists(pptx_path), True, "End-to-end PPTX created")
    expect(eng.slide_count >= 3, True, "End-to-end generates multiple slides")

# ── Summary ────────────────────────────────────────────────────────────
print()
print("=" * 60)
print(f"Passed: {passed} / {total}")
if passed == total:
    print("✓ All importer tests passed.")
    sys.exit(0)
print(f"✗ {total - passed} test(s) failed.")
sys.exit(1)
