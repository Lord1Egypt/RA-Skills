#!/usr/bin/env python3

# ============================================================
# ⚠️ SAMPLE DATA NOTICE / 示例数据声明
# This template contains EXAMPLE / ILLUSTRATIVE data values.
# Users must configure their own data sources (yfinance, Guosen API, etc.)
# to generate plans with real market data.
# 本模板包含的数据为示例/演示用途。
# 用户需自行配置数据源（yfinance、国信API等）以获取真实市场数据。
# ============================================================

"""
💱 Forex Trading Plan Generator v7.0 FINAL / 外汇交易方案生成器
Ai-Wife Global Investing / 媳妇智投全球通用版
Bilingual (CN/EN) Edition / 中英双语版

Usage / 用法:
  python3 template_fx_v7.py "EUR/USD"
  python3 template_fx_v7.py "GBP/USD"
  python3 template_fx_v7.py "USD/JPY"
"""

import datetime
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

def create_forex_plan(symbol="EUR/USD", date_str=None):
    """Generate bilingual forex trading plan / 生成中英双语外汇交易方案"""
    if date_str is None:
        date_str = datetime.datetime.now().strftime("%Y-%m-%d")
    
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    
    # ==================== TITLE / 标题 ====================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"💱 {symbol} Forex Reference Framework / 外汇参考框架")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Version: v7.0 FINAL | Date / 制作日期: {date_str}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("⚠️ 教育用途 | Educational Use Only")
    run.font.size = Pt(8)
    run.italic = True
    
    doc.add_paragraph()
    
    # ==================== HEADER / 方案头部 ====================
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Light Grid Accent 1'
    cells = tbl.rows[0].cells
    cells[0].paragraphs[0].add_run(f'Item / {_bt.item_label}').bold = True
    cells[1].paragraphs[0].add_run('Content / 内容').bold = True
    
    rows_data = [
        ('Currency Pair / 货币对', symbol),
        ('Category / 品种分类', 'Major / Cross / 主要货币对 / 交叉盘'),
        ('Market / 交易市场', 'Global OTC / 全球场外交易（24H）'),
        ('Trading Hours / 交易时间', 'Mon 05:00 — Sat 05:00 EST / 24-hour rolling'),
        ('Pip Size / 最小波动', '1 pip = 0.0001 (most majors / 多数主要货币对)'),
        ('Daily Volume / 日交易量', '~$6.6 trillion USD (world\'s largest market / 全球最大)'),
        ('Date / 日期', date_str),
        ('Type / 方案类型', 'Swing / Intraday / Trend / 波段/日内/趋势'),
    ]
    for item, val in rows_data:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(item).bold = True
        row.cells[1].paragraphs[0].add_run(val)
    
    # ==================== 一、Market Review / 行情回顾 ====================
    doc.add_heading('一、Market Review / 行情回顾', level=2)
    
    doc.add_heading('1.1 Close Data / 收盘数据', level=3)
    tbl = doc.add_table(rows=1, cols=7)
    tbl.style = 'Light Grid Accent 1'
    hdr = ['Pair / 货币对', 'Last / 最新价', 'Change(pips)', 'Change(%) / 涨跌幅', 'High / 最高', 'Low / 最低', 'Range / 振幅']
    for i, h in enumerate(hdr):
        tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    row = tbl.add_row()
    for i in range(7):
        row.cells[i].paragraphs[0].add_run('—')
    
    doc.add_heading('1.2 Recent Trend / 近期走势', level=3)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Light Grid Accent 1'
    hdr = ['Period / 时间段', 'Price Range / 价格区间', 'Change / 涨跌幅', 'Character / 走势特征']
    for i, h in enumerate(hdr):
        tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    periods = ['5-Day / 近5日', '20-Day / 近20日', '60-Day / 近60日']
    for pp in periods:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(pp)
        for j in range(1, 4):
            row.cells[j].paragraphs[0].add_run('—')
    
    doc.add_heading('Summary / 走势总结', level=4)
    p = doc.add_paragraph('—')
    
    # ==================== 二、Market Correlation / 市场联动分析 ====================
    doc.add_heading('二、Market Correlation / 市场联动分析', level=2)
    
    doc.add_heading('2.1 US Dollar Index (DXY) / 美元指数', level=3)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Light Grid Accent 1'
    hdr = [f'Item / {_bt.item_label}', 'Value / 数值', 'Daily Change / 日变化', 'Trend / 趋势']
    for i, h in enumerate(hdr):
        tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    row = tbl.add_row()
    for i in range(4):
        row.cells[i].paragraphs[0].add_run('—')
    
    doc.add_heading('2.2 Related Pairs / 关联品种', level=3)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Light Grid Accent 1'
    hdr = ['Pair / 品种', 'Price / 最新价', 'Trend / 走势', 'Correlation / 相关性']
    for i, h in enumerate(hdr):
        tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    for r in ['EUR/USD', 'GBP/USD', 'USD/JPY', 'USD/CHF', 'XAU/USD (Gold/黄金)']:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(r)
        for j in range(1, 4):
            row.cells[j].paragraphs[0].add_run('—')
    
    doc.add_heading('2.3 Correlation Conclusion / 联动结论', level=3)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Light Grid Accent 1'
    hdr = ['Dimension / 维度', 'Judgment / 判断', 'Note / 说明']
    for i, h in enumerate(hdr):
        tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    for dim in ['Correlation Strength / 相关性强度', 'Direction / 联动方向', 'Trading Reference / 操作参考']:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(dim).bold = True
        row.cells[1].paragraphs[0].add_run('—')
        row.cells[2].paragraphs[0].add_run('—')
    
    # ==================== 三、Macro Analysis / 宏观影响分析 ====================
    doc.add_heading('三、Macro Analysis / 宏观影响分析', level=2)
    
    doc.add_heading('3.1 Bearish Factors / 利空因素 🔴', level=3)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Light Grid Accent 1'
    hdr = ['Factor / 因素', 'Detail / 详情', 'Impact / 影响程度', 'Duration / 持续性']
    for i, h in enumerate(hdr):
        tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    for _ in range(2):
        row = tbl.add_row()
        for i in range(4):
            row.cells[i].paragraphs[0].add_run('—')
    
    doc.add_heading('3.2 Bullish Factors / 利多因素 🟢', level=3)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Light Grid Accent 1'
    hdr = ['Factor / 因素', 'Detail / 详情', 'Impact / 影响程度', 'Duration / 持续性']
    for i, h in enumerate(hdr):
        tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    for _ in range(2):
        row = tbl.add_row()
        for i in range(4):
            row.cells[i].paragraphs[0].add_run('—')
    
    doc.add_heading('3.3 Macro Conclusion / 宏观结论', level=3)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Light Grid Accent 1'
    cells = tbl.rows[0].cells
    cells[0].paragraphs[0].add_run('Dimension / 维度').bold = True
    cells[1].paragraphs[0].add_run('Judgment / 判断').bold = True
    for dim in ['Bull vs Bear / 多空对比', 'Short-term Outlook / 短期展望', 'Mid-term Outlook / 中期展望', 'Key Driver / 核心驱动']:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(dim).bold = True
        row.cells[1].paragraphs[0].add_run('—')
    
    # ==================== 四、Central Bank Policy / 央行政策分析 ====================
    doc.add_heading('四、Central Bank Policy / 央行政策分析 🏦', level=2)
    
    doc.add_heading('4.1 Central Bank Latest / 主要央行最新动态', level=3)
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = 'Light Grid Accent 1'
    hdr = ['Central Bank / 央行', 'Rate / 当前利率', 'Last Change / 上次变动', 'Next Meeting / 下次决议', 'Market Expectation / 市场预期']
    for i, h in enumerate(hdr):
        tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    for cb in ['FOMC / 美联储', 'ECB / 欧洲央行', 'BOJ / BOE / Other / 其他相关央行']:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(cb)
        for j in range(1, 5):
            row.cells[j].paragraphs[0].add_run('—')
    
    doc.add_heading('4.2 Rate Path Outlook / 利率预期路径', level=3)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Light Grid Accent 1'
    hdr = ['Timeline / 时间', 'FOMC Expectation / 美联储预期', 'Other CB / 其他央行预期']
    for i, h in enumerate(hdr):
        tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    for period in ['1 Month / 1个月', '3 Months / 3个月', '6 Months / 6个月', '12 Months / 12个月']:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(period)
        row.cells[1].paragraphs[0].add_run('—')
        row.cells[2].paragraphs[0].add_run('—')
    
    doc.add_heading('4.3 Interest Rate Spread / 利差分析', level=3)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Light Grid Accent 1'
    hdr = ['Dimension / 维度', 'Value / 数值', 'Trend / 趋势']
    for i, h in enumerate(hdr):
        tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    for dim in ['US-EU Spread (2Y) / 美-欧利差(2Y)', 'US-EU Spread (10Y) / 美-欧利差(10Y)', 'Spread Impact / 利差对汇率影响']:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(dim).bold = True
        row.cells[1].paragraphs[0].add_run('—')
        row.cells[2].paragraphs[0].add_run('—')
    
    # ==================== 五、Economic Data / 经济数据分析 ====================
    doc.add_heading('五、Economic Data / 经济数据分析 📊', level=2)
    
    doc.add_heading('5.1 Latest Releases / 最新发布数据', level=3)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Light Grid Accent 1'
    hdr = ['Indicator / 经济指标', 'Actual / 实际值', 'Expected / 预期值', 'Previous / 前值']
    for i, h in enumerate(hdr):
        tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    for ind in ['CPI YoY / CPI年率', 'GDP QoQ / GDP季率', 'Manufacturing PMI / 制造业PMI', 'Services PMI / 服务业PMI', 'Unemployment Rate / 失业率', 'Retail Sales / 零售销售']:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(ind)
        for j in range(1, 4):
            row.cells[j].paragraphs[0].add_run('—')
    
    doc.add_heading('5.2 Cross-Country Comparison / 两国数据对比', level=3)
    tbl = doc.add_table(rows=1, cols=7)
    tbl.style = 'Light Grid Accent 1'
    hdr = ['Indicator / 指标', 'US / 美国', 'EU/Other / 欧元区/对方', 'Spread / 差值', 'Trend / 趋势', 'FX Impact / 汇率影响', 'Direction / 方向']
    for i, h in enumerate(hdr):
        tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    for ind in ['GDP Growth / GDP增速', 'CPI Inflation / CPI通胀', 'Unemployment / 失业率', 'Policy Rate / 基准利率']:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(ind).bold = True
        for j in range(1, 7):
            row.cells[j].paragraphs[0].add_run('—')
    
    doc.add_heading('5.3 Economic Conclusion / 经济结论', level=3)
    p = doc.add_paragraph('—')
    
    # ==================== 六、Technical Analysis / 技术分析 ====================
    doc.add_heading('六、Technical Analysis / 技术分析 📈', level=2)
    
    doc.add_heading('6.1 Key Levels / 关键价位', level=3)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Light Grid Accent 1'
    hdr = ['Type / 类型', 'Level / 价位', 'Note / 说明']
    for i, h in enumerate(hdr):
        tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    for level in ['Strong Resistance R3 / 强压力', 'Medium Resistance R2 / 中压力', 'Weak Resistance R1 / 弱压力', 
                  'Current / 当前价', 'Weak Support S1 / 弱支撑', 'Medium Support S2 / 中支撑', 'Strong Support S3 / 强支撑']:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(level).bold = True
        row.cells[1].paragraphs[0].add_run('—')
        row.cells[2].paragraphs[0].add_run('—')
    
    doc.add_heading('6.2 Technical Indicators / 技术指标', level=3)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Light Grid Accent 1'
    hdr = ['Indicator / 指标', 'Status / 状态', 'Signal / 信号']
    for i, h in enumerate(hdr):
        tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    indicators = [
        ('MA System (20/50/200) / 均线系统', '—', '—'),
        ('MACD (12,26,9)', '—', '—'),
        ('RSI (14)', '—', '—'),
        ('Bollinger Bands (20,2) / 布林带', '—', '—'),
        ('Volume / Open Interest / 成交量/持仓量', '—', '—'),
        ('Pivot Points / 枢轴点', '—', '—'),
    ]
    for ind, state, sig in indicators:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(ind).bold = True
        row.cells[1].paragraphs[0].add_run(state)
        row.cells[2].paragraphs[0].add_run(sig)
    
    doc.add_heading('6.3 Technical Conclusion / 技术结论', level=3)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Light Grid Accent 1'
    cells = tbl.rows[0].cells
    cells[0].paragraphs[0].add_run('Dimension / 维度').bold = True
    cells[1].paragraphs[0].add_run('Judgment / 判断').bold = True
    for dim in ['Trend Direction / 趋势方向', 'Trend Strength / 趋势强度', 'Key Level / 关键位置', 'Divergence / 背离信号']:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(dim).bold = True
        row.cells[1].paragraphs[0].add_run('—')
    
    # ==================== 七、CFTC Positioning / CFTC持仓分析 ====================
    doc.add_heading('七、CFTC Positioning & Fund Flow / CFTC持仓 & 资金动向', level=2)
    
    doc.add_heading('7.1 CFTC Commitments of Traders / CFTC持仓数据', level=3)
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = 'Light Grid Accent 1'
    hdr = ['Category / 类别', 'Long / 多头', 'Short / 空头', 'Net / 净持仓', 'Change / 变化']
    for i, h in enumerate(hdr):
        tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    for cat in ['Speculative Net / 投机性净持仓', 'Commercial Hedge / 商业套保', 'Total Open Interest / 总持仓']:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(cat).bold = True
        for j in range(1, 5):
            row.cells[j].paragraphs[0].add_run('—')
    
    doc.add_heading('7.2 Positioning Conclusion / 持仓结论', level=3)
    p = doc.add_paragraph('—')
    
    # ==================== 八、Market Sentiment / 市场情绪 ====================
    doc.add_heading('八、Market Sentiment & Flow / 市场情绪与资金流 🌊', level=2)
    
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Light Grid Accent 1'
    hdr = ['Dimension / 维度', 'Value / 数值/状态', 'Interpretation / 解读']
    for i, h in enumerate(hdr):
        tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    for dim in ['VIX / 恐慌指数', 'Risk Appetite / 风险偏好', 'Long/Short Ratio / 多空比', 'Retail Sentiment / 散户情绪', 'Institutional Sentiment / 机构情绪']:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(dim).bold = True
        row.cells[1].paragraphs[0].add_run('—')
        row.cells[2].paragraphs[0].add_run('—')
    
    # ==================== 九、Trading Plan / 操作方案 ====================
    doc.add_heading('九、Trading Plan / 操作方案 🎯', level=2)
    
    p = doc.add_paragraph()
    run = p.add_run('💱 Forex: 24H rolling, T+0 both-way, high leverage (20:1~50:1) / 24小时滚动，T+0双向，高杠杆')
    run.italic = True
    run.font.size = Pt(9)
    
    doc.add_heading('9.1 Overall Strategy / 总体策略', level=3)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Light Grid Accent 1'
    cells = tbl.rows[0].cells
    cells[0].paragraphs[0].add_run(f'Item / {_bt.item_label}').bold = True
    cells[1].paragraphs[0].add_run('Content / 内容').bold = True
    for item in ['Core Range / 核心区间', 'Direction / 主要方向', 'Direction Judgment / 方向判断', 'Strategy / 操作思路', 'Position Size / 仓位建议']:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(item).bold = True
        row.cells[1].paragraphs[0].add_run('—')
    
    doc.add_heading('9.2 ⬆️ Long Plan / 做多方案', level=3)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Light Grid Accent 1'
    cells = tbl.rows[0].cells
    cells[0].paragraphs[0].add_run(f'Item / {_bt.item_label}').bold = True
    cells[1].paragraphs[0].add_run('Content / 内容').bold = True
    for item in ['Scenario / 适用场景', 'Entry Condition / 入场条件', 'Entry Price / 入场价位', 'Stop Loss / 止损价位', 
                 'Target 1 (T1) / 第一目标', 'Target 2 (T2) / 第二目标', 'Position Size / 仓位控制', 'Risk/Reward / 盈亏比']:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(item).bold = True
        row.cells[1].paragraphs[0].add_run('—')
    
    doc.add_heading('9.3 ⬇️ Short Plan / 做空方案', level=3)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Light Grid Accent 1'
    cells = tbl.rows[0].cells
    cells[0].paragraphs[0].add_run(f'Item / {_bt.item_label}').bold = True
    cells[1].paragraphs[0].add_run('Content / 内容').bold = True
    for item in ['Scenario / 适用场景', 'Entry Condition / 入场条件', 'Entry Price / 入场价位', 'Stop Loss / 止损价位',
                 'Target 1 (T1) / 第一目标', 'Target 2 (T2) / 第二目标', 'Position Size / 仓位控制', 'Risk/Reward / 盈亏比']:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(item).bold = True
        row.cells[1].paragraphs[0].add_run('—')
    
    doc.add_heading('9.4 Trading Session Strategy / 交易时段策略', level=3)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Light Grid Accent 1'
    hdr = ['Session / 交易时段', 'Beijing Time / 北京时间', 'Liquidity / 流动性', 'Strategy / 策略建议']
    for i, h in enumerate(hdr):
        tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    sessions = [
        ('Asian Session / 亚洲盘', '06:00-15:00', 'Low-Med / 中低', '—'),
        ('European Session / 欧洲盘', '15:00-00:00', 'High / 高', '—'),
        ('US Session / 美洲盘', '20:00-05:00', 'Highest / 最高', '—'),
        ('Overlap (EU+US) / 欧美重叠段', '20:00-00:00', 'Extreme / 极高', '—'),
    ]
    for sess, time, liq, rec in sessions:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(sess).bold = True
        row.cells[1].paragraphs[0].add_run(time)
        row.cells[2].paragraphs[0].add_run(liq)
        row.cells[3].paragraphs[0].add_run(rec)
    
    # ==================== 十、Risk Management / 风险管理 ====================
    doc.add_heading('十、Risk Management / 风险管理 ⚠️', level=2)
    
    doc.add_heading('10.1 Forex-Specific Risks / 外汇特殊风险', level=3)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Light Grid Accent 1'
    hdr = ['Risk Type / 风险类型', 'Description / 说明', 'Mitigation / 应对措施']
    for i, h in enumerate(hdr):
        tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    risks = [
        ('Central Bank Surprise / 央行意外政策', 'Rate decision exceeds expectations', 'Set calendar alerts / 设事件预警'),
        ('Gap Risk / 汇率跳空', 'NFP/CPI release volatility', 'Reduce positions before data / 数据前减仓'),
        ('Leverage Risk / 杠杆风险', 'High leverage amplifies losses', 'Keep ≤10:1 leverage / 控制杠杆'),
        ('Liquidity Risk / 流动性枯竭', 'Holiday / thin session', 'Avoid low-liquidity hours'),
        ('Carry Trade Unwind / 套利平仓', 'Yen/CHF funding panic', 'Monitor VIX & risk sentiment'),
    ]
    for risk, desc, action in risks:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(f'⚠️ {risk}').bold = True
        row.cells[1].paragraphs[0].add_run(desc)
        row.cells[2].paragraphs[0].add_run(action)
    
    doc.add_heading('10.2 Money Management / 资金管理', level=3)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Light Grid Accent 1'
    cells = tbl.rows[0].cells
    cells[0].paragraphs[0].add_run(f'Item / {_bt.item_label}').bold = True
    cells[1].paragraphs[0].add_run('Suggestion / 建议').bold = True
    for item in ['Account Size / 账户总资金', 'Max Risk per Trade / 单笔最大风险', 'Suggested Leverage / 建议杠杆', 'Max Drawdown / 最大回撤容忍', 'Daily Loss Limit / 日亏损上限']:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(item).bold = True
        row.cells[1].paragraphs[0].add_run('—')
    
    # ==================== 十一、Weekly Events / 本周关注 ====================
    doc.add_heading('十一、Weekly Events Calendar / 本周关注事件 📅', level=2)
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = 'Light Grid Accent 1'
    hdr = ['Time / 时间', 'Country / 国家', 'Event / 事件', 'Importance / 重要程度', 'Expected Impact / 预期影响']
    for i, h in enumerate(hdr):
        tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    for _ in range(5):
        row = tbl.add_row()
        for i in range(5):
            row.cells[i].paragraphs[0].add_run('—')
    
    # ==================== 十二、Summary / 方案总结 ====================
    doc.add_heading('十二、Summary / 方案总结 ✅', level=2)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Light Grid Accent 1'
    cells = tbl.rows[0].cells
    cells[0].paragraphs[0].add_run('Dimension / 维度').bold = True
    cells[1].paragraphs[0].add_run('Conclusion / 结论').bold = True
    for dim in ['Macro View / 大势判断', 'Core Logic / 核心逻辑', 'Trading Strategy / 操作思路', 'Position / 仓位建议', 'Stop-Loss Discipline / 止损纪律']:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(dim).bold = True
        row.cells[1].paragraphs[0].add_run('—')
    
    # ==================== One-Liner / 一句话总结 ====================
    doc.add_heading('📌 One-Sentence Summary / 一句话总结', level=3)
    p = doc.add_paragraph('—')
    run = p.runs[0]
    run.bold = True
    
    # ==================== Appendix / 附录 ====================
    doc.add_heading('Appendix: Forex Basics / 附录：外汇基础须知', level=2)
    
    doc.add_heading('Major Currency Pairs / 主要货币对分类', level=3)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Light Grid Accent 1'
    hdr = ['Category / 类别', 'Pairs / 货币对', 'Features / 特点']
    for i, h in enumerate(hdr):
        tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    pairs = [
        ('Major Pairs / 主要货币对', 'EUR/USD, USD/JPY, GBP/USD, USD/CHF', 'Highest liquidity, lowest spread'),
        ('Cross Pairs / 交叉盘', 'EUR/GBP, EUR/JPY, GBP/JPY, EUR/CHF', 'No USD involved / 无美元'),
        ('Commodity Currencies / 商品货币', 'AUD/USD, NZD/USD, USD/CAD', 'Tied to commodity prices'),
        ('EM Pairs / 新兴市场', 'USD/TRY, USD/ZAR, USD/MXN, USD/BRL', 'High volatility, high carry'),
    ]
    for cat, pairs_text, note in pairs:
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(cat).bold = True
        row.cells[1].paragraphs[0].add_run(pairs_text)
        row.cells[2].paragraphs[0].add_run(note)
    
    doc.add_heading('Trading Hours (Beijing Time) / 交易时间（北京时间）', level=3)
    p = doc.add_paragraph('📅 Forex 24H market, main sessions / 外汇24小时交易，主要时段：')
    doc.add_paragraph('  06:00-15:00  Asian / 亚洲盘（Tokyo/Singapore/HK）')
    doc.add_paragraph('  15:00-00:00  European / 欧洲盘（London/Frankfurt）')
    doc.add_paragraph('  20:00-05:00  US / 美洲盘（New York）')
    doc.add_paragraph('  20:00-00:00  Overlap / 欧美重叠段（highest liquidity / 流动性最高）')
    
    doc.add_heading('Risk Disclaimer / 风险声明', level=3)
    p = doc.add_paragraph()
    run = p.add_run(
        'This plan is for reference and educational purposes only, not investment advice. '
        'Forex margin trading carries significant risk; high leverage can amplify both gains and losses. '
        'Trade according to your own risk tolerance and strictly follow risk management. '
        '本方案仅供参考学习，不构成投资建议。外汇保证金交易风险巨大，请根据自身情况谨慎决策。'
    )
    run.font.size = Pt(8)
    run.italic = True
    
    # ==================== FOOTER / 尾部 ====================
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('—— 媳妇智投全球通用版 | Ai-Wife Global Investing ——')
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Date / 日期：{date_str} | Version / 版本：v7.0 FINAL')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('⚠️ AIGC Content — For Reference Only / AIGC生成内容，仅供参考')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    
    return doc


def main():
    import sys
    
    symbol = "EUR/USD"
    if len(sys.argv) >= 2:
        symbol = sys.argv[1]
    
    date_str = datetime.datetime.now().strftime("%Y-%m-%d")
    
    doc = create_forex_plan(symbol, date_str)
    
    safe_symbol = symbol.replace('/', '_')
    filename = f"Forex_FX_{safe_symbol}_{date_str}_v7.0.docx"
    
    doc.save(filename)
    print(f"✅ Bilingual Forex Plan Generated / 双语外汇方案生成完成！")
    print(f"   📄 {filename}")
    print(f"   💱 Pair / 品种: {symbol}")
    print(f"   📅 Date / 日期: {date_str}")
    print(f"   📊 Chapters / 章节: 12 (Bilingual CN/EN / 中英双语)")
    print(f"   🏷️  Version / 版本: v7.0 FINAL")


if __name__ == "__main__":
    main()
