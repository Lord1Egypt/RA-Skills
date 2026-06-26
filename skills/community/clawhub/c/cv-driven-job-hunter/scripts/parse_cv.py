#!/usr/bin/env python3
"""
parse_cv.py — Extrae texto de un CV (PDF/DOCX/TXT/MD) y produce un perfil
estructurado en JSON con secciones, skills, años de experiencia, idiomas, etc.

Uso:
    python3 parse_cv.py path/to/cv.pdf --out path/to/profile.json
    python3 parse_cv.py path/to/cv.docx --out profile.json --raw raw.txt
    python3 parse_cv.py cv.pdf --out profile.json --config data/config.json

El --config es opcional y se usa para overridear datos que la heurística
no logra detectar bien (especialmente niveles de idioma).

Si falta una dependencia opcional (pdfminer.six, python-docx), el script
imprime cómo instalarla y termina con código != 0 — no auto-instala.
"""

import argparse
import json
import re
import sys
from datetime import datetime, timezone
from pathlib import Path

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")


SECTION_PATTERNS = {
    "summary": [r"^\s*#{0,6}\s*(summary|profile|about|resumen|perfil)\s*:?\s*$"],
    "experience": [r"^\s*#{0,6}\s*(experience|work\s+experience|employment|experiencia(\s+laboral)?|trabajo|professional\s+experience)\s*:?\s*$"],
    "education": [r"^\s*#{0,6}\s*(education|educaci[oó]n|estudios|formaci[oó]n)\s*:?\s*$"],
    "skills": [r"^\s*#{0,6}\s*(skills|technical\s+skills|tecnolog[ií]as|habilidades|stack|competencias)\s*:?\s*$"],
    "languages": [r"^\s*#{0,6}\s*(languages|idiomas|lenguajes)\s*:?\s*$"],
    "projects": [r"^\s*#{0,6}\s*(projects|proyectos|side\s+projects)\s*:?\s*$"],
    "certifications": [r"^\s*#{0,6}\s*(certifications?|certificaciones?|certificados?)\s*:?\s*$"],
    "publications": [r"^\s*#{0,6}\s*(publications?|publicaciones?|papers?|charlas?|talks?)\s*:?\s*$"],
    "contact": [r"^\s*#{0,6}\s*(contact|contacto|info|datos\s+personales)\s*:?\s*$"],
}

SKILL_KEYWORDS = {
    "languages": [
        "python", "javascript", "typescript", "go", "golang", "rust", "java",
        "kotlin", "swift", "ruby", "php", "c#", "c\\+\\+", "scala", "elixir",
        "clojure", "haskell", "ocaml", "lua", "dart", "r ", "matlab", "perl",
        "bash", "shell", "powershell", "groovy",
    ],
    "frontend": [
        "react", "vue", "angular", "svelte", "next\\.js", "nextjs", "nuxt",
        "remix", "astro", "tailwind", "redux", "zustand", "rtk", "vite",
        "webpack", "html", "css", "sass", "scss", "graphql", "apollo",
    ],
    "backend": [
        "node\\.?js", "express", "fastify", "nest\\.?js", "django", "flask",
        "fastapi", "rails", "spring", "spring boot", "laravel", "phoenix",
        "actix", "axum", "gin", "echo", "fiber", "grpc", "rest api",
        "graphql api", "soap",
    ],
    "data": [
        "postgres", "postgresql", "mysql", "mariadb", "sqlite", "mongodb",
        "redis", "elasticsearch", "opensearch", "clickhouse", "snowflake",
        "bigquery", "redshift", "duckdb", "dynamodb", "cassandra", "kafka",
        "rabbitmq", "pulsar", "kinesis", "spark", "pyspark", "databricks",
        "airflow", "dbt", "etl", "elt",
    ],
    "cloud_devops": [
        "aws", "gcp", "google cloud", "azure", "kubernetes", "k8s", "docker",
        "openshift", "terraform", "pulumi", "ansible", "helm", "argo", "argocd",
        "jenkins", "github actions", "gitlab ci", "circleci", "prometheus",
        "grafana", "datadog", "sentry", "opentelemetry", "istio", "linkerd",
        "service mesh", "ci/cd", "devops", "sre", "site reliability",
        "observability", "gitops", "infrastructure as code", "iac",
    ],
    "ml_ai": [
        "pytorch", "tensorflow", "scikit-learn", "sklearn", "xgboost",
        "lightgbm", "huggingface", "hugging face", "langchain", "llamaindex",
        "openai", "anthropic", "claude", "gpt", "gpt-4", "gpt-5", "gemini",
        "llm", "llms", "large language model", "large language models",
        "agentic", "agentic ai", "agents", "ai agent", "ai agents",
        "agent architecture", "agent architectures", "multi-agent",
        "mcp", "model context protocol",
        "rag", "retrieval augmented", "retrieval-augmented",
        "vector db", "vector database", "vector search", "semantic search",
        "embeddings", "fine-tuning", "fine tuning", "prompt engineering",
        "pinecone", "weaviate", "qdrant", "chroma", "milvus",
        "agéntico", "agentico", "agéntica", "agentica",
        "agéntic", "arquitecturas agénticas", "arquitecturas agenticas",
        "agentes", "ingeniería de prompts", "ingenieria de prompts",
        "modelos de lenguaje", "ajuste fino",
        "inteligencia artificial generativa", "ia generativa",
    ],
    "architecture": [
        "microservices", "monolith", "event-driven", "event driven",
        "domain-driven design", "ddd", "hexagonal", "clean architecture",
        "enterprise architecture", "solution architecture",
        "solutions architecture", "api strategy", "api design",
        "api management", "api gateway", "system design",
        "distributed systems", "cqrs", "event sourcing", "saga pattern",
        "service-oriented", "soa",
        "arquitectura empresarial", "arquitectura de soluciones",
        "arquitectura de software", "arquitectura cloud",
        "arquitectura híbrida", "arquitectura hibrida",
        "microservicios", "estrategia api", "estrategia de api",
        "diseño de api", "diseno de api", "gestión de api", "gestion de api",
        "diseño de sistemas", "diseno de sistemas",
        "sistemas distribuidos", "diseño dirigido por dominio",
    ],
    "governance": [
        "ai governance", "responsible ai", "model governance",
        "ml governance", "data governance", "data privacy", "gdpr",
        "hipaa", "soc 2", "soc2", "iso 27001", "iso27001",
        "risk management", "compliance",
        "gobierno de ia", "gobernanza", "gobernanza de ia",
        "gobierno de datos", "gobernanza de datos",
        "ia responsable", "privacidad de datos", "cumplimiento normativo",
        "gestión de riesgos", "gestion de riesgos",
    ],
    "leadership": [
        "technical leadership", "tech lead", "engineering management",
        "team lead", "mentoring", "mentorship", "hiring", "interviewing",
        "team building", "cross-functional", "stakeholder management",
        "liderazgo técnico", "liderazgo tecnico",
        "líder técnico", "lider tecnico",
        "mentoría", "mentoria", "mentorizar",
        "gestión de equipos", "gestion de equipos",
        "contratación", "contratacion", "entrevistas técnicas",
    ],
    "testing": [
        "pytest", "jest", "vitest", "cypress", "playwright", "selenium",
        "junit", "rspec", "tdd", "bdd", "unit testing", "integration testing",
        "e2e testing", "load testing",
    ],
}

LANGUAGE_LEVELS = {
    "native": [
        "native", "nativo", "nativa", "lengua materna", "mother tongue",
        "native speaker", "hablante nativo", "first language",
        "native or bilingual", "bilingual proficiency",
    ],
    "c2": [
        "c2", "proficient", "fluent", "fluido", "fluida",
        "bilingüe", "bilingual",
        "full professional proficiency", "full professional",
        "mastery", "dominio",
    ],
    "c1": [
        "c1", "advanced", "avanzado", "avanzada",
        "professional working proficiency", "professional working",
        "high proficiency",
    ],
    "b2": [
        "b2", "upper intermediate", "intermedio alto", "upper-intermediate",
        "competent",
    ],
    "b1": [
        "b1", "intermediate", "intermedio",
        "limited working proficiency", "limited working",
        "conversational", "conversacional", "working knowledge",
    ],
    "a2": [
        "a2", "elementary", "básico", "basico",
        "elementary proficiency", "basic",
    ],
    "a1": [
        "a1", "beginner", "principiante", "novice",
    ],
}

KNOWN_LANGUAGES = {
    "english": ["english", "inglés", "ingles"],
    "spanish": ["spanish", "español", "espanol", "castellano"],
    "portuguese": ["portuguese", "portugués", "portugues"],
    "french": ["french", "francés", "frances"],
    "german": ["german", "alemán", "aleman"],
    "italian": ["italian", "italiano"],
    "mandarin": ["mandarin", "chinese", "chino"],
    "japanese": ["japanese", "japonés", "japones"],
}

MONTHS = {
    "january": 1, "february": 2, "march": 3, "april": 4, "may": 5,
    "june": 6, "july": 7, "august": 8, "september": 9, "october": 10,
    "november": 11, "december": 12,
    "jan": 1, "feb": 2, "mar": 3, "apr": 4,
    "jun": 6, "jul": 7, "aug": 8, "sep": 9, "sept": 9, "oct": 10, "nov": 11, "dec": 12,
    "enero": 1, "febrero": 2, "marzo": 3, "abril": 4, "mayo": 5,
    "junio": 6, "julio": 7, "agosto": 8, "septiembre": 9, "setiembre": 9,
    "octubre": 10, "noviembre": 11, "diciembre": 12,
    "ene": 1, "abr": 4, "ago": 8, "dic": 12,
}

PRESENT_WORDS = ["present", "actualidad", "current", "hoy", "now", "today", "presente", "actual"]


def extract_text_pdf(path: Path) -> str:
    try:
        from pdfminer.high_level import extract_text
    except ImportError:
        sys.exit(
            "Falta pdfminer.six. Instalá con:\n  pip install -r requirements.txt\n"
        )
    return extract_text(str(path))


def extract_text_docx(path: Path) -> str:
    """
    Extrae texto de un .docx caminando todo el XML del body — captura párrafos,
    tablas (incluso anidadas), text boxes, contenido de SDT, headers y footers.
    Más robusto que solo iterar doc.paragraphs + doc.tables.
    """
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError:
        sys.exit(
            "Falta python-docx. Instalá con:\n  pip install -r requirements.txt\n"
        )
    doc = Document(str(path))

    parts: list[str] = []

    def walk(elem) -> None:
        tag = elem.tag
        if tag == qn("w:p"):
            for child in elem:
                walk(child)
            parts.append("\n")
            return
        if tag == qn("w:tab"):
            parts.append("\t")
            return
        if tag == qn("w:br") or tag == qn("w:cr"):
            parts.append("\n")
            return
        if tag == qn("w:t"):
            if elem.text:
                parts.append(elem.text)
            return
        for child in elem:
            walk(child)

    walk(doc.element.body)

    for section in doc.sections:
        for hdr_ftr in (section.header, section.footer):
            if hdr_ftr is None:
                continue
            try:
                walk(hdr_ftr._element)
            except Exception:
                pass

    text = "".join(parts)

    if not text.strip():
        para_parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    para_parts.append(cell.text)
        text = "\n".join(para_parts)

    return text


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_text_pdf(path)
    if suffix == ".docx":
        return extract_text_docx(path)
    if suffix == ".doc":
        sys.exit(
            "Formato .doc (Word legacy) no soportado. Convertí a .docx, .pdf o .txt."
        )
    if suffix in (".txt", ".md", ".markdown", ".rst"):
        return path.read_text(encoding="utf-8", errors="replace")
    sys.exit(f"Formato no soportado: {suffix}. Usá PDF, DOCX, TXT o MD.")


def split_sections(text: str) -> dict:
    lines = text.splitlines()
    sections: dict = {"_preamble": []}
    current = "_preamble"
    for raw in lines:
        line = raw.strip()
        if not line:
            sections.setdefault(current, []).append("")
            continue
        matched = None
        for name, patterns in SECTION_PATTERNS.items():
            for pat in patterns:
                if re.match(pat, line, re.IGNORECASE):
                    matched = name
                    break
            if matched:
                break
        if matched:
            current = matched
            sections.setdefault(current, [])
        else:
            sections.setdefault(current, []).append(raw)
    return {k: "\n".join(v).strip() for k, v in sections.items() if "\n".join(v).strip()}


def detect_skills(text: str) -> dict:
    text_lower = text.lower()
    detected: dict = {}
    for category, kws in SKILL_KEYWORDS.items():
        hits = []
        for kw in kws:
            pattern = r"(?<![\w.])" + kw + r"(?![\w])"
            if re.search(pattern, text_lower):
                normalized = kw.replace("\\.", ".").replace("\\+", "+").strip()
                hits.append(normalized)
        if hits:
            detected[category] = sorted(set(hits))
    return detected


def _detect_level_in_window(window: str) -> str | None:
    window_l = window.lower()
    for lvl, indicators in LANGUAGE_LEVELS.items():
        sorted_inds = sorted(indicators, key=len, reverse=True)
        for ind in sorted_inds:
            if re.search(r"\b" + re.escape(ind) + r"\b", window_l):
                return lvl
    return None


def detect_languages(text: str, override: dict | None = None) -> dict:
    """
    Detecta idiomas y nivel. Si hay sección 'languages'/'idiomas', escanea
    por línea (el nivel del idioma N solo cuenta si está en su misma línea,
    no contamina al N+1). Si no hay sección, escanea el texto completo con
    ventana acotada (±60 chars).
    Si el nivel queda 'unspecified' o no se detecta, aplica el override del
    config (si existe).
    """
    found: dict = {}
    sections = split_sections(text)
    lang_section = sections.get("languages", "")

    if lang_section:
        for raw_line in lang_section.splitlines():
            line = raw_line.strip()
            if not line:
                continue
            line_l = line.lower()
            for lang_key, names in KNOWN_LANGUAGES.items():
                if found.get(lang_key) and found[lang_key] != "unspecified":
                    continue
                if any(re.search(r"\b" + re.escape(n) + r"\b", line_l) for n in names):
                    lvl = _detect_level_in_window(line_l)
                    found[lang_key] = lvl or "unspecified"

    for lang_key, names in KNOWN_LANGUAGES.items():
        if lang_key in found:
            continue
        for name in names:
            matched = False
            for m in re.finditer(r"\b" + re.escape(name) + r"\b", text, re.IGNORECASE):
                start, end = m.span()
                w_start = max(0, start - 60)
                w_end = min(len(text), end + 60)
                window = text[w_start:w_end]
                lvl = _detect_level_in_window(window)
                if lvl:
                    found[lang_key] = lvl
                    matched = True
                    break
                if lang_key not in found:
                    found[lang_key] = "unspecified"
                    matched = True
            if matched and found.get(lang_key) and found[lang_key] != "unspecified":
                break

    if override:
        for lk, lv in override.items():
            if found.get(lk) in (None, "unspecified"):
                found[lk] = lv
    return found


def parse_languages_override(raw) -> dict:
    """
    Acepta dos formas:
        {"english": "c1", "spanish": "native"}
        ["en-c1", "es-native"]
    Devuelve dict normalizado por lang_key (english/spanish/...).
    """
    if not raw:
        return {}
    short_to_key = {
        "en": "english", "es": "spanish", "pt": "portuguese", "fr": "french",
        "de": "german", "it": "italian", "zh": "mandarin", "ja": "japanese",
    }
    if isinstance(raw, dict):
        out = {}
        for k, v in raw.items():
            kl = k.lower().strip()
            if kl in KNOWN_LANGUAGES:
                out[kl] = str(v).lower().strip()
            elif kl in short_to_key:
                out[short_to_key[kl]] = str(v).lower().strip()
        return out
    if isinstance(raw, list):
        out = {}
        for item in raw:
            if not isinstance(item, str) or "-" not in item:
                continue
            short, lvl = item.split("-", 1)
            short = short.lower().strip()
            if short in short_to_key:
                out[short_to_key[short]] = lvl.lower().strip()
            elif short in KNOWN_LANGUAGES:
                out[short] = lvl.lower().strip()
        return out
    return {}


_MONTH_ALT = "|".join(sorted(MONTHS.keys(), key=len, reverse=True))
_PRESENT_ALT = "|".join(PRESENT_WORDS)
_YEAR = r"(?:19|20)\d{2}"

DATE_RANGE_RE = re.compile(
    rf"""
    (
        (?:[A-Za-zñéí]+\.?\s+{_YEAR})     # "Aug 2025", "agosto 2025", "Ago. 2025"
      | (?:\d{{1,2}}[/\-.]{_YEAR})        # "08/2025", "8-2025"
      | (?:{_YEAR}[/\-.]\d{{1,2}})        # "2025-08"
      | {_YEAR}                            # "2025"
    )
    \s*(?:[-–—‐]|\bto\b|\bhasta\b|\bal?\b|→)\s*
    (
        (?:[A-Za-zñéí]+\.?\s+{_YEAR})
      | (?:\d{{1,2}}[/\-.]{_YEAR})
      | (?:{_YEAR}[/\-.]\d{{1,2}})
      | {_YEAR}
      | {_PRESENT_ALT}
    )
    """,
    re.IGNORECASE | re.VERBOSE,
)


def _parse_date_token(s: str, end_default_month: int) -> tuple[int, int] | None:
    """
    Parsea un date token y devuelve (year, month). end_default_month decide
    qué mes usar cuando solo se especifica año (1 para start, 12 para end).
    """
    s = s.strip().rstrip(".").lower()
    if s in PRESENT_WORDS:
        now = datetime.now()
        return now.year, now.month
    m = re.match(rf"^([a-zñéí]+)\.?[\s,.\-/]+({_YEAR})$", s)
    if m and m.group(1) in MONTHS:
        return int(m.group(2)), MONTHS[m.group(1)]
    m = re.match(rf"^(\d{{1,2}})[/\-.]({_YEAR})$", s)
    if m:
        month, year = int(m.group(1)), int(m.group(2))
        if 1 <= month <= 12:
            return year, month
    m = re.match(rf"^({_YEAR})[/\-.](\d{{1,2}})$", s)
    if m:
        year, month = int(m.group(1)), int(m.group(2))
        if 1 <= month <= 12:
            return year, month
    m = re.match(rf"^({_YEAR})$", s)
    if m:
        return int(m.group(1)), end_default_month
    return None


def _scan_intervals(scan_text: str, horizon_year: int) -> list[tuple[int, int]]:
    intervals: list[tuple[int, int]] = []
    for m in DATE_RANGE_RE.finditer(scan_text):
        start = _parse_date_token(m.group(1), end_default_month=1)
        end = _parse_date_token(m.group(2), end_default_month=12)
        if not start or not end:
            continue
        s_y, s_m = start
        e_y, e_m = end
        if s_y > horizon_year or e_y > horizon_year:
            continue
        s_idx = s_y * 12 + (s_m - 1)
        e_idx = e_y * 12 + (e_m - 1)
        if s_idx > e_idx:
            continue
        intervals.append((s_idx, e_idx))
    return intervals


def estimate_years_experience(text: str) -> float:
    """
    Suma rangos de fechas SOLO de la sección de experiencia. Calcula en meses
    para precisión, mergea overlaps, y devuelve años como float.
    Soporta: 'Aug 2025 - Present', 'Jan 2023 - July 2025', '01/2023 - 07/2025',
    '2023 - 2025', y mezclas. Meses en inglés y español, full y abreviados,
    con o sin punto ('Ago.', 'Ene.').
    Fallback: si la sección 'experience' está truncada (típico en PDFs
    multi-columna donde pdfminer linealiza mal el orden de lectura) y el
    resto del texto tiene muchas más fechas, escanea el texto completo
    excluyendo la sección 'education'.
    """
    sections = split_sections(text)
    exp = sections.get("experience", "")
    edu = sections.get("education", "")
    now = datetime.now()
    horizon_year = now.year + 1

    text_minus_edu = text.replace(edu, "") if edu else text

    if exp:
        exp_intervals = _scan_intervals(exp, horizon_year)
        all_intervals = _scan_intervals(text_minus_edu, horizon_year)
        if len(all_intervals) >= 2 * max(1, len(exp_intervals)) and len(all_intervals) >= 3:
            intervals = all_intervals
        else:
            intervals = exp_intervals
    else:
        intervals = _scan_intervals(text_minus_edu, horizon_year)

    if not intervals:
        return 0.0

    intervals.sort()
    merged: list[list[int]] = [list(intervals[0])]
    for s, e in intervals[1:]:
        if s <= merged[-1][1] + 1:
            merged[-1][1] = max(merged[-1][1], e)
        else:
            merged.append([s, e])
    total_months = sum((e - s + 1) for s, e in merged)
    return round(total_months / 12.0, 1)


EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}")
URL_RE = re.compile(r"https?://[^\s)]+")
PHONE_RE = re.compile(r"(\+?\d[\d\s().-]{7,}\d)")


def _looks_like_phone(candidate: str) -> bool:
    """Filtra rangos de fechas y otros falsos positivos del PHONE_RE."""
    digit_count = sum(1 for c in candidate if c.isdigit())
    if digit_count < 8:
        return False
    if "/" in candidate:
        return False
    digits_only = "".join(c for c in candidate if c.isdigit())
    if re.match(r"^(19|20)\d{2}(19|20)\d{2}", digits_only):
        return False
    return True


def detect_contact(text: str) -> dict:
    head = "\n".join(text.splitlines()[:30])
    phones = []
    for p in PHONE_RE.findall(head):
        cleaned = p.strip()
        if _looks_like_phone(cleaned):
            phones.append(cleaned)
    return {
        "emails": list(set(EMAIL_RE.findall(head))),
        "urls": list(set(URL_RE.findall(head))),
        "phones": list(set(phones))[:3],
    }


TITLE_NOUNS = {
    "engineer", "developer", "architect", "scientist", "analyst", "consultant",
    "specialist", "designer", "manager", "lead", "head", "director", "advisor",
    "officer", "researcher", "owner", "principal", "founder", "co-founder",
    "cofounder", "cto", "cio", "ciso", "vp", "evp", "svp", "swe", "sde",
    "programmer", "administrator", "strategist", "evangelist",
    "ingeniero", "ingeniera", "desarrollador", "desarrolladora",
    "arquitecto", "arquitecta", "científico", "cientifico",
    "analista", "consultor", "consultora", "especialista",
    "diseñador", "disenador", "diseñadora", "gerente",
    "líder", "lider", "jefe", "jefa", "directora",
    "asesor", "asesora", "investigador", "investigadora",
    "fundador", "fundadora", "responsable",
}

TITLE_DISQUALIFIERS = {
    "for", "with", "to", "the", "of", "and", "as", "in", "on", "at",
    "from", "by", "into", "during", "across", "throughout", "their", "our",
    "para", "con", "del", "los", "las", "una", "uno", "que",
    "initiatives", "projects", "teams", "iniciativas", "proyectos",
    "equipos", "definición", "definicion", "soluciones",
}


def detect_titles(text: str) -> list:
    """
    Detecta títulos de roles a nivel de línea. Requiere que la línea termine
    en un sustantivo de rol conocido (engineer, architect, lead, manager, ...)
    y que no contenga conectores típicos de oración ('for', 'the', 'as', etc.).
    Esto evita capturar fragmentos como 'lead architecture initiatives for the'.
    """
    titles: set[str] = set()
    line_re = re.compile(r"^[A-Za-zÁÉÍÓÚÑáéíóúñ][\w/+&.\- ]{1,68}$")
    for raw in text.splitlines():
        line = raw.strip().rstrip(".,;:")
        if not line or len(line) > 70:
            continue
        if not line_re.match(line):
            continue
        words = line.split()
        if not (1 <= len(words) <= 6):
            continue
        words_l = [w.lower().rstrip(".,;:") for w in words]
        last = words_l[-1]
        if last not in TITLE_NOUNS and not any(w in TITLE_NOUNS for w in words_l):
            continue
        if any(w in TITLE_DISQUALIFIERS for w in words_l):
            continue
        if sum(1 for w in words if w[0].islower() and not w.lower() in {"of", "and", "the"}) > len(words) // 2:
            continue
        titles.add(line)
    return sorted(titles)[:15]


def main() -> int:
    parser = argparse.ArgumentParser(description="Parse CV → structured profile JSON")
    parser.add_argument("cv_path", type=Path, help="Path al CV (PDF/DOCX/TXT/MD)")
    parser.add_argument("--out", type=Path, required=True, help="Path al JSON de salida")
    parser.add_argument("--raw", type=Path, help="Path opcional para guardar el texto crudo")
    parser.add_argument("--config", type=Path, help="config.json — usado para overrides (idiomas)")
    args = parser.parse_args()

    if not args.cv_path.exists():
        sys.exit(f"No existe el CV: {args.cv_path}")

    text = extract_text(args.cv_path)
    if not text.strip():
        sys.exit("El CV se parseó pero está vacío. Revisá el archivo.")

    if args.raw:
        args.raw.parent.mkdir(parents=True, exist_ok=True)
        args.raw.write_text(text, encoding="utf-8")

    lang_override = {}
    if args.config and args.config.exists():
        try:
            cfg = json.loads(args.config.read_text(encoding="utf-8"))
            lang_override = parse_languages_override(
                (cfg.get("candidate") or {}).get("languages")
            )
        except (json.JSONDecodeError, OSError):
            pass

    sections = split_sections(text)
    profile = {
        "source_file": str(args.cv_path),
        "parsed_at": datetime.now(timezone.utc).isoformat(timespec="seconds"),
        "raw_chars": len(text),
        "contact": detect_contact(text),
        "sections": {k: v for k, v in sections.items() if k != "_preamble"},
        "preamble": sections.get("_preamble", "")[:1000],
        "skills": detect_skills(text),
        "languages": detect_languages(text, override=lang_override),
        "estimated_years_experience": estimate_years_experience(text),
        "detected_titles": detect_titles(text),
    }

    args.out.parent.mkdir(parents=True, exist_ok=True)
    args.out.write_text(json.dumps(profile, indent=2, ensure_ascii=False), encoding="utf-8")

    print(f"OK — perfil escrito en {args.out}")
    print(f"   chars: {profile['raw_chars']}")
    print(f"   secciones: {list(profile['sections'].keys())}")
    print(f"   skills: {sum(len(v) for v in profile['skills'].values())} keywords")
    print(f"   idiomas: {profile['languages']}")
    print(f"   años de experiencia estimados: {profile['estimated_years_experience']}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
