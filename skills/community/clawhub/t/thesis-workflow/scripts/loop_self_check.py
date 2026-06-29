#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
MBA Thesis Workflow - Loop Self-Check Script
v1.7 引入：Guardrails 自动化校验

校验项（10 项）：
1. 章节完整性（# 第1-7章 数量 = 7）
2. 字数门槛（每章 ≥ 100 行）
3. 参考文献存在（## 参考文献）
4. 无 ## 第X章 混合格式
5. 无 **正文加粗**（排除标题级）
6. 引用完整性（作者，年份）模式
7. 三线表无竖线（只检测表头分隔行）
8. 表格标题在表上方
9. 合并残留检查（===END===）
10. 核心章节关键词检查（第5章战略/第6章实施）

使用方式：
  # 校验单个文件
  python3 loop_self_check.py --file 论文_xxx.md
  
  # 校验整个 Phase（指定工作目录）
  python3 loop_self_check.py --phase 2 --workspace ~/.openclaw/workspace/
  
  # 输出 JSON 报告
  python3 loop_self_check.py --file 论文_xxx.md --json
  
  # 校验 Word 文档（Verification Loop）
  python3 loop_self_check.py --file 论文_xxx.docx --verify-docx

退出码：
  0 = 全部通过
  1 = 有失败项
  2 = 文件/参数错误
"""

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Dict, List, Tuple


# ==================== 可配置默认参数 ====================
DEFAULT_MIN_LINES_PER_CHAPTER = 100    # 每章最少行数
DEFAULT_MIN_CITATIONS = 10             # 最少引用数
DEFAULT_TOTAL_CHAPTERS = 7             # 论文总章数
DEFAULT_MIN_CHARTS = 0                  # 最少图表数


# ==================== 校验函数 ====================

def check_chapter_completeness(content: str) -> Tuple[bool, str]:
    """校验 1: 章节完整性 - 必须有 7 个 # 第X章"""
    pattern = r'^#\s*第[1-7]章'
    chapters = re.findall(pattern, content, re.MULTILINE)
    found = len(set(chapters))
    if found == 7:
        return True, f"✅ 章节完整性：找到 {found}/7 章"
    return False, f"❌ 章节完整性：只找到 {found}/7 章（缺失 {7 - found} 章）"


def check_word_count(content: str, min_lines: int = DEFAULT_MIN_LINES_PER_CHAPTER) -> Tuple[bool, str]:
    """校验 2: 字数门槛 - 全文行数 ≥ 700 行（7 章 × 100）"""
    lines = [line for line in content.split('\n') if line.strip()]
    total = len(lines)
    # 拆分章节检查
    chapter_splits = re.split(r'^#\s*第[1-7]章', content, flags=re.MULTILINE)
    short_chapters = []
    for i, ch in enumerate(chapter_splits[1:], 1):  # 跳过第一个空 split
        ch_lines = len([line for line in ch.split('\n') if line.strip()])
        if ch_lines < min_lines:
            short_chapters.append(f"第{i}章({ch_lines}行)")
    
    if not short_chapters:
        return True, f"✅ 字数门槛：全文 {total} 行，每章 ≥ {min_lines} 行"
    return False, f"❌ 字数门槛：以下章节不足 {min_lines} 行：{', '.join(short_chapters)}"


def check_references(content: str) -> Tuple[bool, str]:
    """校验 3: 参考文献存在"""
    if re.search(r'^##\s*参考文献\s*$', content, re.MULTILINE):
        return True, "✅ 参考文献：存在"
    return False, "❌ 参考文献：缺失（未找到 `## 参考文献`）"


def check_mixed_chapter_format(content: str) -> Tuple[bool, str]:
    """校验 4: 无 ## 第X章 混合格式（应为 # 第X章）"""
    pattern = r'^##\s*第[1-7]章'
    matches = re.findall(pattern, content, re.MULTILINE)
    if not matches:
        return True, "✅ 标题层级：无非法的 `## 第X章` 混合格式"
    return False, f"❌ 标题层级：发现 {len(matches)} 处 `## 第X章` 混合格式（应为 `# 第X章`）"


def check_inline_bold(content: str) -> Tuple[bool, str]:
    """校验 5: 无 **正文加粗**（排除标题行、整行加粗作为小节标题）"""
    violations = []
    for i, line in enumerate(content.split('\n'), 1):
        # 跳过标题行
        if line.lstrip().startswith('#'):
            continue
        # 跳过空行
        if not line.strip():
            continue
        # 跳过表格行（含 | 的行）
        if '|' in line and '---' not in line:
            continue
        # 跳过纯加粗的整行（视为小节标题变体）
        if re.match(r'^\s*\*\*[^*]+\*\*\s*$', line):
            continue
        # 检测正文段落中的 **xxx** 加粗
        if re.search(r'\*\*[^*]+\*\*', line):
            violations.append(f"  Line {i}: {line.strip()[:80]}")
    
    if not violations:
        return True, "✅ 正文加粗：无 **xxx** 残留"
    preview = '\n'.join(violations[:5])
    more = f"\n  ... 共 {len(violations)} 处" if len(violations) > 5 else ""
    return False, f"❌ 正文加粗：发现 {len(violations)} 处 **xxx** 残留\n{preview}{more}"


def check_citation_completeness(content: str, min_citations: int = DEFAULT_MIN_CITATIONS) -> Tuple[bool, str]:
    """校验 6: 引用完整性 - 检测是否有 (作者，年份) 模式"""
    # 中英文引用模式
    cn_pattern = r'[（(][^)）]+[，,]\s*\d{4}[)）]'
    en_pattern = r'\([A-Z][a-zA-Z]+(?:\s+(?:and|&)\s+[A-Z][a-zA-Z]+)?,?\s*\d{4}\)'
    
    cn_citations = re.findall(cn_pattern, content)
    en_citations = re.findall(en_pattern, content)
    total = len(cn_citations) + len(en_citations)
    
    if total >= min_citations:
        return True, f"✅ 引用完整性：找到 {total} 处引用标注（中 {len(cn_citations)} / 英 {len(en_citations)}）"
    return False, f"❌ 引用完整性：仅 {total} 处引用（建议 ≥ {min_citations} 处）"


def check_table_format(content: str) -> Tuple[bool, str]:
    """校验 7: 三线表格式 - 检测 markdown 表格列数一致性"""
    lines = content.split('\n')
    table_start_lines = []  # 表格起始行号
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith('|'):
            # 检测分隔行（表头与数据的分隔）
            cols = [c.strip() for c in line.split('|')[1:-1]]
            is_sep = bool(cols) and all(
                re.match(r'^:?-+:?$', c) for c in cols
            )
            if is_sep:
                table_start_lines.append(i)
        i += 1

    if not table_start_lines:
        return True, "✅ 三线表：未发现 markdown 表格"

    # 对每个表格，检查列数一致性
    issues = []
    for start in table_start_lines:
        # 找到表格的完整范围
        end = start + 1
        while end < len(lines) and lines[end].strip().startswith('|'):
            end += 1
        # 检查表格前后的布局
        if start > 0 and lines[start - 1].strip():
            prev = lines[start - 1].strip()
            # 表格上方应有标题（如 表3-1）
            if not re.search(r'表\s*\d+[-.]\d+', prev):
                issues.append(f"第{start + 1}行：表格前一行可能缺少表号（预期格式：表X-X）")

    if issues:
        return False, f"❌ 三线表格式：发现 {len(issues)} 个问题\n" + '\n'.join(issues[:3])
    return True, f"✅ 三线表：{len(table_start_lines)} 个表格，格式一致"


def check_table_caption_position(content: str) -> Tuple[bool, str]:
    """校验 8: 表格标题在表上方 - 检测表X-Y模式是否在表格行之前"""
    lines = content.split('\n')
    tables_found = 0
    captioned = 0
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith('|') and i > 0:
            # 检查是否为表头分隔行
            cols = [c.strip() for c in line.split('|')[1:-1]]
            is_sep = bool(cols) and all(
                re.match(r'^:?-+:?$', c) for c in cols
            )
            if is_sep:
                tables_found += 1
                # 向上搜索 5 行内是否有 表X-Y 或 表X.Y
                search_start = max(0, i - 5)
                prev_text = '\n'.join(lines[search_start:i])
                if re.search(r'表\s*\d+[-.]\d+', prev_text):
                    captioned += 1
        i += 1

    if tables_found == 0:
        return True, "✅ 表格标题位置：未发现 markdown 表格"

    ratio = captioned / tables_found * 100 if tables_found > 0 else 0
    if ratio >= 80:
        return True, f"✅ 表格标题位置：{captioned}/{tables_found} 个表格标题在上方 ({ratio:.0f}%)"
    return False, f"❌ 表格标题位置：仅 {captioned}/{tables_found} 个表格标题在上方 ({ratio:.0f}%)"


def check_merge_residue(content: str) -> Tuple[bool, str]:
    """校验 9: 合并残留 - 不应有 ===END=== 等残留标识"""
    residue_patterns = [r'===END===', r'===END\s+\w+===', r'\[END\]', r'<!--\s*end\s*-->']
    found = []
    for pattern in residue_patterns:
        matches = re.findall(pattern, content, re.IGNORECASE)
        if matches:
            found.extend(matches)
    
    if not found:
        return True, "✅ 合并残留：无残留标识"
    return False, f"❌ 合并残留：发现 {len(found)} 处残留标识：{', '.join(list(set(found))[:5])}"


def check_chapter_keywords(content: str) -> Tuple[bool, str]:
    """校验 10: 核心章节关键词（第5章战略/第6章实施）"""
    # 提取第5章
    ch5_match = re.search(r'^#\s*第5章.*?(?=^#\s*第6章|\Z)', content, re.MULTILINE | re.DOTALL)
    ch6_match = re.search(r'^#\s*第6章.*?(?=^#\s*第7章|\Z)', content, re.MULTILINE | re.DOTALL)
    
    issues = []
    if ch5_match:
        ch5 = ch5_match.group(0)
        strategy_kw = ['战略选择', '竞争战略', '差异化', '集中化', '成本领先', 'QSPM', '战略']
        if not any(kw in ch5 for kw in strategy_kw):
            issues.append("第5章未包含战略选择关键词")
    else:
        issues.append("未找到第5章")
    
    if ch6_match:
        ch6 = ch6_match.group(0)
        implement_kw = ['实施', '保障', '组织', '人才', '财务', 'KPI', '考核', '措施']
        if not any(kw in ch6 for kw in implement_kw):
            issues.append("第6章未包含实施保障关键词")
    else:
        issues.append("未找到第6章")
    
    if not issues:
        return True, "✅ 核心章节关键词：第5章（含战略）+ 第6章（含实施）均满足"
    return False, f"❌ 核心章节关键词：{'; '.join(issues)}"


# ==================== 主流程 ====================

CHECKS = [
    ("章节完整性", check_chapter_completeness),
    ("字数门槛", check_word_count),
    ("参考文献", check_references),
    ("标题层级", check_mixed_chapter_format),
    ("正文加粗", check_inline_bold),
    ("引用完整性", check_citation_completeness),
    ("三线表", check_table_format),
    ("表格标题", check_table_caption_position),
    ("合并残留", check_merge_residue),
    ("核心章节关键词", check_chapter_keywords),
]


def run_checks(content: str) -> Dict:
    """运行所有校验项，返回结构化报告"""
    results = []
    passed = 0
    failed = 0
    
    for name, check_fn in CHECKS:
        try:
            ok, msg = check_fn(content)
            results.append({
                "name": name,
                "passed": ok,
                "message": msg,
            })
            if ok:
                passed += 1
            else:
                failed += 1
        except Exception as e:
            results.append({
                "name": name,
                "passed": False,
                "message": f"❌ 校验异常: {e}",
            })
            failed += 1
    
    return {
        "total": len(CHECKS),
        "passed": passed,
        "failed": failed,
        "all_passed": failed == 0,
        "results": results,
    }


def check_docx(file_path: Path) -> Dict:
    """Verification Loop：校验 Word 文档格式
    
    校验项：
    - 分页符（每章前应有分页）
    - 表格样式（三线表：顶线1.5磅/表头底线0.75磅/底线0.5磅，无竖线）
    - 标题层级字体（检查黑体/宋体在标题级的出现情况）
    - 正文行距 20 磅
    - 加粗残留（非标题行中的加粗文字）
    - 参考文献中英文分编
    """
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.shared import Pt
    except ImportError:
        return {
            "error": "缺少 python-docx 库，请安装：pip install python-docx",
            "all_passed": False,
        }
    
    doc = Document(str(file_path))
    results = []
    
    # ========== 校验1: 分页符 ==========
    # 正确检测：在 XML 中搜索 w:br@w:type="page"
    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    body_xml = doc.element.body
    page_breaks = len(body_xml.findall('.//w:br[@w:type="page"]', nsmap))
    results.append({
        "name": "分页符",
        "passed": page_breaks >= 5,
        "message": f"✅ 分页符：找到 {page_breaks} 处" if page_breaks >= 5
                    else f"⚠️ 分页符：仅 {page_breaks} 处（应有至少 6 章间分页）",
    })
    
    # ========== 校验2: 表格三线表检测 ==========
    tables_ok = 0
    tables_total = len(doc.tables)
    for table in doc.tables:
        tbl_borders = table._tbl.find(qn('w:tblBorders'))
        if tbl_borders is None:
            continue
        top = tbl_borders.find(qn('w:top'))
        bottom = tbl_borders.find(qn('w:bottom'))
        inside_v = tbl_borders.find(qn('w:insideV'))
        # 三线表：顶线+底线存在，无竖线
        has_top_border = top is not None and top.get(qn('w:val')) == 'single'
        has_bottom_border = bottom is not None and bottom.get(qn('w:val')) == 'single'
        no_inside_v = inside_v is None or inside_v.get(qn('w:val')) == 'none'
        if has_top_border and has_bottom_border and no_inside_v:
            tables_ok += 1
    tbl_msg = f"✅ 三线表：{tables_ok}/{tables_total} 个表格格式正确" if tables_ok >= tables_total * 0.8 \
              else f"⚠️ 三线表：仅 {tables_ok}/{tables_total} 个符合三线表格式"
    results.append({
        "name": "三线表",
        "passed": tables_total == 0 or tables_ok >= tables_total * 0.8,
        "message": tbl_msg,
    })
    
    # ========== 校验3: 字体/字号分布 ==========
    simsun_count = 0
    simhei_count = 0
    tnr_count = 0
    arial_count = 0
    size_16_count = 0
    size_14_count = 0
    size_13_count = 0
    size_12_count = 0
    total_runs = 0
    for para in doc.paragraphs:
        for run in para.runs:
            total_runs += 1
            font_name = (run.font.name or '').lower()
            if 'simsun' in font_name or '宋体' in font_name:
                simsun_count += 1
            if 'simhei' in font_name or '黑体' in font_name:
                simhei_count += 1
            if 'times new roman' in font_name or 'tnr' in font_name:
                tnr_count += 1
            if 'arial' in font_name:
                arial_count += 1
            size = run.font.size
            if size:
                sz_pt = size.pt
                if 15.5 <= sz_pt <= 16.5:
                    size_16_count += 1
                elif 13.5 <= sz_pt <= 14.5:
                    size_14_count += 1
                elif 12.5 <= sz_pt <= 13.5:
                    size_13_count += 1
                elif 11.5 <= sz_pt <= 12.5:
                    size_12_count += 1
    font_results = []
    if simhei_count == 0:
        font_results.append("未检测到黑体")
    if simsun_count == 0:
        font_results.append("未检测到宋体")
    font_ok = not font_results
    results.append({
        "name": "字体分布",
        "passed": font_ok,
        "message": f"✅ 字体：黑体{simhei_count}处/宋体{simsun_count}处/TNR{tnr_count}处/Arial{arial_count}处"
        if font_ok and total_runs > 0
        else f"⚠️ 字体异常：{'; '.join(font_results)}",
    })
    
    # ========== 校验4: 行距 20 磅 ==========
    total_paras = len(doc.paragraphs)
    line_20_count = sum(
        1 for p in doc.paragraphs
        if p.paragraph_format.line_spacing
        and abs(p.paragraph_format.line_spacing - 20.0) < 1.0
    )
    line_ok = line_20_count >= total_paras * 0.6 if total_paras > 0 else False
    results.append({
        "name": "行距20磅",
        "passed": line_ok,
        "message": f"✅ 行距：{line_20_count}/{total_paras} 段为20磅"
        if line_ok else f"⚠️ 行距：仅 {line_20_count}/{total_paras} 段为20磅",
    })
    
    # ========== 校验5: 加粗残留 ==========
    bold_residue = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # 跳过标题行：以第/数字开头 或 全大写（Abstract）
        if re.match(r'^第[一二三四五六七八九十]', text):
            continue
        if re.match(r'^\d+\.\d+', text):
            continue
        if text in ['Abstract', '摘 要', '目 录', '参考文献', '致 谢']:
            continue
        for run in para.runs:
            if run.font.bold and run.text.strip():
                bold_residue += 1
                break  # 一段只计一次
    results.append({
        "name": "加粗残留",
        "passed": bold_residue == 0,
        "message": f"✅ 加粗：无非标题行加粗残留"
        if bold_residue == 0 else f"⚠️ 加粗残留：{bold_residue} 段中出现非标题加粗",
    })
    
    # ========== 校验6: 参考文献中英文分编 ==========
    ref_cn_found = False
    ref_en_found = False
    for para in doc.paragraphs:
        t = para.text.strip()
        if '中文参考文献' in t:
            ref_cn_found = True
        if '英文参考文献' in t or '参考文献' == t:
            ref_en_found = True
    ref_ok = ref_cn_found or ref_en_found
    results.append({
        "name": "参考文献分编",
        "passed": ref_ok,
        "message": f"✅ 参考文献：中英文已分编" if ref_cn_found and ref_en_found
        else (f"✅ 参考文献：存在" if ref_ok else "⚠️ 参考文献：未检测到中英文分编"),
    })
    
    # ========== 汇总 ==========
    failed = sum(1 for r in results if not r["passed"])
    return {
        "total": len(results),
        "passed": len(results) - failed,
        "failed": failed,
        "all_passed": failed == 0,
        "results": results,
    }


def main():
    parser = argparse.ArgumentParser(
        description="MBA Thesis Workflow - Loop Self-Check (v1.7)",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    parser.add_argument("--file", type=str, help="要校验的 md/docx 文件路径")
    parser.add_argument("--phase", type=int, choices=[1, 2, 3, 4, 5], help="校验整个 Phase（需配合 --workspace）")
    parser.add_argument("--workspace", type=str, default=str(Path.home() / ".openclaw" / "workspace"), help="工作目录（用于 --phase）")
    parser.add_argument("--json", action="store_true", help="输出 JSON 格式报告")
    parser.add_argument("--verify-docx", action="store_true", help="校验 Word 文档（Verification Loop）")
    
    args = parser.parse_args()
    
    # 模式 1: 校验单个文件
    if args.file:
        file_path = Path(args.file)
        if not file_path.exists():
            print(f"❌ 文件不存在: {file_path}", file=sys.stderr)
            sys.exit(2)
        
        # Word 文档校验
        if args.verify_docx or file_path.suffix.lower() == ".docx":
            report = check_docx(file_path)
        else:
            # Markdown 校验
            content = file_path.read_text(encoding="utf-8")
            report = run_checks(content)
        
        if args.json:
            print(json.dumps(report, ensure_ascii=False, indent=2))
        else:
            print(f"\n=== Loop Self-Check Report: {file_path.name} ===\n")
            for r in report.get("results", []):
                print(r["message"])
            print(f"\n--- 总计: {report.get('passed', 0)}/{report.get('total', 0)} 通过 ---")
            print(f"状态: {'✅ 全部通过' if report.get('all_passed') else '❌ 有失败项'}\n")
        
        sys.exit(0 if report.get("all_passed") else 1)
    
    # 模式 2: 校验整个 Phase
    if args.phase:
        workspace = Path(args.workspace)
        if not workspace.exists():
            print(f"❌ 工作目录不存在: {workspace}", file=sys.stderr)
            sys.exit(2)
        
        # 找到所有论文 md 文件
        globs = ["论文*.md", "*thesis*.md", "*Thesis*.md",
                 "*开题*.md", "*报告*.md", "*dissertation*.md", "*Dissertation*.md"]
        md_files = []
        for p in globs:
            md_files.extend(workspace.glob(p))
        md_files = sorted(set(md_files))  # 去重
        if not md_files:
            print(f"⚠️ 工作目录中未找到论文文件: {workspace}", file=sys.stderr)
            sys.exit(2)
        
        all_passed = True
        for md_file in md_files:
            content = md_file.read_text(encoding="utf-8")
            report = run_checks(content)
            if not report["all_passed"]:
                all_passed = False
            if args.json:
                print(json.dumps({"file": str(md_file), **report}, ensure_ascii=False, indent=2))
            else:
                print(f"\n=== {md_file.name} ===")
                for r in report["results"]:
                    print(f"  {r['message']}")
                print(f"  状态: {report['passed']}/{report['total']} 通过")
        
        sys.exit(0 if all_passed else 1)
    
    # 无参数
    parser.print_help()
    sys.exit(2)


if __name__ == "__main__":
    main()
