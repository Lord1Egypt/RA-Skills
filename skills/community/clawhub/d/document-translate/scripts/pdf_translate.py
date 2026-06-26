#!/usr/bin/env python3
"""
PDF translate script — translates English text in PDF to target language.
Supports two methods:
  --method word: Extract text, build translated DOCX, export to PDF (best quality)
  --method direct: Overlay translated text on original PDF (losses formatting)
Usage: python3 pdf_translate.py <input.pdf> [--output <output.pdf>] [--method word|direct]
"""

import argparse
import subprocess
from pathlib import Path

# Check available libraries
try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def extract_text_pdf(input_path):
    """Extract text from PDF using pdfplumber."""
    if not HAS_PDFPLUMBER:
        raise RuntimeError("pdfplumber not installed: pip install pdfplumber")

    import pdfplumber
    pages = []
    with pdfplumber.open(input_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ''
            pages.append(text)
    return pages


def translate_text_pages(pages, translations):
    """Apply translation mapping to extracted pages."""
    translated = []
    for page_text in pages:
        for eng, chn in translations:
            page_text = page_text.replace(eng, chn)
        translated.append(page_text)
    return translated


def build_docx_from_pages(pages, output_path):
    """Build a DOCX from translated page text."""
    if not HAS_DOCX:
        raise RuntimeError("python-docx not installed: pip install python-docx")

    import docx
    doc = docx.Document()
    for i, page_text in enumerate(pages):
        doc.add_heading(f'Page {i+1}', level=2)
        doc.add_paragraph(page_text)
    doc.save(output_path)
    print(f"DOCX saved: {output_path}")


def pdf_to_docx_method(input_path, output_path, translations):
    """Method A: PDF -> text -> DOCX -> PDF via Word."""
    pages = extract_text_pdf(input_path)
    translated_pages = translate_text_pages(pages, translations)

    # Build intermediate DOCX
    tmp_docx = Path(output_path).with_suffix('.docx')
    build_docx_from_pages(translated_pages, str(tmp_docx))

    # Convert DOCX to PDF (requires LibreOffice)
    try:
        result = subprocess.run(
            ['soffice', '--headless', '--convert-to', 'pdf', '--outdir',
             Path(output_path).parent, str(tmp_docx)],
            capture_output=True, text=True, timeout=120
        )
        if result.returncode == 0:
            converted = Path(output_path).parent / f"{Path(tmp_docx).stem}.pdf"
            if converted.exists() and str(converted) != str(output_path):
                Path(converted).rename(output_path)
            print(f"PDF saved: {output_path}")
        else:
            print(f"LibreOffice conversion failed: {result.stderr}")
            print(f"DOCX intermediate saved at: {tmp_docx}")
    except FileNotFoundError:
        print("LibreOffice not found. DOCX intermediate saved; convert to PDF manually.")
        print(f"Path: {tmp_docx}")


def direct_method(input_path, output_path, translations):
    """Method B: Direct text overlay (preserves layout, modifies PDF in place)."""
    if not HAS_PDFPLUMBER:
        raise RuntimeError("pdfplumber required for direct method")

    import pdfplumber

    with pdfplumber.open(input_path) as pdf:
        for page in pdf.pages:
            words = page.extract_words()
            for word in words:
                text = word.get('text', '')
                for eng, chn in translations:
                    if eng in text:
                        # Note: pdfplumber doesn't support write, use pypdf2 or reportlab
                        pass

    print("Direct method not fully implemented. Use --method word.")
    print(f"Intermediate DOCX approach recommended.")


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Translate PDF files')
    parser.add_argument('input', help='Input PDF file path')
    parser.add_argument('--output', '-o', help='Output PDF file path')
    parser.add_argument('--method', choices=['word', 'direct'], default='word',
                        help='Translation method: word (recommended) or direct')
    args = parser.parse_args()

    if args.output is None:
        stem = Path(args.input).stem
        ext = '.pdf' if args.method == 'direct' else '.pdf'
        args.output = str(Path(args.input).parent / f"{stem}-translated{ext}")

    if args.method == 'word':
        # Load translations (simplified — extend with project-specific terms)
        translations = [
            ("Summary", "摘要"),
            ("Introduction", "简介"),
            ("Results", "结果"),
            ("Conclusion", "结论"),
        ]
        pdf_to_docx_method(args.input, args.output, translations)
    else:
        direct_method(args.input, args.output, [])