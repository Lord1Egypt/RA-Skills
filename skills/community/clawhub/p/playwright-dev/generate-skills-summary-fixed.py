# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 创建文档
doc = Document()

# 设置默认字体为微软雅黑
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
style.font.size = Pt(10)
style.font.color.rgb = RGBColor(0, 0, 0)

# 标题
heading = doc.add_heading('阿福 Skills 汇总文档', level=0)
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
heading_run = heading.runs[0]
heading_run.font.name = 'Microsoft YaHei'
heading_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
heading_run.font.size = Pt(18)
heading_run.font.bold = True
heading_run.font.color.rgb = RGBColor(0, 51, 102)

# 文档信息
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_run = info.add_run('文档维护：阿福（AI 助理） | 更新时间：2026-03-07 00:25 | 版本：v1.0')
info_run.font.name = 'Microsoft YaHei'
info_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
info_run.font.size = Pt(9)
info_run.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph('_' * 80)

# Skills 总览
doc.add_heading('Skills 总览', level=1)
doc.add_paragraph('当前共有 2 个 Skills：')

# 表格
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Skill 名称'
hdr_cells[1].text = '一句话概述'
hdr_cells[2].text = '输出格式'

# 加粗表头并设置字体
for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run.font.bold = True
            run.font.size = Pt(10)

# 添加数据
skills_data = [
    ('豆包专家点评', '自动处理豆包会话，生成专家级点评 HTML 网页', 'HTML + 语音'),
    ('项目知识专家', '以行业专家视角构建项目知识库，输出结构化飞书文档', '飞书文档 + 批注')
]

for name, desc, output in skills_data:
    row_cells = table.add_row().cells
    row_cells[0].text = name
    row_cells[1].text = desc
    row_cells[2].text = output
    # 设置每行字体
    for cell in row_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

doc.add_paragraph()

# Skill 1
doc.add_heading('Skill 1: 豆包专家点评系统', level=1)
doc.add_paragraph('一句话概述：自动处理豆包会话内容，生成包含知识架构图、深度洞察和行动建议的专家级点评 HTML 网页。')

doc.add_heading('完整流程（6 步）', level=2)

steps1 = [
    ('步骤 1：保存原始内容', 'doubao-sessions/日期 - 序号.md'),
    ('步骤 2：更新 worklog.txt', '添加今日记录'),
    ('步骤 3：生成专家点评 HTML', '6 章节标准（专家评分/核心观点/深度洞察/知识架构/对比分析/行动建议）'),
    ('步骤 4：Chrome 自动打开', '独立窗口显示'),
    ('步骤 5：TTS 语音生成', 'MP3 到 Temp 目录'),
    ('步骤 6：飞书发送 + 自动播放', 'filePath 参数 + Start-Process')
]

for step, desc in steps1:
    p = doc.add_paragraph(style='List Number')
    step_run = p.add_run(step)
    step_run.font.name = 'Microsoft YaHei'
    step_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    step_run.bold = True
    
    desc_p = doc.add_paragraph(desc, style='List Bullet')
    for run in desc_p.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

# Skill 2
doc.add_heading('Skill 2: 项目知识库构建专家', level=1)
doc.add_paragraph('一句话概述：基于用户提供的信息，以行业专家视角构建项目知识库，输出结构化飞书文档。')

doc.add_heading('完整流程（8 步迭代）', level=2)

steps2 = [
    '意图识别 - 无需触发词，理解真实意图',
    '专家视角分析 - Critical Thinking + 行业对标',
    '补全架构 - 业务/技术/组织三大架构',
    '飞书文档生成（v1.0）- 3 部分结构，分块写入',
    '发送批注清单 - 5 种类型（疑问/风险/建议/待补充/确认）',
    '等待用户回复 - 解析每条批注回复',
    '修订文档（v1.1）- 更新内容 + 版本号 + 修订记录',
    '循环迭代 - 直到所有批注处理完成'
]

for step in steps2:
    p = doc.add_paragraph(step, style='List Number')
    for run in p.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

doc.add_heading('版本号规范', level=2)
version_info = [
    '格式：v{主版本}.{次版本}.{修订号}',
    'v1.0.0 - 初稿',
    'v1.1.0 - 小修订',
    'v2.0.0 - 重大更新'
]

for info in version_info:
    p = doc.add_paragraph(info, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

# 用户偏好
doc.add_heading('用户偏好（通用）', level=1)

preferences = [
    '视觉学习者 - 结构化图表',
    '自动播放 - TTS 后自动播放',
    '金字塔原理 - 结论先行',
    'Critical Thinking - 主动思考',
    'MECE 法则 - 不重不漏'
]

for pref in preferences:
    p = doc.add_paragraph(pref, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

doc.add_paragraph()
doc.add_paragraph('_' * 80)
footer = doc.add_paragraph('文档同步：本文档已同步到飞书和 Word 双格式')
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in footer.runs:
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

# 保存文档 - 使用英文路径
save_path = r'C:\Users\Xiabi\.openclaw\workspace\skills\Afu-Skills-Summary.docx'
doc.save(save_path)
print(f'Word document saved to: {save_path}')
print(f'File size: {len(open(save_path, "rb").read())} bytes')

# 验证文件
from docx import Document as ReadDoc
test_doc = ReadDoc(save_path)
print(f'Verification: Document has {len(test_doc.paragraphs)} paragraphs')
print(f'First paragraph text: {test_doc.paragraphs[0].text[:50]}')
