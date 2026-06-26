#!/usr/bin/env python3
"""
PDF 转 Word 脚本
用法: python pdf_to_word.py <pdf_path> <output.docx>
"""
import sys
import argparse
from pathlib import Path

try:
    import fitz
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Error: Missing dependencies. Run: pip install pymupdf python-docx")
    sys.exit(1)


def pdf_to_word(pdf_path: str, output_path: str) -> bool:
    """将 PDF 转换为 Word 文档"""
    doc = Document()
    
    # 添加标题
    title = doc.add_heading(Path(pdf_path).stem, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    pdf_doc = fitz.open(pdf_path)
    
    for page_num, page in enumerate(pdf_doc, 1):
        # 添加页码标题
        doc.add_heading(f"第 {page_num} 页", level=2)
        
        # 提取文本
        text = page.get_text()
        
        if text.strip():
            # 按段落添加文本
            for paragraph in text.split('\n'):
                if paragraph.strip():
                    p = doc.add_paragraph(paragraph)
        else:
            doc.add_paragraph(f"[第 {page_num} 页为图片或扫描件，建议使用 OCR 提取]")
        
        # 提取图片
        images = page.get_images()
        if images:
            for img_index, img in enumerate(images):
                xref = img[0]
                base_image = pdf_doc.extract_image(xref)
                image_data = base_image["image"]
                image_ext = base_image["ext"]
                
                # 临时保存图片
                temp_img = f"temp_img_{page_num}_{img_index}.{image_ext}"
                with open(temp_img, "wb") as f:
                    f.write(image_data)
                
                # 添加到 Word
                try:
                    doc.add_picture(temp_img, width=Inches(6))
                    # 删除临时文件
                    Path(temp_img).unlink()
                except Exception as e:
                    print(f"Warning: 无法添加图片 {img_index + 1}: {e}")
    
    pdf_doc.close()
    
    # 保存 Word 文档
    doc.save(output_path)
    print(f"转换完成: {output_path}")
    return True


def main():
    parser = argparse.ArgumentParser(description='PDF 转 Word')
    parser.add_argument('pdf_path', help='输入 PDF 文件路径')
    parser.add_argument('output_path', help='输出 Word 文件路径')
    
    args = parser.parse_args()
    
    try:
        pdf_to_word(args.pdf_path, args.output_path)
    except Exception as e:
        print(f"Error: {e}")
        sys.exit(1)


if __name__ == '__main__':
    main()