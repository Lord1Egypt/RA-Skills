#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
格式精修脚本 - 按TOOLS.md规范精修文档格式
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

def format_polish(filepath):
    doc = Document(filepath)
    AVAILABLE_WIDTH = 8800  # A4可用宽度(twips)
    
    # 4.1 全文左对齐+紧凑排版
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            if '1' in para.style.name and '2' not in para.style.name:
                para.paragraph_format.space_before = Pt(12)
                para.paragraph_format.space_after = Pt(4)
            elif '2' in para.style.name:
                para.paragraph_format.space_before = Pt(8)
                para.paragraph_format.space_after = Pt(3)
            para.paragraph_format.line_spacing = 1.1
        else:
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.line_spacing = 1.15
            for run in para.runs:
                if run.font.size is None or run.font.size > Pt(11):
                    run.font.size = Pt(10.5)
    
    # 4.2 表格列宽按内容比例分配
    for table in doc.tables:
        tbl = table._tbl
        col_count = len(table.columns)
        if col_count == 0:
            continue
        
        # 统一表格总宽
        tblPr = tbl.tblPr
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{AVAILABLE_WIDTH}" w:type="dxa"/>')
            tblPr.append(tblW)
        else:
            tblW.set(qn('w:w'), str(AVAILABLE_WIDTH))
            tblW.set(qn('w:type'), 'dxa')
        
        tblLayout = tblPr.find(qn('w:tblLayout'))
        if tblLayout is None:
            tblLayout = parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>')
            tblPr.append(tblLayout)
        else:
            tblLayout.set(qn('w:type'), 'fixed')
        
        # 测量表头宽度（表头不可换行）
        header_row = table.rows[0]
        header_widths = []
        for ci, cell in enumerate(header_row.cells):
            text = cell.text.strip()
            w = sum(240 if ord(ch) > 127 else 120 for ch in text) + 200
            header_widths.append(w)
        
        # 测量内容最大宽度
        col_max_widths = [0] * col_count
        for row in table.rows:
            for ci, cell in enumerate(row.cells):
                if ci >= col_count:
                    break
                text = cell.text.strip()
                w = sum(240 if ord(ch) > 127 else 120 for ch in text) + 200
                col_max_widths[ci] = max(col_max_widths[ci], w)
        
        # 期望宽度 = max(表头需求, 内容需求*0.65)
        desired = []
        for ci in range(col_count):
            content_need = col_max_widths[ci] * 0.65
            header_need = header_widths[ci] * 1.3
            desired.append(max(content_need, header_need, 900))
        
        # 归一化到可用宽度
        total_d = sum(desired)
        col_widths = [int(d / total_d * AVAILABLE_WIDTH) for d in desired]
        
        ABS_MIN = 900  # 每列最低45pt
        for ci in range(col_count):
            if col_widths[ci] < ABS_MIN:
                col_widths[ci] = ABS_MIN
        
        diff = AVAILABLE_WIDTH - sum(col_widths)
        widest = col_widths.index(max(col_widths))
        col_widths[widest] += diff
        
        # 应用列宽
        grid = tbl.find(qn('w:tblGrid'))
        if grid is not None:
            for ci, gridCol in enumerate(grid.findall(qn('w:gridCol'))):
                if ci < col_count:
                    gridCol.set(qn('w:w'), str(col_widths[ci]))
        
        for row in table.rows:
            for ci, cell in enumerate(row.cells):
                if ci >= col_count:
                    break
                tc = cell._tc
                tcPr = tc.tcPr
                if tcPr is None:
                    tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}/>')
                    tc.insert(0, tcPr)
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is None:
                    tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{col_widths[ci]}" w:type="dxa"/>')
                    tcPr.append(tcW)
                else:
                    tcW.set(qn('w:w'), str(col_widths[ci]))
                    tcW.set(qn('w:type'), 'dxa')
        
        # 表格内文字紧凑
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.paragraph_format.space_before = Pt(1)
                    para.paragraph_format.space_after = Pt(1)
                    para.paragraph_format.line_spacing = 1.0
                    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.save(filepath)
    print("✅ 格式精修完成")

if __name__ == "__main__":
    filepath = "河北青山鼎信技术尽调报告.docx"
    print(f"正在精修格式：{filepath}")
    format_polish(filepath)
