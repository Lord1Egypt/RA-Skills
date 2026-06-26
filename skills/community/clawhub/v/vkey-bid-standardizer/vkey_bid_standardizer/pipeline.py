"""vkey-bid-standardizer.pipeline

4 步流水线编排：
- run_step_renumber
- run_step_fix
- run_step_auto_number
- run_full（默认 all）
- review（只读审计）
- convert_md（MD→docx）
"""
import os
import shutil
from dataclasses import dataclass, field
from typing import List, Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT

from .core import (
    replace_para_text,
    reset_run_font_overrides,
    reset_paragraph_overrides,
    set_style_font,
    apply_paragraph_format,
    get_paragraph_overrides,
)
from .patterns import (
    load_patterns,
    match_pattern,
    resolve_number,
    build_number,
)
from .styles import setup_page, setup_heading_styles, setup_normal_style


# ═══════════════════════════════════════════
#  结果对象
# ═══════════════════════════════════════════

@dataclass
class StepResult:
    output: str
    renumbered: int = 0
    fixed: int = 0
    auto_numbered: int = 0
    warnings: List[str] = field(default_factory=list)

    def __str__(self):
        parts = [f'output={self.output}']
        if self.renumbered:
            parts.append(f'renumbered={self.renumbered}')
        if self.fixed:
            parts.append(f'fixed={self.fixed}')
        if self.auto_numbered:
            parts.append(f'auto_numbered={self.auto_numbered}')
        if self.warnings:
            parts.append(f'warnings={len(self.warnings)}')
        return 'StepResult(' + ', '.join(parts) + ')'


@dataclass
class PipelineResult(StepResult):
    steps: List[str] = field(default_factory=list)


# ════════════════════════════════════
#  Step 1: 重编号
# ════════════════════════════════════

def run_step_renumber(input_path: str, output_path: Optional[str], profile: dict) -> StepResult:
    """识别并重写手动编号（一、/（一）/1./1.1.1.1. 等 → 阿拉伯层级）。

    失败降级：未匹配段落原样保留 + warning。
    """
    if output_path is None:
        output_path = input_path.replace('.docx', '_renum.docx')

    patterns = load_patterns(profile)
    h1_rules = patterns.get('h1', [])
    h2_rules = patterns.get('h2', [])
    h3_rules = patterns.get('h3', [])
    h4_rules = patterns.get('h4', [])
    h5_rules = patterns.get('h5', [])

    doc = Document(input_path)
    counters = [0, 0, 0, 0, 0]
    style_rule_map = {
        'Heading 1': (h1_rules, 0),
        'Heading 2': (h2_rules, 1),
        'Heading 3': (h3_rules, 2),
        'Heading 4': (h4_rules, 3),
        'Heading 5': (h5_rules, 4),
    }

    renumbered = 0
    warnings: List[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name
        if style not in style_rule_map:
            continue

        rules, level = style_rule_map[style]
        rule, m, title = match_pattern(rules, text)
        if rule is None:
            warnings.append(f'[{style}] 未识别编号: {text[:50]}')
            continue

        num_val = resolve_number(rule, m)
        if num_val <= 0:
            warnings.append(f'[{style}] 编号解析为 0: {text[:50]}')
            continue

        rule_name = rule.get('name', '?')
        if rule_name in ('chapter_cn', 'chapter_arabic'):
            counters[0] = num_val
            for i in range(1, 5):
                counters[i] = 0
        elif rule_name == 'section_cn':
            counters[0] = max(counters[0], 1)
            counters[1] = num_val
            for i in range(2, 5):
                counters[i] = 0
        elif rule_name == 'article_cn':
            counters[0] = max(counters[0], 1)
            counters[1] = max(counters[1], 1)
            counters[2] = num_val
            for i in range(3, 5):
                counters[i] = 0
        else:
            counters[level] = num_val
            for i in range(level + 1, 5):
                counters[i] = 0

        if level == 0:
            new_text = title
        else:
            new_text = build_number(counters, level + 1) + ' ' + title

        if text != new_text:
            renumbered += 1
            replace_para_text(para, new_text)

    doc.save(output_path)
    return StepResult(output=output_path, renumbered=renumbered, warnings=warnings)


# ════════════════════════════════════
#  Step 2: 修复样式
# ════════════════════════════════════

def run_step_fix(input_path: str, output_path: Optional[str], profile: dict) -> StepResult:
    """规范化样式（页面/Heading/Normal）+ 清段级/run 级覆盖。"""
    if output_path is None:
        output_path = input_path.replace('.docx', '_fixed.docx')

    doc = Document(input_path)

    # 页面
    setup_page(doc, profile['page'])

    # Heading 1-5 样式
    setup_heading_styles(doc, profile)

    # Normal 样式
    setup_normal_style(doc, profile)

    # 清 Normal run 级字号覆盖
    runs_cleared = reset_run_font_overrides(
        doc, ('Normal', 'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5')
    )

    # 清 Normal + Heading 段级 pPr 覆盖
    ppr_cleared = reset_paragraph_overrides(
        doc, ('Normal', 'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5')
    )

    # 修复表格
    tables_fixed = _fix_tables(doc, profile)

    doc.save(output_path)
    return StepResult(
        output=output_path,
        fixed=runs_cleared + ppr_cleared + tables_fixed,
    )


def _fix_tables(doc, profile: dict) -> int:
    """表格：边框 + 底纹 + 表头居中 + 内容左对齐 + 字体。"""
    tbl_cfg = profile.get('tables') or profile.get('table', {})
    hdr = tbl_cfg.get('header', {})
    body = tbl_cfg.get('body', {})
    hdr_align = hdr.get('alignment', 'center')
    body_align = body.get('alignment', 'left')

    align_map = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
    }

    count = 0
    for table in doc.tables:
        # 边框
        _set_table_borders(table, tbl_cfg)
        for ri, row in enumerate(table.rows):
            for cell in row.cells:
                if ri == 0:
                    _shade_cell(cell, hdr.get('fill', 'D9D9D9'))
                    _set_cell_align(cell, align_map.get(hdr_align, WD_ALIGN_PARAGRAPH.LEFT))
                else:
                    _set_cell_align(cell, align_map.get(body_align, WD_ALIGN_PARAGRAPH.LEFT))
                for para in cell.paragraphs:
                    # 表格内段落不允许首行缩进（必须强制段级 w:ind firstLine=0，
                    # 否则会继承 Normal 样式的 2 字符缩进）
                    pPr = para._element.get_or_add_pPr()
                    ind = pPr.find(qn('w:ind'))
                    if ind is None:
                        ind = OxmlElement('w:ind')
                        pPr.append(ind)
                    for k in ('w:firstLine', 'w:firstLineChars',
                              'w:hanging', 'w:hangingChars'):
                        if ind.get(qn(k)) is not None:
                            del ind.attrib[qn(k)]
                    ind.set(qn('w:firstLine'), '0')
                    ind.set(qn('w:firstLineChars'), '0')
                    para.paragraph_format.first_line_indent = Pt(0)
                    for run in para.runs:
                        if not run.text.strip():
                            continue
                        if ri == 0:
                            run.font.name = hdr.get('font_west', 'Times New Roman')
                            rPr = run._element.get_or_add_rPr()
                            rFonts = rPr.find(qn('w:rFonts'))
                            if rFonts is None:
                                rFonts = OxmlElement('w:rFonts')
                                rPr.insert(0, rFonts)
                            rFonts.set(qn('w:eastAsia'), hdr['font_east'])
                            rFonts.set(qn('w:ascii'), hdr.get('font_west', 'Times New Roman'))
                            rFonts.set(qn('w:hAnsi'), hdr.get('font_west', 'Times New Roman'))
                            rFonts.set(qn('w:cs'), hdr.get('font_west', 'Times New Roman'))
                            run.font.size = Pt(hdr['size'])
                            run.font.bold = True
                        else:
                            run.font.name = body.get('font_west', 'FangSong')
                            rPr = run._element.get_or_add_rPr()
                            rFonts = rPr.find(qn('w:rFonts'))
                            if rFonts is None:
                                rFonts = OxmlElement('w:rFonts')
                                rPr.insert(0, rFonts)
                            rFonts.set(qn('w:eastAsia'), body['font_east'])
                            rFonts.set(qn('w:ascii'), body.get('font_west', 'FangSong'))
                            rFonts.set(qn('w:hAnsi'), body.get('font_west', 'FangSong'))
                            rFonts.set(qn('w:cs'), body.get('font_west', 'FangSong'))
                            run.font.size = Pt(body['size'])
                        count += 1
    return count


def _set_table_borders(table, table_cfg: dict) -> None:
    """设置表格边框（外 1.5pt / 内 0.5pt / 黑色）。"""
    def sz(v, default):
        try:
            return int(round(float(str(v).rstrip('pt').rstrip('PT')) * 8))
        except (ValueError, AttributeError):
            return int(default * 8)

    outer_sz = sz(table_cfg.get('border_outer'), 1.5)
    inner_sz = sz(table_cfg.get('border_inner'), 0.5)
    color = table_cfg.get('border_color', '#000000').lstrip('#')

    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    for e in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(e)
    borders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        e = OxmlElement(f'w:{side}')
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), str(outer_sz))
        e.set(qn('w:space'), '0')
        e.set(qn('w:color'), color)
        borders.append(e)
    for side in ('insideH', 'insideV'):
        e = OxmlElement(f'w:{side}')
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), str(inner_sz))
        e.set(qn('w:space'), '0')
        e.set(qn('w:color'), color)
        borders.append(e)
    tblPr.append(borders)


def _shade_cell(cell, color: str = 'D9D9D9') -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    for e in tcPr.findall(qn('w:shd')):
        tcPr.remove(e)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)


def _set_cell_align(cell, align) -> None:
    for para in cell.paragraphs:
        para.alignment = align


# ════════════════════════════════════
#  Step 3: 自动编号（Word 多级列表）
# ════════════════════════════════════

def run_step_auto_number(input_path: str, output_path: Optional[str], profile: dict) -> StepResult:
    """将 Heading 1-5 样式绑定到 Word 多级列表（5 级：1./1.1/1.1.1/1.1.1.1/1.1.1.1.1）。

    会自动剥离段落中已写好的手动编号（被多级列表替换）。
    """
    if output_path is None:
        output_path = input_path.replace('.docx', '_auto.docx')

    doc = Document(input_path)

    # 剥离手动编号（Heading 1-5 段落）
    style_patterns = {
        'Heading 1': re_compile(r'^[一二三四五六七八九十百千]+[、．.]\s*'),
        'Heading 2': re_compile(r'^\d+\.\d+\s+'),
        'Heading 3': re_compile(r'^\d+\.\d+\.\d+\s+'),
        'Heading 4': re_compile(r'^\d+\.\d+\.\d+\.\d+\s+'),
        'Heading 5': re_compile(r'^\d+\.\d+\.\d+\.\d+\.\d+\s+'),
    }
    stripped = 0
    for para in doc.paragraphs:
        pat = style_patterns.get(para.style.name)
        if pat is None:
            continue
        m = pat.match(para.text)
        if not m:
            continue
        remaining = para.text[m.end():]
        runs = para.runs
        if runs:
            for r in runs:
                r.text = ''
            runs[0].text = remaining
        else:
            para.add_run(remaining)
        stripped += 1

    # 添加多级列表定义
    num_id = _add_multilevel_numbering(doc)

    # 绑定样式
    bound = 0
    for level, name in enumerate(['heading 1', 'heading 2', 'heading 3', 'heading 4', 'heading 5']):
        if _link_style_to_list(doc, name, level, num_id):
            bound += 1

    doc.save(output_path)
    return StepResult(output=output_path, auto_numbered=bound)


import re as _re_mod
def re_compile(pattern):
    return _re_mod.compile(pattern)


def _add_multilevel_numbering(doc) -> str:
    """添加 5 级多级列表定义到 numbering.xml，返回 numId（默认 '100'）。"""
    ABSTRACT_NUM_ID = '100'
    NUM_ID = '100'

    numbering = doc.part.numbering_part.element

    # 删除已有 abstractNum=100 / num=100（允许重入）
    for an in list(numbering.findall(qn('w:abstractNum'))):
        if an.get(qn('w:abstractNumId')) == ABSTRACT_NUM_ID:
            numbering.remove(an)
    for n in list(numbering.findall(qn('w:num'))):
        if n.get(qn('w:numId')) == NUM_ID:
            numbering.remove(n)

    abstract_num = OxmlElement('w:abstractNum')
    abstract_num.set(qn('w:abstractNumId'), ABSTRACT_NUM_ID)

    mlt = OxmlElement('w:multiLevelType')
    mlt.set(qn('w:val'), 'multilevel')
    abstract_num.append(mlt)

    abstract_num.append(_make_lvl(0, 'decimal', '%1.'))
    abstract_num.append(_make_lvl(1, 'decimal', '%1.%2'))
    abstract_num.append(_make_lvl(2, 'decimal', '%1.%2.%3'))
    abstract_num.append(_make_lvl(3, 'decimal', '%1.%2.%3.%4'))
    abstract_num.append(_make_lvl(4, 'decimal', '%1.%2.%3.%4.%5'))

    first_num = numbering.find(qn('w:num'))
    if first_num is not None:
        first_num.addprevious(abstract_num)
    else:
        numbering.append(abstract_num)

    num = OxmlElement('w:num')
    num.set(qn('w:numId'), NUM_ID)
    anid = OxmlElement('w:abstractNumId')
    anid.set(qn('w:val'), ABSTRACT_NUM_ID)
    num.append(anid)
    numbering.append(num)

    return NUM_ID


def _make_lvl(ilvl: int, num_fmt: str, lvl_text: str, suff: str = 'space'):
    """构造 w:lvl 元素。"""
    lvl = OxmlElement('w:lvl')
    lvl.set(qn('w:ilvl'), str(ilvl))
    lvl.set(qn('w:tplc'), '0809000F')

    start_el = OxmlElement('w:start')
    start_el.set(qn('w:val'), '1')
    lvl.append(start_el)

    nf = OxmlElement('w:numFmt')
    nf.set(qn('w:val'), num_fmt)
    lvl.append(nf)

    suff_el = OxmlElement('w:suff')
    suff_el.set(qn('w:val'), suff)
    lvl.append(suff_el)

    lt = OxmlElement('w:lvlText')
    lt.set(qn('w:val'), lvl_text)
    lvl.append(lt)

    ljc = OxmlElement('w:lvlJc')
    ljc.set(qn('w:val'), 'left')
    lvl.append(ljc)

    ppr = OxmlElement('w:pPr')
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '0')
    ind.set(qn('w:hanging'), '0')
    ppr.append(ind)
    lvl.append(ppr)

    return lvl


def _link_style_to_list(doc, target_name: str, ilvl: int, num_id: str) -> bool:
    """将样式链接到多级列表的某个级别。"""
    styles_xml = doc.styles.element
    target = None
    for s in styles_xml.findall(qn('w:style')):
        name_el = s.find(qn('w:name'))
        if name_el is None:
            continue
        if name_el.get(qn('w:val')) == target_name:
            target = s
            break
    if target is None:
        return False

    pPr = target.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        target.append(pPr)

    for old in pPr.findall(qn('w:numPr')):
        pPr.remove(old)

    numPr = OxmlElement('w:numPr')
    ilvl_el = OxmlElement('w:ilvl')
    ilvl_el.set(qn('w:val'), str(ilvl))
    numId_el = OxmlElement('w:numId')
    numId_el.set(qn('w:val'), str(num_id))
    numPr.append(ilvl_el)
    numPr.append(numId_el)
    pPr.insert(0, numPr)
    return True


# ════════════════════════════════════
#  全流水线
# ════════════════════════════════════

def run_full(input_path: str, output_path: Optional[str], profile: dict,
             do_renumber: bool = True, do_fix: bool = True, do_auto_number: bool = True,
             dry_run: bool = False, backup: bool = False) -> PipelineResult:
    """全流水线：renumber → fix → auto_number。

    Args:
        input_path: 输入 docx
        output_path: 最终输出（None → 自动在同目录生成 _all.docx）
        do_* 子步骤开关
        dry_run: True 时只输出计划，不写文件
        backup: True 时把 input 备份到 input.bak
    """
    if output_path is None:
        output_path = input_path.replace('.docx', '_all.docx')

    if backup and not dry_run:
        bak = input_path + '.bak'
        if not os.path.exists(bak):
            shutil.copy2(input_path, bak)

    result = PipelineResult(output=output_path, steps=[])
    cur = input_path
    intermediate = None
    tmp_paths: List[str] = []  # 追踪所有临时文件，异常时清理

    try:
        if do_renumber:
            if dry_run:
                result.steps.append('renumber (skipped, dry-run)')
            else:
                intermediate = output_path + '.tmp1'
                r = run_step_renumber(cur, intermediate, profile)
                result.renumbered = r.renumbered
                result.warnings.extend(r.warnings)
                cur = intermediate
                tmp_paths.append(intermediate)
                result.steps.append(f'renumber ({r.renumbered})')

        if do_fix:
            if dry_run:
                result.steps.append('fix (skipped, dry-run)')
            else:
                intermediate = output_path + '.tmp2'
                r = run_step_fix(cur, intermediate, profile)
                result.fixed = r.fixed
                cur = intermediate
                tmp_paths.append(intermediate)
                result.steps.append(f'fix ({r.fixed})')

        if do_auto_number:
            if dry_run:
                result.steps.append('auto-number (skipped, dry-run)')
            else:
                intermediate = output_path + '.tmp3'
                r = run_step_auto_number(cur, intermediate, profile)
                result.auto_numbered = r.auto_numbered
                cur = intermediate
                tmp_paths.append(intermediate)
                result.steps.append(f'auto-number ({r.auto_numbered})')

        if not dry_run and cur != output_path and os.path.exists(cur):
            if os.path.exists(output_path):
                os.remove(output_path)
            shutil.move(cur, output_path)
            # move 成功后，原 tmp 路径已不存在，从清理列表移除
            if cur in tmp_paths:
                tmp_paths.remove(cur)
    finally:
        # 清理所有未消费的 tmp 文件（异常或提前退出场景）
        for p in tmp_paths:
            try:
                if os.path.exists(p):
                    os.remove(p)
            except OSError:
                pass

    return result


# ════════════════════════════════════
#  审计
# ════════════════════════════════════

def review(input_path: str, profile: dict) -> dict:
    """只读审计：检查样式/字体/缩进/编号，返回结构化报告。"""
    doc = Document(input_path)
    issues: List[str] = []
    stats = {
        'Heading 1': 0, 'Heading 2': 0, 'Heading 3': 0,
        'Heading 4': 0, 'Heading 5': 0, 'Normal': 0, 'Other': 0,
    }
    overrides = {'spacing': 0, 'ind': 0, 'jc': 0, 'outlineLvl': 0, 'numPr': 0}

    for para in doc.paragraphs:
        style = para.style.name
        if style in stats:
            stats[style] += 1
        else:
            stats['Other'] += 1

        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            for tag in ('w:spacing', 'w:ind', 'w:jc', 'w:outlineLvl', 'w:numPr'):
                if pPr.find(qn(tag)) is not None:
                    overrides[tag[2:]] += 1
                    if style.startswith('Heading'):
                        issues.append(f'[{style}] 段级 {tag[2:]} 覆盖: {para.text[:30]}')

    return {
        'input': input_path,
        'paragraphs': sum(stats.values()),
        'by_style': stats,
        'overrides': overrides,
        'issues': issues,
        'issue_count': len(issues),
    }


# ════════════════════════════════════
#  MD → docx
# ════════════════════════════════════

def convert_md(input_md: str, output_docx: Optional[str], profile: dict) -> StepResult:
    """Markdown → docx 转换。解析部分从 markdown_parser 包导入。"""
    from .markdown_parser import classify_line, parse_md_table, strip_md_format

    if output_docx is None:
        output_docx = input_md.replace('.md', '.docx')

    with open(input_md, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    setup_page(doc, profile['page'])
    setup_heading_styles(doc, profile)
    setup_normal_style(doc, profile)

    i = 0
    while i < len(lines):
        line = lines[i]
        line_type, content = classify_line(line)

        if line_type == 'empty':
            i += 1
            continue

        if line_type == 'h1':
            b = profile['body']
            p = doc.add_paragraph(style='Normal')
            run = p.add_run(content)
            _set_style_font_for_run(run, '黑体', '黑体', 18, bold=True)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = Pt(b['line_spacing'])
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(16)

        elif line_type in ('h2', 'h3', 'h4', 'h5', 'h6'):
            level = int(line_type[1]) - 1  # h2 → 1
            h_cfg = next((x for x in profile['headings'] if x['level'] == level), None)
            if h_cfg:
                p = doc.add_paragraph(content, style=f'Heading {level}')
                p.paragraph_format.space_before = Pt(h_cfg['space_before'])
                p.paragraph_format.space_after = Pt(h_cfg['space_after'])

        elif line_type == 'table_row':
            _build_table(doc, lines, i, profile)
            _, next_i = parse_md_table(lines, i)
            i = next_i
            continue

        elif line_type == 'body':
            content = _re_sub_leading_list(content)
            _add_body_para(doc, content, profile)

        i += 1

    doc.save(output_docx)
    return StepResult(output=output_docx, fixed=1)


def _re_sub_leading_list(text: str) -> str:
    import re as _re
    text = _re.sub(r'^[-*]\s+', '', text)
    text = _re.sub(r'^\d+[.\)]\s+', '', text)
    return text


def _set_style_font_for_run(run, west_font: str, east_font: str, size: float, bold: bool = False) -> None:
    """为单个 run 设置字体（与 core.set_style_font 类似但作用于 run）。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    run.font.name = west_font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    for attr in ['w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs']:
        rFonts.set(qn(attr), west_font if attr in ('w:ascii', 'w:hAnsi', 'w:cs') else east_font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def _add_body_para(doc, text: str, profile: dict) -> None:
    b = profile['body']
    p = doc.add_paragraph(style='Normal')
    run = p.add_run(text)
    _set_style_font_for_run(run, b['font_west'], b['font_east'], b['size'])
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(b['line_spacing'])
    p.paragraph_format.space_before = Pt(b['space_before'])
    p.paragraph_format.space_after = Pt(b['space_after'])
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(2 * b['size'] * 0.8)


def _build_table(doc, lines, start_idx: int, profile: dict) -> None:
    """构造 docx 表格（与原 apply_md_to_docx 同款）。"""
    from .markdown_parser import parse_md_table, strip_md_format
    rows, _ = parse_md_table(lines, start_idx)
    if len(rows) < 2:
        return
    tbl_cfg = profile.get('tables') or profile.get('table', {})
    hdr_cfg = tbl_cfg.get('header', {})
    body_cfg = tbl_cfg.get('body', {})

    num_rows = len(rows)
    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table, tbl_cfg)

    for ri, row_data in enumerate(rows):
        for ci in range(num_cols):
            cell_text = row_data[ci] if ci < len(row_data) else ''
            cell_text = strip_md_format(cell_text)
            cell = table.rows[ri].cells[ci]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(cell_text)
            if ri == 0:
                _set_style_font_for_run(
                    run,
                    hdr_cfg.get('font_west', 'Times New Roman'),
                    hdr_cfg.get('font_east', '黑体'),
                    hdr_cfg.get('size', 10.5),
                    bold=True,
                )
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                _shade_cell(cell, hdr_cfg.get('fill', 'D9D9D9'))
            else:
                _set_style_font_for_run(
                    run,
                    body_cfg.get('font_west', 'FangSong'),
                    body_cfg.get('font_east', '仿宋'),
                    body_cfg.get('size', 10.5),
                )
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
