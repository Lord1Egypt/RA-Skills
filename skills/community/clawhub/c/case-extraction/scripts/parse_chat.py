"""
聊天记录解析器
支持 .txt 和 .docx 格式的聊天记录

用法：
    python parse_chat.py <文件路径>

输出：
    解析后的纯文本内容
"""

import sys
import re
import os


def parse_txt(file_path):
    """解析 TXT 格式聊天记录"""
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    return content


def parse_docx(file_path):
    """解析 DOCX 格式聊天记录"""
    try:
        from docx import Document
    except ImportError:
        print("请先安装 python-docx: pip install python-docx")
        sys.exit(1)
    
    doc = Document(file_path)
    lines = []
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text)
    
    # 也读取表格
    for table in doc.tables:
        for row in table.rows:
            row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                lines.append(row_text)
    
    return '\n'.join(lines)


def clean_chat_content(content):
    """
    清洗聊天记录内容，提取关键信息
    
    策略：
    1. 去除时间戳
    2. 去除重复的空行
    3. 识别并标记发言人
    """
    
    lines = content.split('\n')
    cleaned = []
    
    # 常见时间戳模式
    time_patterns = [
        r'^\d{4}-\d{2}-\d{2}\s+\d{2}:\d{2}:\d{2}',  # 2024-01-01 12:00:00
        r'^\d{2}:\d{2}:\d{2}',  # 12:00:00
        r'^\d{2}:\d{2}',  # 12:00
        r'^\[\d{2}:\d{2}\]',  # [12:00]
        r'^\(\d{2}:\d{2}\)',  # (12:00)
    ]
    
    speaker_pattern = r'^([^：:\s]+)[:：]\s*(.+)'  # 发言人: 内容
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # 去除时间戳
        for pattern in time_patterns:
            line = re.sub(pattern, '', line)
        
        line = line.strip()
        if not line:
            continue
        
        # 尝试识别发言人
        speaker_match = re.match(speaker_pattern, line)
        if speaker_match:
            speaker, msg = speaker_match.groups()
            cleaned.append(f"[{speaker}] {msg}")
        else:
            cleaned.append(line)
    
    return '\n'.join(cleaned)


def extract_key_info(content):
    """
    从聊天记录中提取关键信息
    
    提取：
    - 业务场景
    - 核心问题
    - 解决方案
    - 关键结果
    """
    
    # 关键词模式
    patterns = {
        'scenario': [r'场景[：:]\s*(.+)', r'情况[：:]\s*(.+)', r'背景[：:]\s*(.+)'],
        'problem': [r'问题[：:]\s*(.+)', r'挑战[：:]\s*(.+)', r'困难[：:]\s*(.+)'],
        'solution': [r'解决[方案]?[：:]\s*(.+)', r'处理[方式]?[：:]\s*(.+)', r'怎么做的[：:]\s*(.+)'],
        'result': [r'结果[：:]\s*(.+)', r'效果[：:]\s*(.+)', r'业绩[：:]\s*(.+)', r'达成[：:]\s*(.+)'],
    }
    
    extracted = {}
    
    for key, keyword_patterns in patterns.items():
        for pattern in keyword_patterns:
            match = re.search(pattern, content)
            if match:
                extracted[key] = match.group(1)
                break
    
    return extracted


def main():
    if len(sys.argv) < 2:
        print("用法: python parse_chat.py <文件路径>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    
    if not os.path.exists(file_path):
        print(f"文件不存在: {file_path}")
        sys.exit(1)
    
    # 根据扩展名选择解析方式
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.txt':
        content = parse_txt(file_path)
    elif ext in ['.docx', '.doc']:
        content = parse_docx(file_path)
    else:
        print(f"不支持的格式: {ext}")
        print("支持: .txt, .docx, .doc")
        sys.exit(1)
    
    # 清洗内容
    cleaned = clean_chat_content(content)
    
    # 提取关键信息
    key_info = extract_key_info(cleaned)
    
    # 输出结果
    print("=" * 50)
    print("解析完成")
    print("=" * 50)
    
    if key_info:
        print("\n【关键信息提取】")
        for key, value in key_info.items():
            print(f"  {key}: {value}")
    
    print("\n【清洗后的内容】")
    print("-" * 50)
    print(cleaned)
    print("-" * 50)
    
    # 保存清洗后的内容到临时文件
    temp_path = os.path.join(os.path.dirname(file_path), '_cleaned_temp.txt')
    with open(temp_path, 'w', encoding='utf-8') as f:
        f.write(cleaned)
    print(f"\n清洗后的内容已保存到: {temp_path}")
    
    return cleaned, key_info


if __name__ == '__main__':
    main()
