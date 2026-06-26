# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 创建文档
doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(10)

# 标题
heading = doc.add_heading('感知与行动中心 - 专家解决方案 v2.1', level=0)
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
heading_run = heading.runs[0]
heading_run.font.name = 'Microsoft YaHei'
heading_run.font.size = Pt(16)
heading_run.font.bold = True

# 文档信息
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.LEFT
info.add_run('文档维护：阿福（AI 助理） | 更新频率：每周（周五） | 下次更新：2026-03-13\n')
info.add_run('版本号：v2.1（2026-03-06 修订）\n')
info.add_run('本文档由阿福自动生成，基于项目知识库和行业最佳实践')

doc.add_paragraph('_' * 80)

# 第一部分
doc.add_heading('📋 项目基本信息摘要', level=1)

doc.add_heading('项目概述', level=2)
doc.add_paragraph('分期实施：', style='Intense Quote')
doc.add_paragraph('• 一期（2026-03 至 2026-06）：感知侧 - 数据采集、监控看板、数据治理', style='List Bullet')
doc.add_paragraph('• 二期（2026-07 至 2026-09）：行动侧 - 风险预警、工单闭环、组织保障', style='List Bullet')

doc.add_paragraph('一期目标：', style='Intense Quote')
doc.add_paragraph('• 115 家供应商产能数据采集完成', style='List Bullet')
doc.add_paragraph('• 6 大核心指标监控看板上线', style='List Bullet')
doc.add_paragraph('• 数据字段标准化（23 个字段）', style='List Bullet')
doc.add_paragraph('• 风险识别规则配置', style='List Bullet')

# 保存文档
save_path = r'C:\Users\Xiabi\.openclaw\workspace\感知与行动中心 - 专家解决方案 v2.1.docx'
doc.save(save_path)
print(f'Word document saved to: {save_path}')
