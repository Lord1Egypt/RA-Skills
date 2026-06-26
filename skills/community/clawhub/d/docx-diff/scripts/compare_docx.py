#!/usr/bin/env python3
"""
docx-diff: 对比两份 Word 文档，找出文本、图片、格式/样式的差异
输出：高亮标注的 .docx + 纯文本报告 + HTML 可视化报告

用法：
    python compare_docx.py --file1 旧文档.docx --file2 新文档.docx
    python compare_docx.py --file1 旧.docx --file2 新.docx --output ./结果 --threshold 0.6
"""

import argparse
import base64
import difflib
import hashlib
import html as html_lib
import json
import os
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor, Pt


# ── 颜色 ─────────────────────────────────────────────────
COLOR_DELETED = RGBColor(0xFF, 0xCC, 0xCC)
COLOR_ADDED   = RGBColor(0xCC, 0xFF, 0xCC)


# ── 数据结构 ──────────────────────────────────────────────
@dataclass
class ParaDiff:
    tag: str          # equal | delete | insert | replace
    old_text: str = ''
    new_text: str = ''
    old_style: dict = field(default_factory=dict)
    new_style: dict = field(default_factory=dict)
    style_changes: list = field(default_factory=list)  # 样式变化描述

@dataclass
class ImageDiff:
    tag: str          # deleted | added | unchanged
    md5: str = ''
    ref: str = ''
    data: bytes = b'' # 图片二进制（用于 HTML 内嵌）

@dataclass
class StyleDiff:
    para_index: int
    text_preview: str
    changes: list     # [描述字符串]


# ── 工具函数 ──────────────────────────────────────────────
def normalize(text: str) -> str:
    text = text.strip()
    text = re.sub(r'\s+', ' ', text)
    return text


def extract_para_style(para) -> dict:
    """提取段落的样式信息：字体、字号、粗体、斜体、颜色、对齐、段落样式名"""
    style = {
        'style_name': para.style.name if para.style else '',
        'alignment': str(para.alignment),
        'bold': None,
        'italic': None,
        'font_size': None,
        'font_name': None,
        'font_color': None,
    }
    for run in para.runs:
        if run.bold is not None and style['bold'] is None:
            style['bold'] = run.bold
        if run.italic is not None and style['italic'] is None:
            style['italic'] = run.italic
        if run.font.size is not None and style['font_size'] is None:
            style['font_size'] = run.font.size.pt if run.font.size else None
        if run.font.name is not None and style['font_name'] is None:
            style['font_name'] = run.font.name
        if run.font.color and run.font.color.rgb and style['font_color'] is None:
            style['font_color'] = str(run.font.color.rgb)
    return style


def diff_styles(old: dict, new: dict) -> list:
    """对比两个样式字典，返回变化描述列表"""
    changes = []
    labels = {
        'style_name': '段落样式',
        'alignment': '对齐方式',
        'bold': '粗体',
        'italic': '斜体',
        'font_size': '字号',
        'font_name': '字体',
        'font_color': '字体颜色',
    }
    for key, label in labels.items():
        ov, nv = old.get(key), new.get(key)
        if ov != nv and not (ov is None and nv is None):
            changes.append(f'{label}: {ov!r} → {nv!r}')
    return changes


def extract_paragraphs(doc: Document):
    """返回 [(normalized_text, style_dict)] 列表（跳过空段落，包含表格内容）"""
    result = []
    for p in doc.paragraphs:
        text = normalize(p.text)
        if text:
            result.append((text, extract_para_style(p)))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text = normalize(p.text)
                    if text:
                        result.append((text, extract_para_style(p)))
    return result


def extract_images(doc: Document) -> dict:
    """返回 {md5: (ref, blob)} 字典"""
    images = {}
    for rel in doc.part.rels.values():
        if 'image' in rel.reltype:
            try:
                blob = rel.target_part.blob
                md5 = hashlib.md5(blob).hexdigest()
                images[md5] = (rel.target_ref, blob)
            except Exception:
                pass
    return images


# ── 文本 Diff ─────────────────────────────────────────────
def diff_paragraphs(paras1, paras2, threshold=0.6) -> list[ParaDiff]:
    texts1 = [p[0] for p in paras1]
    texts2 = [p[0] for p in paras2]
    styles1 = {t: s for t, s in paras1}
    styles2 = {t: s for t, s in paras2}

    matcher = difflib.SequenceMatcher(None, texts1, texts2, autojunk=False)
    results = []

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            for t in texts1[i1:i2]:
                os_ = styles1.get(t, {})
                ns_ = styles2.get(t, {})
                sc = diff_styles(os_, ns_)
                results.append(ParaDiff('equal', t, t, os_, ns_, sc))

        elif tag == 'delete':
            for t in texts1[i1:i2]:
                results.append(ParaDiff('delete', t, '', styles1.get(t, {}), {}))

        elif tag == 'insert':
            for t in texts2[j1:j2]:
                results.append(ParaDiff('insert', '', t, {}, styles2.get(t, {})))

        elif tag == 'replace':
            old_block = texts1[i1:i2]
            new_block = texts2[j1:j2]
            used_new = set()
            paired = []

            for old in old_block:
                best_score, best_j = 0, None
                for jj, new in enumerate(new_block):
                    if jj in used_new:
                        continue
                    score = difflib.SequenceMatcher(None, old, new).ratio()
                    if score > best_score:
                        best_score, best_j = score, jj
                if best_j is not None and best_score >= threshold:
                    paired.append((old, new_block[best_j], best_j))
                    used_new.add(best_j)
                else:
                    results.append(ParaDiff('delete', old, '', styles1.get(old, {}), {}))

            for jj, new in enumerate(new_block):
                if jj not in {p[2] for p in paired}:
                    results.append(ParaDiff('insert', '', new, {}, styles2.get(new, {})))

            for old, new, _ in paired:
                os_ = styles1.get(old, {})
                ns_ = styles2.get(new, {})
                sc = diff_styles(os_, ns_)
                results.append(ParaDiff('replace', old, new, os_, ns_, sc))

    return results


# ── 图片 Diff ─────────────────────────────────────────────
def diff_images(doc1, doc2) -> list[ImageDiff]:
    imgs1 = extract_images(doc1)
    imgs2 = extract_images(doc2)
    results = []
    for md5, (ref, blob) in imgs1.items():
        if md5 not in imgs2:
            results.append(ImageDiff('deleted', md5, ref, blob))
        else:
            results.append(ImageDiff('unchanged', md5, ref, blob))
    for md5, (ref, blob) in imgs2.items():
        if md5 not in imgs1:
            results.append(ImageDiff('added', md5, ref, blob))
    return results


# ── 输出：高亮 DOCX ───────────────────────────────────────
def highlight_paragraph(para, color: RGBColor):
    hex_color = '{:02X}{:02X}{:02X}'.format(color[0], color[1], color[2])
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)


def add_label(para, label: str, rgb: RGBColor):
    run = para.add_run(label + ' ')
    run.bold = True
    run.font.color.rgb = rgb


def build_docx(para_diffs: list[ParaDiff], image_diffs: list[ImageDiff], stats, img_counts, out_path: Path):
    doc = Document()
    doc.add_heading('文档 Diff 报告', level=1)

    doc.add_heading('一、文本变化', level=2)
    stats = {'equal': 0, 'delete': 0, 'insert': 0, 'replace': 0}

    for d in para_diffs:
        stats[d.tag] += 1
        if d.tag == 'equal':
            p = doc.add_paragraph(d.old_text)
            if d.style_changes:
                note = doc.add_paragraph()
                r = note.add_run('  ↳ 样式变化: ' + '；'.join(d.style_changes))
                r.italic = True
                r.font.color.rgb = RGBColor(0x88, 0x88, 0x00)
        elif d.tag == 'delete':
            p = doc.add_paragraph()
            add_label(p, '[-]', RGBColor(0xCC, 0x00, 0x00))
            p.add_run(d.old_text)
            highlight_paragraph(p, COLOR_DELETED)
        elif d.tag == 'insert':
            p = doc.add_paragraph()
            add_label(p, '[+]', RGBColor(0x00, 0x88, 0x00))
            p.add_run(d.new_text)
            highlight_paragraph(p, COLOR_ADDED)
        elif d.tag == 'replace':
            p_old = doc.add_paragraph()
            add_label(p_old, '[修改前]', RGBColor(0xCC, 0x00, 0x00))
            p_old.add_run(d.old_text)
            highlight_paragraph(p_old, COLOR_DELETED)
            p_new = doc.add_paragraph()
            add_label(p_new, '[修改后]', RGBColor(0x00, 0x88, 0x00))
            p_new.add_run(d.new_text)
            highlight_paragraph(p_new, COLOR_ADDED)
            if d.style_changes:
                note = doc.add_paragraph()
                r = note.add_run('  ↳ 样式变化: ' + '；'.join(d.style_changes))
                r.italic = True
                r.font.color.rgb = RGBColor(0x88, 0x88, 0x00)

    doc.add_heading('二、图片变化', level=2)
    img_counts = {'deleted': 0, 'added': 0, 'unchanged': 0}
    for img in image_diffs:
        img_counts[img.tag] += 1
        if img.tag == 'deleted':
            p = doc.add_paragraph()
            add_label(p, '[-] 图片删除', RGBColor(0xCC, 0x00, 0x00))
            p.add_run(f' (MD5: {img.md5[:8]}…)')
            highlight_paragraph(p, COLOR_DELETED)
        elif img.tag == 'added':
            p = doc.add_paragraph()
            add_label(p, '[+] 图片新增', RGBColor(0x00, 0x88, 0x00))
            p.add_run(f' (MD5: {img.md5[:8]}…)')
            highlight_paragraph(p, COLOR_ADDED)

    if all(v == 0 for v in img_counts.values() if v != img_counts['unchanged']):
        doc.add_paragraph('图片无变化。')

    doc.save(str(out_path))
    return stats, img_counts


# ── 输出：HTML 报告 ───────────────────────────────────────
def img_to_b64(blob: bytes, ref: str) -> str:
    ext = Path(ref).suffix.lower().lstrip('.') or 'png'
    mime = {'jpg': 'jpeg', 'jpeg': 'jpeg', 'png': 'png', 'gif': 'gif', 'bmp': 'bmp'}.get(ext, 'png')
    return f"data:image/{mime};base64,{base64.b64encode(blob).decode()}"


def inline_diff_html(old: str, new: str) -> tuple[str, str]:
    """字符级 diff，返回 (old_html, new_html) 带高亮"""
    matcher = difflib.SequenceMatcher(None, old, new)
    old_html, new_html = [], []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        old_chunk = html_lib.escape(old[i1:i2])
        new_chunk = html_lib.escape(new[j1:j2])
        if tag == 'equal':
            old_html.append(old_chunk)
            new_html.append(new_chunk)
        elif tag == 'delete':
            old_html.append(f'<mark class="del">{old_chunk}</mark>')
        elif tag == 'insert':
            new_html.append(f'<mark class="ins">{new_chunk}</mark>')
        elif tag == 'replace':
            old_html.append(f'<mark class="del">{old_chunk}</mark>')
            new_html.append(f'<mark class="ins">{new_chunk}</mark>')
    return ''.join(old_html), ''.join(new_html)


def build_html(para_diffs, image_diffs, stats, img_counts, file1, file2, out_path: Path):
    rows = []
    for d in para_diffs:
        if d.tag == 'equal' and not d.style_changes:
            row = f'''<tr class="eq">
  <td class="tag">＝</td>
  <td>{html_lib.escape(d.old_text)}</td>
  <td>{html_lib.escape(d.new_text)}</td>
  <td class="style-info"></td>
</tr>'''
        elif d.tag == 'delete':
            row = f'''<tr class="del">
  <td class="tag">－</td>
  <td class="del-cell">{html_lib.escape(d.old_text)}</td>
  <td></td>
  <td class="style-info"></td>
</tr>'''
        elif d.tag == 'insert':
            row = f'''<tr class="ins">
  <td class="tag">＋</td>
  <td></td>
  <td class="ins-cell">{html_lib.escape(d.new_text)}</td>
  <td class="style-info"></td>
</tr>'''
        elif d.tag == 'replace':
            old_h, new_h = inline_diff_html(d.old_text, d.new_text)
            style_html = ''
            if d.style_changes:
                items = ''.join(f'<li>{html_lib.escape(c)}</li>' for c in d.style_changes)
                style_html = f'<ul class="style-list">{items}</ul>'
            row = f'''<tr class="rep">
  <td class="tag">～</td>
  <td class="del-cell">{old_h}</td>
  <td class="ins-cell">{new_h}</td>
  <td class="style-info">{style_html}</td>
</tr>'''
        elif d.tag == 'equal' and d.style_changes:
            items = ''.join(f'<li>{html_lib.escape(c)}</li>' for c in d.style_changes)
            row = f'''<tr class="style-only">
  <td class="tag">≈</td>
  <td>{html_lib.escape(d.old_text)}</td>
  <td>{html_lib.escape(d.new_text)}</td>
  <td class="style-info"><ul class="style-list">{items}</ul></td>
</tr>'''
        rows.append(row)

    # 图片部分
    img_rows = []
    for img in image_diffs:
        if img.tag == 'unchanged':
            continue
        tag_label = '－ 图片删除' if img.tag == 'deleted' else '＋ 图片新增'
        tag_cls = 'del' if img.tag == 'deleted' else 'ins'
        try:
            img_src = img_to_b64(img.data, img.ref)
            img_tag = f'<img src="{img_src}" style="max-width:200px;max-height:200px;border:1px solid #ccc">'
        except Exception:
            img_tag = '(无法预览)'
        img_rows.append(f'''<tr class="{tag_cls}">
  <td class="tag">{tag_label}</td>
  <td colspan="2">{img_tag}<br><small>MD5: {img.md5[:16]}…</small></td>
  <td></td>
</tr>''')

    img_section = ''
    if img_rows:
        img_section = f'''
    <h2>图片变化</h2>
    <table>
      <thead><tr><th>类型</th><th colspan="2">图片预览</th><th>备注</th></tr></thead>
      <tbody>{''.join(img_rows)}</tbody>
    </table>'''

    html = f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<title>文档 Diff 报告</title>
<style>
  * {{ box-sizing: border-box; margin: 0; padding: 0; }}
  body {{ font-family: 'PingFang SC', 'Microsoft YaHei', sans-serif; font-size: 14px;
         background: #f5f5f5; color: #333; padding: 24px; }}
  h1 {{ font-size: 22px; margin-bottom: 8px; color: #222; }}
  h2 {{ font-size: 16px; margin: 24px 0 8px; color: #444; border-left: 4px solid #4a9eff; padding-left: 8px; }}
  .meta {{ color: #666; font-size: 12px; margin-bottom: 16px; }}
  .stats {{ display: flex; gap: 12px; margin-bottom: 20px; flex-wrap: wrap; }}
  .stat {{ background: #fff; border-radius: 8px; padding: 10px 16px; text-align: center;
           box-shadow: 0 1px 4px rgba(0,0,0,.08); min-width: 90px; }}
  .stat .n {{ font-size: 24px; font-weight: bold; }}
  .stat .l {{ font-size: 11px; color: #888; margin-top: 2px; }}
  .n-del {{ color: #e55; }}
  .n-ins {{ color: #2a2; }}
  .n-rep {{ color: #c70; }}
  .n-eq  {{ color: #888; }}
  table {{ width: 100%; border-collapse: collapse; background: #fff;
           box-shadow: 0 1px 4px rgba(0,0,0,.08); border-radius: 8px; overflow: hidden; }}
  th {{ background: #4a9eff; color: #fff; padding: 10px 12px; text-align: left; font-size: 13px; }}
  td {{ padding: 8px 12px; vertical-align: top; border-bottom: 1px solid #f0f0f0;
       line-height: 1.6; word-break: break-all; }}
  .tag {{ width: 60px; font-weight: bold; text-align: center; white-space: nowrap; }}
  tr.eq td {{ background: #fafafa; color: #888; }}
  tr.del td.del-cell {{ background: #fff0f0; }}
  tr.ins td.ins-cell {{ background: #f0fff0; }}
  tr.rep td.del-cell {{ background: #fff0f0; }}
  tr.rep td.ins-cell {{ background: #f0fff0; }}
  tr.style-only td {{ background: #fffbe6; }}
  mark.del {{ background: #ffaaaa; border-radius: 2px; }}
  mark.ins {{ background: #aaffaa; border-radius: 2px; }}
  .style-info {{ font-size: 12px; color: #888; width: 180px; }}
  .style-list {{ padding-left: 14px; }}
  .style-list li {{ margin-bottom: 2px; }}
  .legend {{ display: flex; gap: 16px; margin-bottom: 16px; font-size: 12px; flex-wrap: wrap; }}
  .legend span {{ display: flex; align-items: center; gap: 4px; }}
  .dot {{ width: 12px; height: 12px; border-radius: 2px; display: inline-block; }}
  .filter-bar {{ margin-bottom: 12px; display: flex; gap: 8px; flex-wrap: wrap; }}
  button {{ padding: 5px 12px; border: 1px solid #ddd; border-radius: 4px; cursor: pointer;
            background: #fff; font-size: 12px; transition: background .15s; }}
  button:hover {{ background: #e8f0fe; }}
  button.active {{ background: #4a9eff; color: #fff; border-color: #4a9eff; }}
</style>
</head>
<body>
<h1>📄 文档 Diff 报告</h1>
<div class="meta">
  旧文档：<b>{html_lib.escape(str(file1))}</b> &nbsp;→&nbsp; 新文档：<b>{html_lib.escape(str(file2))}</b>
</div>

<div class="stats">
  <div class="stat"><div class="n n-eq">{stats['equal']}</div><div class="l">未变化</div></div>
  <div class="stat"><div class="n n-del">{stats['delete']}</div><div class="l">删除段落</div></div>
  <div class="stat"><div class="n n-ins">{stats['insert']}</div><div class="l">新增段落</div></div>
  <div class="stat"><div class="n n-rep">{stats['replace']}</div><div class="l">修改段落</div></div>
  <div class="stat"><div class="n n-del">{img_counts['deleted']}</div><div class="l">删除图片</div></div>
  <div class="stat"><div class="n n-ins">{img_counts['added']}</div><div class="l">新增图片</div></div>
</div>

<div class="legend">
  <span><span class="dot" style="background:#ffaaaa"></span> 删除内容</span>
  <span><span class="dot" style="background:#aaffaa"></span> 新增内容</span>
  <span><span class="dot" style="background:#fffbe6"></span> 仅样式变化</span>
</div>

<div class="filter-bar">
  <b>筛选：</b>
  <button class="active" onclick="filter('all')">全部</button>
  <button onclick="filter('del')">仅删除</button>
  <button onclick="filter('ins')">仅新增</button>
  <button onclick="filter('rep')">仅修改</button>
  <button onclick="filter('style-only')">仅样式</button>
  <button onclick="filter('eq')">仅相同</button>
</div>

<h2>文本变化</h2>
<table id="main-table">
  <thead>
    <tr>
      <th>类型</th>
      <th style="width:38%">旧文档</th>
      <th style="width:38%">新文档</th>
      <th>样式变化</th>
    </tr>
  </thead>
  <tbody>
    {''.join(rows)}
  </tbody>
</table>

{img_section}

<script>
function filter(cls) {{
  document.querySelectorAll('.filter-bar button').forEach(b => b.classList.remove('active'));
  event.target.classList.add('active');
  document.querySelectorAll('#main-table tbody tr').forEach(tr => {{
    if (cls === 'all') {{ tr.style.display = ''; }}
    else {{ tr.style.display = tr.classList.contains(cls) ? '' : 'none'; }}
  }});
}}
</script>
</body>
</html>'''

    out_path.write_text(html, encoding='utf-8')


# ── 输出：纯文本报告 ──────────────────────────────────────
def build_txt(para_diffs, image_diffs, stats, img_counts, file1, file2, out_path: Path):
    lines = ['=' * 60, '文档 Diff 报告',
             f'旧文档: {file1}', f'新文档: {file2}', '=' * 60,
             f'统计：未变化 {stats["equal"]} | 删除 {stats["delete"]} | '
             f'新增 {stats["insert"]} | 修改 {stats["replace"]}',
             f'图片：删除 {img_counts["deleted"]} | 新增 {img_counts["added"]}', '']
    for d in para_diffs:
        if d.tag == 'equal':
            if d.style_changes:
                lines.append(f'[样式] {d.old_text}')
                for c in d.style_changes:
                    lines.append(f'       ↳ {c}')
            else:
                lines.append(f'  {d.old_text}')
        elif d.tag == 'delete':
            lines.append(f'[-] {d.old_text}')
        elif d.tag == 'insert':
            lines.append(f'[+] {d.new_text}')
        elif d.tag == 'replace':
            lines.append(f'[修改前] {d.old_text}')
            lines.append(f'[修改后] {d.new_text}')
            for c in d.style_changes:
                lines.append(f'  样式变化: {c}')
        lines.append('')
    for img in image_diffs:
        if img.tag == 'deleted':
            lines.append(f'[-] 图片删除 (MD5: {img.md5[:16]}…)')
        elif img.tag == 'added':
            lines.append(f'[+] 图片新增 (MD5: {img.md5[:16]}…)')
    out_path.write_text('\n'.join(lines), encoding='utf-8')


# ── 主程序 ────────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser(description='DOCX Diff：文本 + 图片 + 样式全面对比')
    parser.add_argument('--file1',     required=True)
    parser.add_argument('--file2',     required=True)
    parser.add_argument('--output',    default=None)
    parser.add_argument('--threshold', type=float, default=0.6,
                        help='段落配对相似度阈值 0~1（默认 0.6）')
    parser.add_argument('--mode',      choices=['all', 'text', 'image'], default='all')
    args = parser.parse_args()

    file1, file2 = Path(args.file1), Path(args.file2)
    for f in [file1, file2]:
        if not f.exists():
            print(f'错误：文件不存在 → {f}', file=sys.stderr); sys.exit(1)
        if f.suffix.lower() != '.docx':
            print(f'错误：仅支持 .docx → {f}', file=sys.stderr); sys.exit(1)

    out_dir = Path(args.output) if args.output else file2.parent
    out_dir.mkdir(parents=True, exist_ok=True)
    stem = file2.stem

    print('正在读取文档...')
    doc1, doc2 = Document(str(file1)), Document(str(file2))

    para_diffs, image_diffs = [], []
    stats = {'equal': 0, 'delete': 0, 'insert': 0, 'replace': 0}
    img_counts = {'deleted': 0, 'added': 0, 'unchanged': 0}

    if args.mode in ('all', 'text'):
        print('正在对比段落与样式...')
        p1 = extract_paragraphs(doc1)
        p2 = extract_paragraphs(doc2)
        para_diffs = diff_paragraphs(p1, p2, args.threshold)
        for d in para_diffs:
            stats[d.tag] += 1

    if args.mode in ('all', 'image'):
        print('正在对比图片...')
        image_diffs = diff_images(doc1, doc2)
        for img in image_diffs:
            img_counts[img.tag] += 1

    out_docx = out_dir / f'{stem}_diff.docx'
    out_html = out_dir / f'{stem}_diff.html'
    out_txt  = out_dir / f'{stem}_diff.txt'

    print(f'生成高亮文档 → {out_docx}')
    build_docx(para_diffs, image_diffs, stats, img_counts, out_docx)
    print(f'生成 HTML 报告 → {out_html}')
    build_html(para_diffs, image_diffs, stats, img_counts, file1, file2, out_html)
    print(f'生成文本报告 → {out_txt}')
    build_txt(para_diffs, image_diffs, stats, img_counts, file1, file2, out_txt)

    print()
    print('── 统计 ───────────────────────────')
    print(f'  未变化：{stats["equal"]:>4} 段')
    print(f'  删  除：{stats["delete"]:>4} 段')
    print(f'  新  增：{stats["insert"]:>4} 段')
    print(f'  修  改：{stats["replace"]:>4} 段')
    print(f'  图片删除：{img_counts["deleted"]:>3} 张')
    print(f'  图片新增：{img_counts["added"]:>3} 张')
    print('───────────────────────────────────')
    print(f'完成！输出目录：{out_dir}')


if __name__ == '__main__':
    main()
