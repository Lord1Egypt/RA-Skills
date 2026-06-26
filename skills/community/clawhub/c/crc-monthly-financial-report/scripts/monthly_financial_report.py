#!/usr/bin/env python3
"""Generate a monthly financial report from a fixed-format Excel draft."""

from __future__ import annotations

import argparse
import json
import math
import os
import re
import sys
import zipfile
from dataclasses import dataclass
from decimal import Decimal, InvalidOperation, ROUND_HALF_UP
from pathlib import Path
from typing import Any, Iterable
from urllib.parse import quote
from xml.sax.saxutils import escape

try:
    from openpyxl import load_workbook
except ImportError:  # pragma: no cover - handled by CLI error path
    load_workbook = None

try:
    import xlrd
except ImportError:  # pragma: no cover - optional .xls support
    xlrd = None

try:
    from docx import Document
    from docx.shared import RGBColor
except ImportError:  # pragma: no cover - handled by CLI error path
    Document = None
    RGBColor = None


KEY_METRICS = ("营业收入", "利润总额", "净利润")
BALANCE_ASSET_METRICS = ("资产总额", "总资产", "负债总额", "总负债", "所有者权益", "净资产")
RATIO_METRICS = ("资产负债率", "杠杆率", "ROE", "ROA", "管理费用率", "关注资产率", "不良资产率")
SECTION_NAMES = {"资产负债（亿元）", "损益（万元）", "主要财务指标"}
MONTH_RE = re.compile(r"(20\d{2})年0?([1-9]|1[0-2])月")
RED_MARK_RE = re.compile(r"(【[^】]+】)")
NEWLINE = "\n"

REASON_PLACEHOLDERS = {
    "营业收入_positive": "【营业收入增长主要由于业务规模扩大及市场需求增加影响】",
    "营业收入_negative": "【营业收入下降主要由于市场需求放缓及竞争加剧影响】",
    "利润总额_positive": "【利润总额增长主要由于业务规模扩大及盈利能力提升影响】",
    "利润总额_negative": "【利润总额同比下降主要由于业务结构变化及成本费用上升影响】",
}


class ReportError(RuntimeError):
    """A user-facing report generation error."""


@dataclass
class SheetData:
    name: str
    rows: list[list[Any]]


@dataclass
class ParsedSheet:
    sheet_name: str
    period: str
    rows: list[dict[str, str]]


def _to_decimal(value: Any) -> Decimal | None:
    if value is None:
        return None
    if isinstance(value, bool):
        return None
    if isinstance(value, (int, float)):
        if isinstance(value, float) and not math.isfinite(value):
            return None
        return Decimal(str(value))
    text = str(value).strip().replace(",", "").replace("%", "").replace("％", "")
    if not text:
        return None
    if text.lower() in {"nan", "none", "null", "n/a", "na", "-"} or text == "/":
        return None
    try:
        value = Decimal(text)
        return value if value.is_finite() else None
    except (InvalidOperation, ValueError, ArithmeticError):
        return None


def _quantize(value: Decimal | None, pattern: str = "0.01") -> Decimal | None:
    if value is None or not value.is_finite():
        return None
    try:
        return value.quantize(Decimal(pattern), rounding=ROUND_HALF_UP)
    except (InvalidOperation, ValueError, ArithmeticError):
        return None


def _fmt_pct(value: Decimal | None, places: str = "0.1") -> str:
    normalized = _quantize(value, "0.01")
    if normalized is None:
        return "未识别"
    rounded = normalized.quantize(Decimal(places), rounding=ROUND_HALF_UP)
    return f"{rounded:.1f}%"


def _fmt_yi(value: Decimal | None) -> str:
    normalized = _quantize(value)
    if normalized is None:
        return "未识别"
    return f"{normalized:.2f}亿"


def _wan_to_yi(value: Decimal | None) -> Decimal | None:
    if value is None:
        return None
    return _quantize(value / Decimal("10000"))


def _fmt_wan_to_yi(value: Decimal | None) -> str:
    return _fmt_yi(_wan_to_yi(value))


def _format_pct_decimal(value: Decimal | None) -> str:
    if value is None:
        return "未识别"
    return _fmt_pct(value * Decimal("100"))


def _format_ratio(value: Decimal | None, name: str) -> str:
    if value is None:
        return "未识别"
    if name == "杠杆率":
        normalized = _quantize(value)
        return f"{normalized:.2f}" if normalized is not None else "未识别"
    normalized = _quantize(value * Decimal("100"))
    return f"{normalized:.2f}%" if normalized is not None else "未识别"


def _section_unit(section: str) -> str:
    return "yi" if section in {"资产负债（亿元）", "主要财务指标"} else "wan"


def _format_amount(value: Decimal | None, section: str) -> str:
    if value is None:
        return "未识别"
    return _fmt_yi(value) if _section_unit(section) == "yi" else _fmt_wan_to_yi(value)


def _format_delta_yi(current: Decimal | None, base: Decimal | None, section: str) -> str:
    if current is None or base is None:
        return "未识别"
    delta = abs(current - base)
    return _fmt_yi(delta) if _section_unit(section) == "yi" else _fmt_wan_to_yi(delta)


def _format_delta_direction(
    current: Decimal | None,
    base: Decimal | None,
    increase_word: str = "增加",
    decrease_word: str = "减少",
) -> str:
    if current is None or base is None:
        return ""
    return increase_word if current >= base else decrease_word


def _cell_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value).strip()


def _trim_leading_blank_cells(cells: Iterable[Any]) -> list[str]:
    values = [_cell_text(cell) for cell in cells]
    while values and not values[0].strip():
        values = values[1:]
    while values and not values[-1].strip():
        values = values[:-1]
    return values


def _load_xlsx(path: Path) -> list[SheetData]:
    if load_workbook is None:
        raise ReportError("缺少 openpyxl，请先执行：pip install openpyxl")
    workbook = load_workbook(path, data_only=True, read_only=True)
    sheets: list[SheetData] = []
    for worksheet in workbook.worksheets:
        rows = [list(row) for row in worksheet.iter_rows(values_only=True)]
        sheets.append(SheetData(worksheet.title, rows))
    return sheets


def _load_xls(path: Path) -> list[SheetData]:
    if xlrd is None:
        raise ReportError("读取 .xls 需要 xlrd，请先执行：pip install xlrd")
    workbook = xlrd.open_workbook(str(path))
    sheets: list[SheetData] = []
    for worksheet in workbook.sheets():
        rows = [worksheet.row_values(i) for i in range(worksheet.nrows)]
        sheets.append(SheetData(worksheet.name, rows))
    return sheets


def load_excel(path: Path) -> list[SheetData]:
    if not path.exists():
        raise ReportError(f"找不到输入文件：{path}")
    suffix = path.suffix.lower()
    if suffix == ".xlsx":
        return _load_xlsx(path)
    if suffix == ".xls":
        return _load_xls(path)
    raise ReportError("仅支持 .xlsx / .xls 月度 Excel 底稿")


def _detect_period_from_rows(sheet: SheetData) -> tuple[int, int] | None:
    candidates: list[tuple[int, int]] = []
    for source in [sheet.name] + [_cell_text(cell) for row in sheet.rows[:5] for cell in row]:
        for match in MONTH_RE.finditer(source):
            candidates.append((int(match.group(1)), int(match.group(2))))
    return max(candidates) if candidates else None


def _format_period(year_month: tuple[int, int] | None, fallback: str = "当期") -> str:
    if not year_month:
        return fallback
    year, month = year_month
    return f"{year}年{month}月"


def _sheet_has_monthly_table(sheet: SheetData) -> bool:
    flat = [_cell_text(cell) for row in sheet.rows[:8] for cell in row]
    return "本月" in flat and "上月" in flat and any(name in flat for name in SECTION_NAMES)


def choose_sheet(sheets: list[SheetData], sheet_name: str = "", report_period: str = "") -> SheetData:
    if sheet_name:
        for sheet in sheets:
            if sheet.name == sheet_name:
                return sheet
        raise ReportError(f"找不到指定工作表：{sheet_name}")

    period_match = MONTH_RE.search(report_period or "")
    if period_match:
        expected = (int(period_match.group(1)), int(period_match.group(2)))
        for sheet in sheets:
            if _detect_period_from_rows(sheet) == expected:
                return sheet

    monthly_sheets = [sheet for sheet in sheets if _sheet_has_monthly_table(sheet)]
    if not monthly_sheets:
        raise ReportError("未识别出固定格式的月度财务数据工作表")
    return max(monthly_sheets, key=lambda sheet: _detect_period_from_rows(sheet) or (0, 0))


def parse_sheet(sheet: SheetData, report_period: str = "") -> ParsedSheet:
    period = report_period.strip() or _format_period(_detect_period_from_rows(sheet))
    parsed_rows: list[dict[str, str]] = []
    current_section = ""

    for raw_row in sheet.rows:
        cells = _trim_leading_blank_cells(raw_row)
        if not cells:
            continue
        name = cells[0]
        if name in SECTION_NAMES:
            current_section = name
            continue
        if name in {"项目=", "项目", "本月", "月度", "累计至本月"}:
            continue
        if len(cells) < 11 or not current_section:
            continue
        parsed_rows.append(
            {
                "section": current_section,
                "indicator": name,
                "current": cells[1],
                "previous": cells[2],
                "mom_rate": cells[3],
                "yoy_month": cells[4],
                "single_yoy_rate": cells[5],
                "ytd": cells[6],
                "ytd_compare": cells[7],
                "ytd_growth": cells[8],
                "target": cells[9],
                "budget_rate": cells[10],
            }
        )
    if not parsed_rows:
        raise ReportError(f"工作表 {sheet.name} 中没有识别到月度财务指标明细")
    return ParsedSheet(sheet.name, period, parsed_rows)


def build_metrics(rows: list[dict[str, str]]) -> dict[str, dict[str, str]]:
    metrics: dict[str, dict[str, str]] = {}
    for row in rows:
        section = row["section"]
        indicator = row["indicator"]
        current = _to_decimal(row["current"])
        previous = _to_decimal(row["previous"])
        yoy_month = _to_decimal(row["yoy_month"])
        mom_rate = _to_decimal(row["mom_rate"])
        single_yoy_rate = _to_decimal(row["single_yoy_rate"])
        ytd = _to_decimal(row["ytd"])
        ytd_compare = _to_decimal(row["ytd_compare"])
        ytd_growth = _to_decimal(row["ytd_growth"])
        budget_rate = _to_decimal(row["budget_rate"])

        if indicator in KEY_METRICS:
            metrics[indicator] = {
                "current_text": _format_amount(current, section),
                "ytd_text": _format_amount(ytd, section),
                "mom_growth_text": _format_pct_decimal(mom_rate),
                "single_yoy_growth_text": _format_pct_decimal(single_yoy_rate),
                "ytd_growth_text": _format_pct_decimal(ytd_growth),
                "budget_rate_text": _format_pct_decimal(budget_rate) if budget_rate is not None else "",
                "mom_delta_text": _format_delta_yi(current, previous, section),
                "single_yoy_delta_text": _format_delta_yi(current, yoy_month, section),
                "mom_delta_direction": "增加" if current is not None and previous is not None and current >= previous else "减少",
                "single_yoy_delta_direction": "增加" if current is not None and yoy_month is not None and current >= yoy_month else "减少",
            }
        elif indicator in BALANCE_ASSET_METRICS:
            normalized_indicator = {
                "资产总额": "总资产",
                "总资产": "总资产",
                "负债总额": "总负债",
                "总负债": "总负债",
                "所有者权益": "所有者权益",
                "净资产": "所有者权益",
            }[indicator]
            metrics[normalized_indicator] = {
                "current_text": _format_amount(current, section),
                "compare_delta_text": _format_delta_yi(current, ytd_compare, section),
                "compare_delta_direction": _format_delta_direction(
                    current,
                    ytd_compare,
                    "增长" if normalized_indicator == "所有者权益" else "增加",
                    "减少",
                ),
            }
        elif indicator == "经营活动现金流":
            metrics[indicator] = {
                "current_text": _format_amount(current, section),
                "single_yoy_growth_text": _format_pct_decimal(single_yoy_rate),
            }
        elif indicator in RATIO_METRICS:
            source_value = ytd if ytd is not None else current
            metrics[indicator] = {"current_text": _format_ratio(source_value, indicator)}
    if not metrics:
        raise ReportError("未能从底稿中提取核心财务指标")
    return metrics


def _growth_verb(label: str, pct_text: str) -> str:
    if pct_text in {"未识别", ""}:
        return f"{label}未识别"
    if pct_text.startswith("-"):
        anti = "减少" if "环比" in label else "下降"
        return f"{label.replace('增长', anti)}{pct_text.lstrip('-')}"
    return f"{label}{pct_text}"


def _delta_verb(prefix: str, direction: str, delta_text: str) -> str:
    if delta_text in {"未识别", ""}:
        return f"{prefix}{direction}未识别"
    return f"{prefix}{direction}{delta_text}"


def _build_metric_line(index_text: str, name: str, item: dict[str, str]) -> str:
    first_sentence = f"{index_text}、累计{name}为{item['ytd_text']}"
    if item["ytd_growth_text"] != "未识别":
        first_sentence += f"，{_growth_verb('同比增长', item['ytd_growth_text'])}"
    if item["budget_rate_text"]:
        first_sentence += f"，预算完成率为{item['budget_rate_text']}"
    if name in {"营业收入", "利润总额"}:
        suffix = "positive" if item["ytd_growth_text"] not in {"未识别", ""} and not item["ytd_growth_text"].startswith("-") else "negative"
        first_sentence += f"，{REASON_PLACEHOLDERS.get(name + '_' + suffix, f'【{name}变动原因待补充】')}"
    first_sentence += "。"

    if name == "营业收入":
        second_sentence = (
            f"单月{name}为{item['current_text']}，"
            f"{_growth_verb('环比增长', item['mom_growth_text'])}，"
            f"{_growth_verb('同比增长', item['single_yoy_growth_text'])}。"
        )
    else:
        second_sentence = (
            f"单月{name}为{item['current_text']}，"
            f"{_delta_verb('环比', item['mom_delta_direction'], item['mom_delta_text'])}，"
            f"{_delta_verb('同比', item['single_yoy_delta_direction'], item['single_yoy_delta_text'])}。"
        )
    return first_sentence + second_sentence


def _build_balance_line(metrics: dict[str, dict[str, str]]) -> str:
    parts: list[str] = []
    for source_name, display_name in (("总资产", "总资产"), ("总负债", "总负债"), ("所有者权益", "所有者权益")):
        item = metrics.get(source_name)
        if item:
            default_verb = "增长" if display_name == "所有者权益" else "增加"
            verb = item.get("compare_delta_direction") or default_verb
            parts.append(f"{display_name}为{item['current_text']}，较年初{verb}{item['compare_delta_text']}。")

    ratio_parts: list[str] = []
    for ratio_name in RATIO_METRICS:
        item = metrics.get(ratio_name)
        if item:
            ratio_parts.append(f"{ratio_name}为{item['current_text']}")

    line = "四、" + "".join(parts)
    if ratio_parts:
        line += "；".join(ratio_parts[:4])
        if len(ratio_parts) > 4:
            line += "，" + "，".join(ratio_parts[4:])
        line += "。"

    cash = metrics.get("经营活动现金流")
    if cash:
        line += f"经营活动现金流为{cash['current_text']}，同比增长{cash['single_yoy_growth_text']}。"
    return line


def build_report(period: str, metrics: dict[str, dict[str, str]]) -> str:
    lines = [f"{period}末财务指标汇报如下：", ""]
    numerals = ["一", "二", "三"]
    for index, metric_name in enumerate(KEY_METRICS):
        item = metrics.get(metric_name)
        if item:
            lines.append(_build_metric_line(numerals[index], metric_name, item))
    lines.append(_build_balance_line(metrics))
    return NEWLINE.join(lines).strip()


def _safe_filename(value: str) -> str:
    value = re.sub(r'[<>:"/\\|?*]+', "_", value).strip()
    return value or "monthly_financial_report"


def write_markdown(path: Path, report_text: str) -> None:
    path.write_text(report_text + NEWLINE, encoding="utf-8")


def write_docx(path: Path, report_text: str) -> None:
    if Document is None or RGBColor is None:
        write_basic_docx(path, report_text)
        return
    document = Document()
    for line in report_text.splitlines():
        paragraph = document.add_paragraph()
        if not line:
            continue
        for part in RED_MARK_RE.split(line):
            if not part:
                continue
            run = paragraph.add_run(part)
            if RED_MARK_RE.fullmatch(part):
                run.font.color.rgb = RGBColor(212, 56, 13)
    document.save(path)


def _w_text(value: str) -> str:
    return escape(value, {'"': "&quot;", "'": "&apos;"})


def _docx_run(text: str, red: bool = False) -> str:
    props = ""
    if red:
        props = '<w:rPr><w:color w:val="D4380D"/></w:rPr>'
    return f'<w:r>{props}<w:t xml:space="preserve">{_w_text(text)}</w:t></w:r>'


def _docx_paragraph(line: str) -> str:
    runs: list[str] = []
    if line:
        for part in RED_MARK_RE.split(line):
            if not part:
                continue
            runs.append(_docx_run(part, bool(RED_MARK_RE.fullmatch(part))))
    return "<w:p>" + "".join(runs) + "</w:p>"


def write_basic_docx(path: Path, report_text: str) -> None:
    """Write a minimal Word document without third-party docx libraries."""
    paragraphs = "".join(_docx_paragraph(line) for line in report_text.splitlines())
    document_xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    {paragraphs}
    <w:sectPr>
      <w:pgSz w:w="11906" w:h="16838"/>
      <w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440" w:header="720" w:footer="720" w:gutter="0"/>
    </w:sectPr>
  </w:body>
</w:document>
'''
    content_types = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>
'''
    rels = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>
'''
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", content_types)
        archive.writestr("_rels/.rels", rels)
        archive.writestr("word/document.xml", document_xml)


def _platform_file_path(path: Path) -> str:
    text = str(path).replace("\\", "/")
    home_workspace = "/home/node/workspace"
    if text == "/workspace" or text.startswith("/workspace/"):
        return text
    if text == home_workspace:
        return "/workspace"
    if text.startswith(f"{home_workspace}/"):
        return f"/workspace{text[len(home_workspace):]}"
    return ""


def _build_word_download_url(path: Path, instance_id: str = "", download_base_url: str = "") -> str:
    instance_id = (instance_id or os.environ.get("CLAWMATE_INSTANCE_ID", "")).strip()
    if not instance_id:
        return ""
    base_url = (
        download_base_url
        or os.environ.get("CLAWMATE_PUBLIC_ORIGIN", "")
        or os.environ.get("OPENCLAW_PUBLIC_ORIGIN", "")
    ).strip().rstrip("/")
    download_path = _platform_file_path(path)
    if not download_path:
        return ""
    api_path = f"/api/instances/{quote(instance_id, safe='')}/files/download?path={quote(download_path, safe='')}"
    return f"{base_url}{api_path}" if base_url else api_path


def generate_report(
    input_path: Path,
    output_dir: Path,
    report_period: str = "",
    sheet_name: str = "",
    output_stem: str = "",
    instance_id: str = "",
    download_base_url: str = "",
) -> dict[str, Any]:
    sheets = load_excel(input_path)
    sheet = choose_sheet(sheets, sheet_name=sheet_name, report_period=report_period)
    parsed = parse_sheet(sheet, report_period=report_period)
    metrics = build_metrics(parsed.rows)
    report_text = build_report(parsed.period, metrics)

    output_dir.mkdir(parents=True, exist_ok=True)
    stem = _safe_filename(output_stem or f"月度财务报告_{parsed.period}")
    markdown_path = output_dir / f"{stem}.md"
    docx_path = output_dir / f"{stem}.docx"
    json_path = output_dir / f"{stem}.json"

    write_markdown(markdown_path, report_text)
    write_docx(docx_path, report_text)
    word_download_url = _build_word_download_url(
        docx_path,
        instance_id=instance_id,
        download_base_url=download_base_url,
    )

    summary = {
        "ok": True,
        "input": str(input_path),
        "sheet": parsed.sheet_name,
        "period": parsed.period,
        "output_markdown": str(markdown_path),
        "output_docx": str(docx_path),
        "output_json": str(json_path),
        "metrics_found": sorted(metrics.keys()),
        "report_text": report_text,
    }
    if word_download_url:
        summary["word_download_url"] = word_download_url
        summary["word_download_markdown"] = f"[下载 Word 报告]({word_download_url})"
    json_path.write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    return summary


def parse_args(argv: list[str]) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate a monthly financial report and Word file from an Excel draft.")
    parser.add_argument("input_excel", help="月度 Excel 底稿路径，支持 .xlsx；.xls 需安装 xlrd")
    parser.add_argument("-o", "--output-dir", default="", help="输出目录；默认使用输入文件同级目录下的 output")
    parser.add_argument("--period", default="", help="报告期间，例如 2026年5月；不填则自动从工作表标题识别")
    parser.add_argument("--sheet", default="", help="指定工作表名；不填则自动选择最新月度工作表")
    parser.add_argument("--output-stem", default="", help="输出文件名前缀；不填则使用 月度财务报告_<期间>")
    parser.add_argument("--instance-id", default="", help="平台实例 ID；默认读取 CLAWMATE_INSTANCE_ID，用于生成 Word 下载链接")
    parser.add_argument("--download-base-url", default="", help="平台外部访问地址；默认读取 CLAWMATE_PUBLIC_ORIGIN，用于生成 Word 下载链接")
    return parser.parse_args(argv)


def main(argv: list[str] | None = None) -> int:
    args = parse_args(argv or sys.argv[1:])
    input_path = Path(args.input_excel).expanduser().resolve()
    output_dir = Path(args.output_dir).expanduser().resolve() if args.output_dir else input_path.parent / "output"
    try:
        summary = generate_report(
            input_path=input_path,
            output_dir=output_dir,
            report_period=args.period,
            sheet_name=args.sheet,
            output_stem=args.output_stem,
            instance_id=args.instance_id,
            download_base_url=args.download_base_url,
        )
    except ReportError as exc:
        print(json.dumps({"ok": False, "error": str(exc)}, ensure_ascii=False), file=sys.stderr)
        return 2
    except Exception as exc:  # pragma: no cover - final CLI guard
        print(json.dumps({"ok": False, "error": f"处理失败：{exc}"}, ensure_ascii=False), file=sys.stderr)
        return 1

    print(json.dumps(summary, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
