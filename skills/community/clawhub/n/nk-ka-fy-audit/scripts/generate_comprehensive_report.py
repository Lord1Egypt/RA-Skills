#!/usr/bin/env python3
"""
KA返佣审计 - 完整报告生成器（2026-05-11 更新版）
生成Word报告 + Excel明细报告
满足2026-05-11最新报告输出要求（9条修改全部落实）

变更记录：
  2026-05-11: 
  1. 审计结论陈述改为"经审计集团客户部...发现疑似多支付返佣...待审计人员进一步核实"
  2. 核心发现下列出所有有差异的月份及差异金额
  3. 删除审计范围、无差异月份等描述
  4. 删除"已核查历史差额调整"相关内容
  5. 删除"月度返佣明细表"
  6. TOP10保留累计+当月
  7. "提请重点关注事项"下增加所有真实差异明细
  8. 删除"SQL手工核算"计算方式
  9. 差异判定标准改为：≤1元通过，>1元差异；不体现"历史差额调整"

Word: 集团客户部FY审计报告（202307~202604）.docx
Excel: 集团客户部FY审计结果明细（202307~202604）.xlsx
"""
import os, sys, sqlite3
from datetime import datetime
from collections import defaultdict

import openpyxl
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import matplotlib.font_manager as fm

# Font setup - use WenQuanYi Micro Hei
WQY_PATH = '/workspace/.fonts/wqy-microhei.ttc'
font_name = 'WenQuanYi Micro Hei'
if os.path.exists(WQY_PATH):
    try:
        fm.fontManager.addfont(WQY_PATH)
        plt.rcParams['font.sans-serif'] = [font_name, 'DejaVu Sans']
        plt.rcParams['axes.unicode_minus'] = False
        print(f"Using font: {font_name}")
    except Exception as e:
        print(f"Font load error: {e}")
        plt.rcParams['font.sans-serif'] = ['DejaVu Sans']
        plt.rcParams['axes.unicode_minus'] = False
else:
    plt.rcParams['font.sans-serif'] = ['DejaVu Sans']
    plt.rcParams['axes.unicode_minus'] = False

# ========== Configuration ==========
DB_PATH = '/workspace/业务数据/集团客户部FY数据.db'
TEMPLATE_PATH = '/workspace/templates/返佣审计结果-模板.xlsx'
OUT_DIR = '/workspace/output'
os.makedirs(OUT_DIR, exist_ok=True)

AUDIT_PERIOD = '202307~202604'
LABEL_PERIOD = '2023年07月 至 2026年04月'
MONTH_COUNT = 34

WORD_OUTPUT = os.path.join(OUT_DIR, f'集团客户部FY审计报告（{AUDIT_PERIOD}）.docx')
EXCEL_OUTPUT = os.path.join(OUT_DIR, f'集团客户部FY审计结果明细（{AUDIT_PERIOD}）.xlsx')

# Wiki placeholders
WIKI_NODE_URL = 'https://sqb.feishu.cn/wiki/QguJwqZWyiQUzAkKjq2cMZwInwb'

# Color constants
DARK_BLUE = RGBColor(0x1A, 0x47, 0x7A)
BLUE = RGBColor(0x2B, 0x6C, 0xB5)
RED = RGBColor(0xCC, 0x33, 0x33)
GREEN = RGBColor(0x33, 0x99, 0x33)
GRAY = RGBColor(0x66, 0x66, 0x66)

# ========== Data Loading ==========
conn = sqlite3.connect(DB_PATH)
c = conn.cursor()

c.execute('SELECT * FROM monthly_audit ORDER BY month')
monthly_rows = c.fetchall()
MONTHLY = {}
for r in monthly_rows:
    MONTHLY[r[0]] = {
        'recalc': r[1], 'summary': r[2],
        'diff': r[3], 'real_diff': r[4], 'brands': r[5]
    }
all_months = sorted(MONTHLY.keys())

c.execute('SELECT month, brand, raw_amount, recalc_amount, summary_amount, diff, is_real_diff, diff_note FROM brand_audit ORDER BY month, brand')
brand_rows = c.fetchall()

brand_series = defaultdict(list)
for r in brand_rows:
    brand_series[r[1]].append((r[0], r[4], r[2]))

# High-growth brands
high_growth = []
for brand, series_data in brand_series.items():
    sorted_data = sorted(series_data, key=lambda x: x[0])
    for i in range(1, len(sorted_data)):
        prev_m, prev_s, prev_raw = sorted_data[i-1]
        curr_m, curr_s, curr_raw = sorted_data[i]
        if prev_s and prev_s > 0:
            growth_pct = (curr_s - prev_s) / prev_s * 100
            if growth_pct >= 50 and curr_s >= 10000:
                high_growth.append((curr_m, brand, prev_s, curr_s, round(growth_pct, 1)))
high_growth.sort(key=lambda x: -x[4])

# Top brands (cumulative)
c.execute('SELECT brand, ROUND(SUM(COALESCE(summary_amount,0)),2) as total FROM brand_audit GROUP BY brand ORDER BY total DESC')
top_brands = c.fetchall()

# Real differences: only diff > 1 (not "历史差额调整")
real_diffs = [(r[0], r[1], r[5], r[6], r[7]) for r in brand_rows if r[6] == 1 and r[5] > 1]
real_diffs.sort(key=lambda x: -abs(x[2]))

# Real diff detail - use brand-level data since no merchant_sn available
c.execute("""
    SELECT month, brand, raw_amount, recalc_amount, summary_amount, diff, diff_note
    FROM brand_audit
    WHERE is_real_diff=1 AND diff > 1
    ORDER BY month, brand
""")
brand_detail_rows = c.fetchall()

monthly_totals = {m: MONTHLY[m]['summary'] for m in all_months}
monthly_brands = {m: MONTHLY[m]['brands'] for m in all_months}

# Summary stats
total_real_diff = sum(d[2] for d in real_diffs)  # all real diffs > 1
diff_months = set(d[0] for d in real_diffs)
pass_months = MONTH_COUNT - len(diff_months)

conn.close()

print(f"审计数据总览:")
print(f"  月份: {MONTH_COUNT} ({all_months[0]}~{all_months[-1]})")
print(f"  品牌: {len(top_brands)}")
print(f"  真实差异品牌-月份: {len(real_diffs)} 条, 涉及 {len(diff_months)} 个月")
print(f"  总真实差异: {total_real_diff:+,.2f} 元")

# ========== Charts ==========
def generate_trend_chart():
    fig, ax = plt.subplots(figsize=(12, 4.5))
    months = list(monthly_totals.keys())
    amounts = [monthly_totals[m] for m in months]
    labels = [m[2:] for m in months]
    x = range(len(months))
    bars = ax.bar(x, amounts, color='#2B6CB5', alpha=0.8, width=0.6, zorder=3)
    for i, (bar, val) in enumerate(zip(bars, amounts)):
        if val > 0:
            ax.text(bar.get_x() + bar.get_width()/2., bar.get_height() + 15000,
                    f'{val/10000:.1f}万', ha='center', va='bottom', fontsize=6, color='#333333', rotation=90)
    ax.set_xlabel('月份', fontsize=10)
    ax.set_ylabel('返佣金额（万元）', fontsize=10)
    ax.set_xticks(x)
    ax.set_xticklabels(labels, fontsize=7, rotation=45)
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v, _: f'{v/10000:.0f}'))
    ax.yaxis.set_major_locator(mticker.MaxNLocator(6))
    ax.grid(axis='y', alpha=0.3, zorder=0)
    ax.set_axisbelow(True)
    avg_val = sum(amounts) / len(amounts)
    ax.axhline(y=avg_val, color='#E53E3E', linestyle='--', linewidth=0.8, alpha=0.6, zorder=4)
    ax.text(len(months)-1, avg_val + 10000, f'均值 {avg_val/10000:.1f}万',
            fontsize=7, color='#E53E3E', ha='right')
    plt.tight_layout()
    chart_path = os.path.join(OUT_DIR, 'trend_chart.png')
    plt.savefig(chart_path, dpi=200, bbox_inches='tight')
    plt.close()
    return chart_path

def generate_top_brands_chart():
    fig, ax = plt.subplots(figsize=(10, 4))
    top10 = top_brands[:10]
    brands = [t[0] for t in top10]
    amounts = [t[1] for t in top10]
    short_brands = [b[:8] + '...' if len(b) > 8 else b for b in brands]
    colors = ['#2B6CB5', '#3182CE', '#4299E1', '#63B3ED', '#90CDF4',
              '#E53E3E', '#DD6B20', '#D69E2E', '#38A169', '#319795']
    bars = ax.barh(range(len(short_brands)), amounts, color=colors[:len(short_brands)], height=0.6)
    ax.set_yticks(range(len(short_brands)))
    ax.set_yticklabels(short_brands, fontsize=8)
    ax.set_xlabel('返佣总额（万元）', fontsize=9)
    ax.xaxis.set_major_formatter(mticker.FuncFormatter(lambda v, _: f'{v/10000:.0f}'))
    ax.invert_yaxis()
    for i, (bar, val) in enumerate(zip(bars, amounts)):
        ax.text(bar.get_width() + 10000, bar.get_y() + bar.get_height()/2.,
                f'{val/10000:.2f}万', ha='left', va='center', fontsize=7, color='#333333')
    ax.grid(axis='x', alpha=0.3)
    ax.set_axisbelow(True)
    plt.tight_layout()
    chart_path = os.path.join(OUT_DIR, 'top_brands_chart.png')
    plt.savefig(chart_path, dpi=200, bbox_inches='tight')
    plt.close()
    return chart_path

print("生成图表...")
trend_chart = generate_trend_chart()
top_brands_chart = generate_top_brands_chart()


# ====================================================================
# WORD REPORT
# ====================================================================
print("生成Word报告...")
doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '2B6CB5')
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

# --- Cover Page ---
for _ in range(6):
    doc.add_paragraph()
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('集团客户部FY审计报告')
run.font.size = Pt(26)
run.font.color.rgb = DARK_BLUE
run.bold = True

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub_p.add_run(f'审计期间：{LABEL_PERIOD}')
run.font.size = Pt(14)
run.font.color.rgb = GRAY

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run(f'报告日期：{datetime.now().strftime("%Y年%m月%d日")}')
run.font.size = Pt(11)
run.font.color.rgb = GRAY

doc.add_paragraph()
dept_p = doc.add_paragraph()
dept_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = dept_p.add_run('内部控制与审计部')
run.font.size = Pt(12)
run.font.color.rgb = BLUE

doc.add_page_break()

# ===== Section 1: 审计结论 =====
doc.add_heading('一、审计结论', level=1)

# 1.1 Opening statement (user's exact wording)
p = doc.add_paragraph()
run = p.add_run(
    f'经审计集团客户部2023年7月至2026年4月返佣计算表，发现疑似多支付返佣 {total_real_diff:+,.2f} 元'
    f'（实际支付金额高于审计重算的应发金额），待审计人员进一步核实。'
)
run.font.color.rgb = RED
run.bold = True
run.font.size = Pt(11)

doc.add_paragraph()

# 1.2 Core findings: list ALL months/品牌 with differences
doc.add_heading('核心发现', level=2)
p = doc.add_paragraph(f'审计发现以下{len(diff_months)}个月份共{len(real_diffs)}条差异记录：')

# Create table of all diff months
diff_table = doc.add_table(rows=len(real_diffs) + 1, cols=4, style='Light Grid Accent 1')
diff_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['月份', '品牌', '差异金额（元）', '差异原因']):
    diff_table.cell(0, i).text = h
    for p2 in diff_table.cell(0, i).paragraphs:
        for r2 in p2.runs:
            r2.bold = True
            r2.font.size = Pt(9)
for i, (m, brand, diff_amt, is_real, note) in enumerate(real_diffs):
    diff_table.cell(i+1, 0).text = m
    diff_table.cell(i+1, 1).text = brand
    diff_table.cell(i+1, 2).text = f'{diff_amt:+,.2f}'
    # Determine likely cause
    if brand == '中快餐饮':
        cause = '中快渠道间连返佣纳入差异'
    elif diff_amt > 1000:
        cause = '待核实具体原因'
    else:
        cause = ''
    diff_table.cell(i+1, 3).text = cause
    for p2 in diff_table.cell(i+1, 2).paragraphs:
        for r2 in p2.runs:
            r2.font.color.rgb = RED

doc.add_paragraph()

# ===== Section 2: 数据分析 =====
doc.add_heading('二、数据分析', level=1)

doc.add_heading('2.1 月度返佣趋势', level=2)
total_summary = sum(monthly_totals.values())
avg_monthly = total_summary / len(monthly_totals) if monthly_totals else 0
p = doc.add_paragraph()
run = p.add_run(
    f'审计期间{MONTH_COUNT}个月累计返佣金额约{total_summary/10000:.0f}万元，'
    f'月均返佣约{avg_monthly/10000:.1f}万元。'
)

doc.add_picture(trend_chart, width=Inches(6))

doc.add_page_break()

# 2.2 Top brands - cumulative
latest_m = all_months[-1]
c2 = sqlite3.connect(DB_PATH)
c_cur = c2.cursor()
c_cur.execute('SELECT brand, summary_amount FROM brand_audit WHERE month=? ORDER BY summary_amount DESC', (latest_m,))
latest_brands = c_cur.fetchall()
latest_total = sum(b[1] for b in latest_brands)
c2.close()

doc.add_heading(f'2.2 主要返佣品牌TOP10（累计 {AUDIT_PERIOD}）', level=2)
top10_total_cum = sum(t[1] for t in top_brands[:10])
grand_total = sum(t[1] for t in top_brands)
pct_cum = top10_total_cum / grand_total * 100 if grand_total else 0
p = doc.add_paragraph()
run = p.add_run(f'TOP10品牌合计返佣约{top10_total_cum/10000:.0f}万元，占全部品牌返佣总额的{pct_cum:.1f}%。')

doc.add_picture(top_brands_chart, width=Inches(5.5))

doc.add_paragraph()
bt = doc.add_table(rows=11, cols=3, style='Light Grid Accent 1')
for i, h in enumerate(['排名', '品牌', '返佣总额（元）']):
    bt.cell(0, i).text = h
    for p in bt.cell(0, i).paragraphs:
        for r in p.runs:
            r.bold = True
for i, (brand, total) in enumerate(top_brands[:10]):
    bt.cell(i+1, 0).text = str(i+1)
    bt.cell(i+1, 1).text = brand
    bt.cell(i+1, 2).text = f'{total:,.2f}'

doc.add_paragraph()

# 2.3 Top brands - current month
doc.add_heading(f'2.3 主要返佣品牌TOP10（{latest_m}月度）', level=2)
fig2, ax2 = plt.subplots(figsize=(10, 4))
l_top10 = latest_brands[:10]
l_brands = [b[0] for b in l_top10]
l_amounts = [b[1] for b in l_top10]
l_short = [b[:8] + '...' if len(b) > 8 else b for b in l_brands]
colors2 = ['#2B6CB5', '#3182CE', '#4299E1', '#63B3ED', '#90CDF4',
           '#E53E3E', '#DD6B20', '#D69E2E', '#38A169', '#319795']
bars2 = ax2.barh(range(len(l_short)), l_amounts, color=colors2[:len(l_short)], height=0.6)
ax2.set_yticks(range(len(l_short)))
ax2.set_yticklabels(l_short, fontsize=8)
ax2.set_xlabel('返佣金额（万元）', fontsize=9)
ax2.xaxis.set_major_formatter(mticker.FuncFormatter(lambda v, _: f'{v/10000:.0f}'))
ax2.invert_yaxis()
for i, (bar, val) in enumerate(zip(bars2, l_amounts)):
    ax2.text(bar.get_width() + 500, bar.get_y() + bar.get_height()/2.,
            f'{val/10000:.2f}万', ha='left', va='center', fontsize=7, color='#333333')
ax2.grid(axis='x', alpha=0.3)
ax2.set_axisbelow(True)
plt.tight_layout()
latest_chart_path = os.path.join(OUT_DIR, 'latest_top_brands_chart.png')
plt.savefig(latest_chart_path, dpi=200, bbox_inches='tight')
plt.close()

l_total = sum(l_amounts)
l_pct = l_total / latest_total * 100 if latest_total else 0
p = doc.add_paragraph()
run = p.add_run(f'{latest_m}月TOP10品牌合计返佣约{l_total/10000:.1f}万元，占当月全部品牌的{l_pct:.1f}%。')

doc.add_picture(latest_chart_path, width=Inches(5.5))

doc.add_paragraph()
bt2 = doc.add_table(rows=11, cols=4, style='Light Grid Accent 1')
for i, h in enumerate(['排名', '品牌', f'{latest_m}月返佣（元）', '占当月比例']):
    bt2.cell(0, i).text = h
    for p in bt2.cell(0, i).paragraphs:
        for r in p.runs:
            r.bold = True
for i, (brand, amt) in enumerate(l_top10):
    bt2.cell(i+1, 0).text = str(i+1)
    bt2.cell(i+1, 1).text = brand
    bt2.cell(i+1, 2).text = f'{amt:,.2f}'
    pct2 = amt / latest_total * 100
    bt2.cell(i+1, 3).text = f'{pct2:.1f}%'

doc.add_page_break()

# ===== Section 3: 提请审计人员重点核实事项 =====
doc.add_heading('三、提请审计人员重点核实事项', level=1)

# 3.1 待核实的疑似差异
doc.add_heading('3.1 待核实的疑似差异', level=2)
if brand_detail_rows:
    p = doc.add_paragraph(f'审计发现以下{len(brand_detail_rows)}个品牌-月份存在疑似多支付返佣，明细如下：')
    
    detail_table = doc.add_table(rows=len(brand_detail_rows) + 1, cols=4, style='Light Grid Accent 1')
    detail_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['月份', '品牌', '审计重算金额（元）', '汇总表金额（元）']):
        detail_table.cell(0, i).text = h
        for p2 in detail_table.cell(0, i).paragraphs:
            for r2 in p2.runs:
                r2.bold = True
                r2.font.size = Pt(8)
    
    for i, rd in enumerate(brand_detail_rows):
        m, brand, raw_amt, recalc_amt, summary_amt, diff, note = rd
        detail_table.cell(i+1, 0).text = m
        detail_table.cell(i+1, 1).text = brand
        detail_table.cell(i+1, 2).text = f'{recalc_amt:,.2f}'
        detail_table.cell(i+1, 3).text = f'{summary_amt:,.2f}'
        for p2 in detail_table.cell(i+1, 3).paragraphs:
            for r2 in p2.runs:
                r2.font.color.rgb = RED
                r2.bold = True
                
else:
    doc.add_paragraph('审计期间未发现需要重点核实的差异。')

doc.add_paragraph()

# 3.2 返佣规则变动（原3.1）
doc.add_heading('3.2 返佣规则变动', level=2)
try:
    twb = openpyxl.load_workbook(TEMPLATE_PATH, data_only=True)
    tws = twb['返佣规则变动']
    rule_rows = []
    for tr in range(2, tws.max_row + 1):
        tm = tws.cell(tr, 1).value
        if tm and str(tm).strip() == latest_m:
            rule_rows.append([
                str(tws.cell(tr, c).value or '') for c in range(2, 7)
            ])
    twb.close()
    if rule_rows:
        p = doc.add_paragraph()
        run = p.add_run(f'{latest_m}月共有{len(rule_rows)}条返佣规则变动，明细如下：')
        rt = doc.add_table(rows=len(rule_rows) + 1, cols=5, style='Light Grid Accent 1')
        for ri, h in enumerate(['品牌', '返佣项目', '原返佣比例', '现返佣比例', '备注']):
            rt.cell(0, ri).text = h
            for p2 in rt.cell(0, ri).paragraphs:
                for r2 in p2.runs:
                    r2.bold = True
                    r2.font.size = Pt(8)
        for ri, row in enumerate(rule_rows):
            for ci, v in enumerate(row):
                rt.cell(ri+1, ci).text = v
                for p2 in rt.cell(ri+1, ci).paragraphs:
                    for r2 in p2.runs:
                        r2.font.size = Pt(8)
    else:
        doc.add_paragraph(f'{latest_m}月无返佣规则变动记录。')
except Exception as e:
    doc.add_paragraph(f'（无法读取规则变动数据）')

doc.add_paragraph()

# 3.3 返佣增长异常品牌预警（原3.2）
doc.add_heading('3.3 返佣增长异常品牌预警', level=2)
p = doc.add_paragraph()
run = p.add_run(
    f'以下{len(high_growth)}个品牌-月份出现返佣金额环比增长>=50%且当月返佣>=10,000元的情况，'
    f'建议审计人员重点关注核实增长原因：')
if high_growth:
    reason_conn = sqlite3.connect(DB_PATH)
    reason_c = reason_conn.cursor()
    growth_reasons = {}
    for m, b, prev_s, curr_s, pct in high_growth:
        reason_c.execute('SELECT raw_amount FROM brand_audit WHERE month=? AND brand=?', (m, b))
        row = reason_c.fetchone()
        curr_r = float(row[0]) if row else 0
        prev_m = str(int(m) - 1) if m[4:] != '01' else str(int(m[:4]) - 1) + '12'
        reason_c.execute('SELECT raw_amount FROM brand_audit WHERE month=? AND brand=?', (prev_m, b))
        prow = reason_c.fetchone()
        prev_r = float(prow[0]) if prow else 0
        if prev_r > 0 and curr_r > 0:
            r_growth = (curr_r - prev_r) / prev_r * 100
            if r_growth > 50:
                growth_reasons[(m, b)] = '交易金额大幅上升带动'
            elif r_growth > 0:
                growth_reasons[(m, b)] = '交易金额略有增长'
            else:
                growth_reasons[(m, b)] = '返佣规则或比例调整'
        else:
            growth_reasons[(m, b)] = '新增商户或返佣规则变更'
    reason_conn.close()
    
    at = doc.add_table(rows=min(len(high_growth), 30) + 1, cols=6, style='Light Grid Accent 1')
    for i, h in enumerate(['月份', '品牌', '上月金额（元）', '本月金额（元）', '增长比例', '初步分析']):
        at.cell(0, i).text = h
        for p2 in at.cell(0, i).paragraphs:
            for r2 in p2.runs:
                r2.bold = True
                r2.font.size = Pt(8)
    for i, (m, b, prev, curr, pct) in enumerate(high_growth[:30]):
        at.cell(i+1, 0).text = m
        at.cell(i+1, 1).text = b
        at.cell(i+1, 2).text = f'{prev:,.2f}'
        at.cell(i+1, 3).text = f'{curr:,.2f}'
        at.cell(i+1, 4).text = f'+{pct}%'
        at.cell(i+1, 5).text = growth_reasons.get((m, b), '请核实')
        if pct > 100:
            for p2 in at.cell(i+1, 4).paragraphs:
                for r2 in p2.runs:
                    r2.font.color.rgb = RED
                    r2.bold = True

doc.add_page_break()

# ===== Section 4: 审计结果明细及数据来源 =====
doc.add_heading('四、审计结果明细及数据来源', level=1)

doc.add_heading('4.1 审计结果明细表（Excel）', level=2)
p = doc.add_paragraph()
run = p.add_run(f'《集团客户部FY审计结果明细（{AUDIT_PERIOD}）》')
run.bold = True
p.add_run('\n该Excel文件包含完整的品牌级审计数据，可按品牌、月份逐项查询核实。')
p = doc.add_paragraph()
run = p.add_run('下载链接（待上传飞书后更新）：')

doc.add_heading('4.2 审计数据库', level=2)
p = doc.add_paragraph()
run = p.add_run('《集团客户部FY数据.db》')
run.bold = True
p.add_run('\n包含所有原始数据、审计重算结果，支持SQL查询和深度分析。')

doc.add_heading('4.3 审计技能', level=2)
p = doc.add_paragraph()
run = p.add_run('「内控-集团客户部FY审计」')
run.bold = True
p.add_run('\n审计工作流基于此AI技能自动执行，支持增量审计、独立重算和结构化报告输出。')

doc.add_page_break()

# ===== Section 5: 审计方法论 =====
doc.add_heading('五、审计方法论', level=1)

doc.add_heading('5.1 审计方法', level=2)
for mp in [
    '本审计采用"独立重算"方法：从原始交易数据和返佣比例直接计算应发返佣金额，与汇总表逐月、逐品牌比对。',
    '不直接信任Excel计算表中的公式结果，而是按照统一的业务理解规则重新计算。',
    '所有原始数据均入库存储（SQLite），支持溯源查询和多次重算。',
    '差异判定标准：|差额|≤1元视为通过；|差额|>1元视为差异，需重点关注。',
]:
    doc.add_paragraph(mp, style='List Bullet')

doc.add_heading('5.2 数据来源', level=2)
for sp in [
    f'数据期间：{LABEL_PERIOD}，共{MONTH_COUNT}个月份的KA返佣Excel文件。',
    '原始数据来源：集团客户部提供的各月KA返佣结算表。',
    '数据库存储路径：/workspace/业务数据/集团客户部FY数据.db',
]:
    doc.add_paragraph(sp, style='List Bullet')

doc.add_heading('5.3 计算方式说明', level=2)
ct = doc.add_table(rows=6, cols=2, style='Light Grid Accent 1')
ct.cell(0, 0).text = '计算方式'
ct.cell(0, 1).text = '说明'
for r in ct.rows[0].cells:
    for p in r.paragraphs:
        for r2 in p.runs:
            r2.bold = True
for i, (m, d) in enumerate([
    ('常规返佣', 'ROUND(交易金额 × 返佣比例, 2) 每通道独立计算后求和'),
    ('中快系列', '直接使用原始total_rebate（包含支付宝间连返佣）'),
    ('拓展(202307版)', 'ROUND(交易金额 × (结算费率-商户返佣比例-成本费率) × 渠道分润比例, 2)'),
    ('志华/拓展人志华', 'ROUND(交易金额 × 收益率 × 渠道分润比例, 2) 四通道独立计算后求和'),
    ('异常/问题表', '按备注规则判定是否入账，仅符合条件计入'),
]):
    ct.cell(i+1, 0).text = m
    ct.cell(i+1, 1).text = d

doc.add_heading('5.4 差异判定标准', level=2)
st = doc.add_table(rows=3, cols=2, style='Light Grid Accent 1')
st.cell(0, 0).text = '差额范围'
st.cell(0, 1).text = '判定'
for r in st.rows[0].cells:
    for p in r.paragraphs:
        for r2 in p.runs:
            r2.bold = True
for i, (rng, result) in enumerate([
    ('|差额| ≤ 1元', '通过'),
    ('|差额| > 1元', '差异（需重点关注）'),
]):
    st.cell(i+1, 0).text = rng
    st.cell(i+1, 1).text = result

doc.add_paragraph()

# Save Word
doc.save(WORD_OUTPUT)
print(f"✅ Word报告已保存: {WORD_OUTPUT}")


# ====================================================================
# EXCEL REPORT
# ====================================================================
print("生成Excel报告...")
wb = openpyxl.load_workbook(TEMPLATE_PATH)

c2 = sqlite3.connect(DB_PATH).cursor()

# --- Audit Summary ---
ws = wb['审计摘要']
items = [
    ('审计时间范围', LABEL_PERIOD),
    ('审计月份数', f'{MONTH_COUNT}个月'),
    ('涉及品牌总数', f'{len(top_brands)}个'),
    ('通过月份数（|差额|≤1元）', f'{pass_months}个月'),
    ('差异月份数（|差额|>1元）', f'{len(diff_months)}个月'),
    ('汇总表总额', f'{sum(m["summary"] for m in MONTHLY.values()):,.2f}元'),
    ('重算总额', f'{sum(m["recalc"] for m in MONTHLY.values()):,.2f}元'),
    ('疑似多支付返佣金额', f'{total_real_diff:+,.2f}元'),
]
for i, (k, v) in enumerate(items):
    ws.cell(row=i+2, column=1, value=k)
    ws.cell(row=i+2, column=2, value=v)

# --- Monthly Trend ---
ws = wb['月度趋势']
ws.cell(row=1, column=1, value='月份')
ws.cell(row=1, column=2, value='返佣金额（元）')
ws.cell(row=1, column=3, value='品牌数')
ws.cell(row=1, column=4, value='差异（元）')
ws.cell(row=1, column=5, value='状态')
for i, m in enumerate(all_months):
    md = MONTHLY[m]
    st = '通过' if abs(md['real_diff']) <= 1 else '差异'
    ws.cell(row=i+2, column=1, value=m)
    ws.cell(row=i+2, column=2, value=md['summary'])
    ws.cell(row=i+2, column=3, value=md['brands'])
    ws.cell(row=i+2, column=4, value=md['diff'])
    ws.cell(row=i+2, column=5, value=st)
for rr in range(len(all_months)+2, 50):
    for cc in range(1, 6):
        ws.cell(row=rr, column=cc, value='')

# --- Detail Sheet ---
# Rename sheet: remove 历史差额调整 references
ws_d = wb['审计结果-多发明细']
ws_d.title = '审计结果-差异明细'
ws_d = wb['审计结果-差异明细']

detail_headers = ['月份', '品牌', '原始表金额', '审计重算金额', '汇总表金额', '差额', '审计备注']
for i, h in enumerate(detail_headers):
    ws_d.cell(row=1, column=i+1, value=h)

detail_rows = []
for r in brand_rows:
    m, b, raw_amt, recalc_amt, summary_amt, diff, is_real, note = r
    if diff > 1:
        detail_rows.append({
            'month': m, 'brand': b, 'raw': raw_amt,
            'recalc': recalc_amt, 'summary': summary_amt,
            'diff': diff, 'note': note
        })

detail_rows.sort(key=lambda x: (x['month'], -abs(x['diff'])))

row = 2
for d in detail_rows:
    ws_d.cell(row=row, column=1, value=d['month'])
    ws_d.cell(row=row, column=2, value=d['brand'])
    ws_d.cell(row=row, column=3, value=d['raw'])
    ws_d.cell(row=row, column=4, value=d['recalc'])
    ws_d.cell(row=row, column=5, value=d['summary'])
    ws_d.cell(row=row, column=6, value=d['diff'] if abs(d['diff']) > 0.01 else 0)
    ws_d.cell(row=row, column=7, value=d['note'] if d['note'] else '')
    row += 1
for rr in range(row, 550):
    for cc in range(1, 8):
        ws_d.cell(row=rr, column=cc, value='')

# --- Brand Growth Alert sheet ---
ws_alert = wb['品牌增长预警']
ws_alert.cell(row=1, column=1, value='月份')
ws_alert.cell(row=1, column=2, value='品牌')
ws_alert.cell(row=1, column=3, value='上月金额（元）')
ws_alert.cell(row=1, column=4, value='本月金额（元）')
ws_alert.cell(row=1, column=5, value='增长比例')
ws_alert.cell(row=1, column=6, value='备注')
for i, (m, b, prev, curr, pct) in enumerate(high_growth):
    ws_alert.cell(row=i+2, column=1, value=m)
    ws_alert.cell(row=i+2, column=2, value=b)
    ws_alert.cell(row=i+2, column=3, value=prev)
    ws_alert.cell(row=i+2, column=4, value=curr)
    ws_alert.cell(row=i+2, column=5, value=f'+{pct}%')
    ws_alert.cell(row=i+2, column=6, value=f'环比增长{pct}%，请核实增长原因')
for rr in range(len(high_growth)+2, 260):
    for cc in range(1, 9):
        ws_alert.cell(row=rr, column=cc, value='')

c2.close()
wb.save(EXCEL_OUTPUT)
print(f"✅ Excel报告已保存: {EXCEL_OUTPUT}")

# === Summary ===
print(f"\n{'='*60}")
print(f"报告生成完成！")
print(f"  Word: {WORD_OUTPUT}")
print(f"  Excel: {EXCEL_OUTPUT}")
print(f"{'='*60}")
print(f"\n审计数据摘要:")
print(f"  月份数: {MONTH_COUNT}")
print(f"  品牌数: {len(top_brands)}")
print(f"  通过/差异: {pass_months}/{len(diff_months)}月")
print(f"  疑似多支付返佣: {total_real_diff:+,.2f}元")
print(f"  高增长预警: {len(high_growth)}条")