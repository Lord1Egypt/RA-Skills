#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import re
import sys
import tempfile
import zipfile
from pathlib import Path
from typing import Any

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt
except ImportError as exc:  # pragma: no cover - environment guard
    raise SystemExit(
        "Missing python-docx. Run with the Codex workspace Python or install python-docx."
    ) from exc

try:
    from lxml import etree
except ImportError as exc:  # pragma: no cover - environment guard
    raise SystemExit("Missing lxml. Run with the Codex workspace Python or install lxml.") from exc


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}


def wqn(local: str) -> str:
    return f"{{{W_NS}}}{local}"


def load_spec(path: Path) -> dict[str, Any]:
    data = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(data, dict):
        raise ValueError("Review JSON root must be an object.")
    if not isinstance(data.get("references"), list) or not data["references"]:
        raise ValueError("Review JSON must contain a non-empty references list.")
    if not isinstance(data.get("paragraphs"), list) or not data["paragraphs"]:
        raise ValueError("Review JSON must contain a non-empty paragraphs list.")
    return data


def normalize_references(raw_refs: list[dict[str, Any]]) -> list[dict[str, str]]:
    refs: list[dict[str, str]] = []
    for idx, raw in enumerate(raw_refs, start=1):
        if not isinstance(raw, dict):
            raise ValueError(f"Reference {idx} must be an object.")
        reference_text = str(raw.get("text") or raw.get("formatted") or raw.get("gbt") or "").strip()
        if not reference_text:
            raise ValueError(f"Reference {idx} is missing 'text', 'formatted', or 'gbt'.")
        anchor = str(raw.get("anchor") or f"_RefBib{idx:03d}")
        if not re.match(r"^_RefBib[0-9]{3,}$", anchor):
            raise ValueError(
                f"Reference {idx} anchor '{anchor}' must look like _RefBib001."
            )
        refs.append({"anchor": anchor, "gbt": reference_text})
    return refs


def set_run_font(run, font_name: str, east_asia_font: str, size_pt: float, bold: bool = False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)
    run.font.size = Pt(size_pt)
    run.font.bold = bold


def field_result_props(
    font_name: str,
    east_asia_font: str,
    size_half_points: int,
    superscript: bool = True,
) -> OxmlElement:
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), font_name)
    fonts.set(qn("w:hAnsi"), font_name)
    fonts.set(qn("w:eastAsia"), east_asia_font)
    r_pr.append(fonts)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), str(size_half_points))
    r_pr.append(size)
    if superscript:
        vert = OxmlElement("w:vertAlign")
        vert.set(qn("w:val"), "superscript")
        r_pr.append(vert)
    return r_pr


def make_literal_run(
    text: str,
    font_name: str,
    east_asia_font: str,
    size_half_points: int,
    superscript: bool = True,
) -> OxmlElement:
    run = OxmlElement("w:r")
    run.append(field_result_props(font_name, east_asia_font, size_half_points, superscript))
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    return run


def make_ref_field_runs(
    anchor: str,
    display: str,
    font_name: str,
    east_asia_font: str,
    size_half_points: int,
    switch: str = "",
) -> list[OxmlElement]:
    runs: list[OxmlElement] = []

    r_begin = OxmlElement("w:r")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    r_begin.append(fld_begin)
    runs.append(r_begin)

    r_instr = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = f" REF {anchor} {switch} \\h " if switch else f" REF {anchor} \\h "
    r_instr.append(instr)
    runs.append(r_instr)

    r_sep = OxmlElement("w:r")
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    r_sep.append(fld_sep)
    runs.append(r_sep)

    r_result = OxmlElement("w:r")
    r_result.append(field_result_props(font_name, east_asia_font, size_half_points, True))
    t = OxmlElement("w:t")
    t.text = display
    r_result.append(t)
    runs.append(r_result)

    r_end = OxmlElement("w:r")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r_end.append(fld_end)
    runs.append(r_end)

    return runs


def citation_terms(cite: Any, collapse: bool | None, min_range: int) -> list[tuple[int, int]]:
    if isinstance(cite, int):
        return [(cite, cite)]
    if not isinstance(cite, list) or not cite:
        raise ValueError(f"Invalid citation value: {cite!r}")
    indices = []
    for item in cite:
        if not isinstance(item, int):
            raise ValueError(f"Citation list items must be integers: {cite!r}")
        indices.append(item)
    if collapse is False:
        return [(idx, idx) for idx in indices]

    terms: list[tuple[int, int]] = []
    run_start = indices[0]
    prev = indices[0]
    for idx in indices[1:]:
        if idx == prev + 1:
            prev = idx
            continue
        terms.extend(expand_or_collapse(run_start, prev, min_range))
        run_start = prev = idx
    terms.extend(expand_or_collapse(run_start, prev, min_range))
    return terms


def expand_or_collapse(start: int, end: int, min_range: int) -> list[tuple[int, int]]:
    if end - start + 1 >= min_range:
        return [(start, end)]
    return [(idx, idx) for idx in range(start, end + 1)]


def append_citation(paragraph, cite_obj: dict[str, Any], refs: list[dict[str, str]], style: dict[str, Any]):
    cite = cite_obj.get("cite")
    collapse = cite_obj.get("collapse")
    min_range = int(cite_obj.get("min_range", style.get("min_range", 3)))
    terms = citation_terms(cite, collapse, min_range)

    font_name = style["citation_font"]
    east_asia = style["east_asia_font"]
    size_hp = int(round(float(style["body_size_pt"]) * 2))

    paragraph._p.append(make_literal_run("[", font_name, east_asia, size_hp, True))
    for term_index, (start, end) in enumerate(terms):
        if term_index:
            paragraph._p.append(make_literal_run(",", font_name, east_asia, size_hp, True))
        for idx in (start, end) if start != end else (start,):
            if idx < 1 or idx > len(refs):
                raise ValueError(f"Citation index {idx} is outside references list.")
            anchor = refs[idx - 1]["anchor"]
            for run in make_ref_field_runs(anchor, str(idx), font_name, east_asia, size_hp):
                paragraph._p.append(run)
            if start != end and idx == start:
                paragraph._p.append(make_literal_run("-", font_name, east_asia, size_hp, True))
    paragraph._p.append(make_literal_run("]", font_name, east_asia, size_hp, True))


def add_review_paragraph(doc: Document, parts: list[Any], refs: list[dict[str, str]], style: dict[str, Any]):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(float(style["first_line_indent_pt"]))
    p.paragraph_format.line_spacing = float(style["line_spacing"])
    p.paragraph_format.space_after = Pt(float(style["body_space_after_pt"]))
    for part in parts:
        if isinstance(part, str):
            run = p.add_run(part)
            set_run_font(
                run,
                style["body_font"],
                style["east_asia_font"],
                float(style["body_size_pt"]),
            )
        elif isinstance(part, dict) and "cite" in part:
            append_citation(p, part, refs, style)
        else:
            raise ValueError(f"Paragraph part must be string or citation object: {part!r}")


def append_hidden_bookmarked_number(
    paragraph,
    text: str,
    bookmark_name: str,
    bookmark_id: int,
    style: dict[str, Any],
):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), bookmark_name)
    paragraph._p.append(start)

    run = paragraph.add_run(text)
    set_run_font(
        run,
        style["body_font"],
        style["east_asia_font"],
        float(style["reference_size_pt"]),
    )
    vanish = OxmlElement("w:vanish")
    run._r.get_or_add_rPr().append(vanish)

    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.append(end)


def build_base_docx(spec: dict[str, Any], refs: list[dict[str, str]], out: Path):
    style = {
        "body_font": "SimSun",
        "east_asia_font": "宋体",
        "citation_font": "Times New Roman",
        "body_size_pt": 12,
        "reference_size_pt": 10.5,
        "title_size_pt": 16,
        "reference_heading_size_pt": 14,
        "line_spacing": 1.5,
        "reference_line_spacing": 1.25,
        "first_line_indent_pt": 24,
        "body_space_after_pt": 6,
        "reference_space_after_pt": 3,
        "min_range": 3,
    }
    style.update(spec.get("style") or {})
    if "body_east_asia_font" in style and "east_asia_font" not in (spec.get("style") or {}):
        style["east_asia_font"] = style["body_east_asia_font"]

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = style["body_font"]
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), style["east_asia_font"])
    normal.font.size = Pt(float(style["body_size_pt"]))

    for section in doc.sections:
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    title_run = title.add_run(str(spec.get("title") or "文献综述"))
    set_run_font(
        title_run,
        style["body_font"],
        style["east_asia_font"],
        float(style["title_size_pt"]),
        bold=True,
    )

    for paragraph_parts in spec["paragraphs"]:
        if not isinstance(paragraph_parts, list):
            raise ValueError("Each paragraph must be a list of string/citation parts.")
        add_review_paragraph(doc, paragraph_parts, refs, style)

    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)
    heading_run = heading.add_run(str(spec.get("references_heading") or "参考文献"))
    set_run_font(
        heading_run,
        style["body_font"],
        style["east_asia_font"],
        float(style["reference_heading_size_pt"]),
        bold=True,
    )

    for idx, ref in enumerate(refs, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(21)
        p.paragraph_format.first_line_indent = Pt(-21)
        p.paragraph_format.line_spacing = float(style["reference_line_spacing"])
        p.paragraph_format.space_after = Pt(float(style["reference_space_after_pt"]))
        append_hidden_bookmarked_number(p, str(idx), ref["anchor"], idx, style)
        ref_run = p.add_run(ref["gbt"])
        set_run_font(
            ref_run,
            style["body_font"],
            style["east_asia_font"],
            float(style["reference_size_pt"]),
        )

    doc.save(out)


def unzip_docx(src: Path, dest: Path):
    with zipfile.ZipFile(src) as zf:
        zf.extractall(dest)


def zip_docx(src_dir: Path, dest: Path):
    if dest.exists():
        dest.unlink()
    with zipfile.ZipFile(dest, "w", zipfile.ZIP_DEFLATED) as zf:
        for path in src_dir.rglob("*"):
            if path.is_file():
                zf.write(path, path.relative_to(src_dir).as_posix())


def ensure_bracket_numbering(work: Path) -> str:
    numbering_path = work / "word" / "numbering.xml"
    parser = etree.XMLParser(remove_blank_text=False)
    tree = etree.parse(str(numbering_path), parser)
    root = tree.getroot()

    existing_abstract_ids = [
        int(el.get(wqn("abstractNumId")))
        for el in root.findall("w:abstractNum", namespaces=NS)
        if (el.get(wqn("abstractNumId")) or "").isdigit()
    ]
    existing_num_ids = [
        int(el.get(wqn("numId")))
        for el in root.findall("w:num", namespaces=NS)
        if (el.get(wqn("numId")) or "").isdigit()
    ]
    abstract_id = max(existing_abstract_ids, default=0) + 1
    num_id = max(existing_num_ids, default=0) + 1

    abstract = etree.Element(wqn("abstractNum"))
    abstract.set(wqn("abstractNumId"), str(abstract_id))
    etree.SubElement(abstract, wqn("nsid")).set(wqn("val"), "5A5A5A5A")
    etree.SubElement(abstract, wqn("multiLevelType")).set(wqn("val"), "singleLevel")

    lvl = etree.SubElement(abstract, wqn("lvl"))
    lvl.set(wqn("ilvl"), "0")
    etree.SubElement(lvl, wqn("start")).set(wqn("val"), "1")
    etree.SubElement(lvl, wqn("numFmt")).set(wqn("val"), "decimal")
    etree.SubElement(lvl, wqn("lvlText")).set(wqn("val"), "[%1]")
    etree.SubElement(lvl, wqn("lvlJc")).set(wqn("val"), "left")
    etree.SubElement(lvl, wqn("suff")).set(wqn("val"), "space")
    p_pr = etree.SubElement(lvl, wqn("pPr"))
    ind = etree.SubElement(p_pr, wqn("ind"))
    ind.set(wqn("left"), "420")
    ind.set(wqn("hanging"), "420")

    root.append(abstract)
    num = etree.Element(wqn("num"))
    num.set(wqn("numId"), str(num_id))
    etree.SubElement(num, wqn("abstractNumId")).set(wqn("val"), str(abstract_id))
    root.append(num)

    tree.write(str(numbering_path), xml_declaration=True, encoding="UTF-8", standalone="yes")
    return str(num_id)


def apply_reference_numbering(work: Path, num_id: str) -> int:
    document_path = work / "word" / "document.xml"
    parser = etree.XMLParser(remove_blank_text=False)
    tree = etree.parse(str(document_path), parser)
    root = tree.getroot()
    numbered_paragraph_ids: set[int] = set()

    for bookmark in root.findall(".//w:bookmarkStart", namespaces=NS):
        name = bookmark.get(wqn("name")) or ""
        if not name.startswith("_RefBib"):
            continue
        p = bookmark.getparent()
        if p is None or p.tag != wqn("p"):
            continue
        p_pr = p.find("w:pPr", namespaces=NS)
        if p_pr is None:
            p_pr = etree.Element(wqn("pPr"))
            p.insert(0, p_pr)
        num_pr = p_pr.find("w:numPr", namespaces=NS)
        if num_pr is None:
            num_pr = etree.Element(wqn("numPr"))
            p_pr.insert(0, num_pr)
        ilvl = num_pr.find("w:ilvl", namespaces=NS)
        if ilvl is None:
            ilvl = etree.SubElement(num_pr, wqn("ilvl"))
        ilvl.set(wqn("val"), "0")
        num_id_el = num_pr.find("w:numId", namespaces=NS)
        if num_id_el is None:
            num_id_el = etree.SubElement(num_pr, wqn("numId"))
        num_id_el.set(wqn("val"), num_id)
        numbered_paragraph_ids.add(id(p))

    tree.write(str(document_path), xml_declaration=True, encoding="UTF-8", standalone="yes")
    return len(numbered_paragraph_ids)


def enable_field_updates(work: Path):
    settings_path = work / "word" / "settings.xml"
    if not settings_path.exists():
        return
    parser = etree.XMLParser(remove_blank_text=False)
    tree = etree.parse(str(settings_path), parser)
    root = tree.getroot()
    update = root.find("w:updateFields", namespaces=NS)
    if update is None:
        update = etree.SubElement(root, wqn("updateFields"))
    update.set(wqn("val"), "true")
    tree.write(str(settings_path), xml_declaration=True, encoding="UTF-8", standalone="yes")


def finish_ooxml(out: Path) -> int:
    with tempfile.TemporaryDirectory() as tmp:
        work = Path(tmp) / "docx"
        work.mkdir()
        unzip_docx(out, work)
        num_id = ensure_bracket_numbering(work)
        numbered = apply_reference_numbering(work, num_id)
        enable_field_updates(work)
        zip_docx(work, out)
    return numbered


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(
        description="Build a DOCX literature review with REF citations and auto-numbered bibliography."
    )
    parser.add_argument("review_json", type=Path, help="Review JSON spec.")
    parser.add_argument("--out", type=Path, required=True, help="Output .docx path.")
    args = parser.parse_args(argv)

    spec = load_spec(args.review_json)
    refs = normalize_references(spec["references"])
    args.out.parent.mkdir(parents=True, exist_ok=True)
    build_base_docx(spec, refs, args.out)
    numbered = finish_ooxml(args.out)
    print(f"wrote={args.out}")
    print(f"references={len(refs)}")
    print(f"auto_numbered_bibliography_paragraphs={numbered}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
