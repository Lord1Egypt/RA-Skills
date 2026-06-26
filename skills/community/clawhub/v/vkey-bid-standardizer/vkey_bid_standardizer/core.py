"""vkey-bid-standardizer.core

通用 OOXML 操作：run 文本替换、run/paragraph 覆盖清理、样式字体设置等。
"""
from typing import Iterable

from docx.document import Document as DocumentClass
from docx.styles.style import _ParagraphStyle
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt


# ═══════════════════════════════════════════
#  段落 run 文本操作
# ═══════════════════════════════════════════

def replace_para_text(para, new_text: str) -> None:
    """替换段落文本，保留段落级样式（heading / font 由 style 决定）。"""
    runs = para.runs
    if not runs:
        para.add_run(new_text)
        return
    for r in runs:
        r.text = ''
    runs[0].text = new_text


# ═══════════════════════════════════════════
#  覆盖清理（让样式定义完全生效）
# ═══════════════════════════════════════════

def reset_run_font_overrides(doc: DocumentClass, style_names: Iterable[str] = ('Normal',)) -> int:
    """清除指定样式段落 run 级的字号/字体覆盖（w:sz / w:szCs / w:rFonts），
    让样式定义生效。

    Returns:
        被清除的元素数。
    """
    targets = ('w:sz', 'w:szCs', 'w:rFonts')
    count = 0
    for para in doc.paragraphs:
        if para.style.name not in style_names:
            continue
        for run in para.runs:
            rPr = run._element.get_or_add_rPr()
            for tag in targets:
                el = rPr.find(qn(tag))
                if el is not None:
                    rPr.remove(el)
                    count += 1
    return count


def reset_paragraph_overrides(
    doc: DocumentClass,
    style_names: Iterable[str] = ('Normal', 'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5'),
    fields: Iterable[str] = (
        'w:spacing', 'w:ind', 'w:jc', 'w:outlineLvl',
        'w:numPr', 'w:keepNext', 'w:keepLines', 'w:pageBreakBefore',
    ),
) -> int:
    """清除指定样式段落段级 pPr 覆盖，让样式定义完全生效。

    默认清掉 7 个常见覆盖字段：
    - w:spacing     段前/段后/行距
    - w:ind         左右缩进/首行缩进
    - w:jc          对齐方式
    - w:outlineLvl  大纲级别
    - w:numPr       编号引用（避免与 apply_auto_numbering 冲突）
    - w:keepNext / w:keepLines / w:pageBreakBefore  分页/段控制

    Returns:
        被清除的元素数。
    """
    count = 0
    for para in doc.paragraphs:
        if para.style.name not in style_names:
            continue
        pPr = para._element.find(qn('w:pPr'))
        if pPr is None:
            continue
        for tag in fields:
            el = pPr.find(qn(tag))
            if el is not None:
                pPr.remove(el)
                count += 1
    return count


# ═══════════════════════════════════════════
#  样式字体设置
# ═══════════════════════════════════════════

def set_style_font(style: _ParagraphStyle, west_font: str, east_font: str, size: float, bold: bool = True) -> None:
    """设置样式字体，清除主题引用并显式设置黑色 #000000。"""
    rpr = style.element.get_or_add_rPr()
    for tag in ('w:rFonts', 'w:color', 'w:szCs', 'w:bCs'):
        old = rpr.find(qn(tag))
        if old is not None:
            rpr.remove(old)
    rfonts = OxmlElement('w:rFonts')
    rfonts.set(qn('w:ascii'), west_font)
    rfonts.set(qn('w:hAnsi'), west_font)
    rfonts.set(qn('w:eastAsia'), east_font)
    rfonts.set(qn('w:cs'), west_font)
    rpr.insert(0, rfonts)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '000000')
    rpr.append(color)
    style.font.name = west_font
    style.font.size = Pt(size)
    style.font.bold = bold


# ═══════════════════════════════════════════
#  段落格式应用
# ═══════════════════════════════════════════

def apply_paragraph_format(style: _ParagraphStyle, config: dict) -> None:
    """根据 config 字典应用段前/段后/行距/对齐/首行缩进。"""
    from docx.enum.text import WD_LINE_SPACING

    pf = style.paragraph_format
    if 'space_before' in config:
        pf.space_before = Pt(config['space_before'])
    if 'space_after' in config:
        pf.space_after = Pt(config['space_after'])
    if 'line_spacing' in config:
        unit = config.get('line_spacing_unit', 'multiple')
        if unit == 'exact':
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = Pt(config['line_spacing'])
        else:
            pf.line_spacing = config['line_spacing']
    if 'alignment' in config:
        align_map = {'left': 0, 'center': 1, 'right': 2, 'justify': 3}
        pf.alignment = align_map.get(config['alignment'], 0)
    if 'first_line_chars' in config and 'size' in config:
        # GB/T 9704 公文格式：首行缩进 N 字符 = N × 中文字号 pt
        indent_pt = config['first_line_chars'] * config['size']
        pf.first_line_indent = Pt(indent_pt)
        # 直接写 OOXML w:ind 的 firstLine + firstLineChars，避免 python-docx
        # 默认把 firstLineChars 写成 1 导致只缩进 1 字符
        pPr_el = style.element.find(qn('w:pPr'))
        if pPr_el is None:
            pPr_el = OxmlElement('w:pPr')
            style.element.append(pPr_el)
        ind = pPr_el.find(qn('w:ind'))
        if ind is None:
            ind = OxmlElement('w:ind')
            pPr_el.append(ind)
        ind.set(qn('w:firstLine'), str(int(round(indent_pt * 20))))  # twips (1/20 pt)
        # OOXML 规范：firstLineChars 单位是 1/100 字符；2 字符必须写 200
        ind.set(qn('w:firstLineChars'), str(int(config['first_line_chars']) * 100))
        # leftChars 留默认（0），让样式控制 left 缩进
        if ind.get(qn('w:left')) is None:
            ind.set(qn('w:left'), '0')
            ind.set(qn('w:leftChars'), '0')


def get_paragraph_overrides(para) -> dict:
    """读取段落 pPr 覆盖（用于诊断/审计）。"""
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        return {}
    result = {}
    for child in pPr:
        tag = child.tag.split('}')[-1]
        result[tag] = {k.split('}')[-1]: v for k, v in child.attrib.items()}
    return result
