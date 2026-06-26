"""
generate_docx.py — DOCX报告生成参考模板（v6.1 consulting-report-generator）
基于 python-docx 生成专业Word文档，Theme Contract 风格
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# Theme Contract (deepblue)
THEME = {
    "primary": RGBColor(0x05, 0x1C, 0x2C),
    "secondary": RGBColor(0x33, 0x33, 0x33),
    "accent": RGBColor(0x00, 0x6B, 0xA6),
    "light": RGBColor(0xF2, 0xF2, 0xF2),
    "bg": RGBColor(0xFF, 0xFF, 0xFF),
    "font_cn": "Microsoft YaHei",
    "font_en": "Arial",
}

def set_font(run, size=11, color=THEME["secondary"], bold=False, font_name=None):
    """统一设置字体"""
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.name = font_name or THEME["font_en"]
    # 设置东亚字体
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{THEME["font_cn"]}"/>')
        rPr.append(rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), THEME["font_cn"])


def add_title(doc, text, level=0):
    """添加标题（0=封面主标题, 1=章节, 2=子节）"""
    if level == 0:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_before = Pt(72)
        p.space_after = Pt(24)
        run = p.add_run(text)
        set_font(run, size=26, color=THEME["primary"], bold=True)
        return p
    else:
        heading = doc.add_heading(text, level=level)
        for run in heading.runs:
            set_font(run, size={1: 18, 2: 14}[level], color=THEME["primary"], bold=True)
        return heading


def add_body(doc, text, size=11, bold=False):
    """添加正文段落"""
    p = doc.add_paragraph()
    p.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    return p


def add_bullet(doc, items):
    """添加列表"""
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        for run in p.runs:
            set_font(run, size=11)
    return doc


def add_table_simple(doc, headers, rows):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'
    
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_font(run, size=9, color=THEME["bg"], bold=True)
        # 表头背景色
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="051C2C"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    
    # 数据行
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_font(run, size=9)
            if ri % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>')
                cell._tc.get_or_add_tcPr().append(shading)
    
    return table


def add_page_break(doc):
    """分页"""
    doc.add_page_break()


def generate_report(title, sections, output_path):
    """
    生成完整DOCX报告
    
    参数:
        title: 报告标题
        sections: [{"heading": str, "content": [text|table|list], ...}]
        output_path: 输出路径
    """
    doc = Document()
    
    # 封面
    add_title(doc, title, level=0)
    doc.add_paragraph()
    for section in sections:
        add_page_break(doc)
        add_title(doc, section["heading"], level=1)
        
        for item in section.get("content", []):
            if isinstance(item, str):
                add_body(doc, item)
            elif isinstance(item, list):
                add_bullet(doc, item)
    
    doc.save(output_path)
    print(f"✅ DOCX报告已生成: {output_path}")
    return output_path


if __name__ == "__main__":
    # 示例
    generate_report(
        title="精益生产体系深度解析",
        sections=[
            {"heading": "一、精益生产概述", "content": ["精益生产源于丰田..."]},
            {"heading": "二、七大浪费", "content": ["TIMWOOD...", ["生产过剩", "等待"]]},
        ],
        output_path="/tmp/report.docx"
    )
