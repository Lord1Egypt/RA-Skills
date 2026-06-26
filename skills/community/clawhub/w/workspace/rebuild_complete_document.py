#!/usr/bin/env python3
"""
Rebuild Complete Document with Consistent Styles
Creates a new document with all 7 chapters and uniform formatting
"""

import sys
import os
from pathlib import Path
import subprocess
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE

def create_clean_chapter4():
    """Create cleaned version of chapter 4"""
    source_file = os.path.expanduser("~/Documents/PersonalProjects/report-agent/outputs/chapters/ch04.md")
    output_file = "chapter4_clean.txt"
    
    if not os.path.exists(source_file):
        print(f"❌ Chapter 4 source not found: {source_file}")
        return False
    
    try:
        with open(source_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Clean markdown formatting
        lines = content.split('\n')
        cleaned_lines = []
        
        for line in lines:
            # Remove markdown headers
            if line.startswith('# '):
                cleaned_lines.append(line[2:])
            elif line.startswith('## '):
                cleaned_lines.append(line[3:])
            elif line.startswith('### '):
                cleaned_lines.append(line[4:])
            else:
                cleaned_lines.append(line)
        
        cleaned_content = '\n'.join(cleaned_lines)
        
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(cleaned_content)
        
        print(f"✅ Created: {output_file}")
        return True
        
    except Exception as e:
        print(f"❌ Error creating chapter 4: {e}")
        return False

def read_chapter_text(chapter_num):
    """Read cleaned chapter text from file"""
    if chapter_num == 4:
        filename = "chapter4_clean.txt"
    else:
        filename = f"chapter{chapter_num}_clean.txt"
    
    if not os.path.exists(filename):
        print(f"❌ Chapter file not found: {filename}")
        return None
    
    with open(filename, 'r', encoding='utf-8') as f:
        return f.read()

def apply_uniform_styles(doc):
    """Apply uniform styles to document"""
    print("   Applying uniform styles...")
    
    # Define styles (simplified - actual implementation would need more detail)
    styles = doc.styles
    
    # Ensure Heading 1 style exists
    if 'Heading 1' not in styles:
        heading1 = styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        heading1.font.size = Pt(16)
        heading1.font.bold = True
        heading1.paragraph_format.space_after = Pt(12)
    
    # Ensure Heading 2 style exists  
    if 'Heading 2' not in styles:
        heading2 = styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
        heading2.font.size = Pt(14)
        heading2.font.bold = True
        heading2.paragraph_format.space_after = Pt(8)
    
    # Ensure Normal style
    if 'Normal' not in styles:
        normal = styles['Normal']
        normal.font.size = Pt(12)
        normal.paragraph_format.line_spacing = 1.5
        normal.paragraph_format.space_after = Pt(6)
    
    return True

def create_complete_document():
    """Create complete document with all 7 chapters"""
    print("📝 Creating complete document with all 7 chapters...")
    
    try:
        # Create new document
        doc = Document()
        
        # Apply uniform styles
        apply_uniform_styles(doc)
        
        # Define all chapters
        chapters = [
            (1, "فصل ۱: مقدمه و بیان مسئله"),
            (2, "فصل ۲: مروری بر ادبیات موضوع"),
            (3, "فصل ۳: معماری سیستم"),
            (4, "فصل ۴: ساخت کلاستر Kubernetes"),
            (5, "فصل ۵: پیاده‌سازی لایه سرویس‌های در سطح کاربر"),
            (6, "فصل ۶: پایش، سخت‌سازی امنیتی، آزمون‌ها و تحویل نهایی"),
            (7, "فصل ۷: خلصه، نتیجه‌گیری و کارهای آینده")
        ]
        
        # Add each chapter
        for chapter_num, chapter_title in chapters:
            print(f"   Adding {chapter_title}")
            
            # Add page break (except for first chapter)
            if chapter_num > 1:
                doc.add_page_break()
            
            # Add chapter title
            title_para = doc.add_paragraph(chapter_title)
            title_para.style = doc.styles['Heading 1']
            
            # Add chapter content
            content = read_chapter_text(chapter_num)
            if content:
                # For now, add placeholder text
                # In real implementation, we would parse and structure the content
                if chapter_num in [1, 2, 3]:
                    # These chapters are already in the original document
                    # We'll add placeholder indicating they need to be copied from original
                    placeholder = doc.add_paragraph(f"[محتوای فصل {chapter_num} از سند اصلی کپی شود]")
                    placeholder.style = doc.styles['Normal']
                else:
                    # For chapters 4-7, add the content we have
                    paragraphs = content.split('\n\n')
                    for para_text in paragraphs:
                        if para_text.strip():
                            # Check for section headers
                            if chapter_num == 4 and any(f"۴.{i}" in para_text for i in range(1, 8)):
                                section_para = doc.add_paragraph(para_text)
                                section_para.style = doc.styles['Heading 2']
                            elif chapter_num == 5 and any(f"۵.{i}" in para_text for i in range(1, 9)):
                                section_para = doc.add_paragraph(para_text)
                                section_para.style = doc.styles['Heading 2']
                            elif chapter_num == 6 and any(f"۶.{i}" in para_text for i in range(1, 6)):
                                section_para = doc.add_paragraph(para_text)
                                section_para.style = doc.styles['Heading 2']
                            elif chapter_num == 7 and any(f"۷-{i}" in para_text for i in range(1, 4)):
                                section_para = doc.add_paragraph(para_text)
                                section_para.style = doc.styles['Heading 2']
                            else:
                                # Regular paragraph
                                para = doc.add_paragraph(para_text)
                                para.style = doc.styles['Normal']
            
            # Add spacing
            doc.add_paragraph()
        
        # Save final document
        final_name = "BSc-project-report-COMPLETE.docx"
        print(f"   Saving complete document: {final_name}")
        doc.save(final_name)
        
        print(f"✅ Complete document created: {final_name}")
        return final_name
        
    except Exception as e:
        print(f"❌ Error creating document: {e}")
        import traceback
        traceback.print_exc()
        return None

def open_chrome_with_url():
    """Open Chrome with the ChatGPT conversation URL"""
    url = "https://chatgpt.com/g/g-p-68e4e4213d388191b568c928ecf4eb30-university/c/69e76af8-7b18-8387-9139-d13d13b9818c"
    
    print(f"🌐 Opening Chrome with URL: {url}")
    
    try:
        # Try to open in Chrome
        subprocess.run(["open", "-a", "Google Chrome", url])
        print("✅ Chrome opened with the conversation")
        return True
    except:
        try:
            # Try with default browser
            subprocess.run(["open", url])
            print("✅ Browser opened with the conversation")
            return True
        except Exception as e:
            print(f"❌ Could not open browser: {e}")
            return False

def main():
    """Main function"""
    print("=" * 60)
    print("REBUILD COMPLETE DOCUMENT WITH UNIFORM STYLES")
    print("=" * 60)
    
    print("\n🔍 Checking and preparing files...")
    
    # Create chapter 4 if needed
    if not os.path.exists("chapter4_clean.txt"):
        print("Creating Chapter 4 clean file...")
        create_clean_chapter4()
    
    # Check all required files
    required_files = []
    for i in range(4, 8):  # Chapters 4-7
        if i == 4:
            filename = "chapter4_clean.txt"
        else:
            filename = f"chapter{i}_clean.txt"
        required_files.append(filename)
    
    missing_files = []
    for file in required_files:
        if os.path.exists(file):
            print(f"   ✅ {file}")
        else:
            print(f"   ❌ {file}")
            missing_files.append(file)
    
    if missing_files:
        print(f"\n❌ Missing files: {missing_files}")
        print("Please ensure all chapter files are available")
        return
    
    # Create complete document
    print("\n🚀 Building complete document...")
    final_doc = create_complete_document()
    
    if final_doc:
        print("\n✅ DOCUMENT REBUILT SUCCESSFULLY!")
        print(f"\n📄 Final document: {final_doc}")
        
        print("\n📋 NEXT STEPS:")
        print("1. Open the ChatGPT conversation to get manifests")
        print("2. Fix Chapter 3 issues with those manifests")
        print("3. Copy Chapters 1-3 from your original edited document")
        print("4. Apply uniform styles to all chapters")
        print("5. Add manifests and references as per guidelines")
        
        # Ask to open Chrome
        response = input("\nOpen ChatGPT conversation in Chrome now? (y/n): ").strip().lower()
        if response == 'y':
            open_chrome_with_url()
            
            print("\n📝 Please share the manifest content from the conversation")
            print("   I'll use it to fix Chapter 3 and complete the package")
    else:
        print("\n❌ Failed to create complete document")

if __name__ == "__main__":
    main()