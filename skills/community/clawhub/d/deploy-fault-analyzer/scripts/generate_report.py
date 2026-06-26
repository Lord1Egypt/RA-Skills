#!/usr/bin/env python3
"""
部署故障分析报告生成器 v2.0
新增:
  - parse_deploy_json_log()    → JSON 日志自动解析
  - categorize_error_message() → 中文错误模式自动分类
  - deduplicate_errors()       → 错误去重归并
  - generate_multi_fault_report() → 多故障批量生成（含总览页）
用法:
  python3 generate_report.py --json deploy_logs.json --output output.docx
  python3 generate_report.py --data fault_data.json --output output.docx
  python3 generate_report.py --data faults.json --output output.docx --batch
  或由 Hermes 在 execute_code 中 import 调用
"""

import json
import re
import sys
from pathlib import Path
from datetime import datetime
from typing import Optional


# ═══════════════════════════════════════════════════════════════
#  解析与分类
# ═══════════════════════════════════════════════════════════════

def parse_deploy_json_log(filepath: str) -> dict:
    """解析部署平台 JSON 日志格式。
    期望格式: {generated_at, total_logs, total_errors, logs: [{type, message, timestamp, task_id}]}
    返回: {errors, warnings, infos, total, meta}
    """
    with open(filepath, 'r') as f:
        data = json.load(f)

    logs = data.get('logs', [])
    if not logs:
        raise ValueError("No 'logs' array found in JSON — 不是标准部署日志格式")

    errors = [l for l in logs if l.get('type') == 'error']
    warnings = [l for l in logs if l.get('type') == 'warning']
    infos = [l for l in logs if l.get('type') == 'info']

    return {
        'errors': errors,
        'warnings': warnings,
        'infos': infos,
        'total': len(logs),
        'meta': {
            'generated_at': data.get('generated_at', ''),
            'total_logs': data.get('total_logs', len(logs)),
            'total_errors': data.get('total_errors', len(errors)),
        },
    }


def parse_task_id(task_id: Optional[str]) -> dict:
    """解析 task_id → 任务阶段和类型。

    PREP_UPLOAD_33B3DC → {'phase': 'PREP', 'type': 'UPLOAD', 'id': '33B3DC'}
    CHECK_DEPLOY_F5B45D → {'phase': 'CHECK', 'type': 'DEPLOY', 'id': 'F5B45D'}
    None → {'phase': '全局', 'type': '系统', 'id': ''}
    """
    if not task_id:
        return {'phase': '全局', 'type': '系统', 'id': ''}
    parts = task_id.split('_')
    if len(parts) >= 2:
        return {
            'phase': parts[0],
            'type': parts[1],
            'id': parts[-1] if len(parts) > 2 else '',
        }
    return {'phase': task_id, 'type': 'UNKNOWN', 'id': ''}


def categorize_error_message(msg: str) -> str:
    """根据中文/英文关键词自动归类错误。
    返回: '数据异常' | '配置错误' | '资源不足' | '网络/连接故障' |
          '权限问题' | '服务异常' | '依赖缺失' | '用户操作' | '待分类'
    """
    patterns = [
        (r'日期格式不正确|日期格式错误|格式yyyy-mm-dd', '数据异常'),
        (r'资源池名称不正确|名称不一致', '配置错误'),
        (r'未填写|必填项.*未填|必填.*缺失|必填项', '配置错误'),
        (r'缺失|不存在|not found|No such file|no such file', '资源不足'),
        (r'格式不正确.*IP|IP地址格式|格式不合法', '配置错误'),
        (r'already installed|is already|冲突|conflict', '服务异常'),
        (r'已终止|用户.*终止|用户.*取消', '用户操作'),
        (r'Connection refused|timeout|unreachable|连接.*拒绝', '网络/连接故障'),
        (r'Permission denied|权限不足|访问被拒', '权限问题'),
        (r'OOM|out of memory|磁盘空间不足|No space', '资源不足'),
        (r'ModuleNotFoundError|ImportError|依赖.*缺失', '依赖缺失'),
        (r'Verifying|verification|校验失败|验证失败', '服务异常'),
        (r'host.*not found|DNS.*fail|解析失败', '网络/连接故障'),
        (r'bms\S*\.pri\.|hostname.*invalid', '配置错误'),
    ]
    for pattern, category in patterns:
        if re.search(pattern, msg):
            return category
    return '待分类'


def deduplicate_errors(errors: list) -> list:
    """归并重复错误 → 根因列表。
    每项: {category, count, task_ids (set), sample (第一条原始错误)}
    
    三级去重策略：
    1. 消息归一化（去时间戳、去具体数值）→ 模式特征
    2. 按类别+语义核心去重
    3. 按关键短语同级合并（同类别 + 同短语 → 归并）
    """
    # 1. 过滤用户操作类
    real_faults = [e for e in errors
                   if '用户已终止' not in e.get('message', '')
                   and '用户终止' not in e.get('message', '')]

    # 2. 消息归一化 → 模式特征
    def _msg_pattern(msg: str) -> str:
        """激进归一化：去所有时间戳、数字、路径 → 保留语义骨架"""
        msg = re.sub(r'\[.*?\]', '', msg)
        msg = re.sub(r'\d{4}-\d{2}-\d{2}\s+\d{2}:\d{2}:\d{2}[,.]?\d*', '', msg)
        msg = re.sub(r'\d{2}:\d{2}:\d{2}[,.]?\d*', '', msg)
        msg = re.sub(r'[A-F0-9]{6,}', 'HEXID', msg)
        msg = re.sub(r'sheet\d+', 'sheetN', msg)
        msg = re.sub(r'\d{4,}', 'NUM', msg)
        msg = re.sub(r'/\S*\.\w+', '/PATH', msg)
        msg = re.sub(r'\w+Impl\.py', 'MODULE.py', msg)
        msg = re.sub(r'\w+\.py', 'MODULE.py', msg)
        msg = re.sub(r'\s+', ' ', msg).strip()
        return msg

    def _extract_key_phrase(msg: str) -> str:
        """从消息中提取关键短语用于同级合并"""
        phrases = [
            '日期格式不正确', '日期格式错误',
            '资源池名称不正确', '名称不一致',
            '未填写', '必填项',
            'already installed', 'is already',
            '缺失',
        ]
        for p in phrases:
            if p in msg:
                return p
        return msg[:40]  # fallback

    # 第一级 + 第二级
    seen = {}
    for e in real_faults:
        msg = e.get('message', '')
        cat = categorize_error_message(msg)
        # 提取语义核心
        pat = _msg_pattern(msg)
        semantic = re.sub(r'(HEXID|NUM|MODULE\.py|/PATH|sheetN)', '', pat)
        semantic = re.sub(r'\s+', ' ', semantic).strip()
        key_phrase = _extract_key_phrase(semantic)
        
        merge_key = f"{cat}:{key_phrase}"
        if merge_key not in seen:
            seen[merge_key] = {
                'category': cat,
                'count': 1,
                'task_ids': {e.get('task_id', '?')},
                'sample': e,
                'samples': [e.get('message', '')[:120]],
            }
        else:
            seen[merge_key]['count'] += 1
            seen[merge_key]['task_ids'].add(e.get('task_id', '?'))
            # 保留前3个不同示例
            if len(seen[merge_key]['samples']) < 3:
                sample_text = e.get('message', '')[:120]
                if sample_text not in seen[merge_key]['samples']:
                    seen[merge_key]['samples'].append(sample_text)

    return sorted(seen.values(), key=lambda x: -x['count'])


# ═══════════════════════════════════════════════════════════════
#  Word 文档生成
# ═══════════════════════════════════════════════════════════════

def _build_document_base():
    """创建基础文档，设置默认样式"""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    return doc


def _write_fault_section(doc, fault_data: dict, idx: int = None, total: int = None):
    """写入单个故障的五段式分析"""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Section header
    label = f'故障 {idx}：{fault_data["error_type"]}' if idx else fault_data['error_type']
    doc.add_heading(label, level=1)

    # 信息子表
    sub = doc.add_table(rows=3, cols=2, style='Light Grid Accent 1')
    for i, (k, v) in enumerate([
        ('故障级别', fault_data.get('fault_level', 'P?')),
        ('故障时间', fault_data.get('fault_time', '—')),
        ('错误类别', fault_data.get('error_category', '待分类')),
    ]):
        sub.rows[i].cells[0].text = k
        sub.rows[i].cells[1].text = v

    doc.add_paragraph('')

    # 一、故障概述
    doc.add_heading('故障概述', level=2)
    for label in ['affected_component', 'affected_scope', 'related_config']:
        val = fault_data.get(label, '')
        if val:
            doc.add_paragraph(f'{label}：{val}' if label == 'related_config'
                              else f'影响组件：{val}' if label == 'affected_component'
                              else f'影响范围：{val}')

    # 二、故障详情
    doc.add_heading('故障详情', level=2)
    doc.add_heading('错误消息', level=3)
    doc.add_paragraph(fault_data.get('error_message', '—'))

    stack = fault_data.get('stack_trace', '')
    if stack:
        doc.add_heading('调用链', level=3)
        for line in stack.split('\n')[:15]:
            doc.add_paragraph(line, style='No Spacing')

    # 三、根因分析
    doc.add_heading('根因分析', level=2)
    doc.add_paragraph(fault_data.get('root_cause', '待进一步分析'))

    evidence = fault_data.get('evidence_lines', [])
    if evidence:
        doc.add_heading('关键证据', level=3)
        for i, line in enumerate(evidence, 1):
            doc.add_paragraph(f'{i}. {line}')

    # 四、解决方案
    doc.add_heading('解决方案', level=2)
    for step in fault_data.get('solution_steps', []):
        doc.add_paragraph(step, style='List Number')

    verif = fault_data.get('verification', '')
    if verif:
        doc.add_heading('验证方法', level=3)
        doc.add_paragraph(verif)

    # 五、预防措施
    doc.add_heading('预防措施', level=2)
    doc.add_paragraph(fault_data.get('prevention', '—'))


def generate_fault_report(output_path: str, fault_data: dict) -> str:
    """生成单故障报告（五段式标准结构）"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = _build_document_base()

    title = doc.add_heading('部署故障分析及解决方案', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')

    _write_fault_section(doc, fault_data)

    doc.save(output_path)
    return output_path


def generate_multi_fault_report(output_path: str, faults: list) -> str:
    """生成多故障综合报告（总览 + 逐项 + 综合建议）"""
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = _build_document_base()
    now = datetime.now()

    # ═══ 封面总览 ═══
    title = doc.add_heading('部署故障分析及解决方案', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')

    doc.add_heading('总体概况', level=1)
    summary_table = doc.add_table(rows=5, cols=2, style='Light Grid Accent 1')
    top_level = max((f.get('fault_level', 'P3') for f in faults), key=lambda x: int(x[1]))
    summary_data = [
        ('文档编号', f'DOC-{now.strftime("%Y%m%d")}-MULTI'),
        ('分析时间', now.strftime('%Y-%m-%d %H:%M:%S')),
        ('归并根因数', str(len(faults))),
        ('最高故障级别', top_level),
        ('分析引擎', 'deploy-fault-analyzer v2.0'),
    ]
    for i, (k, v) in enumerate(summary_data):
        summary_table.rows[i].cells[0].text = k
        summary_table.rows[i].cells[1].text = v

    doc.add_paragraph('')

    # ═══ 错误全景图 ═══
    doc.add_heading('错误全景图', level=1)
    pano = doc.add_table(rows=len(faults)+1, cols=5, style='Light Grid Accent 1')
    for j, h in enumerate(['#', '故障', '类别', '级别', '出现次数']):
        pano.rows[0].cells[j].text = h
        for p in pano.rows[0].cells[j].paragraphs:
            for run in p.runs:
                run.bold = True

    for i, f in enumerate(faults, 1):
        row_data = [
            f'F{i}',
            f.get('error_type', '?')[:40],
            f.get('error_category', '?'),
            f.get('fault_level', '?'),
            str(f.get('count', f.get('出现次数', '?'))),
        ]
        for j, val in enumerate(row_data):
            pano.rows[i].cells[j].text = val

    doc.add_page_break()

    # ═══ 逐项分析 ═══
    for idx, f in enumerate(faults, 1):
        _write_fault_section(doc, f, idx=idx, total=len(faults))
        if idx < len(faults):
            doc.add_paragraph('─' * 50)
            doc.add_paragraph('')

    # ═══ 综合建议 ═══
    doc.add_page_break()
    doc.add_heading('综合建议与整改清单', level=1)
    
    # Collect all prevention items
    all_preventions = []
    for i, f in enumerate(faults, 1):
        prev = f.get('prevention', '')
        if prev and prev != '—':
            all_preventions.append(f'（F{i}）{prev}')
    
    for i, item in enumerate(all_preventions, 1):
        doc.add_paragraph(f'{i}. {item}', style='List Number')

    doc.save(output_path)
    return output_path


# ═══════════════════════════════════════════════════════════════
#  Excel 知识库
# ═══════════════════════════════════════════════════════════════

def append_to_excel(xlsx_path: str, record: dict) -> str:
    """追加一条记录到 Excel 知识库"""
    import openpyxl

    xlsx_path = Path(xlsx_path)
    if xlsx_path.exists():
        wb = openpyxl.load_workbook(xlsx_path)
        ws = wb.active
    else:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.append(['序号', '记录时间', '故障来源', '故障类别', '故障级别',
                    '错误摘要', '影响组件', '根因分析', '解决方案摘要', '关联文档'])

    next_id = ws.max_row  # 含表头，正好做序号
    ws.append([
        next_id,
        record.get('record_time', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
        record.get('fault_source', ''),
        record.get('error_category', ''),
        record.get('fault_level', ''),
        record.get('error_summary', ''),
        record.get('affected_component', ''),
        record.get('root_cause', ''),
        record.get('solution_summary', ''),
        record.get('doc_path', ''),
    ])

    wb.save(xlsx_path)
    return str(xlsx_path)


def batch_append_excel(xlsx_path: str, records: list) -> str:
    """批量追加多条记录"""
    for rec in records:
        append_to_excel(xlsx_path, rec)
    return str(xlsx_path)


# ═══════════════════════════════════════════════════════════════
#  CLI
# ═══════════════════════════════════════════════════════════════

if __name__ == '__main__':
    import argparse
    parser = argparse.ArgumentParser(description='部署故障分析报告生成器 v2.0')
    parser.add_argument('--json', help='JSON 日志文件路径（直接解析）')
    parser.add_argument('--data', help='fault_data JSON 文件路径')
    parser.add_argument('--output', required=True, help='输出 .docx 路径')
    parser.add_argument('--xlsx', default='', help='Problems .xlsx 路径')
    parser.add_argument('--batch', action='store_true', help='批量多故障模式')
    args = parser.parse_args()

    if args.json:
        # 自动解析 JSON 日志
        parsed = parse_deploy_json_log(args.json)
        root_causes = deduplicate_errors(parsed['errors'])
        print(f"解析完成: {parsed['total']} 条日志, {len(parsed['errors'])} 错误, "
              f"归并 {len(root_causes)} 个根因")

        # 构建 fault_data 列表
        faults = []
        for rc in root_causes:
            sample = rc['sample']
            task_info = parse_task_id(sample.get('task_id'))
            faults.append({
                'fault_time': sample.get('timestamp', ''),
                'error_type': f"{task_info['phase']}_{task_info['type']}: {rc['category']}",
                'error_message': sample.get('message', '')[:200],
                'stack_trace': '',
                'affected_component': '/home/conf',
                'affected_scope': '部署流程',
                'related_config': '',
                'error_category': rc['category'],
                'fault_level': 'P1' if rc['count'] > 5 else 'P2',
                'root_cause': f"（共{rc['count']}条重复）从 task_ids={rc['task_ids']} 提取",
                'evidence_lines': [],
                'solution_steps': ['待 AI 分析引擎补充'],
                'verification': '',
                'prevention': '',
                'count': rc['count'],
            })

        with open(args.output.replace('.docx', '.json'), 'w') as f:
            json.dump(faults, f, ensure_ascii=False, indent=2)

        if len(faults) > 1:
            result = generate_multi_fault_report(args.output, faults)
        else:
            result = generate_fault_report(args.output, faults[0])
        print(f'报告生成: {result}')

    elif args.data:
        with open(args.data, 'r') as f:
            fault_data = json.load(f)

        if args.batch or isinstance(fault_data, list):
            faults = fault_data if isinstance(fault_data, list) else [fault_data]
            if len(faults) > 1:
                result = generate_multi_fault_report(args.output, faults)
            else:
                result = generate_fault_report(args.output, faults[0])
        else:
            result = generate_fault_report(args.output, fault_data)
        print(f'报告生成: {result}')

    if args.xlsx:
        xlsx_result = batch_append_excel(args.xlsx, faults)
        print(f'Excel 更新: {xlsx_result}')
