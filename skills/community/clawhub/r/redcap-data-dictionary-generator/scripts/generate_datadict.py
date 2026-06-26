#!/usr/bin/env python3
"""
REDCap CRF 数据字典生成器
将 Word 文档（CRF/方案）转换为 CSV 格式的 REDCap 数据字典

Section Header 规则：仅第一个字段填写，后续同组字段留空

依赖: python-docx, lxml
安装: pip install python-docx lxml
"""

import csv
import io
import re
import sys
import argparse
from pathlib import Path

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# REDCap 数据字典 CSV 列名（按标准顺序）
REDCAP_HEADERS = [
    "Variable / Field Name",
    "Form Name",
    "Section Header",
    "Field Type",
    "Field Label",
    "Choices, Calculations, OR Slider Labels",
    "Field Note",
    "Text Validation Type OR Show Slider Number",
    "Text Validation Min",
    "Text Validation Max",
    "Identifier?",
    "Branching Logic (Show field only if...)",
    "Required Field?",
    "Custom Alignment",
    "Question Number (surveys only)",
    "Matrix Group Name"
]


def make_var_name(prefix, q_num):
    """生成变量名"""
    return f"{prefix}_{q_num}"


def add_field(fields, var_name, form, section_header, ftype, label, 
              choices='', note='', required='', qnum='', identifier=''):
    """
    添加字段到列表
    
    ⚠️ 重要：section_header 仅第一个字段填写，后续同组字段留空
    调用时每个分组的第一个字段传 section_header，后续字段传 ''
    """
    fields.append({
        'Variable / Field Name': var_name,
        'Form Name': form,
        'Section Header': section_header,  # Only first field gets section header
        'Field Type': ftype,
        'Field Label': label,
        'Choices, Calculations, OR Slider Labels': choices,
        'Field Note': note,
        'Text Validation Type OR Show Slider Number': '',
        'Text Validation Min': '',
        'Text Validation Max': '',
        'Identifier?': identifier,
        'Branching Logic (Show field only if...)': '',
        'Required Field?': required,
        'Custom Alignment': '',
        'Question Number (surveys only)': qnum,
        'Matrix Group Name': ''
    })


def parse_docx(docx_path):
    """
    解析 DOCX 文件，返回字段列表
    
    Section Header 规则：每个分组仅第一个字段填写 Section Header，
    后续同组字段该列留空。
    """
    if not DOCX_AVAILABLE:
        raise ImportError("需要安装 python-docx: pip install python-docx")
    
    doc = Document(docx_path)
    fields = []
    
    # ========== 第一部分：基本职业情况（从段落解析）==========
    form1 = "basic_info"
    section1 = "人口学特征"
    
    basic_items = [
        ('bi_1', '您的性别', 'radio', '1, 男 | 2, 女', 'y', '1'),
        ('bi_2', '您的年龄', 'radio', '1, 30岁及以下 | 2, 31-40 | 3, 41-50 | 4, 51-60 | 5, 60岁以上', '', '2'),
        ('bi_3', '您的户籍', 'radio', '1, 本市 | 2, 外省市', '', '3'),
        ('bi_4', '您的婚姻状况', 'radio', '1, 未婚 | 2, 已婚 | 3, 离婚 | 4, 丧偶 | 5, 其他', '', '4'),
        ('bi_5', '您的最高学历', 'radio', '1, 博/硕士 | 2, 本科 | 3, 大专(高职) | 4, 中专(高中)', '', '5'),
        ('bi_6', '您目前的专业职称', 'radio', '1, 正高 | 2, 副高 | 3, 中级 | 4, 初级', '', '6'),
        ('bi_7', '您获得当前专业职称距今已有多久', 'radio', '1, 2年及以下 | 2, 3-5年 | 3, 6-8年 | 4, 9年及以上', '', '7'),
        ('bi_8', '您目前从事的工作', 'radio', '1, 医生 | 2, 护理 | 3, 医技人员 | 4, 行政后勤 | 5, 其他', '', '8'),
        ('bi_9', '您所在科室', 'radio', '1, 内科 | 2, 外科 | 3, 妇产科 | 4, 儿科 | 5, 急诊科或ICU | 6, 医辅科室 | 7, 行政科室 | 8, 后勤科室', '', '9'),
        ('bi_10', '您的工龄', 'radio', '1, 5年及以下 | 2, 6-10年 | 3, 11-20年 | 4, 21-30年 | 5, 30年以上', '', '10'),
        ('bi_11', '您的编制情况', 'radio', '1, 在编职工 | 2, 非在编职工 | 3, 派遣人员 | 4, 其他', '', '11'),
        ('bi_12', '您所在医院', 'radio', '1, 综合性医院 | 2, 专科医院', '', '12'),
        ('bi_13', '您所在医院的级别', 'radio', '1, 三甲 | 2, 三乙 | 3, 二甲 | 4, 二甲以下', '', '13'),
        ('bi_14', '您认为自己现在是否处于亚健康状态？', 'radio', '1, 是 | 2, 否 | 3, 说不清', '', '14'),
        ('bi_15', '您现在是否患有以下慢性病？（可多选）', 'checkbox', '0, 均没有 | 1, 高血压 | 2, 慢性胃肠炎 | 3, 糖尿病 | 4, 类风湿关节炎 | 5, 脑血管病 | 6, 椎间盘疾病 | 7, 慢性阻塞性肺病 | 8, 高脂血症 | 9, 胆结石胆囊炎 | 10, 脂肪肝 | 11, 肿瘤 | 12, 其它', '', '15'),
    ]
    
    first_in_section = True
    for var, label, ftype, choices, req, qnum in basic_items:
        section = section1 if first_in_section else ''
        add_field(fields, var, form1, section, ftype, label, choices, required=req, qnum=qnum)
        first_in_section = False
    
    # ========== 第二部分：职业压力主观感受（表格1）==========
    form2 = "job_stress"
    section2 = "职业压力（Likert量表）"
    
    # 反向计分题号
    reverse_qs = {1, 2, 4, 8, 10, 13, 15, 17, 19, 21, 22, 24, 25, 26, 29, 31}
    likert_forward = "1, 非常符合 | 2, 基本符合 | 3, 有点不符合 | 4, 非常不符合"
    likert_reverse = "1, 非常不符合 | 2, 有点不符合 | 3, 基本符合 | 4, 非常符合"
    
    first_in_section = True
    for row in doc.tables[0].rows[1:32]:
        cells = [c.text.strip() for c in row.cells]
        if cells[0].isdigit():
            q_num = int(cells[0])
            question = cells[1]
            choices = likert_reverse if q_num in reverse_qs else likert_forward
            section = section2 if first_in_section else ''
            note = '反向计分' if q_num in reverse_qs else ''
            add_field(fields, f"js_{q_num}", form2, section, 'radio', question, 
                     choices, note, qnum=str(q_num))
            first_in_section = False
    
    # ========== 第三部分：工作状态的主观感受（表格2）==========
    form3 = "work_status"
    section3 = "工作状态（频率量表）"
    
    freq_choices = "0, 从未发生 | 1, 一年几次或更少 | 2, 一月一次或更少 | 3, 一月几次 | 4, 一周一次 | 5, 一周几次 | 6, 每天发生"
    
    first_in_section = True
    for row in doc.tables[1].rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        if cells[0].isdigit() and cells[0] and '从未发生' not in cells[1]:
            q_num = int(cells[0])
            question = cells[1]
            section = section3 if first_in_section else ''
            add_field(fields, f"ws_{q_num}", form3, section, 'radio', question, 
                     freq_choices, qnum=str(q_num))
            first_in_section = False
    
    # ========== 第四部分：职业心理健康评价 ===========
    form4 = "mental_health"
    section4 = "心理健康评价与调适需求"
    
    q4_items = [
        ('mh_1', '我对目前工作压力总体感受为', 'radio', '1, 基本没有 | 2, 比较小 | 3, 一般 | 4, 比较大 | 5, 非常大'),
        ('mh_2', '在过去一年内，您是否能有效应对压力？', 'radio', '1, 完全能 | 2, 大多数时间能 | 3, 少数时间能 | 4, 完全不能'),
        ('mh_3', '我目前工作压力来自（可多选，最多可选3项）', 'checkbox', 'A, 工作强度大 | B, 工作内容单调乏味 | C, 工作难度大 | D, 工作环境差 | E, 工作与生活角色转换困难 | F, 对管理制度不认同 | G, 对组织文化不认同 | H, 与同事相处困难时有冲突 | I, 与领导相处困难 | J, 社会不理解和媒体负面报道 | K, 医患矛盾和患者投诉 | L, 职业发展前途迷茫 | M, 薪酬不能体现工作价值'),
        ('mh_4', '我目前生活压力来自（可多选，最多可选3项）', 'checkbox', 'A, 购房与贷款压力 | B, 婚恋压力 | C, 情感问题 | D, 赡养老人压力 | E, 照顾小孩压力 | F, 人际交往压力 | G, 自身健康问题 | H, 家人身体健康 | I, 子女教育及就业 | J, 家庭财务收支压力 | K, 生活与工作协调'),
        ('mh_5', '当压力很大时，我会通过哪些方式调适（可多选，最多可选2项）', 'checkbox', 'A, 自我调节（运动、美食、睡觉等） | B, 找家人朋友聊天倾诉 | C, 参加单位文体活动或心理调适技巧培训 | D, 找上级领导或工会寻求帮助 | E, 寻找专业心理咨询师辅导 | F, 寻找专科医生咨询'),
        ('mh_6', '您所在医院一般通过什么方式缓解职工心理压力（可多选，最多可选3项）', 'checkbox', 'A, 个别谈心 | B, 小组座谈 | C, 组织开展文体活动 | D, 心理疏导活动（巴林特小组、正念减压等） | E, 其他'),
        ('mh_7', '您希望医院采取何种方式来让员工认识和缓解工作压力（可多选，最多可选3项）', 'checkbox', 'A, 完善职业规划和晋升通道 | B, 完善激励制度（薪酬、福利等） | C, 健全管理制度 | D, 畅通职工参政议政渠道 | E, 改善工作环境和条件 | F, 提供心理疏导及咨询 | G, 提供心理专题讲座或课程 | H, 提供文体活动解压'),
        ('mh_8', '对于H-EAP的了解程度', 'radio', 'A, 从未听说 | B, 听说过但不了解 | C, 了解一些但没参与过 | D, 了解并参与过'),
        ('mh_9', '您认为医院是否有必要开展系统性的员工心理关爱活动？', 'radio', 'A, 有必要 | B, 如果有当然好 | C, 说不清 | D, 没有必要'),
        ('mh_10', '若医院开展EAP项目，您最希望获得哪方面帮助（可多选，最多可选3项）', 'checkbox', 'A, 个人成长规划 | B, 个人专业素养 | C, 危机事务援助 | D, 休闲娱乐活动 | E, 医疗保健服务 | F, 儿童/老人照顾 | G, 课程培训或沙龙 | H, 心理疏导'),
    ]
    
    first_in_section = True
    for var, label, ftype, choices in q4_items:
        section = section4 if first_in_section else ''
        add_field(fields, var, form4, section, ftype, label, choices)
        first_in_section = False
    
    return fields


def generate_csv(fields, output_path=None):
    """生成 CSV 内容"""
    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=REDCAP_HEADERS)
    writer.writeheader()
    
    for field in fields:
        writer.writerow(field)
    
    csv_content = output.getvalue()
    
    if output_path:
        with open(output_path, 'w', encoding='utf-8-sig', newline='') as f:
            f.write(csv_content)
        return output_path
    
    return csv_content


def main():
    parser = argparse.ArgumentParser(description='REDCap CRF 数据字典生成器')
    parser.add_argument('file_path', help='Word 文档路径 (.docx)')
    parser.add_argument('-o', '--output', help='输出 CSV 文件路径')
    parser.add_argument('--preview', action='store_true', help='预览前20行')
    
    args = parser.parse_args()
    
    file_path = Path(args.file_path)
    
    if not file_path.exists():
        print(f"错误: 文件不存在 - {args.file_path}")
        sys.exit(1)
    
    print(f"正在解析: {file_path}")
    
    try:
        fields = parse_docx(str(file_path))
        print(f"解析完成，共 {len(fields)} 个字段")
        
        output_path = args.output or str(file_path.with_suffix('_data_dict.csv'))
        result_path = generate_csv(fields, output_path)
        
        print(f"CSV 已生成: {result_path}")
        
        # 统计各表单字段数
        forms = {}
        for f in fields:
            form = f['Form Name']
            forms[form] = forms.get(form, 0) + 1
        print("各表单字段:", forms)
        
        if args.preview:
            csv_content = generate_csv(fields)
            lines = csv_content.split('\n')
            print("\n=== 预览（前20行）===")
            for line in lines[:20]:
                cols = line.split(',')
                if len(cols) >= 4:
                    print(f"{cols[0][:15]:15} | {cols[1][:12]:12} | {cols[2][:20]:20} | {cols[3][:8]:8}")
        
    except Exception as e:
        print(f"处理失败: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()