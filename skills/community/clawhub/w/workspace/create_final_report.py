#!/usr/bin/env python3
"""
Create final formatted report from COMPLETE document
"""

import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_final_report():
    print("📝 Creating final formatted report...")
    
    # Load the COMPLETE document
    source_doc = "BSc-project-report-COMPLETE.docx"
    if not os.path.exists(source_doc):
        print(f"❌ Source document not found: {source_doc}")
        return False
    
    print(f"📄 Loading: {source_doc}")
    try:
        doc = Document(source_doc)
        print(f"✅ Loaded document with {len(doc.paragraphs)} paragraphs")
    except Exception as e:
        print(f"❌ Error loading document: {e}")
        return False
    
    # Apply formatting to all paragraphs
    print("🎨 Applying formatting...")
    
    # Set font size and justification
    for para in doc.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in para.runs:
            run.font.size = Pt(14)
    
    # Save final document
    output_file = "BSc-project-report-FINAL-SUBMISSION.docx"
    doc.save(output_file)
    
    print(f"\n🎉 Final report saved: {output_file}")
    print(f"📏 File size: {os.path.getsize(output_file):,} bytes")
    
    # Verification
    print("\n🔍 Verification:")
    chapters_found = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith("فصل") or "فصل" in text[:20]:
            chapters_found.append(text)
    
    print(f"✅ Chapters found: {len(chapters_found)}")
    for i, chapter in enumerate(chapters_found[:10]):
        print(f"   {i+1}. {chapter}")
    
    if len(chapters_found) >= 7:
        print("✅ ALL 7 CHAPTERS PRESENT!")
    else:
        print(f"⚠️  Only {len(chapters_found)} chapters found")
    
    return True

if __name__ == "__main__":
    create_final_report()
