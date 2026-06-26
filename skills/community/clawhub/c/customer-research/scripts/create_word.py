#!/usr/bin/env python3
"""
生成客户调研 Word 报告
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys
from datetime import datetime

def create_word_report(customer_name, results, output_path):
    """创建 Word 调研报告
    
    参数：
        customer_name: 客户名称
        results: dict，键为调研维度，值为搜索结果列表
                 必须包含：AI 相关领导发言、AI 相关招标记录、数据相关内容、基本情况与最新动态
        output_path: 输出文件路径
    """
    
    doc = Document()
    
    # 标题
    title = doc.add_heading(f'客户调研报告：{customer_name}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 基本信息
    doc.add_paragraph(f'调研时间：{datetime.now().strftime("%Y 年 %m 月 %d 日")}')
    doc.add_paragraph('调研人：销售智能体')
    doc.add_paragraph()
    
    # 执行摘要
    doc.add_heading('📌 执行摘要', level=1)
    
    # 统计有内容的维度
    filled_dimensions = {k: v for k, v in results.items() if v and any(v)}
    doc.add_paragraph(f'本次调研围绕 {customer_name} 的 AI 相关领导发言、招标记录、数据相关内容及基本情况展开，共获取 {len(filled_dimensions)} 个维度的信息。')
    doc.add_paragraph()
    
    # 各维度详情
    sections = [
        ("1️⃣ AI 相关领导发言", "AI 相关领导发言"),
        ("2️⃣ AI 相关招标记录", "AI 相关招标记录"),
        ("3️⃣ 数据相关内容", "数据相关内容"),
        ("4️⃣ 基本情况与最新动态", "基本情况与最新动态")
    ]
    
    for emoji_name, key_name in sections:
        doc.add_heading(emoji_name, level=2)
        items = results.get(key_name, [])
        if items and any(items):
            for i, item in enumerate(items, 1):
                p = doc.add_paragraph(style='List Number')
                # 清理文本，只保留前 300 字符
                clean_text = item[:300].replace('\n', ' ').strip()
                p.add_run(clean_text)
        else:
            doc.add_paragraph('暂无相关公开信息，需通过拜访确认', style='Intense Quote')
        doc.add_paragraph()
    
    # 销售机会分析
    doc.add_heading('🎯 销售机会分析', level=1)
    doc.add_paragraph('注：以下内容基于调研结果推断，需拜访确认。', style='Intense Quote')
    doc.add_paragraph()
    
    doc.add_paragraph('决策链（需拜访确认）：', style='Heading 2')
    doc.add_paragraph('• EB（经济决策人）：待拜访确认')
    doc.add_paragraph('• Technical Buyer：信息科主任/CTO（待拜访确认）')
    doc.add_paragraph('• Champion：待识别')
    
    doc.add_paragraph('切入建议：', style='Heading 2')
    doc.add_paragraph('• 首访目标：信息科/IT 部门负责人，了解现有系统情况', style='List Bullet')
    doc.add_paragraph('• 切入点：结合调研发现的具体痛点展开', style='List Bullet')
    doc.add_paragraph('• MEDDIC 行动：确认预算、决策流程、培养 Champion', style='List Bullet')
    
    # 信息来源
    doc.add_heading('📎 信息来源', level=1)
    doc.add_paragraph('本报告基于公开网络搜索整理，主要来源包括：')
    doc.add_paragraph('• 百度、Bing 等搜索引擎', style='List Bullet')
    doc.add_paragraph('• 客户官网及公开新闻报道', style='List Bullet')
    doc.add_paragraph('• 中国政府采购网、各省市公共资源交易中心（招标信息）', style='List Bullet')
    doc.add_paragraph('')
    doc.add_paragraph('注：所有信息均来自公开渠道，决策链等推断内容需通过拜访确认。')
    
    # 保存文档
    doc.save(output_path)
    print(f"Word 报告已生成：{output_path}")
    return output_path

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("用法：python create_word.py <客户名称> <输出路径>")
        print("注意：results 数据应通过 agent framework 搜索后传入")
        sys.exit(1)
    
    customer_name = sys.argv[1]
    output_path = sys.argv[2]
    
    # 空的 results 字典，由调用方通过 agent 搜索后传入
    # 禁止在脚本中硬编码客户信息
    results = {
        "AI 相关领导发言": [],
        "AI 相关招标记录": [],
        "数据相关内容": [],
        "基本情况与最新动态": []
    }
    
    create_word_report(customer_name, results, output_path)