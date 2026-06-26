#!/usr/bin/env python3
"""
Proper document merge with correct encoding for Persian text
"""

import os
import sys
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def read_text_file(filename):
    """Read text file with proper encoding"""
    try:
        with open(filename, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        print(f"❌ Error reading {filename}: {e}")
        return None

def create_proper_document():
    """Create a proper Word document with all 7 chapters"""
    
    print("📝 Creating proper document with correct encoding...")
    
    # Check all required files
    source_doc = "BSc-project-report-ORIGINAL.docx"
    if not os.path.exists(source_doc):
        print(f"❌ Source document not found: {source_doc}")
        return False
    
    # Load the original document (chapters 1-4)
    print(f"📄 Loading original document: {source_doc}")
    try:
        doc = Document(source_doc)
        print(f"✅ Loaded document with {len(doc.paragraphs)} paragraphs")
    except Exception as e:
        print(f"❌ Error loading document: {e}")
        return False
    
    # Chapters to add
    chapters = [
        (5, "فصل ۵: پیاده‌سازی لایه سرویس‌های در سطح کاربر", "chapter5_clean.txt"),
        (6, "فصل ۶: پایش، سخت‌سازی امنیتی، آزمون‌ها و تحویل نهایی", "chapter6_clean.txt"),
        (7, "فصل ۷: خلصه، نتیجه‌گیری و کارهای آینده", "chapter7_clean.txt")
    ]
    
    # Add each chapter
    for chap_num, title, filename in chapters:
        print(f"\n📖 Processing Chapter {chap_num}: {title}")
        
        # Read chapter content
        content = read_text_file(filename)
        if content is None:
            print(f"❌ Skipping Chapter {chap_num}")
            continue
        
        print(f"✅ Read {len(content)} characters from {filename}")
        
        # Add page break
        doc.add_page_break()
        
        # Add chapter title
        title_para = doc.add_paragraph(title)
        title_para.style = 'Heading 1'
        title_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add content
        # Split content into paragraphs
        lines = content.split('\n')
        for line in lines:
            if line.strip():  # Skip empty lines
                para = doc.add_paragraph(line.strip())
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                # Set font size
                for run in para.runs:
                    run.font.size = Pt(14)
        
        print(f"✅ Added Chapter {chap_num}")
    
    # Save the new document
    output_file = "BSc-project-report-PROPER.docx"
    doc.save(output_file)
    
    print(f"\n🎉 Document saved: {output_file}")
    print(f"📏 File size: {os.path.getsize(output_file):,} bytes")
    
    return True

def check_current_document():
    """Check current document state"""
    print("\n🔍 Checking current document state...")
    
    files = [
        "BSc-project-report-ORIGINAL.docx",
        "BSc-project-report-FINAL-UNIVERSITY.docx",
        "BSc-project-report-PROPER.docx"
    ]
    
    for f in files:
        if os.path.exists(f):
            size = os.path.getsize(f)
            print(f"📄 {f}: {size:,} bytes")
        else:
            print(f"❌ {f}: Not found")
    
    # Check text files
    print("\n📋 Text files:")
    for chap in [5, 6, 7]:
        f = f"chapter{chap}_clean.txt"
        if os.path.exists(f):
            size = os.path.getsize(f)
            print(f"✅ {f}: {size:,} bytes")
        else:
            print(f"❌ {f}: Missing")

def main():
    print("🔧 PROPER DOCUMENT MERGE TOOL")
    print("==============================")
    
    # Check current state
    check_current_document()
    
    print("\n" + "="*50)
    print("🚀 Creating properly encoded document...")
    print("="*50)
    
    # Create proper document
    success = create_proper_document()
    
    if success:
        print("\n" + "="*50)
        print("✅ SUCCESS: Proper document created!")
        print("="*50)
        
        print("\n🎯 NEXT STEPS:")
        print("1. Open BSc-project-report-PROPER.docx in Word")
        print("2. Verify all 7 chapters are present")
        print("3. Verify Persian text displays correctly")
        print("4. Apply final formatting if needed")
        print("5. Save as final version")
        
        print("\n⚠️  IMPORTANT:")
        print("- This uses python-docx library for proper encoding")
        print("- Persian text should display correctly")
        print("- Formatting (justification, font size) is applied")
    else:
        print("\n❌ FAILED: Could not create proper document")
        print("\n🔧 ALTERNATIVE SOLUTION:")
        print("1. Open original document in Word")
        print("2. Open each text file in TextEdit (not terminal)")
        print("3. Copy from TextEdit → Paste in Word")
        print("4. This preserves encoding better")

if __name__ == "__main__":
    main()