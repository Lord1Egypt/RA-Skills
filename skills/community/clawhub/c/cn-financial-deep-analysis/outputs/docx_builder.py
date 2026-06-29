#!/usr/bin/env python3
"""
cn-financial-deep-analysis Word 文档生成插件

将 Markdown 格式的研报转换为 Word 文档，可选嵌入图表。

用法：
    python3 docx_builder.py --input report.md --output report.docx
    python3 docx_builder.py --input report.md --output report.docx --charts-dir ./charts/
"""

import argparse
import os
import re
import sys

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def _set_cell_shading(cell, color_hex: str):
    """设置单元格底色。"""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex,
        qn("w:val"): "clear",
    })
    shading.append(shading_elm)


def _add_styled_paragraph(doc, text: str, style: str, bold: bool = False, font_size: int = None):
    """添加带样式的段落。"""
    p = doc.add_paragraph(text, style=style)
    if bold or font_size:
        for run in p.runs:
            if bold:
                run.bold = True
            if font_size:
                run.font.size = Pt(font_size)
    return p


def build_docx(md_path: str, output_path: str, charts_dir: str = None):
    """将 Markdown 研报转为 Word 文档。"""

    with open(md_path, "r", encoding="utf-8") as f:
        md_content = f.read()

    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    lines = md_content.split("\n")
    i = 0

    in_table = False
    table_rows = []
    in_code_block = False

    while i < len(lines):
        line = lines[i]

        # 代码块处理
        if line.strip().startswith("```"):
            in_code_block = not in_code_block
            i += 1
            continue
        if in_code_block:
            i += 1
            continue

        # 表格处理
        if line.strip().startswith("|") and line.strip().endswith("|"):
            if not in_table:
                in_table = True
                table_rows = []

            # 跳过分隔行
            if re.match(r"^\|[\s\-:|]+\|$", line.strip()):
                i += 1
                continue

            cells = [c.strip() for c in line.strip().split("|")[1:-1]]
            table_rows.append(cells)
            i += 1
            continue
        else:
            # 结束表格，渲染
            if in_table and table_rows:
                col_count = max(len(r) for r in table_rows)
                table = doc.add_table(rows=len(table_rows), cols=col_count)
                table.style = "Table Grid"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                for ri, row_data in enumerate(table_rows):
                    for ci, cell_text in enumerate(row_data):
                        if ci < col_count:
                            cell = table.rows[ri].cells[ci]
                            cell.text = cell_text
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(9)
                            # 表头着色
                            if ri == 0:
                                _set_cell_shading(cell, "2563EB")
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.color.rgb = RGBColor(255, 255, 255)
                                        run.bold = True
                doc.add_paragraph()  # 表后空行
                table_rows = []
                in_table = False
            # 继续处理当前行

        # 标题
        if line.startswith("# ") and not line.startswith("## "):
            _add_styled_paragraph(doc, line[2:].strip(), "Normal", bold=True, font_size=20)
            doc.add_paragraph()
        elif line.startswith("## "):
            _add_styled_paragraph(doc, line[3:].strip(), "Normal", bold=True, font_size=16)
        elif line.startswith("### "):
            _add_styled_paragraph(doc, line[4:].strip(), "Normal", bold=True, font_size=13)
        elif line.startswith("#### "):
            _add_styled_paragraph(doc, line[5:].strip(), "Normal", bold=True, font_size=11)
        elif line.strip().startswith("- ") or line.strip().startswith("* "):
            text = re.sub(r"^[\-\*]\s+", "", line.strip())
            p = doc.add_paragraph(text, style="List Bullet")
            for run in p.runs:
                run.font.size = Pt(10)
        elif line.strip() and not line.startswith("|"):
            p = doc.add_paragraph(line.strip())
            for run in p.runs:
                run.font.size = Pt(10)

        i += 1

    # 处理文件末尾未关闭的表格
    if in_table and table_rows:
        col_count = max(len(r) for r in table_rows)
        table = doc.add_table(rows=len(table_rows), cols=col_count)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ri, row_data in enumerate(table_rows):
            for ci, cell_text in enumerate(row_data):
                if ci < col_count:
                    cell = table.rows[ri].cells[ci]
                    cell.text = cell_text
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)
                    if ri == 0:
                        _set_cell_shading(cell, "2563EB")
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.color.rgb = RGBColor(255, 255, 255)
                                run.bold = True

    # 嵌入图表
    if charts_dir and os.path.isdir(charts_dir):
        chart_files = sorted([f for f in os.listdir(charts_dir) if f.endswith(".png")])
        if chart_files:
            doc.add_page_break()
            _add_styled_paragraph(doc, "附：财务图表", "Normal", bold=True, font_size=16)
            for cf in chart_files:
                filepath = os.path.join(charts_dir, cf)
                doc.add_paragraph(cf.replace(".png", ""), style="Normal").runs[0].bold = True
                doc.add_picture(filepath, width=Inches(5.5))
                doc.add_paragraph()
    elif charts_dir:
        print(f"警告：图表目录不存在 {charts_dir}", file=sys.stderr)

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    print(f"已生成 Word 文档：{output_path}")


def main():
    parser = argparse.ArgumentParser(description="将 Markdown 研报转为 Word 文档")
    parser.add_argument("--input", "-i", required=True, help="Markdown 研报文件路径")
    parser.add_argument("--output", "-o", required=True, help="Word 文档输出路径")
    parser.add_argument("--charts-dir", "-c", default=None, help="图表目录（可选，嵌入图表）")
    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"错误：输入文件不存在 {args.input}", file=sys.stderr)
        sys.exit(1)

    build_docx(args.input, args.output, args.charts_dir)


if __name__ == "__main__":
    main()
