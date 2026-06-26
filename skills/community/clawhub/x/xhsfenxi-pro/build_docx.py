"""
build_docx.py — 通用黑体 Word 生成器
将 Markdown 文件转换为全黑体样式的 Word 文档
支持：标题层级 / 表格 / 列表 / 引用 / 代码块 / 分割线

用法（命令行）：
    python3 build_docx.py <md_path> <out_path> [title] [subtitle] [meta]

用法（Python）：
    from xhscosmoskill.scripts.build_docx import build_word
    build_word("/tmp/report.md", "/tmp/report.docx",
               title="B 治愈系代表", subtitle="爆款选题公式")
"""
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("请先安装：pip install python-docx")
    sys.exit(1)

# ── 配色 ──────────────────────────────────────────────────────
C_BRAND  = RGBColor(0x1A, 0x1A, 0x1A)
C_H1     = RGBColor(0x0D, 0x0D, 0x0D)
C_H2     = RGBColor(0x1A, 0x1A, 0x1A)
C_ACCENT = RGBColor(0x00, 0x80, 0x60)   # Shopify 绿
C_LIGHT  = RGBColor(0x6B, 0x72, 0x80)
C_TH_BG  = RGBColor(0xF3, 0xF4, 0xF6)
C_LINE   = RGBColor(0x22, 0xC5, 0x5E)
HEITI    = "黑体"


# ── 工具函数 ──────────────────────────────────────────────────

def set_font(run, size, bold=False, color=None, italic=False):
    """全黑体设置（中英文统一）"""
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    rf  = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts')
        rPr.insert(0, rf)
    for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rf.set(qn(attr), HEITI)


def clean(s):
    for p, r in [
        (r'\*\*(.+?)\*\*', r'\1'), (r'\*(.+?)\*', r'\1'),
        (r'\[(.+?)\]\(.+?\)', r'\1'), (r'`(.+?)`', r'\1'),
        (r'^#{1,6}\s+', ''),
    ]:
        s = re.sub(p, r, s)
    return s.strip()


def set_cell_bg(cell, rgb):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}')
    tcPr.append(shd)


def add_rule(doc, color=C_LINE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(6)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), f'{color[0]:02X}{color[1]:02X}{color[2]:02X}')
    pBdr.append(bot)
    pPr.append(pBdr)


# ── 核心构建函数 ──────────────────────────────────────────────

def build_word(md_path: str, out_path: str,
               title: str = "", subtitle: str = "", meta: str = ""):
    """
    将 Markdown 转换为黑体 Word 文档

    参数：
        md_path   — 输入 Markdown 文件路径
        out_path  — 输出 .docx 文件路径
        title     — 封面大标题（可选）
        subtitle  — 封面副标题（可选）
        meta      — 封面元数据行（可选）
    """
    md_text = Path(md_path).read_text(encoding="utf-8")

    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3.0)
        sec.right_margin  = Cm(3.0)

    # 默认样式
    nml = doc.styles['Normal']
    nml.font.name = HEITI
    nml.font.size = Pt(10.5)
    nml._element.rPr.rFonts.set(qn('w:eastAsia'), HEITI)

    # 封面
    if title:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(title), 28, bold=True, color=C_ACCENT)

    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(subtitle), 14, color=C_LIGHT)

    if meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(meta), 9, color=C_LIGHT)

    if title:
        doc.add_page_break()

    # 解析 Markdown
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        # H1
        if re.match(r'^# [^#]', line):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after  = Pt(4)
            set_font(p.add_run(clean(line)), 20, bold=True, color=C_H1)
            i += 1; continue

        # H2
        if re.match(r'^## [^#]', line):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after  = Pt(2)
            set_font(p.add_run(clean(line)), 13, bold=True, color=C_H2)
            add_rule(doc)
            i += 1; continue

        # H3
        if re.match(r'^### ', line):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(3)
            set_font(p.add_run(clean(line)), 11, bold=True, color=C_ACCENT)
            i += 1; continue

        # 引用
        if line.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Cm(1)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            set_font(p.add_run(line[2:].strip()), 10.5, italic=True, color=C_LIGHT)
            i += 1; continue

        # 代码块
        if line.startswith('```'):
            i += 1
            code = []
            while i < len(lines) and not lines[i].startswith('```'):
                code.append(lines[i])
                i += 1
            p = doc.add_paragraph('\n'.join(code))
            p.paragraph_format.left_indent  = Cm(0.8)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            for run in p.runs:
                set_font(run, 9, color=RGBColor(0x37, 0x37, 0x37))
            i += 1; continue

        # 水平线
        if line.strip() == '---':
            add_rule(doc, C_TH_BG)
            i += 1; continue

        # 表格
        if line.strip().startswith('|'):
            tbl_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                if not re.match(r'^\|[\s|:-]+$', lines[i].strip()):
                    tbl_lines.append(lines[i])
                i += 1
            if not tbl_lines:
                continue
            rows = [
                [c.strip() for c in tl.split('|')[1:-1]]
                for tl in tbl_lines
            ]
            rows = [r for r in rows if any(c for c in r)]
            if not rows:
                continue
            ncols = max(len(r) for r in rows)
            tbl = doc.add_table(rows=len(rows), cols=ncols)
            tbl.style     = 'Table Grid'
            tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
            for ri, row in enumerate(rows):
                for ci in range(ncols):
                    cell = tbl.cell(ri, ci)
                    txt  = clean(row[ci]) if ci < len(row) else ''
                    p    = cell.paragraphs[0]
                    p.clear()
                    run  = p.add_run(txt)
                    set_font(run, 9.5,
                             bold=(ri == 0),
                             color=C_H2 if ri == 0 else C_BRAND)
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after  = Pt(2)
                    if ri == 0:
                        set_cell_bg(cell, C_TH_BG)
            doc.add_paragraph()
            continue

        # 列表
        if re.match(r'^[-*] |^\d+\. ', line):
            txt = re.sub(r'^[-*] |^\d+\. ', '', line)
            p   = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent  = Cm(0.8)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            set_font(p.add_run(clean(txt)), 10.5, color=C_BRAND)
            i += 1; continue

        # 空行
        if not line.strip():
            i += 1; continue

        # 普通段落
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(4)
        set_font(p.add_run(clean(line)), 10.5, color=C_BRAND)
        i += 1

    doc.save(out_path)
    print(f"✅ Word 已生成：{out_path}")
    return out_path


# ── 命令行入口 ────────────────────────────────────────────────

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("用法：python3 build_docx.py <md_path> <out_path> [title] [subtitle] [meta]")
        sys.exit(1)

    md   = sys.argv[1]
    out  = sys.argv[2]
    ttl  = sys.argv[3] if len(sys.argv) > 3 else ""
    sub  = sys.argv[4] if len(sys.argv) > 4 else ""
    meta = sys.argv[5] if len(sys.argv) > 5 else ""

    build_word(md, out, ttl, sub, meta)
