"""
Render a markdown-style gongwen draft into a GB/T 9704-2012 compliant .docx.

Usage:
    python render_docx.py --input draft.md --out final.docx --type tongzhi

Draft format (simple — one section per line, blank lines separate paragraphs):

    # 标题 / Title
    主送机关：XX、YY：
    （正文段落 1）
    （正文段落 2）
    一、一级标题
    （此处缩进 2 字符的正文）
    （一）二级标题
    1.三级标题
    ...
    特此通知。
    附件：1. xxx
          2. yyy
    -- 落款 --
    XX 机关
    二〇二六年五月十三日

Dependencies: python-docx
"""

import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ImportError:
    sys.exit("Missing python-docx. Run: pip install python-docx")


FONT_FANGSONG = "仿宋_GB2312"
FONT_HEITI = "黑体"
FONT_KAITI = "楷体_GB2312"
FONT_XIAOBIAOSONG = "方正小标宋简体"  # fallback to 宋体 if not installed
PT_SAN = Pt(16)        # 三号 = 16pt
PT_ER = Pt(22)         # 二号 = 22pt (for title)
LINE_SPACING_PT = Pt(28.8)  # 固定值 28-29 磅


def _set_font(run, name=FONT_FANGSONG, size=PT_SAN, bold=False):
    run.font.name = name
    run.font.size = size
    run.font.bold = bold
    # East Asian font binding
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)


def _set_para(p, first_line_indent=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=LINE_SPACING_PT):
    pf = p.paragraph_format
    pf.alignment = alignment
    pf.line_spacing = line_spacing  # absolute pt = 固定值
    pf.line_spacing_rule = 4  # WD_LINE_SPACING.EXACTLY
    if first_line_indent:
        pf.first_line_indent = Pt(32)  # ~2 chars of 三号字
    else:
        pf.first_line_indent = Pt(0)


def render(draft_path: str, out_path: str, doc_type: str = "general"):
    text = Path(draft_path).read_text(encoding="utf-8")
    doc = Document()

    # Page setup per GB/T 9704
    section = doc.sections[0]
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue

        if line.startswith("# "):
            # Title — 二号 小标宋体 居中
            p = doc.add_paragraph()
            _set_para(p, first_line_indent=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            run = p.add_run(line[2:].strip())
            _set_font(run, name=FONT_XIAOBIAOSONG, size=PT_ER, bold=True)
        elif line.startswith("主送") or (line.endswith("：") and not line.startswith(("一", "二", "三", "四", "五", "六", "七", "八", "九", "十", "（"))):
            # 主送机关 — 仿宋三号 顶格
            p = doc.add_paragraph()
            _set_para(p, first_line_indent=False)
            content = line.replace("主送机关：", "").replace("主送：", "")
            run = p.add_run(content)
            _set_font(run)
        elif re.match(r"^[一二三四五六七八九十]、", line):
            # 一级标题 — 黑体三号
            p = doc.add_paragraph()
            _set_para(p)
            run = p.add_run(line)
            _set_font(run, name=FONT_HEITI, size=PT_SAN, bold=False)
        elif re.match(r"^（[一二三四五六七八九十]+）", line):
            # 二级标题 — 楷体三号 加粗
            p = doc.add_paragraph()
            _set_para(p)
            run = p.add_run(line)
            _set_font(run, name=FONT_KAITI, size=PT_SAN, bold=True)
        elif re.match(r"^\d+\.", line):
            # 三级标题 — 仿宋三号 加粗
            p = doc.add_paragraph()
            _set_para(p)
            run = p.add_run(line)
            _set_font(run, name=FONT_FANGSONG, size=PT_SAN, bold=True)
        elif line.startswith("-- 落款 --") or line.startswith("--落款--"):
            # Skip marker
            i += 1
            # Next two lines are 机关 + 日期, right-aligned
            for _ in range(2):
                if i < len(lines):
                    txt = lines[i].rstrip()
                    if txt:
                        p = doc.add_paragraph()
                        _set_para(p, first_line_indent=False, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
                        run = p.add_run(txt + "    ")  # space for stamp
                        _set_font(run)
                    i += 1
            continue
        elif line.startswith("附件"):
            # 附件 — 仿宋三号
            p = doc.add_paragraph()
            _set_para(p, first_line_indent=False)
            run = p.add_run(line)
            _set_font(run)
        else:
            # 正文 — 仿宋三号 首行缩进
            p = doc.add_paragraph()
            _set_para(p)
            run = p.add_run(line)
            _set_font(run)

        i += 1

    doc.save(out_path)
    print(f"Wrote {out_path}")


def validate(docx_path: str):
    """Quick format compliance checks."""
    doc = Document(docx_path)
    issues = []

    # Margins
    s = doc.sections[0]
    if abs(s.top_margin.cm - 3.7) > 0.15:
        issues.append(f"Top margin {s.top_margin.cm:.2f}cm, expected 3.7cm")
    if abs(s.bottom_margin.cm - 3.5) > 0.15:
        issues.append(f"Bottom margin {s.bottom_margin.cm:.2f}cm, expected 3.5cm")
    if abs(s.left_margin.cm - 2.8) > 0.15:
        issues.append(f"Left margin {s.left_margin.cm:.2f}cm, expected 2.8cm")
    if abs(s.right_margin.cm - 2.6) > 0.15:
        issues.append(f"Right margin {s.right_margin.cm:.2f}cm, expected 2.6cm")

    # Body font / size (sample first few paragraphs)
    body_runs = []
    for p in doc.paragraphs[:20]:
        for r in p.runs:
            if r.text.strip():
                body_runs.append((r.font.name, r.font.size))
    seen_fangsong = any("仿宋" in (n or "") for n, _ in body_runs)
    if not seen_fangsong:
        issues.append("No 仿宋 font found in first 20 paragraphs — body should be 仿宋_GB2312")
    seen_san = any(s and abs(s.pt - 16) < 0.5 for _, s in body_runs)
    if not seen_san:
        issues.append("No 三号 (16pt) text found — body should be 三号")

    if not issues:
        print("OK — basic format compliance checks pass.")
    else:
        print("Issues:")
        for it in issues:
            print(f"  - {it}")


def main():
    p = argparse.ArgumentParser()
    p.add_argument("--input", help="Draft markdown file")
    p.add_argument("--out", default="final.docx")
    p.add_argument("--type", default="general", help="tongzhi/qingshi/baogao/han/jiyao/pifu/general")
    p.add_argument("--validate", help="Validate an existing .docx for GB/T 9704 compliance")
    args = p.parse_args()

    if args.validate:
        validate(args.validate)
    elif args.input:
        render(args.input, args.out, args.type)
    else:
        p.error("Need --input FILE or --validate FILE")


if __name__ == "__main__":
    main()
