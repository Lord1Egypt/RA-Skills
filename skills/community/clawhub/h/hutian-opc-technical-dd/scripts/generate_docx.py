#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成菲瑞药业技术尽调报告的docx版本
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

def add_heading(doc, text, level):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text):
    """添加段落"""
    p = doc.add_paragraph(text)
    return p

def add_table_from_list(doc, headers, rows):
    """从列表添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    
    # 添加表头
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header
        hdr_cells[idx].paragraphs[0].runs[0].bold = True
    
    # 添加数据行
    for row_idx, row in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row):
            row_cells[col_idx].text = str(cell_text)
    
    return table

def main():
    doc = Document()
    
    # 设置文档标题
    title = doc.add_heading('湖北菲瑞生物药业有限公司技术尽调报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加文档信息
    doc.add_paragraph()
    info_para = doc.add_paragraph()
    info_para.add_run('报告日期：').bold = True
    info_para.add_run('2026年5月7日\n')
    info_para.add_run('尽调对象：').bold = True
    info_para.add_run('湖北菲瑞生物药业有限公司\n')
    info_para.add_run('尽调级别：').bold = True
    info_para.add_run('全面技术尽调（四层分析法）\n')
    info_para.add_run('报告密级：').bold = True
    info_para.add_run('内部资料')
    
    doc.add_paragraph()
    doc.add_paragraph('_' * 60)
    
    # 摘要与核心结论
    add_heading(doc, '摘要与核心结论', 1)
    
    add_heading(doc, '一句话结论',  2)
    doc.add_paragraph(
        '湖北菲瑞生物药业有限公司是一家从事冻干闪释技术应用的创业期企业，技术有一定基础但存在严重夸大宣传的问题。"世界唯二、中国唯一"等核心说法与事实严重不符，市场规模数据存在数千倍夸大。'
    )
    
    add_heading(doc, '投资建议等级', 2)
    doc.add_paragraph()
    
    # 评分表格
    headers = ['评估维度', '评分（1-10）', '核心发现']
    rows = [
        ['技术壁垒', '5.0', '技术有一定积累，但存在明显竞争对手，5-8年领先优势存疑'],
        ['团队实力', '4.0', '20余人研发团队支撑130种药品储备存疑，员工人数与声称不符'],
        ['市场前景', '6.0', '冻干闪释技术有真实市场，但"万亿市场"严重夸大'],
        ['财务健康', '4.0', '营收5000万/估值7亿，P/S达14倍，估值偏高'],
        ['商业化能力', '5.0', '现有业务以OEM代工为主，药品上市周期长'],
        ['风险可控性', '3.0', '夸大宣传、竞争加剧、药品申报周期等多重风险'],
    ]
    add_table_from_list(doc, headers, rows)
    
    doc.add_paragraph()
    
    # 第一层
    add_heading(doc, '第一层：信息提取与核实', 1)
    
    add_heading(doc, '1.1 公司基本信息核实', 2)
    headers = ['声明内容', '融资计划书', '核实结果', '说明']
    rows = [
        ['公司成立时间', '2019年11月', '⚠️ 存疑', '企查查显示湖北菲瑞生物药业有限公司成立于2022年12月28日'],
        ['注册资本', '2596万元', '✅ 基本属实', '企查查显示2596.9232万元，实缴一致'],
        ['员工人数', '近100人', '❌ 严重不符', '企查查显示2024年参保人数仅4人'],
        ['办公地址', '鄂州市', '✅ 属实', '鄂州市鄂城区樊口街道旭光大道18号'],
        ['瞪羚企业', '2025年7月获评', '✅ 属实', '湖北省2025年入库瞪羚企业'],
    ]
    add_table_from_list(doc, headers, rows)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('关键矛盾：').bold = True
    p.add_run('融资计划书声称成立于2019年11月，但企查查显示主体公司成立于2022年12月28日。湖北日报报道提及"菲瑞生物集团成立于2019年11月"，可能存在集团与子公司混淆的问题。')
    
    p2 = doc.add_paragraph()
    p2.add_run('员工人数严重矛盾：').bold = True
    p2.add_run('融资计划书声称近100人，但企查查显示2024年参保人数仅4人。这与声称的"生产工人50余人、研发人员20余人"严重不符。')
    
    add_heading(doc, '1.2 专利情况核实', 2)
    headers = ['声明内容', '融资计划书', '核实结果', '说明']
    rows = [
        ['专利总数', '18项', '⚠️ 13项', '企查查显示专利13项（部分实质审查中）'],
        ['发明专利', '8项', '⚠️ 存疑', '部分专利如布洛芬状态为"实质审查"'],
        ['实用新型', '7项', '✅ 吻合', '-'],
        ['软件著作权', '3项', '✅ 吻合', '-'],
    ]
    add_table_from_list(doc, headers, rows)
    
    # 第二层
    add_heading(doc, '第二层：交叉验证与质疑', 1)
    
    add_heading(doc, '2.1 "世界唯二、中国唯一"说法验证', 2)
    p = doc.add_paragraph()
    p.add_run('❌ 该说法与事实严重不符').bold = True
    
    doc.add_paragraph('存在以下竞争对手：')
    headers = ['企业', '技术/产品', '规模/实力']
    rows = [
        ['陕西量子高科', 'Qsorb®冻干闪释技术', '中国冻干闪释技术开创者和引领者，拥有24项核心专利，成立于2002年'],
        ['齐鲁制药', '奥氮平、利培酮等冻干口崩片', '5个已上市品规，2025年销售额1.38亿元'],
        ['扬子江药业', '利培酮等冻干口崩片', '4个已上市品规'],
        ['石药集团', '阿立哌唑、左乙拉西坦等', '3个已上市品规，已有FDA申报案例'],
        ['华润双鹤', '齐拉西酮、奥氮平等', '3个品规'],
        ['常州伟博海泰', '冻干闪释片代工', '6万㎡基地，年产能10亿片'],
    ]
    add_table_from_list(doc, headers, rows)
    
    add_heading(doc, '2.2 市场规模数据验证', 2)
    headers = ['融资计划书声称', '实际情况', '差距']
    rows = [
        ['市场规模达"万亿元以上"', '2025年中国口腔崩解片市场规模约12.3亿元', '约800倍'],
        ['冻干口崩制剂', '2025年约4.82-7.59亿元', '差距数千倍'],
    ]
    add_table_from_list(doc, headers, rows)
    
    # 第三层
    add_heading(doc, '第三层：深度风险分析', 1)
    
    add_heading(doc, '3.1 技术风险', 2)
    headers = ['风险类型', '风险等级', '具体描述']
    rows = [
        ['技术壁垒夸大', '高', '"5-8年领先优势"基于"世界唯二、中国唯一"，但实际存在陕西量子高科等强劲竞争对手'],
        ['竞争对手突破', '高', '齐鲁、石药等上市公司具备更强研发和资金实力，可能快速追赶'],
        ['专利数量不足', '中', '实际13项专利 vs 声称18项，部分专利仍在实质审查'],
    ]
    add_table_from_list(doc, headers, rows)
    
    add_heading(doc, '3.2 商业风险', 2)
    headers = ['风险类型', '风险等级', '具体描述']
    rows = [
        ['估值偏高', '高', 'PS达14倍（5000万营收/7亿估值），远高于行业可比公司'],
        ['OEM模式天花板', '高', '代工模式毛利率低，难以建立品牌护城河'],
        ['产能瓶颈', '中', '目前仅1条营养品产线，交货期3个月+'],
    ]
    add_table_from_list(doc, headers, rows)
    
    add_heading(doc, '3.3 财务风险', 2)
    headers = ['风险类型', '风险等级', '具体描述']
    rows = [
        ['营收规模小', '高', '年营收5000万，抵御风险能力有限'],
        ['资金缺口大', '高', '计划2029年前投入7.65亿元，资金来源不明'],
        ['盈利不确定性', '高', '2025年预计净利润750万，净利率仅15%'],
    ]
    add_table_from_list(doc, headers, rows)
    
    add_heading(doc, '3.4 药品申报风险', 2)
    headers = ['风险类型', '风险等级', '具体描述']
    rows = [
        ['申报周期长', '高', '2类药5-6年，3类药2-3年，时间窗口存在不确定性'],
        ['竞争激烈', '高', '齐鲁、石药等已布局类似产品'],
        ['研发投入大', '高', '每个3类药投入1000万，2类药4500万'],
    ]
    add_table_from_list(doc, headers, rows)
    
    # 第四层
    add_heading(doc, '第四层：投资建议与评分', 1)
    
    add_heading(doc, '4.1 综合评分', 2)
    headers = ['评估维度', '评分（1-10）', '说明']
    rows = [
        ['技术壁垒', '5.0', '技术有一定积累，但非独有，存在陕西量子高科等强劲竞争对手'],
        ['团队实力', '4.0', '员工人数与声称严重不符，研发人员数量支撑130种药品储备存疑'],
        ['市场前景', '6.0', '冻干闪释技术有真实市场空间，但万亿说法严重夸大'],
        ['财务健康', '4.0', '营收规模小，估值偏高，投入产出比待验证'],
        ['商业化能力', '5.0', 'OEM模式可行，药品上市周期长，IPO目标激进'],
        ['风险可控性', '3.0', '多重风险叠加，夸大宣传问题突出'],
    ]
    add_table_from_list(doc, headers, rows)
    
    add_heading(doc, '4.2 投资建议', 2)
    p = doc.add_paragraph()
    p.add_run('谨慎推荐（下调观察）').bold = True
    
    doc.add_paragraph('推荐理由：')
    reasons = [
        '冻干闪释技术确有一定市场应用价值，符合老龄化社会的用药需求趋势',
        '瞪羚企业认证、ISO认证等资质具备一定可信度',
        '湖北地方政府产业基金背书',
        '已有营养品OEM业务实现营收',
    ]
    for r in reasons:
        doc.add_paragraph(f'• {r}')
    
    doc.add_paragraph()
    doc.add_paragraph('谨慎理由：')
    concerns = [
        '❌ "世界唯二、中国唯一"说法严重失实：陕西量子高科明确声称自己是该领域的开创者和引领者',
        '❌ 市场规模夸大数千倍：万亿 vs 实际12亿，相差过于悬殊',
        '❌ 员工人数严重不符：声称100人 vs 实际参保4人，差异无法合理解释',
        '❌ 专利数量存在夸大：18项 vs 实际13项',
        '❌ 竞争对手众多：齐鲁、石药、扬子江等多家上市公司均有冻干口崩片布局',
        '❌ 估值偏高：PS达14倍，远高于行业水平',
        '❌ IPO目标激进：2029年上市，但药品申报周期长，不确定性大',
    ]
    for c in concerns:
        doc.add_paragraph(f'• {c}')
    
    add_heading(doc, '4.3 建议关注指标', 2)
    headers = ['关注事项', '优先级', '说明']
    rows = [
        ['实际员工人数和团队构成', '极高', '要求提供社保缴纳明细、员工名册'],
        ['药品备案号真实性', '极高', '要求提供备案证原件或药监局官网截图'],
        ['技术评估报告原件', '高', '要求提供鄂技交评字【2023】第S130号原件'],
        ['财务数据核实', '高', '要求提供经审计的财务报表'],
        ['陕西量子高科对比分析', '高', '详细了解竞争对手的技术、产品、商业模式'],
        ['药品申报进度', '高', '跟踪CTR20243287、CTR20244976审评进展'],
    ]
    add_table_from_list(doc, headers, rows)
    
    # 附录
    add_heading(doc, '附录', 1)
    
    add_heading(doc, '附录A：主要竞争对手概况', 2)
    
    p = doc.add_paragraph()
    p.add_run('陕西量子高科药业有限公司').bold = True
    details = [
        '成立时间：2002年',
        '注册资本：1.52亿元',
        '员工人数：108人（专业技术人员75人）',
        '核心技术：Qsorb®冻干闪释技术',
        '专利：29项发明专利（含PCT）',
        '产能：单线3亿片/年',
        '认证：2008年获得中国首张药品（冻干）片剂GMP证书',
    ]
    for d in details:
        doc.add_paragraph(f'• {d}')
    
    add_heading(doc, '附录B：冻干口崩片市场规模数据', 2)
    headers = ['年份', '口腔崩解片市场规模（亿元）', '冻干口崩制剂规模（亿元）']
    rows = [
        ['2025', '12.30', '约4.82-7.59'],
        ['2026E', '13.05', '约8.05'],
        ['2027E', '13.85', '约9.48'],
        ['2028E', '14.70', '约10.07'],
        ['2029E', '15.60', '约10.71'],
        ['2030E', '16.55', '约11.39'],
    ]
    add_table_from_list(doc, headers, rows)
    
    p = doc.add_paragraph()
    p.add_run('数据来源：').bold = True
    p.add_run('博研咨询、行业研报')
    
    add_heading(doc, '附录C：需要进一步核实的事项清单', 2)
    headers = ['序号', '核实事项', '优先级', '说明']
    rows = [
        ['1', '实际员工人数和团队构成', '极高', '声称100人 vs 参保4人，差异需解释'],
        ['2', 'CTR20243287、CTR20244976备案号', '极高', '要求提供原件或官方截图'],
        ['3', '鄂技交评字【2023】第S130号报告', '高', '要求提供原件'],
        ['4', '药品生产许可证', '高', '核实是否具备人用药生产资质'],
        ['5', '经审计的财务报表', '高', '核实营收5000万的真实性'],
        ['6', '1.2亿元实际投入明细', '高', '核实资金使用情况'],
        ['7', '与陕西量子高科的技术差异', '高', '需专业评估'],
        ['8', '128种产品冷冻曲线摸索', '中', '核实具体进展'],
    ]
    add_table_from_list(doc, headers, rows)
    
    # 页脚
    doc.add_paragraph()
    doc.add_paragraph('_' * 60)
    footer = doc.add_paragraph()
    footer.add_run('报告完成时间：').bold = True
    footer.add_run('2026年5月7日\n')
    footer.add_run('报告编制人：').bold = True
    footer.add_run('技术尽调团队\n')
    footer.add_run('免责声明：').bold = True
    footer.add_run('本报告基于公开信息和有限核查编制，仅供参考，不构成投资建议。建议投资前进行更深入的尽职调查。')
    
    # 保存文档
    import os
    output_dir = os.path.dirname(os.path.abspath(__file__))
    output_path = os.path.join(output_dir, '菲瑞药业技术尽调报告.docx')
    doc.save(output_path)
    print(f"docx文档已生成：{output_path}")

if __name__ == '__main__':
    main()
