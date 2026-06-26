#!/usr/bin/env python3
"""
build_docx.py - 将Markdown方案文档和配图合并为正式DOCX

用法：
  python3 build_docx.py <markdown_file> <images_dir> <output.docx> [--title <title>]

示例：
  python3 build_docx.py solution.md ./imgs/ solution.docx --title "XX系统解决方案建议书"

依赖：pip install python-docx --break-system-packages
"""

import sys, os, re, argparse

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx --break-system-packages")
    sys.exit(1)


def set_run_font(run, name='仿宋', size=Pt(14), bold=False):
    run.font.name = name
    run.font.size = size
    run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), name)


def add_heading_styled(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, '黑体', Pt({1:22, 2:16, 3:14}.get(level, 14)), True)
    return h


def add_para(doc, text, bold=False, size=14, indent=True, font_name='仿宋', align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, font_name, Pt(size), bold)
    p.paragraph_format.line_spacing = Pt(28)
    if indent and not bold:
        p.paragraph_format.first_line_indent = Pt(28)
    if align:
        p.alignment = align
    return p


def add_image(doc, img_path, width_cm=15):
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Cm(width_cm))
        return p
    return None


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, '楷体', Pt(10.5))
    p.paragraph_format.space_after = Pt(12)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h.strip())
        set_run_font(run, '黑体', Pt(10.5), True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = ''
            run = cell.paragraphs[0].add_run(val.strip())
            set_run_font(run, '仿宋', Pt(10.5))
    return table


# Image embedding map: keywords in heading → image filename
IMAGE_MAP = [
    ('架构',      '01_architecture.png', '图X-1 技术架构图'),
    ('流程',      '02_process_flow.png', '图X-2 核心业务流程'),
    ('路线',      '03_roadmap.png',      '图X-3 实施路线图'),
    ('回报',      '04_roi.png',          '图X-4 投资回报分析'),
    ('价值',      '05_value_map.png',    '图X-5 痛点-功能-价值映射图'),
    ('映射',      '05_value_map.png',    '图X-5 痛点-功能-价值映射图'),
    ('安全',      '06_security.png',     '图X-6 安全与合规体系架构'),
]


def parse_markdown(doc, lines, img_dir):
    """解析Markdown行，生成docx元素"""
    i = 0
    chapter_num = 0
    fig_num = 0

    while i < len(lines):
        line = lines[i].rstrip()

        if not line:
            i += 1
            continue

        # Headings
        if line.startswith('#'):
            m = re.match(r'^(#{1,4})\s+(.+)$', line)
            if m:
                level = len(m.group(1))
                text = m.group(2).strip()

                if level == 1:
                    chapter_num += 1
                    fig_num = 0

                if level <= 3:
                    add_heading_styled(doc, text, level)
                else:
                    add_para(doc, text, bold=True, size=14, indent=False)

                # Auto-insert images based on heading keywords
                for kw, img_file, caption_tpl in IMAGE_MAP:
                    if kw in text and level <= 2:
                        img_path = os.path.join(img_dir, img_file)
                        if os.path.exists(img_path):
                            fig_num += 1
                            caption = caption_tpl.replace('X', str(chapter_num)).replace(
                                str(chapter_num) + '-' + str(fig_num),
                                f'{chapter_num}-{fig_num}')
                            caption = f'图{chapter_num}-{fig_num} {caption_tpl.split(" ", 1)[-1]}' if ' ' in caption_tpl else caption
                            add_image(doc, img_path)
                            add_caption(doc, caption)
                        break

            i += 1
            continue

        # Tables
        if '|' in line and line.strip().startswith('|'):
            table_lines = []
            while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            if len(table_lines) >= 2:
                headers = [c.strip() for c in table_lines[0].split('|') if c.strip()]
                data_rows = []
                for tl in table_lines[2:]:
                    cells = [c.strip() for c in tl.split('|') if c.strip()]
                    if cells:
                        while len(cells) < len(headers):
                            cells.append('')
                        data_rows.append(cells[:len(headers)])
                if headers and data_rows:
                    add_table(doc, headers, data_rows)
            continue

        # Code blocks
        if line.startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            p.paragraph_format.left_indent = Cm(1)
            continue

        # Horizontal rule
        if line.startswith('---'):
            i += 1
            continue

        # Bold lines
        if line.startswith('**') and '**' in line[2:]:
            clean = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
            clean = re.sub(r'`(.+?)`', r'\1', clean)
            add_para(doc, clean, bold=True, font_name='黑体', indent=False)
            i += 1
            continue

        # Regular paragraph
        clean = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
        clean = re.sub(r'`(.+?)`', r'\1', clean)
        clean = re.sub(r'^(•|-)\s+', '● ', clean)
        if clean.strip():
            add_para(doc, clean)
        i += 1


def build_docx(md_path, img_dir, output_path, title=None):
    """主构建函数"""
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Normal style
    style = doc.styles['Normal']
    style.font.name = '仿宋'
    style.font.size = Pt(14)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    # Title page
    if title:
        doc.add_paragraph()
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for line in title.split('\n'):
            run = p.add_run(line)
            set_run_font(run, '黑体', Pt(36), True)
            p.add_run('\n')
        doc.add_paragraph()
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('V1.0  |  2026年4月')
        set_run_font(run, '仿宋', Pt(14))
        doc.add_page_break()

    # Parse markdown
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    parse_markdown(doc, lines, img_dir)

    # Back cover
    doc.add_page_break()
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('— END —')
    set_run_font(run, '黑体', Pt(24), True)

    doc.save(output_path)
    size_kb = os.path.getsize(output_path) // 1024
    print(f'✓ DOCX已生成: {output_path} ({size_kb}KB)')


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Markdown + 配图 → DOCX')
    parser.add_argument('markdown', help='Markdown方案文件')
    parser.add_argument('images', help='配图目录')
    parser.add_argument('output', help='输出DOCX路径')
    parser.add_argument('--title', default=None, help='封面标题（支持换行符\\n）')
    args = parser.parse_args()

    if not os.path.exists(args.markdown):
        print(f'ERROR: Markdown文件不存在: {args.markdown}')
        sys.exit(1)

    title = args.title.replace('\\n', '\n') if args.title else None
    build_docx(args.markdown, args.images, args.output, title)
