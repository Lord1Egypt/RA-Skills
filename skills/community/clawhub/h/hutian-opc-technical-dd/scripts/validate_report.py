#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
验证脚本 - 检查文档质量
"""

from docx import Document

def validate_docx(filepath):
    doc = Document(filepath)
    errors = []
    
    # 检查1：表格空行
    for i, table in enumerate(doc.tables):
        for j, row in enumerate(table.rows):
            if all(cell.text.strip() == '' for cell in row.cells):
                errors.append(f"表格{i+1}行{j}：整行为空")
    
    # 检查2：空单元格占比（>10%异常）
    for i, table in enumerate(doc.tables):
        total = sum(len(row.cells) for row in table.rows)
        empty = sum(1 for row in table.rows for cell in row.cells if not cell.text.strip())
        if total > 0 and empty / total > 0.1:
            errors.append(f"表格{i+1}：空格占比{empty/total*100:.0f}%")
    
    # 检查3：空表（仅表头）
    for i, table in enumerate(doc.tables):
        if len(table.rows) <= 1:
            errors.append(f"表格{i+1}：仅有表头，无数据行")
    
    # 检查4：空章节（标题下无内容）
    headings_with_content = set()
    all_headings = set()
    current_heading = None
    for el in doc.element.body:
        tag = el.tag.split('}')[-1]
        if tag == 'p':
            for para in doc.paragraphs:
                if para._element is el:
                    if para.style.name.startswith('Heading'):
                        current_heading = para.text.strip()
                        if current_heading:
                            all_headings.add(current_heading)
                    elif para.text.strip() and current_heading:
                        headings_with_content.add(current_heading)
                    break
        elif tag == 'tbl' and current_heading:
            headings_with_content.add(current_heading)
    empty_sections = all_headings - headings_with_content
    for h in empty_sections:
        errors.append(f"空章节：{h}")
    
    # 统计信息
    table_count = len(doc.tables)
    empty_row_count = sum(1 for table in doc.tables for row in table.rows if all(cell.text.strip() == '' for cell in row.cells))
    empty_section_count = len(empty_sections)
    
    if errors:
        print("❌ 自检未通过：")
        for e in errors[:20]:  # 只显示前20个错误
            print(f"  - {e}")
        if len(errors) > 20:
            print(f"  ... 还有 {len(errors)-20} 个错误")
        return False, {
            'table_count': table_count,
            'empty_row_count': empty_row_count,
            'empty_section_count': empty_section_count,
            'errors': errors
        }
    else:
        print(f"✅ 自检通过：{table_count}表，{empty_row_count}空行，{empty_section_count}空章节")
        return True, {
            'table_count': table_count,
            'empty_row_count': empty_row_count,
            'empty_section_count': empty_section_count,
            'errors': []
        }

if __name__ == "__main__":
    filepath = "河北青山鼎信技术尽调报告.docx"
    print(f"正在验证：{filepath}")
    print("=" * 50)
    result, stats = validate_docx(filepath)
    print("=" * 50)
    print(f"表格数：{stats['table_count']}")
    print(f"空行数：{stats['empty_row_count']}")
    print(f"空章节数：{stats['empty_section_count']}")
