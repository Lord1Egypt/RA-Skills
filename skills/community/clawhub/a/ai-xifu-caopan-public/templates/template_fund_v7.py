#!/usr/bin/env python3

# ============================================================
# ⚠️ SAMPLE DATA NOTICE / 示例数据声明
# This template contains EXAMPLE / ILLUSTRATIVE data values.
# Users must configure their own data sources (yfinance, Guosen API, etc.)
# to generate plans with real market data.
# 本模板包含的数据为示例/演示用途。
# 用户需自行配置数据源（yfinance、国信API等）以获取真实市场数据。
# ============================================================

# -*- coding: utf-8 -*-
"""随机基金方案 - 创业板成长ETF华夏 (159967) / Random Fund Plan - ChiNext Growth ETF ChinaAMC (159967)"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import sys
sys.path.insert(0, os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'global_market'))
from geopolitical_risk import get_geopolitical_risks, format_geopolitical_section
from language_detection import detect_input_lang, match_market_lang, native_name, BilingualText

today = "2026-05-23"

# ── 🌍 大叔亲授·审题+质检机制（默认中文） ──
_input_lang = "zh"
_dest_lang = "zh"
_bt = BilingualText(input_lang=_input_lang, dest_lang=_dest_lang)

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'SimHei'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_title(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'SimHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
        if level <= 1:
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return h

def add_para(text, bold=False, size=11, align=None, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'SimHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    return p

def add_table(headers, rows, header_bg='2F5496', alt_bg='D6E4F0'):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'SimHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
        set_cell_shading(cell, header_bg)
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri+1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = 'SimHei'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
            if ri % 2 == 1:
                set_cell_shading(cell, alt_bg)
    doc.add_paragraph()
    return table

def safe(text):
    rep = {
        '\U0001f4ca': '[方案]', '\U0001f4c8': '[上涨]', '\U0001f4c9': '[下跌]',
        '\U0001f4cc': '[注意]', '\u26a0\ufe0f': '[注意]', '\u26a0': '[注意]',
        '\u2705': '[是]', '\u274c': '[否]',
        '\U0001f7e2': '[低]', '\U0001f7e1': '[中]', '\U0001f7e0': '[较高]', '\U0001f534': '[高]',
        '\u2b50': '[星]', '\U0001f60d': '',
        '\U0001f4b0': '[资金]', '\U0001f3b2': '[概率]',
        '\U0001f9e0': '[思维]', '\U0001f4a1': '[提示]',
        '\U0001f4dd': '[笔记]', '\U0001f3e6': '[基金]', '\U0001f4e6': '[期货]',
        '\U0001f680': '[上涨]', '\U0001f525': '[热门]',
        '\U0001f44d': '[好]', '\U0001f440': '[关注]',
        '\U0001f64c': '', '\U0001f64f': '', '\ufe0f': '',
    }
    for k, v in rep.items():
        text = text.replace(k, v)
    return text

# ===== AIGC头部 / AIGC Header =====
aigc_top = doc.add_paragraph()
aigc_top.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = aigc_top.add_run('[AIGC生成内容，仅供参考，不构成投资建议] / [AIGC-generated content for reference only, not investment advice]')
run.font.size = Pt(10)
run.bold = True
run.font.color.rgb = RGBColor(0xcc, 0x33, 0x33)
run.font.name = 'SimHei'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

# ===== TITLE =====
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run(safe('创业板成长ETF华夏(159967) 投资方案 / Investment Plan'))
run.font.size = Pt(20)
run.bold = True
run.font.name = 'SimHei'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

ver_p = doc.add_paragraph()
ver_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ver_p.add_run(safe(f'版本 / Version：v7.0 FINAL | 制作日期 / Date：{today} | 品种 / Product：创业板成长ETF(159967.SZ)'))
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.name = 'SimHei'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
doc.add_paragraph()

# ===== 方案头部 / Plan Header =====
add_title(safe('方案头部 / Plan Header'), level=2)
add_table([f'项目 / {_bt.item_label}', '内容 / Details'], [
    ['基金名称 / Fund Name', '创业板成长ETF华夏 / ChiNext Growth ETF ChinaAMC'],
    ['基金代码 / Fund Code', '159967.SZ'],
    ['交易所 / Exchange', '深圳证券交易所 / Shenzhen Stock Exchange'],
    ['跟踪指数 / Tracking Index', '创成长指数(399296) / ChiNext Growth Index (399296)'],
    ['基金规模 / Fund Size', '约41.4亿元 / Approx. 4.14 Billion CNY'],
    ['制作日期 / Date Prepared', '2026-05-23(周六 / Saturday)'],
    ['最新估值(IOPV) / Latest IOPV', '0.8846元 / 0.8846 CNY'],
    ['最新价格 / Latest Price', '0.885元 / 0.885 CNY'],
    ['今日涨跌幅 / Daily Change', '+5.61%(涨 / Up 0.047元)'],
    ['交易制度 / Trading Rule', 'T+1(ETF场内交易，当日买入最早次日卖出) / T+1 (ETF on-exchange, buy today sell tomorrow earliest)'],
    ['基金类型 / Fund Type', '股票型ETF/成长风格/创业板 / Equity ETF / Growth Style / ChiNext'],
])

# ===== 一、行情回顾 / Market Review =====
add_title(safe('一、行情回顾 / 1. Market Review'), level=1)
add_title('1.1 收盘数据 / Closing Data', level=2)
add_table(['名称 / Name', '收盘价 / Close', '日涨跌 / Daily Change', '涨跌幅 / Change%', '成交量 / Volume', '成交额 / Turnover', '振幅 / Amplitude'], [
    ['创业板成长ETF / ChiNext Growth ETF', '0.885', '+0.047', '+5.61%', '1182.56万手', '10.29亿', '4.53%'],
])

add_title('1.2 近5日走势 / Recent 5-Day Trend', level=2)
add_table(['日期 / Date', '开盘 / Open', '最高 / High', '最低 / Low', '收盘 / Close', '涨跌幅 / Change%', '成交额 / Turnover'], [
    ['05-18(一/Mon)', '0.857', '0.883', '0.857', '0.870', '0.00%', '6.66亿'],
    ['05-19(二/Tue)', '0.861', '0.866', '0.829', '0.860', '-1.15%', '6.98亿'],
    ['05-20(三/Wed)', '0.850', '0.877', '0.850', '0.867', '+0.81%', '14.48亿'],
    ['05-21(四/Thu)', '0.878', '0.890', '0.838', '0.838', '-3.34%', '10.24亿'],
    ['05-22(五/Fri)', '0.850', '0.888', '0.850', '0.885', '+5.61%', '10.29亿'],
])
add_para(safe('[重要/Important] 近5日累计上涨1.72%，波动较大。05-21大跌3.34%后05-22大涨5.61%反包，做多动能强劲。/ Cumulative +1.72% over 5 days with high volatility. Plunged 3.34% on 05-21 then surged 5.61% on 05-22 engulfing the loss — strong bullish momentum.'), bold=True)

# ===== 二、跟踪指数分析 / Tracking Index Analysis =====
add_title(safe('二、跟踪指数分析 / 2. Tracking Index Analysis'), level=1)
add_title('2.1 创成长指数行情 / ChiNext Growth Index Market', level=2)
add_table(['指数 / Index', '最新 / Latest', '涨跌幅 / Change%', 'ETF偏离 / ETF Deviation'], [
    ['创成长(399296) / ChiNext Growth (399296)', '8655.81', '+5.08%', 'ETF涨幅5.61%，略偏强 / ETF +5.61%, slightly stronger'],
])
add_para(safe('跟踪分析：05-22创成长指数上涨5.08%，ETF涨幅5.61%略高于指数，溢价率约0.04%，基本合理。/ Analysis: On 05-22 the ChiNext Growth Index rose 5.08%, ETF rose 5.61% slightly higher with ~0.04% premium — generally reasonable.'))

add_title('2.2 指数成分风格 / Index Component Style', level=2)
add_table(['维度 / Dimension', '说明 / Description'], [
    ['选股逻辑 / Selection Logic', '从创业板中选取成长性突出、动量效应强的50只股票 / Selects 50 stocks from ChiNext with outstanding growth and strong momentum effect'],
    ['行业分布 / Sector Distribution', '新能源、医药生物、电子、计算机等成长行业为主 / New Energy, Pharma, Electronics, IT and other growth sectors'],
    ['风格特征 / Style Characteristics', '高成长、高波动、高弹性 / High growth, high volatility, high beta'],
    ['波动率 / Volatility', '年化波动约30%，属于高波动指数 / Annualized volatility ~30%, high-volatility index'],
])

# ===== 三、市场环境分析 / Market Environment Analysis =====
add_title(safe('三、市场环境分析 / 3. Market Environment Analysis'), level=1)
add_title('3.1 大盘环境 / Market Environment', level=2)
add_table(['指数 / Index', '最新 / Latest', '涨跌幅 / Change%', '对基金影响 / Impact on Fund'], [
    ['上证指数 / Shanghai Composite', '4112.90', '+0.87%', '[偏多/Bullish] 大盘回暖提供支撑 / Market rebound provides support'],
    ['深证成指 / Shenzhen Component', '15597.30', '+2.30%', '[偏多/Bullish] 深市明显强势 / Shenzhen notably strong'],
    ['创业板指 / ChiNext Index', '3938.50', '+2.84%', '[偏多/Bullish] 创业板领涨，利好成长ETF / ChiNext leads, bullish for growth ETF'],
    ['创成长 / ChiNext Growth', '8655.81', '+5.08%', '[偏多/Bullish] 成长风格明显占优 / Growth style clearly dominant'],
])
add_para(safe('大盘分析：05-22三大指数全面上涨，创业板指涨2.84%领涨，成长风格明显占优。市场资金从价值股切换至成长股的迹象明显，利好创业板成长ETF。/ Market Analysis: All three major indices rose on 05-22, with ChiNext leading at +2.84%. Growth style clearly dominates. Capital rotation from value to growth stocks is evident — bullish for ChiNext Growth ETF.'))

add_title('3.2 风格与板块 / Style & Sector', level=2)
add_table(['维度 / Dimension', '要点 / Key Points', '影响 / Impact'], [
    ['市场风格 / Market Style', '成长风格占优，科技+创业板领涨 / Growth style dominates, Tech + ChiNext lead', '[偏多/Bullish]'],
    ['成交量 / Volume', '两市成交约2.9万亿，放量明显 / Combined turnover ~2.9 trillion, notably expanded', '[偏多/Bullish]'],
    ['北向资金 / Northbound Funds', '05-22北向资金大幅流入成长股 / Northbound funds heavily流入 growth stocks on 05-22', '[偏多/Bullish]'],
    ['板块轮动 / Sector Rotation', '科技/消费电子/信创板块活跃 / Tech/Consumer Electronics/IT Innovation sectors active', '[偏多/Bullish]'],
    ['政策面 / Policy', '科技自主创新政策利好频出 / Frequent tech self-innovation policy tailwinds', '[偏多/Bullish]'],
])

# ===== 🌍 地缘政治分析 / Geopolitical Analysis =====
add_title(safe('四、全球地缘政治博弈 / 4. Global Geopolitical Conflict'), level=1)
try:
    geo_analysis = get_geopolitical_risks("A股")
    lines = format_geopolitical_section(geo_analysis).split('\n')
    for line in lines:
        if line.strip():
            add_para(safe(line))
except:
    add_para(safe('地缘政治数据加载中...'))
doc.add_page_break()

# ===== 五、大盘风险预警 / Market Risk Warning =====
add_title(safe('五、大盘风险预警 / 5. Market Risk Warning'), level=1)
add_title('4.1 预警信号 / Warning Signals', level=2)
add_table(['序号 / No.', '信号 / Signal', '状态 / Status', '说明 / Description'], [
    ['1', '指数跌破关键支撑 / Index breaks key support', '[未触发/Not Triggered]', '创业板指3938仍在3800上方 / ChiNext 3938 still above 3800'],
    ['2', '成交量异常 / Abnormal volume', '[未触发/Not Triggered]', '放量属于正常回暖 / Volume expansion is normal recovery'],
    ['3', '外围市场大跌 / Overseas market crash', '[未触发/Not Triggered]', '美欧市场平稳 / US and European markets stable'],
    ['4', '北向持续流出 / Sustained Northbound outflow', '[未触发/Not Triggered]', '北向已转为流入成长股 / Northbound reversed to inflows into growth stocks'],
    ['5', '板块轮动过快 / Excessive sector rotation', '[未触发/Not Triggered]', '科技成长热点持续 / Tech growth theme sustained'],
    ['6', '个基普跌 / Widespread fund decline', '[触发/Triggered]', '05-21曾大跌，但已反包 / Dropped sharply on 05-21 but has since engulfed'],
    ['7', '政策利空 / Policy headwind', '[未触发/Not Triggered]', '科技政策利好为主 / Mostly favorable tech policies'],
    ['8', '经济数据不及预期 / Economic data misses expectations', '[触发/Triggered]', '社零偏弱但成长股受影响小 / Retail sales weak but growth stocks less affected'],
    ['9', '情绪过度乐观 / Excessively optimistic sentiment', '[未触发/Not Triggered]', '市场情绪适中 / Market sentiment moderate'],
    ['10', 'ETF溢价过高 / Excessive ETF premium', '[未触发/Not Triggered]', '溢价仅0.04%，正常 / Premium only 0.04%, normal'],
])
add_para(safe('信号触发统计：2/10个触发。综合风险等级：[低]风险。当前成长风格环境友好。/ Signals triggered: 2/10. Composite risk level: [Low]. Growth-style environment is favorable.'))

# ===== 六、技术面分析 / Technical Analysis =====
add_title(safe('六、技术面分析 / 5. Technical Analysis'), level=1)
add_title('5.1 关键价位 / Key Price Levels', level=2)
add_table(['强阻力 / Strong Resistance', '次阻力 / Secondary Resistance', '当前价 / Current', '次支撑 / Secondary Support', '强支撑 / Strong Support'], [
    ['0.92-0.95', '0.89-0.90', '0.885', '0.85-0.86', '0.82-0.83'],
])

add_title('5.2 量价指标 / Volume & Price Indicators', level=2)
add_table(['指标 / Indicator', '数值 / Value', '信号 / Signal'], [
    ['近5日均量 / 5-Day Avg Volume', '约11.15亿手', '放量突破，做多信号 / Breakout on volume, bullish signal'],
    ['05-22成交额 / 05-22 Turnover', '10.29亿', '放量+5.61%，突破形态 / Volume +5.61%, breakout pattern'],
    ['溢价率 / Premium Rate', '0.04%', '合理范围，无套利空间 / Reasonable, no arbitrage opportunity'],
    ['换手率 / Turnover Rate', '25.28%', '交投活跃，流动性充足 / Active trading, ample liquidity'],
])

add_title('5.3 趋势判断 / Trend Analysis', level=2)
add_table(['指标 / Indicator', '判断 / Judgment', '信号 / Signal'], [
    ['短期趋势(5日) / Short-term (5D)', '0.870至0.885', '震荡向上 / Consolidating upward'],
    ['中期趋势(20日) / Medium-term (20D)', '均线约0.85附近 / MA ~0.85', '价格站上均线 / Price above MA'],
    ['MACD', '05-22大涨后有望金叉 / Possible golden cross after 05-22 surge', '关注确认信号 / Watch for confirmation'],
    ['BOLL布林带 / Bollinger Bands', '价格突破中轨 / Price broke above mid-band', '偏强，关注上轨0.90 / Bullish, watch upper band 0.90'],
])

# ===== 七、走势预判 / Trend Forecast =====
add_title(safe('七、走势预判 / 6. Trend Forecast'), level=1)
add_title('6.1 大盘预判 / Market Forecast', level=2)
add_table(['指数 / Index', '05-25(一)预判 / Mon Forecast', '05-26(二)预判 / Tue Forecast', '置信度 / Confidence'], [
    ['创业板指 / ChiNext Index', '震荡偏强 3880-3980 / Consolidating bullish 3880-3980', '震荡偏强 3900-4000 / Consolidating bullish 3900-4000', '[中/Med] 3星/Stars'],
    ['创成长 / ChiNext Growth', '震荡偏强 8500-8800 / Consolidating bullish 8500-8800', '震荡偏强 8550-8900 / Consolidating bullish 8550-8900', '[中/Med] 3星/Stars'],
])

add_title('6.2 ETF预判 / ETF Forecast', level=2)
add_table(['日期 / Date', '走势 / Trend', '区间 / Range', '置信度 / Confidence', '主要依据 / Key Basis'], [
    ['05-25(一/Mon)', '震荡偏强 / Consolidating bullish', '0.86-0.90', '[中/Med] 3星/Stars', '05-22放量反包，短期做多动能延续，0.90附近有获利回吐压力 / Volume engulfing on 05-22, short-term bullish momentum continues; profit-taking pressure near 0.90'],
    ['05-26(二/Tue)', '震荡整理 / Consolidating', '0.86-0.91', '[中低/Med-Low] 2星/Stars', '连续反弹后需整固，突破0.90则打开上涨空间至0.92-0.95 / Needs consolidation after consecutive rebounds; breaking 0.90 opens upside to 0.92-0.95'],
])
add_para(safe('[重要/Important] 核心逻辑：ETF 05-22大涨5.61%反包前日跌幅，短期趋势偏强。成长风格占优的环境下，突破0.90则打开上涨空间。/ Core logic: ETF surged 5.61% on 05-22 engulfing the previous day\'s loss; short-term trend is bullish. In a growth-style-dominant environment, breaking 0.90 opens upside room.'), bold=True)

# ===== 八、基金基本面 / Fund Fundamentals =====
add_title(safe('八、基金基本面 / 7. Fund Fundamentals'), level=1)
add_title('7.1 基金概况 / Fund Overview', level=2)
add_table([f'项目 / {_bt.item_label}', '内容 / Details'], [
    ['基金管理人 / Fund Manager', '华夏基金管理有限公司 / China Asset Management Co., Ltd.'],
    ['成立日期 / Inception Date', '2019年6月 / June 2019'],
    ['基金规模 / Fund Size', '约41.4亿元 / Approx. 4.14 Billion CNY'],
    ['管理费率 / Management Fee', '0.50%/年 / 0.50% p.a.'],
    ['托管费率 / Custody Fee', '0.10%/年 / 0.10% p.a.'],
    ['跟踪指数 / Tracking Index', '创成长指数(399296) / ChiNext Growth Index (399296)'],
    ['今年以来涨幅 / YTD Return', '+39.15%'],
])

# ===== 九、同类ETF对比 / Comparable ETF Comparison =====
add_title(safe('九、同类ETF对比 / 8. Comparable ETF Comparison'), level=1)
add_table(['ETF名称 / ETF Name', '代码 / Code', '涨跌幅 / Change%', '跟踪指数 / Tracked Index', '规模 / Size'], [
    ['创业板成长ETF华夏 / ChiNext Growth ETF ChinaAMC', '159967', '+5.61%', '创成长 / ChiNext Growth', '41.4亿'],
    ['创业板ETF易方达 / ChiNext ETF E Fund', '159915', '+2.84%', '创业板指 / ChiNext Index', '约200亿'],
    ['创业板50ETF华安 / ChiNext 50 ETF HuaAn', '159949', '+3.20%', '创业板50 / ChiNext 50', '约100亿'],
])
add_para(safe('同类对比：159967涨幅+5.61%远高于159915(+2.84%)和159949(+3.20%)，创成长指数选股策略在当下市场表现优异。/ Comparison: 159967 +5.61% far outperforms 159915 (+2.84%) and 159949 (+3.20%); the ChiNext Growth Index selection strategy excels in the current market.'))

# ===== 十、操作建议 / Operation Suggestions =====
add_title(safe('十、操作建议 / 9. Operation Suggestions'), level=1)
add_title('9.1 操作参考 / Operation Reference', level=2)
add_table(['操作类型 / Operation Type', '条件 / Condition', '仓位 / Position', '说明 / Description'], [
    ['分批建仓 / Phased Entry', '回调至0.85-0.86企稳 / Pullback to 0.85-0.86 with stabilization', '20-30%', 'ETF适合定投式分批入场 / ETF suits DCA-style phased entry'],
    ['加仓 / Add Position', '突破0.90且站稳 / Break above 0.90 and hold', '加至40-50% / Increase to 40-50%', '确认突破后加仓 / Add after breakout confirmation'],
    ['定投 / DCA', '每周固定时间买入 / Buy at fixed weekly time', '每期固定金额 / Fixed amount per period', 'ETF适合长期定投，降低波动 / ETF suits long-term DCA to reduce volatility'],
    ['持有观望 / Hold & Watch', '价格在0.85上方 / Price above 0.85', '持有 / Hold', '等待趋势延续 / Wait for trend continuation'],
])
p = doc.add_paragraph()
run = p.add_run(safe('[重要/Important] 注意：ETF场内交易为T+1，当日买入最早次日才能卖出。ETF适合中长期配置，不建议频繁短线交易。/ Note: ETF on-exchange trading is T+1 — buy today, sell tomorrow earliest. ETFs suit medium-to-long-term allocation; frequent short-term trading not recommended.'))
run.font.size = Pt(11)
run.bold = True
run.font.color.rgb = RGBColor(0xcc, 0x33, 0x33)
run.font.name = 'SimHei'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

add_para(safe('建议：创业板成长ETF受益于成长风格占优，适合中长期定投。逢回调0.85-0.86分批建仓，中长期持有。注意0.90附近获利回吐压力。/ Suggestion: ChiNext Growth ETF benefits from the growth-style dominance, suitable for medium-to-long-term DCA. Enter in phases on pullbacks to 0.85-0.86, hold medium-to-long term. Watch for profit-taking pressure near 0.90.'))

# ===== 十、风险提示 / Risk Warning =====
add_title(safe('十、风险提示 / 10. Risk Warning'), level=1)
add_table(['风险类型 / Risk Type', '说明 / Description', '应对 / Countermeasure'], [
    ['市场风险 / Market Risk', '创业板波动大，年化波动约30% / ChiNext is highly volatile, annualized volatility ~30%', '控制仓位，分散投资 / Control position size, diversify'],
    ['风格切换风险 / Style Rotation Risk', '成长可能切换为价值风格 / Growth may rotate to value', '关注市场风格变化 / Monitor style shifts'],
    ['跟踪误差 / Tracking Error', 'ETF可能偏离指数表现 / ETF may deviate from index performance', '避免高溢价买入 / Avoid buying at high premium'],
    ['政策风险 / Policy Risk', '科技/医药行业政策变化 / Tech/Pharma policy changes', '关注政策风向 / Watch policy direction'],
])

# ===== 十一、综合评级 / Comprehensive Rating =====
add_title(safe('十一、综合评级 / 11. Comprehensive Rating'), level=1)
add_table(['维度 / Dimension', '评分 / Rating', '说明 / Description'], [
    ['技术/趋势 / Technical/Trend', 'B+', '放量突破反包，短期偏强，0.90阻力 / Volume breakout engulfing, short-term bullish, resistance at 0.90'],
    ['跟踪指数 / Tracking Index', 'A-', '创成长+5.08%领涨，成长因子有效 / ChiNext Growth +5.08% leads; growth factor effective'],
    ['基金基本面 / Fund Fundamentals', 'A', '华夏基金运作成熟，费率合理 / ChinaAMC well-established, reasonable fees'],
    ['市场环境 / Market Environment', 'B+', '成长风格占优，北向流入 / Growth style dominant, Northbound inflows'],
    ['综合评级 / Composite Rating', 'B+/A-', '[低/Low]风险，震荡偏多，分批布局 / Low risk, consolidating bullish, phase in'],
])
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run(safe('[一句话总结 / One-Sentence Summary] 创业板成长ETF 05-22大涨5.61%反包，短期偏强。成长风格占优环境下，创成长指数跑赢创业板指近一倍。05-25预计0.86-0.90震荡偏强，建议回调0.85-0.86分批建仓20-30%，突破0.90加仓至40-50%。ETF适合中长期定投。[低]风险。/ ChiNext Growth ETF surged 5.61% on 05-22 engulfing losses; short-term bullish. In a growth-dominant environment, the ChiNext Growth Index nearly doubled the ChiNext Index\'s return. 05-25 expected range 0.86-0.90 consolidating bullish. Recommended: phase in 20-30% on pullback to 0.85-0.86, add to 40-50% on break above 0.90. ETF suits long-term DCA. [Low] risk.'))
run.font.size = Pt(12)
run.bold = True
run.font.name = 'SimHei'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

# ===== AIGC声明 / AIGC Disclaimer =====
doc.add_paragraph()
aigc_title = doc.add_paragraph()
aigc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = aigc_title.add_run('[AIGC声明 / AIGC Disclaimer]')
run.font.size = Pt(11)
run.bold = True
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.name = 'SimHei'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

aigc_text = doc.add_paragraph()
aigc_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = aigc_text.add_run('本文档由AI辅助生成，数据来源于公开市场信息。本文档仅供参考学习，不构成投资建议。投资有风险，入市需谨慎。基金的过往业绩不预示未来表现。\nThis document is AI-assisted and based on publicly available market data. It is for reference and learning purposes only and does not constitute investment advice. Investing involves risk. Past performance of a fund does not indicate future results.')
run.font.size = Pt(9)
run.font.italic = True
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.name = 'SimHei'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

doc.add_paragraph()

for d in ['[注意] 本方案仅供参考，不构成投资建议。基金有风险，投资需谨慎。\n[Note] This plan is for reference only and does not constitute investment advice. Funds carry risk; invest with caution.',
          '[注意] ETF交易为T+1，当日买入最早次日才能卖出。\n[Note] ETF trading is T+1 — buy today, sell tomorrow earliest.',
          '[注意] 基金的过往业绩并不预示其未来表现，投资者应理性判断。\n[Note] Past performance does not guarantee future results. Investors should exercise rational judgment.',
          '[注意] 建议采用定投方式分批入场，控制单次投入比例。\n[Note] It is recommended to enter in phases via DCA, controlling single-instance allocation.']:
    p = doc.add_paragraph()
    run = p.add_run(safe(d))
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.name = 'SimHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

# ===== 中英双语风险声明 / Bilingual Risk Disclaimer =====
doc.add_paragraph()
risk_title = doc.add_paragraph()
risk_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = risk_title.add_run('【风险声明 / Risk Disclaimer】')
run.font.size = Pt(10)
run.bold = True
run.font.color.rgb = RGBColor(0x99, 0x33, 0x33)
run.font.name = 'SimHei'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

risk_disclaimers = [
    ('投资有风险，入市需谨慎。本方案中的所有数据、分析和建议均基于公开信息和AI模型生成，不构成任何形式的投资建议、要约或承诺。',
     'Investing involves risk. All data, analysis, and suggestions in this plan are generated based on public information and AI models, and do not constitute any form of investment advice, offer, or commitment.'),
    ('本方案中的预测和预判仅为基于现有信息的推演，实际市场走势可能因多种因素（包括但不限于政策变化、宏观经济波动、突发事件等）而与预测产生重大偏差。',
     'All forecasts and projections in this plan are extrapolations based on available information. Actual market movements may deviate materially due to various factors including but not limited to policy changes, macroeconomic fluctuations, and unforeseen events.'),
    ('过往表现不代表未来收益。任何投资决策均应基于投资者自身的风险承受能力、投资目标和财务状况，并在必要时咨询专业投资顾问。',
     'Past performance is not indicative of future returns. All investment decisions should be based on the investor\'s own risk tolerance, investment objectives, and financial situation, with consultation from professional investment advisors when necessary.'),
    ('本方案的作者及生成平台不对因使用本方案中的信息而导致的任何直接或间接损失承担责任。',
     'The author and generating platform of this plan assume no liability for any direct or indirect losses arising from the use of information contained herein.'),
]

for cn, en in risk_disclaimers:
    p = doc.add_paragraph()
    run = p.add_run(f'{cn}\n{en}')
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.name = 'SimHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

doc.add_paragraph()
sig_p = doc.add_paragraph()
sig_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sig_p.add_run('媳妇智投全球通用版出品，必属精品 / Crafted by Xifu Smart Invest Global Edition — Excellence Guaranteed')
run.font.size = Pt(14)
run.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
run.font.name = 'SimHei'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

aigc_btm = doc.add_paragraph()
aigc_btm.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = aigc_btm.add_run('[AIGC标识 / AIGC Label] 本方案由AI生成，内容仅供参考，不构成投资建议。\nThis plan is AI-generated, for reference only, and does not constitute investment advice.')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.name = 'SimHei'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

output_path = "/home/sandbox/.openclaw/workspace/创业板成长ETF_159967_投资方案_2026-05-23.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
print(f"Size: {os.path.getsize(output_path)} bytes")

# Verify
doc2 = Document(output_path)
bad = sum(1 for p in doc2.paragraphs if 1 <= len(p.text.strip()) <= 2 and p.text.strip() not in ['-', '--', '*', ''])
emoji_count = 0
for p in doc2.paragraphs:
    for ch in p.text:
        if ord(ch) > 0xFFFF or (0x2600 <= ord(ch) <= 0x27BF) or (0x1F300 <= ord(ch) <= 0x1F9FF):
            emoji_count += 1
for t in doc2.tables:
    for r in t.rows:
        for c in r.cells:
            for ch in c.text:
                if ord(ch) > 0xFFFF or (0x2600 <= ord(ch) <= 0x27BF) or (0x1F300 <= ord(ch) <= 0x1F9FF):
                    emoji_count += 1

print(f"Paras: {len(doc2.paragraphs)}, Tables: {len(doc2.tables)}")
print(f"Single-char: {bad}, Emoji: {emoji_count}")

text_all = ''
for p in doc2.paragraphs:
    text_all += p.text
for t in doc2.tables:
    for r in t.rows:
        for c in r.cells:
            text_all += c.text

print(f"  T+1: {'[是]' if 'T+1' in text_all else '[否]'}")
print(f"  AIGC: {'[是]' if 'AIGC' in text_all else '[否]'}")
print(f"  Signature: {'[是]' if '媳妇智投全球通用版' in text_all else '[否]'}")
print(f"  Fund warning: {'[是]' if '过往业绩' in text_all else '[否]'}")
print(f"  Bilingual Risk Disclaimer: {'[是]' if 'Risk Disclaimer' in text_all else '[否]'}")
print(f"  Bilingual: {'[是]' if 'Investment Plan' in text_all else '[否]'}")

print("\n=== 章节结构 / Section Structure ===")
for p in doc2.paragraphs:
    t = p.text.strip()
    if t.startswith(('一、','二、','三、','四、','五、','六、','七、','八、','九、','十、','十一、')):
        print(f"  {t[:70]}")
