#!/usr/bin/env python3
"""将简历文件转换为 AI 可读的简历包。"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import shutil
import subprocess
import sys
import tempfile
from dataclasses import asdict, dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Callable


SUPPORTED_TEXT_SUFFIXES = {".txt", ".md", ".markdown"}
SUPPORTED_DOC_SUFFIXES = {".pdf", ".docx", ".txt", ".md", ".markdown"}
TRUNCATE_ERROR_CHARS = 5000


class ExtractionError(RuntimeError):
    def __init__(self, message: str, parser_attempts: list["ParserAttempt"] | None = None):
        super().__init__(message)
        self.parser_attempts = parser_attempts or []


@dataclass
class ParserAttempt:
    parser: str
    status: str
    started_at: str
    finished_at: str
    detail: str | None = None
    error: str | None = None


@dataclass
class TextStats:
    character_count: int
    non_whitespace_character_count: int
    line_count: int
    page_count: int | None
    empty_page_count: int | None
    suspected_scanned: bool


@dataclass
class ExtractionResult:
    markdown: str
    raw: dict[str, Any]
    parser: str
    quality: str
    warnings: list[str]
    parser_attempts: list[ParserAttempt]


def utc_now() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat()


def truncate_detail(text: str, limit: int = TRUNCATE_ERROR_CHARS) -> str:
    text = text.strip()
    if len(text) <= limit:
        return text
    return text[:limit] + f"\n... <truncated {len(text) - limit} chars>"


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as fh:
        for chunk in iter(lambda: fh.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def find_command(command_name: str) -> str | None:
    found = shutil.which(command_name)
    if found:
        return found
    user_local = Path.home() / ".local" / "bin" / command_name
    if user_local.exists() and user_local.is_file():
        return str(user_local)
    return None


def read_text(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(errors="replace")


def normalize_text(text: str) -> str:
    text = text.replace("\x00", "")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    return text.strip()


def text_to_markdown(text: str, title: str = "简历文本") -> str:
    text = normalize_text(text)
    if not text:
        return ""
    if text.lstrip().startswith("#"):
        return text + "\n"
    return f"# {title}\n\n{text}\n"


def find_output_file(output_target: Path, suffix: str) -> Path | None:
    if output_target.is_file():
        return output_target
    if output_target.is_dir():
        matches = sorted(output_target.rglob(f"*{suffix}"))
        return matches[0] if matches else None
    return None


def import_docling_python(path: Path) -> ExtractionResult:
    try:
        from docling.document_converter import DocumentConverter  # type: ignore
    except Exception as exc:  # pragma: no cover - 依赖可选包
        raise ExtractionError(f"docling Python 包不可用：{exc}") from exc

    try:
        converter = DocumentConverter()
        result = converter.convert(str(path))
        document = result.document
        markdown = document.export_to_markdown()
        try:
            raw_doc = document.export_to_dict()
        except Exception:
            raw_doc = {"note": "Docling 文档对象无法导出为 dict。"}
    except Exception as exc:  # pragma: no cover - 依赖可选包
        raise ExtractionError(f"docling Python 转换失败：{exc}") from exc

    return ExtractionResult(
        markdown=text_to_markdown(markdown, "简历"),
        raw={"format": "docling", "document": raw_doc},
        parser="docling-python",
        quality="high",
        warnings=[],
        parser_attempts=[],
    )


def run_docling_cli(path: Path, timeout_seconds: int = 180) -> ExtractionResult:
    docling = find_command("docling")
    if not docling:
        raise ExtractionError("docling CLI 不可用")

    with tempfile.TemporaryDirectory(prefix="hr-docling-") as tmp:
        tmp_dir = Path(tmp)
        common_args = [
            docling,
            "--no-ocr",
            "--no-tables",
            "--output",
            str(tmp_dir),
            str(path),
        ]

        md_proc = subprocess.run(
            [common_args[0], "--to", "md", *common_args[1:]],
            text=True,
            capture_output=True,
            timeout=timeout_seconds,
            check=False,
        )
        if md_proc.returncode != 0:
            raise ExtractionError(md_proc.stderr.strip() or "docling CLI Markdown 转换失败")

        markdown_file = find_output_file(tmp_dir, ".md")
        markdown = markdown_file.read_text(encoding="utf-8") if markdown_file else md_proc.stdout

        raw_doc: dict[str, Any] = {}
        json_proc = subprocess.run(
            [common_args[0], "--to", "json", *common_args[1:]],
            text=True,
            capture_output=True,
            timeout=timeout_seconds,
            check=False,
        )
        if json_proc.returncode == 0:
            json_file = find_output_file(tmp_dir, ".json")
            raw_text = json_file.read_text(encoding="utf-8") if json_file else json_proc.stdout
            try:
                raw_doc = json.loads(raw_text) if raw_text.strip() else {}
            except json.JSONDecodeError:
                raw_doc = {"raw_json_text": raw_text}

    return ExtractionResult(
        markdown=text_to_markdown(markdown, "简历"),
        raw={"format": "docling", "document": raw_doc},
        parser="docling-cli",
        quality="high",
        warnings=["Docling CLI 已关闭 OCR 和表格模型，以避免运行时下载 OCR 模型。"],
        parser_attempts=[],
    )


def extract_pdf_with_pypdf(path: Path) -> ExtractionResult:
    reader_cls = None
    module_name = ""
    try:
        from pypdf import PdfReader  # type: ignore

        reader_cls = PdfReader
        module_name = "pypdf"
    except Exception:
        try:
            from PyPDF2 import PdfReader  # type: ignore

            reader_cls = PdfReader
            module_name = "PyPDF2"
        except Exception as exc:
            raise ExtractionError(f"pypdf/PyPDF2 不可用：{exc}") from exc

    try:
        reader = reader_cls(str(path))
        pages = [
            {"page": idx, "text": page.extract_text() or ""}
            for idx, page in enumerate(reader.pages, start=1)
        ]
    except Exception as exc:
        raise ExtractionError(f"{module_name} 提取失败：{exc}") from exc

    text = "\n\n".join(f"## 第 {page['page']} 页\n\n{page['text']}" for page in pages)
    return ExtractionResult(
        markdown=text_to_markdown(text, "简历"),
        raw={"format": "pdf-text", "pages": pages},
        parser=module_name,
        quality="fallback",
        warnings=["已使用回退解析器。版式、表格和扫描页可能不完整。"],
        parser_attempts=[],
    )


def extract_pdf_with_pymupdf(path: Path) -> ExtractionResult:
    try:
        import fitz  # type: ignore
    except Exception as exc:
        raise ExtractionError(f"PyMuPDF 不可用：{exc}") from exc

    pages = []
    try:
        with fitz.open(path) as doc:  # type: ignore[attr-defined]
            for idx, page in enumerate(doc, start=1):
                pages.append({"page": idx, "text": page.get_text("text")})
    except Exception as exc:
        raise ExtractionError(f"PyMuPDF 提取失败：{exc}") from exc

    text = "\n\n".join(f"## 第 {page['page']} 页\n\n{page['text']}" for page in pages)
    return ExtractionResult(
        markdown=text_to_markdown(text, "简历"),
        raw={"format": "pdf-text", "pages": pages},
        parser="pymupdf",
        quality="fallback",
        warnings=["已使用回退解析器。版式、表格和扫描页可能不完整。"],
        parser_attempts=[],
    )


def extract_pdf_with_pdftotext(path: Path) -> ExtractionResult:
    pdftotext = find_command("pdftotext")
    if not pdftotext:
        raise ExtractionError("pdftotext CLI 不可用")

    proc = subprocess.run(
        [pdftotext, "-layout", str(path), "-"],
        text=True,
        capture_output=True,
        timeout=120,
        check=False,
    )
    if proc.returncode != 0:
        raise ExtractionError(proc.stderr.strip() or "pdftotext 提取失败")

    return ExtractionResult(
        markdown=text_to_markdown(proc.stdout, "简历"),
        raw={"format": "pdf-text", "text": proc.stdout},
        parser="pdftotext",
        quality="fallback",
        warnings=["已使用回退解析器。版式、表格和扫描页可能不完整。"],
        parser_attempts=[],
    )


def extract_docx_with_python_docx(path: Path) -> ExtractionResult:
    try:
        import docx  # type: ignore
    except Exception as exc:
        raise ExtractionError(f"python-docx 不可用：{exc}") from exc

    try:
        document = docx.Document(str(path))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        tables = []
        for table in document.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            tables.append(rows)
    except Exception as exc:
        raise ExtractionError(f"python-docx 提取失败：{exc}") from exc

    table_text = []
    for table in tables:
        for row in table:
            table_text.append(" | ".join(row))
    text = "\n\n".join(paragraphs + table_text)
    return ExtractionResult(
        markdown=text_to_markdown(text, "简历"),
        raw={"format": "docx-text", "paragraphs": paragraphs, "tables": tables},
        parser="python-docx",
        quality="fallback",
        warnings=["已使用回退解析器。复杂格式可能不完整。"],
        parser_attempts=[],
    )


def extract_plain_text(path: Path) -> ExtractionResult:
    text = read_text(path)
    return ExtractionResult(
        markdown=text_to_markdown(text, "简历"),
        raw={"format": "plain-text", "text": text},
        parser="plain-text",
        quality="source",
        warnings=[],
        parser_attempts=[],
    )


def local_attempts_for_suffix(suffix: str) -> list[Callable[[Path], ExtractionResult]]:
    if suffix == ".pdf":
        return [extract_pdf_with_pypdf, extract_pdf_with_pymupdf, extract_pdf_with_pdftotext]
    if suffix == ".docx":
        return [extract_docx_with_python_docx]
    if suffix in SUPPORTED_TEXT_SUFFIXES:
        return [extract_plain_text]
    return []


def docling_attempts() -> list[Callable[[Path], ExtractionResult]]:
    return [import_docling_python, run_docling_cli]


def attempt_name(attempt: Callable[[Path], ExtractionResult]) -> str:
    return attempt.__name__.removeprefix("extract_").removeprefix("run_").removeprefix("import_")


def build_attempts(path: Path, parser: str, allow_model_downloads: bool) -> list[Callable[[Path], ExtractionResult]]:
    suffix = path.suffix.lower()
    local_attempts = local_attempts_for_suffix(suffix)
    if parser == "local":
        return local_attempts
    if parser == "docling":
        return docling_attempts() + local_attempts
    if allow_model_downloads:
        return docling_attempts() + local_attempts
    return local_attempts


def extract_resume(path: Path, parser: str, allow_model_downloads: bool = False) -> ExtractionResult:
    suffix = path.suffix.lower()
    attempts = build_attempts(path, parser, allow_model_downloads)
    if not attempts and parser == "auto":
        raise ExtractionError(f"不支持的简历文件类型：{suffix or '<none>'}")

    parser_attempts: list[ParserAttempt] = []
    errors = []
    for attempt in attempts:
        started_at = utc_now()
        name = attempt_name(attempt)
        try:
            result = attempt(path)
            if not normalize_text(result.markdown):
                error = f"{name}: 生成了空文本"
                errors.append(error)
                parser_attempts.append(
                    ParserAttempt(name, "failed", started_at, utc_now(), error="生成了空文本")
                )
                continue
            parser_attempts.append(
                ParserAttempt(name, "succeeded", started_at, utc_now(), detail=f"parser={result.parser}")
            )
            result.parser_attempts = parser_attempts
            return result
        except ExtractionError as exc:
            error = truncate_detail(str(exc))
            errors.append(f"{name}: {error}")
            parser_attempts.append(ParserAttempt(name, "failed", started_at, utc_now(), error=error))

    if parser == "docling" and not errors:
        raise ExtractionError("已请求 Docling，但没有 Docling 转换器成功执行。", parser_attempts)
    raise ExtractionError("没有解析器成功执行：\n- " + "\n- ".join(errors), parser_attempts)


def first_match(pattern: str, text: str, flags: int = 0) -> str | None:
    match = re.search(pattern, text, flags)
    return match.group(1).strip() if match else None


def detect_name(text: str) -> str | None:
    explicit = first_match(r"(?:姓名|Name)\s*[:：]\s*([^\n\r,，|/]+)", text, re.IGNORECASE)
    if explicit:
        return explicit.strip()

    stop_words = {
        "个人简历",
        "简历",
        "求职意向",
        "教育经历",
        "教育背景",
        "工作经历",
        "项目经历",
        "专业技能",
        "获奖情况",
        "科研经历",
        "resume",
        "curriculum vitae",
        "education",
        "experience",
        "skills",
    }
    for raw_line in text.splitlines()[:80]:
        line = raw_line.strip(" #\t:-|，,")
        lowered = line.lower()
        if not line or lowered in stop_words:
            continue
        if "@" in line or re.search(r"\d{5,}", line):
            continue
        spaced_chinese = re.fullmatch(r"[\u4e00-\u9fff](?:\s+[\u4e00-\u9fff]){1,3}", line)
        if spaced_chinese:
            return re.sub(r"\s+", "", line)
        chinese_name = re.fullmatch(r"[\u4e00-\u9fff]{2,4}", line)
        english_name = re.fullmatch(r"[A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3}", line)
        if chinese_name or english_name:
            return line
    return None


def detect_skills(text: str) -> list[str]:
    skill_terms = [
        "Python",
        "Java",
        "C++",
        "C#",
        "Go",
        "Rust",
        "JavaScript",
        "TypeScript",
        "React",
        "Vue",
        "Node.js",
        "SQL",
        "MySQL",
        "PostgreSQL",
        "Redis",
        "Kafka",
        "Docker",
        "Kubernetes",
        "Linux",
        "PyTorch",
        "TensorFlow",
        "scikit-learn",
        "LLM",
        "NLP",
        "RAG",
        "GNN",
        "推荐系统",
        "机器学习",
        "深度学习",
        "数据分析",
    ]
    found = []
    for term in skill_terms:
        if re.search(re.escape(term), text, re.IGNORECASE):
            found.append(term)
    return found


def detect_education_lines(text: str) -> list[str]:
    keywords = ("大学", "学院", "本科", "硕士", "博士", "学士", "MBA", "University", "College", "Bachelor", "Master", "PhD")
    lines = []
    for line in text.splitlines():
        compact = line.strip()
        if compact and any(keyword.lower() in compact.lower() for keyword in keywords):
            lines.append(compact)
    return lines[:12]


def build_candidate_draft(markdown: str, source_name: str, source_sha256: str) -> dict[str, Any]:
    text = re.sub(r"^#+\s*", "", markdown, flags=re.MULTILINE)
    email = first_match(r"([A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,})", text, re.IGNORECASE)
    phone = first_match(r"((?:\+?86[-\s]?)?1[3-9]\d[-\s]?\d{4}[-\s]?\d{4})", text)

    return {
        "schema_version": "resume-candidate-draft/v0.1",
        "source": {
            "file_name": source_name,
            "sha256": source_sha256,
        },
        "identity": {
            "name": detect_name(text),
            "phone": phone,
            "email": email,
        },
        "profile": {
            "current_company": None,
            "years_of_experience": None,
            "education_evidence": detect_education_lines(text),
            "skill_mentions": detect_skills(text),
        },
        "resume": {
            "summary": None,
            "markdown_file": "resume.md",
            "raw_json_file": "resume.raw.json",
        },
        "extraction_notes": [
            "此草稿由确定性规则生成，可能不完整。",
            "请结合简历 Markdown 和提取提示词生成最终结构化字段。",
            "不要编造简历中不存在的字段。",
        ],
    }


def compute_text_stats(markdown: str, raw: dict[str, Any], source_suffix: str) -> TextStats:
    text = normalize_text(re.sub(r"^#+\s*", "", markdown, flags=re.MULTILINE))
    pages = raw.get("pages")
    page_count: int | None = None
    empty_page_count: int | None = None
    if isinstance(pages, list):
        page_count = len(pages)
        empty_page_count = sum(1 for page in pages if not normalize_text(str(page.get("text", ""))))
    elif source_suffix == ".pdf":
        page_numbers = {int(match) for match in re.findall(r"## 第 (\d+) 页", markdown)}
        if page_numbers:
            page_count = max(page_numbers)
            empty_page_count = None

    non_ws = len(re.sub(r"\s+", "", text))
    suspected_scanned = source_suffix == ".pdf" and (
        non_ws < 120 or (page_count is not None and empty_page_count is not None and page_count > 0 and empty_page_count / page_count >= 0.5)
    )
    return TextStats(
        character_count=len(text),
        non_whitespace_character_count=non_ws,
        line_count=len(text.splitlines()) if text else 0,
        page_count=page_count,
        empty_page_count=empty_page_count,
        suspected_scanned=suspected_scanned,
    )


def build_review(candidate_draft: dict[str, Any], result: ExtractionResult, stats: TextStats) -> tuple[bool, list[str], list[str]]:
    missing = []
    identity = candidate_draft.get("identity", {})
    for field in ("name", "phone", "email"):
        if not identity.get(field):
            missing.append(field)

    reasons = []
    if result.quality == "fallback":
        reasons.append("使用了回退解析器，版式、表格和时间线可能不完整。")
    if stats.suspected_scanned:
        reasons.append("PDF 可读文本较少，疑似扫描件或图片型简历。")
    if missing:
        reasons.append("关键字段缺失或未能可靠识别：" + ", ".join(missing))
    return bool(reasons), reasons, missing


def write_json(path: Path, payload: dict[str, Any]) -> None:
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def build_output_dir(source: Path, out_dir: Path | None) -> Path:
    if out_dir:
        return out_dir
    stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    return source.parent / f"{source.stem}.resume_bundle.{stamp}"


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("resume_file", help="PDF、DOCX、TXT 或 Markdown 简历的路径。")
    parser.add_argument("--out-dir", help="写入简历包的目录。")
    parser.add_argument(
        "--parser",
        default="auto",
        choices=["auto", "local", "docling"],
        help="解析器策略。auto 默认只使用本地安全解析器；docling 会先尝试高保真解析再回退。",
    )
    parser.add_argument(
        "--allow-model-downloads",
        action="store_true",
        help="允许 auto 策略优先尝试可能下载模型的 Docling 高保真解析。",
    )
    parser.add_argument("--no-copy-original", action="store_true", help="不要将原始文件复制到简历包中。")
    args = parser.parse_args()

    source = Path(args.resume_file).expanduser().resolve()
    if not source.exists() or not source.is_file():
        print(f"未找到简历文件：{source}", file=sys.stderr)
        return 2
    if source.suffix.lower() not in SUPPORTED_DOC_SUFFIXES and args.parser != "docling":
        print(f"未使用 Docling 时不支持该文件类型：{source.suffix}", file=sys.stderr)
        return 2

    out_dir = build_output_dir(source, Path(args.out_dir).expanduser().resolve() if args.out_dir else None)
    out_dir.mkdir(parents=True, exist_ok=True)

    file_hash = sha256_file(source)
    started_at = utc_now()
    try:
        result = extract_resume(source, args.parser, allow_model_downloads=args.allow_model_downloads)
    except ExtractionError as exc:
        write_json(
            out_dir / "extraction_report.json",
            {
                "schema_version": "resume-extraction-report/v0.1",
                "source_file": str(source),
                "source_sha256": file_hash,
                "started_at": started_at,
                "finished_at": utc_now(),
                "status": "failed",
                "error": truncate_detail(str(exc)),
                "parser_attempts": [asdict(attempt) for attempt in exc.parser_attempts],
            },
        )
        print(str(exc), file=sys.stderr)
        return 1

    original_name = f"original{source.suffix.lower()}"
    if not args.no_copy_original:
        shutil.copy2(source, out_dir / original_name)

    markdown = normalize_text(result.markdown) + "\n"
    (out_dir / "resume.md").write_text(markdown, encoding="utf-8")
    write_json(out_dir / "resume.raw.json", result.raw)
    candidate_draft = build_candidate_draft(markdown, source.name, file_hash)
    write_json(out_dir / "candidate_draft.json", candidate_draft)

    text_stats = compute_text_stats(markdown, result.raw, source.suffix.lower())
    review_required, review_reasons, missing_fields = build_review(candidate_draft, result, text_stats)
    report = {
        "schema_version": "resume-extraction-report/v0.1",
        "source_file": str(source),
        "source_sha256": file_hash,
        "original_copy": original_name if not args.no_copy_original else None,
        "started_at": started_at,
        "finished_at": utc_now(),
        "status": "succeeded",
        "parser": result.parser,
        "quality": result.quality,
        "warnings": result.warnings,
        "parser_attempts": [asdict(attempt) for attempt in result.parser_attempts],
        "text_stats": asdict(text_stats),
        "review_required": review_required,
        "review_reasons": review_reasons,
        "missing_or_ambiguous_fields": missing_fields,
        "outputs": {
            "markdown": "resume.md",
            "raw_json": "resume.raw.json",
            "candidate_draft": "candidate_draft.json",
        },
        "privacy": {
            "contains_personal_data": True,
            "handling_note": "请将此简历包视为机密招聘数据。",
        },
    }
    write_json(out_dir / "extraction_report.json", report)
    write_json(
        out_dir / "manifest.json",
        {
            "schema_version": "resume-bundle-manifest/v0.1",
            "created_at": report["finished_at"],
            "files": sorted(item.name for item in out_dir.iterdir() if item.is_file()),
        },
    )

    print(
        json.dumps(
            {
                "status": "succeeded",
                "bundle_dir": str(out_dir),
                "parser": result.parser,
                "quality": result.quality,
                "review_required": review_required,
            },
            ensure_ascii=False,
        )
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
