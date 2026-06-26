"""
DOCX → outline importer (uses python-docx).

Extracts paragraphs and tables from a Word document and maps them to
DeckCraft page types. Each top-level section (delimited by blank paragraphs
or heading style) becomes one slide.

This is a simpler heuristic than the PDF importer — DOCX structure is
usually cleaner, so the default behavior is to treat each Heading 1 as a
new section starting with a content slide.
"""
from typing import Dict, Any, List, Optional
import re
import docx

from .base import (
    looks_like_toc, looks_like_section_divider,
    looks_like_table, looks_like_stats,
    extract_bullets, extract_stats, extract_heading,
    make_outline_page,
)


def docx_to_outline(
    docx_path: str,
    theme: str = "business",
    canvas: str = "16:9",
    max_pages: int = 200,
) -> Dict[str, Any]:
    """Convert a Word document to a DeckCraft outline dict.

    The conversion uses the document's heading structure:
    - First non-empty paragraph → "cover" slide
    - Each "Heading 1" starts a new section
    - Content under each heading becomes bullets in a "content" slide
    - Tables are extracted into "table" slides
    - Pages with stat-like lines (number + label) become "stat_cards"

    Args:
        docx_path: Path to the .docx file.
        theme: Theme name (default: "business").
        canvas: Canvas preset (default: "16:9").
        max_pages: Max slides to generate (safety limit).

    Returns:
        Outline dict with 'pages' list.
    """
    doc = docx.Document(docx_path)
    body_elements = list(doc.element.body)
    pages: List[Dict[str, Any]] = []

    # Pass 1: collect paragraphs and tables in document order
    items = []
    for el in body_elements:
        tag = el.tag.split("}")[-1]
        if tag == "p":
            # Find the corresponding paragraph object
            for p in doc.paragraphs:
                if p._element is el:
                    text = p.text.strip()
                    style = p.style.name if p.style else ""
                    items.append({"kind": "p", "text": text, "style": style})
                    break
        elif tag == "tbl":
            for t in doc.tables:
                if t._element is el:
                    rows = []
                    for r in t.rows:
                        rows.append([c.text.strip() for c in r.cells])
                    if rows:
                        items.append({"kind": "tbl", "rows": rows})
                    break

    # Pass 2: split into sections by Heading 1
    sections = []  # list of {title, content: [(kind, data)]}
    current = {"title": "", "content": []}

    for item in items:
        if item["kind"] == "p":
            style = item.get("style", "")
            text = item["text"]
            if style.startswith("Heading 1") or style == "Title":
                # Start new section
                if current["title"] or current["content"]:
                    sections.append(current)
                current = {"title": text or "(Untitled)", "content": []}
                continue
            if not text:
                # Blank line — section break if we have content
                if current["content"]:
                    sections.append(current)
                    current = {"title": "", "content": []}
                continue
            current["content"].append({"kind": "p", "text": text, "style": style})
        else:
            current["content"].append(item)

    if current["title"] or current["content"]:
        sections.append(current)

    if not sections:
        return {
            "theme": theme,
            "canvas": canvas,
            "pages": [make_outline_page("cover", title="Imported from DOCX")],
        }

    # Pass 3: convert sections to pages
    for i, section in enumerate(sections):
        if len(pages) >= max_pages:
            break

        if i == 0 and not section["title"]:
            # First section without a heading — treat as cover
            text = _flatten_text(section["content"])
            title = section["title"] or extract_heading(text) or "Imported Document"
            pages.append(make_outline_page("cover", title=title))
            continue

        if i == 0:
            # Use first section as cover (title only)
            pages.append(make_outline_page("cover", title=section["title"]))
            continue

        # Check for table
        table_data = next((c for c in section["content"] if c["kind"] == "tbl"), None)
        if table_data is not None:
            headers = table_data["rows"][0] if table_data["rows"] else []
            rows = table_data["rows"][1:] if len(table_data["rows"]) > 1 else []
            pages.append(make_outline_page(
                "table", title=section["title"],
                headers=headers, rows=rows,
            ))
            continue

        text = _flatten_text(section["content"])

        # Check for TOC (agenda-like numbered list)
        # Include the heading text since "Agenda" heading + numbered list is also a TOC
        combined = section["title"] + "\n" + text if section["title"] else text
        if looks_like_toc(combined):
            items = []
            for line in text.split("\n"):
                line = line.strip()
                m = re.match(r"^\s*(\d+|[IVX]+)[\s\.\)、]+(.+)$", line)
                if m and 5 < len(m.group(2).strip()) < 80:
                    items.append([m.group(1), m.group(2).strip(), ""])
            if not items:
                # Fallback: all non-empty short lines
                for line in text.split("\n"):
                    line = line.strip()
                    if 5 < len(line) < 80 and not line.startswith("-"):
                        items.append([str(len(items) + 1), line, ""])
            pages.append(make_outline_page("toc", items=items[:8]))
            continue

        # Check for stats (but only if heading is NOT a TOC/Agenda heading)
        if not re.match(r"^\s*(agenda|contents?|目录|大纲|议程)\b", section["title"], re.IGNORECASE):
            stats = extract_stats(text)
            if stats and len(stats) >= 2:
                pages.append(make_outline_page(
                    "stat_cards", title=section["title"], stats=stats,
                ))
                continue

        # Check for section divider (very short body)
        if looks_like_section_divider(text) and len(text.strip()) < 60:
            pages.append(make_outline_page("section", title=section["title"]))
            continue

        # Default: content slide
        bullets = extract_bullets(text)
        if not bullets:
            bullets = [l for l in text.split("\n") if 5 < len(l.strip()) < 200][:8]
        pages.append(make_outline_page(
            "content", title=section["title"] or "Section",
            bullets=bullets,
        ))

    # Always add a closing if the last page isn't already one
    if pages and pages[-1].get("type") != "closing":
        pages.append(make_outline_page("closing", title="Thank You", message=""))

    return {
        "theme": theme,
        "canvas": canvas,
        "pages": pages,
    }


def _flatten_text(items: List[dict]) -> str:
    """Flatten a list of section items into a single text blob."""
    parts = []
    for it in items:
        if it["kind"] == "p":
            parts.append(it["text"])
    return "\n".join(parts)


__all__ = ["docx_to_outline"]
