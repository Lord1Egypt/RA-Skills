"""
docx-to-md.py  (v2 — security-hardened)
Converts a Word (.docx) file to Markdown.
- Maps heading styles → # / ## / ###
- Converts tables → Markdown table syntax
- Falls back to Claude vision for embedded images (only with --allow-vision)

Usage:
    python docx-to-md.py <input.docx> <output.md> [--allow-vision]

Dependencies installed to /tmp/office_md_deps/ (isolated, pinned versions).
"""

import sys, base64, json, urllib.request, subprocess
from pathlib import Path

# ── isolated dependency install ───────────────────────────────────────────────
_DEP_DIR = Path("/tmp/office_md_deps")
_DEP_DIR.mkdir(exist_ok=True)
subprocess.run(
    [
        sys.executable, "-m", "pip", "install", "--quiet",
        "--target", str(_DEP_DIR),
        "python-docx==1.1.2",
    ],
    check=True,
)
if str(_DEP_DIR) not in sys.path:
    sys.path.insert(0, str(_DEP_DIR))

from docx import Document
from docx.oxml.ns import qn


# ── vision helper (only called when --allow-vision is set) ────────────────────
def call_vision(b64_data: str, media_type: str = "image/png") -> str:
    payload = {
        "model": "claude-sonnet-4-20250514",
        "max_tokens": 2000,
        "messages": [{
            "role": "user",
            "content": [
                {"type": "image", "source": {"type": "base64", "media_type": media_type, "data": b64_data}},
                {"type": "text", "text": (
                    "Extract ALL text and structured content from this image into clean Markdown. "
                    "Preserve headings, lists, and tables. "
                    "Output ONLY the extracted markdown — no preamble."
                )}
            ]
        }]
    }
    req = urllib.request.Request(
        "https://api.anthropic.com/v1/messages",
        data=json.dumps(payload).encode("utf-8"),
        headers={"Content-Type": "application/json", "anthropic-version": "2023-06-01"},
        method="POST",
    )
    try:
        with urllib.request.urlopen(req) as resp:
            return json.loads(resp.read())["content"][0]["text"]
    except Exception as e:
        return f"*(vision extraction failed: {e})*"


# ── main ──────────────────────────────────────────────────────────────────────
def convert(src: str, dst: str, allow_vision: bool = False):
    doc = Document(src)
    md_lines = []
    vision_count = 0
    vision_skipped = 0

    for para in doc.paragraphs:
        style = para.style.name.lower()
        text  = para.text.strip()

        if not text:
            # Check for inline images
            for run in para.runs:
                for blip in run._element.iter(qn("a:blip")):
                    embed = blip.get(qn("r:embed"))
                    if embed and embed in doc.part.rels:
                        if not allow_vision:
                            vision_skipped += 1
                            md_lines.append(
                                "*(embedded image — vision extraction skipped; "
                                "re-run with --allow-vision after user confirmation)*"
                            )
                            continue
                        img_part = doc.part.rels[embed].target_part
                        mt  = img_part.content_type
                        b64 = base64.standard_b64encode(img_part.blob).decode("utf-8")
                        md_lines.append(call_vision(b64, mt))
                        vision_count += 1
            md_lines.append("")
            continue

        if   "heading 1" in style: md_lines.append(f"# {text}")
        elif "heading 2" in style: md_lines.append(f"## {text}")
        elif "heading 3" in style: md_lines.append(f"### {text}")
        elif "list"      in style: md_lines.append(f"- {text}")
        else:                      md_lines.append(text)

    for table in doc.tables:
        header = [cell.text.strip() for cell in table.rows[0].cells]
        md_lines.append("| " + " | ".join(header) + " |")
        md_lines.append("| " + " | ".join("---" for _ in header) + " |")
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            md_lines.append("| " + " | ".join(cells) + " |")
        md_lines.append("")

    Path(dst).write_text("\n".join(md_lines), encoding="utf-8")
    if vision_skipped:
        print(
            f"  VISION_REQUIRED: {vision_skipped} embedded image(s) detected. "
            f"Re-run with --allow-vision after user confirmation."
        )
    print(f"  Saved → {dst}  ({len(doc.paragraphs)} paragraphs, {vision_count} images via vision)")


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: docx-to-md.py <input.docx> <output.md> [--allow-vision]")
        sys.exit(1)
    allow_vision = "--allow-vision" in sys.argv
    convert(sys.argv[1], sys.argv[2], allow_vision=allow_vision)
