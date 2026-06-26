#!/usr/bin/env python3
"""
tex2docx — LaTeX to Word (.docx) converter.
Requires: pandoc (system) + python-docx + lxml + pypandoc_binary.

Usage: python tex2docx.py input.tex [output.docx]
"""
import sys, os, re, shutil
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
import pypandoc
from lxml import etree

_CENTER = WD_ALIGN_PARAGRAPH.CENTER
NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NS_M = 'http://schemas.openxmlformats.org/officeDocument/2006/math'


def clean(t):
    t = re.sub(r"\\'(\w)", r'\1', t)
    t = re.sub(r'\\`(\w)', r'\1', t)
    t = re.sub(r'\\\^(\w)', r'\1', t)
    t = re.sub(r'\\"(\w)', r'\1', t)
    t = re.sub(r'\\textquoteright', "'", t)
    t = re.sub(r'\\ldots', '...', t)
    t = re.sub(r'\\([%_&#{}])', r'\1', t)
    t = re.sub(r'~', ' ', t)
    t = re.sub(r'\\[,;\\ ]', ' ', t)
    for cmd in ['hline', 'midrule', 'toprule', 'bottomrule']:
        t = re.sub(r'\\' + cmd, '', t)
    t = re.sub(r'\\cline\{[^}]*\}', '', t)
    t = re.sub(r'\\(?:textbf|text|mathrm|mathbf|textit|textsf|texttt)\{(.*?)\}', r'\1', t)
    t = re.sub(r'\\\\', '\n', t)
    t = re.sub(r'\$', '', t)
    t = re.sub(r'\s+', ' ', t).strip()
    return t


def find_image(rel, fig_dir):
    base = os.path.basename(rel)
    noext = re.sub(r'\.\w+$', '', base)
    for ext in ['.png', '.pdf']:
        p = os.path.join(fig_dir, noext + ext)
        if os.path.exists(p):
            return p
    return None


def main():
    if len(sys.argv) < 2:
        print("Usage: python tex2docx.py input.tex [output.docx]")
        sys.exit(1)

    tex_path = os.path.abspath(sys.argv[1])
    base = os.path.splitext(tex_path)[0]
    out_path = os.path.abspath(sys.argv[2]) if len(sys.argv) > 2 else base + '.docx'
    fig_dir = os.path.join(os.path.dirname(tex_path), 'figures')
    tmp_pandoc = base + '_tmp_pandoc.docx'

    tex = open(tex_path, encoding='utf-8').read()
    tex = re.sub(r'(?<!\\)%.*$', '', tex, flags=re.MULTILINE)

    # === Step 1: Pandoc ===
    print("[1/3] Pandoc conversion...")
    cwd = os.getcwd()
    os.chdir(os.path.dirname(tex_path))
    pypandoc.convert_file(
        os.path.basename(tex_path), 'docx',
        outputfile=os.path.basename(tmp_pandoc),
        extra_args=['--from=latex', '--to=docx', '--wrap=preserve', '--resource-path=.:figures']
    )
    os.chdir(cwd)
    print(f"  Temp: {os.path.getsize(tmp_pandoc)} bytes")

    omml_paras = []
    for p in Document(tmp_pandoc).paragraphs:
        if p._element.findall(f'.//{{{NS_M}}}oMath'):
            elem = etree.fromstring(etree.tostring(p._element))
            pPr = elem.find(f'{{{NS_W}}}pPr')
            if pPr is None:
                pPr = etree.SubElement(elem, f'{{{NS_W}}}pPr')
            jc = pPr.find(f'{{{NS_W}}}jc')
            if jc is None:
                jc = etree.SubElement(pPr, f'{{{NS_W}}}jc')
            jc.set(f'{{{NS_W}}}val', 'center')
            omml_paras.append(elem)
    os.remove(tmp_pandoc)
    print(f"  OMML equations: {len(omml_paras)}")

    # === Extract data ===
    body_m = re.search(r'\\begin\{document\}(.*)\\end\{document\}', tex, re.DOTALL)
    body = body_m.group(1) if body_m else tex

    title_m = re.search(r'\\title\{(.*?)\}', tex, re.DOTALL)
    title = ''
    if title_m:
        raw = title_m.group(1).replace('\\large ', '').replace('\\\\', ' ')
        title = clean(raw)

    author_m = re.search(r'\\IEEEauthorblockN\{(.*?)\}', tex)
    author = author_m.group(1) if author_m else ''

    abs_m = re.search(r'\\begin\{abstract\}(.*)\\end\{abstract\}', body, re.DOTALL)
    abstract = clean(abs_m.group(1)) if abs_m else ''

    kw_m = re.search(r'\\begin\{IEEEkeywords\}(.*)\\end\{IEEEkeywords\}', body, re.DOTALL)
    keywords = clean(kw_m.group(1)) if kw_m else ''

    tables = []
    for m in re.finditer(
        r'\\begin\{table\}.*?\\caption\{(.*?)\}.*?\\begin\{tabular\}.*?\n(.*?)\\end\{tabular\}.*?\\end\{table\}',
        tex, re.DOTALL
    ):
        rows = []
        for rr in re.split(r'\\\\', m.group(2)):
            cells = [clean(c) for c in re.split(r'&', rr) if clean(c)]
            if cells:
                rows.append(cells)
        tables.append({'caption': clean(m.group(1)), 'data': rows})

    bib_m = re.search(r'\\begin\{thebibliography\}(.*)\\end\{thebibliography\}', tex, re.DOTALL)
    refs = []
    if bib_m:
        entries = re.split(r'\\bibitem\{(.*?)\}', bib_m.group(0))
        for j in range(1, len(entries) - 1, 2):
            refs.append({'id': entries[j], 'text': re.sub(r'\s+', ' ', clean(entries[j + 1]))})

    figs = []
    for m in re.finditer(r'\\includegraphics(?:\[.*?\])?\{(.*?)\}', tex):
        path = m.group(1).strip()
        after = tex[m.end():m.end() + 800]
        cap_m = re.search(r'\\caption\{(.*?)\}', after)
        figs.append({'path': path, 'caption': clean(cap_m.group(1)) if cap_m else ''})

    print(f"  {len(tables)} tables, {len(refs)} refs, {len(figs)} figures")

    # === Build DOCX ===
    print("[2/3] Building document...")
    doc = Document()

    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(1.91)
    sec.right_margin = Cm(1.91)
    sec._sectPr.append(parse_xml(f'<w:cols {nsdecls("w")} w:num="2" w:space="480"/>'))

    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.0

    # Build \ref->roman numeral mapping
    # Parse tex to find all \section{...} and their following \label{...}
    label_to_roman = {}
    sec_count = 0
    for line in tex.split('\n'):
        s = line.strip()
        if re.match(r'\\section\{[^}]*\}', s) and not s.startswith('\\section*'):
            sec_count += 1
        lm = re.search(r'\\label\{([^}]*)\}', s)
        if lm:
            roman = ''
            if sec_count > 0:
                n = sec_count
                vals = [(1000,'M'),(900,'CM'),(500,'D'),(400,'CD'),(100,'C'),(90,'XC'),
                        (50,'L'),(40,'XL'),(10,'X'),(9,'IX'),(5,'V'),(4,'IV'),(1,'I')]
                for v, l in vals:
                    while n >= v:
                        roman += l
                        n -= v
            label_to_roman[lm.group(1)] = roman

    print(f"  {len(label_to_roman)} labels mapped to roman numerals")

    def P(text, size=10, bold=False, italic=False, align=None, sb=0, sa=3):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        if sb: p.paragraph_format.space_before = Pt(sb)
        if sa: p.paragraph_format.space_after = Pt(sa)
        if align: p.alignment = align
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        if bold: run.font.bold = True
        if italic: run.font.italic = True
        return p

    def add_table(caption, data):
        P(f'Table {caption}', 9, bold=True, align=_CENTER, sb=8, sa=3)
        nc = max(len(r) for r in data)
        tbl = doc.add_table(rows=len(data), cols=nc)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        borders = OxmlElement('w:tblBorders')
        for bn in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            b = OxmlElement(f'w:{bn}')
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), '4')
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), '000000')
            borders.append(b)
        tbl._tbl.tblPr.append(borders)
        for ri, row in enumerate(data):
            for ci in range(min(nc, len(row))):
                cell = tbl.rows[ri].cells[ci]
                cell.text = ''
                cp = cell.paragraphs[0]
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cp.add_run(row[ci])
                run.font.name = 'Times New Roman'
                run.font.size = Pt(8)
                if ri == 0: run.font.bold = True
        P('', 2, sa=4)

    def add_image(path, caption):
        found = find_image(path, fig_dir)
        if found:
            fp = doc.add_paragraph()
            fp.alignment = _CENTER
            fp.paragraph_format.space_before = Pt(6)
            fp.paragraph_format.space_after = Pt(2)
            fp.add_run().add_picture(found, width=Inches(3.3))
        if caption:
            P(f'Fig. {caption}', 8, italic=True, align=_CENTER, sb=2, sa=6)

    def add_eq(oi):
        if oi < len(omml_paras):
            # Insert after the last paragraph to maintain ordering
            if doc.paragraphs:
                doc.paragraphs[-1]._element.addnext(omml_paras[oi])
            else:
                doc.element.body.append(omml_paras[oi])
            return oi + 1
        return oi

    # ---- Header ----
    P(title, 14, bold=True, align=_CENTER, sa=4)
    P(author, 10, align=_CENTER, sa=2)
    P('State Grid Quanzhou Power Supply Company\nFengze, Quanzhou, Fujian, China\nEmail: huangchuzhi81@gmail.com',
      9, italic=True, align=_CENTER, sa=12)

    ap = doc.add_paragraph()
    ap.paragraph_format.space_before = Pt(6)
    ap.paragraph_format.space_after = Pt(3)
    r = ap.add_run('Abstract\u2014')
    r.font.size = Pt(9); r.font.bold = True; r.font.name = 'Times New Roman'
    r = ap.add_run(abstract)
    r.font.size = Pt(9); r.font.name = 'Times New Roman'

    kp = doc.add_paragraph()
    kp.paragraph_format.space_after = Pt(12)
    r = kp.add_run('Keywords\u2014')
    r.font.size = Pt(9); r.font.bold = True; r.font.name = 'Times New Roman'
    r = kp.add_run(keywords)
    r.font.size = Pt(9); r.font.name = 'Times New Roman'

    # ---- Body ----
    bc = body
    for env in ['abstract', 'IEEEkeywords']:
        bc = re.sub(rf'\\begin{{{env}}}.*?\\end{{{env}}}', '', bc, flags=re.DOTALL)
    bc = re.sub(r'\\maketitle', '', bc)
    bc = re.sub(r'\\label\{[^}]*\}', '', bc)

    oi, ti, fi = 0, 0, 0
    ie, ig, it, ib = False, False, False, False
    eb, tb = [], []

    def flush():
        nonlocal tb
        if not tb: return
        rw = re.sub(r'\s+', ' ', ' '.join(tb)).strip()
        tb = []
        if not rw: return
        if rw == '[htbp]' or 'htbp' in rw: return
        if rw.startswith('\\') and len(rw) < 40: return
        if re.match(r'^\{.*\}$', rw) and len(rw) < 10: return
        if '&' in rw or rw.startswith('\\title{') or rw.startswith('\\author{'): return
        # Filter LaTeX environment commands
        if re.match(r'\\\\begin\\{|\\\\end\\{|\\\\item', rw): return
        rw = re.sub(r'\\ref\{([^}]*)\}', lambda m: label_to_roman.get(m.group(1), m.group(1)), rw)
        rw = re.sub(r'\\cite\{([^}]*)\}', r'[\1]', rw)
        rw = clean(rw)
        if rw: P(rw, 9, sa=3)

    for line in bc.split('\n'):
        s = line.strip()
        if not s: flush(); continue
        m = re.match(r'\\(?:sub)*section\*?\{(.*?)\}', s)
        if m: flush(); P(clean(m.group(1)), 10, bold=True, sb=12, sa=6); continue
        if re.match(r'\\begin\{figure\}', s): flush(); ig = True; continue
        if ig and re.match(r'\\end\{figure\}', s):
            ig = False
            if fi < len(figs): add_image(figs[fi]['path'], figs[fi]['caption']); fi += 1
            continue
        if ig: continue
        if re.match(r'\\begin\{table\}', s): flush(); it = True; continue
        if it and re.match(r'\\end\{table\}', s):
            it = False
            if ti < len(tables): add_table(tables[ti]['caption'], tables[ti]['data']); ti += 1
            continue
        if it: continue
        if re.match(r'\\begin\{(equation|align|equation\*|align\*)\}', s):
            flush(); ie = True; eb = []; continue
        if ie:
            if re.search(r'\\end\{(equation|align|equation\*|align\*)\}', s):
                ie = False; oi = add_eq(oi); eb = []; continue
            eb.append(s); continue
        if re.match(r'\\begin\{thebibliography\}', s):
            flush(); ib = True; P('REFERENCES', 10, bold=True, sb=12, sa=6); continue
        if ib and re.match(r'\\end\{thebibliography\}', s): ib = False; continue
        if ib:
            m = re.search(r'\\bibitem\{(.*?)\}', s)
            if m:
                rt = clean(s[m.end():]).strip()
                if rt:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(1)
                    p.paragraph_format.left_indent = Cm(0.5)
                    p.paragraph_format.first_line_indent = Cm(-0.5)
                    p.paragraph_format.line_spacing = 1.0
                    r1 = p.add_run(f'[{m.group(1)}] ')
                    r1.font.name = 'Times New Roman'; r1.font.size = Pt(8); r1.font.bold = True
                    r2 = p.add_run(rt)
                    r2.font.name = 'Times New Roman'; r2.font.size = Pt(8)
                continue
            ct = clean(s)
            if ct and doc.paragraphs and doc.paragraphs[-1].text.strip().startswith('[b') and len(doc.paragraphs[-1].runs) >= 2:
                doc.paragraphs[-1].runs[-1].text += ' ' + ct
            continue
        tb.append(s)
    flush()

    # Indent body paragraphs
    for p in doc.paragraphs:
        t = (p.text or '').strip()
        if not t or p.alignment == _CENTER: continue
        if any(t.startswith(x) for x in ['Abstract', 'Keywords', 'Fig.', 'Table']): continue
        if t == 'REFERENCES' or t.startswith('[b'): continue
        pPr = p._element.find(f'{{{NS_W}}}pPr')
        if pPr is None: pPr = etree.SubElement(p._element, f'{{{NS_W}}}pPr')
        old = pPr.find(f'{{{NS_W}}}ind')
        if old is not None: pPr.remove(old)
        ind = etree.SubElement(pPr, f'{{{NS_W}}}ind')
        ind.set(f'{{{NS_W}}}firstLine', '360')

    doc.save(out_path)
    print(f"[3/3] Saved: {out_path} ({os.path.getsize(out_path)} bytes)")
    print("Done!")


if __name__ == '__main__':
    main()
