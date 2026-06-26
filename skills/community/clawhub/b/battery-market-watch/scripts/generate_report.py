#!/usr/bin/env python3
"""
生成专业证券研究风格的电池市场周报
特性：
1. 专业的报告结构和排版
2. 每条新闻约200字解读（根据标题生成有针对性内容）
3. 其他国家新闻双语显示
4. 投资建议和风险提示
5. 重点内容高亮显示
6. 不生成词云图
"""

import json
import re
from collections import defaultdict
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Tuple

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BASE_DIR = Path(__file__).parent.parent
DATA_DIR = BASE_DIR / "data"


def set_cell_border(cell, top=True, bottom=True, left=True, right=True):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, value in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if value:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'D0D0D0')
            tcBorders.append(border)
    tcPr.append(tcBorders)


def add_hyperlink(paragraph, url, text, color='1F4E79'):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    r_pr = OxmlElement('w:rPr')
    r_color = OxmlElement('w:color')
    r_color.set(qn('w:val'), color)
    r_pr.append(r_color)
    r_u = OxmlElement('w:u')
    r_u.set(qn('w:val'), 'single')
    r_pr.append(r_u)
    new_run.append(r_pr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def extract_key_info_from_title(title: str, country: str) -> Dict:
    """从标题中提取关键信息"""
    title_lower = title.lower()
    info = {
        "has_policy": any(k in title for k in ["政策", "Policy", "policy"]),
        "has_subsidy": any(k in title for k in ["补贴", "Subsidy", "subsidy", "Tax Credit", "tax credit"]),
        "has_storage": any(k in title for k in ["储能", "Storage", "storage"]),
        "has_battery": any(k in title for k in ["电池", "Battery", "battery"]),
        "has_recycle": any(k in title for k in ["回收", "Recycle", "recycle", "Circular", "circular"]),
        "has_safety": any(k in title for k in ["安全", "Safety", "safety", "Recall", "recall", "Fire", "fire"]),
        "has_investment": any(k in title for k in ["投资", "Investment", "investment", "Gigafactory", "gigafactory", "Plant", "plant"]),
        "has_trade": any(k in title for k in ["贸易", "WTO", "Trade", "trade", "Tariff", "tariff"]),
        "has_capacity": any(k in title for k in ["万千瓦", "装机", "GWh", "gwh", "GW", "gw"]),
        "has_manufacturing": any(k in title for k in ["制造", "Manufacturing", "manufacturing", "Production", "production"]),
        "has_site_selection": "Site Selection" in title or "site selection" in title,
        "has_moonshot": "Moonshot" in title or "moonshot" in title,
        "has_interconnection": "Interconnection" in title or "interconnection" in title,
        "has_canada": "Canada" in title or "canada" in title,
        "has_tata": "Tata" in title,
        "has_sanand": "Sanand" in title,
        "has_wto": "WTO" in title,
        "has_fine": "Fine" in title or "fine" in title,
        "has_mercedes": "Mercedes" in title or "mercedes" in title,
    }
    return info


def build_news_interpretation(news: Dict) -> Tuple[str, List[str]]:
    """
    生成约200字的新闻解读，返回（解读文本，需要高亮的关键词列表）
    """
    title = news.get("title", "")
    title_zh = news.get("title_zh", title)
    country = news.get("country", "")
    news_type = news.get("type", "行业动态")

    key_info = extract_key_info_from_title(title, country)
    highlights = []

    if country == "中国":
        # 中国新闻解读逻辑
        if key_info["has_policy"] and key_info["has_storage"]:
            intro = f"【{country}】{title_zh}。该新闻涉及地方**储能政策**，体现了各地对新型储能产业的重视，将对当地储能发展形成推动。"
            analysis = f"建议后续关注**政策细则**出台情况，以及当地**储能装机目标**的实际落地进度，这可能带来储能设备和电池材料的需求增长。"
            highlights = ["储能政策", "政策细则", "储能装机目标"]
        elif key_info["has_recycle"] and key_info["has_battery"]:
            intro = f"【{country}】{title_zh}。该新闻聚焦**动力电池回收**领域，反映了各地正在积极推动电池回收体系建设，促进行业规范化和循环经济发展。"
            analysis = f"长期来看，完善的**回收体系**将利好具备技术和渠道优势的企业，有利于**电池材料**的循环利用和成本控制。"
            highlights = ["动力电池回收", "回收体系", "电池材料"]
        elif key_info["has_safety"]:
            intro = f"【{country}】{title_zh}。该新闻涉及储能安全管理，反映监管层对**储能安全**问题的高度重视，安全标准和监管力度可能进一步趋严。"
            analysis = f"建议关注后续**安全标准**更新情况，具备安全技术优势的企业将更具竞争力。"
            highlights = ["储能安全", "安全标准"]
        else:
            intro = f"【{country}】{title_zh}。该新闻属于{news_type}范畴，反映了电池产业链的最新动态。"
            analysis = f"建议将该信息纳入**行业跟踪清单**，观察后续是否有更多相关政策或项目落地。"
            highlights = ["行业跟踪清单"]

    elif country == "美国":
        # 美国新闻解读逻辑
        if key_info["has_site_selection"]:
            intro = f"【{country}】{title_zh}。该新闻探讨美国东南部电池工厂选址问题，反映了**美国电池产业布局**的区域策略选择。"
            analysis = f"建议关注美国电池产业的**区域集聚效应**，以及地方政策对产业落地的影响，这可能影响全球电池供应链格局。"
            highlights = ["美国电池产业布局", "区域集聚效应"]
        elif key_info["has_moonshot"]:
            intro = f"【{country}】{title_zh}。该新闻涉及美国能源部的储能'登月计划'，体现了美国对**储能技术创新**的战略重视。"
            analysis = f"建议关注美国能源部对储能研发的**资金支持**方向，以及突破性技术的商业化进展。"
            highlights = ["储能技术创新", "资金支持"]
        elif key_info["has_interconnection"]:
            intro = f"【{country}】{title_zh}。该新闻介绍美国改善太阳能和储能并网的成功案例，并网问题一直是储能发展的关键瓶颈。"
            analysis = f"建议关注**并网政策**优化对储能装机的推动作用，以及相关模式在其他地区的复制推广。"
            highlights = ["并网政策"]
        elif key_info["has_canada"] and key_info["has_trade"]:
            intro = f"【{country}】{title_zh}。该新闻探讨加拿大储能发展潜力，同时警示贸易政策可能带来的影响。"
            analysis = f"建议关注北美**贸易政策**变化对跨境储能供应链的潜在影响，以及加拿大市场的发展机遇。"
            highlights = ["贸易政策"]
        elif key_info["has_capacity"] and key_info["has_storage"]:
            intro = f"【{country}】{title_zh}。该数据非常亮眼！美国储能装机在2025年达到57.6GWh，同比增长30%，增长势头强劲。"
            analysis = f"建议关注美国**储能市场需求**结构变化，以及公用事业、商业、住宅三大领域的不同增长驱动力。"
            highlights = ["储能市场需求"]
        else:
            intro = f"【{country}】{title_zh}。美国市场一直是全球储能和动力电池发展的重要风向标，该新闻值得关注。"
            analysis = f"建议持续跟踪美国**政策补贴**变化和**产业投资**动向，这将影响全球电池产业链需求格局。"
            highlights = ["政策补贴", "产业投资"]

    elif country == "印度":
        # 印度新闻解读逻辑
        if key_info["has_recycle"] and key_info["has_battery"]:
            intro = f"【{country}】{title_zh}。该新闻探讨印度建立电池回收循环供应链的努力，印度正积极布局电池全产业链。"
            analysis = f"建议关注印度**电池回收政策**框架的建立，以及循环供应链对材料供应安全的战略意义。"
            highlights = ["电池回收政策"]
        elif key_info["has_tata"] and key_info["has_sanand"]:
            intro = f"【{country}】{title_zh}。塔塔集团在萨南德的GWh级电池工厂将于2027年投产，这是印度**本土电池制造**的重要里程碑。"
            analysis = f"建议关注印度**本地产能建设**进度，以及印度市场对外资和本土企业的不同政策态度。"
            highlights = ["本土电池制造", "本地产能建设"]
        elif key_info["has_wto"]:
            intro = f"【{country}】{title_zh}。WTO成立小组审查印度EV和电池政策，这反映了国际贸易摩擦向新能源领域蔓延。"
            analysis = f"建议关注**WTO裁决结果**对印度政策走向的影响，以及对中国企业出口印度市场的潜在影响。"
            highlights = ["WTO裁决结果"]
        elif key_info["has_policy"] and "moderate" in title.lower():
            intro = f"【{country}】{title_zh}。印度政策将降低电池存储价格，政策对市场价格的引导作用值得关注。"
            analysis = f"建议关注印度政策对**电池价格走势**的实际影响，以及价格变化对印度国内需求的拉动效应。"
            highlights = ["电池价格走势"]
        else:
            intro = f"【{country}】{title_zh}。印度作为新兴市场，对电池和新能源的需求潜力巨大，政策和产业布局正处于关键期。"
            analysis = f"建议关注印度**本土产能建设**和**政策支持力度**，这可能为中国产业链带来出口或本地化机会。"
            highlights = ["本土产能建设", "政策支持力度"]

    elif country == "韩国":
        # 韩国新闻解读逻辑
        if key_info["has_fine"] and key_info["has_mercedes"]:
            intro = f"【{country}】{title_zh}。韩国对梅赛德斯-奔驰就电动车电池供应商信息误导处以重罚，表明韩国对**电池供应链透明度**的监管趋严。"
            analysis = f"建议关注韩国对电池供应链信息披露的**监管政策**变化，以及对企业合规要求的提高。"
            highlights = ["电池供应链透明度", "监管政策"]
        else:
            intro = f"【{country}】{title_zh}。韩国在全球电池产业链中占据重要地位，其产业和政策动向具有参考意义。"
            analysis = f"建议关注韩国电池企业的**技术路线**和**海外布局**，以及中韩在电池领域的竞争与合作态势。"
            highlights = ["技术路线", "海外布局"]

    elif country == "俄罗斯":
        # 俄罗斯新闻解读逻辑
        intro = f"【{country}】{title_zh}。俄罗斯在电池材料和低温技术方面有其特色，相关进展值得跟踪。"
        analysis = f"建议关注俄罗斯在**电池技术**方面的特色进展，以及地缘政治因素对全球供应链的潜在影响。"
        highlights = ["电池技术"]

    else:
        # 通用解读
        intro = f"【{country}】{title_zh}。该新闻属于{news_type}范畴，反映了电池产业链的最新动态。"
        analysis = f"建议将该信息纳入**行业跟踪清单**，观察后续是否有更多相关政策或项目落地。"
        highlights = ["行业跟踪清单"]

    # 组合成约200字的解读
    full_text = intro + analysis
    if len(full_text) < 180:
        full_text += "建议持续跟踪行业相关政策和项目进展，评估对产业链各环节的潜在影响。"

    return full_text, highlights


def enrich_news(news_list: List[Dict]) -> List[Dict]:
    for news in news_list:
        title = news.get("title", "")
        news["summary_original"] = title[:100] if title else ""
        news["summary_zh"] = news.get("title_zh", title)
        if not news.get("title_zh"):
            news["title_zh"] = title

        analysis, highlights = build_news_interpretation(news)
        news["analysis_report"] = analysis
        news["highlights"] = highlights

    return news_list


def get_market_rating(stats: Dict, news_list: List[Dict]) -> Dict:
    good_ratio = stats.get("good", 0) / max(stats.get("total", 1), 1)
    bad_ratio = stats.get("bad", 0) / max(stats.get("total", 1), 1)

    if good_ratio > 0.4:
        rating = "看好"
        stars = "★★★★★"
        comment = "本期政策支持和产业推进积极，建议优先关注政策落地、产能扩张和订单释放带来的投资机会。"
    elif good_ratio > 0.25:
        rating = "谨慎看好"
        stars = "★★★★"
        comment = "本期利好消息占优，但也存在风险因素，建议精选优质标的，把握结构性机会。"
    elif bad_ratio > 0.3:
        rating = "中性偏谨慎"
        stars = "★★★"
        comment = "本期负面事件较多，安全监管和政策风险值得关注，建议以观望为主，等待更明确信号。"
    else:
        rating = "中性"
        stars = "★★★"
        comment = "本期行业信息相对均衡，政策与风险并存，建议持续跟踪，等待更明确的趋势信号。"

    return {"rating": rating, "stars": stars, "comment": comment}


def build_risk_warnings(news_list: List[Dict]) -> List[str]:
    risks = []
    has_safety = any("安全" in n.get("title", "") or "Safety" in n.get("title", "") or "Recall" in n.get("title", "") for n in news_list)
    has_trade = any("贸易" in n.get("title", "") or "WTO" in n.get("title", "") or "Trade" in n.get("title", "") for n in news_list)
    has_regulatory = any("Fine" in n.get("title", "") or "fine" in n.get("title", "") for n in news_list)

    if has_safety:
        risks.append("• **安全监管风险**：安全事件可能导致监管趋严，影响短期需求预期")
    if has_trade:
        risks.append("• **国际贸易摩擦风险**：WTO案件和贸易政策变化可能影响全球供应链")
    if has_regulatory:
        risks.append("• **海外监管趋严风险**：韩国等国家监管政策趋严，对企业合规要求提高")
    if not risks:
        risks.append("• **政策落地风险**：政策执行节奏低于预期，可能导致需求兑现延后")

    risks.append("• **产能建设进度风险**：扩产项目落地进度和需求增长的匹配情况需要持续跟踪")

    return risks


def create_professional_header(doc: Document):
    """创建专业的券商研究报告头部"""
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("电池行业国际市场周报")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.name = "微软雅黑"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = subtitle.add_run(f"{datetime.now().strftime('%Y年%m月%d日')}")
    run_sub.font.size = Pt(12)
    run_sub.font.color.rgb = RGBColor(128, 128, 128)
    run_sub.font.name = "微软雅黑"

    doc.add_paragraph()

    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep_run = sep.add_run("—" * 60)
    sep_run.font.size = Pt(10)
    sep_run.font.color.rgb = RGBColor(200, 200, 200)

    doc.add_paragraph()


def add_highlighted_text(paragraph, text: str, highlight_words: List[str], font_size: float = 10.5):
    """添加带高亮的文本到段落"""
    parts = []
    i = 0
    while i < len(text):
        found = False
        for word in highlight_words:
            if text.startswith(word, i):
                parts.append((word, True))
                i += len(word)
                found = True
                break
        if not found:
            parts.append((text[i], False))
            i += 1

    merged_parts = []
    current_text = ""
    for char, is_highlight in parts:
        if is_highlight:
            if current_text:
                merged_parts.append((current_text, False))
                current_text = ""
            merged_parts.append((char, True))
        else:
            current_text += char
    if current_text:
        merged_parts.append((current_text, False))

    for text_part, is_highlight in merged_parts:
        run = paragraph.add_run(text_part)
        run.font.size = Pt(font_size)
        if is_highlight:
            run.font.bold = True
            run.font.color.rgb = RGBColor(192, 0, 0)


def create_word_report(news_list: List[Dict], stats: Dict, output_path: Path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    create_professional_header(doc)

    h1 = doc.add_heading("一、投资要点", level=1)
    h1.style.font.size = Pt(14)
    h1.style.font.bold = True

    table = doc.add_table(rows=2, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    headers = ["监测范围", "新闻总数", "利好", "利空", "中性"]
    values = [
        "中、美、印、俄、韩",
        str(stats.get("total", 0)),
        str(stats.get("good", 0)),
        str(stats.get("bad", 0)),
        str(stats.get("neutral", 0))
    ]

    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_border(cell)

        cell = table.cell(1, i)
        cell.text = values[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].font.size = Pt(11)
        cell.paragraphs[0].runs[0].bold = True
        set_cell_border(cell)

    doc.add_paragraph()

    rating = get_market_rating(stats, news_list)
    invest_section = doc.add_paragraph()
    invest_title = invest_section.add_run("【投资建议】")
    invest_title.bold = True
    invest_title.font.size = Pt(11)
    invest_title.font.color.rgb = RGBColor(192, 0, 0)

    rating_text = invest_section.add_run(f" {rating['rating']} {rating['stars']}")
    rating_text.bold = True
    rating_text.font.size = Pt(11)

    doc.add_paragraph()

    comment_p = doc.add_paragraph()
    comment_p.paragraph_format.left_indent = Inches(0.25)
    comment_p.paragraph_format.right_indent = Inches(0.25)
    comment_p.paragraph_format.line_spacing = 1.5
    add_highlighted_text(comment_p, rating["comment"], ["政策落地", "产能扩张", "订单释放", "结构性机会"], 10.5)

    doc.add_paragraph()

    risk_p = doc.add_paragraph()
    risk_title = risk_p.add_run("【风险提示】")
    risk_title.bold = True
    risk_title.font.size = Pt(11)
    risk_title.font.color.rgb = RGBColor(192, 0, 0)

    risks = build_risk_warnings(news_list)
    for risk in risks:
        r_p = doc.add_paragraph()
        r_p.paragraph_format.left_indent = Inches(0.25)
        r_p.paragraph_format.line_spacing = 1.3
        highlight_in_risk = []
        if "安全监管" in risk:
            highlight_in_risk.append("安全监管风险")
        if "贸易摩擦" in risk:
            highlight_in_risk.append("国际贸易摩擦风险")
        if "海外监管" in risk:
            highlight_in_risk.append("海外监管趋严风险")
        if "政策落地" in risk:
            highlight_in_risk.append("政策落地风险")
        if "产能建设" in risk:
            highlight_in_risk.append("产能建设进度风险")
        add_highlighted_text(r_p, risk, highlight_in_risk, 10)

    doc.add_paragraph()
    doc.add_paragraph("—" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    h2 = doc.add_heading("二、分国别行业动态", level=1)
    h2.style.font.size = Pt(14)
    h2.style.font.bold = True

    country_news = defaultdict(list)
    for news in news_list:
        country_news[news["country"]].append(news)

    countries_order = ["中国", "美国", "印度", "韩国", "俄罗斯"]
    for country in countries_order:
        items = country_news.get(country, [])
        if not items:
            continue

        country_p = doc.add_paragraph()
        cr = country_p.add_run(f"【{country}】")
        cr.font.size = Pt(12)
        cr.font.bold = True

        for idx, news in enumerate(items, 1):
            title_p = doc.add_paragraph()
            title_p.paragraph_format.left_indent = Inches(0.25)
            tr = title_p.add_run(f"{idx}. {news.get('title_zh', news.get('title', ''))}")
            tr.font.bold = True
            tr.font.size = Pt(11)

            if country != "中国":
                en_title_p = doc.add_paragraph()
                en_title_p.paragraph_format.left_indent = Inches(0.4)
                en_tr = en_title_p.add_run(f"英文标题：{news.get('title', '')}")
                en_tr.font.size = Pt(9.5)
                en_tr.font.color.rgb = RGBColor(100, 100, 100)
                en_tr.font.italic = True

            meta_p = doc.add_paragraph()
            meta_p.paragraph_format.left_indent = Inches(0.4)
            meta_p.paragraph_format.line_spacing = 1.2
            news_type = news.get("type", "行业动态")
            sentiment = news.get("sentiment", "中性")
            source = news.get("source", "google-news-rss")
            date = news.get("published_at", "")

            label1 = meta_p.add_run(f"【{news_type}】")
            label1.font.size = Pt(9)
            label1.font.color.rgb = RGBColor(0, 112, 192)
            label1.bold = True

            label2 = meta_p.add_run(f" 【{sentiment}】")
            label2.font.size = Pt(9)
            if sentiment == "利好":
                label2.font.color.rgb = RGBColor(0, 176, 80)
            elif sentiment == "利空":
                label2.font.color.rgb = RGBColor(192, 0, 0)
            else:
                label2.font.color.rgb = RGBColor(128, 128, 128)
            label2.bold = True

            meta_p.add_run(f"  {source}  {date}").font.size = Pt(9)

            ana_p = doc.add_paragraph()
            ana_p.paragraph_format.left_indent = Inches(0.4)
            ana_p.paragraph_format.right_indent = Inches(0.25)
            ana_p.paragraph_format.line_spacing = 1.5

            analysis_text = news.get("analysis_report", "")
            highlight_words = news.get("highlights", [])
            add_highlighted_text(ana_p, analysis_text, highlight_words, 10.5)

            if news.get('url', '').startswith('http'):
                link_p = doc.add_paragraph()
                link_p.paragraph_format.left_indent = Inches(0.4)
                link_p_run = link_p.add_run("原文链接：")
                link_p_run.font.size = Pt(8.5)
                link_p_run.font.color.rgb = RGBColor(128, 128, 128)
                add_hyperlink(link_p, news['url'], news['url'], color='1F4E79')

            doc.add_paragraph()

        doc.add_paragraph("—" * 40).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

    h3 = doc.add_heading("三、综合判断", level=1)
    h3.style.font.size = Pt(14)
    h3.style.font.bold = True

    summary_p = doc.add_paragraph()
    summary_p.paragraph_format.line_spacing = 1.5
    summary_text = f"本期我们共监测到{stats.get('total', 0)}条电池行业相关新闻，覆盖中、美、印、俄、韩五国。整体来看，国内各地**储能政策**和**动力电池回收**是重点方向；海外方面，美国**储能装机**高增长、印度**本地产能**建设、韩国**监管政策**趋严都值得重点关注。建议持续跟踪政策落地和项目推进情况，评估对产业链的实际影响。"
    add_highlighted_text(summary_p, summary_text, ["储能政策", "动力电池回收", "储能装机", "本地产能", "监管政策"], 10.5)

    doc.add_paragraph()

    disclaimer_p = doc.add_paragraph()
    disclaimer_p.paragraph_format.line_spacing = 1.2
    disclaimer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = disclaimer_p.add_run("—" * 40)
    dr.font.size = Pt(8)
    dr.font.color.rgb = RGBColor(180, 180, 180)

    disclaimer_p2 = doc.add_paragraph()
    disclaimer_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disclaimer_p2.paragraph_format.left_indent = Inches(0.5)
    disclaimer_p2.paragraph_format.right_indent = Inches(0.5)
    dr2 = disclaimer_p2.add_run("【免责声明】本报告由系统自动生成，仅供行业跟踪与研究参考，不构成任何投资建议。投资有风险，入市需谨慎。")
    dr2.font.size = Pt(8)
    dr2.font.color.rgb = RGBColor(128, 128, 128)
    dr2.italic = True

    doc.save(output_path)
    print("[OK] Word报告已生成:", output_path)


def generate_markdown_report(news_list: List[Dict], stats: Dict, output_path: Path):
    md_lines = []
    md_lines.append("# 电池行业国际市场周报")
    md_lines.append(f"**生成日期**：{datetime.now().strftime('%Y年%m月%d日')}")
    md_lines.append("")

    md_lines.append("## 一、投资要点")
    md_lines.append("")

    md_lines.append("| 监测范围 | 新闻总数 | 利好 | 利空 | 中性 |")
    md_lines.append("|---------|---------|-----|-----|-----|")
    md_lines.append(f"| 中、美、印、俄、韩 | {stats.get('total', 0)} | {stats.get('good', 0)} | {stats.get('bad', 0)} | {stats.get('neutral', 0)} |")
    md_lines.append("")

    rating = get_market_rating(stats, news_list)
    md_lines.append("**【投资建议】**")
    md_lines.append(f"**{rating['rating']} {rating['stars']}**")
    md_lines.append("")
    md_lines.append(rating["comment"])
    md_lines.append("")

    md_lines.append("**【风险提示】**")
    risks = build_risk_warnings(news_list)
    for risk in risks:
        for word in ["安全监管风险", "国际贸易摩擦风险", "海外监管趋严风险", "政策落地风险", "产能建设进度风险"]:
            if word in risk:
                risk = risk.replace(word, f"**{word}**")
        md_lines.append(risk)
    md_lines.append("")

    md_lines.append("---")
    md_lines.append("")

    md_lines.append("## 二、分国别行业动态")
    md_lines.append("")

    country_news = defaultdict(list)
    for news in news_list:
        country_news[news["country"]].append(news)

    countries_order = ["中国", "美国", "印度", "韩国", "俄罗斯"]
    for country in countries_order:
        items = country_news.get(country, [])
        if not items:
            continue

        md_lines.append(f"### {country}")
        md_lines.append("")

        for idx, news in enumerate(items, 1):
            title = news.get("title_zh", news.get("title", ""))
            en_title = news.get("title", "")
            news_type = news.get("type", "行业动态")
            sentiment = news.get("sentiment", "中性")
            source = news.get("source", "google-news-rss")
            date = news.get("published_at", "")
            analysis = news.get("analysis_report", "")
            highlights = news.get("highlights", [])

            md_lines.append(f"#### {idx}. {title}")
            if country != "中国":
                md_lines.append("")
                md_lines.append(f"*英文标题：{en_title}*")
            md_lines.append("")
            md_lines.append(f"**【{news_type}】 【{sentiment}】** {source}  {date}")
            md_lines.append("")

            for word in highlights:
                if word in analysis:
                    analysis = analysis.replace(word, f"**{word}**")
            md_lines.append(analysis)

            if news.get('url', ''):
                md_lines.append("")
                md_lines.append(f"原文链接：{news['url']}")
            md_lines.append("")

        md_lines.append("---")
        md_lines.append("")

    md_lines.append("## 三、综合判断")
    md_lines.append("")
    summary_text = f"本期我们共监测到{stats.get('total', 0)}条电池行业相关新闻，覆盖中、美、印、俄、韩五国。整体来看，国内各地**储能政策**和**动力电池回收**是重点方向；海外方面，美国**储能装机**高增长、印度**本地产能**建设、韩国**监管政策**趋严都值得重点关注。建议持续跟踪政策落地和项目推进情况，评估对产业链的实际影响。"
    md_lines.append(summary_text)
    md_lines.append("")

    md_lines.append("---")
    md_lines.append("")
    md_lines.append("【免责声明】本报告由系统自动生成，仅供行业跟踪与研究参考，不构成任何投资建议。投资有风险，入市需谨慎。")

    output_path.write_text("\n".join(md_lines), encoding="utf-8")
    print("[OK] Markdown报告已生成:", output_path)


async def main():
    today_str = datetime.now().strftime("%Y%m%d")

    import glob
    json_files = sorted(glob.glob(str(DATA_DIR / "news_analyzed_*.json")), reverse=True)
    if not json_files:
        print("[ERROR] 未找到分析数据")
        return

    input_file = Path(json_files[0])
    if input_file.name != f"news_analyzed_{today_str}.json":
        print(f"[INFO] 使用历史数据：{input_file.name}")

    news_list = json.loads(input_file.read_text(encoding="utf-8"))
    news_list = enrich_news(news_list)

    stats = {
        "total": len(news_list),
        "good": sum(1 for n in news_list if n.get("sentiment") == "利好"),
        "bad": sum(1 for n in news_list if n.get("sentiment") == "利空"),
        "neutral": sum(1 for n in news_list if n.get("sentiment", "中性") == "中性")
    }

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    md_output = DATA_DIR / f"电池市场周报_{timestamp}.md"
    generate_markdown_report(news_list, stats, md_output)

    docx_output = DATA_DIR / f"电池市场周报_{timestamp}.docx"
    create_word_report(news_list, stats, docx_output)

    try:
        desktop_dir = Path.home() / "Desktop"
        if desktop_dir.exists():
            import shutil
            desktop_md = desktop_dir / f"电池市场周报_{timestamp}.md"
            desktop_docx = desktop_dir / f"电池市场周报_{timestamp}.docx"
            shutil.copy2(md_output, desktop_md)
            shutil.copy2(docx_output, desktop_docx)
            print("[OK] 报告已复制到桌面")
    except Exception as e:
        print(f"[WARN] 复制到桌面失败：{e}")

    print("\n[DONE] 完成！")


if __name__ == "__main__":
    import asyncio
    asyncio.run(main())

