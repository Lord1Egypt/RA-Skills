#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DOCX Compare / Duplicate Check v1.0.2
比较两个 Word 文档的重复文本和重复图片。

v1.0.1 — 修复图片批注注入：双 XML 声明、新条目遗漏写入 ZIP
v1.0.0 — 初始版本

功能:
  1. 精确匹配：标准化后完全相同的段落
  2. 模糊匹配：difflib 相似度 ≥ 阈值（默认 95%）
  3. 图片查重：MD5 哈希完全一致，在重复图片上追加 **Word 批注**

输出:
  - *_标记重复.docx      → 重复文字标黄色
  - *_标记图片重复.docx   → 重复图片附加 Word 批注（右侧气泡）
  - *_vs_*_重复报告.txt  → 详细文字报告

用法:
  python compare_docx.py --file1 a.docx --file2 b.docx
  python compare_docx.py --file1 a.docx --file2 b.docx --output ./out --mode image
"""

import argparse
import difflib
import hashlib
import os
import re
import sys
import shutil
from copy import deepcopy
from io import BytesIO
from zipfile import ZipFile
from lxml import etree

# ── 依赖检查 ──────────────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit("❌ 缺少 python-docx 库，请运行: pip install python-docx")


# ── 命名空间 ──────────────────────────────────────────────────────
W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
R = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
A = 'http://schemas.openxmlformats.org/drawingml/2006/main'
WP = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
PIC = 'http://schemas.openxmlformats.org/drawingml/2006/picture'
CT_RELS = 'http://schemas.openxmlformats.org/package/2006/relationships'
CT_WORDPROC = 'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml'
REL_COMMENTS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments'

NSMAP = {
    'w': W, 'r': R, 'a': A, 'wp': WP, 'pic': PIC,
}
NSMAP_COMMENT = {
    'w': W, 'r': R,
}


# ── 工具函数 ──────────────────────────────────────────────────────

def normalize_text(text: str) -> str:
    text = re.sub(r'[\s\u3000\xa0]+', ' ', text)
    # 清除标点符号（用字符替换避免逃逸序列问题）
    for ch in '，。！？、；：""''《》（）【】{}.,!?;:()"\'<>':
        text = text.replace(ch, '')
    return text.strip().lower()


def md5_of_bytes(data: bytes) -> str:
    return hashlib.md5(data).hexdigest()


def extract_images_from_docx(docx_path: str) -> dict[str, list]:
    """
    提取所有内嵌图片。
    返回 {md5: [{"path": "word/media/...", "size": int, "index": int}, ...]}
    """
    images = {}
    try:
        with ZipFile(docx_path, 'r') as z:
            media_files = [n for n in z.namelist() if n.startswith('word/media/')]
            for idx, mf in enumerate(media_files):
                data = z.read(mf)
                h = md5_of_bytes(data)
                images.setdefault(h, []).append({
                    "path": mf, "size": len(data), "index": idx,
                })
    except Exception as e:
        print(f"  ⚠ 读取图片失败 ({docx_path}): {e}")
    return images


def _parse_rels(rels_bytes: bytes) -> dict[str, str]:
    """解析 .rels 文件，返回 {rId: target}"""
    root = etree.fromstring(rels_bytes)
    mapping = {}
    for rel in root:
        rid = rel.get('Id')
        target = rel.get('Target', '')
        mapping[rid] = target
    return mapping


def _make_comment_element(comment_id: int,
                          text: str = "🔴 此图片在另一文档中存在重复"):
    """构造 w:comment XML 元素"""
    c = etree.Element(f'{{{W}}}comment', nsmap=NSMAP_COMMENT)
    c.set(f'{{{W}}}id', str(comment_id))
    c.set(f'{{{W}}}author', '查重工具')
    c.set(f'{{{W}}}date', '2024-01-01T00:00:00Z')
    c.set(f'{{{W}}}initials', 'CT')

    p = etree.SubElement(c, f'{{{W}}}p')

    # run with commentReference
    r1 = etree.SubElement(p, f'{{{W}}}r')
    rPr1 = etree.SubElement(r1, f'{{{W}}}rPr')
    rStyle1 = etree.SubElement(rPr1, f'{{{W}}}rStyle')
    rStyle1.set(f'{{{W}}}val', 'CommentReference')
    etree.SubElement(r1, f'{{{W}}}commentReference')

    # run with text
    r2 = etree.SubElement(p, f'{{{W}}}r')
    rPr2 = etree.SubElement(r2, f'{{{W}}}rPr')
    rStyle2 = etree.SubElement(rPr2, f'{{{W}}}rStyle')
    rStyle2.set(f'{{{W}}}val', 'CommentText')
    t = etree.SubElement(r2, f'{{{W}}}t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text

    return c


# ── 核心：对重复图片加批注 ───────────────────────────────────────

def mark_duplicate_images_with_comments(docx_path: str,
                                        duplicate_md5s: set[str],
                                        output_path: str) -> bool:
    """
    遍历 docx 中的 inline/anchor 图片，将重复图片的所在 run 包裹上
    commentRangeStart / commentRangeEnd / commentReference 并写入 comments.xml。

    用 lxml 处理，可以处理包含中文路径的 ZIP 条目。
    """
    try:
        # ── 1. 计算所有图片的 MD5 ──
        all_images_md5: dict[str, str] = {}
        with ZipFile(docx_path, 'r') as z:
            media_files = [n for n in z.namelist() if n.startswith('word/media/')]
            for mf in media_files:
                all_images_md5[mf] = md5_of_bytes(z.read(mf))

        dup_media = {p for p, h in all_images_md5.items() if h in duplicate_md5s}
        if not dup_media:
            shutil.copy2(docx_path, output_path)
            print(f"  ℹ 文档中未发现重复图片，直接复制原文件")
            return True

        # ── 2. 读取 ZIP 全部内容 ──
        zip_data: dict[str, bytes] = {}
        with ZipFile(docx_path, 'r') as z:
            for name in z.namelist():
                zip_data[name] = z.read(name)

        # ── 3. 解析关系，建立 rId → media_path 映射 ──
        rels_path = 'word/_rels/document.xml.rels'
        if rels_path not in zip_data:
            print(f"  ⚠ 未找到 {rels_path}")
            shutil.copy2(docx_path, output_path)
            return True

        rid_to_media = {}
        rels_root = etree.fromstring(zip_data[rels_path])
        for rel in rels_root:
            rid = rel.get('Id')
            target = rel.get('Target', '')
            if target.startswith('media/'):
                rid_to_media[rid] = f'word/{target}'

        # ── 4. 解析 document.xml，找出重复图片所在的 run ──
        doc_root = etree.fromstring(zip_data['word/document.xml'])
        # 注册命名空间前缀，使序列化保留前缀
        for prefix, uri in NSMAP.items():
            etree.register_namespace(prefix, uri)

        # 找到所有包含重复图片的 run 元素 (去重)
        target_run_set = set()
        for p_elem in doc_root.iter(f'{{{W}}}p'):
            for r_elem in p_elem.iter(f'{{{W}}}r'):
                for blip in r_elem.iter(f'{{{A}}}blip'):
                    embed = blip.get(f'{{{R}}}embed')
                    if embed and embed in rid_to_media:
                        mp = rid_to_media[embed]
                        if mp in dup_media:
                            target_run_set.add((p_elem, r_elem))

        if not target_run_set:
            shutil.copy2(docx_path, output_path)
            print(f"  ℹ 文档中未找到重复图片的 inline 引用")
            return True

        # ── 5. 对每个目标 run 添加批注标记 ──
        # 需要确保 comment id 不冲突：检查现有 comments.xml
        existing_comments_count = 0
        if 'word/comments.xml' in zip_data:
            comments_root = etree.fromstring(zip_data['word/comments.xml'])
            existing_comments = comments_root.findall(f'{{{W}}}comment')
            existing_comments_count = len(existing_comments)

        comment_counter = [existing_comments_count]  # mutable for closure

        # 先收集所有要插入的 comment id，再统一构建 comments.xml
        comment_markers = []  # list of (para_elem, run_elem, comment_id)

        for p_elem, r_elem in target_run_set:
            cid = comment_counter[0]
            comment_counter[0] += 1

            # 创建 commentRangeStart
            crs = etree.Element(f'{{{W}}}commentRangeStart')
            crs.set(f'{{{W}}}id', str(cid))

            # 创建 commentRangeEnd
            cre = etree.Element(f'{{{W}}}commentRangeEnd')
            cre.set(f'{{{W}}}id', str(cid))

            # 创建 commentReference run
            ref_r = etree.Element(f'{{{W}}}r')
            ref_rPr = etree.SubElement(ref_r, f'{{{W}}}rPr')
            ref_rStyle = etree.SubElement(ref_rPr, f'{{{W}}}rStyle')
            ref_rStyle.set(f'{{{W}}}val', 'CommentReference')
            etree.SubElement(ref_r, f'{{{W}}}commentReference').set(
                f'{{{W}}}id', str(cid))

            comment_markers.append((p_elem, r_elem, cid, crs, cre, ref_r))

        # 从后往前插入（保持索引正确）
        # 在每个 run 前插入 commentRangeStart，run 后插入 commentRangeEnd + commentReference
        for p_elem, r_elem, cid, crs, cre, ref_r in reversed(comment_markers):
            # run 后插入 commentRangeEnd
            r_elem.addnext(cre)
            # 再后面插入 commentReference run
            cre.addnext(ref_r)
            # run 前插入 commentRangeStart
            r_elem.addprevious(crs)

        # ── 6. 构建/更新 comments.xml ──
        if 'word/comments.xml' in zip_data:
            # 追加到现有注释
            comments_root = etree.fromstring(zip_data['word/comments.xml'])
        else:
            comments_root = etree.Element(f'{{{W}}}comments', nsmap=NSMAP_COMMENT)

        for p_elem, r_elem, cid, _, _, _ in comment_markers:
            c_elem = _make_comment_element(cid)
            comments_root.append(c_elem)

        # ── 7. 序列化更新后的 XML ──
        # document.xml
        doc_decl = b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        # ⚠️ etree.tostring(xml_declaration=False, standalone=True) still outputs
        # the declaration with single quotes if standalone is set.
        # Don't give standalone=True here — we manually prepend the declaration.
        zip_data['word/document.xml'] = doc_decl + etree.tostring(
            doc_root, xml_declaration=False, encoding='UTF-8')

        # comments.xml
        zip_data['word/comments.xml'] = doc_decl + etree.tostring(
            comments_root, xml_declaration=False, encoding='UTF-8')

        # ── 8. 更新关系文件（添加 comments 关系）──
        # 检查是否已有 comments 关系
        has_comments_rel = False
        for rel in rels_root:
            t = rel.get('Type', '')
            if t == REL_COMMENTS:
                has_comments_rel = True
                break

        if not has_comments_rel:
            # 找个可用的 rId
            max_id = 0
            for rel in rels_root:
                rid = rel.get('Id', '')
                if rid.startswith('rId'):
                    try:
                        num = int(rid[3:])
                        max_id = max(max_id, num)
                    except ValueError:
                        pass
            new_rid = f'rId{max_id + 1}'
            new_rel = etree.SubElement(rels_root, f'{{{CT_RELS}}}Relationship')
            new_rel.set('Id', new_rid)
            new_rel.set('Type', REL_COMMENTS)
            new_rel.set('Target', 'comments.xml')
            zip_data[rels_path] = doc_decl + etree.tostring(
                rels_root, xml_declaration=False, encoding='UTF-8')

        # ── 9. 更新 [Content_Types].xml ──
        ct_path = '[Content_Types].xml'
        if ct_path in zip_data and not has_comments_rel:
            ct_root = etree.fromstring(zip_data[ct_path])
            ns_ct = 'http://schemas.openxmlformats.org/package/2006/content-types'
            # 检查是否已存在
            has_ct = False
            for override in ct_root.iter(f'{{{ns_ct}}}Override'):
                if override.get('PartName', '') == '/word/comments.xml':
                    has_ct = True
                    break
            if not has_ct:
                override = etree.SubElement(
                    ct_root, f'{{{ns_ct}}}Override')
                override.set('PartName', '/word/comments.xml')
                override.set('ContentType', CT_WORDPROC)
                zip_data[ct_path] = doc_decl + etree.tostring(
                    ct_root, xml_declaration=False, encoding='UTF-8')

        # ── 10. 写出新 ZIP ──
        with ZipFile(output_path, 'w') as zout:
            # 保证不改变压缩方式
            with ZipFile(docx_path, 'r') as zin:
                original_names = set(zin.namelist())
                for name in zin.namelist():
                    info = zin.getinfo(name)
                    if name in zip_data:
                        zout.writestr(info, zip_data[name])
                    else:
                        zout.writestr(info, zin.read(name))
            # 再追加 zip_data 中新增的条目（如 word/comments.xml）
            for name in zip_data:
                if name not in original_names:
                    zout.writestr(name, zip_data[name])

        print(f"  ✅ 已为 {len(comment_markers)} 处重复图片添加 Word 批注")
        return True

    except Exception as e:
        import traceback
        print(f"  ⚠ 图片批注处理失败: {e}")
        traceback.print_exc()
        # fallback: 复制原文件
        try:
            shutil.copy2(docx_path, output_path)
        except Exception:
            pass
        return False


# ── 段落提取 ──────────────────────────────────────────────────────

def extract_paragraphs(doc: Document) -> list[dict]:
    paras = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text:
            paras.append({
                "index": i, "text": text,
                "normalized": normalize_text(text),
                "paragraph": p,
            })
    return paras


def extract_tables_text(doc: Document) -> list[str]:
    texts = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    texts.append(t)
    return texts


# ── 精确匹配 ──────────────────────────────────────────────────────

def find_exact_duplicates(paras1: list[dict],
                          paras2: list[dict],
                          min_len: int = 3) -> list[tuple]:
    results = []
    norm2_texts = [(p2["normalized"], p2) for p2 in paras2]
    for p1 in paras1:
        norm1 = p1["normalized"]
        if len(norm1) < min_len:
            continue
        for n2, p2 in norm2_texts:
            if norm1 == n2:
                results.append((p1, p2, "exact"))
    return results


# ── 模糊匹配 ──────────────────────────────────────────────────────

def find_fuzzy_duplicates(paras1: list[dict], paras2: list[dict],
                          cutoff: float = 0.95) -> list[tuple]:
    results = []
    for p1 in paras1:
        best_ratio = 0.0
        best_p2 = None
        for p2 in paras2:
            ratio = difflib.SequenceMatcher(
                None, p1["normalized"], p2["normalized"]
            ).ratio()
            if ratio > best_ratio:
                best_ratio = ratio
                best_p2 = p2
        if best_ratio >= cutoff and best_p2 is not None:
            results.append((p1, best_p2, round(best_ratio, 4)))
    return results


# ── 图片查重 ──────────────────────────────────────────────────────

def find_duplicate_images(imgs1: dict,
                          imgs2: dict) -> list[dict]:
    common_md5 = set(imgs1.keys()) & set(imgs2.keys())
    results = []
    for m5 in common_md5:
        results.append({
            "md5": m5,
            "in_file1": imgs1[m5],
            "in_file2": imgs2[m5],
            "total_occurrences": len(imgs1[m5]) + len(imgs2[m5]),
        })
    return results


# ── 文字标注输出 ──────────────────────────────────────────────────

def _apply_yellow_highlight(paragraph):
    for run in paragraph.runs:
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), 'FFFF00')
        run._element.get_or_add_rPr().append(shading)


def save_marked_text_docx(original_path: str, duplicate_info: list,
                          output_path: str,
                          use_second_index: bool = False) -> bool:
    try:
        doc = Document(original_path)
        if use_second_index:
            dup_indices = set(info[1]["index"] for info in duplicate_info)
        else:
            dup_indices = set(info[0]["index"] for info in duplicate_info)
        dup_indices.discard(-1)
        for i, p in enumerate(doc.paragraphs):
            if i in dup_indices and p.text.strip():
                _apply_yellow_highlight(p)
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"  ⚠ 标记文档保存失败 ({output_path}): {e}")
        return False


# ── 报告生成 ──────────────────────────────────────────────────────

def build_report(f1: str, f2: str,
                 exact_dups: list, fuzzy_dups: list,
                 img_dups: list[dict],
                 total1: int, total2: int) -> str:
    lines = []
    lines.append("=" * 60)
    lines.append("  Word 文档查重报告")
    lines.append("=" * 60)
    lines.append(f"  {os.path.basename(f1)}  ↔  {os.path.basename(f2)}")
    lines.append(f"  文档1段落数: {total1}")
    lines.append(f"  文档2段落数: {total2}")
    lines.append("")

    lines.append("─" * 60)
    lines.append(f"【精确匹配】共 {len(exact_dups)} 组重复段落")
    lines.append("─" * 60)
    for p1, p2, _ in exact_dups[:30]:
        lines.append(f"  📄 文档1 §{p1['index']}: {p1['text'][:80]}")
        lines.append(f"  📄 文档2 §{p2['index']}: {p2['text'][:80]}")
        lines.append("")
    if len(exact_dups) > 30:
        lines.append(f"  ... 还有 {len(exact_dups) - 30} 组 (仅显示前30)")

    lines.append("─" * 60)
    lines.append(f"【模糊匹配】共 {len(fuzzy_dups)} 组 (阈值 ≥95%)")
    lines.append("─" * 60)
    for p1, p2, ratio in fuzzy_dups[:20]:
        lines.append(f"  🔗 相似度 {ratio*100:.1f}%")
        lines.append(f"     文档1 §{p1['index']}: {p1['text'][:80]}")
        lines.append(f"     文档2 §{p2['index']}: {p2['text'][:80]}")
        lines.append("")
    if len(fuzzy_dups) > 20:
        lines.append(f"  ... 还有 {len(fuzzy_dups) - 20} 组 (仅显示前20)")

    lines.append("─" * 60)
    lines.append(f"【图片重复】共 {len(img_dups)} 组")
    lines.append("─" * 60)
    for dup in img_dups:
        f1_imgs = ", ".join(i["path"] for i in dup["in_file1"])
        f2_imgs = ", ".join(i["path"] for i in dup["in_file2"])
        lines.append(f"  🖼  MD5: {dup['md5'][:12]}...")
        lines.append(f"      文档1: {f1_imgs}")
        lines.append(f"      文档2: {f2_imgs}")
        lines.append("")
    lines.append(f"  💡 图片重复已标记在 *_标记图片重复.docx 中 (Word 批注方式)")

    lines.append("")
    lines.append("=" * 60)
    lines.append("  报告结束")
    lines.append("=" * 60)
    return "\n".join(lines)


# ── 主入口 ────────────────────────────────────────────────────────

def run_compare(file1: str, file2: str,
                output_dir: str | None = None,
                mode: str = "all",
                threshold: float = 0.95,
                verbose: bool = True) -> dict:
    for f in [file1, file2]:
        if not os.path.isfile(f):
            raise FileNotFoundError(f"文件不存在: {f}")

    if output_dir is None:
        output_dir = os.path.dirname(os.path.abspath(file1))
    os.makedirs(output_dir, exist_ok=True)

    base1 = os.path.splitext(os.path.basename(file1))[0]
    base2 = os.path.splitext(os.path.basename(file2))[0]

    if verbose:
        print(f"📖 读取文档...")
        print(f"   {file1}")
        print(f"   {file2}")

    doc1 = Document(file1)
    doc2 = Document(file2)

    # 段落提取 (文本模式)
    paras1 = paras2 = []
    if mode in ("all", "text"):
        paras1 = extract_paragraphs(doc1)
        paras2 = extract_paragraphs(doc2)
        if verbose:
            print(f"  文档1: {len(paras1)} 段, 文档2: {len(paras2)} 段")

        # 表格文本
        tables1 = extract_tables_text(doc1)
        tables2 = extract_tables_text(doc2)
        from_ = {"index": -1, "text": "", "normalized": "", "paragraph": None}
        table_paras1 = [{**from_, "text": t, "normalized": normalize_text(t)}
                        for t in tables1]
        table_paras2 = [{**from_, "text": t, "normalized": normalize_text(t)}
                        for t in tables2]
        all_paras1 = paras1 + table_paras1
        all_paras2 = paras2 + table_paras2
    else:
        all_paras1 = paras1
        all_paras2 = paras2

    # 精确匹配
    exact_dups = []
    if mode in ("all", "text"):
        exact_dups = find_exact_duplicates(all_paras1, all_paras2)
        if verbose:
            print(f"\n📐 精确匹配: {len(exact_dups)} 组")

    # 模糊匹配
    fuzzy_dups = []
    if mode in ("all", "text"):
        fuzzy_dups = find_fuzzy_duplicates(all_paras1, all_paras2,
                                            cutoff=threshold)
        exact_set = set()
        for p1, p2, _ in exact_dups:
            exact_set.add((p1["normalized"], p2["normalized"]))
        fuzzy_dups = [(p1, p2, r) for p1, p2, r in fuzzy_dups
                      if (p1["normalized"], p2["normalized"]) not in exact_set]
        if verbose:
            print(f"🔄 模糊匹配 (≥{threshold*100:.0f}%): {len(fuzzy_dups)} 组")

    # 图片查重
    img_dups = []
    dup_md5s = set()
    if mode in ("all", "image"):
        imgs1 = extract_images_from_docx(file1)
        imgs2 = extract_images_from_docx(file2)
        cnt1 = sum(len(v) for v in imgs1.values())
        cnt2 = sum(len(v) for v in imgs2.values())
        if verbose:
            print(f"\n🖼  图片查重: 文档1={cnt1}张, 文档2={cnt2}张")
        img_dups = find_duplicate_images(imgs1, imgs2)
        dup_md5s = set(d["md5"] for d in img_dups)
        if verbose:
            print(f"      重复图片: {len(img_dups)} 组 ({len(dup_md5s)} 个独立MD5)")

    # ── 输出：合并标记（文字黄底 + 图片批注，统一输出到 {base}_标记重复.docx）──
    out_marked = (None, None)
    has_text = exact_dups and mode in ("all", "text")
    has_img  = img_dups and dup_md5s and mode in ("all", "image")

    if has_text or has_img:
        path1 = os.path.join(output_dir, f"{base1}_标记重复.docx")
        path2 = os.path.join(output_dir, f"{base2}_标记重复.docx")

        # 步骤 1/2：文字标记（黄底）→ 临时文件
        if has_text:
            temp1 = os.path.join(output_dir, f"__tmp_{base1}.docx")
            temp2 = os.path.join(output_dir, f"__tmp_{base2}.docx")
            all_text_dups = exact_dups + fuzzy_dups
            ok1 = save_marked_text_docx(file1, all_text_dups, temp1)
            ok2 = save_marked_text_docx(file2, all_text_dups, temp2,
                                        use_second_index=True)
            src1, src2 = temp1, temp2
            if not ok1: src1 = file1
            if not ok2: src2 = file2
            if verbose and ok1: print(f"\n📝 文本重复 → 黄底标记完成")
        else:
            src1, src2 = file1, file2

        # 步骤 2/2：图片批注（追加到临时/原始文件上，直接输出最终文件）
        if has_img:
            if verbose: print(f"💬 图片重复 → 正在添加批注...")
            ok1 = mark_duplicate_images_with_comments(src1, dup_md5s, path1)
            ok2 = mark_duplicate_images_with_comments(src2, dup_md5s, path2)
        else:
            shutil.copy2(src1, path1)
            shutil.copy2(src2, path2)
            ok1, ok2 = True, True

        # 清理临时文件
        if has_text:
            for t in [temp1, temp2]:
                try:
                    if t: os.remove(t)
                except: pass

        out_marked = (path1 if (ok1 or not has_img) else None,
                      path2 if (ok2 or not has_img) else None)
        tags = []
        if has_text: tags.append("文字黄底")
        if has_img:  tags.append("图片批注")
        if verbose:
            print(f"  ✅ {' + '.join(tags)} → {path1}")
            print(f"                          {path2}")

    # ── 报告 ──
    report = build_report(file1, file2, exact_dups, fuzzy_dups, img_dups,
                          len(paras1), len(paras2))
    report_path = os.path.join(output_dir,
                               f"{base1}_vs_{base2}_重复报告.txt")
    with open(report_path, "w", encoding="utf-8") as f:
        f.write(report)
    if verbose:
        print(f"\n📄 报告 → {report_path}")
        print(f"\n{'=' * 50}")
        print(f"  对比完成！")
        print(f"  精确: {len(exact_dups)} | 模糊: {len(fuzzy_dups)} | 图片重复: {len(img_dups)}")
        print(f"{'=' * 50}")

    return {
        "exact_matches": len(exact_dups),
        "fuzzy_matches": len(fuzzy_dups),
        "image_duplicates": len(img_dups),
        "report_path": report_path,
        "text_marked_files": out_marked,
        "image_marked_files": out_marked,
    }


# ── CLI ───────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="DOCX 文档查重 — 比较两个 Word 文档的重复文本和重复图片"
    )
    parser.add_argument("--file1", required=True, help="第一个 Word 文档 (.docx)")
    parser.add_argument("--file2", required=True, help="第二个 Word 文档 (.docx)")
    parser.add_argument("--output", "-o", help="输出目录（默认与 file1 同目录）")
    parser.add_argument("--mode", choices=["all", "text", "image"],
                        default="all", help="检查模式")
    parser.add_argument("--threshold", type=float, default=0.95,
                        help="模糊匹配相似度阈值 (0~1)")
    parser.add_argument("--quiet", "-q", action="store_true",
                        help="静默模式")

    args = parser.parse_args()
    if args.threshold < 0 or args.threshold > 1:
        sys.exit("❌ --threshold 必须在 0~1 之间")

    try:
        result = run_compare(args.file1, args.file2, args.output,
                              mode=args.mode, threshold=args.threshold,
                              verbose=not args.quiet)
        if args.quiet and result["report_path"]:
            print(result["report_path"])
    except FileNotFoundError as e:
        sys.exit(f"❌ {e}")
    except Exception as e:
        import traceback
        traceback.print_exc()
        sys.exit(f"❌ 运行时错误: {e}")


if __name__ == "__main__":
    main()
