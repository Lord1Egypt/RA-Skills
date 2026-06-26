# -*- coding: utf-8 -*-
"""
机关公文生成脚本
基于 GB/T 9704-2012 党政机关公文格式标准

注意：正文中的中文引号会被 normalize_chinese_quotes() 自动处理为弯引号
"""

from docx import Document
from docx.shared import Pt, Cm, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================
# 【行距常量】- 28.5磅
# ============================================
LINE_SPACING = Pt(28.5)

# ============================================
# 【配置区】- 修改以下内容生成不同文档
# ============================================

CONFIG = {
    # 发文字号（设为 None 则不显示）
    # 年份必须使用六角括号 〔〕，如 陕XX〔2026〕28号
    "docNumber": "陕XX〔2025〕XX号",

    # 发文字号对齐方式：left / center / right
    "docNumberAlign": "center",

    # 公文标题
    "title": "关于做好全国模范单位奖牌颁发工作的通知",

    # 主送机关
    "mainRecipients": "各设区市、各高校：",

    # 署名（发文机关全称）
    "author": "陕西省XXXX",

    # 日期（年月日写全）
    "date": "2025年X月X日",

    # 附件列表（设为空列表则不显示）
    "attachments": ["全国模范单位奖牌颁发情况统计表"],

    # ── 联系方式（三种配置方式，任选其一）──
    # 方式一：简单字符串（自动用中文括号包裹，首行缩进）
    # "contact": "联系人：XXX 189XXXXXXXX，地址：陕西省西安市XXXXXXXX"
    #
    # 方式二：结构化字典（自动格式化，支持多字段）
    # "contact": {"联系人": "XXX", "电话": "02986515321", "地址": "陕西省西安市XXX"}
    #
    # 方式三：设为 None 则不显示联系方式
    "contact": "联系人：XXX 189XXXXXXXX，电子邮箱：XXXXXXXX@qq.com，邮寄地址：陕西省西安市XXXXXXXX",

    # 联系方式格式选项
    "contactFormat": {
        # 是否用中文括号包裹整段联系方式（默认 True）
        "bracket": True,
        # 字段间分隔符（结构化模式生效，默认 "  "）
        "separator": "  ",
    },

    # 是否显示页码（格式：— N —，半角宋体14pt，单页右空一字，双页左空一字）
    "showPageNumber": True,

    # 输出文件名
    "outputFile": "公文.docx"
}

# ============================================
# 【正文内容区】- 在这里填入实际内容
# ============================================

CONTENTS = [
    # 格式说明：
    # ("h1", "一、一级标题")
    # ("h2", "（一）二级标题")
    # ("normal", "正文内容")
    # ("blank", None)  # 空行

    ("h1", "一、充分认识奖牌颁发工作的重要意义"),
    ("normal", "全国模范单位奖牌是表彰先进、树立典型的重要载体，做好奖牌颁发工作对于弘扬精神、激发会员积极性具有重要意义。各地要高度重视，精心组织，确保奖牌颁发工作顺利完成。"),
    ("h1", "二、获奖名单"),
    ("normal", "本次表彰共涉及我省XX家全国模范单位，涵盖社区、企业、学校等多种类型。"),
    ("h1", "三、领取颁发要求"),
    ("h2", "（一）领取方式"),
    ("normal", "获奖单位请于2025年X月X日前联系所在设区市领取奖牌。"),
    ("h2", "（二）颁发时间"),
    ("normal", "请于2025年X月X日前完成奖牌颁发工作，并填写奖牌领取确认表。"),
    ("h1", "四、工作要求"),
    ("normal", "各地要加强领导，明确责任，确保奖牌颁发工作规范有序。要以此次颁奖为契机，进一步推动事业发展。"),
    ("h1", "五、联系方式"),
    ("normal", "联系人：XXX    联系电话：XXXX-XXXXXXXX"),
]

# ============================================
# 【辅助函数】
# ============================================
# 【引号标准化】- 将 ASCII 直引号还原为中文弯引号
# 原因：AI 在生成脚本时容易将用户输入的中文弯引号转为 ASCII 直引号，
#       本函数在文本写入 docx 前自动还原，确保公文格式正确。
# ============================================
LEFT_DOUBLE = "\u201C"   # \u201c 中文左双弯引号
RIGHT_DOUBLE = "\u201D"  # \u201d 中文右双弯引号
LEFT_SINGLE = "\u2018"  # \u2018 中文左单弯引号
RIGHT_SINGLE = "\u2019" # \u2019 中文右单弯引号

def normalize_chinese_quotes(text):
    """
    将文本中的 ASCII 直引号交替还原为中文弯引号。

    原理：ASCII 直引号 " 和 ' 是对称字符（同一字符既是左引号也是右引号），
    无法直接区分左右。因此采用交替替换策略：第1、3、5...个双引号 → 左双引号；
    第2、4、6...个双引号 → 右双引号。单引号同理，独立计数。
    这是中文排版中配对引号的标准做法。
    """
    if not text:
        return text

    result = []
    double_count = 0  # 双引号计数（独立）
    single_count = 0  # 单引号计数（独立）
    i = 0
    while i < len(text):
        c = text[i]
        if c == '"':
            if double_count % 2 == 0:
                result.append(LEFT_DOUBLE)   # 奇数个双引号 → 左
            else:
                result.append(RIGHT_DOUBLE)  # 偶数个双引号 → 右
            double_count += 1
        elif c == "'":
            if single_count % 2 == 0:
                result.append(LEFT_SINGLE)   # 奇数个单引号 → 左
            else:
                result.append(RIGHT_SINGLE)  # 偶数个单引号 → 右
            single_count += 1
        else:
            result.append(c)
        i += 1
    return ''.join(result)

# ============================================

def set_run_font(run, font_name="仿宋_GB2312", font_size=16, bold=False):
    """
    设置文本字体（run 级别强制声明）。
    
    在 run._element.rPr 上完整设置 ascii / hAnsi / eastAsia 三个字体槽位，
    确保 Word 渲染时不会因样式继承导致字体"溢出"。
    """
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    # 中文字体：显式设置 eastAsia 槽位
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    # 西文字体同步，防止 Word 回退到 Normal 样式字体
    run._element.rPr.rFonts.set(qn('w:ascii'), font_name)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), font_name)


def create_paragraph(doc, text, font_name="仿宋_GB2312", font_size=16,
                     alignment=WD_ALIGN_PARAGRAPH.LEFT,
                     first_line_indent=Twips(640), bold=False,
                     space_before=0, space_after=0,
                     right_indent=None):
    """
    创建段落

    参数:
        font_name: 字体名称
        font_size: 字号（pt）
        alignment: 对齐方式
        first_line_indent: 首行缩进，Twips(640)=2个汉字，0表示无缩进
        bold: 是否加粗
        right_indent: 右缩进（Twips），用于落款距右4汉字
    """
    p = doc.add_paragraph()
    # 用 paragraph_format.alignment 确保覆盖 Normal 样式
    p.paragraph_format.alignment = alignment

    # 设置段落间距
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)

    # 设置首行缩进
    if first_line_indent and first_line_indent > 0:
        p.paragraph_format.first_line_indent = first_line_indent
    else:
        p.paragraph_format.first_line_indent = Cm(0)

    # 设置右缩进（用于落款距右4汉字）
    if right_indent is not None:
        p.paragraph_format.right_indent = right_indent

    # 添加文本（自动还原中文弯引号）
    if text:
        text = normalize_chinese_quotes(text)
        run = p.add_run(text)
        set_run_font(run, font_name, font_size, bold)

    return p


def create_blank_line(doc):
    """创建空行"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p


def add_content(doc, item):
    """
    根据类型添加正文内容。
    
    二级标题（h2）特殊处理：
    - 检测文本是否包含正文内容（同一段落或换行分隔）
    - 将标题部分和正文部分拆分为独立 run，各自强制声明字体
    - 标题部分用楷体_GB2312，正文部分用仿宋_GB2312
    - 同段落模式：在同一 run 序列中通过 line break 分隔
    - 换行模式：正文另起一个段落
    """
    content_type, text = item

    if content_type == "blank":
        return create_blank_line(doc)

    elif content_type == "h1":
        # 一级标题：黑体，不加粗
        return create_paragraph(doc, text,
                                font_name="黑体",
                                font_size=16,
                                first_line_indent=Twips(640),
                                bold=False)

    elif content_type == "h2":
        return add_h2_with_body(doc, text)

    elif content_type == "normal":
        # 正文：仿宋，首行缩进，左对齐
        return create_paragraph(doc, text,
                                font_name="仿宋_GB2312",
                                font_size=16,
                                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                first_line_indent=Twips(640),
                                bold=False)

    else:
        return create_paragraph(doc, text)


def split_h2_title_and_body(text):
    """
    拆分二级标题文本，识别标题部分和正文部分。
    
    支持两种格式：
    1. 同段落模式 —— 标题和正文在同一行，第一个句号后跟正文
       例: "（一）创新学习机制，激发内生动力。改变以往……"
       → title: "（一）创新学习机制，激发内生动力。"
       → body:  "改变以往……"
    
    2. 换行模式 —— 标题和正文之间有换行符
       例: "（一）创新学习机制，激发内生动力\n改变以往……"
       → title: "（一）创新学习机制，激发内生动力"
       → body:  "改变以往……"
    
    3. 纯标题模式 —— 没有正文内容
       例: "（一）创新学习机制，激发内生动力"
       → title: "（一）创新学习机制，激发内生动力"
       → body:  None
    
    检测策略：
    - 如果文本包含 \\n → 换行模式
    - 如果第一个句号（。或.）后面还有内容 → 同段落模式
    - 否则 → 纯标题模式
    
    边界条件：
    - 以"一是""二是"等开头的句子视为正文（排除标题末尾标点后直接跟正文）
    - 句号紧跟右括号"）"的情况正确处理
    """
    if not text:
        return text, None
    
    # 模式1：换行分隔
    if '\n' in text:
        parts = text.split('\n', 1)
        return parts[0].strip(), parts[1].strip()
    
    # 模式2/3：在同一段落内查找标题和正文的分界点
    # 找第一个句号（中文句号或英文句号）
    for i, c in enumerate(text):
        if c in ('。', '.'):
            # 句号后面还有文字内容 → 同段落模式
            remaining = text[i + 1:].strip()
            if remaining:
                return text[:i + 1], remaining
            else:
                # 句号在末尾，没有正文 → 纯标题模式
                return text, None
    
    # 没有句号 → 纯标题模式
    return text, None


def add_h2_with_body(doc, text):
    """
    添加二级标题段落，自动拆分标题和正文部分并分别设置字体。
    
    标题部分：楷体_GB2312 16pt
    正文部分：仿宋_GB2312 16pt
    
    - 同段落模式：标题 run + line break + 正文 run（同一行显示，字体各自独立）
    - 换行模式：标题段落 + 正文段落（各占一行）
    - 纯标题模式：仅标题段落（向后兼容）
    """
    title_text, body_text = split_h2_title_and_body(text)
    
    # 引号标准化
    title_text = normalize_chinese_quotes(title_text)
    
    if body_text is not None:
        body_text = normalize_chinese_quotes(body_text)
    
    # ── 纯标题模式 ──
    if body_text is None:
        return create_paragraph(doc, title_text,
                                font_name="楷体_GB2312",
                                font_size=16,
                                first_line_indent=Twips(640),
                                bold=False)
    
    # ── 换行模式：标题和正文各自独立段落 ──
    if '\n' in text:
        # 标题段落
        create_paragraph(doc, title_text,
                        font_name="楷体_GB2312",
                        font_size=16,
                        first_line_indent=Twips(640),
                        bold=False)
        # 正文段落
        return create_paragraph(doc, body_text,
                               font_name="仿宋_GB2312",
                               font_size=16,
                               alignment=WD_ALIGN_PARAGRAPH.LEFT,
                               first_line_indent=Twips(640),
                               bold=False)
    
    # ── 同段落模式：标题和正文在同一行，通过独立 run 保证字体隔离 ──
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Twips(640)
    
    # Run 1：标题部分（楷体）
    run_title = p.add_run(title_text)
    set_run_font(run_title, "楷体_GB2312", 16, bold=False)
    
    # Run 2：正文部分（仿宋）
    run_body = p.add_run(body_text)
    set_run_font(run_body, "仿宋_GB2312", 16, bold=False)
    
    return p


def set_page_margins(doc, top=Cm(3.7), bottom=Cm(3.5),
                     left=Cm(2.8), right=Cm(2.6),
                     header_distance=Cm(1.5), footer_distance=Cm(2.5)):
    """设置页边距及页眉页脚距离"""
    sections = doc.sections
    for section in sections:
        section.top_margin = top
        section.bottom_margin = bottom
        section.left_margin = left
        section.right_margin = right
        section.header_distance = header_distance
        section.footer_distance = footer_distance


def format_contact(config):
    """
    格式化联系方式。

    支持两种输入：
    1. 字符串：直接使用，如 "联系人：XXX 电话：XXX"
    2. 字典：自动格式化，如 {"联系人": "XXX", "电话": "XXX"}
    3. None：不显示

    返回格式化后的字符串（不含括号，括号由调用方添加）。
    """
    contact = config.get("contact")
    fmt_opts = config.get("contactFormat", {})

    if contact is None:
        return None

    if isinstance(contact, dict):
        # 结构化模式：按字典顺序拼接
        separator = fmt_opts.get("separator", "  ")
        parts = []
        for key, value in contact.items():
            parts.append(key + "：" + str(value))
        return separator.join(parts)

    # 字符串模式：直接使用
    return contact


def add_page_number(doc):
    """
    添加符合 GB/T 9704-2012 标准的页码。

    标准规定：
    - 格式：— N —（半角破折号，页码数字，半角破折号）
    - 字体：半角宋体（宋体），4号（14pt）
    - 位置：版心下边缘之下一行
    - 奇数页（单页）：页码右对齐，右侧空一个字符
    - 偶数页（双页）：页码左对齐，左侧空一个字符
    - 段前段后间距为0，行间距单倍间距

    实现方式：
    - 利用 Word 奇偶页不同页脚功能（differentFirstPage=False, oddAndEvenPages=True）
    - 奇数页页脚：右对齐，右侧Tab后插入页码域
    - 偶数页页脚：左对齐，左侧Tab前插入页码域
    - 数字和破折号均使用4号（14pt）宋体

    GB/T 9704-2012 6.7：
    "页码一般用阿拉伯数字标注，置于版心下边缘之下，数字左右各放一条一字线"
    "单页码居右空一字，双页码居左空一字"
    """
    section = doc.sections[0]
    # 启用奇偶页不同页脚
    sectPr = section._sectPr
    titlePg = OxmlElement('w:titlePg')
    sectPr.append(titlePg)

    # 设置奇偶页不同
    evenAndOdd = OxmlElement('w:evenAndOddHeaders')
    sectPr.append(evenAndOdd)

    def make_page_number_fldChar(run_elem, fld_type="current"):
        """创建页码域字段（ARABIC 为当前页，NUMPAGES 为总页数）"""
        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE \\* ARABIC ' if fld_type == "current" else ' NUMPAGES \\* ARABIC '

        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')

        run_elem.append(fldChar_begin)
        run_elem.append(instrText)
        run_elem.append(fldChar_end)

    def set_footer_run_font(run_elem, text=None, size_pt=14):
        """设置页脚run的字体为宋体14pt"""
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), '宋体')
        rFonts.set(qn('w:hAnsi'), '宋体')
        rFonts.set(qn('w:eastAsia'), '宋体')
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(size_pt * 2))  # half-points
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(size_pt * 2))
        rPr.append(rFonts)
        rPr.append(sz)
        rPr.append(szCs)
        run_elem.append(rPr)
        if text is not None:
            t = OxmlElement('w:t')
            t.set(qn('xml:space'), 'preserve')
            t.text = text
            run_elem.append(t)

    def build_footer_paragraph(footer, alignment):
        """在 footer 中创建一个段落并返回其 pPr
        段前段后间距为0，行间距单倍间距，无首行缩进
        """
        # 清空默认段落
        for p in footer.paragraphs:
            p._element.getparent().remove(p._element)

        pPr_elem = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), alignment)  # 'right' or 'left'
        # 页脚行距设为单倍行距（28pt），段前段后间距为0
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:line'), '240')  # 单倍行距 twips值
        spacing.set(qn('w:lineRule'), 'auto')  # 自动行距（单倍）
        spacing.set(qn('w:before'), '0')  # 段前间距0
        spacing.set(qn('w:after'), '0')  # 段后间距0
        # 明确清除首行缩进和左缩进
        ind = OxmlElement('w:ind')
        ind.set(qn('w:firstLine'), '0')
        ind.set(qn('w:left'), '0')
        pPr.append(jc)
        pPr.append(spacing)
        pPr.append(ind)
        pPr_elem.append(pPr)
        footer._element.append(pPr_elem)
        return pPr_elem

    ONE_IDEOGRAPH_SPACE = '\u3000'  # 全角空格，等于一个汉字宽

    # ── 奇数页页脚：右对齐，右侧空一字 ──────────────────────────────
    odd_footer = section.footer
    p_odd = build_footer_paragraph(odd_footer, 'right')

    # run: "— "  4号一字线
    r1 = OxmlElement('w:r')
    set_footer_run_font(r1, text='— ', size_pt=14)
    p_odd.append(r1)
    # run: 页码域 4号
    r2 = OxmlElement('w:r')
    set_footer_run_font(r2, size_pt=14)
    make_page_number_fldChar(r2)
    p_odd.append(r2)
    # run: " —" + 全角空格（右空一字） 4号一字线
    r3 = OxmlElement('w:r')
    set_footer_run_font(r3, text=' —' + ONE_IDEOGRAPH_SPACE, size_pt=14)
    p_odd.append(r3)

    # ── 偶数页页脚：左对齐，左侧空一字 ──────────────────────────────
    even_footer = section.even_page_footer
    p_even = build_footer_paragraph(even_footer, 'left')

    # run: 全角空格（左空一字）+ "— " 4号一字线
    r4 = OxmlElement('w:r')
    set_footer_run_font(r4, text=ONE_IDEOGRAPH_SPACE + '— ', size_pt=14)
    p_even.append(r4)
    # run: 页码域 4号
    r5 = OxmlElement('w:r')
    set_footer_run_font(r5, size_pt=14)
    make_page_number_fldChar(r5)
    p_even.append(r5)
    # run: " —" 4号一字线
    r6 = OxmlElement('w:r')
    set_footer_run_font(r6, text=' —', size_pt=14)
    p_even.append(r6)


# ============================================
# 【主函数】
# ============================================

def generate_document():
    print("=== 机关公文生成器 ===")
    print("标题:", CONFIG["title"])
    print("发文字号:", CONFIG["docNumber"])
    print("主送机关:", CONFIG["mainRecipients"])
    print("输出:", CONFIG["outputFile"])
    print("====================\n")

    # 创建文档
    doc = Document()

    # 设置页边距（GB/T 9704-2012）
    set_page_margins(doc)

    # ── 1. 发文字号 ─────────────────────────────────────────────────
    # GB/T 9704-2012 6.1.2：发文机关标志下方，居中或左对齐
    # 年份必须使用六角括号 〔〕
    if CONFIG["docNumber"]:
        align_str = CONFIG.get("docNumberAlign", "center")
        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }
        doc_align = align_map.get(align_str, WD_ALIGN_PARAGRAPH.CENTER)
        create_paragraph(doc, CONFIG["docNumber"],
                        alignment=doc_align,
                        first_line_indent=0)

    # ── 2. 标题前空两行 ──────────────────────────────────────────────
    create_blank_line(doc)
    create_blank_line(doc)

    # ── 3. 标题 - 居中，方正小标宋，不加粗 ─────────────────────────
    create_paragraph(doc, CONFIG["title"],
                     font_name="方正小标宋简体",
                     font_size=22,
                     alignment=WD_ALIGN_PARAGRAPH.CENTER,
                     first_line_indent=0,
                     bold=False)

    # ── 4. 标题和主送机关之间空一行 ─────────────────────────────────
    create_blank_line(doc)

    # ── 5. 主送机关（顶格，不缩进）─────────────────────────────────
    if CONFIG["mainRecipients"]:
        create_paragraph(doc, CONFIG["mainRecipients"],
                        first_line_indent=0)

    # ── 6. 正文 ─────────────────────────────────────────────────────
    for item in CONTENTS:
        add_content(doc, item)

    # ── 7. 附件说明 ─────────────────────────────────────────────────
    if CONFIG["attachments"] and len(CONFIG["attachments"]) > 0:
        create_blank_line(doc)
        if len(CONFIG["attachments"]) == 1:
            # 单个附件：附件：xxx
            create_paragraph(doc, "附件：" + CONFIG["attachments"][0],
                             first_line_indent=Twips(640))
        else:
            # 多个附件：附件：1.xxx 换行 2.xxx
            create_paragraph(doc, "附件：1." + CONFIG["attachments"][0],
                             first_line_indent=Twips(640))
            for i, att in enumerate(CONFIG["attachments"][1:], start=2):
                create_paragraph(doc, str(i) + "." + att,
                                 first_line_indent=Twips(640))

    # ── 8. 落款和成文日期 ──────────────────────────────────────────
    # GB/T 9704-2012 7.3.5.1：署名和成文日期右对齐，右侧空4汉字
    # 实现方式：右对齐 + 右缩进 Twips(1280)（4个汉字 ≈ 4 × 320twips）
    SIGN_RIGHT_INDENT = Twips(1280)

    # 落款前空3行
    create_blank_line(doc)
    create_blank_line(doc)
    create_blank_line(doc)

    if CONFIG["author"]:
        create_paragraph(doc, CONFIG["author"],
                        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                        first_line_indent=0,
                        right_indent=SIGN_RIGHT_INDENT)

    if CONFIG["date"]:
        create_paragraph(doc, CONFIG["date"],
                        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                        first_line_indent=0,
                        right_indent=SIGN_RIGHT_INDENT)

    # ── 9. 联系方式 ─────────────────────────────────────────────────
    # 位置：落款日期下方空1行
    # 格式：首行缩进2字符，仿宋_GB2312 16pt
    # 支持字符串模式和结构化字典模式
    contact_text = format_contact(CONFIG)
    if contact_text:
        create_blank_line(doc)

        fmt_opts = CONFIG.get("contactFormat", {})
        use_bracket = fmt_opts.get("bracket", True)
        if use_bracket:
            contact_text = "（" + contact_text + "）"

        create_paragraph(doc, contact_text,
                         alignment=WD_ALIGN_PARAGRAPH.LEFT,
                         first_line_indent=Twips(640))

    # ── 10. 页码（GB/T 9704-2012）───────────────────────────────────
    if CONFIG.get("showPageNumber", True):
        add_page_number(doc)

    # 保存文档
    doc.save(CONFIG["outputFile"])
    print("生成成功:", CONFIG["outputFile"])


if __name__ == "__main__":
    generate_document()
